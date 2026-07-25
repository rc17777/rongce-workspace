import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

DESKTOP = r'C:\Users\scrccpa\Desktop'

SECTIONS = [
    (r'融策制度汇编-人力资源篇.docx', '人力资源篇'),
    (r'融策制度汇编-财务管理篇.docx', '财务管理篇'),
    (r'融策制度汇编-业务部管理篇.docx', '业务部管理篇'),
    (r'融策制度汇编-业务质控篇.docx', '业务质控篇'),
    (r'融策制度汇编-行政综合篇.docx', '行政综合篇'),
]

# ======================================
# Step 1: Create master document with cover
# ======================================
master = Document()

# Use the first section's cover page as template
first = Document(os.path.join(DESKTOP, SECTIONS[0][0]))

# Copy cover page elements (from first doc's first page - stop at first page break or section start)
print('Building cover and TOC...')
cover_count = 0
for p in first.paragraphs:
    master.element.body.append(copy.deepcopy(p._element))
    cover_count += 1
    t = p.text.strip()
    # Stop after the overview/TOC section
    if '制度体系概览' in t:
        # include the overview items too
        continue
    # Check if we've passed the TOC - look for the first actual document heading
    if t.startswith('四川融策') and '制度' in t and cover_count > 10:
        # This is likely a document title - we've passed the overview
        # Remove this last one, we don't want it in the cover section
        last_elem = master.element.body[-1]
        master.element.body.remove(last_elem)
        break

print(f'  Cover/TOC: {cover_count} paragraphs from first doc')

# ======================================
# Step 2: Process each section, appending content
# ======================================
for filename, label in SECTIONS:
    print(f'\nProcessing: {label} ({filename})')
    section_doc = Document(os.path.join(DESKTOP, filename))
    
    # Find the section body content - skip cover page and TOC
    # The section files each have their own cover + TOC + "XX篇 · 制度体系概览"
    # We need to skip those and start from the actual document content
    
    started = False
    skipped = 0
    added = 0
    
    for p in section_doc.paragraphs:
        t = p.text.strip()
        
        # Find the first actual document in the section
        # Documents start with patterns like "四川融策薪酬管理制度" or "RC-HR-001"
        if not started:
            is_bold = any(r.bold for r in p.runs if r.bold)
            # Skip cover page materials
            if '四川融策' in t and is_bold and skipped < 5:
                skipped += 1
                continue
            if '管理制度汇编' in t or '——' in t or '制度体系概览' in t or '目' == t.strip() or '目录' in t:
                skipped += 1
                continue
            if '四川融策会计师事务所有限公司' in t or '四川融策工程咨询有限公司' in t:
                skipped += 1
                continue
            if t.startswith('二〇') and len(t) < 20:
                skipped += 1
                continue
            
            # Now check if we're at the start of actual content
            if t and (t.startswith('四川融策') or t.startswith('RC-') or t.startswith('【RC-')):
                started = True
        
        if started:
            master.element.body.append(copy.deepcopy(p._element))
            added += 1
    
    print(f'  Skipped: {skipped}, Added: {added} paragraphs')
    
    # Add a page break between sections
    if label != SECTIONS[-1][1]:  # not the last section
        pb = etree.SubElement(master.element.body, qn('w:p'))
        r = etree.SubElement(pb, qn('w:r'))
        br = etree.SubElement(r, qn('w:br'))
        br.set(qn('w:type'), 'page')

# ======================================
# Step 3: Fix the TOC
# ======================================
# Now we need to update the master TOC to include all sections.
# The cover page has a TOC from the first section only.
# Let's find and rebuild it.

# Find the TOC section in the master doc
toc_section_start = None
toc_section_end = None
for i, p in enumerate(master.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '目    录':
        toc_section_start = i
    if toc_section_start and is_bold and t == '行政综合篇' and '章节数' in master.paragraphs[i-1].text if i > 0 else False:
        pass  # this might not work
    
    # Find TOC end: the next page after "行政综合篇" (or last TOC item)
    # TOC ends at the first template-like section marker
    if toc_section_start and i > toc_section_start:
        if is_bold and ('人力资源篇' in t or '业务部管理篇' in t or '业务质控篇' in t or '行政综合篇' in t):
            toc_last = i

# Actually, let's find and replace the TOC manually
# The TOC paragraphs are between "目    录" and the first section heading
# "人力资源篇" is typically the first section in the body

toc_start_idx = None
toc_end_idx = None
body_section_idx = None

for i, p in enumerate(master.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '目    录':
        toc_start_idx = i
    # The first body section marker after TOC
    if toc_start_idx and i > toc_start_idx + 1 and is_bold and t == '人力资源篇':
        toc_end_idx = i  # TOC items end before this
        body_section_idx = i
        break

print(f'\nTOC: start={toc_start_idx}, end={toc_end_idx}')

if toc_start_idx and toc_end_idx:
    # Remove old TOC entries (paragraphs between toc_start+1 and toc_end-1)
    toc_entries_start = toc_start_idx + 1
    toc_entries_end = toc_end_idx - 1
    
    # Collect paragraph elements to remove
    to_remove = []
    for i in range(toc_entries_start, toc_entries_end + 1):
        to_remove.append(master.paragraphs[i]._element)
    
    for elem in to_remove:
        elem.getparent().remove(elem)
    
    print(f'  Removed {len(to_remove)} old TOC entries')
    
    # Now insert new TOC entries
    # Target: the paragraph before 人力资源篇 in the TOC area
    # After removal, 人力资源篇 (body section) hasn't moved
    
    # Find 人力资源篇 again
    hr_elem = None
    for p in master.paragraphs:
        t = p.text.strip()
        is_bold = any(r.bold for r in p.runs if r.bold)
        if is_bold and t == '人力资源篇':
            hr_elem = p._element
            break
    
    if hr_elem:
        # Create TOC entries for each section
        toc_labels = ['人力资源篇', '财务管理篇', '业务部管理篇', '业务质控篇', '行政综合篇']
        
        # Use the first section's 人力资源篇 as style template
        ref_elem = hr_elem
        
        for label in reversed(toc_labels):
            new_p = copy.deepcopy(ref_elem)
            # Update text
            for r_elem in new_p.findall(qn('w:r')):
                for t_elem in r_elem.findall(qn('w:t')):
                    t_elem.text = label
                    break
            hr_elem.addprevious(new_p)
        
        print(f'  Inserted {len(toc_labels)} TOC entries')

# ======================================
# Step 4: Save
# ======================================
output = os.path.join(DESKTOP, '融策公司制度体系（完整版）.docx')
master.save(output)
print(f'\nSaved: {output}')
print(f'Size: {os.path.getsize(output):,} bytes')
print(f'Paragraphs: {len(master.paragraphs)}')
