import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy, os
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DESKTOP = r'C:\Users\scrccpa\Desktop'
OUTPUT = os.path.join(DESKTOP, '融策公司制度体系（完整版）.docx')

SECTIONS = [
    (r'融策制度汇编-人力资源篇.docx', '人力资源篇'),
    (r'融策制度汇编-财务管理篇.docx', '财务管理篇'),
    (r'融策制度汇编-业务部管理篇.docx', '业务部管理篇'),
    (r'融策制度汇编-业务质控篇.docx', '业务质控篇'),
    (r'融策制度汇编-行政综合篇.docx', '行政综合篇'),
]

# Open first section as base
base_path = os.path.join(DESKTOP, SECTIONS[0][0])
master = Document(base_path)
body = master.element.body

print(f'Base: {SECTIONS[0][1]} - {len(master.paragraphs)} paragraphs')

# Collect all paragraphs from subsequent sections and append to master
total_paras = len(master.paragraphs)
for filename, label in SECTIONS[1:]:
    filepath = os.path.join(DESKTOP, filename)
    sec_doc = Document(filepath)
    sec_body = sec_doc.element.body
    
    count = 0
    # Copy all paragraph elements from section to master
    for child in sec_body:
        if child.tag == qn('w:p'):
            body.append(copy.deepcopy(child))
            count += 1
        elif child.tag == qn('w:tbl'):
            body.append(copy.deepcopy(child))
    
    print(f'  + {label}: {count} elements')
    total_paras += count

print(f'\nTotal: {total_paras} elements')

master.save(OUTPUT)
print(f'Saved: {OUTPUT}')
print(f'Size: {os.path.getsize(OUTPUT):,} bytes')
