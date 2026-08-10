import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx')

# Check TOC section
print("=== TOC (first 60 paragraphs, bold only) ===")
for i, p in enumerate(doc.paragraphs[:60]):
    t = p.text.strip()
    if not t:
        continue
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold:
        print(f'  [{i}] {t[:80]}')

print("\n=== Section headings (all bold+large paragraphs) ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    is_bold = any(r.bold for r in p.runs if r.bold)
    is_large = any(r.font.size and r.font.size >= 140000 for r in p.runs if r.font.size)
    if is_bold and is_large and len(t) < 30:
        print(f'  [{i}] {t[:80]}')

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
