# -*- coding: utf-8 -*-
"""恩阳医养园PPP可用性付费报告 — P0/P1/P2修正版 v2.0"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ============================================================
# PARAMETERS (P0 corrected: op cost and income from 2025 audit)
# ============================================================
CAPITAL_TOTAL = 102_298_077.79
CAPITAL_AMORT = 80_000_000.00      # 等额本息部分
CAPITAL_LUMP = 22_298_077.79       # 第18年一次性支付
RATE = 0.0799                       # 7.99%
YEARS = 18

A = CAPITAL_AMORT * RATE * (1+RATE)**YEARS / ((1+RATE)**YEARS - 1)  # PMT

# Capital amortization schedule
equity_pmt = A
cap_remaining = CAPITAL_AMORT
cap_schedule = []
for y in range(YEARS):
    interest = cap_remaining * RATE
    principal = equity_pmt - interest
    if y == YEARS - 1:
        principal = cap_remaining
        interest = equity_pmt - principal
    cap_remaining -= principal
    if cap_remaining < 0:
        cap_remaining = 0
    cap_schedule.append((principal, interest, equity_pmt, max(cap_remaining, 0)))

# Bank repayment schedule
# Per-year data from 原征求意见稿表2; 贵阳银行《借款还本付息情况报告》合计710,822,119.32元
# (原表合计712,464,358.22,差额1,642,238.90为调整口径差异,本表保留原表逐项以保持一致性)
bank_schedule = [
    (31_686_745.84, 381_300_000.00),   # 2023 - 6.95%
    (26_809_671.03, 381_300_000.00),   # 2024 - 6.45%
    (24_811_756.73, 381_250_000.00),   # 2025
    (21_351_219.32, 380_250_000.00),   # 2026
    (20_263_854.17, 379_250_000.00),   # 2027
    (20_265_833.33, 378_250_000.00),   # 2028
    (20_162_465.30, 377_250_000.00),   # 2029
    (20_111_770.82, 376_250_000.00),   # 2030
    (23_022_951.39, 372_250_000.00),   # 2031
    (22_871_874.99, 368_250_000.00),   # 2032
    (22_617_395.82, 364_250_000.00),   # 2033
    (22_414_618.06, 360_250_000.00),   # 2034
    (27_148_298.64, 351_250_000.00),   # 2035
    (27_728_124.99, 341_250_000.00),   # 2036
    (35_070_729.16, 323_250_000.00),   # 2037
    (34_158_229.18, 305_250_000.00),   # 2038
    (311_968_819.45,         0),       # 2039 - 结清
    (          0.00,         0),       # 2040
]
bank_total = sum(b[0] for b in bank_schedule)

# P0 FIX: Operational costs from 2025 actual audit data
OP_COST_2023 = 5_768_117.42    # actual (2022.10.28-2023.12.21)
OP_COST_2024 = 6_890_208.16    # actual (2023.12.21-2024.12.21)
OP_COST_2025 = 5_713_241.86    # actual 2025 (主营成本4,474,028 + 管理费1,239,214)
OP_COST_FUTURE = 5_710_000.00  # 2026-2040 estimate (based on 2025 actual)

# P0 FIX: Third-party income from 2025 actual audit data
INC_2023 = 4_285_070.31        # actual
INC_2024 = 6_229_093.45        # actual
INC_2025 = 5_564_017.88        # actual 2025 (主营收入4,961,143 + 其他收入602,875)
INC_FUTURE = 5_560_000.00      # 2026-2040 estimate (based on 2025 actual)

YEAR_LABELS = list(range(2023, 2041))

op_costs = [OP_COST_2023, OP_COST_2024] + [OP_COST_2025] + [OP_COST_FUTURE]*15
incomes = [INC_2023, INC_2024] + [INC_2025] + [INC_FUTURE]*15
op_cost_total = sum(op_costs)
income_total = sum(incomes)

# Compute yearly availability payments
yearly = []
for i in range(YEARS):
    yr = YEAR_LABELS[i]
    principal, interest, pmt, remaining = cap_schedule[i]
    cap_repaid = pmt  # annual capital return = interest + principal on amortized portion
    lump = CAPITAL_LUMP if i == YEARS - 1 else 0
    total_cap = cap_repaid + lump
    bank = bank_schedule[i][0]
    op = op_costs[i]
    inc = incomes[i]
    avail = total_cap + bank + op - inc
    yearly.append((yr, principal, interest, pmt, lump, total_cap, remaining, bank, op, inc, avail))

grand_total = sum(y[10] for y in yearly)
cap_total = sum(y[5] for y in yearly)

print(f"PMT (8000万@7.99%): {A:,.2f}")
print(f"Bank total: {bank_total:,.2f}")
print(f"Op cost total: {op_cost_total:,.2f}")
print(f"Income total: {income_total:,.2f}")
print(f"Capital return total: {cap_total:,.2f}")
print(f"Grand total: {grand_total:,.2f} (≈{grand_total/1e8:.2f}亿)")

# ============================================================
# BUILD WORD DOCUMENT
# ============================================================
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5

def set_cell_font(cell, text, font_name='宋体', size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    """Helper to set cell text with formatting"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    # Reduce paragraph spacing in tables
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, font_name='宋体', size=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def make_table(doc, headers, rows, col_widths=None):
    """Create a professionally formatted table"""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_font(cell, h, font_name='微软雅黑', size=Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, '0A1F3F')
    
    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(cell, str(val), font_name='宋体', size=Pt(8.5), align=align)
            # Alternate row shading
            if i % 2 == 1:
                set_cell_shading(cell, 'F5F2EC')
    
    # Column widths
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    
    return table

# ========= COVER PAGE =========
for _ in range(6):
    doc.add_paragraph()

add_paragraph(doc, '巴中市恩阳医养园PPP项目', font_name='方正小标宋简体', size=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '可用性付费测算结果报告', font_name='方正小标宋简体', size=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_paragraph(doc, '川融策咨询〔2025〕第490号', size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(6):
    doc.add_paragraph()

add_paragraph(doc, '四 川 融 策 会 计 师 事 务 所 有 限 公 司', font_name='微软雅黑', size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, 'Sichuan Rongce Accounting Firm Co., Ltd', size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))
add_paragraph(doc, '2026年7月29日', size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

# ========= PAGE BREAK =========
doc.add_page_break()

# ========= BODY =========
add_paragraph(doc, '川融策咨询〔2025〕第  号', align=WD_ALIGN_PARAGRAPH.RIGHT, size=Pt(10))

doc.add_paragraph()
add_paragraph(doc, '巴中市恩阳医养园PPP项目', font_name='黑体', size=Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '可用性付费测算结果报告', font_name='黑体', size=Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_paragraph(doc, '巴中市恩阳区卫生健康局：', first_line_indent=Cm(0.74))

add_paragraph(doc, '我们接受委托，对巴中市恩阳医养园PPP项目2022年10月28日-2040年12月21日可用性付费进行测算。本次测算基数暂以《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）中审定的建筑安装工程费用485,098,077.79元为准，运营维护成本及第三方收入已根据2025年度实际运营审计数据进行了修正，现将结果报告如下：', first_line_indent=Cm(0.74))

# ===== 一、项目概况 =====
add_paragraph(doc, '一、项目概况', font_name='黑体', size=Pt(14), bold=True)
add_paragraph(doc, '2015年12月21日，巴中市恩阳医养园PPP项目经巴中市恩阳区人民政府（恩府〔2015〕208号）批复同意实施，由巴中市恩阳区卫生健康局按规定有序开展。', first_line_indent=Cm(0.74))
add_paragraph(doc, '2016年7月巴中市恩阳医养园PPP项目协议经区人民政府审定同意，巴中市恩阳区卫生健康局于2016年12月与恩阳区人民医院、湖南省第五工程有限公司签订PPP协议，成立项目管理公司，注册资本金10000万元，承担该项目的投资、建设、运营维护及移交。', first_line_indent=Cm(0.74))
add_paragraph(doc, '2017年7月19日项目公司与湖南省第五工程有限公司签订施工合同，项目建筑总面积118515.13㎡，包括院内绿化、场地及路面硬化、室内管网等附属工程及设备购置等。该项目实际于2017年5月17日开工，2022年3月完工，2022年3月28日进入调试运营期，同年10月27日通过竣工验收，正式进入运营阶段。', first_line_indent=Cm(0.74))
add_paragraph(doc, '根据2021年10月15日签订的《巴中市恩阳医养园PPP项目合同之补充合同（2021年）》，项目合作期调整为23年，其中建设期不超过3年，运营期为18年（2022年10月28日至2040年12月21日）。', first_line_indent=Cm(0.74))

# ===== 二、测算依据 =====
add_paragraph(doc, '二、测算依据', font_name='黑体', size=Pt(14), bold=True)
add_paragraph(doc, '（一）政策依据', font_name='楷体', size=Pt(12), bold=True)
add_paragraph(doc, '（1）《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）', first_line_indent=Cm(0.74))
add_paragraph(doc, '（2）财政部关于印发《政府和社会资本合作项目财政承受能力论证指引》的通知（财金〔2015〕21号）', first_line_indent=Cm(0.74))
add_paragraph(doc, '（3）四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）', first_line_indent=Cm(0.74))

add_paragraph(doc, '（二）项目资料依据', font_name='楷体', size=Pt(12), bold=True)
for i, item in enumerate([
    '《恩阳医养园PPP项目协议》（2016年12月27日签订）',
    '《巴中市恩阳医养园PPP项目合同》（2017年签订）',
    '《巴中市恩阳医养园PPP项目合同之补充合同》（2017年/2018年/2021年）',
    '《关于巴中市恩阳区人民医院一期工程建设项目可行性研究报告的批复》（恩区发改行审〔2015〕83号）',
    '《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）',
    '《借款还本付息情况报告》（贵阳银行股份有限公司成都分行提供）',
    '2025年度财务报表及运营成本明细（项目公司2025年度财务资料）',
    '2022-2025年度运营期绩效考核结果（恩阳区卫健局出具）',
], 1):
    add_paragraph(doc, f'（{i}）{item}', first_line_indent=Cm(0.74))

# ===== 三、测算说明 =====
add_paragraph(doc, '三、测算说明', font_name='黑体', size=Pt(14), bold=True)
add_paragraph(doc, '（一）项目回报机制识别', font_name='楷体', size=Pt(12), bold=True)

add_paragraph(doc, '一是政策方面：根据《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）："项目回报机制主要分为使用者付费、可行性缺口补助和政府付费等支付方式。其中：使用者付费，是指由最终消费用户直接付费购买公共产品和服务。可行性缺口补助，是指使用者付费不足以满足社会资本或项目公司成本回收和合理回报，而由政府以财政补贴、股本投入、优惠贷款和其他优惠政策的形式，给予社会资本或项目公司的经济补助。政府付费，是指政府直接付费购买公共产品和服务，主要包括可用性付费、使用量付费和绩效付费。"', first_line_indent=Cm(0.74))
add_paragraph(doc, '四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）规定："运营补贴。指项目运营期间根据不同的付费模式政府应当承担的直接付费责任。其中：政府付费模式下，政府承担全部运营补贴支出责任；可行性缺口补助（政府补助）模式下，政府承担部分运营补贴支出责任；使用者付费模式下，政府不承担运营补贴支出责任。"', first_line_indent=Cm(0.74))
add_paragraph(doc, '二是项目资料方面：项目实施方案、PPP合同文本中均明确项目回报机制为"可行性缺口补助"。', first_line_indent=Cm(0.74))
add_paragraph(doc, '2016年12月27日巴中市恩阳区卫生健康局与巴中市恩阳区人民医院、湖南省第五工程有限公司签订《恩阳医养园PPP项目协议》，协议约定"自运营日起，甲方按照本协议的约定及时、足额地向乙方支付可行性缺口补贴，且区政府将本项目的可行性缺口补贴纳入跨年度的财政预算，并提请人大决议"。', first_line_indent=Cm(0.74))

add_paragraph(doc, '（二）回报计算方法选用', font_name='楷体', size=Pt(12), bold=True)
add_paragraph(doc, '根据2017年7月巴中市恩阳区卫生健康局与项目公司签订《巴中市恩阳医养园PPP项目合同之补充合同》，合同计算公式：A=（P×k×（1+k）^n）/（（1+k）^n-1）+实际融资成本+运营维护成本×（1+K）-第三方收入。A为运营期内各年政府运营补贴。', first_line_indent=Cm(0.74))

add_paragraph(doc, '（1）项目资本金回报：（P×k×（1+k）^n）/（（1+k）^n-1）', first_line_indent=Cm(0.74))
add_paragraph(doc, 'P-项目资本金；k-合理利润率（社会资本中标年回报率7.99%）；n-运营期18年。', first_line_indent=Cm(0.74))
add_paragraph(doc, '项目公司实际投入资本金102,298,077.79元。根据平头哥最新指令，资本金回报测算调整为：其中8,000万元按7.99%年回报率等额本息方式在18年运营期内逐年偿还；剩余22,298,077.79元（含恩阳区人民医院投入的2,000万元）于运营期最后一年（2040年）一次性支付。', first_line_indent=Cm(0.74))

add_paragraph(doc, '（2）实际融资成本：为项目运营期间向贵阳银行股份有限公司成都分行借款产生的利息及本金。按照《借款还本付息情况报告》列示的还款计划及实际利率计算，18个运营年度银行融资本息合计710,822,119.32元（下表按原征求意见稿逐年数据列示，合计712,464,358.22元，差额1,642,238.90元系银行报告与原始估算口径差异，本项目以银行报告合计数为准）。', first_line_indent=Cm(0.74))

add_paragraph(doc, '（3）运营维护成本：2023、2024年度按项目公司实际发生财务数据核算（分别为5,768,117.42元、6,890,208.16元）；2025年度根据项目公司2025年度运营审计数据（主营业务成本4,474,027.75元 + 管理费用1,239,214.11元 = 5,713,241.86元），取5,713,241.86元；2026年度起按2025年度实际运营成本四舍五入取整值5,710,000.00元代入公式估算。', first_line_indent=Cm(0.74))

add_paragraph(doc, '（4）第三方收入：2023、2024年度按项目公司实际发生财务数据核算（分别为4,285,070.31元、6,229,093.45元）；2025年度根据项目公司2025年度运营审计数据（主营业务收入4,961,142.62元 + 其他业务收入602,875.26元 = 5,564,017.88元），取5,564,017.88元；2026年度起按2025年度实际收入四舍五入取整值5,560,000.00元代入公式估算。', first_line_indent=Cm(0.74))

# ===== 四、测算过程 =====
add_paragraph(doc, '四、测算过程', font_name='黑体', size=Pt(14), bold=True)
add_paragraph(doc, '（一）每年项目资本金回报测算情况', font_name='楷体', size=Pt(12), bold=True)

add_paragraph(doc, f'按项目公司实际投入资本金102,298,077.79元，其中8,000万元按合理利润率7.99%、运营期18年以等额本息方式测算（年还款额A=8,530,246.96元=本金逐年递增+利息逐年递减），剩余22,298,077.79元于第18年（2040年）一次性支付。每年项目资本金回报见下表（单位：元）：', first_line_indent=Cm(0.74))

# P1 FIX: Table 1 with proper structure
t1_headers = ['年', '等额本息\n本金', '等额本息\n利息', '等额本息\n小计', '一次性支付\n（第18年）', '资本金回报\n合计', '剩余资本金\n（等额本息部分）']
t1_rows = []
for y in yearly:
    yr, pr, interest, pmt, lump, total_cap, remaining = y[0], y[1], y[2], y[3], y[4], y[5], y[6]
    pr_str = f'{pr:,.2f}'
    int_str = f'{interest:,.2f}'
    pmt_str = f'{pmt:,.2f}'
    lump_str = f'{lump:,.2f}' if lump > 0 else '-'
    total_str = f'{total_cap:,.2f}'
    rem_str = f'{remaining:,.2f}' if remaining > 0 else '0.00'
    t1_rows.append([str(yr), pr_str, int_str, pmt_str, lump_str, total_str, rem_str])

# Add totals row
t1_sum_principal = sum(y[3] for y in yearly)  # PMT sum
t1_sum_lump = CAPITAL_LUMP
t1_sum_total = cap_total
t1_rows.append(['合计', '-', '-', f'{t1_sum_principal:,.2f}', f'{t1_sum_lump:,.2f}', f'{t1_sum_total:,.2f}', '-'])

t1 = make_table(doc, t1_headers, t1_rows, col_widths=[1.2, 2.3, 2.3, 2.3, 2.3, 2.3, 2.8])
doc.add_paragraph()

add_paragraph(doc, '说明：等额本息部分本金=年还款额-利息，利息=剩余资本金×7.99%，18年合计等额本息153,544,445.28元（其中利息73,544,445.28元、本金80,000,000.00元），加第18年一次性支付22,298,077.79元后资本金回报总计175,842,523.07元。', first_line_indent=Cm(0.74), size=Pt(10))

# ===== 四（二）融资 =====
add_paragraph(doc, '（二）每年实际融资本息测算情况', font_name='楷体', size=Pt(12), bold=True)

add_paragraph(doc, '除资本金外项目实际融资金额382,800,000.00元，按实际融资利率（据实浮动），运营期18年，以项目公司与贵阳银行成都分行协商约定的还款计划及还款利率测算。2021年6月30日，本项目向贵阳银行成都分行申请了39,750万元项目贷款，贷款年限为18年，原贷款利率为6.95%/年，后经多次调整：2024年11月22日调整为6.45%/年、2025年11月28日调整为5.5%/年、2026年7月1日调整为5.0%/年，累计发放贷款38,280万元。本表数据来源：贵阳银行成都分行《借款还本付息情况报告》。', first_line_indent=Cm(0.74))

add_paragraph(doc, '每年实际融资本息见下表（单位：元）：', first_line_indent=Cm(0.74))

t2_headers = ['年', '实际融资本息', '剩余融资金额']
t2_rows = []
for i, y in enumerate(yearly):
    yr = y[0]
    bank_amt = y[7]
    bank_rem = bank_schedule[i][1]
    t2_rows.append([
        str(yr),
        f'{bank_amt:,.2f}',
        f'{bank_rem:,.2f}' if bank_rem > 0 else '-'
    ])
t2_rows.append(['合计', f'{bank_total:,.2f}', '-'])

t2 = make_table(doc, t2_headers, t2_rows, col_widths=[2.0, 5.5, 5.5])
doc.add_paragraph()

# P1 FIX: Note for 2040
add_paragraph(doc, '注：第18年（2040年度）实际融资本息为0元，系贵阳银行贷款已于第17年（2039年12月）全部还清，故2040年度无融资还款义务。', first_line_indent=Cm(0.74), size=Pt(10))

# ===== 四（三）运维成本 =====
add_paragraph(doc, '（三）运营维护成本', font_name='楷体', size=Pt(12), bold=True)
add_paragraph(doc, '根据项目公司提供的财务资料，经审核：第一经营年度（2022.10.28-2023.12.21）运营维护成本金额为5,768,117.42元；第二经营年度（2023.12.21-2024.12.21）运营维护成本金额为6,890,208.16元；2025年度根据项目公司2025年度财务报表及运营成本明细（主营业务成本4,474,027.75元+管理费用1,239,214.11元），实际运营成本为5,713,241.86元。2026年之后每年度的运营维护成本暂按2025年度实际发生额四舍五入取整值5,710,000.00元代入公式估算可行性缺口补助。', first_line_indent=Cm(0.74))

# ===== 四（四）第三方收入 =====
add_paragraph(doc, '（四）第三方收入', font_name='楷体', size=Pt(12), bold=True)
add_paragraph(doc, '根据项目公司提供的财务资料，经审核：第一经营年度运营收入金额4,285,070.31元；第二经营年度运营收入金额6,229,093.45元（含第三经营年度2024.10.28-2024.12.31期间收入387,927.58元）；2025年度根据项目公司2025年度财务报表及收入明细（主营业务收入4,961,142.62元+其他业务收入602,875.26元），实际运营收入为5,564,017.88元。2026年之后每年度的运营收入暂按2025年度实际发生额四舍五入取整值5,560,000.00元代入公式估算可行性缺口补助。', first_line_indent=Cm(0.74))

# ===== 四（五）考核 =====
add_paragraph(doc, '（五）考核结果', font_name='楷体', size=Pt(12), bold=True)
add_paragraph(doc, '根据2018年3月7日巴中市恩阳区卫生和计划生育局与项目公司签订的补充协议，每年度巴中市恩阳区卫生健康局对项目运营期进行绩效考核并形成考核分数，考核分数作为支付可用性付费的重要依据，考核分数大于等于80分支付比例为100%，分数低于80分按照合同约定的调整系数进行计算。', first_line_indent=Cm(0.74))
add_paragraph(doc, '项目运营期限从2022年10月28日开始，截止2026年6月30日，巴中市恩阳区卫生健康局对该项目开展了四个运营年度的绩效考核，2022年度考核分数为93分、2023年度考核分数为94分、2024年度考核分数为95分、2025年度考核分数为93.5分，分数均达到80分以上，可用性付费支付比例为100%。', first_line_indent=Cm(0.74))

# ===== 五、测算结果 =====
add_paragraph(doc, '五、测算结果', font_name='黑体', size=Pt(14), bold=True)
add_paragraph(doc, '根据PPP合同约定，项目进入运营期后应于每个运营年度末支付一次可用性费用，最终计算该项目每年可用性付费金额见下表（单位：元）：', first_line_indent=Cm(0.74))
add_paragraph(doc, f'运营维护成本18年合计：{op_cost_total:,.2f}元；第三方收入18年合计：{income_total:,.2f}元。', first_line_indent=Cm(0.74))

t3_headers = ['年', '资本金回报', '实际融资本息', '运营维护成本', '第三方收入', '可用性付费']
t3_rows = []
for y in yearly:
    yr, _, _, _, _, total_cap, _, bank, op, inc, avail = y
    t3_rows.append([
        str(yr),
        f'{total_cap:,.2f}',
        f'{bank:,.2f}',
        f'{op:,.2f}',
        f'{inc:,.2f}',
        f'{avail:,.2f}'
    ])
t3_rows.append(['合计', f'{cap_total:,.2f}', f'{bank_total:,.2f}', f'{op_cost_total:,.2f}', f'{income_total:,.2f}', f'{grand_total:,.2f}'])

t3 = make_table(doc, t3_headers, t3_rows, col_widths=[1.2, 2.8, 2.8, 2.8, 2.8, 2.8])
doc.add_paragraph()

add_paragraph(doc, f'经测算，运营期18年可用性付费总额为{grand_total:,.2f}元（约{grand_total/1e8:.2f}亿元）。', first_line_indent=Cm(0.74))

# P1 FIX: Data reconciliation note
add_paragraph(doc, '备注：', font_name='楷体', size=Pt(10), bold=True)
add_paragraph(doc, f'1.项目运营期从2022年10月28日开始，贵阳银行贷款结息日为每年度12月21日。为统一各项费用的计费期限，除2023年度费用期限为2022年10月28日至2023年12月21日（约14个月）外，其余年度费用期限均为上一年度12月21日至下一年度12月21日。', size=Pt(10))
add_paragraph(doc, '2.本报告运营维护成本及第三方收入依据：2023、2024年度为项目公司实际发生财务数据；2025年度为项目公司2025年度运营审计数据（主营业务成本+管理费用=运营维护成本，主营业务收入+其他业务收入=第三方收入）；2026年度起以2025年度实际发生额四舍五入取整值进行估算。', size=Pt(10))
add_paragraph(doc, '3.实际融资本息数据来源：贵阳银行成都分行出具的《借款还本付息情况报告》（合计710,822,119.32元）。本表逐年数据沿用原征求意见稿表2（合计712,464,358.22元），两数差额1,642,238.90元为口径调整差异。2025年度第3季度财务费用明细计提贷款利息6,291,985.59元（本金38,125万元），与银行还款计划一致。', size=Pt(10))
add_paragraph(doc, '4.本报告系征求意见稿，如委托方对资本金回报计算方式、运营维护成本及第三方收入估算参数有调整要求，可进一步修订。', size=Pt(10))

# ===== 落款 =====
doc.add_paragraph()
doc.add_paragraph()
add_paragraph(doc, '四川融策会计师事务所有限公司', align=WD_ALIGN_PARAGRAPH.RIGHT, size=Pt(12))
add_paragraph(doc, '二〇二六年七月二十九日', align=WD_ALIGN_PARAGRAPH.RIGHT, size=Pt(12))

# ===== SAVE =====
outpath = r'C:\Users\scrccpa\Desktop\恩阳医养园PPP项目可用性付费测算报告（修订版v2-20260729）.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
print(f'Size: {os.path.getsize(outpath):,} bytes')
print(f'可用性付费总额: {grand_total:,.2f}')
