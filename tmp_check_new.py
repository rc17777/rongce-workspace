import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

filename = r'C:\Users\scrccpa\Desktop\融策制度汇编-业务部管理篇.docx'
doc = Document(filename)

print(f'File: {filename}')
print(f'Paragraphs: {len(doc.paragraphs)}')
print()

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    is_bold = any(r.bold for r in p.runs if r.bold)
    is_large = any(r.font.size and r.font.size >= 140000 for r in p.runs if r.font.size)
    if is_bold or is_large:
        print(f'  [{i}] {t[:120]}')
