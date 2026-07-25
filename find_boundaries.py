import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from docx import Document
from docx.oxml.ns import qn

COMPLETE = r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx'

doc = Document(COMPLETE)
body = doc.element.body

# Find 业务质控篇 section heading (deep in doc)
biz_qc_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '业务质控篇' and i > 500:
        biz_qc_idx = i
        break

# Find end of 财务管理篇 - last paragraph of 可分配利润核算细则
# Look for "第八章 附则" that's part of the last FIN doc, then find the paragraph AFTER it
# The inserted content starts with the reversed RC-BIZ content
# Look for the first BIZ doc heading that appears out of place (before 业务质控篇)

# Strategy: find the last occurrence of something from 财务管理篇 
# The 可分配利润核算细则's 附则 should end with something like "附件"/"签署页"
# Then the next non-empty stuff is the insert

# Let me scan backwards from biz_qc_idx to find where 可分配利润核算细则 ends
insert_start = None
for i in range(biz_qc_idx - 1, 0, -1):
    t = doc.paragraphs[i].text.strip()
    if '可分配利润核算细则' in t or '第八章 附则' in t:
        # This is part of the last FIN doc - find the end of it
        # Walk forward to find where it ends
        for j in range(i, biz_qc_idx):
            t2 = doc.paragraphs[j].text.strip()
            # The inserted content starts when we see BIZ document content
            if '业务部' in t2 and any(r.bold for r in doc.paragraphs[j].runs if r.bold):
                insert_start = j
                break
        if insert_start:
            break

# If we couldn't find it by BIZ marker, try another approach
if insert_start is None:
    # Find the first element that looks like the reversed RC-BIZ content
    # RC-BIZ docs have "RC-BIZ-00" in their heading
    for i in range(biz_qc_idx - 1, 0, -1):
        t = doc.paragraphs[i].text.strip()
        if 'RC-BIZ-008' in t or '业务部经营数据统计与分析制度' in t:
            # Walk back to find the start of reversed RC-BIZ-008 content
            # In reverse order, the chapters come before the title
            for j in range(i - 1, 0, -1):
                t2 = doc.paragraphs[j].text.strip()
                if '可分配利润核算细则' in t2:
                    insert_start = j + 1
                    break
            if insert_start is None:
                insert_start = i - 20  # rough estimate
            break

if insert_start is None:
    print('ERROR: Could not find insertion start')
    sys.exit(1)

insert_end = biz_qc_idx - 1

print(f'业务质控篇 at: {biz_qc_idx}')
print(f'Insert start: {insert_start}')
print(f'Insert end: {insert_end}')
print(f'Block size: {insert_end - insert_start + 1} paragraphs')

# Verify by showing first and last paragraphs of the block
print(f'\nFirst 3 paragraphs in block:')
for i in range(insert_start, min(insert_start + 3, insert_end + 1)):
    print(f'  [{i}] {doc.paragraphs[i].text.strip()[:90]}')
print(f'\nLast 3 paragraphs in block:')
for i in range(max(insert_start, insert_end - 2), insert_end + 1):
    print(f'  [{i}] {doc.paragraphs[i].text.strip()[:90]}')
