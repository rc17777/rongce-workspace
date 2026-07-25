import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

COMPLETE = r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx'

doc = Document(COMPLETE)
body = doc.element.body

# Fix: remove the old "人力资源篇 · 制度体系概览" page (which was copied from the section file)
# and insert proper TOC entries before 人力资源篇 section heading

# Find key positions
toc_heading_idx = None
hr_section_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '目    录' and toc_heading_idx is None:
        toc_heading_idx = i
    if '人力资源篇 · 制度体系概览' in t and i > 10 and i < 100:
        # This is the old overview page from section file - we'll replace its TOC
        pass
    if is_bold and t == '人力资源篇' and i < 100 and toc_heading_idx:
        hr_section_idx = i  # This is the TOC entry

print(f'TOC heading: {toc_heading_idx}')
print(f'HR TOC entry: {hr_section_idx}')

# ============================================
# Rebuild TOC: remove old entries after "目    录"
# and replace with 5 section entries
# ============================================

# Find old overview page (人力资源篇 · 制度体系概览) and remove it
overview_idx = None
for i, p in enumerate(doc.paragraphs):
    if '人力资源篇 · 制度体系概览' in p.text:
        overview_idx = i
        break

if overview_idx:
    # Remove from overview to just before the first actual content
    # The first document is "四川融策薪酬管理制度"
    first_doc_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        is_bold = any(r.bold for r in p.runs if r.bold)
        if is_bold and '四川融策薪酬管理制度' == t:
            # Find the first occurrence (the title, not the code block)
            first_doc_idx = i
            break
    
    if first_doc_idx:
        print(f'Overview at {overview_idx}, first doc at {first_doc_idx}')
        # Remove the overview page (from overview_idx to first_doc_idx-1)
        to_remove = []
        for i in range(overview_idx, first_doc_idx):
            to_remove.append(doc.paragraphs[i]._element)
        
        for elem in to_remove:
            elem.getparent().remove(elem)
        
        print(f'Removed {len(to_remove)} overview page paragraphs')

# Now find the TOC area after the removal
# The TOC is between "目    录" heading and before the first section heading
# After removal, we need to find 人力资源篇 again

toc_items_start = None
toc_items_end = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '目    录':
        toc_items_start = i + 1
    if toc_items_start and i > toc_items_start and is_bold and t == '人力资源篇':
        toc_items_end = i
        break

print(f'TOC items: {toc_items_start} to {toc_items_end}')

if toc_items_start and toc_items_end:
    # Remove old TOC items
    to_remove = []
    for i in range(toc_items_start, toc_items_end):
        to_remove.append(doc.paragraphs[i]._element)
    
    for elem in to_remove:
        elem.getparent().remove(elem)
    
    print(f'Removed {len(to_remove)} old TOC items')
    
    # Find 人力资源篇 again (it's now at a shifted index)
    hr_elem = None
    for p in doc.paragraphs:
        t = p.text.strip()
        is_bold = any(r.bold for r in p.runs if r.bold)
        if is_bold and t == '人力资源篇':
            hr_elem = p._element
            break
    
    if hr_elem:
        toc_labels = ['人力资源篇', '财务管理篇', '业务部管理篇', '业务质控篇', '行政综合篇']
        for label in reversed(toc_labels):
            new_p = copy.deepcopy(hr_elem)
            for r_elem in new_p.findall(qn('w:r')):
                for t_elem in r_elem.findall(qn('w:t')):
                    t_elem.text = label
                    break
            hr_elem.addprevious(new_p)
        
        print(f'Inserted {len(toc_labels)} TOC entries')

# ============================================
# Save
# ============================================
doc.save(COMPLETE)
import os
print(f'\nSaved: {COMPLETE}')
print(f'Size: {os.path.getsize(COMPLETE):,} bytes')

# Quick verify
print('\n=== Final TOC ===')
doc2 = Document(COMPLETE)
for i, p in enumerate(doc2.paragraphs[:30]):
    t = p.text.strip()
    if not t:
        continue
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and i >= 20:
        print(f'  [{i}] {t[:80]}')
