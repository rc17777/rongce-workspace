#!/usr/bin/env python3
"""Extract text from the dike project audit report and run TextQualityFilter."""
import sys, os, glob
sys.path.insert(0, '.')

# Find the actual file using glob (avoids encoding issues with inline paths)
base = r'C:\Users\15528\Desktop\报告初稿-大渡河右岸金川县沙耳段堤防项目--定7.10'
files = glob.glob(os.path.join(base, '**', '*.docx'), recursive=True)
print(f'Found {len(files)} docx files:')
for f in files:
    print(f'  {f} ({os.path.getsize(f)} bytes)')

# Find the main report (竣工结算审核报告)
main_report = None
for f in files:
    if '审核报告' in os.path.basename(f) and '竣工结算' in os.path.basename(f):
        main_report = f
        break

if not main_report:
    # Fallback: take the largest docx
    files.sort(key=lambda f: os.path.getsize(f), reverse=True)
    main_report = files[0]

print(f'\nSelected: {os.path.basename(main_report)}')

# Extract text
import docx
doc = docx.Document(main_report)

text_parts = []
for p in doc.paragraphs:
    if p.text.strip():
        text_parts.append(p.text)

# Extract tables
for i, table in enumerate(doc.tables):
    text_parts.append(f'\n【表格{i+1}】')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        text_parts.append(' | '.join(cells))

full_text = '\n'.join(text_parts)

# Save extracted text
out_dir = r'C:\Users\15528\.openclaw\workspace-main\reports'
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'extracted_dike_report.txt')
with open(out_path, 'w', encoding='utf-8') as f:
    f.write(full_text)

# Also save first 5000 chars for preview
print(f'\nExtracted: {len(full_text)} chars, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')
print(f'Saved to: {out_path}')
print(f'\n=== First 2000 chars preview ===')
print(full_text[:2000])
