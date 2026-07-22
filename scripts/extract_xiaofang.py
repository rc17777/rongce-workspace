# -*- coding: utf-8 -*-
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import load_workbook

ROOT = Path(r"C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722\成都市郫都区消防救援局2025年度部门预算项目")
REPORT = next(ROOT.glob("*.docx"))
WP = next(ROOT.glob("*.xlsx"))

def vb(val): return "" if val is None else str(val).strip()

doc = Document(REPORT)
paras = []
for i, p in enumerate(doc.paragraphs, 1):
    if not p.text.strip(): continue
    ppr = p._p.pPr
    lr = ppr.spacing.get(qn("w:lineRule")) if ppr is not None and ppr.spacing is not None else ""
    am = {WD_ALIGN_PARAGRAPH.CENTER:"居中", WD_ALIGN_PARAGRAPH.LEFT:"左对齐", WD_ALIGN_PARAGRAPH.RIGHT:"右对齐", WD_ALIGN_PARAGRAPH.JUSTIFY:"两端对齐"}
    runs = []
    for r in p.runs:
        if not r.text.strip(): continue
        f = r.font; ea = r._element.rPr.rFonts.get(qn("w:eastAsia")) if r._element.rPr is not None and r._element.rPr.rFonts is not None else None
        runs.append({"t":r.text,"font":f.name or "","ea":ea or "","sz":round(f.size.pt,1) if f.size else None,"b":bool(f.bold)})
    paras.append({"i":i,"text":p.text.strip(),"a":am.get(p.alignment,"未显式"),"lr":lr,"runs":runs})

tabs = [[[vb(c.text) for c in row.cells] for row in t.rows] for t in doc.tables]

wb = load_workbook(WP, data_only=False)
sheets = []
for s in wb.worksheets:
    rows = [[vb(c.value) for c in row] for row in s.iter_rows()]
    sheets.append({"name":s.title,"rows":[r for r in rows if any(v not in ("","") for v in r)]})

for p in paras:
    t = p["text"]
    if any(k in t for k in ["一、","二、","三、","四、","五、","2025","得分","复核","偏离","问题","建议","消防"]):
        fonts = "/".join(sorted({(r["ea"] or r["font"]) for r in p["runs"] if (r["ea"] or r["font"])}))
        print(f"[{p['i']}] align={p['a']} font={fonts or '--'} | {t[:130]}")
print("\n=== SHEETS ===")
for s in sheets:
    print(f"\n--- {s['name']} ({len(s['rows'])} rows) ---")
    for i,r in enumerate(s["rows"]): print(f"  R{i+1}: {' | '.join(r[:10])}")
for i,t in enumerate(tabs,1):
    if i <= 3:
        print(f"\n--- Table {i} ({len(t)} rows) ---")
        for r in t[:3]: print(f"  {' | '.join(r[:8])}")
