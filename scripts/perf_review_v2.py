# -*- coding: utf-8 -*-
"""
绩效自评复核 V2 — 每个单位一个Excel，逐项目逐指标14条标准检查
"""
import sys, os, re, json
from pathlib import Path
from collections import defaultdict, Counter
from copy import copy

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, numbers
from openpyxl.utils import get_column_letter

BASE = Path(r'C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722')
DESKTOP = Path(r'C:\Users\scrccpa\Desktop')

# ============================================================
# 样式
# ============================================================
HEADER_FONT = Font(name='黑体', size=11, bold=True, color='FFFFFF')
HEADER_FILL = PatternFill(start_color='0A1F3F', end_color='0A1F3F', fill_type='solid')
HEADER_ALIGN = Alignment(horizontal='center', vertical='center', wrap_text=True)
BODY_FONT = Font(name='仿宋', size=10)
BODY_ALIGN = Alignment(vertical='center', wrap_text=True)
THIN_BORDER = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)
P0_FILL = PatternFill(start_color='FFB3B3', end_color='FFB3B3', fill_type='solid')  # 红 - 致命
P1_FILL = PatternFill(start_color='FFE5A3', end_color='FFE5A3', fill_type='solid')  # 橙 - 重要
P2_FILL = PatternFill(start_color='E8E8E8', end_color='E8E8E8', fill_type='solid')  # 灰 - 优化
OK_FILL  = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')  # 绿 - 通过

SUMMARY_HEADERS = ['序号','标准','项目','指标','指标性质','指标值','权重','自评得分','复核得分','问题描述','严重度','建议']

# ============================================================
# 工具函数
# ============================================================
def safe_str(v):
    if v is None: return ''
    return str(v).strip()

def safe_float(v, default=None):
    if v is None: return default
    try: return float(v)
    except: return default

def is_project_sheet(ws, sn):
    """判断是否为项目评分表（基于表头检测，非名称白名单）"""
    if ws.max_row < 5:
        return False
    # 汇总/清单类sheet名称跳过
    skip_names = {'自评复核报告表', '目标完成，偏离度', '汇总', '差值表'}
    if sn in skip_names:
        return False
    # 检查是否有标准项目评分表表头: 一级指标|二级指标|三级指标
    for ri in range(1, min(10, ws.max_row + 1)):
        c1 = safe_str(ws.cell(row=ri, column=1).value)
        c2 = safe_str(ws.cell(row=ri, column=2).value)
        c3 = safe_str(ws.cell(row=ri, column=3).value)
        if '一级指标' in c1 and '二级指标' in c2 and '三级指标' in c3:
            # 确认有复核得分列
            c11 = safe_str(ws.cell(row=ri, column=11).value)
            if '复核得分' in c11 or '复核' in c11:
                return True
    return False

def extract_project_name(ws):
    """从xlsx sheet提取项目名"""
    # Row 2 usually contains project name
    for ri in range(1, min(5, ws.max_row + 1)):
        cell_val = safe_str(ws.cell(row=ri, column=1).value)
        if cell_val and len(cell_val) > 8 and ('项目' in cell_val or '经费' in cell_val or '资金' in cell_val):
            # Clean up: remove prefix like "51011723T000007943391-"
            cleaned = re.sub(r'^\d+[A-Z]?\d*-?', '', cell_val)
            if len(cleaned) > 5:
                return cleaned
            return cell_val
    # Fallback: get merged cell value from row 2
    val = safe_str(ws.cell(row=2, column=1).value)
    if val:
        cleaned = re.sub(r'^\d+[A-Z]?\d*-?', '', val)
        return cleaned if len(cleaned) > 3 else val
    return ''

def parse_header_row(ws):
    """定位表头行并映射列"""
    for ri in range(1, min(10, ws.max_row + 1)):
        row_vals = [safe_str(ws.cell(row=ri, column=c).value) for c in range(1, 14)]
        if '三级指标' in row_vals[2] and ('权重' in ''.join(row_vals) or '自评得分' in ''.join(row_vals)):
            return ri, {c: safe_str(ws.cell(row=ri, column=c).value) for c in range(1, 14)}
    return None, {}

def read_indicator_rows(ws, header_row):
    """读取所有指标行数据"""
    rows = []
    for ri in range(header_row + 1, ws.max_row + 1):
        indicator = safe_str(ws.cell(row=ri, column=3).value)
        if not indicator or indicator == 'None':
            continue
        # Check for 合计 row
        level1 = safe_str(ws.cell(row=ri, column=1).value)
        if '合计' in level1 or '合计' in indicator:
            continue
        
        # 预算执行行（特殊，但也要检查）
        row = {
            'row': ri,
            '一级指标': level1,
            '二级指标': safe_str(ws.cell(row=ri, column=2).value),
            '三级指标': indicator,
            '指标性质': safe_str(ws.cell(row=ri, column=4).value),
            '指标值': ws.cell(row=ri, column=5).value,
            '度量单位': safe_str(ws.cell(row=ri, column=6).value),
            '权重': ws.cell(row=ri, column=7).value,
            '自评完成值': ws.cell(row=ri, column=8).value,
            '自评得分': ws.cell(row=ri, column=9).value,
            '复核完成值': ws.cell(row=ri, column=10).value,
            '复核得分': ws.cell(row=ri, column=11).value,
            '偏离度': ws.cell(row=ri, column=12).value,
            '备注': safe_str(ws.cell(row=ri, column=13).value),
        }
        rows.append(row)
    return rows

# ============================================================
# 14条标准检查（逐指标）
# ============================================================

def check_criteria(indicator_row, all_rows_in_project, project_name):
    """
    对单个指标行执行全部14条标准中适用的检查
    返回: list of issues
    """
    issues = []
    r = indicator_row
    
    ind  = r['三级指标']
    lev1 = r['一级指标']
    lev2 = r['二级指标']
    nat  = r['指标性质']
    tval = r['指标值']
    unit = r['度量单位']
    wgt  = safe_float(r['权重'], 0)
    sv   = r['自评完成值']
    ss   = safe_float(r['自评得分'], 0)
    rv   = r['复核完成值']
    rs   = safe_float(r['复核得分'], -1)
    rem  = r['备注']
    dev  = r['偏离度']
    
    is_deducted = (rs >= 0 and rs < ss)  # 复核扣了分
    is_zero = (rs == 0 and ss > 0)       # 直接打0
    is_blank = (rs < 0)                   # 复核得分留空
    is_budget_exec = ('预算执行' in ind)
    
    # 复核得分留空但有备注说明 → 特殊标记
    if is_blank and rem:
        issues.append({
            '标准': '复核得分留空',
            '问题': f'复核得分留空，但备注有说明: "{rem[:60]}"',
            '严重度': 'P1',
            '建议': '复核得分不应留空，需补填或确认是否打0分',
            '备注原文': rem[:80]
        })
    
    # ---- 标准1: 没得资料证明的，打0分 ----
    if is_zero and rem:
        # 检查备注是否说明无资料
        no_data_kw = ['无资料', '无佐证', '未提供佐证', '未提供资料', '无证明', '无数据', 
                      '无法提供', '无可供', '缺失', '无相关', '未收集']
        has_no_data = any(kw in rem for kw in no_data_kw)
        has_no_data_alt = ('佐证' in rem and ('未' in rem or '无' in rem))
        
        # 排除：备注解释了因资金/调整/客观原因未实施（非资料缺失）
        work_not_done = any(kw in rem for kw in ['无剩余资金', '无资金', '未实施', '不再实施',
                                                   '调整未', '未再组织', '已取消', '已调整',
                                                   '未开展', '已停止', '暂停'])
        
        if not has_no_data and not has_no_data_alt and not work_not_done:
            issues.append({
                '标准': '标准1',
                '问题': f'复核得分=0但备注未明确说明"无佐证资料/无法提供佐证"',
                '严重度': 'P2',
                '建议': '备注应明确写"无佐证资料，按复核口径打0分"',
                '备注原文': rem[:80]
            })
        elif work_not_done and not has_no_data:
            # 因工作未实施打0分是合理的，但备注写法可优化
            issues.append({
                '标准': '标准1',
                '问题': f'复核得分=0，备注说明"{rem[:30]}..."（非"无资料"口径，建议统一为"项目未实施，无佐证资料"）',
                '严重度': 'P2',
                '建议': '统一备注口径，建议写"项目未实施，无佐证资料，按复核口径打0分"',
                '备注原文': rem[:80]
            })
    
    # ---- 标准2: 指标设置跟项目无关，打0分 ----
    if '无关' in rem or '与项目不相关' in rem or '不相关' in rem:
        if rs > 0:
            issues.append({
                '标准': '标准2',
                '问题': f'指标与项目无关但复核未打0分(复核得分={rs})',
                '严重度': 'P0',
                '建议': '与项目无关的指标按口径应打0分',
                '备注原文': rem[:80]
            })
    
    # ---- 标准3: 指标不具有可考查性，酌情扣分 ----
    # 检查扣分比例是否与严重程度匹配
    unverifiable_kw = ['无可考查', '无法考查', '不具考核性', '不具可考', '空洞']
    if any(kw in rem for kw in unverifiable_kw):
        deducted_pct = (ss - rs) / max(ss, 1) * 100 if ss > 0 else 0
        # 检查备注中的扣分说明
        if '100%' in rem or '扣100%' in rem:
            # 严重：应该全扣
            if deducted_pct < 95:
                issues.append({
                    '标准': '标准3',
                    '问题': f'备注标注"严重/扣100%"但实际只扣{deducted_pct:.0f}%(自评{ss}→复核{rs})',
                    '严重度': 'P0',
                    '建议': '严重无考查性应按100%扣分',
                    '备注原文': rem[:80]
                })
        elif '60%' in rem or '扣60%' in rem:
            if deducted_pct > 0 and abs(deducted_pct - 60) > 10:
                issues.append({
                    '标准': '标准3',
                    '问题': f'备注标注"扣60%"但实际扣{deducted_pct:.0f}%，比例不一致',
                    '严重度': 'P1',
                    '建议': '确认扣分比例与标准一致',
                    '备注原文': rem[:80]
                })
        elif '30%' in rem or '扣30%' in rem:
            if deducted_pct > 0 and abs(deducted_pct - 30) > 10:
                issues.append({
                    '标准': '标准3',
                    '问题': f'备注标注"扣30%"但实际扣{deducted_pct:.0f}%，比例不一致',
                    '严重度': 'P2',
                    '建议': '确认扣分比例与标准一致',
                    '备注原文': rem[:80]
                })
    
    # ---- 标准4: 满意度无资料的，打0分 ----
    is_satisfaction = ('满意' in ind or '满意度' in ind or 
                       ('满意' in lev1) or ('满意' in lev2))
    if is_satisfaction and is_deducted:
        # 检查是否无资料
        if any(kw in rem for kw in ['无资料', '无佐证', '未提供', '无数据', '无调查', '未调查']):
            if rs > 0:
                issues.append({
                    '标准': '标准4',
                    '问题': f'满意度指标无资料但复核未打0分(复核得分={rs})',
                    '严重度': 'P0',
                    '建议': '满意度无资料应按口径打0分',
                    '备注原文': rem[:80]
                })
    
    # ---- 标准5: 完成值超过130%，不扣分 ----
    if is_deducted and rs >= 0:
        # 检查偏离度或完成值
        try:
            dev_val = safe_float(dev)
            if dev_val and dev_val > 1.3:
                issues.append({
                    '标准': '标准5',
                    '问题': f'完成值偏离度{dev_val:.0%}>130%，按口径不扣分但复核扣了分',
                    '严重度': 'P1',
                    '建议': '完成值超130%不扣分，需修正复核评分',
                    '备注原文': rem[:80]
                })
        except:
            pass
    
    # ---- 标准6: 数量指标设为定性指标，但与项目相关，不扣分 ----
    if '数量指标' in lev2 and nat == '定性' and is_deducted:
        if '无关' not in rem and '不相关' not in rem:
            # 定性数量指标与项目相关不应扣分
            issues.append({
                '标准': '标准6',
                '问题': f'数量指标设为定性但与项目相关，按口径不扣分但复核扣了分',
                '严重度': 'P1',
                '建议': '定性数量指标与项目相关不扣分，应按比例复核评分',
                '备注原文': rem[:80]
            })
    
    # ---- 标准7: 多目标任务，仅个别未完成，时效指标酌情扣30% ----
    if '时效' in lev2 and is_deducted:
        deducted_pct = (ss - rs) / max(ss, 1) * 100 if ss > 0 else 0
        if deducted_pct > 0 and abs(deducted_pct - 30) > 10:
            issues.append({
                '标准': '标准7',
                '问题': f'时效指标扣{deducted_pct:.0f}%，按口径应为"酌情扣30%"，差异较大',
                '严重度': 'P1',
                '建议': '多目标任务仅个别未完成，时效指标酌情扣30%',
                '备注原文': rem[:80]
            })
    
    # ---- 标准8: 效益指标设为定性，虽产生效益但无具体标准，不扣分 ----
    is_benefit = ('效益' in lev1 or '效益指标' in lev1 or '效益' in lev2)
    if is_benefit and nat == '定性' and is_deducted:
        # 检查是否描述了产生了效益
        has_benefit = any(kw in rem for kw in ['产生', '有', '达到', '实现', '完成', '良好', 
                                                '合格', '有效', '明显', '提升', '改善', '增强',
                                                '协助', '解决', '促进'])
        if has_benefit:
            issues.append({
                '标准': '标准8',
                '问题': f'效益定性指标产生效益但复核扣分(自评{ss}→复核{rs})',
                '严重度': 'P1',
                '建议': '效益定性指标产生效益不扣分，除非完全无效益',
                '备注原文': rem[:80]
            })
    
    # ---- 标准9: 产出指标、效益指标重复设置 ----
    # 在项目级别检查（外部处理）
    
    # ---- 标准10: 三级指标与二级指标定义不匹配 ----
    money_in_qty = ('资金支付' in ind or '资金拨付' in ind or '预算执行' in ind or 
                    '资金执行' in ind)
    if '数量指标' in lev2 and money_in_qty:
        if rs > 0:
            issues.append({
                '标准': '标准10',
                '问题': f'数量指标三级设为"资金支付类"({ind})，未反映公共产品或服务数量，应打0分',
                '严重度': 'P0',
                '建议': '三级指标与二级指标定义不匹配，按口径打0分',
                '备注原文': rem[:80]
            })
        else:
            # 已经0分但备注未说明原因
            if '不匹配' not in rem and '数量' not in rem:
                issues.append({
                    '标准': '标准10',
                    '问题': f'数量指标三级设为"资金支付类"已打0分，但备注未说明原因',
                    '严重度': 'P2',
                    '建议': '备注应写明"三级指标与二级指标定义不匹配"',
                    '备注原文': rem[:80]
                })
    
    # ---- 标准14: 自评复核扣分原因描述是否符合口径 ----
    if is_deducted and rs >= 0:
        deducted_pct = (ss - rs) / max(ss, 1) * 100
        
        # 14a: 有扣分但备注为空
        if not rem:
            issues.append({
                '标准': '标准14',
                '问题': f'复核扣分{deducted_pct:.0f}%(自评{ss}→复核{rs})，但备注为空，未说明扣分原因',
                '严重度': 'P1',
                '建议': '每项扣分必须写明原因，引用对应复核口径条款',
                '备注原文': ''
            })
        else:
            # 14b: 备注过于笼统（短备注且无具体规则引用）
            vague_kw = ['酌情', '指标设置', '指标内容']
            has_specific_rule = any(kw in rem for kw in ['扣', '无资料', '无佐证', '不匹配', '定性', '数量', '不具可考', '不具考核', '未提供'])
            if any(kw in rem for kw in vague_kw) and len(rem) < 15 and not has_specific_rule:
                issues.append({
                    '标准': '标准14',
                    '问题': f'备注过于笼统: "{rem}"，未引用具体口径条款',
                    '严重度': 'P2',
                    '建议': '扣分原因应写明：①对应哪条复核口径 ②具体扣分比例 ③判定依据',
                    '备注原文': rem[:80]
                })
            
            # 14c: 备注写了扣分比例但实际不一致
            pct_in_remark = re.findall(r'扣\s*(\d+)\s*%', rem)
            if pct_in_remark:
                stated_pct = float(pct_in_remark[0])
                if abs(deducted_pct - stated_pct) > 15:
                    issues.append({
                        '标准': '标准14',
                        '问题': f'备注写"扣{stated_pct:.0f}%"但实际扣{deducted_pct:.0f}%，比例不一致',
                        '严重度': 'P1',
                        '建议': f'修正备注中扣分比例或修正复核得分，保持一致性',
                        '备注原文': rem[:80]
                    })
            
            # 14d: 备注提到某条口径规则但实际得分与规则不符
            if ('无佐证' in rem or '未提供佐证' in rem or '无资料' in rem) and rs > 0:
                issues.append({
                    '标准': '标准14',
                    '问题': f'备注写"{rem[:30]}..."说明无佐证资料但复核未打0分(得分={rs})，违反口径第1条',
                    '严重度': 'P0',
                    '建议': '无佐证资料应按口径第1条打0分',
                    '备注原文': rem[:80]
                })
            if ('无关' in rem or '不相关' in rem) and rs > 0:
                issues.append({
                    '标准': '标准14',
                    '问题': f'备注提到"无关/不相关"但复核未打0分(得分={rs})，违反口径第2条',
                    '严重度': 'P0',
                    '建议': '指标与项目无关应按口径第2条打0分',
                    '备注原文': rem[:80]
                })
            if ('未调查' in rem or '未问卷' in rem or '未进行满意度' in rem):
                if ('满意' in ind or '满意度' in ind) and rs > 0:
                    issues.append({
                        '标准': '标准14',
                        '问题': f'满意度指标备注提到"{rem[:30]}..."但复核未打0分(得分={rs})，违反口径第4条',
                        '严重度': 'P0',
                        '建议': '满意度无资料应按口径第4条打0分',
                        '备注原文': rem[:80]
                    })
    
    return issues


def check_deduction_reasons(indicator_row):
    """标准14快捷版: 独立检查扣分原因（用于主循环）"""
    # Already integrated into check_criteria above
    return []

def check_dup_indicators(all_rows, project_name):
    """标准9: 检查产出/效益指标重复"""
    issues = []
    
    # 按指标名分组
    indicator_map = defaultdict(list)
    for r in all_rows:
        ind = r['三级指标']
        lev1 = r['一级指标']
        if ind:
            indicator_map[ind].append(lev1)
    
    for ind, lev1_list in indicator_map.items():
        if len(lev1_list) >= 2:
            has_output = any('产出' in l for l in lev1_list)
            has_benefit = any('效益' in l for l in lev1_list)
            if has_output and has_benefit:
                issues.append({
                    '标准': '标准9',
                    '问题': f'产出指标与效益指标重复设置:"{ind}"，应按规则只计其一，另一打0分',
                    '严重度': 'P1',
                    '建议': '判断更适用于产出还是效益，保留一个正常评分，另一个打0分',
                    '备注原文': ''
                })
    
    return issues


# ============================================================
# 格式检查
# ============================================================
def check_docx_format(docx_path):
    """标准13: 报告格式检查"""
    issues = []
    doc = Document(docx_path)
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        # 标题居中（封面页）
        is_cover = (i < 30)
        
        # 一级标题: 一、二、三、...
        if re.match(r'^[一二三四五六七八九十]、', text) and len(text) < 50:
            # 应使用黑体
            for r in p.runs:
                if r.text.strip():
                    fname = (r.font.name or '')
                    if '黑体' not in fname and '黑體' not in fname:
                        issues.append({
                            '位置': f'正文P{i}', 
                            '问题': f'一级标题"{text[:25]}"字体应为黑体，实际={fname}',
                            '严重度': 'P1'
                        })
                        break
        
        # 二级标题: （一）（二）...
        if re.match(r'^（[一二三四五六七八九十\d]+）', text) and len(text) < 50:
            for r in p.runs:
                if r.text.strip():
                    fname = (r.font.name or '')
                    if '楷体' not in fname and '楷體' not in fname:
                        issues.append({
                            '位置': f'正文P{i}',
                            '问题': f'二级标题"{text[:25]}"字体应为楷体GB2312加粗，实际={fname}',
                            '严重度': 'P1'
                        })
                        break
                    if not r.font.bold:
                        issues.append({
                            '位置': f'正文P{i}',
                            '问题': f'二级标题"{text[:25]}"应为加粗',
                            '严重度': 'P2'
                        })
        
        # 正文段落检查（非标题）
        is_heading = re.match(r'^[一二三四五六七八九十]、|^（[^）]+）|^\d+[\.\s]', text)
        if not is_heading and len(text) > 10:
            # 检查字体
            for r in p.runs:
                if r.text.strip() and len(r.text.strip()) > 2:
                    fname = (r.font.name or '')
                    is_digit = re.match(r'^[\d\s.,，。%+\-（）()/]+$', r.text.strip())
                    if not is_digit:
                        if '仿宋' not in fname and 'fangsong' not in fname.lower() and fname:
                            issues.append({
                                '位置': f'正文P{i}',
                                '问题': f'正文字体应为仿宋GB2312，实际={fname}',
                                '严重度': 'P2'
                            })
                            break
            
            # 检查行距（只检查第一个有文本的run所在段落）
            pf = p.paragraph_format
            if pf.line_spacing and isinstance(pf.line_spacing, int):
                ls_pt = pf.line_spacing / 635  # EMU to pt
                if abs(ls_pt - 30) > 3:
                    issues.append({
                        '位置': f'正文P{i}',
                        '问题': f'行距应为30磅固定值，实际≈{ls_pt:.0f}磅',
                        '严重度': 'P2'
                    })
                    break  # 只报一次
    
    # 页面设置
    for si, section in enumerate(doc.sections):
        pw = section.page_width
        ph = section.page_height
        # A4 ≈ 1190700 EMU (21cm) × 1683800 EMU (29.7cm)
        if pw and abs(pw - 1190700) > 50000:
            issues.append({
                '位置': f'页面设置S{si}',
                '问题': f'页面宽度{pw/35000:.1f}cm，非A4标准(21cm)',
                '严重度': 'P1'
            })
        if ph and abs(ph - 1683800) > 50000:
            issues.append({
                '位置': f'页面设置S{si}',
                '问题': f'页面高度{ph/35000:.1f}cm，非A4标准(29.7cm)',
                '严重度': 'P1'
            })
    
    return issues


# ============================================================
# 报告↔底稿一致性
# ============================================================
def check_report_worksheet_consistency(docx_path, xlsx_sheets):
    """标准11: 报告与盖章底稿信息一致性"""
    issues = []
    doc = Document(docx_path)
    
    # Extract project list from docx Table 0 (项目选取表), excluding 合计 rows
    docx_projects = []
    for t in doc.tables:
        if t.rows:
            h = [c.text.strip() for c in t.rows[0].cells]
            if '项目名称' in str(h) or ('序号' in str(h[0]) and '项目' in str(h[1] if len(h) > 1 else '')):
                for row in t.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    # Skip 合计/小计 rows
                    first_cell = cells[0] if cells else ''
                    if '合计' in first_cell or '小计' in first_cell:
                        continue
                    if len(cells) > 1 and cells[1] and cells[1] != '项目名称':
                        docx_projects.append(cells[1])
                break
    
    # Compare counts
    if docx_projects and xlsx_sheets:
        if len(docx_projects) != len(xlsx_sheets):
            issues.append({
                '标准': '标准11',
                '问题': f'报告选取项目数({len(docx_projects)})≠盖章底稿项目数({len(xlsx_sheets)})',
                '严重度': 'P0',
                '建议': '核实哪边多/少了项目，修正一致'
            })
    
    return issues


# ============================================================
# 报告逻辑检查
# ============================================================
def check_report_logic(docx_path):
    """标准12: 报告内容逻辑问题"""
    issues = []
    doc = Document(docx_path)
    full = '\n'.join([p.text for p in doc.paragraphs])
    
    # 检查指标分布百分比合计
    # Pattern: "完全一致...占XX%" "一般偏差...占XX%"
    pcts = re.findall(r'占\s*(\d+\.?\d*)\s*%', full)
    if len(pcts) >= 3:
        try:
            total = sum(float(x) for x in pcts[:4])  # Take first 4
            if abs(total - 100) > 5:
                issues.append({
                    '标准': '标准12',
                    '问题': f'指标分布百分比合计≈{total:.1f}%≠100%（{pcts[:4]}）',
                    '严重度': 'P1',
                    '建议': '核实各分类指标计数，确保百分比合计100%'
                })
        except:
            pass
    
    # 检查项目编号是否连续
    # (Skip - pattern varies)
    
    # 检查问题数与建议数匹配
    if '四、' in full and '五、' in full:
        part4 = full.split('四、')[1]
        part5 = full.split('五、')[1] if '五、' in full else ''
        issue_items = len(re.findall(r'[（(][一二三四五六七八九十\d]+[）)]', part4.split('五、')[0] if '五、' in part4 else part4))
        suggest_items = len(re.findall(r'[（(][一二三四五六七八九十\d]+[）)]', part5))
        if issue_items > 0 and suggest_items > 0:
            if suggest_items < issue_items - 1:
                issues.append({
                    '标准': '标准12',
                    '问题': f'建议条数({suggest_items})明显少于问题条数({issue_items})，可能存在建议覆盖不全',
                    '严重度': 'P1',
                    '建议': '每类问题应有对应建议，检查是否遗漏'
                })
    
    return issues


# ============================================================
# 主流程
# ============================================================

def process_unit(unit_dir):
    """处理一个单位"""
    unit_name = unit_dir.name
    print(f'\n{"="*60}')
    print(f'📋 {unit_name}')
    
    # 找文件
    docx_files = list(unit_dir.glob('*.docx'))
    xlsx_files = list(unit_dir.glob('*.xlsx'))
    if not docx_files or not xlsx_files:
        print(f'  ⚠ 缺少文件，跳过')
        return None
    
    docx_path = str(docx_files[0])
    xlsx_path = str(xlsx_files[0])
    
    # 读取xlsx
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    
    # 筛选项目sheet
    project_sheets = [(sn, wb[sn]) for sn in wb.sheetnames if is_project_sheet(wb[sn], sn)]
    
    all_issues = []  # 所有问题汇总
    
    # === 11: 报告与底稿一致性 ===
    s11 = check_report_worksheet_consistency(docx_path, [sn for sn, _ in project_sheets])
    for iss in s11:
        iss['项目'] = '全局'
        iss['指标'] = ''
        iss['备注原文'] = ''
    all_issues.extend(s11)
    print(f'  标准11 一致性: {len(s11)}个问题')
    
    # === 12: 逻辑检查 ===
    s12 = check_report_logic(docx_path)
    for iss in s12:
        iss['项目'] = '全局'
        iss['指标'] = ''
        iss['备注原文'] = ''
    all_issues.extend(s12)
    print(f'  标准12 逻辑: {len(s12)}个问题')
    
    # === 13: 格式检查 ===
    s13 = check_docx_format(docx_path)
    # Filter to P0/P1 only for main summary
    s13_main = [x for x in s13 if x['严重度'] in ('P0', 'P1')]
    for iss in s13_main:
        iss['标准'] = '标准13'
        iss['项目'] = '全局'
        iss['指标'] = iss.get('位置', '')
        iss['建议'] = '按格式要求修正'
        iss['备注原文'] = ''
    all_issues.extend(s13_main)
    print(f'  标准13 格式: {len(s13_main)}个P0/P1问题（共{len(s13)}个）')
    
    # === 逐项目检查（标准1-10, 14） ===
    project_results = {}  # {project_name: [issues]}
    
    for sn, ws in project_sheets:
        pname = extract_project_name(ws)
        if not pname:
            pname = sn
        
        print(f'  项目 [{sn}] {pname[:50]}...')
        
        header_row, _ = parse_header_row(ws)
        if header_row is None:
            print(f'    ⚠ 找不到表头，跳过')
            continue
        
        indicator_rows = read_indicator_rows(ws, header_row)
        if not indicator_rows:
            print(f'    ⚠ 无指标数据')
            continue
        
        proj_issues = []
        
        # 标准1-8, 10, 14（逐指标）
        for r in indicator_rows:
            iss = check_criteria(r, indicator_rows, pname)
            for i in iss:
                i['项目'] = pname
                i['指标'] = r['三级指标']
                i['指标性质'] = r['指标性质']
                i['指标值'] = str(r['指标值']) if r['指标值'] is not None else ''
                i['权重'] = str(r['权重']) if r['权重'] is not None else ''
                i['自评得分'] = str(r['自评得分']) if r['自评得分'] is not None else ''
                i['复核得分'] = str(r['复核得分']) if r['复核得分'] is not None else ''
            proj_issues.extend(iss)
        
        # 标准9：重复指标
        s9 = check_dup_indicators(indicator_rows, pname)
        for i in s9:
            i['项目'] = pname
            i['指标'] = i['问题'].split('"')[1] if '"' in i['问题'] else ''
            i['指标性质'] = ''
            i['指标值'] = ''
            i['权重'] = ''
            i['自评得分'] = ''
            i['复核得分'] = ''
            i['备注原文'] = ''
        proj_issues.extend(s9)
        
        project_results[pname] = proj_issues
        all_issues.extend(proj_issues)
        
        print(f'    {len(indicator_rows)}个指标, {len(proj_issues)}个问题')
    
    # === 生成Excel ===
    safe_name = re.sub(r'[\\/*?:\"<>|]', '_', unit_name)[:40]
    out_path = DESKTOP / f'绩效自评复核_{safe_name}.xlsx'
    
    out_wb = openpyxl.Workbook()
    
    # --- Sheet 1: 复核总览 ---
    ws_summary = out_wb.active
    ws_summary.title = '复核总览'
    
    # 标题
    ws_summary.merge_cells('A1:L1')
    title_cell = ws_summary.cell(row=1, column=1, value=f'{unit_name} — 绩效自评复核报告')
    title_cell.font = Font(name='黑体', size=14, bold=True)
    title_cell.alignment = Alignment(horizontal='center')
    
    # 统计区
    row = 3
    p0_count = sum(1 for x in all_issues if x['严重度'] == 'P0')
    p1_count = sum(1 for x in all_issues if x['严重度'] == 'P1')
    p2_count = sum(1 for x in all_issues if x['严重度'] == 'P2')
    
    stats = [
        ('复核日期', '2026年7月22日'),
        ('复核项目数', len(project_results)),
        ('复核指标数', sum(len(read_indicator_rows(ws, parse_header_row(ws)[0])) 
                        for sn, ws in project_sheets if parse_header_row(ws)[0])),
        ('', ''),
        ('P0 致命', f'{p0_count} 项 — 提交前必须修改'),
        ('P1 重要', f'{p1_count} 项 — 建议修改'),
        ('P2 优化', f'{p2_count} 项 — 可优化'),
        ('合计', f'{len(all_issues)} 项'),
    ]
    for label, val in stats:
        c1 = ws_summary.cell(row=row, column=1, value=label)
        c1.font = Font(name='黑体', size=11, bold=True)
        c2 = ws_summary.cell(row=row, column=2, value=val)
        c2.font = BODY_FONT
        row += 1
    
    row += 1
    ws_summary.cell(row=row, column=1, value='=== 复核结果明细 ===').font = Font(name='黑体', size=11, bold=True)
    row += 1
    
    # 表头
    for c, h in enumerate(SUMMARY_HEADERS, 1):
        cell = ws_summary.cell(row=row, column=c, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = HEADER_ALIGN
        cell.border = THIN_BORDER
    row += 1
    
    # 数据
    p0_list = [x for x in all_issues if x['严重度'] == 'P0']
    p1_list = [x for x in all_issues if x['严重度'] == 'P1']
    p2_list = [x for x in all_issues if x['严重度'] == 'P2']
    sorted_issues = p0_list + p1_list + p2_list
    
    for idx, iss in enumerate(sorted_issues, 1):
        vals = [
            idx,
            iss.get('标准', ''),
            iss.get('项目', ''),
            iss.get('指标', ''),
            iss.get('指标性质', ''),
            iss.get('指标值', ''),
            iss.get('权重', ''),
            iss.get('自评得分', ''),
            iss.get('复核得分', ''),
            iss.get('问题', ''),
            iss.get('严重度', ''),
            iss.get('建议', ''),
        ]
        for c, v in enumerate(vals, 1):
            cell = ws_summary.cell(row=row, column=c, value=str(v))
            cell.font = BODY_FONT
            cell.alignment = BODY_ALIGN
            cell.border = THIN_BORDER
            sev = iss.get('严重度', '')
            if sev == 'P0': cell.fill = P0_FILL
            elif sev == 'P1': cell.fill = P1_FILL
        # Append remark if exists
        if iss.get('备注原文'):
            ws_summary.cell(row=row, column=12).value += f'\n[备注原文] {iss["备注原文"][:100]}'
        row += 1
    
    # Column widths
    widths = [5, 10, 28, 30, 8, 12, 6, 8, 8, 55, 6, 40]
    for c, w in enumerate(widths, 1):
        ws_summary.column_dimensions[get_column_letter(c)].width = w
    ws_summary.freeze_panes = 'A2'
    ws_summary.auto_filter.ref = f'A1:{get_column_letter(len(SUMMARY_HEADERS))}{row-1}'
    
    # --- 逐项目Sheet ---
    for pname, pissues in project_results.items():
        safe_pname = re.sub(r'[\\/*?:\"<>|\[\]]', '_', pname)[:28]
        ws = out_wb.create_sheet(safe_pname)
        
        # Project headers
        proj_headers = ['序号','标准','指标','指标性质','指标值','权重','自评得分','复核得分','问题描述','严重度','建议','备注原文']
        for c, h in enumerate(proj_headers, 1):
            cell = ws.cell(row=1, column=c, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = HEADER_ALIGN
            cell.border = THIN_BORDER
        
        # Sort issues
        p0s = [x for x in pissues if x['严重度'] == 'P0']
        p1s = [x for x in pissues if x['严重度'] == 'P1']
        p2s = [x for x in pissues if x['严重度'] == 'P2']
        sorted_p = p0s + p1s + p2s
        
        for idx, iss in enumerate(sorted_p, 1):
            vals = [
                idx,
                iss.get('标准', ''),
                iss.get('指标', ''),
                iss.get('指标性质', ''),
                iss.get('指标值', ''),
                iss.get('权重', ''),
                iss.get('自评得分', ''),
                iss.get('复核得分', ''),
                iss.get('问题', ''),
                iss.get('严重度', ''),
                iss.get('建议', ''),
                iss.get('备注原文', '')[:120],
            ]
            for c, v in enumerate(vals, 1):
                cell = ws.cell(row=idx+1, column=c, value=str(v))
                cell.font = BODY_FONT
                cell.alignment = BODY_ALIGN
                cell.border = THIN_BORDER
                sev = iss.get('严重度', '')
                if sev == 'P0': cell.fill = P0_FILL
                elif sev == 'P1': cell.fill = P1_FILL
        
        # Column widths
        pw2 = [5, 10, 30, 8, 12, 6, 8, 8, 55, 6, 40, 40]
        for c, w in enumerate(pw2, 1):
            ws.column_dimensions[get_column_letter(c)].width = w
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:{get_column_letter(len(proj_headers))}{max(2, len(sorted_p)+1)}'
    
    # --- 格式检查Sheet ---
    ws_fmt = out_wb.create_sheet('格式检查详情')
    fmt_headers = ['序号','位置','问题描述','严重度','建议']
    for c, h in enumerate(fmt_headers, 1):
        cell = ws_fmt.cell(row=1, column=c, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = HEADER_ALIGN
        cell.border = THIN_BORDER
    
    for idx, iss in enumerate(s13, 1):
        vals = [idx, iss.get('位置', ''), iss.get('问题', ''), iss.get('严重度', ''), 
                '按格式规范修正' if iss['严重度'] in ('P0','P1') else '可优化']
        for c, v in enumerate(vals, 1):
            cell = ws_fmt.cell(row=idx+1, column=c, value=str(v))
            cell.font = BODY_FONT
            cell.alignment = BODY_ALIGN
            cell.border = THIN_BORDER
            sev = iss.get('严重度', '')
            if sev == 'P0': cell.fill = P0_FILL
            elif sev == 'P1': cell.fill = P1_FILL
    
    fmt_w = [5, 20, 60, 6, 30]
    for c, w in enumerate(fmt_w, 1):
        ws_fmt.column_dimensions[get_column_letter(c)].width = w
    ws_fmt.freeze_panes = 'A2'
    
    # Move 格式检查 to last position
    out_wb.move_sheet('格式检查详情', offset=len(out_wb.sheetnames)-1)
    
    out_wb.save(str(out_path))
    print(f'  ✅ 已保存: {out_path.name}')
    
    return {
        'unit': unit_name,
        'file': out_path.name,
        'projects': len(project_results),
        'indicators': sum(len(read_indicator_rows(ws, parse_header_row(ws)[0])) 
                         for sn, ws in project_sheets if parse_header_row(ws)[0]),
        'P0': p0_count,
        'P1': p1_count,
        'P2': p2_count,
        'total': len(all_issues),
    }


# ============================================================
# MAIN
# ============================================================
def main():
    print('绩效自评复核 V2')
    print('=' * 60)
    
    summaries = []
    for unit_dir in sorted(BASE.iterdir()):
        if unit_dir.is_dir():
            result = process_unit(unit_dir)
            if result:
                summaries.append(result)
    
    # 清理旧的汇总文件
    old = DESKTOP / '绩效自评复核结果.xlsx'
    if old.exists():
        old.unlink()
    
    print(f'\n{"="*60}')
    print('📊 全部完成！')
    print(f'{"单位":35s} {"项目":>4s} {"指标":>4s} {"P0":>4s} {"P1":>4s} {"P2":>4s} {"文件"}')
    print('-' * 90)
    for s in summaries:
        print(f'{s["unit"][:35]:35s} {s["projects"]:>4d} {s["indicators"]:>4d} {s["P0"]:>4d} {s["P1"]:>4d} {s["P2"]:>4d} {s["file"]}')
    print('-' * 90)
    total_p0 = sum(s['P0'] for s in summaries)
    total_p1 = sum(s['P1'] for s in summaries)
    total_p2 = sum(s['P2'] for s in summaries)
    total_all = sum(s['total'] for s in summaries)
    print(f'{"合计":35s} {sum(s["projects"] for s in summaries):>4d} {sum(s["indicators"] for s in summaries):>4d} {total_p0:>4d} {total_p1:>4d} {total_p2:>4d}')
    print(f'\n总计: {total_all}个问题 (P0:{total_p0} P1:{total_p1} P2:{total_p2})')
    print(f'输出目录: {DESKTOP}')

if __name__ == '__main__':
    main()
