# -*- coding: utf-8 -*-
"""融策制度体系 → 正式Word文档生成器"""
import os, re, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\scrccpa\.openclaw\workspace\output\新制度体系'
DESKTOP = r'C:\Users\scrccpa\Desktop'

# ===== FIXES =====
FIXES = {
    '02-绩效考核管理制度.md': [
        # D等强制分布 → 指导性
        (r'D等比例原则上控制在.*?无D等。', 'D等为"不称职"等级，原则上作为绩效改进的参考指标，不设硬性比例强制分布。如部门全体员工绩效考核得分均≥60分（合格线），该季度/年度可无D等。任何D等评定必须基于员工实际工作表现是否未达到岗位基本要求。'),
        # 补充产假保护
        ('考核结果申诉', '特殊情形考核\n\n第二十一条 处于法定产假、哺乳假、工伤停工留薪期、医疗期内的员工，考核按以下方式处理：（1）整个考核周期均在上述期间的，不参与当期考核排名，绩效工资按基数100%发放；（2）部分在保护期的，按实际在岗时间可量化的工作指标考核，缺少数据的维度不扣分。\n\n考核结果申诉'),
        # 补充民主程序
        ('本制度经管理层审议通过后施行', '本制度经职工代表大会（或全体员工大会）讨论通过，听取工会或职工代表意见后，由管理层审议发布施行'),
        # 取消360度评价
        (r'360度协作评分.*?团队协作', '团队协作评分由项目经理和部门负责人根据日常观察直接评分，取消360度互评机制'),
        # 季度→半年度
        (r'季度考核', '半年度考核'),
        (r'每季度', '每半年'),
        # 流程压缩
        (r'20个工作日', '5个工作日'),
        # 连续D等完整程序
        ('改进期满仍为D的.*?解除劳动合同', '改进期满仍为D的，公司可依据《劳动合同法》第40条第（二）项规定：（1）对员工进行专业技术培训或合理调岗；（2）培训或调岗后经考核仍不能胜任工作的，提前30日以书面形式通知员工或额外支付一个月工资后，解除劳动合同，并依法支付经济补偿'),
    ],
    '03-员工手册.md': [
        # 试用期延长 → 删除
        (r'考核不合格的.*?不得延长试用期。', '考核不合格的，公司应在试用期内依据《劳动合同法》第39条第（一）项规定，以书面形式通知员工解除劳动合同，并列明不符合录用条件的具体事实和证据。'),
        # 晚婚假 → 删除
        (r'晚婚.*?增加\d+天.*?。', '法定婚假3天，婚假期间工资全额发放。'),
        # 监控权限
        ('因业务需要.*?合理管理', '公司因下列特定目的可对工作邮箱和工作电脑中的工作文件进行管理：（1）离职交接的数据保全；（2）合规审计或内部调查；（3）应司法机关要求的数据调取；（4）信息安全事件应急响应。管理范围限于与工作直接相关的文件'),
        # 民主程序
        ('经民主程序通过', '经全体员工大会讨论通过，与职工代表平等协商确定'),
    ],
    '06-财务报销管理制度.md': [
        # 360天认证 → 删除
        (r'专票认证.*?360.*?天.*?。', '增值税专用发票应通过增值税发票综合服务平台进行用途确认（勾选认证），取消认证期限。'),
        # 费用标准调整权
        ('费用标准由管理层.*?调整', '费用标准调整由财务部门根据物价水平和公司经营情况提出方案，经职工代表大会或全体员工讨论后，由管理层审批发布'),
    ],
    '07-审计质量控制制度.md': [
        # 绩效100%扣 → 月≤20%
        (r'扣减当年全部绩效工资', '扣减绩效工资（按月执行，每月扣减不超过当月工资的20%，且扣除后不低于最低工资标准）'),
        (r'扣减.*?50%.*?绩效', '扣减绩效工资（按月执行，每月扣减不超过当月工资的20%）'),
        # 向监管部门报告
        ('需要向监管部门报告的.*?决定', '依法属于强制报告事项的，质控负责人应在发现后5个工作日内向财政部门和注册会计师协会报告'),
    ],
    '08-造价咨询质量控制制度.md': [
        # 同07
        (r'扣减当年全部绩效工资', '扣减绩效工资（按月执行，每月扣减不超过当月工资的20%，且扣除后不低于最低工资标准）'),
    ],
    '10-招聘与入职管理制度.md': [
        # 试用期延长
        (r'基本合格.*?延长试用期.*?法定上限', '试用期考核不合格的，在试用期届满前依据《劳动合同法》第39条第（一）项解除劳动合同'),
    ],
    '13-项目收入确认与回款管理制度.md': [
        # 坏账奖金追回
        (r'已发放的予以追回', '已发放的项目奖金不予追回。未发放的不再发放'),
        # 坏账追责
        ('因项目经理未及时跟踪.*?扣减相应', '项目经理未按预警机制履行催收职责的，依执业责任追究制度处理'),
    ],
    '22-执业责任追究制度.md': [
        # 绩效100% → 月≤20%
        (r'扣减全年绩效奖金', '扣减绩效奖金（按月执行，每月扣减不超过当月工资的20%，且扣除后不低于最低工资标准）'),
        # 调岗降级 → 协商一致
        ('降级或调岗', '经与员工协商一致后降级或调岗'),
        # 拒绝违法指令
        ('受上级指令违规', '员工有权拒绝明显违法的指令。受上级指令违规'),
    ],
    '12-职级晋升管理制度.md': [
        # 降级 → 协商
        ('有下列情形之一的，.*?降级', '有下列情形之一的，公司可与员工协商降级，经双方书面确认后执行'),
    ],
    '14-预算管理制度.md': [
        # 简化
        (r'11月.*?1月初下达', '12月中旬各部门提交预算方案 → 12月下旬管理层审议 → 1月初下达执行'),
        (r'增量预算法.*?弹性预算法', '采用增量预算法（以上年实际为基数，结合业务变化调整），新增业务可采用零基预算'),
        (r'月度跟踪\+季度分析\+半年度总结\+年度总结', '月度简要跟踪+半年度书面报告'),
    ],
    '04-项目管理规范.md': [
        # C类简化
        (r'所有项目均须', 'A类和B类项目须执行'),
    ],
}

def apply_fixes(filename, content):
    if filename not in FIXES:
        return content
    for old, new in FIXES[filename]:
        content = re.sub(old, new, content, flags=re.DOTALL)
    return content

def read_md(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_heading_safe(doc, text, level):
    """Add heading with proper Chinese font"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def set_paragraph_font(para, font_name='仿宋_GB2312', size=Pt(14)):
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size

def parse_md_to_docx(doc, md_content, title):
    """Parse markdown content to docx with proper formatting"""
    lines = md_content.split('\n')
    in_table = False
    in_code = False
    current_para = doc.add_paragraph()
    
    for line in lines:
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            if not in_table and not in_code:
                current_para = doc.add_paragraph()
            continue
        
        # Headers
        if stripped.startswith('# ') and not stripped.startswith('## '):
            h = doc.add_heading(stripped[2:], level=1)
            for run in h.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            continue
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=2)
            for run in h.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            continue
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            for run in h.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            continue
        
        # Tables (simple detection)
        if '|' in stripped and stripped.count('|') >= 2:
            if not in_table:
                in_table = True
                # Skip separator line
                if '---' in stripped or ':-' in stripped:
                    continue
                # Start table with header
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                for i, cell_text in enumerate(cells):
                    cell = table.rows[0].cells[i]
                    cell.text = cell_text
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                            run.font.name = '黑体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                continue
            
            # Table data row
            if '---' in stripped or ':-' in stripped:
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            row = table.add_row()
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = cell_text
                    for p in row.cells[i].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.name = '仿宋'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            continue
        else:
            if in_table:
                in_table = False
                doc.add_paragraph()
        
        # Code blocks
        if stripped.startswith('```'):
            in_code = not in_code
            continue
        
        if in_code:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue
        
        # Normal paragraph
        # Bold markers
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
        clean = re.sub(r'__(.+?)__', r'\1', clean)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'_(.+?)_', r'\1', clean)
        
        # List items
        if re.match(r'^[\d]+[\.\)、]', clean) or stripped.startswith('- ') or stripped.startswith('* '):
            clean = re.sub(r'^[\-\*\d]+[\.\)、\s]*', '', clean, count=1)
        
        p = doc.add_paragraph(clean)
        set_paragraph_font(p)
    
    return doc

def add_page_number(doc):
    """Add page numbers to footer"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        
        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

def setup_section(doc):
    """Setup page layout"""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run('四川融策会计师/工程咨询有限公司')
    hr.font.name = '宋体'
    hr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(128, 128, 128)

def add_cover(doc):
    """Add cover page"""
    # Empty lines for spacing
    for _ in range(6):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0A, 0x1F, 0x3F)
    run.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('制度体系文件')
    run2.font.name = '黑体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run2.font.size = Pt(36)
    run2.font.color.rgb = RGBColor(0x0A, 0x1F, 0x3F)
    run2.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('【合订本】')
    run3.font.name = '黑体'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run3.font.size = Pt(18)
    
    doc.add_paragraph()
    
    info_lines = [
        '版本号：V1.0',
        '编制日期：2026年7月22日',
        '审核人：__________    批准人：__________',
        '',
        '四川融策会计师事务所有限公司',
        '四川融策工程咨询有限公司',
        '',
        '成都 · 2026年7月'
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(14)
    
    # Page break
    doc.add_page_break()

def generate_word():
    print('开始生成Word文档...')
    
    files = sorted([f for f in os.listdir(SRC) if f.endswith('.md')])
    print(f'共{len(files)}个制度文件')
    
    doc = Document()
    setup_section(doc)
    add_cover(doc)
    
    # Title page for TOC
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目    录')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True
    doc.add_paragraph()
    
    # Table of Contents
    categories = {
        '架构与总纲': ['00', '05'],
        '公司治理': ['26', '27', '09'],
        '人力资源': ['01', '02', '03', '10', '11', '12'],
        '财务管理': ['06', '13', '14', '15', '16'],
        '业务运营': ['04', '17', '18', '19', '20'],
        '质量控制': ['07', '08', '21', '22'],
        '行政支持': ['23', '24', '25', '28', '29'],
        '专项制度': ['30', '31', '32', '33'],
    }
    
    for cat, prefixes in categories.items():
        p = doc.add_paragraph()
        run = p.add_run(f'\n{cat}')
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.bold = True
        
        for f in files:
            for pre in prefixes:
                if f.startswith(pre):
                    title = f.replace('.md', '').replace(pre+'-', '', 1)
                    p2 = doc.add_paragraph(f'    {title}')
                    for r in p2.runs:
                        r.font.name = '仿宋'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                        r.font.size = Pt(12)
                    break
    
    doc.add_page_break()
    
    # Content pages
    for i, f in enumerate(files):
        print(f'  [{i+1}/{len(files)}] {f}')
        filepath = os.path.join(SRC, f)
        content = read_md(filepath)
        content = apply_fixes(f, content)
        
        # Extract title
        title = f.replace('.md', '')
        # Remove prefix like "00-" or "01-"
        if title[:2].isdigit() and title[2:3] == '-':
            title = title[3:]
        
        # Add title page for this document
        doc.add_page_break()
        
        # Parse content
        parse_md_to_docx(doc, content, title)
    
    # Page numbers
    add_page_number(doc)
    
    # Save
    output_path = os.path.join(DESKTOP, '融策公司制度体系（完整版）.docx')
    doc.save(output_path)
    print(f'\n✅ 完整版已保存: {output_path}')
    return output_path

if __name__ == '__main__':
    try:
        result = generate_word()
        print(f'文件大小: {os.path.getsize(result):,} bytes')
    except Exception as e:
        print(f'❌ 错误: {e}')
        import traceback
        traceback.print_exc()
