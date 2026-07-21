# -*- coding: utf-8 -*-
"""
融策·券商风整报告 Word 生成器
============================
生成正式 Word 研报版本，配套嵌入高清图表 PNG。

用法：
    python -X utf8 scripts/generate_securities_report_word.py output/融策_券商风整报告_Word版.docx
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
CHART_DIR = ROOT / 'output' / 'charts_pro'
BRAND_DIR = Path(r'E:\2026\宣传册\融策logo')
LOGO_PATH = BRAND_DIR / '融策logo.png'
LETTERHEAD_PATH = BRAND_DIR / '融策抬头.jpg'

NAVY = RGBColor(6, 26, 51)
NAVY2 = RGBColor(10, 42, 74)
TEAL = RGBColor(26, 111, 120)
GOLD = RGBColor(184, 138, 68)
GOLD2 = RGBColor(214, 176, 113)
INK = RGBColor(31, 41, 51)
GRAY = RGBColor(83, 97, 109)
MUTED = RGBColor(135, 147, 160)
PANEL = RGBColor(244, 246, 248)
WHITE = RGBColor(255, 255, 255)
WARM = RGBColor(245, 242, 236)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color='D9DEE5', size='6'):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    p_pr.append(shd)


def add_run(paragraph, text, size=10.5, color=INK, bold=False):
    run = paragraph.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    return run


def add_para(doc, text='', style=None, size=10.5, color=INK, bold=False, align=None, space_after=6):
    # 兼容 add_para(doc, text, size, color) 这种简写，避免把字号误当样式名。
    if isinstance(style, (int, float)):
        actual_size = style
        actual_color = size if isinstance(size, RGBColor) else color
        style = None
        size = actual_size
        color = actual_color
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, size, color, bold)
    return p


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles['Normal'].font.name = 'Microsoft YaHei'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    styles['Normal'].font.size = Pt(10.5)
    styles['Normal'].font.color.rgb = INK

    h1 = styles['Heading 1']
    h1.font.name = 'Microsoft YaHei'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = WHITE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    shade_paragraph_style(h1, '061A33')

    h2 = styles['Heading 2']
    h2.font.name = 'Microsoft YaHei'
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = TEAL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name = 'Microsoft YaHei'
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = INK

    chart_title = styles.add_style('RC_ChartTitle', WD_STYLE_TYPE.PARAGRAPH)
    chart_title.font.name = 'Microsoft YaHei'
    chart_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    chart_title.font.size = Pt(10)
    chart_title.font.bold = True
    chart_title.font.color.rgb = NAVY
    chart_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_title.paragraph_format.space_before = Pt(10)
    chart_title.paragraph_format.space_after = Pt(4)

    source = styles.add_style('RC_Source', WD_STYLE_TYPE.PARAGRAPH)
    source.font.name = 'Microsoft YaHei'
    source._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    source.font.size = Pt(8)
    source.font.italic = True
    source.font.color.rgb = MUTED
    source.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.space_after = Pt(14)


def shade_paragraph_style(style, fill):
    p_pr = style._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    p_pr.append(shd)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if LOGO_PATH.exists():
            run = hp.add_run()
            run.add_picture(str(LOGO_PATH), width=Cm(0.55))
            add_run(hp, '  ', 8, MUTED)
        add_run(hp, '融策·审盾研究 | 券商风研报母版', 8, MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(fp, '四川融策会计师事务所 / 四川融策工程咨询公司', 8, MUTED)


def add_cover(doc):
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Cm(2.0))
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(14.0)
    left, right = table.rows[0].cells
    set_cell_shading(left, '061A33')
    set_cell_border(left, '061A33')
    set_cell_border(right, 'FFFFFF')
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '融\n策\n审\n盾', 18, GOLD2, True)

    add_para(doc, '行业深度报告 | AI审计中台系列', size=11, color=GOLD, bold=True, space_after=28)
    add_para(doc, '融策AI审计中台：\n重塑政府审计的数据穿透力', size=28, color=NAVY, bold=True, space_after=14)
    add_para(doc, '——从信息化核查到智能交叉验证的0到1跃迁', size=16, color=TEAL, space_after=28)

    p = add_para(doc, '本报告复刻券商深度研报的结构逻辑，以“观点先行、图表作证、判断落地”为核心，形成适用于绩效评价、经责审计、专项资金、工程决算等业务线的标准化研报 Word 版母版。', size=11.5, color=INK, space_after=22)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(1.2)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.RIGHT
    rows = [('分析师', '融策右护卫'), ('报告版本', 'Word母版 v1.0'), ('报告日期', '2026年7月21日'), ('适用场景', '绩效评价 / 经责审计 / 工程决算 / 专项资金')]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = meta.cell(i, j)
            cell.text = ''
            set_cell_shading(cell, 'F4F6F8' if j == 0 else 'FFFFFF')
            set_cell_border(cell)
            p = cell.paragraphs[0]
            add_run(p, val, 9, GRAY if j == 0 else INK, bold=(j == 0))
    doc.add_page_break()


def add_core_summary(doc):
    doc.add_heading('核心观点', level=1)
    points = [
        ('不是做“漂亮图表”，而是做证据链页面', '每页必须有一个明确判断，并用KPI、图表、右侧解释和来源共同支撑。'),
        ('研报模板应服务融策业务线，而非照搬券商金融话术', '绩效评价、经责审计、工程决算、专项资金各自需要不同的数据口径与图表母版。'),
        ('Word用于正式归档，PPTX用于客户汇报和内部打磨', '正式交付需要正文逻辑、图表、表格、依据、免责声明完整闭环。'),
    ]
    for idx, (title, body) in enumerate(points, 1):
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        t.columns[0].width = Cm(1.2)
        t.columns[1].width = Cm(14.2)
        set_cell_shading(t.cell(0, 0), '061A33')
        set_cell_shading(t.cell(0, 1), 'F4F6F8')
        set_cell_border(t.cell(0, 0), '061A33')
        set_cell_border(t.cell(0, 1), 'D9DEE5')
        p0 = t.cell(0, 0).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p0, f'{idx:02d}', 16, GOLD2, True)
        p1 = t.cell(0, 1).paragraphs[0]
        add_run(p1, title + '\n', 11.5, NAVY, True)
        add_run(p1, body, 10, INK)
        add_para(doc, '', space_after=4)
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading('目录', level=1)
    items = [
        ('一、行业痛点：传统审计报告为何不够像研报', '4'),
        ('二、图文分析：财政收入修复斜率与税收弹性', '5'),
        ('三、数据表格：业务线模板化程度与AI适配度', '6'),
        ('四、方法论：观点-指标-图表-解释-行动', '7'),
        ('五、结论与行动建议', '8'),
        ('免责声明', '9'),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        add_run(p, title, 11, INK)
        add_run(p, ' ' * 20 + page, 11, MUTED)
    add_para(doc, '注：正式出具前可使用 Word“引用-目录”自动更新目录。', 8.5, MUTED)
    doc.add_page_break()


def add_kpi_row(doc, kpis):
    table = doc.add_table(rows=1, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, k in enumerate(kpis):
        cell = table.cell(0, i)
        set_cell_shading(cell, 'F4F6F8')
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        add_run(p, k[0] + '\n', 8.5, GRAY)
        add_run(p, k[1], 15, k[2], True)
    add_para(doc, '', space_after=4)


def add_chart(doc, filename, title, source):
    path = CHART_DIR / filename
    doc.add_paragraph(title, style='RC_ChartTitle')
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(16.2))
    else:
        add_para(doc, f'[图表文件缺失：{path}]', 9, MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph(f'资料来源：{source}', style='RC_Source')


def add_body_pages(doc):
    if LETTERHEAD_PATH.exists():
        p_head = doc.add_paragraph()
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head.add_run().add_picture(str(LETTERHEAD_PATH), width=Cm(16.2))
        p_head.paragraph_format.space_after = Pt(10)
    doc.add_heading('一、行业痛点：传统审计报告为何不够像研报', level=1)
    add_para(doc, '传统审计报告的问题并非没有内容，而是信息组织方式过于“底稿化”：先堆事实，再列问题，最后给建议。券商研报的强项在于先给判断，再组织证据，让读者在第一页就知道结论，在正文中逐步建立信任。', 10.5)
    add_para(doc, '对融策而言，研报化不是包装，而是把复杂审计发现转译成客户能快速理解和决策的表达方式。', 10.5, NAVY, True)
    doc.add_heading('1.1 报告页的标准闭环', level=2)
    add_para(doc, '每一页正文应形成“观点-指标-图表-解释-行动”的闭环：观点回答结论，指标压住重点，图表提供证据，解释说明风险，行动落到管理建议。', 10.5)

    doc.add_heading('二、图文分析：财政收入修复斜率与税收弹性', level=1)
    add_kpi_row(doc, [('2025E预算收入', '22.8万亿', NAVY), ('2024-2025E增量', '+0.8万亿', GOLD), ('税收占比', '83.3%', TEAL)])
    add_chart(doc, '01_pro_line_fiscal.png', '图1：全国一般公共预算收入与税收收入走势（亿元）', '财政部，Wind，融策会计师事务所整理')
    doc.add_heading('2.1 核心判断', level=2)
    for text in [
        '财政收入修复并非线性扩张，税基质量和房地产链条回暖仍决定后续弹性。',
        '税收收入占比维持高位，说明非税收入拉动空间有限，应重点关注税源真实性。',
        '若预算收入增速明显高于经济增速，应回查一次性收入、非税缴库和跨期调节。',
    ]:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        add_run(p, '• ', 11, GOLD, True)
        add_run(p, text, 10.5, INK)

    doc.add_heading('三、数据表格：业务线模板化程度与AI适配度', level=1)
    table = doc.add_table(rows=6, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['业务线', '数据结构化', '报告标准化', 'AI适配度', '优先级']
    rows = [
        ['绩效评价', '中', '高', '高', 'P0'],
        ['工程决算', '高', '中', '高', 'P0'],
        ['经责审计', '中', '中', '中高', 'P1'],
        ['专项资金', '中高', '中', '高', 'P1'],
        ['资产清查', '高', '高', '中', 'P2'],
    ]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, '061A33')
        set_cell_border(cell, '061A33')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, 9, WHITE, True)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            set_cell_shading(cell, 'F4F6F8' if i % 2 == 0 else 'FFFFFF')
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, val, 9, INK)
    doc.add_paragraph('资料来源：融策项目台账，内部复盘', style='RC_Source')

    add_chart(doc, '02_pro_bar_projects.png', '图2：2025年度审计咨询项目类型分布（个）', '融策项目台账，行业访谈，融策AI审计中台')

    doc.add_heading('四、方法论：观点-指标-图表-解释-行动', level=1)
    method = doc.add_table(rows=1, cols=5)
    method.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (title, desc) in enumerate([('观点', '一句话判断'), ('指标', '三个关键数字'), ('图表', '一张主证据图'), ('解释', '三条核心判断'), ('行动', '下一步建议')]):
        cell = method.cell(0, i)
        set_cell_shading(cell, 'F4F6F8')
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title + '\n', 13, NAVY, True)
        add_run(p, desc, 8.5, GRAY)

    doc.add_heading('五、结论与行动建议', level=1)
    for title, body in [
        ('短期动作', '先选绩效评价报告作为样板，固定4类图表：趋势、结构、对比、风险矩阵。'),
        ('中期建设', '建立业务线模板库，形成数据口径、图表样式、判断句库三件套。'),
        ('长期价值', '形成融策自己的“研究型审计报告”品牌，让客户感知从审计服务升级到决策支持。'),
    ]:
        p = add_para(doc, title, 12, NAVY, True, space_after=2)
        shade_paragraph(p, 'F5F2EC')
        add_para(doc, body, 10.5, INK)


def add_disclaimer(doc):
    doc.add_page_break()
    doc.add_heading('免责声明', level=1)
    add_para(doc, '本报告为四川融策会计师事务所及四川融策工程咨询公司内部研报模板示例，所列数据与案例仅用于展示图文协同版式，不构成正式审计结论、投资建议或对外承诺。正式项目报告应以经复核的底稿、取证材料、法规依据和签批流程为准。', 9, GRAY)
    doc.add_heading('使用要求', level=2)
    for text in ['图表数据必须注明来源。', '重大金额、比率和结论必须与底稿交叉核验。', 'AI生成内容必须经项目经理和质控复核后方可用于正式报告。', '对外版本应删除内部方法、模型路由和未公开数据来源。']:
        p = doc.add_paragraph()
        add_run(p, '• ', 10, GOLD, True)
        add_run(p, text, 9.5, GRAY)


def build(output):
    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    add_header_footer(doc)
    add_core_summary(doc)
    add_toc(doc)
    add_body_pages(doc)
    add_disclaimer(doc)
    Path(output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f'✅ Word版整报告已生成: {output}')


if __name__ == '__main__':
    output = sys.argv[1] if len(sys.argv) > 1 else 'output/融策_券商风整报告_Word版.docx'
    build(output)
