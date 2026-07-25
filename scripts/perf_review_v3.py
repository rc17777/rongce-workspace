# -*- coding: utf-8 -*-
"""
绩效自评复核 V3 — 简洁版
输出结构：①总览 ②11-14报告层面 ③1-10指标层面 ④格式清单
"""
import sys, os, re
from pathlib import Path
from collections import defaultdict, Counter
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

BASE = Path(r'C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722')
DESKTOP = Path(r'C:\Users\scrccpa\Desktop')

# ---------- styles ----------
HDR_FONT = Font(name='黑体', size=11, bold=True, color='FFFFFF')
HDR_FILL = PatternFill(start_color='0A1F3F', end_color='0A1F3F', fill_type='solid')
TITLE_FONT = Font(name='黑体', size=14, bold=True)
SEC_FONT = Font(name='黑体', size=12, bold=True, color='0A1F3F')
BODY_FONT = Font(name='仿宋', size=10)
BOLD_FONT = Font(name='仿宋', size=10, bold=True)
BORDER = Border(left=Side('thin'), right=Side('thin'), top=Side('thin'), bottom=Side('thin'))
ALIGN_C = Alignment(horizontal='center', vertical='center', wrap_text=True)
ALIGN_W = Alignment(vertical='center', wrap_text=True)
FILL_P0 = PatternFill(start_color='FF6B6B', end_color='FF6B6B', fill_type='solid')
FILL_P1 = PatternFill(start_color='FFD93D', end_color='FFD93D', fill_type='solid')
FILL_P2 = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
FILL_OK = PatternFill(start_color='6BCB77', end_color='6BCB77', fill_type='solid')
FILL_LIGHT = PatternFill(start_color='F0F4F8', end_color='F0F4F8', fill_type='solid')

def sfv(v): return str(v).strip() if v is not None else ''
def sfn(v, d=None):
    if v is None: return d
    try: return float(v)
    except: return d

# ---------- docx helpers ----------
def read_docx_tables(docx_path):
    """提取docx中关键表格数据"""
    doc = Document(docx_path)
    tables_data = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            cells = [sfv(c.text) for c in row.cells]
            rows.append(cells)
        tables_data.append(rows)
    return tables_data, doc

def check_format_v3(doc):
    """格式检查：只报告明确不符合要求的问题"""
    issues = []
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt: continue
        
        # 一级标题 一、二、...  → 黑体
        if re.match(r'^[一二三四五六七八九十]、', txt) and len(txt) < 60:
            for r in p.runs:
                if r.text.strip():
                    fn = (r.font.name or '')
                    if fn and '黑体' not in fn:
                        issues.append({'类型':'标题字体','位置':f'正文', '详情':f'一级标题"{txt[:20]}"非黑体(={fn})','严重度':'P1'})
                        break
        
        # 二级标题 （一）（二）... → 楷体加粗
        if re.match(r'^（[一二三四五六七八九十\d]+）', txt) and len(txt) < 60:
            for r in p.runs:
                if r.text.strip():
                    fn = (r.font.name or '')
                    if fn and '楷体' not in fn:
                        issues.append({'类型':'标题字体','位置':f'正文', '详情':f'二级标题"{txt[:20]}"非楷体(={fn})','严重度':'P1'})
                        break
    
    # 页面设置
    for si, sec in enumerate(doc.sections):
        if sec.page_width and abs(sec.page_width - 1190700) > 50000:
            issues.append({'类型':'页面','位置':f'第{si+1}节','详情':f'页宽{sec.page_width/35000:.1f}cm≠A4(21cm)','严重度':'P1'})
        if sec.page_height and abs(sec.page_height - 1683800) > 50000:
            issues.append({'类型':'页面','位置':f'第{si+1}节','详情':f'页高{sec.page_height/35000:.1f}cm≠A4(29.7cm)','严重度':'P1'})
    
    return issues

# ---------- xlsx helpers ----------
def is_project_sheet(ws, sn):
    if ws.max_row < 5: return False
    if sn in ('汇总','自评复核报告表','目标完成，偏离度','差值表'): return False
    for ri in range(1, min(8, ws.max_row+1)):
        if '三级指标' in sfv(ws.cell(row=ri, column=3).value) and '复核得分' in sfv(ws.cell(row=ri, column=11).value):
            return True
    return False

def read_xlsx_projects(xlsx_path):
    """读取xlsx所有项目sheet，返回 {sheet_name: {project_name, indicators[]}}"""
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    projects = {}
    for sn in wb.sheetnames:
        ws = wb[sn]
        if not is_project_sheet(ws, sn): continue
        
        # 项目名
        pname = ''
        for ri in range(1, min(5, ws.max_row+1)):
            v = sfv(ws.cell(row=ri, column=1).value)
            if v and len(v) > 8 and not v.startswith('附件'):
                pname = re.sub(r'^\d+[A-Z]?\d*-?', '', v)
                break
        if not pname: pname = sn
        
        # 表头定位
        hrow = None
        for ri in range(1, min(8, ws.max_row+1)):
            if '三级指标' in sfv(ws.cell(row=ri, column=3).value):
                hrow = ri; break
        if hrow is None: continue
        
        # 指标数据
        indicators = []
        for ri in range(hrow+1, ws.max_row+1):
            ind = sfv(ws.cell(row=ri, column=3).value)
            lv1 = sfv(ws.cell(row=ri, column=1).value)
            if not ind or ind == 'None' or '合计' in ind or '合计' in lv1: continue
            
            lev2 = sfv(ws.cell(row=ri, column=2).value)
            nat  = sfv(ws.cell(row=ri, column=4).value)
            tval = ws.cell(row=ri, column=5).value
            wgt  = sfn(ws.cell(row=ri, column=7).value, 0)
            sv   = ws.cell(row=ri, column=8).value
            ss   = sfn(ws.cell(row=ri, column=9).value, 0)
            rv   = ws.cell(row=ri, column=10).value
            rs   = sfn(ws.cell(row=ri, column=11).value, -1)
            rem  = sfv(ws.cell(row=ri, column=13).value)
            
            indicators.append({
                '一级指标': lv1, '二级指标': lev2, '三级指标': ind,
                '指标性质': nat, '指标值': tval, '权重': wgt,
                '自评完成值': sv, '自评得分': ss,
                '复核完成值': rv, '复核得分': rs, '备注': rem,
            })
        
        projects[pname] = indicators
    
    return projects

# ---------- 14条标准检查 ----------
def check_indicators(indicators, pname):
    """逐指标检查标准1-10和标准14，返回issues列表"""
    issues = []
    
    for r in indicators:
        ind = r['三级指标']; lev1 = r['一级指标']; lev2 = r['二级指标']
        nat = r['指标性质']; wgt = r['权重']; ss = r['自评得分']
        rs = r['复核得分']; rem = r['备注']
        
        is_deducted = (rs >= 0 and rs < ss)
        is_zero = (rs == 0 and ss > 0)
        is_blank = (rs < 0)
        deducted_pct = (ss - rs) / max(ss, 1) * 100 if is_deducted else 0
        
        # --- 标准1: 没资料=0分 ---
        if is_zero and rem:
            has_no_data = any(kw in rem for kw in ['无资料','无佐证','未提供佐证','未提供资料','无法提供','无相关'])
            work_not_done = any(kw in rem for kw in ['无剩余资金','无资金','未实施','不再实施','调整未','已取消','已调整','未再组织'])
            if not has_no_data and not work_not_done:
                issues.append({'标准':'1','指标':ind,'问题':f'打0分但备注未说明"无佐证资料"','严重度':'P2','建议':'统一备注口径','备注原文':rem[:60]})
        
        # --- 标准2: 无关=0分 ---
        if ('无关' in rem or '不相关' in rem) and rs > 0:
            issues.append({'标准':'2','指标':ind,'问题':f'指标与项目无关但未打0分(得分={rs})','严重度':'P0','建议':'按口径第2条打0分','备注原文':rem[:60]})
        
        # --- 标准3: 可考查性 ---
        unverifiable = any(kw in rem for kw in ['不具考核','不具可考','无可考查','空洞'])
        if unverifiable and is_deducted:
            # 检查扣分比例与标注是否一致
            if '扣100%' in rem or '100%' in rem:
                if deducted_pct < 90:
                    issues.append({'标准':'3','指标':ind,'问题':f'标注扣100%但实际只扣{deducted_pct:.0f}%','严重度':'P1','建议':'确认扣分比例','备注原文':rem[:60]})
            elif '扣60%' in rem:
                if abs(deducted_pct - 60) > 15:
                    issues.append({'标准':'3','指标':ind,'问题':f'标注扣60%但实际扣{deducted_pct:.0f}%','严重度':'P1','建议':'确认扣分比例','备注原文':rem[:60]})
            elif '扣30%' in rem:
                if abs(deducted_pct - 30) > 15:
                    issues.append({'标准':'3','指标':ind,'问题':f'标注扣30%但实际扣{deducted_pct:.0f}%','严重度':'P2','建议':'确认扣分比例','备注原文':rem[:60]})
        
        # --- 标准4: 满意度无资料=0分 ---
        if ('满意' in ind or '满意' in lev2) and is_deducted:
            if any(kw in rem for kw in ['未调查','未问卷','无资料','无调查','未进行']):
                if rs > 0:
                    issues.append({'标准':'4','指标':ind,'问题':f'满意度无资料但复核未打0分(得分={rs})','严重度':'P0','建议':'按口径第4条打0分','备注原文':rem[:60]})
        
        # --- 标准5: 完成>130%不扣分（跳过，数据有限）---
        
        # --- 标准6: 数量定性但相关不扣分 ---
        if '数量' in lev2 and nat == '定性' and is_deducted and '无关' not in rem:
            issues.append({'标准':'6','指标':ind,'问题':'数量指标定性但与项目相关，不应扣分','严重度':'P1','建议':'按比例复核评分，不扣分','备注原文':rem[:60]})
        
        # --- 标准7: 时效酌情扣30% ---
        if '时效' in lev2 and is_deducted and deducted_pct > 0:
            if deducted_pct > 45:
                issues.append({'标准':'7','指标':ind,'问题':f'时效指标扣{deducted_pct:.0f}%（远超口径30%）','严重度':'P1','建议':'多目标仅个别未完成，按口径扣30%','备注原文':rem[:60]})
        
        # --- 标准8: 效益定性产生效益不扣分 ---
        if ('效益' in lev1 or '效益' in lev2) and nat == '定性' and is_deducted:
            has_benefit = any(kw in rem for kw in ['产生','达到','实现','完成','有效','提升','改善','增强','协助','促进','良好','合格'])
            if has_benefit:
                issues.append({'标准':'8','指标':ind,'问题':f'效益定性产生效益但被扣分(自评{ss}→复核{rs})','严重度':'P1','建议':'按口径第8条不扣分','备注原文':rem[:60]})
        
        # --- 标准10: 三级不匹配二级 ---
        if '数量指标' in lev2 and any(kw in ind for kw in ['资金支付','资金拨付','预算执行','资金执行']):
            if rs > 0:
                issues.append({'标准':'10','指标':ind,'问题':f'数量指标三级为资金支付类({ind})，不反映公共产品或服务数量','严重度':'P0','建议':'按口径第10条打0分','备注原文':rem[:60]})
        
        # --- 标准14: 扣分原因合规性 ---
        if is_deducted:
            if not rem:
                issues.append({'标准':'14','指标':ind,'问题':f'扣{deducted_pct:.0f}%但备注为空，无扣分原因','严重度':'P1','建议':'必须填写扣分原因和依据口径','备注原文':''})
            else:
                # 14c: 备注写了"无佐证"但没打0分
                if ('无佐证' in rem or '无资料' in rem) and rs > 0:
                    issues.append({'标准':'14','指标':ind,'问题':f'备注说明无佐证资料但复核未打0分(={rs})','严重度':'P0','建议':'按口径第1条打0分','备注原文':rem[:60]})
                # 14c: 扣分比例写了但不一致
                pcts = re.findall(r'扣\s*(\d+)\s*%', rem)
                if pcts and abs(deducted_pct - float(pcts[0])) > 15:
                    issues.append({'标准':'14','指标':ind,'问题':f'备注写扣{pcts[0]}%但实际扣{deducted_pct:.0f}%','严重度':'P1','建议':'修正备注或复核得分','备注原文':rem[:60]})
        # 复核得分留空
        if is_blank and rem:
            issues.append({'标准':'14','指标':ind,'问题':f'复核得分留空，备注:"{rem[:40]}"','严重度':'P1','建议':'复核得分不应留空','备注原文':rem[:60]})
    
    # --- 标准9: 重复指标（项目级） ---
    ind_groups = defaultdict(list)
    for r in indicators:
        ind_groups[r['三级指标']].append(r['一级指标'])
    for ind, lv1s in ind_groups.items():
        if len(lv1s) >= 2:
            has_out = any('产出' in l for l in lv1s)
            has_ben = any('效益' in l for l in lv1s)
            if has_out and has_ben:
                issues.append({'标准':'9','指标':ind,'问题':'产出与效益指标重复设置，应只计其一、另一打0分','严重度':'P1','建议':'判断更适用产出/效益，保留一个正常评分','备注原文':''})
    
    return issues

# ---------- 报告层面检查 (11-13) ----------
def check_report_level(tables_data, doc):
    rpt = {'标准11':[], '标准12':[], '标准13':[]}
    
    # 标准11: 报告表vs底稿（由外部对比，这里只标记"待外部对比"）
    # 标准12: 逻辑
    full = '\n'.join([p.text for p in doc.paragraphs])
    pcts = re.findall(r'占\s*(\d+\.?\d*)\s*%', full)
    if len(pcts) >= 3:
        total = sum(float(x) for x in pcts[:4])
        if abs(total - 100) > 5:
            rpt['标准12'].append(f'指标分布百分比合计≈{total:.1f}%≠100%')
    
    # 标准13: 格式
    rpt['标准13'] = check_format_v3(doc)
    
    return rpt

# ---------- 主流程 ----------
def process_unit(unit_dir):
    unit = unit_dir.name
    docx_f = list(unit_dir.glob('*.docx'))
    xlsx_f = list(unit_dir.glob('*.xlsx'))
    if not docx_f or not xlsx_f: return None
    
    docx_p = str(docx_f[0]); xlsx_p = str(xlsx_f[0])
    tables_data, doc = read_docx_tables(docx_p)
    projects = read_xlsx_projects(xlsx_p)
    
    # 报告层面
    rpt = check_report_level(tables_data, doc)
    
    # 标准11: docx项目数 vs xlsx项目数
    docx_proj_count = 0
    for tab in tables_data:
        if tab and '项目名称' in str(tab[0]):
            docx_proj_count = sum(1 for r in tab[1:] if '合计' not in sfv(r[0]) and '小计' not in sfv(r[0]))
    xlsx_proj_count = len(projects)
    
    s11_ok = (docx_proj_count == xlsx_proj_count)
    if not s11_ok:
        rpt['标准11'].append(f'报告选取{docx_proj_count}个项目，盖章底稿{xlsx_proj_count}个sheet')
    
    # 指标层面
    all_issues = []
    proj_summary = []
    total_indicators = 0
    
    for pname, indicators in projects.items():
        iss = check_indicators(indicators, pname)
        # Tag each issue with project name
        for i in iss:
            i['项目'] = pname
        all_issues.extend(iss)
        total_indicators += len(indicators)
        
        p0 = sum(1 for i in iss if i['严重度']=='P0')
        p1 = sum(1 for i in iss if i['严重度']=='P1')
        p2 = sum(1 for i in iss if i['严重度']=='P2')
        proj_summary.append((pname, len(indicators), p0, p1, p2))
    
    # --- 生成Excel ---
    safe_u = re.sub(r'[\\/*?:\"<>|]','_', unit)[:35]
    out = DESKTOP / f'绩效复核_{safe_u}.xlsx'
    wb = openpyxl.Workbook()
    
    # ============ Sheet 1: 复核结论 ============
    ws = wb.active; ws.title = '复核结论'
    
    row = 1
    ws.merge_cells('A1:F1')
    c = ws.cell(row=1, column=1, value=f'绩效自评复核报告 — {unit}')
    c.font = TITLE_FONT; c.alignment = Alignment(horizontal='center')
    row += 2
    
    # ---- 一、四项核心结论 ----
    ws.merge_cells(f'A{row}:F{row}')
    ws.cell(row=row, column=1, value='一、四项复核结论').font = SEC_FONT
    row += 1
    
    conclusions = [
        ('11. 报告↔盖章底稿一致', s11_ok,
         '通过 ✅' if s11_ok else f'⚠ 报告{docx_proj_count}个≠底稿{xlsx_proj_count}个'),
        ('12. 报告内容逻辑', len(rpt['标准12'])==0,
         '通过 ✅' if len(rpt['标准12'])==0 else f'⚠ {len(rpt["标准12"])}个问题'),
        ('13. 报告格式合规', len(rpt['标准13'])==0,
         '通过 ✅' if len(rpt['标准13'])==0 else f'⚠ {len(rpt["标准13"])}个问题'),
        ('14. 扣分原因合规', sum(1 for i in all_issues if i['标准']=='14')==0,
         f'通过 ✅' if sum(1 for i in all_issues if i['标准']=='14')==0 else f'⚠ {sum(1 for i in all_issues if i["标准"]=="14")}个问题'),
    ]
    
    conc_headers = ['复核项','状态','结论']
    for ci, h in enumerate(conc_headers, 1):
        cell = ws.cell(row=row, column=ci, value=h)
        cell.font = HDR_FONT; cell.fill = HDR_FILL; cell.alignment = ALIGN_C; cell.border = BORDER
    row += 1
    
    for item, ok, msg in conclusions:
        ws.cell(row=row, column=1, value=item).font = BOLD_FONT
        ws.cell(row=row, column=2, value='✅' if ok else '⚠️')
        c3 = ws.cell(row=row, column=3, value=msg)
        c3.font = Font(name='仿宋', size=10, color='006400' if ok else 'CC0000')
        for ci in range(1,4): ws.cell(row=row, column=ci).border = BORDER; ws.cell(row=row, column=ci).alignment = ALIGN_W
        row += 1
    
    row += 1
    
    # ---- 二、指标复核统计 ----
    ws.merge_cells(f'A{row}:F{row}')
    ws.cell(row=row, column=1, value='二、指标复核统计').font = SEC_FONT
    row += 1
    
    # 汇总
    p0_total = sum(1 for i in all_issues if i['严重度']=='P0')
    p1_total = sum(1 for i in all_issues if i['严重度']=='P1')
    p2_total = sum(1 for i in all_issues if i['严重度']=='P2')
    
    stats = [
        ('复核项目数', len(projects)),
        ('复核指标数', total_indicators),
        ('P0 致命', f'{p0_total} 项 — 必须修改'),
        ('P1 重要', f'{p1_total} 项 — 建议修改'),
        ('P2 优化', f'{p2_total} 项 — 可优化'),
    ]
    for label, val in stats:
        ws.cell(row=row, column=1, value=label).font = BOLD_FONT
        ws.cell(row=row, column=2, value=val).font = BODY_FONT
        ws.cell(row=row, column=1).border = BORDER; ws.cell(row=row, column=2).border = BORDER
        row += 1
    
    row += 1
    
    # 逐项目摘要
    ws.merge_cells(f'A{row}:F{row}')
    ws.cell(row=row, column=1, value='逐项目情况').font = SEC_FONT
    row += 1
    
    pj_headers = ['序号','项目名称','指标数','P0','P1','P2']
    for ci, h in enumerate(pj_headers, 1):
        cell = ws.cell(row=row, column=ci, value=h)
        cell.font = HDR_FONT; cell.fill = HDR_FILL; cell.alignment = ALIGN_C; cell.border = BORDER
    row += 1
    for idx, (pn, cnt, p0, p1, p2) in enumerate(proj_summary, 1):
        vals = [idx, pn[:40], cnt, p0, p1, p2]
        for ci, v in enumerate(vals, 1):
            cell = ws.cell(row=row, column=ci, value=v)
            cell.font = BODY_FONT; cell.alignment = ALIGN_W; cell.border = BORDER
            if ci >= 4 and v > 0: cell.fill = FILL_P0 if ci==4 else (FILL_P1 if ci==5 else FILL_P2)
        row += 1
    
    row += 2
    
    # ---- 三、P0/P1 问题清单 ----
    if p0_total + p1_total > 0:
        ws.merge_cells(f'A{row}:F{row}')
        ws.cell(row=row, column=1, value='三、需处理问题清单（P0/P1）').font = SEC_FONT
        row += 1
        
        iss_headers = ['序号','标准','项目','指标','问题描述','严重度']
        for ci, h in enumerate(iss_headers, 1):
            cell = ws.cell(row=row, column=ci, value=h)
            cell.font = HDR_FONT; cell.fill = HDR_FILL; cell.alignment = ALIGN_C; cell.border = BORDER
        row += 1
        
        priority_issues = [i for i in all_issues if i['严重度'] in ('P0','P1')]
        priority_issues.sort(key=lambda x: (0 if x['严重度']=='P0' else 1, x['标准']))
        
        for idx, iss in enumerate(priority_issues, 1):
            vals = [idx, iss['标准'], iss.get('项目','')[:25], iss['指标'][:30], iss['问题'], iss['严重度']]
            for ci, v in enumerate(vals, 1):
                cell = ws.cell(row=row, column=ci, value=str(v))
                cell.font = BODY_FONT; cell.alignment = ALIGN_W; cell.border = BORDER
                if iss['严重度'] == 'P0': cell.fill = FILL_P0
                elif iss['严重度'] == 'P1': cell.fill = FILL_P1
            row += 1
    
    # Column widths
    widths = [6, 8, 28, 30, 55, 8]
    for ci, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.freeze_panes = 'A2'
    
    # ============ Sheet 2: 格式问题清单 ============
    ws2 = wb.create_sheet('格式问题')
    fmt_h = ['序号','类型','位置','问题详情','严重度']
    for ci, h in enumerate(fmt_h, 1):
        cell = ws2.cell(row=1, column=ci, value=h)
        cell.font = HDR_FONT; cell.fill = HDR_FILL; cell.alignment = ALIGN_C; cell.border = BORDER
    
    for idx, iss in enumerate(rpt['标准13'], 1):
        vals = [idx, iss.get('类型',''), iss.get('位置',''), iss.get('详情',''), iss.get('严重度','')]
        for ci, v in enumerate(vals, 1):
            cell = ws2.cell(row=idx+1, column=ci, value=str(v))
            cell.font = BODY_FONT; cell.alignment = ALIGN_W; cell.border = BORDER
            if iss.get('严重度') == 'P1': cell.fill = FILL_P1
    
    fmt_w = [6, 12, 12, 55, 8]
    for ci, w in enumerate(fmt_w, 1):
        ws2.column_dimensions[get_column_letter(ci)].width = w
    ws2.freeze_panes = 'A2'
    
    # ============ Sheet 3: 逐项目明细 ============
    for pname, indicators in projects.items():
        safe_p = re.sub(r'[\\/*?:\"<>|\[\]]','_', pname)[:28]
        ws3 = wb.create_sheet(safe_p)
        
        # 指标表
        ind_h = ['序号','一级','二级','三级指标','性质','指标值','权重','自评得分','复核得分','有问题']
        for ci, h in enumerate(ind_h, 1):
            cell = ws3.cell(row=1, column=ci, value=h)
            cell.font = HDR_FONT; cell.fill = HDR_FILL; cell.alignment = ALIGN_C; cell.border = BORDER
        
        proj_issues = [i for i in all_issues if i['指标'] in [r['三级指标'] for r in indicators]]
        
        for idx, r in enumerate(indicators, 1):
            has_issue = any(i['指标'] == r['三级指标'] for i in proj_issues)
            vals = [idx, r['一级指标'][:10], r['二级指标'][:10], r['三级指标'][:30],
                    r['指标性质'], str(r['指标值'])[:15] if r['指标值'] else '',
                    r['权重'], r['自评得分'], r['复核得分'],
                    '⚠' if has_issue else '✓']
            for ci, v in enumerate(vals, 1):
                cell = ws3.cell(row=idx+1, column=ci, value=str(v))
                cell.font = BODY_FONT; cell.alignment = ALIGN_W; cell.border = BORDER
                if ci == 10 and has_issue: cell.fill = FILL_P1
            # 备注
            if r['备注']:
                ws3.cell(row=idx+1, column=11, value=r['备注'][:80]).font = BODY_FONT
        
        # 该项目的issue
        if proj_issues:
            gap_row = len(indicators) + 3
            ws3.cell(row=gap_row, column=1, value='问题详情:').font = BOLD_FONT
            iss_h2 = ['标准','指标','问题','严重度','建议']
            for ci, h in enumerate(iss_h2, 1):
                ws3.cell(row=gap_row+1, column=ci, value=h).font = HDR_FONT
                ws3.cell(row=gap_row+1, column=ci).fill = HDR_FILL
                ws3.cell(row=gap_row+1, column=ci).border = BORDER
            for i2, iss in enumerate(proj_issues):
                vals2 = [iss['标准'], iss['指标'][:30], iss['问题'], iss['严重度'], iss.get('建议','')]
                for ci, v in enumerate(vals2, 1):
                    cell = ws3.cell(row=gap_row+2+i2, column=ci, value=str(v))
                    cell.font = BODY_FONT; cell.border = BORDER
        
        ind_w = [5,8,8,30,6,12,6,8,8,6,40]
        for ci, w in enumerate(ind_w, 1):
            ws3.column_dimensions[get_column_letter(ci)].width = w
        ws3.freeze_panes = 'A2'
    
    wb.save(str(out))
    return {
        'unit': unit, 'file': out.name,
        'projects': len(projects), 'indicators': total_indicators,
        'P0': p0_total, 'P1': p1_total, 'P2': p2_total,
        's11_ok': s11_ok, 's12_ok': len(rpt['标准12'])==0,
        's13_ok': len(rpt['标准13'])==0, 's14_ok': sum(1 for i in all_issues if i['标准']=='14')==0,
        'fmt_count': len(rpt['标准13'])
    }

def main():
    print('绩效自评复核 V3')
    results = []
    for d in sorted(BASE.iterdir()):
        if d.is_dir():
            r = process_unit(d)
            if r: results.append(r)
            print(f'  {r["unit"][:30]:30s} 项目{r["projects"]:2d} 指标{r["indicators"]:3d}  P0:{r["P0"]} P1:{r["P1"]} P2:{r["P2"]}')
    
    # Final summary
    print(f'\n{"="*70}')
    print(f'{"单位":35s} {"11一致":>6s} {"12逻辑":>6s} {"13格式":>6s} {"14原因":>6s}')
    print('-'*60)
    for r in results:
        s11 = '✅' if r['s11_ok'] else '⚠️'
        s12 = '✅' if r['s12_ok'] else '⚠️'
        s13 = '✅' if r['s13_ok'] else f'⚠ {r["fmt_count"]}'
        s14 = '✅' if r['s14_ok'] else '⚠️'
        print(f'{r["unit"][:35]:35s} {s11:>6s} {s12:>6s} {str(s13):>6s} {s14:>6s}')
    
    total = sum(r['P0'] for r in results)
    t1 = sum(r['P1'] for r in results)
    t2 = sum(r['P2'] for r in results)
    print(f'\n总计: P0={total} P1={t1} P2={t2}  | 输出: {DESKTOP}')

if __name__ == '__main__':
    main()
