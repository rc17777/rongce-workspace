#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取融策公司所有制度文件、岗位说明书、招聘文档的文本内容
"""

import os, sys, json
sys.stdout.reconfigure(encoding='utf-8')

# Try to import docx
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import olefile
except ImportError:
    olefile = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

BASE = r"C:\Users\scrccpa\Desktop\公司制度"

files = [
    # 制度文件
    (r"制度\【20220420】四川融策会计师事务所有限公司薪酬奖金及福利待遇制度.docx", "薪酬奖金及福利待遇制度"),
    (r"制度\【制度】员工手册2023.1.1公布（20230504）0616.docx", "员工手册"),
    (r"制度\差旅报销制度20260130（讨论版未定）.docx", "差旅报销制度2026讨论版"),
    (r"制度\差旅报销制度（2025）.docx", "差旅报销制度2025"),
    (r"制度\差旅费借支申请表.doc", "差旅费借支申请表"),
    (r"制度\差旅费附注.xlsx", "差旅费附注"),
    (r"制度\车辆管理办法06docx.docx", "车辆管理办法"),
    # 岗位说明书
    (r"岗位说明书0620\业务专员岗位说明书.docx", "业务专员岗位说明书"),
    (r"岗位说明书0620\业务经理岗位说明书.docx", "业务经理岗位说明书"),
    (r"岗位说明书0620\人事岗位说明书.docx", "人事岗位说明书"),
    (r"岗位说明书0620\复核岗位说明书.docx", "复核岗位说明书"),
    (r"岗位说明书0620\审计助理岗位说明书.doc", "审计助理岗位说明书"),
    (r"岗位说明书0620\审计总监岗位说明书.docx", "审计总监岗位说明书"),
    (r"岗位说明书0620\审计经理助理岗位说明书.doc", "审计经理助理岗位说明书"),
    (r"岗位说明书0620\审计经理岗位说明书.docx", "审计经理岗位说明书"),
    (r"岗位说明书0620\技术总监岗位说明书.docx", "技术总监岗位说明书"),
    (r"岗位说明书0620\招投标助理岗位说明书.docx", "招投标助理岗位说明书"),
    (r"岗位说明书0620\行政专员岗位说明书.docx", "行政专员岗位说明书"),
    (r"岗位说明书0620\财务岗位说明书.doc", "财务岗位说明书"),
    (r"岗位说明书0620\造价员岗位说明书.docx", "造价员岗位说明书"),
    (r"岗位说明书0620\造价部经理岗位说明书.docx", "造价部经理岗位说明书"),
    (r"岗位说明书0620\项目经理岗位说明书.doc", "项目经理岗位说明书"),
    # 招聘
    (r"招聘岗位.docx", "招聘岗位"),
]

def extract_docx(path):
    """Extract text from .docx file"""
    try:
        doc = Document(path)
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract tables
        tables_text = []
        for i, table in enumerate(doc.tables):
            tables_text.append(f"\n--- 表格 {i+1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(cells))
        return "\n".join(paragraphs) + "\n" + "\n".join(tables_text)
    except Exception as e:
        return f"[提取失败: {e}]"

def extract_doc(path):
    """Extract text from legacy .doc file using olefile"""
    if olefile is None:
        return "[无法提取: 未安装olefile]"
    try:
        ole = olefile.OleFileIO(path)
        # Try to read WordDocument stream
        if ole.exists('WordDocument'):
            # Try to extract text from the Word stream
            data = ole.openstream('WordDocument').read()
            # Simple extraction: look for readable text
            text = ""
            for i in range(0, len(data)-1, 2):
                if data[i] >= 0x20 and data[i] <= 0x7e:
                    text += chr(data[i])
                elif data[i] == 0x0d:
                    text += "\n"
            # Also try to read 1Table or 0Table
            for stream_name in ole.listdir():
                name = "/".join(stream_name)
                if 'Table' in name:
                    try:
                        tbl_data = ole.openstream(stream_name).read()
                        # Try to decode as UTF-16LE
                        try:
                            tbl_text = tbl_data.decode('utf-16le', errors='ignore')
                            # Filter printable CJK and ASCII
                            filtered = ''.join(c for c in tbl_text if c.isprintable() or c in '\n\r\t')
                            if len(filtered) > len(text):
                                text = filtered
                        except:
                            pass
                    except:
                        pass
            ole.close()
            return text[:50000] if text else "[提取失败: 未找到可读文本]"
        else:
            ole.close()
            return "[无法提取: 不支持的Word文档格式]"
    except Exception as e:
        return f"[提取失败: {e}]"

def extract_xlsx(path):
    """Extract text from .xlsx file"""
    if openpyxl is None:
        return "[无法提取: 未安装openpyxl]"
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"\n=== 工作表: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"[提取失败: {e}]"

def main():
    output_dir = r"C:\Users\scrccpa\.openclaw\workspace\output\extracted"
    os.makedirs(output_dir, exist_ok=True)
    
    summary = {}
    
    for rel_path, label in files:
        full_path = os.path.join(BASE, rel_path)
        print(f"\n{'='*60}")
        print(f"处理: {label}")
        print(f"路径: {full_path}")
        
        if not os.path.exists(full_path):
            print(f"[文件不存在]")
            summary[rel_path] = {"label": label, "status": "文件不存在", "text": ""}
            continue
        
        ext = os.path.splitext(full_path)[1].lower()
        
        if ext == '.docx':
            text = extract_docx(full_path)
        elif ext == '.doc':
            text = extract_doc(full_path)
        elif ext == '.xlsx':
            text = extract_xlsx(full_path)
        else:
            text = "[不支持的格式]"
        
        # Save extracted text
        safe_name = label.replace('/', '_').replace('\\', '_').replace(' ', '_')
        out_path = os.path.join(output_dir, f"{safe_name}.txt")
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        print(f"提取长度: {len(text)} 字符")
        print(f"已保存: {out_path}")
        
        summary[rel_path] = {"label": label, "status": "ok", "text": text, "len": len(text)}
    
    # Save summary
    summary_path = os.path.join(output_dir, "_summary.json")
    # Don't save full text to JSON (too large), just metadata
    meta_summary = {k: {"label": v["label"], "status": v["status"], "len": v.get("len", 0)} for k, v in summary.items()}
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(meta_summary, f, ensure_ascii=False, indent=2)
    print(f"\n元数据已保存: {summary_path}")
    print(f"\n提取完成。共处理 {len(files)} 个文件。")

if __name__ == '__main__':
    main()