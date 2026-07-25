# -*- coding: utf-8 -*-
"""融策制度分类汇编Word生成器 — 4册"""
import os, re, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\scrccpa\.openclaw\workspace\output\新制度体系'
DESKTOP = r'C:\Users\scrccpa\Desktop'

FIXES = {
    '02-绩效考核管理制度.md': [
        (r'D等比例原则上控制在.*?无D等。', 'D等为"不称职"等级，原则上作为绩效改进的参考指标，不设硬性比例强制分布。'),
        ('考核结果申诉', '特殊情形考核\n\n第二十一条 处于法定产假、哺乳假、工伤停工留薪期、医疗期内的员工，考核按以下方式处理：（1）整个考核周期均在上述期间的，不参与当期考核排名，绩效工资按基数100%发放；（2）部分在保护期的，按实际在岗时间可量化的工作指标考核，缺少数据的维度不扣分。\n\n考核结果申诉'),
        ('本制度经管理层审议通过后施行', '本制度经职工代表大会讨论通过后施行'),
        (r'360度协作评分.*?团队协作', '团队协作评分由项目经理和部门负责人根据日常观察直接评分'),
        (r'季度考核', '半年度考核'),
        (r'每季度', '每半年'),
        (r'20个工作日', '5个工作日'),
        ('改进期满仍为D的.*?解除劳动合同', '改进期满仍为D的，依据《劳动合同法》第40条第（二）项规定处理，依法支付经济补偿'),
    ],
    '03-员工手册.md': [
        (r'考核不合格的.*?不得延长试用期。', '考核不合格的，公司应在试用期内依据《劳动合同法》第39条第（一）项规定解除劳动合同。'),
        (r'晚婚.*?增加\d+天.*?。', '法定婚假3天，婚假期间工资全额发放。'),
    ],
    '06-财务报销管理制度.md': [
        (r'专票认证.*?360.*?天.*?。', '增值税专用发票取消认证期限。'),
    ],
    '07-审计质量控制制度.md': [
        (r'扣减当年全部绩效工资', '扣减绩效工资（月扣≤20%工资）'),
        (r'扣减.*?50%.*?绩效', '扣减绩效工资（月扣≤20%工资）'),
    ],
    '08-造价咨询质量控制制度.md': [
        (r'扣减当年全部绩效工资', '扣减绩效工资（月扣≤20%工资）'),
    ],
    '10-招聘与入职管理制度.md': [
        (r'基本合格.*?延长试用期.*?法定上限', '试用期考核不合格的，依据《劳动合同法》第39条解除劳动合同'),
    ],
    '13-项目收入确认与回款管理制度.md': [
        (r'已发放的予以追回', '已发放的不予追回'),
    ],
    '22-执业责任追究制度.md': [
        (r'扣减全年绩效奖金', '扣减绩效奖金（月扣≤20%工资）'),
    ],
}

def apply_fixes(filename, content):
    if filename not in FIXES: return content
    for old, new in FIXES[filename]:
        content = re.sub(old, new, content, flags=re.DOTALL)
    return content

def read_md(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_cover(doc, book_title, subtitle=''):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('四川融策')
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(32); r.font.color.rgb = RGBColor(0x0A, 0x1F, 0x3F); r.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(book_title)
    r2.font.name = '黑体'; r2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r2.font.size = Pt(28); r2.font.color.rgb = RGBColor(0x0A, 0x1F, 0x3F)
    doc.add_paragraph()
    info = ['版本号：V1.0', '编制日期：2026年7月22日', '审核人：__________  批准人：__________', '', '四川融策会计师事务所有限公司', '四川融策工程咨询有限公司']
    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'); r.font.size = Pt(14)
    doc.add_page_break()

def parse_content(doc, content):
    lines = content.split('\n')
    for line in lines:
        s = line.strip()
        if not s: continue
        if s.startswith('# ') and not s.startswith('## '):
            h = doc.add_heading(s[2:], level=1)
            for r in h.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif s.startswith('## '):
            h = doc.add_heading(s[3:], level=2)
            for r in h.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif s.startswith('### '):
            h = doc.add_heading(s[4:], level=3)
            for r in h.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif s.startswith('```'): continue
        else:
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            p = doc.add_paragraph(clean)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'); r.font.size = Pt(14)

def setup_section(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.8)

def build_book(book_name, filename, file_list):
    print(f'  生成: {filename}')
    doc = Document()
    setup_section(doc)
    add_cover(doc, book_name)
    
    for f in file_list:
        filepath = os.path.join(SRC, f)
        if not os.path.exists(filepath):
            print(f'    ⚠ 跳过(不存在): {f}')
            continue
        content = read_md(filepath)
        content = apply_fixes(f, content)
        doc.add_page_break()
        parse_content(doc, content)
    
    out = os.path.join(DESKTOP, filename)
    doc.save(out)
    return out

# Category definitions
books = [
    ('融策制度汇编-人力资源篇', '融策制度汇编-人力资源篇.docx', [
        '00-制度体系架构.md',
        '01-薪酬管理制度.md', '02-绩效考核管理制度.md', '03-员工手册.md',
        '10-招聘与入职管理制度.md', '11-培训与发展管理制度.md', '12-职级晋升管理制度.md',
    ]),
    ('融策制度汇编-财务管理篇', '融策制度汇编-财务管理篇.docx', [
        '06-财务报销管理制度.md', '13-项目收入确认与回款管理制度.md',
        '14-预算管理制度.md', '15-资金管理制度.md', '16-固定资产管理制度.md',
    ]),
    ('融策制度汇编-业务质控篇', '融策制度汇编-业务质控篇.docx', [
        '04-项目管理规范.md', '17-业务承接与合同管理制度.md', '18-客户关系管理制度.md',
        '19-业务分包管理制度.md', '20-投标管理制度.md',
        '07-审计质量控制制度.md', '08-造价咨询质量控制制度.md',
        '21-三级复核实施细则.md', '22-执业责任追究制度.md',
    ]),
    ('融策制度汇编-行政综合篇', '融策制度汇编-行政综合篇.docx', [
        '26-公司章程-会计师事务所.md', '27-公司章程-工程咨询公司.md', '09-股东会议事规则.md',
        '23-信息安全与保密管理制度.md', '24-印章与证照管理制度.md', '25-档案管理制度.md',
        '28-办公场所管理制度.md', '29-采购管理制度.md',
        '30-数智化建设管理制度.md', '31-业务拓展与创新管理制度.md',
        '32-风险管理制度.md', '33-党建工作制度.md',
        '05-制度发布与版本管理规范.md',
    ]),
]

for title, fname, flist in books:
    out = build_book(title, fname, flist)
    sz = os.path.getsize(out)
    print(f'    ✅ {fname} ({sz:,} bytes)')

print('\n全部完成!')
