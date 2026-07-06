from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import landscape, A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

OUT_DIR = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
WORK_DIR = Path("work/sichuan_rongce_brochure/v3_assets")
DOCX_PATH = OUT_DIR / "四川融策宣传册_完善稿_v3_原风格图文版.docx"
PDF_PATH = OUT_DIR / "四川融策宣传册_完善稿_v3_原风格图文版.pdf"
SOURCE_DIR = Path("work/sichuan_rongce_brochure/baidu_ocr/pages")

BLUE = "0A1F3F"
TEAL = "1A5C6E"
GOLD = "C5955C"
WARM = "F5F2EC"
INK = "2B2B2B"
GRAY = "667085"

PAGES = [
    {
        "kind": "cover",
        "title": "四川融策",
        "en": "SICHUAN RONGCE",
        "subtitle": "政府审计与工程咨询综合服务机构",
        "lead": "谋专业之策  融品质之精",
        "note": "审计 · 绩效 · 财政监督 · 工程咨询 · 数字化分析",
    },
    {
        "kind": "toc",
        "title": "目录",
        "en": "CONTENT",
        "lead": "围绕财政资金、公共资源和工程建设，形成审计、绩效、咨询和数据分析一体化服务。",
        "items": ["01 公司优势", "02 公司介绍", "03 服务体系", "04 政府审计", "05 预算绩效", "06 工程咨询", "07 数字化能力", "08 代表经验", "09 合作价值"],
    },
    {
        "kind": "content",
        "title": "公司优势",
        "en": "COMPANY ADVANTAGE",
        "lead": "融策的优势，不只在资质和年限，更在对政府业务、财政资金、工程项目和基层执行场景的长期理解。",
        "items": ["专业团队：注册会计师、造价师、高级会计师等复合力量协同作业。", "场景经验：长期服务财政、审计、教育、民政、农业农村、交通、国资、医保等领域。", "方法沉淀：形成项目案例库、风险问题库、指标体系和数据识别规则。", "质量控制：重大金额和关键结论复核到数据来源、计算方法和佐证材料。"],
    },
    {
        "kind": "content",
        "title": "公司介绍",
        "en": "COMPANY INTRODUCTION",
        "lead": "四川融策会计师事务所有限公司成立于2000年，是四川省内较早成立的会计师事务所之一。总部位于成都，并设有阿坝州、西藏等服务支点。",
        "items": ["20余年专业服务积累，服务客户400余家。", "业务覆盖政府审计、预算绩效、财政监督、工程咨询和管理咨询。", "坚持客观公正、实事求是、质量优先，重视证据闭环和整改落地。", "以专业判断和数据能力服务财政管理现代化。"],
    },
    {
        "kind": "content",
        "title": "服务体系",
        "en": "BUSINESS SCOPE",
        "lead": "从单项审计到全过程咨询，从结果评价到制度优化，从人工核查到数据化识别，形成六类综合服务。",
        "items": ["政府审计：经济责任、财务收支、专项资金、资产清查、国企审计、竣工决算财务审计。", "预算绩效：事前评估、目标审核、运行监控、重点评价、结果应用。", "财政监督：预算执行、财经纪律、政府采购、资产管理、会计信息质量。", "工程咨询：预算编制、清单控制价、财政评审、结算审核、全过程工程咨询。", "采购审计：采购程序、合同履约、资金支付、围标串标线索。", "管理咨询：内控建设、资产管理、整改提升、流程优化。"],
    },
    {
        "kind": "content",
        "title": "政府审计",
        "en": "GOVERNMENT AUDIT",
        "lead": "面向财政、审计、主管部门和国有企事业单位，融策把账表、合同、项目、资金、资产和责任链条贯通核查。",
        "items": ["经济责任审计：关注重大决策、资金资产、项目建设、内部控制和廉政风险。", "专项资金审计：核查申报、分配、拨付、使用、绩效和结余沉淀。", "财政监督检查：聚焦预算执行、财经纪律、政府采购、资产管理和会计信息质量。", "竣工决算财务审计：核实建设成本归集、资金来源、资产交付、尾工尾款。"],
    },
    {
        "kind": "content",
        "title": "预算绩效管理",
        "en": "BUDGET PERFORMANCE",
        "lead": "围绕预算编制、执行监控、绩效评价和结果应用，帮助财政部门和预算单位把资金使用效果说清楚、评准确、用起来。",
        "items": ["事前绩效评估：必要性、可行性、财政承受能力和预期绩效。", "绩效目标审核：目标完整性、指标可衡量性和预算匹配性。", "绩效运行监控：执行进度、资金支付、产出偏差和风险预警。", "重点绩效评价：政策、部门整体、项目支出和专项资金评价。", "结果应用：整改清单、预算挂钩建议和管理制度优化。"],
    },
    {
        "kind": "content",
        "title": "工程咨询与财政评审",
        "en": "ENGINEERING CONSULTING",
        "lead": "工程咨询业务重点服务投资控制和项目管理，帮助委托方把工程造价、合同履约、资金支付和项目绩效对齐。",
        "items": ["预算编制与财政评审：核工程量、定额套用、材料价格、措施费和取费标准。", "清单及招标控制价：提升招标文件和控制价编制质量。", "结算审核：核查合同条款、变更签证、隐蔽工程、现场工程量和支付资料。", "全过程工程咨询：协同项目前期、招采、实施、验收、结算和绩效评价。"],
    },
    {
        "kind": "content",
        "title": "数字化审计能力",
        "en": "DIGITAL AUDIT",
        "lead": "融策将审计经验沉淀为数据标准、识别规则、疑点模型和底稿模板，提高覆盖面、发现率和证据质量。",
        "items": ["数据标准：整理财务、预算、支付、合同、采购、资产和工程项目字段。", "规则模型：识别重复支付、超预算执行、供应商异常、项目进度异常和资金沉淀。", "穿透核查：形成疑点来源、核查路径、佐证材料、影响金额和整改建议。", "报告复核：校验金额汇总、口径一致、附表闭环和结论依据。"],
    },
    {
        "kind": "content",
        "title": "代表经验与联系",
        "en": "EXPERIENCE & CONTACT",
        "lead": "融策长期服务省、市、县多级财政和主管部门，项目覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。",
        "items": ["代表客户：四川省财政厅、省公安厅交警总队、省民政厅、省农业农村厅、省教育厅等。", "市县经验：达州、绵阳、宜宾、什邡、德阳、康定、九寨沟等地财政项目。", "重点领域：交通、教育、民政、农业农村、医保等民生和公共治理场景。", "电话：028-87659276  邮箱：scrccpa@163.com  网址：www.scrccpa.com"],
    },
]


def rgb(hex_value: str) -> tuple[int, int, int]:
    hex_value = hex_value.strip("#")
    return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))


def get_font(size: int, bold: bool = False):
    paths = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_text_box(draw: ImageDraw.ImageDraw, text: str, x: int, y: int, w: int, font, fill, line_gap: int = 8) -> int:
    line = ""
    lines = []
    for ch in text:
        trial = line + ch
        if draw.textlength(trial, font=font) <= w:
            line = trial
        else:
            if line:
                lines.append(line)
            line = ch
    if line:
        lines.append(line)
    for item in lines:
        draw.text((x, y), item, font=font, fill=fill)
        y += font.size + line_gap
    return y


def load_bg(index: int) -> Image.Image:
    source = SOURCE_DIR / f"page_{index:02d}.png"
    if not source.exists():
        source = Path("work/sichuan_rongce_brochure") / f"page_{index:02d}.png"
    img = Image.open(source).convert("RGB")
    return img.resize((1600, 900))


def soften_background(img: Image.Image, index: int) -> Image.Image:
    if index in {1, 2, 3, 5, 10}:
        return img.filter(ImageFilter.GaussianBlur(radius=0.4))
    overlay = Image.new("RGBA", img.size, (255, 255, 255, 188))
    return Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")


def make_page(page: dict, index: int) -> Path:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    img = soften_background(load_bg(index), index)
    draw = ImageDraw.Draw(img, "RGBA")
    blue = rgb(BLUE)
    gold = rgb(GOLD)
    teal = rgb(TEAL)
    ink = rgb(INK)
    gray = rgb(GRAY)

    if page["kind"] == "cover":
        draw.rectangle((0, 0, 1600, 900), fill=(31, 18, 15, 218))
        draw.text((930, 500), "四川融策会计师事务所有限公司", font=get_font(32, True), fill=gold)
        draw.text((930, 545), page["subtitle"], font=get_font(24), fill=(232, 214, 181))
        draw.text((300, 390), page["lead"], font=get_font(38, True), fill=gold)
        draw.text((300, 460), page["note"], font=get_font(25), fill=(232, 214, 181))
        draw.line((300, 535, 780, 535), fill=gold, width=5)
    elif page["kind"] == "toc":
        draw.rectangle((78, 88, 735, 810), fill=(255, 255, 255, 222))
        draw.rectangle((780, 88, 1525, 810), fill=(255, 255, 255, 218))
        draw.text((135, 140), page["en"], font=get_font(72, True), fill=(78, 70, 68))
        draw.text((138, 235), page["title"], font=get_font(34, True), fill=ink)
        draw_text_box(draw, page["lead"], 830, 170, 600, get_font(30), gray, 10)
        y = 335
        for item in page["items"]:
            no, label = item.split(" ", 1)
            draw.text((150, y), no, font=get_font(34, True), fill=gold)
            draw.text((220, y + 7), label, font=get_font(27, True), fill=ink)
            y += 62
        draw.text((1230, 725), "RONGCE", font=get_font(42, True), fill=gold)
    else:
        draw.rectangle((76, 70, 1528, 828), fill=(255, 255, 255, 226))
        draw.rectangle((76, 70, 1528, 83), fill=gold)
        draw.text((122, 120), page["title"], font=get_font(52, True), fill=gold)
        draw.text((124, 184), page["en"], font=get_font(25, True), fill=(75, 69, 68))
        draw_text_box(draw, page["lead"], 620, 128, 780, get_font(29), gray, 10)

        # Left-side decorative panel keeps the original brochure's visual density.
        draw.rounded_rectangle((120, 290, 515, 710), radius=18, fill=blue + (235,), outline=gold, width=4)
        draw.text((168, 338), "RONGCE", font=get_font(43, True), fill=gold)
        draw.text((168, 397), "专业 · 客观 · 公正", font=get_font(28, True), fill=(255, 255, 255))
        for i, word in enumerate(["审计", "绩效", "监督", "咨询"]):
            x = 165 + (i % 2) * 165
            y = 485 + (i // 2) * 92
            draw.ellipse((x, y, x + 70, y + 70), fill=teal + (255,), outline=gold, width=3)
            draw.text((x + 13, y + 17), word, font=get_font(24, True), fill=(255, 255, 255))

        y = 315
        for item in page["items"]:
            draw.rectangle((610, y + 10, 626, y + 36), fill=gold)
            y = draw_text_box(draw, item, 645, y, 760, get_font(27), ink, 8) + 18
        draw.text((1145, 780), "公开公正  用心服务  诚信为本  服务至上  追求卓越", font=get_font(22), fill=gold)

    out = WORK_DIR / f"page_{index:02d}.png"
    img.save(out, quality=95)
    return out


def register_pdf_font() -> str:
    for path, name in [(r"C:\Windows\Fonts\msyh.ttc", "MSYH"), (r"C:\Windows\Fonts\simhei.ttf", "SIMHEI")]:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont(name, path))
            return name
    return "Helvetica"


def generate_pdf(images: list[Path]) -> None:
    page_w, page_h = landscape(A4)
    c = canvas.Canvas(str(PDF_PATH), pagesize=landscape(A4))
    for idx, img in enumerate(images):
        c.drawImage(str(img), 0, 0, width=page_w, height=page_h, preserveAspectRatio=False, mask="auto")
        if idx != len(images) - 1:
            c.showPage()
    c.save()


def set_run(run, size=11, bold=False, color=INK, font_name="微软雅黑") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def generate_docx(images: list[Path]) -> None:
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v3原风格图文版"
    for idx, (page, image) in enumerate(zip(PAGES, images), start=1):
        if idx > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.orientation = 1
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.top_margin = Cm(1.0)
        sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2)
        sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(page["title"])
        set_run(r, 24, True, BLUE)
        p = doc.add_paragraph()
        r = p.add_run(page.get("subtitle") or page.get("lead", ""))
        set_run(r, 11, False, GRAY, "宋体")
        doc.add_picture(str(image), width=Cm(26.8))
        if page.get("items"):
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell_shading(table.cell(0, 0), BLUE)
            r = table.cell(0, 0).paragraphs[0].add_run("文案要点")
            set_run(r, 10, True, "FFFFFF")
            for item in page["items"]:
                c = table.add_row().cells[0]
                r = c.paragraphs[0].add_run("• " + item)
                set_run(r, 9.5, False, INK, "宋体")
    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    images = [make_page(page, idx) for idx, page in enumerate(PAGES, start=1)]
    generate_pdf(images)
    generate_docx(images)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
