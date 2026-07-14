# -*- coding: utf-8 -*-
"""生成投资人演讲稿Word文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json, os

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

NAVY = RGBColor(0x0B, 0x1D, 0x3A)
GRAY = RGBColor(0x6B, 0x7B, 0x8D)
RED = RGBColor(0xC0, 0x39, 0x2B)

def ap(text, size=11, bold=False, color=None, align=None, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    return p

def body(text, size=11):
    for line in text.strip().split('\n'):
        line = line.strip()
        if not line:
            ap('', size=size, sa=2)
            continue
        if line.startswith('**') and '**' in line[2:]:
            idx = line.index('**', 2)
            prefix = line[2:idx]
            suffix = line[idx+2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.35
            r1 = p.add_run(prefix)
            r1.font.size = Pt(size)
            r1.font.name = '微软雅黑'
            r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r1.bold = True
            if suffix:
                r2 = p.add_run(suffix)
                r2.font.size = Pt(size)
                r2.font.name = '微软雅黑'
                r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            ap(line, size=size, sa=4)

def sep():
    ap('─' * 60, size=8, color=GRAY, sa=10)

def hdr(num, title, dur=''):
    tag = f'━━━ Slide {num} ━━━ {title}'
    if dur: tag += f'（{dur}）'
    ap(tag, size=14, bold=True, color=NAVY, sa=8)

# Load content from JSON
data_path = os.path.join(os.path.dirname(__file__), 'speech_content.json')
with open(data_path, 'r', encoding='utf-8') as f:
    content = json.load(f)

# Title page
ap('', size=11, sa=30)
ap('融策·政府审计AI赋能方法论', size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
ap('投资人演讲脚本', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
ap('配套PPT：融策AI审计赋能方法论_投资人演示.pptx', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
ap('演讲时长：18-22分钟  |  2026年5月', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
ap('知悉范围：创始人+演讲人  |  机密', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
sep()

# Opening checklist
ap('▎ 开场前确认', size=13, bold=True, color=RED, sa=6)
for item in content['checklist']:
    ap(item, size=11, sa=3)
sep()

# Slides
for s in content['slides']:
    hdr(s['num'], s['title'], s.get('duration', ''))
    body(s['text'])
    sep()

# Appendix
doc.add_page_break()
ap('', size=11, sa=10)
ap('▎ 附录：常见投资人问题与应答要点', size=16, bold=True, color=RED, sa=16)
for faq in content['faq']:
    ap(faq['q'], size=13, bold=True, color=NAVY, sa=4)
    body(faq['a'])
    ap('', size=6, sa=4)
sep()

# Footer
ap('本脚本为融策核心机密文档，知悉范围：创始人+演讲人。', size=9, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
ap('严禁以任何形式复制、转发或对外披露。', size=9, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)

output_path = r'C:\Users\scrccpa\Desktop\融策AI审计赋能方法论_投资人演讲稿.docx'
doc.save(output_path)
print(f'Done: {output_path}')
