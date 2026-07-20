"""
绩效评价报告复核器 - 核心引擎
================================
提取报告问题 → RAG检索法规/案例 → 生成复核意见
用法:
    from scripts.report_reviewer import ReviewEngine
    engine = ReviewEngine()
    result = engine.review("path/to/report.docx")
"""
import os, sys, re, json, time
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.rag_vector import hybrid_search, rerank, EMBEDDING_MODEL, INDEX_DIR
from scripts.desensitize import sanitize_for_llm, restore_from_llm

WORKSPACE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ========== 报告解析 ==========
def parse_docx(filepath):
    """解析docx，提取结构化信息"""
    import docx
    doc = docx.Document(filepath)
    
    full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    
    # 提取标题
    title = ''
    for p in doc.paragraphs[:5]:
        if p.text.strip() and len(p.text.strip()) > 5:
            title = p.text.strip()
            break
    
    # 提取表格数据
    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    
    # 提取问题条目（审计报告常见问题模式）
    issues = extract_issues(full_text)
    
    # 提取法规引用
    regulations = extract_regulations(full_text)
    
    return {
        'title': title,
        'full_text': full_text,
        'issues': issues,
        'regulations': regulations,
        'tables_count': len(tables),
        'char_count': len(full_text),
    }

def extract_issues(text):
    """从报告文本提取审计发现问题"""
    issues = []
    
    # 模式1: "问题X：..." 或 "X.问题..."
    patterns = [
        r'[一二三四五六七八九十\d]+[、\.．]\s*(?:问题|事项|情况)[：:]\s*(.{10,200})',
        r'(?:发现|存在|查出)[的]?(.{10,200}?)[。\n]',
        r'(?:违规|违反|不符合|未按|超范围|超标准|虚列|挤占|挪用|截留)(.{10,200}?)[。\n]',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text)
        for m in matches:
            issue_text = m.group(1).strip()
            if len(issue_text) > 10 and issue_text not in [i['text'] for i in issues]:
                issues.append({
                    'text': issue_text,
                    'source': 'regex',
                    'type': classify_issue(issue_text),
                })
    
    # 去重（按前50字）
    seen = set()
    unique_issues = []
    for issue in issues:
        key = issue['text'][:50]
        if key not in seen:
            seen.add(key)
            unique_issues.append(issue)
    
    return unique_issues[:20]  # 最多20个问题

def classify_issue(text):
    """分类问题类型"""
    text_lower = text
    if any(k in text_lower for k in ['串标', '围标', '投标', '招标', '采购']):
        return '采购/招投标'
    elif any(k in text_lower for k in ['预算', '决算', '支出', '收入', '资金']):
        return '预算/资金'
    elif any(k in text_lower for k in ['工程', '项目', '建设', '施工']):
        return '工程项目'
    elif any(k in text_lower for k in ['资产', '固定资产', '存货', '盘点']):
        return '资产管理'
    elif any(k in text_lower for k in ['合同', '协议', '约定']):
        return '合同管理'
    elif any(k in text_lower for k in ['税', '费', '票据', '发票']):
        return '税费票据'
    elif any(k in text_lower for k in ['人员', '编制', '工资', '福利']):
        return '人员经费'
    else:
        return '其他'

def extract_regulations(text):
    """提取报告引用的法规"""
    regulations = []
    
    # 模式: 文号
    patterns = [
        r'([\u4e00-\u9fa5]{2,10}(?:发|字|函|文|号|办|令|公告|通知)\[[\d]{4}\][\d]+号?)',
        r'(财政部令[第\d]+号)',
        r'(审计署令[第\d]+号)',
        r'(《[^》]+》)',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text)
        for m in matches:
            reg = m.group(1).strip()
            if reg not in regulations:
                regulations.append(reg)
    
    return regulations[:20]

# ========== RAG检索 ==========
def search_regulation(issue_text, model=None, top_k=5):
    """为问题检索相关法规"""
    query = f"审计发现：{issue_text} 相关法规依据"
    result = hybrid_search(query, top_k, model=model)
    if result.get('results'):
        result['results'] = rerank(query, result['results'], top_k)
    return result.get('results', [])

def search_case(issue_text, model=None, top_k=5):
    """为问题检索同类案例"""
    query = f"审计案例：{issue_text} 处理处罚"
    result = hybrid_search(query, top_k, model=model)
    if result.get('results'):
        result['results'] = rerank(query, result['results'], top_k)
    return result.get('results', [])

# ========== 复核引擎 ==========
class ReviewEngine:
    def __init__(self):
        from sentence_transformers import SentenceTransformer
        print(f'加载模型: {EMBEDDING_MODEL}...')
        self.model = SentenceTransformer(EMBEDDING_MODEL)
        print('模型就绪')
    
    def review(self, filepath):
        """复核一份报告"""
        print(f'\n复核: {filepath}')
        
        # 1. 解析
        report = parse_docx(filepath)
        print(f'  标题: {report["title"]}')
        print(f'  问题: {len(report["issues"])} 个')
        print(f'  法规: {len(report["regulations"])} 条')
        print(f'  字数: {report["char_count"]}')
        
        # 2. 脱敏
        sanitized_text, mapping = sanitize_for_llm(report['full_text'][:3000])
        
        # 3. 逐问题复核
        review_results = []
        for i, issue in enumerate(report['issues'][:10]):  # 最多10个问题
            print(f'  [{i+1}/{min(10, len(report["issues"]))}] {issue["text"][:50]}...')
            
            # 检索法规
            regs = search_regulation(issue['text'], self.model, 3)
            # 检索案例
            cases = search_case(issue['text'], self.model, 3)
            
            review_results.append({
                'issue': issue,
                'regulations': [{
                    'score': r['score'],
                    'rerank_score': r.get('rerank_score', 0),
                    'source': r['source'],
                    'effective_date': r.get('effective_date', ''),
                    'text': r['text'][:150],
                } for r in regs],
                'cases': [{
                    'score': r['score'],
                    'rerank_score': r.get('rerank_score', 0),
                    'source': r['source'],
                    'text': r['text'][:150],
                } for r in cases],
            })
        
        # 4. 法规时效性检查
        expired_regs = []
        for reg_text in report['regulations']:
            # 检查法规是否在库中
            results = search_regulation(reg_text, self.model, 1)
            if results:
                doc = results[0]
                if doc.get('effective_date'):
                    year = doc['effective_date'][:4]
                    if year.isdigit() and int(year) < 2020:
                        expired_regs.append({
                            'regulation': reg_text,
                            'year': year,
                            'source': doc['source'],
                        })
        
        return {
            'title': report['title'],
            'char_count': report['char_count'],
            'issues_count': len(report['issues']),
            'regulations_count': len(report['regulations']),
            'reviews': review_results,
            'expired_regulations': expired_regs,
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
        }

# ========== CLI ==========
if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print('用法: python scripts/report_reviewer.py <报告.docx>')
        sys.exit(1)
    
    engine = ReviewEngine()
    result = engine.review(sys.argv[1])
    
    print(f'\n{"="*60}')
    print(f'复核报告: {result["title"]}')
    print(f'{"="*60}')
    print(f'问题数: {result["issues_count"]} | 法规数: {result["regulations_count"]}')
    
    if result['expired_regulations']:
        print(f'\n⚠️ 过期法规 ({len(result["expired_regulations"])} 条):')
        for r in result['expired_regulations']:
            print(f'  - {r["regulation"]} ({r["year"]})')
    
    print(f'\n逐问题复核意见:')
    for i, r in enumerate(result['reviews']):
        print(f'\n  [{i+1}] {r["issue"]["text"][:80]}')
        print(f'      类型: {r["issue"]["type"]}')
        if r['regulations']:
            print(f'      相关法规: {r["regulations"][0]["source"]} (score:{r["regulations"][0]["rerank_score"]})')
        if r['cases']:
            print(f'      同类案例: {r["cases"][0]["source"]}')
    
    # 保存结果
    out_path = os.path.join(WORKSPACE, 'output', 'review_result.json')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f'\n结果已保存: {out_path}')