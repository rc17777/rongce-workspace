# -*- coding: utf-8 -*-
"""插入最后3张图 + 再加一段文字，确保12000+字、50%视觉"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, copy

doc_path = r'D:\openclaw-workspace\output\模拟案例一_绩效目标编制辅导与审核工作方案.docx'
chart_dir = r'D:\openclaw-workspace\output\charts'
doc = Document(doc_path)

def bd(tmp, text):
    para = tmp.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run(text)
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    para.paragraph_format.line_spacing = 1.5

def add_img(tmp, filename, width=5.5, caption=''):
    path = os.path.join(chart_dir, filename)
    if not os.path.exists(path):
        print(f'  [MISS] {filename}'); return
    sp = tmp.add_paragraph()
    ip = tmp.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ip.add_run()
    r.add_picture(path, width=Inches(width))
    if caption:
        cp = tmp.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.size = Pt(10); cr.font.name = '楷体'
        cr._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        cr.font.color.rgb = RGBColor(0x66,0x66,0x66)
    tmp.add_paragraph()

def insert_after(doc, keyword, fn):
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            target_elem = para._element
            tmp = Document()
            fn(tmp)
            ins = target_elem
            for tmp_p in tmp.paragraphs:
                cloned = copy.deepcopy(tmp_p._element)
                ins.addnext(cloned)
                ins = cloned
            print(f'  OK {len(tmp.paragraphs)} elem @ ...{keyword[-25:]}')
            return
    print(f'  MISS: {keyword[:30]}')

# 1. 在 "方案二" 后插入 3C原则图
insert_after(doc, '方案二', lambda t: [
    add_img(t, '3c_principle.png', 5.0, '图14：指标值设定 3C 原则——Challenging + Comparable + Credible')
])

# 2. 在 before_after 图后插 side_by_side
insert_after(doc, '图11：指标修改前后对比', lambda t: [
    add_img(t, 'side_by_side.png', 6.0, '图15：修改前后指标结构对比（逐类展示，新增14项标绿）')
])

# 3. 在结语后插入绩效管理闭环图
def add_closing_extra(tmp):
    add_img(tmp, 'performance_cycle.png', 5.5, '图16：预算绩效管理闭环——本次辅导与审核服务的是闭环的第一步')
    bd(tmp, '展望后续工作，本次辅导与审核的成果不应止步于一份绩效目标定稿。建议区交通运输局以本次工作为起点，将绩效目标管理融入日常预算管理流程：每季度对照绩效目标开展一次执行进度自查（绩效监控），年中根据实际情况对绩效目标进行一次动态调整（如需），年度终了时以绩效目标为基准开展自评（绩效评价），将评价结果作为下一年度预算安排和改进管理的重要依据（结果应用），最终形成\"编制-监控-评价-反馈\"的完整绩效管理闭环。四川融策会计师事务所有限公司愿意在后续的绩效监控和绩效评价环节继续提供专业支持，助力区交通运输局建立长效化的预算绩效管理机制。')

insert_after(doc, '为区交通运输局的预算绩效管理贡献专业力量', add_closing_extra)

doc.save(doc_path)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
print(f'=== FINAL ===')
print(f'Chars: {total}  Images: {img_count}  Tables: {len(doc.tables)}')
print(f'Size: {os.path.getsize(doc_path)/1024:.1f} KB')
