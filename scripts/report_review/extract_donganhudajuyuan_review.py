import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import fitz
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

BASE = Path(r'C:\Users\scrccpa\Desktop\新建文件夹')
REPORT = BASE / '关于东安湖大剧院运营补贴专项审计报告.docx'
BASIS_DIR = BASE / '年末补贴申请资料' / '审核依据'
EVIDENCE_DIR = BASE / '年末补贴申请资料' / '佐证资料'
WORKPAPER_DIR = BASE / '底稿-龙泉驿区文广体旅局-东安湖大剧院运营补贴'
OUT = Path(r'C:\Users\scrccpa\.openclaw\workspace\outputs\donganhudajuyuan_review')
OUT.mkdir(parents=True, exist_ok=True)


def clean_text(text: str) -> str:
    text = text.replace('\x00', ' ')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_docx(path: Path):
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for ti, table in enumerate(doc.tables, 1):
        parts.append(f'\n[表格 {ti}]')
        for row in table.rows:
            vals = [cell.text.strip().replace('\n', ' / ') for cell in row.cells]
            if any(vals):
                parts.append(' | '.join(vals))
    return clean_text('\n'.join(parts))


def extract_pdf(path: Path, max_pages=None):
    doc = fitz.open(str(path))
    parts = []
    pages = len(doc)
    limit = pages if max_pages is None else min(pages, max_pages)
    for i in range(limit):
        txt = doc[i].get_text('text') or ''
        if txt.strip():
            parts.append(f'\n--- page {i+1} ---\n{txt}')
    text = clean_text('\n'.join(parts))
    return text, pages


def extract_doc(path: Path):
    # Legacy .doc extraction is environment-dependent. Keep metadata and flag for manual/Word conversion if needed.
    return ''


def iter_files(root: Path):
    if not root.exists():
        return []
    return sorted([p for p in root.rglob('*') if p.is_file()], key=lambda p: str(p))


def safe_name(path: Path):
    s = re.sub(r'[\\/:*?"<>|]+', '_', path.stem)
    return s[:80]


def process_file(path: Path, category: str):
    suffix = path.suffix.lower()
    rec = {
        'category': category,
        'path': str(path),
        'name': path.name,
        'suffix': suffix,
        'size': path.stat().st_size,
        'modified': datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec='seconds'),
        'text_chars': 0,
        'pages': None,
        'status': 'pending',
        'extract_path': None,
        'sample': ''
    }
    text = ''
    try:
        if suffix == '.docx':
            text = extract_docx(path)
            rec['status'] = 'text_extracted'
        elif suffix == '.pdf':
            text, pages = extract_pdf(path)
            rec['pages'] = pages
            rec['status'] = 'text_extracted' if text else 'no_embedded_text_or_scanned'
        elif suffix == '.doc':
            text = extract_doc(path)
            rec['status'] = 'legacy_doc_not_extracted'
        else:
            rec['status'] = 'unsupported_type'
    except Exception as e:
        rec['status'] = f'error: {type(e).__name__}: {e}'
    rec['text_chars'] = len(text)
    rec['sample'] = text[:1000]
    if text:
        out_path = OUT / f'{category}_{safe_name(path)}.txt'
        out_path.write_text(text, encoding='utf-8')
        rec['extract_path'] = str(out_path)
    return rec

records = []
records.append(process_file(REPORT, 'report'))
for p in iter_files(BASIS_DIR):
    records.append(process_file(p, 'basis'))
for p in iter_files(EVIDENCE_DIR):
    records.append(process_file(p, 'evidence'))
for p in iter_files(WORKPAPER_DIR):
    records.append(process_file(p, 'workpaper'))

(OUT / 'file_inventory.json').write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding='utf-8')

md = ['# 东安湖大剧院运营补贴专项审计复核抽取清单', '']
for r in records:
    md.append(f"## {r['category']} | {r['name']}")
    md.append(f"- 路径: `{r['path']}`")
    md.append(f"- 类型: `{r['suffix']}` | 大小: {r['size']} | 页数: {r['pages']} | 状态: {r['status']} | 文本字符: {r['text_chars']}")
    if r['extract_path']:
        md.append(f"- 抽取文本: `{r['extract_path']}`")
    if r['sample']:
        md.append('')
        md.append('```text')
        md.append(r['sample'][:1200])
        md.append('```')
    md.append('')
(OUT / 'extraction_summary.md').write_text('\n'.join(md), encoding='utf-8')
print(json.dumps({'out': str(OUT), 'records': len(records), 'report_chars': records[0]['text_chars']}, ensure_ascii=False, indent=2))
