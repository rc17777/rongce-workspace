from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ROOT = Path(r"C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722\成都市郫都区民政局2025年度部门预算项目")
REPORT = next(ROOT.glob("*.docx"))
WORKPAPER = next(ROOT.glob("*.xlsx"))
OUTPUT = ROOT.parent / "成都市郫都区民政局2025年度部门预算项目绩效自评复核结果.xlsx"


def value_or_blank(value):
    return "" if value is None else str(value).strip()


def font_info(run):
    font = run.font
    east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    return {
        "text": run.text,
        "font": font.name or east_asia or "",
        "east_asia": east_asia or "",
        "size_pt": round(font.size.pt, 1) if font.size else None,
        "bold": bool(font.bold),
    }


def paragraph_info(paragraph, index):
    ppr = paragraph._p.pPr
    spacing = paragraph.paragraph_format.line_spacing
    line_rule = ""
    if ppr is not None and ppr.spacing is not None:
        line_rule = ppr.spacing.get(qn("w:lineRule")) or ""
    align = {
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }.get(paragraph.alignment, "未显式设置")
    runs = [font_info(run) for run in paragraph.runs if run.text.strip()]
    return {
        "index": index,
        "text": paragraph.text.strip(),
        "style": paragraph.style.name if paragraph.style else "",
        "alignment": align,
        "line_spacing": str(spacing) if spacing is not None else "",
        "line_rule": line_rule,
        "runs": runs,
    }


def is_number_token(text):
    return bool(re.fullmatch(r"[0-9,.%+-]+", text.strip()))


def collect_docx():
    doc = Document(REPORT)
    paragraphs = [paragraph_info(p, i + 1) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    tables = []
    for ti, table in enumerate(doc.tables, 1):
        rows = []
        for ri, row in enumerate(table.rows, 1):
            rows.append([value_or_blank(cell.text) for cell in row.cells])
        tables.append({"table": ti, "rows": rows})
    sections = []
    for section in doc.sections:
        sections.append({
            "page_width_cm": round(section.page_width.cm, 2),
            "page_height_cm": round(section.page_height.cm, 2),
        })
    return doc, paragraphs, tables, sections


def collect_xlsx():
    workbook = load_workbook(WORKPAPER, data_only=False)
    data = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows():
            values = [cell.value for cell in row]
            if any(value not in (None, "") for value in values):
                rows.append([value_or_blank(value) for value in values])
        data.append({"sheet": sheet.title, "max_row": sheet.max_row, "max_col": sheet.max_column, "rows": rows})
    return data


def find_rows(table_rows, keywords):
    matches = []
    for row in table_rows:
        joined = " | ".join(row)
        if all(keyword in joined for keyword in keywords):
            matches.append(joined)
    return matches


def main():
    doc, paragraphs, tables, sections = collect_docx()
    sheets = collect_xlsx()

    all_report_text = "\n".join(p["text"] for p in paragraphs)
    all_table_text = "\n".join(" | ".join(row) for table in tables for row in table["rows"])
    all_workpaper_text = "\n".join(" | ".join(row) for sheet in sheets for row in sheet["rows"])

    findings = []
    def add(category, item, finding, standard, position, severity, recommendation, evidence):
        findings.append([category, item, finding, standard, position, severity, recommendation, evidence])

    # Format review: page, title hierarchy, paragraph fonts/line spacing, numeric fonts.
    for section in sections:
        if abs(section["page_width_cm"] - 21.0) <= 0.15 and abs(section["page_height_cm"] - 29.7) <= 0.15:
            add("报告格式", "A4纸张", "符合", "A4型", "页面设置", "通过", "无需修改", f"页面尺寸：{section['page_width_cm']}×{section['page_height_cm']}cm")
        else:
            add("报告格式", "A4纸张", "不符合", "A4型", "页面设置", "P1", "调整为A4纵向页面", f"页面尺寸：{section['page_width_cm']}×{section['page_height_cm']}cm")

    heading_pattern = re.compile(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）|\d+[.．])")
    for p in paragraphs:
        text = p["text"]
        if not heading_pattern.match(text):
            continue
        level = 1 if re.match(r"^[一二三四五六七八九十]+、", text) else 2 if re.match(r"^（[一二三四五六七八九十]+）", text) else 3
        expected_font = {1: "黑体", 2: "楷体GB2312", 3: "仿宋GB2312"}[level]
        if level > 3:
            add("报告格式", "标题层级", "标题超过三级", "最多三级：一、/（一）/1.", f"第{p['index']}段：{text[:50]}", "P1", "合并或下沉四级及以下标题", "标题编号识别结果")
        fonts = "/".join(sorted({(r["east_asia"] or r["font"]) for r in p["runs"] if (r["east_asia"] or r["font"])}))
        if p["alignment"] != "居中" and level == 1:
            add("报告格式", "一级标题对齐", "一级标题未居中", "报告文字部分标题应居中", f"第{p['index']}段：{text[:50]}", "P2", "将该标题设为居中", f"当前对齐：{p['alignment']}")
        if expected_font not in fonts:
            add("报告格式", f"{level}级标题字体", f"字体未见{expected_font}", f"{level}级标题应使用{expected_font}", f"第{p['index']}段：{text[:50]}", "P2", f"统一调整为{expected_font}" + ("并加粗" if level == 2 else ""), f"当前字体：{fonts or '未显式设置'}")

    title = paragraphs[0] if paragraphs else None
    if title:
        title_fonts = "/".join(sorted({(r["east_asia"] or r["font"]) for r in title["runs"] if (r["east_asia"] or r["font"])}))
        title_size = [r["size_pt"] for r in title["runs"] if r["size_pt"]]
        title_bold = all(r["bold"] for r in title["runs"]) if title["runs"] else False
        if title["alignment"] != "居中" or "小标宋" not in title_fonts or not any(abs(size - 22) < 1 for size in title_size) or not title_bold:
            add("报告格式", "主标题", "主标题格式未完全满足要求", "居中、2号小标宋、加粗", f"第{title['index']}段：{title['text'][:60]}", "P2", "统一设置为居中、2号小标宋、加粗", f"对齐={title['alignment']}；字体={title_fonts or '未显式'}；字号={title_size}；加粗={title_bold}")
        else:
            add("报告格式", "主标题", "符合", "居中、2号小标宋、加粗", f"第{title['index']}段", "通过", "无需修改", "样式核验通过")

    body_font_issues = 0
    line_issues = 0
    number_font_issues = 0
    for p in paragraphs:
        if p is title or heading_pattern.match(p["text"]):
            continue
        runs = p["runs"]
        chinese_runs = [r for r in runs if re.search(r"[\u4e00-\u9fff]", r["text"])]
        if chinese_runs:
            bad = [r for r in chinese_runs if "仿宋" not in (r["east_asia"] or r["font"])]
            size_bad = [r for r in chinese_runs if r["size_pt"] is not None and abs(r["size_pt"] - 16) > 1]
            if bad or size_bad:
                body_font_issues += 1
        if p["line_rule"] != "exact" or "381000" not in p["line_spacing"]:
            line_issues += 1
        for run in runs:
            if re.search(r"\d", run["text"]) and not is_number_token(run["text"]):
                continue
            if re.search(r"\d", run["text"]) and "Times New Roman" not in (run["font"] or ""):
                number_font_issues += 1
    if body_font_issues:
        add("报告格式", "正文中文字体/字号", f"发现{body_font_issues}个正文段落未明确符合三号仿宋GB2312", "正文三号仿宋GB2312", "正文段落样式", "P2", "统一应用正文样式，中文为三号仿宋GB2312", "以Word运行字体与字号属性自动核验")
    else:
        add("报告格式", "正文中文字体/字号", "未发现异常", "正文三号仿宋GB2312", "正文段落样式", "通过", "无需修改", "自动核验范围内通过")
    if line_issues:
        add("报告格式", "正文行距", f"发现{line_issues}个正文段落未明确为固定值30磅", "固定值30磅", "正文段落样式", "P2", "统一设置正文为固定值30磅", "以Word段落lineRule/line属性自动核验")
    if number_font_issues:
        add("报告格式", "文中数字字体", f"发现{number_font_issues}处数字运行字体未明确为Times New Roman", "数字使用Times New Roman", "正文及表格数字", "P2", "将数字统一替换为Times New Roman字体", "以Word运行字体属性自动核验")

    # Workpaper/report identity and content cross-check.
    report_name_tokens = re.findall(r"成都市郫都区民政局[^\n]{0,40}", all_report_text)
    paper_name_tokens = re.findall(r"成都市郫都区民政局[^\n]{0,40}", all_workpaper_text)
    if report_name_tokens and paper_name_tokens:
        add("一致性", "报告与底稿主体", "主体名称均出现成都市郫都区民政局", "报告与盖章底稿信息一致", "报告封面及底稿表头", "通过", "仍应人工核验盖章页、日期、版本", f"报告：{report_name_tokens[0][:70]}；底稿：{paper_name_tokens[0][:70]}")
    else:
        add("一致性", "报告与底稿主体", "自动提取未能同时识别主体名称", "报告与盖章底稿信息一致", "报告封面及底稿表头", "P1", "人工核验盖章底稿的单位、年度、项目范围、签章和日期", "无法仅凭文本提取确认盖章信息")

    # Identify score and evidence language. These are focused flags, not automatic deductions.
    zero_score_rows = [line for line in all_workpaper_text.splitlines() if re.search(r"(^|\|)\s*0(?:\.0+)?\s*(\||$)", line)]
    if zero_score_rows:
        add("自评复核", "0分项目依据", f"识别到底稿中{len(zero_score_rows)}行含0分，需要逐项核验是否对应“无资料/无关/满意度无资料/重复设置/定义不匹配”等适用口径", "无资料、指标无关、满意度无资料、重复或定义不匹配时为0分", "自评复核底稿评分表", "P1", "逐项在扣分原因中写明事实、适用口径和资料缺口；不可仅写“资料不足”", "自动扫描到0分行，详见“底稿摘录”工作表")
    else:
        add("自评复核", "0分项目依据", "未自动识别0分行", "适用时必须按明确口径给0分", "自评复核底稿评分表", "P2", "人工核对合并单元格和公式计算结果", "Excel可能使用公式或合并单元格")

    keywords = ["无资料", "无佐证", "未提供", "满意度", "重复", "无关", "定性", "完成值", "130%", "时效", "数量", "效益"]
    keyword_hits = {keyword: all_workpaper_text.count(keyword) for keyword in keywords}
    add("自评复核", "扣分原因描述", "已对底稿中关键口径词进行检索，需结合每条指标原文逐项判断", "扣分原因应对应用户明确的10项复核口径", "自评复核底稿评分表", "P2", "重点核对：无资料=0分；无关=0分；可考查性按100%/60%/30%；满意度无资料=0；完成值超130%不扣；重复指标仅保留一次分值", "关键词命中：" + "；".join(f"{k}{v}处" for k, v in keyword_hits.items()))

    # Internal report logic and evidence references.
    if "复核" not in all_report_text:
        add("报告内容", "复核结论表述", "报告正文未检索到“复核”关键词，需人工确认结论、范围和方法表述是否完整", "报告内容逻辑完整、结论有据", "报告正文", "P1", "补充或核实复核范围、资料依据、评分结果和总体结论", "全文关键词检索结果")
    else:
        add("报告内容", "复核结论表述", "报告正文含复核表述", "报告内容逻辑完整、结论有据", "报告正文", "通过", "仍需人工抽查关键结论对应底稿", "全文关键词检索通过")

    # Workbook output.
    wb = Workbook()
    ws = wb.active
    ws.title = "复核问题清单"
    headers = ["类别", "复核项", "复核发现", "适用标准/口径", "位置", "等级", "处理建议", "核验依据"]
    ws.append(["成都市郫都区民政局2025年度部门预算项目绩效自评复核结果"])
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws["A1"].font = Font(name="Microsoft YaHei", size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="0A1F3F")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28
    ws.append(headers)
    for cell in ws[2]:
        cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A5C6E")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in findings:
        ws.append(row)
    thin = Side(style="thin", color="B7C3CC")
    for row in ws.iter_rows(min_row=3, max_row=ws.max_row, min_col=1, max_col=len(headers)):
        severity = row[5].value
        fill = "FFFFFF" if ws.max_row % 2 else "F5F2EC"
        if severity == "P1": fill = "FCE4D6"
        elif severity == "P2": fill = "FFF2CC"
        elif severity == "通过": fill = "E2F0D9"
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.fill = PatternFill("solid", fgColor=fill)
    widths = [14, 20, 40, 38, 35, 10, 42, 52]
    for i, width in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = width
    ws.freeze_panes = "A3"
    ws.auto_filter.ref = f"A2:H{ws.max_row}"

    summary = wb.create_sheet("复核汇总")
    summary.append(["复核范围", "结论"])
    summary.append(["报告文件", REPORT.name])
    summary.append(["底稿文件", WORKPAPER.name])
    summary.append(["复核重点", "格式规范、报告与底稿一致性、内容逻辑、自评扣分口径"])
    counts = Counter(row[5] for row in findings)
    summary.append(["发现统计", "；".join(f"{key}：{value}" for key, value in counts.items())])
    summary.append(["重要说明", "报告与盖章底稿的签章、手写痕迹及扫描件真实性无法由文本自动提取确认，须人工终核。所有P1/P2项均需回到原件逐项确认。"])
    for cell in summary[1]:
        cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0A1F3F")
    summary.column_dimensions["A"].width = 22
    summary.column_dimensions["B"].width = 100
    for row in summary.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

    extract = wb.create_sheet("底稿摘录")
    extract.append(["来源", "Sheet/表号", "行号", "内容"])
    for cell in extract[1]:
        cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A5C6E")
    for sheet in sheets:
        for index, row in enumerate(sheet["rows"], 1):
            joined = " | ".join(row)
            if any(keyword in joined for keyword in keywords) or re.search(r"(^|\|)\s*0(?:\.0+)?\s*(\||$)", joined):
                extract.append(["Excel底稿", sheet["sheet"], index, joined])
    for table in tables:
        for index, row in enumerate(table["rows"], 1):
            joined = " | ".join(row)
            if any(keyword in joined for keyword in keywords):
                extract.append(["Word报告表格", f"表{table['table']}", index, joined])
    for width, col in zip([15, 22, 10, 120], range(1, 5)):
        extract.column_dimensions[get_column_letter(col)].width = width
    for row in extract.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    extract.freeze_panes = "A2"
    extract.auto_filter.ref = f"A1:D{extract.max_row}"

    wb.save(OUTPUT)
    payload = {"report": str(REPORT), "workpaper": str(WORKPAPER), "output": str(OUTPUT), "findings": findings, "paragraphs": paragraphs, "tables": tables, "sheets": sheets}
    debug = OUTPUT.with_suffix(".json")
    debug.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"output": str(OUTPUT), "debug": str(debug), "counts": counts, "finding_count": len(findings)}, ensure_ascii=False, default=str))

if __name__ == "__main__":
    main()
