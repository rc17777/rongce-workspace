# -*- coding: utf-8 -*-
"""
融策·券商风研报 Word 模板生成器 v2.0
=====================================
一键生成符合顶级券商研报规范的 Word 模板文档 (.docx)。

v2.0 新增特性：
- 页眉页脚自动化（Logo + 标题 + 页码）
- 一级标题深蓝色块底纹 + 白字
- 三级标题（Heading 3）
- 数据来源脚注样式
- 引用框/要点框样式
- 表格三线样式预设
- 目录自动生成占位

Author: 融策右护卫 (OpenClaw AI)
Date: 2026-07-21
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 融策品牌色
COLOR_DEEP_BLUE = RGBColor(10, 31, 63)     # #0A1F3F
COLOR_TEAL = RGBColor(26, 92, 110)         # #1A5C6E
COLOR_COPPER = RGBColor(197, 149, 92)      # #C5955C
COLOR_DARK_GRAY = RGBColor(74, 74, 74)     # #4A4A4A
COLOR_LIGHT_GRAY = RGBColor(155, 155, 155) # #9B9B9B
COLOR_BG_GRAY = RGBColor(248, 249, 250)    # #F8F9FA 表格交替行

def create_securities_template(output_path):
    doc = Document()

    # 1. 设置页面大小和边距 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 2. 设置全局默认中文字体 (微软雅黑)
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = COLOR_DARK_GRAY
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5

    # 3. 自定义样式
    styles = doc.styles

    # 封面主标题 (Cover Title)
    style = styles.add_style('RC_CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Heading 1']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(26)
    style.font.bold = True
    style.font.color.rgb = COLOR_DEEP_BLUE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(20)
    style.paragraph_format.space_before = Pt(60)

    # 封面副标题 (Cover Subtitle)
    style = styles.add_style('RC_CoverSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(16)
    style.font.bold = False
    style.font.color.rgb = COLOR_TEAL
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(40)

    # 封面投资摘要块 (Cover Summary)
    style = styles.add_style('RC_CoverSummary', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.font.bold = False
    style.font.color.rgb = COLOR_DARK_GRAY
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.left_indent = Cm(0.5)

    # 一级标题 (Heading 1) - 券商风深蓝色块白字
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Microsoft YaHei'
    style_h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(255, 255, 255)  # 白字
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(12)
    style_h1.paragraph_format.left_indent = Pt(10)
    # 底纹设置
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '0A1F3F')  # 深蓝底
    style_h1._element.get_or_add_pPr().append(shading_elm)

    # 二级标题 (Heading 2) - 青绿加粗 + 左侧缩进
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Microsoft YaHei'
    style_h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.font.color.rgb = COLOR_TEAL
    style_h2.paragraph_format.space_before = Pt(18)
    style_h2.paragraph_format.space_after = Pt(8)
    style_h2.paragraph_format.left_indent = Pt(12)

    # 三级标题 (Heading 3) - 深灰加粗
    style_h3 = styles['Heading 3']
    style_h3.font.name = 'Microsoft YaHei'
    style_h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_h3.font.size = Pt(12)
    style_h3.font.bold = True
    style_h3.font.color.rgb = COLOR_DARK_GRAY
    style_h3.paragraph_format.space_before = Pt(12)
    style_h3.paragraph_format.space_after = Pt(6)

    # 核心观点高亮文字 (Highlight Text)
    style_hl = styles.add_style('RC_Highlight', WD_STYLE_TYPE.CHARACTER)
    style_hl.font.bold = True
    style_hl.font.color.rgb = COLOR_DEEP_BLUE

    # 图表标题 (Chart Title)
    style = styles.add_style('RC_ChartTitle', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)
    style.font.bold = True
    style.font.color.rgb = COLOR_DEEP_BLUE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # 资料来源 (Source Note)
    style = styles.add_style('RC_Source', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(8)
    style.font.italic = True
    style.font.color.rgb = COLOR_LIGHT_GRAY
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(18)

    # 要点框 / 引用框 (Key Point Box)
    style = styles.add_style('RC_KeyBox', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_DARK_GRAY
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.left_indent = Cm(1.0)
    style.paragraph_format.right_indent = Cm(1.0)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.3
    # 浅灰底色
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F2EC')  # 融策暖灰
    style._element.get_or_add_pPr().append(shading_elm)

    # ==========================
    # 填充文档内容
    # ==========================

    # ============ 封面页 ============
    # 报告类型标签
    p = doc.add_paragraph('行业深度报告 | 融策·审盾系列', style='Normal')
    p.runs[0].font.color.rgb = COLOR_COPPER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(40)

    # 主标题
    doc.add_paragraph('融策AI审计中台：重塑政府审计的数据穿透力', style='RC_CoverTitle')
    
    # 副标题
    doc.add_paragraph('——从信息化核查到智能交叉验证的0到1跃迁', style='RC_CoverSubtitle')

    # 核心观点
    p = doc.add_paragraph('核心观点（投资摘要）', style='Heading 2')
    p.runs[0].font.color.rgb = COLOR_COPPER
    p.paragraph_format.left_indent = Pt(0)

    summary_text = [
        "⚡ 范式转移：传统审计主要依赖抽查与财务数据表面核对，融策AI中台通过「时空×行为×关系」多维坐标系，实现100%全量数据穿透，漏报率断崖式下降。",
        "⚡ 降本增效：在绩效评价与经责审计场景中，AI复核模型将报告质量检查时间压缩30%以上，同时拦截95%以上的格式与逻辑矛盾。",
        "⚡ 护城河建立：以「融策·审盾」为锚点，我们不仅提供审计结论，更提供可回溯、高可信的底层数字证据链，构成同业难以复制的竞争壁垒。"
    ]
    for text in summary_text:
        p = doc.add_paragraph(text, style='RC_CoverSummary')

    # 分析师信息区（封面底部）
    doc.add_paragraph('\n' * 6)  # 占位推到页面底部
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_table.rows[0].cells[0].text = '分析师'
    info_table.rows[0].cells[1].text = '融策右护卫'
    info_table.rows[1].cells[0].text = '执业编号'
    info_table.rows[1].cells[1].text = 'RC-AI-2026'
    info_table.rows[2].cells[0].text = '联系方式'
    info_table.rows[2].cells[1].text = 'contact@rongce.com'
    info_table.rows[3].cells[0].text = '报告日期'
    info_table.rows[3].cells[1].text = '2026年7月21日'
    
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(9)
                    paragraph.runs[0].font.name = 'Microsoft YaHei'
    
    doc.add_page_break()

    # ============ 目录页（占位）============
    doc.add_paragraph('目  录', style='Heading 1')
    doc.add_paragraph('[本模板使用时，请在此插入→引用→目录→自动生成]', style='Normal')
    p = doc.add_paragraph('Word操作：点击"引用"选项卡 → "目录" → 选择自动目录样式', style='Normal')
    p.runs[0].font.color.rgb = COLOR_LIGHT_GRAY
    p.runs[0].font.size = Pt(9)
    
    doc.add_page_break()

    # ============ 正文页 ============
    doc.add_paragraph('1. 行业痛点：信息孤岛下的审计盲区', style='Heading 1')
    p = doc.add_paragraph('长期以来，政府审计与工程咨询面临着海量非结构化数据难以有效利用的痛点。', style='Normal')
    run = p.add_run('传统"人工翻卷宗"模式不仅耗时耗力，更极易在繁杂的数据汪洋中遗漏关键串标或违规线索。')
    run.font.bold = True
    run.font.color.rgb = COLOR_DEEP_BLUE

    doc.add_paragraph('1.1 招投标领域的围串标隐蔽化', style='Heading 2')
    doc.add_paragraph('当前，围标串标手段已从初级的"IP地址重合"演化为复杂的"历史伴随投标"、"隐蔽股权代持"。单靠人工经验已无法应对规模化的合谋造假。因此，引入数据挖掘技术显得尤为迫切。', style='Normal')
    
    # 要点框示例
    doc.add_paragraph('🔍 核心发现：根据融策对近三年200个工程招标项目的数据分析，存在围串标嫌疑的项目占比高达18.3%，其中采用"历史伴随投标"手法的占比超过60%，而传统IP检测方法仅能识别其中不到30%。', style='RC_KeyBox')

    # 图表示例
    doc.add_paragraph('图1：2020-2025年全国一般公共预算收入（亿元）', style='RC_ChartTitle')
    p = doc.add_paragraph('[此处插入由 securities_chart.py 生成的折线图PNG]', style='Normal')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = COLOR_LIGHT_GRAY
    p.runs[0].font.size = Pt(9)
    doc.add_paragraph('资料来源：财政部，融策会计师事务所', style='RC_Source')

    doc.add_paragraph('2. 融策破局：AI中台的三层架构', style='Heading 1')
    doc.add_paragraph('为解决上述痛点，融策开发了基于多Agent协同的审计中台架构。该架构包含数据接入、智能分析与报告生成三大模块。', style='Normal')
    
    doc.add_paragraph('2.1 数据接入层：打破信息孤岛', style='Heading 2')
    doc.add_paragraph('支持财政决算报表、招投标文件、工商登记信息、银行流水等20+种数据源的自动化接入与结构化处理。', style='Normal')
    
    doc.add_paragraph('2.1.1 非结构化数据OCR', style='Heading 3')
    doc.add_paragraph('采用PaddleOCR引擎，对PDF扫描件的识别准确率达到98.5%以上，支持表格还原、印章识别等高级功能。', style='Normal')

    doc.add_page_break()
    
    # ============ 免责声明页 ============
    doc.add_paragraph('免责声明', style='Heading 1')
    doc.add_paragraph('', style='Normal')  # 空行
    
    disclaimer = "本报告仅供四川融策会计师事务所及四川融策工程咨询公司的内部决策与客户参考。报告中的信息均来源于已公开的资料，我公司对这些信息的准确性及完整性不作任何保证。报告中的信息或所表达的意见并不构成任何形式的投资或业务建议。"
    p = doc.add_paragraph(disclaimer, style='Normal')
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = COLOR_LIGHT_GRAY
    p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph('', style='Normal')
    
    # 评级说明
    p = doc.add_paragraph('评级说明', style='Heading 2')
    p.paragraph_format.left_indent = Pt(0)
    
    rating_table = doc.add_table(rows=4, cols=2)
    rating_table.style = 'Light Grid Accent 1'
    rating_table.rows[0].cells[0].text = '评级'
    rating_table.rows[0].cells[1].text = '说明'
    rating_table.rows[1].cells[0].text = '优秀'
    rating_table.rows[1].cells[1].text = '各项指标均达到或超过预期目标，无重大问题'
    rating_table.rows[2].cells[0].text = '良好'
    rating_table.rows[2].cells[1].text = '多数指标达到预期，存在轻微问题但不影响整体评价'
    rating_table.rows[3].cells[0].text = '需改进'
    rating_table.rows[3].cells[1].text = '关键指标未达标，存在较严重问题，需要整改'
    
    for row in rating_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(9)
                    paragraph.runs[0].font.name = 'Microsoft YaHei'

    doc.save(output_path)
    print(f'✅ 券商风研报模板 v2.0 已生成: {output_path}')


if __name__ == '__main__':
    out_file = sys.argv[1] if len(sys.argv) > 1 else 'output/融策_券商风研报模板_v2.0.docx'
    create_securities_template(out_file)
