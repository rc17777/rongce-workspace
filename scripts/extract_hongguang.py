import json, re
from collections import Counter
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import load_workbook

ROOT = Path(r"C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722\成都市郫都区人民政府红光街道办事处部门预算项目")
REPORT = next(ROOT.glob("*.docx"))
WORKPAPER = next(ROOT.glob("*.xlsx"))
OUT = ROOT.parent / "评审数据提取.json"

def value_or_blank(value):
    return "" if value is None else str(value).strip()

def font_info(run):
    font = run.font
    east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    return {"text": run.text, "font": font.name or "", "east_asia": east_asia or "", "size_pt": round(font.size.pt, 1) if font.size else None, "bold": bool(font.bold)}

def paragraph_info(paragraph, index):
    ppr = paragraph._p.pPr
    spacing = paragraph.paragraph_format.line_spacing
    line_rule = ""
    if ppr is not None and ppr.spacing is not None:
        line_rule = ppr.spacing.get(qn("w:lineRule")) or ""
    align_map = {WD_ALIGN_PARAGRAPH.CENTER: "居中", WD_ALIGN_PARAGRAPH.LEFT: "左对齐", WD_ALIGN_PARAGRAPH.RIGHT: "右对齐", WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐"}
    align = align_map.get(paragraph.alignment, "未显式设置")
    return {"index": index, "text": paragraph.text.strip(), "style": paragraph.style.name if paragraph.style else "", "alignment": align, "line_spacing": str(spacing) if spacing is not None else "", "line_rule": line_rule, "runs": [font_info(run) for run in paragraph.runs if run.text.strip()]}

def collect_docx():
    doc = Document(REPORT)
    paragraphs = [paragraph_info(p, i+1) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    tables = []
    for ti, table in enumerate(doc.tables, 1):
        rows = [[value_or_blank(cell.text) for cell in row.cells] for row in table.rows]
        tables.append({"table": ti, "rows": rows})
    sections = [{"page_width_cm": round(s.page_width.cm,2), "page_height_cm": round(s.page_height.cm,2)} for s in doc.sections]
    return doc, paragraphs, tables, sections

def collect_xlsx():
    wb = load_workbook(WORKPAPER, data_only=False)
    sheets = []
    for s in wb.worksheets:
        rows = [[value_or_blank(c.value) for c in row] for row in s.iter_rows()]
        rows = [r for r in rows if any(v not in ("","") for v in r)]
        sheets.append({"sheet": s.title, "max_row": s.max_row, "max_col": s.max_column, "rows": rows})
    return sheets

doc, paragraphs, tables, sections = collect_docx()
sheets = collect_xlsx()

payload = {"report": str(REPORT), "workpaper": str(WORKPAPER), "paragraphs": paragraphs, "tables": tables, "sections": sections, "sheets": sheets}
OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
print(f"OK: {len(paragraphs)} paragraphs, {len(tables)} tables, {len(sheets)} sheets")

# quick summary
for p in paragraphs:
    if any(w in p["text"] for w in ["一、","二、","三、","四、","五、","六、"]):
        fonts = "/".join(sorted({(r["east_asia"] or r["font"]) for r in p["runs"] if (r["east_asia"] or r["font"])}))
        print(f"  H1 [{p['index']}]: {p['text'][:60]}  align={p['alignment']}  font={fonts or '--'}")
for s in sheets:
    print(f"  Sheet [{s['sheet']}]: rows={len(s['rows'])}, max_col={s['max_col']}")
for t in tables:
    if t["table"] <= 3:
        print(f"  Table [{t['table']}]: rows={len(t['rows'])}")
        for r in t["rows"][:2]:
            print(f"    {' | '.join(r[:6])}")
