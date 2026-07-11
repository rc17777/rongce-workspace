#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""融策标书模板 v1.0 — 初稿"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ========== 配色系统 ==========
C_PRIMARY = RGBColor(0x00, 0x33, 0x66)       # 藏蓝 #003366
C_PRIMARY_LIGHT = RGBColor(0xCC, 0xD5, 0xE0) # 浅蓝灰
C_ACCENT = RGBColor(0xD4, 0x8B, 0x28)        # 金色 #D48B28
C_GRAY_BG = RGBColor(0xF2, 0xF2, 0xF2)       # 浅灰背景
C_BORDER = RGBColor(0xCC, 0xCC, 0xCC)         # 边框灰
C_BLACK = RGBColor(0x33, 0x33, 0x33)          # 正文深灰（比纯黑柔和）
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四
style.font.color.rgb = C_BLACK
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

# ——— 辅助函数 ———
def add_heading_styled(text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = C_PRIMARY
        run.font.name = '微软雅黑'
        rpr = run._element.get_or_add_rPr()
        rpr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, color=None, size=None, align=None, space_after=None, font_name=None):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    else: run.font.color.rgb = C_BLACK
    if size: run.font.size = Pt(size)
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    fn = font_name or '宋体'
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    return p

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=None, align=None, font_name=None):
    """设置单元格文字"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold: run.bold = True
    c = color or C_WHITE
    run.font.color.rgb = c
    run.font.size = Pt(size or 10)
    fn = font_name or '微软雅黑'
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if align: p.alignment = align
    else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

def make_style_table(headers, rows, col_widths=None):
    """创建标准化表格"""
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格边框 - 浅灰
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="#CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="#CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="#CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="#CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="#CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="#CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    # 设置列宽（如果指定）
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    
    # 表头行（深蓝底白字）
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '003366')
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=10)
    
    # 数据行（隔行灰底）
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')
            set_cell_text(cell, str(val), bold=False, color=C_BLACK, size=10, 
                         align=WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            # 第一列左对齐 + 加粗
            if c_idx == 0:
                set_cell_text(cell, str(val), bold=True, color=C_BLACK, size=10,
                             align=WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.add_paragraph()  # 表后空行
    return table

# ================================================================
# 模板内容
# ================================================================

# ===== 封面页 =====
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('四 川 融 策')
run.font.size = Pt(36)
run.font.color.rgb = C_PRIMARY
run.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('投标文件')
run.font.size = Pt(42)
run.font.color.rgb = C_PRIMARY
run.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# 分隔线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 30)
run.font.color.rgb = C_PRIMARY
run.font.size = Pt(12)

doc.add_paragraph()

# 项目信息
info_items = [
    '项 目 名 称：________________________',
    '招 标 编 号：________________________',
    '投 标 单 位：四川融策会计师事务所',
    '日      期：____年____月____日',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)
    run.font.color.rgb = C_BLACK
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.line_spacing = 2.0

# 分页
doc.add_page_break()

# ===== 目录页 =====
add_heading_styled('目  录', level=1)

toc_items = [
    ('一、法定代表人身份证明', '1'),
    ('二、授权委托书', '3'),
    ('三、投标函', '5'),
    ('四、报价一览表', '7'),
    ('五、公司简介及资质', '9'),
    ('六、项目实施方案', '13'),
    ('七、项目人员配置', '25'),
    ('八、类似业绩一览表', '30'),
    ('九、服务承诺', '35'),
    ('十、其他附件', '37'),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 页码右对齐
    run2 = p.add_run(f'{"." * 40}{page}')
    run2.font.size = Pt(12)
    run2.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.line_spacing = 2.0

doc.add_page_break()

# ===== 第一章：公司简介（示范页） =====
add_heading_styled('第一章  公司简介', level=1)

add_para('1.1 公司概况', bold=True, color=C_PRIMARY, size=14, space_after=6)
add_para(
    '四川融策会计师事务所（普通合伙）成立于20XX年，是经四川省财政厅批准设立的综合性会计师事务所。'
    '公司注册资本XXX万元，现有从业人员XX人，其中注册会计师XX人。'
    '公司业务涵盖政府审计、绩效评价、资产清查、专项债申报、监督检查、'
    '工程预算编制、财政评审、全过程工程咨询、工程结算审核等专业领域。',
    size=12, space_after=12
)

add_para('1.2 组织架构', bold=True, color=C_PRIMARY, size=14, space_after=6)
add_para('（此处插入组织架构图）', align=WD_ALIGN_PARAGRAPH.CENTER, color=C_PRIMARY_LIGHT, size=12, space_after=12)

add_para('1.3 资质荣誉', bold=True, color=C_PRIMARY, size=14, space_after=6)

# 示范表格1 — 资质列表
make_style_table(
    headers=['序号', '资质名称', '发证机关', '有效期'],
    rows=[
        ('1', '会计师事务所执业证书', '四川省财政厅', '20XX年-20XX年'),
        ('2', 'ISO9001质量管理体系认证', '中国质量认证中心', '20XX年-20XX年'),
        ('3', 'AAA级信用等级证书', 'XX信用评估机构', '20XX年-20XX年'),
        ('4', '工程咨询单位乙级资信', '四川省工程咨询协会', '20XX年-20XX年'),
    ]
)

doc.add_page_break()

# ===== 第二章：项目实施方案（示范页） =====
add_heading_styled('第二章  项目实施方案', level=1)

add_para('2.1 项目理解', bold=True, color=C_PRIMARY, size=14, space_after=6)
add_para(
    '本项目为XXX项目，主要涉及竣工财务决算审核工作。我公司将严格按照'
    '《基本建设项目竣工财务决算管理暂行办法》（财建〔2016〕503号）等政策法规要求，'
    '遵循独立、客观、公正的原则，提供高质量的审计服务。',
    size=12, space_after=12
)

add_para('2.2 工作流程', bold=True, color=C_PRIMARY, size=14, space_after=6)
add_para('（此处插入工作流程图）', align=WD_ALIGN_PARAGRAPH.CENTER, color=C_PRIMARY_LIGHT, size=12, space_after=12)

add_para('2.3 工作内容', bold=True, color=C_PRIMARY, size=14, space_after=6)

# 示范表格2 — 工作内容表
make_style_table(
    headers=['序号', '工作内容', '工作标准', '输出成果'],
    rows=[
        ('1', '资料收集与初步审核', '收集齐全、完整', '资料清单'),
        ('2', '现场盘点与核实', '账实相符', '盘点记录表'),
        ('3', '资金使用合规性审核', '专款专用', '资金流向表'),
        ('4', '二类费用计费复核', '依据国家计费标准', '复核意见表'),
        ('5', '竣工财务决算报告编制', '格式规范、数据准确', '决算审核报告'),
    ]
)

add_para('2.4 人员配置', bold=True, color=C_PRIMARY, size=14, space_after=6)

# 示范表格3 — 人员表
make_style_table(
    headers=['序号', '姓名', '职务/职称', '执业资格', '本项目职责'],
    rows=[
        ('1', '李开', '合伙人/总经理', '注册会计师', '项目总负责人'),
        ('2', 'XXX', '部门经理', '注册会计师', '现场负责人'),
        ('3', 'XXX', '项目经理', '中级会计师', '审计实施'),
        ('4', 'XXX', '审计助理', '初级会计师', '底稿编制'),
    ]
)

doc.add_page_break()

# ===== 第三章：业绩表（示范页） =====
add_heading_styled('第三章  类似业绩', level=1)
add_para('近三年同类项目业绩一览表：', size=12, space_after=6)

# 示范表格4 — 业绩表
make_style_table(
    headers=['序号', '项目名称', '委托单位', '合同金额\n（万元）', '完成\n时间', '服务范围'],
    rows=[
        ('1', '阿坝州S220线安羌镇至茸安乡段\n灾害修复整治工程竣工财务决算审核', '阿坝州财政局', 'XX', '2026', '竣工财务决算审核'),
        ('2', '阿坝州S452线垮沙乡至柯河乡段\n灾害修复整治工程竣工财务决算审核', '阿坝州财政局', 'XX', '2026', '竣工财务决算审核'),
        ('3', '甘孜州XX县20XX年度\n预算执行情况审计', '甘孜州审计局', 'XX', '2025', '预算执行审计'),
        ('4', 'XX市教育局20XX年度\n营养改善计划专项资金审计', 'XX市教育局', 'XX', '2025', '专项资金审计'),
    ]
)

add_para('注：以上为部分代表性业绩，可提供完整清单。', size=10, color=C_PRIMARY_LIGHT, space_after=12)

doc.add_page_break()

# ===== 附录：排版规范说明 =====
add_heading_styled('附录  融策标书排版规范', level=1)

norm_items = [
    ('配色规范', '主色：藏蓝 #003366 | 辅色：浅灰 #F5F5F5 | 点缀：金色 #D48B28'),
    ('字体规范', '标题：微软雅黑 | 正文：宋体 | 英文：Times New Roman'),
    ('字号规范', '一级标题：二号（22pt）| 二级标题：三号（16pt）| 正文：小四（12pt）| 表格：五号（10.5pt）'),
    ('页边距', '上2.5cm 下2.0cm 左2.8cm 右2.8cm'),
    ('行距规范', '正文：1.5倍行距 | 表格内：单倍行距'),
    ('表格规范', '表头深蓝#003366底白字 | 隔行灰底#F5F5F5 | 浅灰边框'),
    ('页码规范', '封面不编号，目录用罗马数字，正文从1开始'),
]

for title, content in norm_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'■ {title}：')
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = C_PRIMARY
    run1.font.name = '微软雅黑'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    run2 = p.add_run(content)
    run2.font.size = Pt(11)
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(8)

# 保存
output_path = r'D:\openclaw-workspace\bid_aba\融策标书模板_初稿.docx'
doc.save(output_path)
print(f"✅ 标书模板初稿已生成: {output_path}")
print(f"   包含页面：封面 → 目录 → 公司简介(含示范表格) → 实施方案(含示范表格) → 业绩表 → 排版规范")  
