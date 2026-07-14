#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""提取校服采购项目所有文件文本 - 使用glob自动发现"""
import pdfplumber
import os
import sys
import glob as glb
from pathlib import Path

# Fix Windows console encoding
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

BASE = r"C:\Users\scrccpa\Desktop\校服"
OUT = r"D:\openclaw-workspace\output\校服分析\txt"
os.makedirs(OUT, exist_ok=True)

# Find all PDFs
all_pdfs = list(Path(BASE).rglob("*.pdf"))
all_docx = list(Path(BASE).rglob("*.docx"))

print(f"Found {len(all_pdfs)} PDFs")
print(f"Found {len(all_docx)} DOCXs")

for p in all_pdfs[:5]:
    print(f"  PDF: {p.name}  ({p.stat().st_size/1024:.0f} KB)")
print("  ...")
for d in all_docx[:5]:
    print(f"  DOCX: {d.name}  ({d.stat().st_size/1024:.0f} KB)")
print("  ...")

def extract_pdf(pdf_path, out_name):
    out_path = os.path.join(OUT, out_name)
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            texts = []
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t:
                    texts.append(f"--- Page {i+1} ---\n{t}")
            full = "\n\n".join(texts)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(full)
            print(f"[OK] PDF: {out_name} ({len(full)} chars, {len(pdf.pages)} pages)")
            return len(full)
    except Exception as e:
        print(f"[FAIL] PDF {out_name}: {e}")
        return 0

def extract_docx(docx_path, out_name):
    out_path = os.path.join(OUT, out_name)
    try:
        from docx import Document
        doc = Document(str(docx_path))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for i, table in enumerate(doc.tables):
            texts.append(f"\n[Table {i+1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                texts.append(" | ".join(cells))
        full = "\n".join(texts)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(full)
        print(f"[OK] DOCX: {out_name} ({len(full)} chars)")
        return len(full)
    except Exception as e:
        print(f"[FAIL] DOCX {out_name}: {e}")
        return 0

# Step 1: Key procurement PDFs
print("\n=== Step 1: 关键招标采购文件 ===")
key_patterns = {
    "01-招标文件.txt": "*招标文件*9.25*.pdf",
    "02-评审结果.txt": "*评审结果*.pdf",
    "03-招标通知书.txt": "*招标通知书*.pdf",
    "04-采购前期纪要.txt": "*前期纪要*.pdf",
    "05-需求公告前公示.txt": "*需求公告前公示*.pdf",
    "06-需求统计.txt": "*需求统计*.pdf",
    "07-采购公告.txt": "*采购公告*.pdf",
    "08-选用方式.txt": "*选用方式*.pdf",
    "09-采购承诺书.txt": "*采购承诺书*.pdf",
    "10-签到表.txt": "*签到表*.pdf",
    "11-合同审计稿.txt": "*合同审计稿*.pdf",
    "12-合同.txt": "*校服采购项目合同*.pdf",
}

for out_name, pattern in key_patterns.items():
    matches = list(Path(BASE).rglob(pattern))
    if matches:
        # Pick the one not in the simpler "校服" mirror dir
        m = matches[0]
        for cand in matches:
            if "2025年校服采购" in str(cand):
                m = cand
                break
        extract_pdf(m, out_name)
    else:
        print(f"[SKIP] {out_name} - not found ({pattern})")

# Step 2: 资格投标文件 (docx only for now)
print("\n=== Step 2: 资格投标文件 ===")
qual_docx = [d for d in all_docx if "资格" in d.name and "投标" in d.name]
for d in qual_docx:
    # Figure out company name from path
    company = d.parent.name if d.parent.name else "unknown"
    extract_docx(d, f"{company}-资格标.txt")

# Step 3: 商务投标文件 (docx only)  
print("\n=== Step 3: 商务投标文件 ===")
biz_docx = [d for d in all_docx if "商务" in d.name and "投标" in d.name]
for d in biz_docx:
    company = d.parent.name if d.parent.name else "unknown"
    extract_docx(d, f"{company}-商务标.txt")

print("\nDone!")
