#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract pricing from 顺华 (pseudo-docx) and 弘博士 (OLE2 .doc)"""
import sys, os, re, struct
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from pathlib import Path
import olefile

BASE = Path(r"C:\Users\scrccpa\Desktop\校服\2025年校服采购\校服\2025年\投标文件\投标文件")

def extract_ole_text(filepath):
    """Extract raw text from OLE2 .doc file"""
    try:
        ole = olefile.OleFileIO(filepath)
        # The 'WordDocument' stream contains the main text
        if ole.exists('WordDocument'):
            word_stream = ole.openstream('WordDocument').read()
        
        # Try to get text from the 1Table or 0Table stream
        # For Word 97-2003, text is usually in the WordDocument stream
        # with FIB (File Information Block) at the start
        
        # Simple approach: extract all readable ASCII and try UTF-16LE
        # The text is stored in the WordDocument stream
        
        # Actually, let's try to use the simpler approach:
        # Extract the 'WordDocument' stream and look for text bytes
        
        # For binary .doc, the text might be stored as:
        # - ASCII (1 byte per char) in the WordDocument stream after the FIB
        # - Unicode (2 bytes per char) in a separate piece table
        
        # Let's try getting all text-like bytes
        all_text = b''
        for stream_name in ole.listdir():
            name = '/'.join(stream_name)
            try:
                data = ole.openstream(stream_name).read()
                # Try to extract Unicode text (UTF-16LE)
                try:
                    decoded = data.decode('utf-16-le', errors='ignore')
                    # Filter out garbage
                    cleaned = ''.join(c for c in decoded if c.isprintable() or c in '\n\r\t')
                    if len(cleaned) > 100:
                        all_text += cleaned.encode('utf-8')
                except:
                    pass
            except:
                pass
        
        ole.close()
        
        # Try the more reliable approach for doc files
        return extract_doc_text_fib(filepath)
    except Exception as e:
        print(f"  OLE extraction failed: {e}")
        return ""

def extract_doc_text_fib(filepath):
    """Extract text from .doc using FIB-based approach"""
    try:
        ole = olefile.OleFileIO(filepath)
        word_stream = ole.openstream('WordDocument').read()
        
        # Parse FIB to find text boundaries
        # FIB starts at offset 0
        # At offset 0x000A (10): flags, bit 0x0200 = fComplex
        # At offset 0x0024 (36): ccpText (count of chars in main text)
        
        flags = struct.unpack_from('<H', word_stream, 0x000A)[0]
        is_complex = bool(flags & 0x0200)
        
        # Magic number at 0x0000 should be 0xA5EC for Word
        magic = struct.unpack_from('<H', word_stream, 0x0000)[0]
        
        # Try getting ccpText (character count) from FIB
        # In complex format, it's at a different offset
        # Simple FIB: ccpText at 0x004C
        # But for simplicity, let's try a brute-force approach
        
        # Extract text by looking for character sequences
        text_parts = []
        
        # Search for UTF-16LE text sequences
        i = 0
        while i < len(word_stream) - 1:
            # Check if this is a printable Unicode char
            try:
                char = word_stream[i:i+2].decode('utf-16-le')
                if char.isprintable() or char in '\n\r\t ':
                    text_parts.append(char)
                    i += 2
                    continue
            except:
                pass
            i += 1
        
        ole.close()
        return ''.join(text_parts)
    except Exception as e:
        print(f"  FIB extraction failed: {e}")
        return ""

def extract_from_docx_pretending(path, label):
    """Extract pricing from a docx file (even if .doc extension)"""
    print(f"\n{'='*60}")
    print(f"📄 {label}")
    
    doc = Document(str(path))
    
    # Find pricing tables
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if any(c for c in cells):
                rows.append(' | '.join(cells))
        table_text = '\n'.join(rows)
        
        if any(kw in table_text for kw in ['开标一览表', '投标报价', '分项报价', '单价', '元/套']):
            print(f"\n[Table {ti+1}]")
            for row in rows:
                print(f"  {row[:200]}")
    
    # Also find pricing paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if re.search(r'[陆柒捌玖拾佰仟万亿零壹贰叁肆伍陆柒捌玖拾元角分整].*[元套]', t) or \
           re.search(r'投标报价.*\d+', t) or \
           re.search(r'[大小]写.*\d+.*元', t):
            print(f"\n[Para] {t[:200]}")

# === MAIN ===

# 顺华 - docx in disguise
for fname, label in [
    ("其他投标文件.doc", "顺华-商务标"),
    ("资格性投标文件.doc", "顺华-资格标"),
]:
    path = BASE / "成都顺华服装有限公司" / fname
    if path.exists():
        extract_from_docx_pretending(path, label)

# 弘博士 - true OLE2 .doc
# Try using the smaller 资格 file first to test extraction
print(f"\n{'='*60}")
print("📄 弘博士 - 尝试OLE2提取")
path = BASE / "弘博士服饰集团有限公司" / "资格性投标文件（电子文件）-弘博士服饰集团有限公司.doc"
text = extract_ole_text(path)
if text:
    # Look for pricing
    price_lines = []
    for line in text.split('\n'):
        line = line.strip()
        if any(kw in line for kw in ['报价', '价格', '元', '投标', '开标', '一览']):
            price_lines.append(line)
    if price_lines:
        print("Found pricing lines:")
        for l in price_lines[:20]:
            print(f"  {l[:200]}")
    else:
        print(f"Extracted {len(text)} chars but no pricing found")
        print("First 500 chars:")
        for c in text[:500]:
            if c.isprintable():
                print(c, end='')
else:
    print("No text extracted from OLE2")
    print("Will try markitdown as fallback...")

print("\nDone!")
