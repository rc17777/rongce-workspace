"""
S220/S452 竣工财务决算审核 - 数据勾稽复核
复核维度：
  1. 报告↔附表勾稽：报告中的金额与附表是否一致
  2. 报告内部勾稽：报告正文之间数据是否一致（概算、送审、审定、核减等）
输出：Excel复核对账表
"""
import sys, os, re, json
sys.stdout.reconfigure(encoding='utf-8')

BASE = r'C:\Users\scrccpa\Desktop\新建文件夹'
PROJECTS = {
    'S220': os.path.join(BASE, 'S220安羌镇至茸安乡段灾害修复整治工程'),
    'S452': os.path.join(BASE, 'S452垮沙乡至柯河乡段灾害修复整治工程'),
}

def extract_xls_text(xls_path):
    """从xls提取所有sheet数据（用python）"""
    import subprocess
    result = subprocess.run(
        [sys.executable, '-c', f'''
import xlrd
book = xlrd.open_workbook(r"{xls_path}", formatting_info=False)
for sn in book.sheet_names():
    sh = book.sheet_by_name(sn)
    print(f"===== SHEET: {{sn}} =====")
    for r in range(sh.nrows):
        row = []
        for c in range(sh.ncols):
            cell = sh.cell(r,c)
            v = cell.value
            if cell.ctype == 2:  # number
                row.append(str(v))
            elif v:
                row.append(str(v).strip())
            else:
                row.append('')
        print('|'.join(row))
        '''],
        capture_output=True, text=True, timeout=30
    )
    return result.stdout

def extract_docx_text(docx_path):
    """从docx提取所有文本"""
    from docx import Document
    doc = Document(docx_path)
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())
    # Also extract tables
    for i, table in enumerate(doc.tables):
        text.append(f"【表格{i+1}】")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            text.append('|'.join(cells))
    return '\n'.join(text)

def parse_amounts(text):
    """从文本中提取所有金额数值（万元）"""
    amounts = {}
    # 找常见的审计金额字段
    patterns = {
        '概算金额': r'概算[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '送审金额': r'送审[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '审定金额': r'审定[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '核减金额': r'核减[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '审核金额': r'审核[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '报审金额': r'报审[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '送审总价': r'送审(?:总价|总额)[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '审定总价': r'审定(?:总价|总额)[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '合同金额': r'合同[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '已支付': r'已支付[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '未支付': r'未支付[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '本次支付': r'本次(?:申请)?支付[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '建筑安装工程': r'建筑安装工程[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '待摊投资': r'待摊投资[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '项目总投资': r'项目(?:总)?投资[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
        '工程费用': r'工程费用[^0-9]*?(\d[\d,.]*\d*万?[元]?)',
    }
    
    for key, pat in patterns.items():
        matches = re.findall(pat, text, re.IGNORECASE)
        if matches:
            amounts[key] = matches
    
    return amounts

def normalize_amount(s):
    """规范化金额为数字（万元）"""
    if isinstance(s, (int, float)):
        return float(s)
    s = s.replace(',', '').replace(' ', '').replace('，', '').strip()
    if s.endswith('万元') or s.endswith('万'):
        s = s.replace('万元', '').replace('万', '')
        return float(s)
    if s.endswith('元'):
        s = s.replace('元', '')
        return float(s) / 10000
    try:
        return float(s)
    except:
        return None

def clean_num(s):
    """清理数字字符串"""
    return s.replace(',', '').replace(' ', '').replace('，', '').strip()

print("="*60)
print("  融策审计复核：报告↔附表数据勾稽对账")
print("="*60)

all_results = []

for proj_id, proj_dir in PROJECTS.items():
    print(f"\n{'='*60}")
    print(f"  {proj_id}: {os.path.basename(proj_dir)}")
    print(f"{'='*60}")
    
    # Find files
    files = os.listdir(proj_dir)
    report_file = None
    summary_note_file = None
    table1_file = None
    table2_file = None
    
    for f in files:
        fp = os.path.join(proj_dir, f)
        if '审核报告' in f or '审核表' in f:
            # Check type
            if f.endswith('.docx'):
                report_file = fp if '审核报告' in f else report_file
                summary_note_file = fp if '说明' in f else summary_note_file
            elif f.endswith('.xls'):
                if '汇总表' in f or '表1' in f or '附表1' in f:
                    table1_file = fp
                elif '明细表' in f or '表2' in f or '附表2' in f:
                    table2_file = fp
    
    # Assign by pattern matching
    for f in files:
        fp = os.path.join(proj_dir, f)
        if f.endswith('.docx'):
            if '审核报告' in f and '说明' not in f:
                report_file = fp
            elif '说明' in f:
                summary_note_file = fp
        elif f.endswith('.xls'):
            if '汇总表' in f or '附表1' in f:
                table1_file = fp
            elif '明细表' in f or '附表2' in f:
                table2_file = fp
    
    print(f"\n  文件清单：")
    print(f"  审核报告: {os.path.basename(report_file) if report_file else '未找到'}")
    print(f"  编制说明: {os.path.basename(summary_note_file) if summary_note_file else '未找到'}")
    print(f"  附表1(汇总): {os.path.basename(table1_file) if table1_file else '未找到'}")
    print(f"  附表2(明细): {os.path.basename(table2_file) if table2_file else '未找到'}")
    
    # ---- Extract data ----
    report_text = ''
    table1_text = ''
    table2_text = ''
    
    if report_file:
        report_text = extract_docx_text(report_file)
        print(f"\n  报告正文长度: {len(report_text)}字符")
    
    if table1_file:
        table1_text = extract_xls_text(table1_file)
    
    if table2_file:
        table2_text = extract_xls_text(table2_file)
    
    # ---- Parse amounts ----
    report_amt = parse_amounts(report_text)
    table1_amt = parse_amounts(table1_text)
    table2_amt = parse_amounts(table2_text)
    
    # ---- 复核1: 报告↔附表勾稽 ----
    print(f"\n  【复核1】报告↔附表勾稽核对")
    print(f"  {'-'*50}")
    
    # Find key numbers in report
    print(f"\n  报告中的金额：")
    for key, vals in report_amt.items():
        for v in vals:
            a = normalize_amount(v)
            if a:
                print(f"    {key}: {v} = {a:.4f}万元")
    
    # Find key numbers in table1
    print(f"\n  附表1(汇总表)中的金额：")
    for key, vals in table1_amt.items():
        for v in vals:
            a = normalize_amount(v)
            if a:
                print(f"    {key}: {v} = {a:.4f}万元")
    
    # Find key numbers in table2
    print(f"\n  附表2(明细表)中的金额：")
    for key, vals in table2_amt.items():
        for v in vals:
            a = normalize_amount(v)
            if a:
                print(f"    {key}: {v} = {a:.4f}万元")
    
    # ---- 复核2: 报告内部数据勾稽 ----
    print(f"\n  【复核2】报告内部数据勾稽")
    print(f"  {'-'*50}")
    
    # 检查概算 = 送审 关系
    if '概算金额' in report_amt and '送审金额' in report_amt:
        gs = normalize_amount(report_amt['概算金额'][0])
        ss = normalize_amount(report_amt['送审金额'][0])
        if gs and ss:
            diff = abs(gs - ss)
            print(f"  概算金额: {gs:.4f}万元")
            print(f"  送审金额: {ss:.4f}万元")
            if diff < 0.01:
                print(f"  ✅ 概算=送审，一致")
            else:
                print(f"  ⚠️ 差异: {diff:.4f}万元")
    
    # 检查送审 = 审定 + 核减
    if '送审金额' in report_amt and '审定金额' in report_amt and '核减金额' in report_amt:
        ss = normalize_amount(report_amt['送审金额'][0])
        ds = normalize_amount(report_amt['审定金额'][0])
        hj = normalize_amount(report_amt['核减金额'][0])
        if ss and ds and hj:
            calc = ds + hj
            diff = abs(ss - calc)
            print(f"  送审({ss:.4f}) = 审定({ds:.4f}) + 核减({hj:.4f}) = {calc:.4f}")
            if diff < 0.01:
                print(f"  ✅ 勾稽一致")
            else:
                print(f"  ⚠️ 差异: {diff:.4f}万元")

print("\n\n复核完成！")
