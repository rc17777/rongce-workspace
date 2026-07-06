from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Flowable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT_DIR = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
DOCX_PATH = OUT_DIR / "四川融策宣传册_完善稿.docx"
PDF_PATH = OUT_DIR / "四川融策宣传册_完善稿.pdf"

BLUE = "0A1F3F"
TEAL = "1A5C6E"
GOLD = "C5955C"
WARM = "F5F2EC"
DARK = "20252B"

PAGES = [
    {
        "title": "四川融策",
        "subtitle": "政府审计与工程咨询综合服务机构",
        "tagline": "谋专业之策，融品质之精",
        "body": [
            "以财政资金、公共资源、工程建设和预算绩效为核心服务场景，提供审计、咨询、评价、监督检查和数字化分析一体化服务。",
            "我们关注的不只是出具报告，而是帮助委托方识别风险、厘清责任、优化管理、形成可落地的改进方案。",
        ],
        "points": ["财政监督检查", "预算绩效管理", "经济责任审计", "工程造价咨询", "数字化审计分析"],
    },
    {
        "title": "公司定位",
        "subtitle": "把传统审计经验转化为可复用的治理能力",
        "body": [
            "四川融策会计师事务所有限公司成立于2000年，是四川省内较早成立的会计师事务所之一。公司总部位于成都，并设有阿坝州、西藏等服务支点。",
            "融策长期服务财政、审计、教育、民政、农业农村、交通、国资、医疗保障等公共部门和企事业单位，逐步形成了政府审计、预算绩效、财政监督、工程咨询协同发展的业务格局。",
        ],
        "points": ["20余年专业积累", "400余家客户服务经验", "50人左右专业团队", "会计、造价、绩效、工程复合能力"],
    },
    {
        "title": "核心优势",
        "subtitle": "专业判断、现场经验、数据能力三位一体",
        "body": [
            "融策的优势不只体现在人员资质，更体现在对政府业务逻辑、财政资金运行规律、项目建设流程和基层执行难点的理解。",
            "我们将制度审查、数据核验、访谈核查、现场踏勘、证据闭环和整改建议贯穿项目全过程，确保发现问题有依据、评价结论可解释、整改建议能执行。",
        ],
        "points": ["注册会计师、造价师、高级会计师等复合团队", "覆盖财政资金、工程项目、政府采购、国资管理等场景", "长期积累项目案例库、问题库和分析模型", "质量控制前置，重大结论复核到数据来源和计算依据"],
    },
    {
        "title": "服务体系",
        "subtitle": "围绕财政管理和公共治理形成六类服务",
        "body": [
            "融策服务体系从单项审计延伸到全过程咨询，从结果评价延伸到制度优化，从人工核查延伸到数据化识别。",
        ],
        "points": [
            "政府审计：经济责任、财务收支、专项资金、资产清查、国企审计、财政监督检查",
            "预算绩效：事前绩效评估、绩效目标审核、运行监控、重点绩效评价、结果应用辅导",
            "工程咨询：预算编制、财政评审、工程结算、竣工决算财务审计、全过程工程咨询",
            "采购与招投标审计：采购程序合规、围标串标线索、合同履约、资金支付核查",
            "管理咨询：内控制度、资产管理、资金管理、整改提升、流程优化",
            "数字化分析：多源数据整理、异常识别、疑点清单、取证路径和审计底稿自动化支持",
        ],
    },
    {
        "title": "重点能力一：政府审计",
        "subtitle": "从真实性、合规性延伸到绩效性和治理性",
        "body": [
            "面向财政、审计、主管部门及国有企事业单位，融策提供经济责任审计、财务收支审计、专项资金审计、资产清查、财政监督检查等服务。",
            "我们的工作重点是把账、表、合同、项目、资金流和管理制度连起来，形成问题事实、责任边界、影响程度和整改路径的完整闭环。",
        ],
        "points": ["经济责任审计：重大决策、资金资产、项目建设、内部控制、廉政风险", "专项资金审计：申报、分配、拨付、使用、绩效、结余沉淀全链条核查", "财政监督检查：预算执行、财经纪律、政府采购、资产管理、会计信息质量", "资产清查：账实核对、权属核验、损失认定、管理建议和制度完善"],
    },
    {
        "title": "重点能力二：全过程预算绩效管理",
        "subtitle": "让财政资金从“花了没有”走向“花得值不值”",
        "body": [
            "融策围绕预算编制、执行监控、绩效评价、结果应用四个环节，为财政部门和预算单位提供全过程预算绩效管理服务。",
            "我们强调指标体系与项目实际匹配，强调评价证据与结论对应，强调问题建议能够反馈到预算安排、项目管理和政策优化。",
        ],
        "points": ["事前绩效评估：必要性、可行性、财政承受能力、预期绩效", "绩效目标审核：目标完整性、指标可衡量性、预算匹配性", "绩效运行监控：执行进度、资金支付、产出偏差、风险预警", "重点绩效评价：政策、部门整体、项目支出、专项资金评价", "结果应用：整改清单、预算挂钩建议、管理制度优化"],
    },
    {
        "title": "重点能力三：工程咨询与财政评审",
        "subtitle": "把工程语言、财务语言和财政管理要求对齐",
        "body": [
            "工程咨询业务覆盖预算编制、招标控制价、工程量清单、财政评审、结算审核、竣工决算财务审计和全过程工程咨询。",
            "融策注重工程造价、合同管理、资金支付、变更签证、项目绩效之间的交叉核验，帮助委托方控制投资风险、规范项目管理。",
        ],
        "points": ["预算编制与财政评审：工程量、定额套用、材料价格、措施费、取费标准", "结算审核：合同条款、变更签证、隐蔽工程、现场核量、支付资料", "全过程咨询：项目前期、招采、实施、验收、结算、绩效评价协同管理", "竣工决算财务审计：建设成本归集、资金来源、资产交付、尾工尾款"],
    },
    {
        "title": "数字化改革服务能力",
        "subtitle": "用数据提高审计覆盖面、发现率和证据质量",
        "body": [
            "面对财政资金规模大、项目类型多、资料分散、现场时间有限等现实问题，融策正在将审计经验沉淀为数据标准、识别规则、疑点模型和作业流程。",
            "通过数据采集清单、字段标准化、异常规则库、穿透核查路径和工作底稿模板，提升项目启动、疑点筛选、现场核查、报告复核的效率。",
        ],
        "points": ["数据标准：财务、预算、支付、合同、采购、资产、工程项目等字段清单", "识别模型：重复支付、超预算执行、供应商异常、项目进度异常、资金沉淀", "证据闭环：疑点来源、核查过程、佐证材料、影响金额、整改建议", "质量复核：金额汇总、口径一致性、附表与正文一致、结论依据可追溯"],
    },
    {
        "title": "代表服务经验",
        "subtitle": "长期服务省、市、县多级财政和主管部门",
        "body": [
            "融策服务对象覆盖四川省财政厅、省级主管部门、市县财政部门、国有企事业单位及工程建设单位，项目类型包括预算绩效管理、财政监督检查、专项资金评价、工程决算和管理咨询等。",
            "原宣传册列示的代表客户包括四川省财政厅、四川省公安厅交通警察总队、四川省民政厅、四川省农业农村厅、四川省退役军人事务厅、四川省教育厅、四川省药品监督管理局，以及达州、绵阳、宜宾、什邡、德阳、康定、九寨沟等地财政部门。",
        ],
        "points": ["省级部门绩效评价与预算绩效管理服务", "市县财政重点项目、政策和部门整体支出绩效评价", "交通、教育、民政、农业农村、医保等重点民生领域审计评价", "工程项目竣工决算、财政评审和结算审核服务"],
    },
    {
        "title": "合作价值",
        "subtitle": "交付一份报告，更交付一套能落地的改进方案",
        "body": [
            "融策坚持客观公正、实事求是、质量优先。我们希望每一个项目都能沉淀为委托方可继续使用的管理工具：一套风险清单、一组指标体系、一批可核查的数据规则、一份明确责任和时限的整改建议。",
            "面向数字化改革和财政管理现代化，融策将持续提升数据分析、行业研究、报告复核和现场核查能力，为政府部门和企事业单位提供更稳健、更深入、更高效率的专业服务。",
        ],
        "points": ["咨询电话：028-87659276", "企业邮箱：scrccpa@163.com", "公司网址：www.scrccpa.com", "四川融策会计师事务所有限公司"],
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run(run, size=11, bold=False, color=DARK, font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_page(doc: Document, page: dict, first: bool = False) -> None:
    if not first:
        doc.add_section(WD_SECTION.NEW_PAGE)
    sec = doc.sections[-1]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(page["title"])
    set_run(r, size=26, bold=True, color=BLUE)
    p = doc.add_paragraph()
    r = p.add_run(page["subtitle"])
    set_run(r, size=13, bold=True, color=GOLD)

    if "tagline" in page:
        p = doc.add_paragraph()
        r = p.add_run(page["tagline"])
        set_run(r, size=18, bold=True, color=TEAL)

    for text in page["body"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run(r, size=11, color=DARK, font="宋体")

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    hdr = table.cell(0, 0)
    set_cell_shading(hdr, BLUE)
    p = hdr.paragraphs[0]
    r = p.add_run("核心内容")
    set_run(r, size=11, bold=True, color="FFFFFF")
    for item in page["points"]:
        row = table.add_row().cells[0]
        p = row.paragraphs[0]
        r = p.add_run("• " + item)
        set_run(r, size=10.5, color=DARK, font="宋体")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = foot.add_run("公开公正  用心服务  诚信为本  服务至上  追求卓越")
    set_run(r, size=8, color=GOLD)


def generate_docx() -> None:
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿"
    for idx, page in enumerate(PAGES):
        add_docx_page(doc, page, first=idx == 0)
    doc.save(DOCX_PATH)


class ColorBar(Flowable):
    def __init__(self, width, height, color):
        super().__init__()
        self.width = width
        self.height = height
        self.color = color

    def draw(self):
        self.canv.setFillColor(self.color)
        self.canv.rect(0, 0, self.width, self.height, stroke=0, fill=1)


def register_fonts() -> tuple[str, str]:
    candidates = [
        (r"C:\Windows\Fonts\msyh.ttc", "MSYH"),
        (r"C:\Windows\Fonts\simhei.ttf", "SIMHEI"),
        (r"C:\Windows\Fonts\simsun.ttc", "SIMSUN"),
    ]
    registered = []
    for path, name in candidates:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont(name, path))
            registered.append(name)
    if not registered:
        return "Helvetica", "Helvetica-Bold"
    return registered[0], registered[0]


def generate_pdf() -> None:
    font, bold_font = register_fonts()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.4 * cm,
    )
    width, _ = A4
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "RongceTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=27,
        leading=34,
        textColor=colors.HexColor("#" + BLUE),
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    subtitle = ParagraphStyle(
        "RongceSubtitle",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#" + GOLD),
        spaceAfter=16,
    )
    body = ParagraphStyle(
        "RongceBody",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.5,
        leading=17,
        textColor=colors.HexColor("#" + DARK),
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    small = ParagraphStyle(
        "RongceSmall",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor("#" + DARK),
    )
    footer = ParagraphStyle(
        "Footer",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#" + GOLD),
        alignment=TA_CENTER,
    )

    story = []
    usable_width = width - 3.4 * cm
    for idx, page in enumerate(PAGES):
        story.append(ColorBar(usable_width, 0.22 * cm, colors.HexColor("#" + GOLD)))
        story.append(Spacer(1, 0.45 * cm))
        story.append(Paragraph(page["title"], title))
        story.append(Paragraph(page["subtitle"], subtitle))
        if "tagline" in page:
            story.append(Paragraph(page["tagline"], ParagraphStyle("Tag", parent=subtitle, fontSize=17, textColor=colors.HexColor("#" + TEAL))))
        for text in page["body"]:
            story.append(Paragraph(text, body))
        data = [[Paragraph("核心内容", ParagraphStyle("Head", parent=small, fontName=bold_font, textColor=colors.white))]]
        data.extend([[Paragraph("• " + item, small)] for item in page["points"]])
        table = Table(data, colWidths=[usable_width])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BLUE)),
            ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#" + WARM)),
            ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#" + GOLD)),
            ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D8C7A9")),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story.append(Spacer(1, 0.25 * cm))
        story.append(table)
        story.append(Spacer(1, 0.35 * cm))
        story.append(Paragraph("公开公正  用心服务  诚信为本  服务至上  追求卓越", footer))
        if idx != len(PAGES) - 1:
            story.append(PageBreak())

    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_docx()
    generate_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
