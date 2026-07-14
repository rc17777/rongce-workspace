#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成融策AI赋能投标章节"""
import sys,os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
DARK = RGBColor(0x0A, 0x1F, 0x3F)
GOLD = RGBColor(0xC5, 0x95, 0x5C)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# ====== 标题 ======
h = doc.add_heading('AI赋能：融策智能审计平台', level=1)
for run in h.runs:
    run.font.color.rgb = DARK

doc.add_paragraph(
    '四川融策会计师事务所自主研发AI审计技术体系，将人工智能深度嵌入政府审计全流程，'
    '实现审计效率与质量的系统性提升。'
)

# ===== 一、技术架构 =====
doc.add_heading('一、融策AI审计技术架构', level=2)
doc.add_paragraph(
    '基于大语言模型（DeepSeek/Llama）+ 多Agent协作框架 + RAG专业知识库的技术底座，'
    '构建覆盖"数据采集→智能分析→风险识别→报告生成→质量复核"的全链条AI审计流水线。'
)

# 技术指标表
t = doc.add_table(5, 3, style='Table Grid')
t.cell(0,0).text = '技术组件'
t.cell(0,1).text = '能力指标'
t.cell(0,2).text = '应用效果'
data = [
    ('RAG审计知识库', '1,235份法规案例·13,635个知识块', '专业知识实时检索，准确率>95%'),
    ('多Agent协同平台', '7个专业AI Agent并行作业', '审计流程自动化率>60%'),
    ('智能分析引擎', '11层围标检测·15维报告复核', '风险识别覆盖率提升300%'),
    ('量化评价体系', '7+1套经责审计指标体系', '评价结论标准化、可追溯'),
]
for i, (a,b,c) in enumerate(data):
    t.cell(i+1,0).text = a
    t.cell(i+1,1).text = b
    t.cell(i+1,2).text = c

for j in range(3):
    cell = t.cell(0,j)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell._element.get_or_add_tcPr().append(
        cell._element.makeelement(qn('w:shd'), {qn('w:fill'): '0A1F3F', qn('w:val'): 'clear'})
    )

# ===== 二、核心产品 =====
doc.add_heading('二、核心AI审计产品', level=2)

doc.add_heading('产品一：智能围标串标检测系统', level=3)
for item in [
    '11层递进式检测体系：报价规律→文本雷同→图片哈希→元数据交叉→打印机型号→文档结构→格式重合率→工商关联→保证金资金链→投标行为异常→意思联络证据',
    '核心策略：无需代理机构配合数据（L3+L4+L5三杀即可定案），破解"代理不给IP"的行业顽疾',
    '应用场景：政府采购招投标审计、工程招标监督、供应商合规审查',
    '典型案例：某县中小学项目招标审计，6家公司元数据指纹一致，文本重合率100%→锁定围标',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('产品二：经责审计量化评价系统', level=3)
for item in [
    '7+1套指标体系：覆盖工程项目/行政事业/商业竞争/公益功能/特定功能/金融/科创+自定义',
    '6种一票否决机制 · 四类调整因素 · 三级损失认定标准',
    '评价过程可追溯、可复核、可审计——告别"拍脑袋打分"',
    '整合模板v2.0已落地，配套Excel量化评价工作表',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('产品三：审计报告AI复核系统', level=3)
for item in [
    '15维度递进式检查：10维正文复核 + 5维三方交叉比对',
    '致命层（提交前必做）：报告↔附表全量勾稽、取证单→报告完整闭环、全链路金额追踪',
    'Output分级：P0重大矛盾 / P1重大遗漏 / P2口径差异',
    '实测效果：食在攒劲餐饮公司审计报告复核，检测P1问题8处、P2问题4处，14项数字勾稽全验证',
]:
    doc.add_paragraph(item, style='List Bullet')

# ===== 三、赋能效果 =====
doc.add_heading('三、AI赋能效果对比', level=2)
t2 = doc.add_table(6, 3, style='Table Grid')
for j, h in enumerate(['审计环节', '传统方式', 'AI赋能方式']):
    t2.cell(0,j).text = h
    for p in t2.cell(0,j).paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    t2.cell(0,j)._element.get_or_add_tcPr().append(
        t2.cell(0,j)._element.makeelement(qn('w:shd'), {qn('w:fill'): '0A1F3F', qn('w:val'): 'clear'})
    )

for i, (a,b,c) in enumerate([
    ('数据采集与清洗', '人工翻阅凭证·手工录入', 'AI自动OCR识别·智能结构化·异常自动标记'),
    ('风险识别', '凭经验逐条筛查', '11层自动检测·全量数据比对·零遗漏'),
    ('勾稽验证', '人工逐项核对', 'AI自动跨表勾稽·即时定位差异'),
    ('报告生成', '逐字撰写·反复修改', 'AI辅助撰写·模板智能填充·一键排版'),
    ('质量复核', '人工抽查·主观判断', '15维自动复核·量化评分·全链路追踪'),
]):
    t2.cell(i+1,0).text = a
    t2.cell(i+1,1).text = b
    t2.cell(i+1,2).text = c

# ===== 四 =====
doc.add_heading('四、为什么选择融策AI审计', level=2)
for a in [
    '✅ 不是概念，是成品：以上三个AI产品均有实际项目验证，不是PPT产品',
    '✅ 自主知识产权：全栈自研，非第三方套壳，数据安全可控',
    '✅ 审计专业基因：AI+CPA双重加持——懂AI的团队很多，懂审计的AI团队极少',
    '✅ 持续迭代：知识库每日更新，模型定期升级，能力持续进化',
    '✅ 一人公司·AI倍增：融策正在注册AI OPC（一人有限责任公司），以"1个CPA + AI = 5人团队产能"为核心理念——同样的预算，选融策得到的是"人+AI"双重服务',
]:
    doc.add_paragraph(a, style='List Bullet')

out = r'D:\openclaw-workspace\bid_aba\AI赋能章节-投标方案.docx'
doc.save(out)
print(f'Done: {out} ({os.path.getsize(out)}B)')
