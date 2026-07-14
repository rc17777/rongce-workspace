# -*- coding: utf-8 -*-
"""四川融策宣传册 v9 —— 可编辑 Word 原生版
使用 python-docx 原生排版，文字可直接编辑，表格和段落可自由调整。
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

OUT = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
DOCX = OUT / "四川融策宣传册_v9_可编辑Word版.docx"

# Colors
C_NW = RGBColor(0x1A, 0x36, 0x5D)   # navy
C_TL = RGBColor(0x3A, 0x7B, 0x8A)   # teal
C_GD = RGBColor(0xD4, 0xA5, 0x74)   # gold
C_MU = RGBColor(0x71, 0x80, 0x96)   # muted
C_IK = RGBColor(0x2D, 0x37, 0x48)   # dark
C_WH = RGBColor(0xFF, 0xFF, 0xFF)   # white
C_BG = RGBColor(0xFA, 0xF8, 0xF4)   # warm white
C_LG = RGBColor(0xE8, 0xE0, 0xD4)   # light gold

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, spec in kwargs.items():
        color, sz, val = spec
        el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:color="{color}" w:space="0"/>')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph(cell_or_doc, text, font_name="微软雅黑", font_size=11, bold=False, color=C_IK, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    """Add a styled paragraph to a cell or document"""
    if hasattr(cell_or_doc, 'add_paragraph'):
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p

def new_section(doc, margins=(1.5, 2.0, 1.5, 2.0)):
    """Add new section with margins (top, bottom, left, right in cm)"""
    sec = doc.add_section()
    sec.top_margin = Cm(margins[0])
    sec.bottom_margin = Cm(margins[1])
    sec.left_margin = Cm(margins[2])
    sec.right_margin = Cm(margins[3])
    # Set page background
    return sec

def colored_table(doc, rows, cols, col_widths=None):
    """Create a borderless table for layout"""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    # Remove all borders by default
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            for edge in ['top','left','bottom','right']:
                el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="none" w:sz="0" w:color="auto" w:space="0"/>')
                tcBorders.append(el)
            tcPr.append(tcBorders)
    return table

# ===================== BUILD DOCUMENT =====================
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===================== PAGE 1: COVER =====================
sec = doc.sections[0]
sec.top_margin = Cm(0)
sec.bottom_margin = Cm(0)
sec.left_margin = Cm(0)
sec.right_margin = Cm(0)
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)

# Full navy background via table
t1 = colored_table(doc, 1, 1, [21.0])
cell = t1.rows[0].cells[0]
set_cell_shading(cell, "1A365D")

# Inner table for layout
add_paragraph(cell, "SICHUAN", font_size=36, bold=True, color=RGBColor(0xFF,0xFF,0xFF), space_before=80, space_after=0)
p = cell.paragraphs[-1]
p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for r in p.runs:
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

add_paragraph(cell, "RONGCE", font_size=56, bold=True, color=C_WH, space_before=0, space_after=10)
add_paragraph(cell, "", font_size=2, space_before=0, space_after=0)
add_paragraph(cell, "谋专业之策  融品质之精", font_size=32, bold=True, color=C_WH, space_before=60, space_after=6)
add_paragraph(cell, "政府审计与工程咨询综合服务机构", font_size=18, color=RGBColor(0xCC,0xCC,0xCC), space_before=6, space_after=20)

# Gold line
p_line = add_paragraph(cell, "━" * 40, font_size=14, color=C_GD, space_before=20, space_after=10)

add_paragraph(cell, "", font_size=2, space_before=0, space_after=0)
add_paragraph(cell, "审计  ·  绩效  ·  财政监督  ·  工程咨询  ·  数字化分析", font_size=14, color=RGBColor(0xAA,0xAA,0xCC), space_before=120, space_after=6)
add_paragraph(cell, "四川融策会计师事务所有限公司", font_size=12, color=RGBColor(0x88,0x88,0xAA), space_before=40, space_after=0)

# ===================== PAGE 2: WHY RONGCE =====================
new_section(doc, (2.5, 2.0, 2.5, 2.5))
add_paragraph(doc, "为什么是融策？", font_size=28, bold=True, color=C_NW, space_before=0, space_after=2)
add_paragraph(doc, "WHY RONGCE", font_size=13, color=C_GD, space_before=0, space_after=6)
add_paragraph(doc, "━" * 30, font_size=8, color=C_GD, space_before=0, space_after=12)

add_paragraph(doc, "财政资金的管理，核心是两个问题：", font_size=18, color=C_MU, space_before=0, space_after=2)
add_paragraph(doc, "\u201c资金到底去哪了？\u201d 和 \u201c这钱花得值不值？\u201d", font_size=20, bold=True, color=C_NW, space_before=0, space_after=4)
add_paragraph(doc, "融策做的事情，就是用专业的方法和数据工具，帮您把这两个问题搞清楚。", font_size=14, color=C_MU, space_before=4, space_after=12)

# Three cards
items = [
    ("\u2460", "不只查账", "合同 · 项目 · 资产\n内控制度 · 决策程序\n账表只是入口"),
    ("\u2461", "不只找问题", "能不能改 · 怎么改\n改得怎么样\n整改落地才是目标"),
    ("\u2462", "不只靠经验", "数据分析 + 现场核查\n用数据扩大覆盖面\n用现场核实关键点"),
]
t2 = colored_table(doc, 2, 3, [5.3, 5.3, 5.3])
for i, (icon, head, desc) in enumerate(items):
    cell_h = t2.rows[0].cells[i]
    cell_h.merge(t2.rows[1].cells[i])  # Actually let's merge properly
    
# Simpler approach: three separate tables side by side
# Use a single row 3-col table with merged rows
# Remove the table and use simpler layout

# Actually, let me just restart with a cleaner approach
# This is getting too complex with tables. Let me use a simpler structure.

# Delete the messed-up table and add text directly
# (python-docx limitation: can't easily delete)

# Let me just add the cards as separate paragraphs with shading
for i, (icon, head, desc) in enumerate(items):
    add_paragraph(doc, f"{icon}  {head}", font_size=18, bold=True, color=C_NW, space_before=10, space_after=4)
    for line in desc.split("\n"):
        add_paragraph(doc, f"    {line}", font_size=12, color=C_MU, space_before=0, space_after=0)
    add_paragraph(doc, "", font_size=2, space_before=6, space_after=0)

# Company brief
add_paragraph(doc, "", font_size=6, space_before=12, space_after=0)
add_paragraph(doc, "┈" * 50, font_size=8, color=C_GD, space_before=0, space_after=6)
for item in [
    "始于 2000 年，四川省内较早成立的会计师事务所之一",
    "长期服务财政、审计、教育、民政、交通、国资、医保等领域",
    "审计 + 绩效 + 财政监督 + 工程咨询协同发展",
    "覆盖四川、西藏、贵州三省"
]:
    add_paragraph(doc, f"  \u25cf  {item}", font_size=12, color=C_IK, space_before=1, space_after=1)

add_paragraph(doc, "", font_size=6, space_before=20, space_after=0)
add_paragraph(doc, "公开公正  用心服务  诚信为本  服务至上  追求卓越", font_size=10, color=C_MU, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=0)

# This approach is messy. Let me redo this properly from scratch with a clean, consistent approach.
# The issue is that python-docx layout is fundamentally different from Pillow image drawing.
# Let me write a clean version that focuses on readability and editability.

print("Part 1 done - continuing with a cleaner approach...")
DOCX.write_text("placeholder")  # Reset

# Actually, let me just write a clean, simple document. The user wants editable text, not pixel-perfect design.
# Let me focus on: clean structure, good typography, proper sections, and readability.

print("Clean approach needed - see v9_clean.py")
