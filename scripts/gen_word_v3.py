#!/usr/bin/env python3
"""Generate Word documents from all 37 policy .md files."""
import os, re, glob
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx.oxml

SRC_DIR = r"C:\Users\scrccpa\.openclaw\workspace\output\新制度体系"
OUT_DIR = r"C:\Users\scrccpa\Desktop"

# Category to Word file mapping
CATEGORIES = {
    "人力资源篇": {
        "pattern": ["01-薪酬管理制度.md", "02-绩效考核管理制度.md", "03-员工手册.md",
                    "10-招聘与入职管理制度.md", "11-培训与发展管理制度.md",
                    "12-职级晋升管理制度.md", "36-高管经营目标责任与绩效考核办法.md",
                    "09-股东会议事规则.md"],
        "output": "融策制度汇编-人力资源篇.docx"
    },
    "财务管理篇": {
        "pattern": ["06-财务报销管理制度.md", "13-项目收入确认与回款管理制度.md",
                    "14-预算管理制度.md", "15-资金管理制度.md",
                    "16-固定资产管理制度.md", "37-可分配利润核算细则.md"],
        "output": "融策制度汇编-财务管理篇.docx"
    },
    "业务质控篇": {
        "pattern": ["04-项目管理规范.md", "17-业务承接与合同管理制度.md",
                    "18-客户关系管理制度.md", "19-业务分包管理制度.md",
                    "20-投标管理制度.md", "34-项目独立核算与分润制度.md",
                    "35-跨部门协同与交叉营销奖励办法.md",
                    "07-审计质量控制制度.md", "08-造价咨询质量控制制度.md",
                    "21-三级复核实施细则.md", "22-执业责任追究制度.md"],
        "output": "融策制度汇编-业务质控篇.docx"
    },
    "行政综合篇": {
        "pattern": ["05-制度发布与版本管理规范.md", "23-信息安全与保密管理制度.md",
                    "24-印章与证照管理制度.md", "25-档案管理制度.md",
                    "26-公司章程-会计师事务所.md", "27-公司章程-工程咨询公司.md",
                    "28-办公场所管理制度.md", "29-采购管理制度.md",
                    "30-数智化建设管理制度.md", "31-业务拓展与创新管理制度.md",
                    "32-风险管理制度.md", "33-党建工作制度.md"],
        "output": "融策制度汇编-行政综合篇.docx"
    }
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color)
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')

def set_font(run, name='宋体', size=Pt(11), bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def parse_md_to_docx(doc, md_content, title):
    """Parse markdown content into docx paragraphs and tables."""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip frontmatter / header metadata lines
        if stripped.startswith('> **制度编号**'):
            # Metadata block - extract and format
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped.lstrip('> '))
            set_font(run, '微软雅黑', Pt(9), color=RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue
        elif stripped.startswith('> **版本**') or stripped.startswith('> **编制日期**') or \
             stripped.startswith('> **适用范围**') or stripped.startswith('> **优先级**') or \
             stripped.startswith('> **替代说明**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped.lstrip('> '))
            set_font(run, '微软雅黑', Pt(9), color=RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue
        elif stripped.startswith('> **审核人**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped.lstrip('> '))
            set_font(run, '微软雅黑', Pt(9), color=RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue
        
        # Code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            set_font(run, 'Consolas', Pt(9), color=RGBColor(0x33, 0x33, 0x33))
            i += 1
            continue
        
        # Horizontal rule
        if stripped == '---' or stripped == '***' or stripped == '___':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add bottom border
            pPr = p._element.get_or_add_pPr()
            pBdr = docx.oxml.OxmlElement('w:pBdr')
            bottom = docx.oxml.OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            set_font(run, '微软雅黑', Pt(15), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            set_font(run, '微软雅黑', Pt(13), bold=True, color=RGBColor(0x1A, 0x5C, 0x6E))
            i += 1
            continue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[4:])
            set_font(run, '微软雅黑', Pt(12), bold=True)
            i += 1
            continue
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[5:])
            set_font(run, '微软雅黑', Pt(11), bold=True)
            i += 1
            continue
        
        # Empty line
        if not stripped:
            if in_table and table_rows:
                # End of table
                rows_data = [r for r in table_rows if r and r[0] != '---']
                if len(rows_data) >= 1:
                    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                    table.style = 'Table Grid'
                    for ri, row in enumerate(rows_data):
                        for ci, cell_text in enumerate(row):
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            run = p.add_run(cell_text.strip())
                            set_font(run, '微软雅黑', Pt(9))
                            if ci < len(row) - 1:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            if ri == 0 and rows_data[0][0] != '---':
                                set_cell_shading(cell, '0A1F3F')
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.bold = True
                doc.add_paragraph()  # spacing
                table_rows = []
                in_table = False
            i += 1
            continue
        
        # Table rows
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not all(c == '---' or c == ':---' or c == '---:' or c == ':---:' or re.match(r'^-{2,}$', c) for c in cells if c.strip()):
                table_rows.append(cells)
                in_table = True
            i += 1
            continue
        
        # Ordered list
        if re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            # Parse bold within list items
            text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), stripped)
            run = p.add_run(text)
            set_font(run, '宋体', Pt(11))
            i += 1
            continue
        
        # Unordered list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), stripped[2:])
            run = p.add_run('• ' + text)
            set_font(run, '宋体', Pt(11))
            i += 1
            continue
        
        # Normal paragraph - handle inline formatting
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Simple bold handling
        parts = re.split(r'(\*\*.+?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, '宋体', Pt(11), bold=True)
            else:
                run = p.add_run(part)
                set_font(run, '宋体', Pt(11))
        i += 1
    
    # Handle trailing table
    if table_rows:
        rows_data = [r for r in table_rows if r and r[0] != '---']
        if len(rows_data) >= 1:
            table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.style = 'Table Grid'
            for ri, row in enumerate(rows_data):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(cell_text.strip())
                    set_font(run, '微软雅黑', Pt(9))
                    if ri == 0:
                        set_cell_shading(cell, '0A1F3F')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True

def add_cover_page(doc, category_name):
    """Add a simple cover page for a category volume."""
    # Empty lines for spacing
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策')
    set_font(run, '微软雅黑', Pt(28), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('管理制度汇编')
    set_font(run, '微软雅黑', Pt(24), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
    
    doc.add_paragraph()
    
    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 18)
    set_font(run, '微软雅黑', Pt(14), color=RGBColor(0x1A, 0x5C, 0x6E))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'— {category_name} —')
    set_font(run, '微软雅黑', Pt(18), bold=True, color=RGBColor(0x1A, 0x5C, 0x6E))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 18)
    set_font(run, '微软雅黑', Pt(14), color=RGBColor(0x1A, 0x5C, 0x6E))
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策会计师事务所有限公司')
    set_font(run, '微软雅黑', Pt(12))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策工程咨询有限公司')
    set_font(run, '微软雅黑', Pt(12))
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('发布日期：2026年7月22日')
    set_font(run, '微软雅黑', Pt(10), color=RGBColor(0x99, 0x99, 0x99))
    
    doc.add_page_break()

def add_toc(doc, entries):
    """Add table of contents."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('目    录')
    set_font(run, '微软雅黑', Pt(16), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for filename, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{title}')
        set_font(run, '微软雅黑', Pt(11))
    
    doc.add_page_break()

def set_page_margins(doc):
    """Set standard page margins."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

def generate_volume(category_name, file_list, output_name):
    """Generate one categorized Word volume."""
    doc = Document()
    set_page_margins(doc)
    
    add_cover_page(doc, category_name)
    
    # Read files and build TOC
    entries = []
    for fname in file_list:
        path = os.path.join(SRC_DIR, fname)
        if not os.path.exists(path):
            print(f'  WARNING: {fname} not found, skipping')
            continue
        with open(path, encoding='utf-8') as f:
            first_line = f.readline().strip().lstrip('# ')
            if not first_line:
                first_line = fname
            entries.append((fname, first_line))
    
    add_toc(doc, entries)
    
    for fname, title in entries:
        path = os.path.join(SRC_DIR, fname)
        with open(path, encoding='utf-8') as f:
            content = f.read()
        
        # Section divider
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        set_font(run, '微软雅黑', Pt(18), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Horizontal line
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = docx.oxml.OxmlElement('w:pBdr')
        bottom = docx.oxml.OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '8')
        bottom.set(qn('w:color'), '1A5C6E')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        parse_md_to_docx(doc, content, title)
        doc.add_page_break()
    
    out_path = os.path.join(OUT_DIR, output_name)
    doc.save(out_path)
    print(f'  → {output_name} ({len(entries)}份制度)')
    return len(entries)

def generate_full_volume():
    """Generate the complete single-volume edition."""
    doc = Document()
    set_page_margins(doc)
    
    # Cover
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策')
    set_font(run, '微软雅黑', Pt(32), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('公司制度体系')
    set_font(run, '微软雅黑', Pt(28), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 20)
    set_font(run, '微软雅黑', Pt(14), color=RGBColor(0x1A, 0x5C, 0x6E))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（完整版）')
    set_font(run, '微软雅黑', Pt(18), bold=True, color=RGBColor(0x1A, 0x5C, 0x6E))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 20)
    set_font(run, '微软雅黑', Pt(14), color=RGBColor(0x1A, 0x5C, 0x6E))
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('共37项制度 · 七大类')
    set_font(run, '微软雅黑', Pt(12), color=RGBColor(0x66, 0x66, 0x66))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026年7月22日发布')
    set_font(run, '微软雅黑', Pt(10), color=RGBColor(0x99, 0x99, 0x99))
    
    doc.add_page_break()
    
    # Full TOC
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('目    录')
    set_font(run, '微软雅黑', Pt(18), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    all_files = sorted([f for f in os.listdir(SRC_DIR) if f.endswith('.md') and f != '00-制度体系架构.md'])
    total = 0
    for category_name, info in CATEGORIES.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(category_name)
        set_font(run, '微软雅黑', Pt(13), bold=True, color=RGBColor(0x1A, 0x5C, 0x6E))
        
        for fname in info['pattern']:
            path = os.path.join(SRC_DIR, fname)
            if os.path.exists(path):
                with open(path, encoding='utf-8') as f:
                    title = f.readline().strip().lstrip('# ')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(f'  {title}')
                set_font(run, '微软雅黑', Pt(10))
                total += 1
    
    doc.add_page_break()
    
    # Content
    for category_name, info in CATEGORIES.items():
        # Category header page
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(category_name)
        set_font(run, '微软雅黑', Pt(20), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 20)
        set_font(run, '微软雅黑', Pt(12), color=RGBColor(0x1A, 0x5C, 0x6E))
        
        doc.add_page_break()
        
        for fname in info['pattern']:
            path = os.path.join(SRC_DIR, fname)
            if not os.path.exists(path):
                continue
            with open(path, encoding='utf-8') as f:
                content = f.read()
            
            first_line = content.split('\n')[0].strip().lstrip('# ')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(first_line)
            set_font(run, '微软雅黑', Pt(16), bold=True, color=RGBColor(0x0A, 0x1F, 0x3F))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            parse_md_to_docx(doc, content, first_line)
            doc.add_page_break()
    
    out_path = os.path.join(OUT_DIR, '融策公司制度体系（完整版）.docx')
    doc.save(out_path)
    print(f'  → 融策公司制度体系（完整版）.docx ({total}份制度)')

if __name__ == '__main__':
    print('开始生成制度Word文档...\n')
    
    total_count = 0
    for cat_name, info in CATEGORIES.items():
        existing = [f for f in info['pattern'] if os.path.exists(os.path.join(SRC_DIR, f))]
        print(f'[{cat_name}] {len(existing)}份文件')
        count = generate_volume(cat_name, existing, info['output'])
        total_count += count
    
    print(f'\n分册合计: {total_count}份')
    print('\n开始生成完整版...')
    generate_full_volume()
    
    print(f'\n全部完成！桌面文件列表:')
    for f in sorted(glob.glob(os.path.join(OUT_DIR, '融策制度*'))):
        size_kb = os.path.getsize(f) / 1024
        print(f'  {os.path.basename(f)} ({size_kb:.0f} KB)')
