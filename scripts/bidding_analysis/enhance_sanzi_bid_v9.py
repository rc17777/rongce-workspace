from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook

SRC = Path(r"D:\openclaw-workspace\projects\西昌三资三化投标\西昌三资三化投标服务方案-完整版-V8f.docx")
XLSX = Path(r"C:\Users\scrccpa\Desktop\三资三化课件\三资三化资料清单+可引用观点摘要.xlsx")
OUT = SRC.with_name("西昌三资三化投标服务方案-完整版-V9-课件资料增强版.docx")
REPORT = SRC.with_name("V9完善说明-课件资料整合清单.md")

THEME_BLUE = "0A1F3F"
THEME_TEAL = "1A5C6E"
THEME_GOLD = "C5955C"
LIGHT = "F5F2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(9)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                set_cell_shading(cell, THEME_BLUE)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor.from_string('FFFFFF')
                        r.font.bold = True
            elif row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT)


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)


def make_paragraph(doc, text='', style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p


def insert_block_before(doc, marker, block_func):
    for p in doc.paragraphs:
        if p.text.strip().startswith(marker):
            tmp = Document()
            block_func(tmp)
            for element in tmp.element.body[:-1]:
                p._p.addprevious(deepcopy(element))
            return True
    raise ValueError(f"marker not found: {marker}")


def append_block(doc, block_func):
    block_func(doc)


def add_para(doc, text, style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0.74) if not style else None
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2 = p.add_run(text[len(bold_lead):])
        r2.font.name = '宋体'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        r = p.add_run(text)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color='FFFFFF')
        set_cell_shading(hdr[i], THEME_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val)
    style_table(table)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def load_quotes():
    wb = load_workbook(XLSX, data_only=True)
    ws = wb['可引用观点摘要']
    quotes = []
    for row in ws.iter_rows(min_row=4, values_only=True):
        if row and row[0]:
            quotes.append({"theme": row[1], "quote": row[2], "use": row[3], "source": row[4]})
    return quotes


def block_policy(doc):
    doc.add_heading('1.13 最新讲义资料补充：从“三资三化”到“大财政体系”的政策深化', level=2)
    add_para(doc, '结合采购人最新提供的讲义课件及实施方案资料，本项目对“三资三化”的理解进一步从单一资产盘活扩展为县域财政资源统筹、国有资本运营和财政可持续能力建设的系统工程。其核心不只是“把资产卖出去”，而是通过清查、确权、确值、分类、运营、收益统筹六个环节，推动国有资源资产化、国有资产资本化、国有资本杠杆化，最终形成财政资源可识别、可计量、可运营、可监管、可循环的闭环机制。')
    add_para(doc, '从最新政策资料看，“三资三化”改革与“大财政”“大统筹”“全口径预算管理”高度同向，要求把依托行政权力、政府信用、国有资源资产取得的收入纳入政府预算统筹，避免资产收益游离于预算之外。对西昌市而言，本项目成果应服务于财政局统筹全市资源资产、培育稳定现金流、提高财政调控能力和防范债务风险四项目标。')
    add_table(doc, ['政策维度', '课件资料提炼', '本项目转化要求'], [
        ['大财政体系', '把分散在部门、平台公司和公共领域的资源资产纳入财政统筹视野。', '建立跨部门资产资源“一本账”，形成财政可统筹的底数基础。'],
        ['大统筹机制', '推动资金、资产、资源联动配置，提升财政资源使用效益。', '形成资产清单、收益清单、问题清单、项目清单“四张清单”。'],
        ['全口径预算', '国有资产资源收益应纳入预算管理，避免体外循环。', '设计收入归集、非税缴库、预算安排和绩效跟踪闭环。'],
        ['市场化运营', '通过经营权配置、租赁、转让、资产证券化等方式提高资产流动性。', '坚持公益属性、合规程序和市场价值三者平衡。'],
    ], [3, 6, 7])


def block_method(doc):
    doc.add_heading('3.16 最新课件转化的“七步闭环”实施法', level=2)
    add_para(doc, '在原“四阶段递进法”基础上，结合最新讲义资料，项目实施进一步细化为“清查摸底—确权确值—分类评价—盘活设计—交易实施—收益统筹—监督闭环”七步工作法。七步法强调每一步均形成可交付成果，避免停留在概念阐释和文字方案层面。')
    add_table(doc, ['步骤', '工作重点', '交付成果', '质量控制点'], [
        ['一、清查摸底', '汇集财政、自然资源、住建、水务、文旅、林草、国资等部门资料，形成全口径底数。', '资产资源摸底清单、资料缺口清单', '账账、账实、账证、账图交叉核验。'],
        ['二、确权确值', '核实权属、用途、限制条件、历史形成原因和可运营边界。', '权属诊断表、估值路径建议表', '权属不清不得直接进入交易环节。'],
        ['三、分类评价', '按公益性、经营性、资源性、潜在收益性、风险等级分类。', '资产分级分类矩阵', '明确能用则用、不用则售、不售则租、能融则融。'],
        ['四、盘活设计', '选择租赁、经营权配置、特许经营、合作开发、资产证券化等路径。', '单项资产盘活方案、项目包方案', '测算现金流、财政贡献和社会影响。'],
        ['五、交易实施', '设计交易结构、公开程序、合同条款、履约监管要求。', '交易实施建议书、协议关键条款清单', '防止低价处置和利益输送。'],
        ['六、收益统筹', '明确收入归集、非税缴库、预算安排和后续用途。', '收益统筹方案、资金闭环台账', '严禁截留、坐支、体外循环。'],
        ['七、监督闭环', '建立动态更新、绩效评价、审计追踪和整改机制。', '后续监管台账、年度评估报告模板', '形成可持续治理机制。'],
    ], [2.2, 5.4, 4.2, 4.2])
    doc.add_heading('3.17 数字化赋能“三资三化”的操作方案', level=2)
    add_para(doc, '最新课件强调“三资三化×数字化”是存量资产盘活的新引擎。本项目拟在不增加采购人系统建设负担的前提下，采用“轻量化数据底座+标准化字段+动态更新机制”的方式，将资产信息、权属信息、经营状态、估值参数、盘活路径、收益归集和风险事项统一纳入可维护台账。')
    add_table(doc, ['数字化模块', '核心功能', '对投标服务的增强价值'], [
        ['资产身份证', '为每项资源资产建立唯一编码，关联位置、权属、用途、照片、图纸和合同。', '解决“底数不清、口径不一、部门重复统计”问题。'],
        ['价值雷达图', '从现金流、区位、权属清晰度、市场需求、公益约束等维度评分。', '支撑盘活优先序排序和项目包筛选。'],
        ['项目储备库', '将成熟度较高资产转化为近期、中期、远期项目储备。', '服务2026年及“十五五”项目谋划。'],
        ['收益跟踪表', '记录出让、租赁、运营、分成、缴库和预算安排情况。', '把资产盘活结果纳入财政管理闭环。'],
    ], [3, 5.5, 6])


def block_lists(doc):
    doc.add_heading('5.6 “四张清单”成果体系升级', level=2)
    add_para(doc, '根据最新讲义资料，本项目清单编制不止形成静态资产台账，还应形成可直接服务决策、交易和监管的“四张清单”。四张清单分别解决“有什么、能做什么、问题在哪、怎么落地”的问题，是从资料整理走向可实施方案的关键桥梁。')
    add_table(doc, ['清单名称', '核心字段', '用途', '对应成果'], [
        ['资产资源清单', '名称、位置、权属、面积/规模、账面价值、现状用途、管理主体。', '回答全市国有资源资产底数问题。', '《国有资产资源整合清单》'],
        ['收益潜力清单', '现有收入、潜在收入、估值依据、现金流稳定性、市场需求。', '回答哪些资产具备盘活价值。', '《资产收益测算表》'],
        ['问题风险清单', '权属瑕疵、历史遗留、公益约束、债务关联、处置限制。', '回答哪些事项需要前置处理。', '《问题处理预案》'],
        ['项目实施清单', '盘活模式、实施主体、交易路径、时间安排、预期收益。', '回答近期能落地哪些项目包。', '《分批实施方案》'],
    ], [3, 5.5, 4.5, 4.5])


def block_risk(doc):
    doc.add_heading('6.17 最新资料补充的七类红线风险与防控措施', level=2)
    add_para(doc, '最新讲义和项目资料进一步提示，“三资三化”项目的风险并不只在资产评估环节，而是贯穿权属确认、交易结构、融资安排、收益归集、公共服务和后续运营全过程。联合体将在原有风险预案基础上增加七类红线风险防控表。')
    add_table(doc, ['风险类别', '典型表现', '防控措施'], [
        ['虚假盘活风险', '仅做账面划转、包装融资，未形成真实经营现金流。', '以现金流、运营主体、合同履约和收益缴库作为实质判断标准。'],
        ['重复融资风险', '同一资产、同一收益权被多头质押或重复打包。', '建立资产融资状态字段，交易前核验抵押、担保、质押和债务关系。'],
        ['收益高估风险', '客流、租金、收费、增长率预测脱离市场现实。', '采用保守、中性、乐观三情景测算，重大项目引入第三方评估。'],
        ['权属不清风险', '产权证、管理权、使用权、经营权边界不一致。', '权属瑕疵资产单独列示，先确权后盘活，不带病交易。'],
        ['公益属性弱化风险', '公共资源商业化后影响群众基本服务获取。', '保留政府监管权、价格约束、公益服务比例和应急接管条款。'],
        ['隐性债务风险', '以资产盘活名义变相举债、承诺回购、固定收益兜底。', '坚持市场化、法治化原则，合同文本严禁财政兜底和明股实债安排。'],
        ['国有资产流失风险', '低估低价出让、非公开交易、利益输送。', '评估前置、公开交易、底价控制、利益冲突回避和审计追踪。'],
    ], [3, 5.8, 7.2])


def block_after(doc):
    doc.add_heading('10.7 后续服务升级：从成果交付到持续运营辅导', level=2)
    add_para(doc, '结合最新课件资料，本项目后续服务将从“答疑式服务”升级为“运营辅导式服务”。在正式成果提交后，联合体将围绕项目过会、交易落地、收益归集、绩效跟踪和政策迭代提供持续支持，确保成果能够转化为财政收入、项目储备和管理能力。')
    add_table(doc, ['服务事项', '服务内容', '服务价值'], [
        ['过会辅导', '协助财政局准备政府常务会、专题会、部门协调会汇报材料和答疑口径。', '提高方案决策通过率。'],
        ['交易辅导', '协助完善交易公告、竞买条件、协议条款、履约监管条款。', '提高盘活项目落地率。'],
        ['收益归集辅导', '协助建立收入缴库、预算安排、绩效评价和审计追踪台账。', '防止收益体外循环。'],
        ['项目储备辅导', '围绕2026年和“十五五”项目谋划，持续更新项目库和融资工具匹配表。', '形成长期财政资源储备。'],
        ['能力建设培训', '面向财政、国资、自然资源、住建、水务、文旅等部门开展专题培训。', '把一次性咨询成果转化为内部治理能力。'],
    ], [3, 6, 6])
    doc.add_heading('附录：本次最新课件资料整合说明', level=1)
    add_para(doc, '本次完善重点吸收采购人提供的《三资三化及存量资产盘活课件》《国有资本运营与存量资产盘活》《乌鲁木齐补充课件》《关于全面推进XX县（区）“三资三化”改革专项行动的实施方案》等资料，并结合既有文献分析、项目方案底稿和参考案例，对投标方案进行了结构化增强。')
    add_table(doc, ['整合方向', '新增内容', '嵌入章节'], [
        ['政策深化', '大财政、大统筹、全口径预算、国有资产收益统筹等政策逻辑。', '第一章'],
        ['实施路径', '七步闭环实施法和数字化赋能方案。', '第三章'],
        ['成果清单', '资产资源、收益潜力、问题风险、项目实施四张清单。', '第五章'],
        ['风险控制', '虚假盘活、重复融资、收益高估、隐性债务等七类红线风险。', '第六章'],
        ['后续服务', '从成果交付延伸至过会、交易、收益归集和项目储备辅导。', '第十章'],
    ], [3, 7, 4])


def apply_base_styles(doc):
    heading2_prefixes = (
        '1.13 最新讲义资料补充', '3.16 最新课件转化', '3.17 数字化赋能',
        '5.6 “四张清单”成果体系升级', '6.17 最新资料补充', '10.7 后续服务升级'
    )
    heading1_prefixes = ('附录：本次最新课件资料整合说明',)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith(heading1_prefixes):
            p.style = 'Heading 1'
        elif text.startswith(heading2_prefixes):
            p.style = 'Heading 2'
        for r in p.runs:
            if r.text:
                if p.style and p.style.name.startswith('Heading'):
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.color.rgb = RGBColor.from_string(THEME_BLUE)
                    r.font.bold = True
                else:
                    if not r.font.name:
                        r.font.name = '宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def main():
    doc = Document(SRC)
    insert_block_before(doc, '第二章', block_policy)
    insert_block_before(doc, '第四章', block_method)
    insert_block_before(doc, '第六章', block_lists)
    insert_block_before(doc, '第七章', block_risk)
    append_block(doc, block_after)
    apply_base_styles(doc)
    doc.save(OUT)

    report = """# V9完善说明：课件资料整合清单

## 输出文件
- `{out}`

## 完善依据
- `C:\\Users\\scrccpa\\Desktop\\三资三化课件\\三资三化资料清单+可引用观点摘要.xlsx`
- `C:\\Users\\scrccpa\\Desktop\\三资三化课件` 下最新 PDF/PPTX/DOCX 课件资料
- 既有西昌“三资三化”投标方案 V8f、文献分析稿和原始参考文献

## 本次增强位置
1. 第一章新增 `1.13 最新讲义资料补充：从“三资三化”到“大财政体系”的政策深化`
2. 第三章新增 `3.16 最新课件转化的“七步闭环”实施法`、`3.17 数字化赋能“三资三化”的操作方案`
3. 第五章新增 `5.6 “四张清单”成果体系升级`
4. 第六章新增 `6.17 最新资料补充的七类红线风险与防控措施`
5. 第十章新增 `10.7 后续服务升级：从成果交付到持续运营辅导`
6. 文末新增 `附录：本次最新课件资料整合说明`

## 内容导向
- 从“资产盘活方案”提升为“大财政体系下的三资统筹改革方案”
- 从“资料整理”提升为“四张清单+七步闭环+收益统筹+监管闭环”
- 从“交付文本”提升为“过会、交易、收益归集、项目储备”的持续运营辅导
""".format(out=str(OUT))
    REPORT.write_text(report, encoding='utf-8')
    print(OUT)
    print(REPORT)

if __name__ == '__main__':
    main()
