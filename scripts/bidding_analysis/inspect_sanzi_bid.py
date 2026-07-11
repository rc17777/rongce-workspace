from pathlib import Path
from docx import Document
from openpyxl import load_workbook

bid = Path(r'D:\openclaw-workspace\projects\西昌三资三化投标\西昌三资三化投标服务方案-完整版-V8f.docx')
xlsx = Path(r'C:\Users\scrccpa\Desktop\三资三化课件\三资三化资料清单+可引用观点摘要.xlsx')
out = Path(r'C:\Users\scrccpa\.openclaw\workspace\output\三资三化投标方案结构检查.txt')
out.parent.mkdir(parents=True, exist_ok=True)
lines = []
doc = Document(bid)
lines.append(f'BID={bid}')
lines.append(f'paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}')
lines.append('## headings/sample')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else ''
    if style.startswith('Heading') or text.startswith(('第一', '第二', '第三', '第四', '第五', '第六', '第七', '第八', '第九', '第十', '一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
        lines.append(f'{i}\t{style}\t{text[:160]}')

wb = load_workbook(xlsx, data_only=True)
lines.append('\n## workbook sheets')
lines.append(str(wb.sheetnames))
for s in ['课件提炼总结','使用建议']:
    ws = wb[s]
    lines.append(f'\n## {s}')
    for row in ws.iter_rows(min_row=4, values_only=True):
        if row and row[0]:
            lines.append(' | '.join(str(x or '') for x in row[:3]))

ws = wb['可引用观点摘要']
lines.append('\n## top quote rows')
for row in ws.iter_rows(min_row=4, max_row=20, values_only=True):
    if row and row[0]:
        lines.append(' | '.join(str(x or '')[:220] for x in row[:5]))

out.write_text('\n'.join(lines), encoding='utf-8')
print(out)
