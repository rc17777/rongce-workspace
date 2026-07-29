# -*- coding: utf-8 -*-
"""
批量 DOCX → Markdown 转换 + 入库
输入: E:\2026\审计方法&政策文件\杂志资料\  (DOCX)
输出: D:\openclaw-workspace\knowledge\11-杂志文献\  (MD)
"""
import os, sys, re, time
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)

SRC_DIR = r'E:\2026\审计方法&政策文件\杂志资料'
DST_DIR = r'D:\openclaw-workspace\knowledge\11-杂志文献'

# 保留的目录结构（只保留杂志名称层级，去掉期号子目录）
MAGAZINE_MAP = {
    '财政监督': '财政监督',
    '中国内部审计': '中国内部审计',
    '中国审计': '中国审计',
    '中国注册会计师': '中国注册会计师',
    '审计观察': '审计观察',
    '四川注册会计师': '四川注册会计师',
}

def extract_text_from_docx(docx_path):
    """从 DOCX 提取正文文本"""
    try:
        doc = Document(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return '\n\n'.join(paragraphs)
    except Exception as e:
        return f"[DOCX解析错误: {e}]"

def extract_tables_from_docx(docx_path):
    """从 DOCX 提取表格为 Markdown"""
    try:
        doc = Document(docx_path)
        tables_md = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
                rows.append(cells)
            if not rows:
                continue
            md = '| ' + ' | '.join(rows[0]) + ' |\n'
            md += '| ' + ' | '.join(['---'] * len(rows[0])) + ' |\n'
            for r in rows[1:]:
                # Pad row to match header length
                r = r + [''] * (len(rows[0]) - len(r))
                md += '| ' + ' | '.join(r[:len(rows[0])]) + ' |\n'
            tables_md.append(md)
        return tables_md
    except:
        return []

def sanitize_filename(name):
    """清理文件名中的非法字符"""
    return re.sub(r'[<>:"/\\|?*]', '', name).strip()

def get_magazine_name(file_rel_path):
    """从相对路径推断杂志名称"""
    for mag in sorted(MAGAZINE_MAP.keys(), key=len, reverse=True):
        if mag in file_rel_path:
            return MAGAZINE_MAP[mag]
    return '其他'

def docx_to_md(docx_path, file_rel_path):
    """单个 DOCX → Markdown 转换"""
    text = extract_text_from_docx(docx_path)
    tables = extract_tables_from_docx(docx_path)
    
    # 提取标题（第一行）
    lines = text.split('\n\n')
    title = lines[0][:80] if lines else os.path.basename(docx_path)
    # Clean up title
    title = title.strip().rstrip('.')
    
    # 检测杂志名称
    mag = get_magazine_name(file_rel_path)
    
    # 构建 YAML frontmatter
    yaml = [
        '---',
        f'title: "{title}"',
        f'magazine: "{mag}"',
        f'source: "{file_rel_path}"',
        f'type: magazine',
        f'processed: {time.strftime("%Y-%m-%d")}',
        '---',
        '',
    ]
    
    # 正文
    md_body = text
    
    # 如果有表格，追加
    if tables:
        md_body += '\n\n---\n\n## 表格\n\n'
        for t in tables:
            md_body += t + '\n\n'
    
    return '\n'.join(yaml) + '\n' + md_body

def main():
    os.makedirs(DST_DIR, exist_ok=True)
    
    # Collect all DOCX files
    docx_files = []
    for root, dirs, files in os.walk(SRC_DIR):
        for f in files:
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                docx_files.append(os.path.join(root, f))
    
    print(f'找到 {len(docx_files)} 个 DOCX 文件\n')
    
    # Process
    converted = 0
    errors = 0
    skipped = 0
    
    for fp in sorted(docx_files):
        rel = os.path.relpath(fp, SRC_DIR)
        mag = get_magazine_name(rel)
        
        # Determine output path
        out_dir = os.path.join(DST_DIR, mag)
        os.makedirs(out_dir, exist_ok=True)
        
        # Output filename
        base = os.path.splitext(os.path.basename(fp))[0]
        out_name = sanitize_filename(base) + '.md'
        out_path = os.path.join(out_dir, out_name)
        
        # Skip if already exists
        if os.path.exists(out_path):
            src_size = os.path.getsize(out_path)
            if src_size > 500:  # Non-empty
                print(f'  ⏭  {mag}/{out_name} (已存在)')
                skipped += 1
                continue
        
        try:
            md_content = docx_to_md(fp, rel)
            with open(out_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            sz = os.path.getsize(out_path)
            status = '✅' if sz > 500 else '⚠️ (空白)'
            print(f'  {status} {mag}/{out_name} ({sz/1024:.1f} KB)')
            converted += 1
        except Exception as e:
            print(f'  ❌ {rel}: {e}')
            errors += 1
    
    print(f'\n{"="*50}')
    print(f'转换完成: {converted} 成功 / {errors} 失败 / {skipped} 跳过')
    print(f'输出目录: {DST_DIR}')
    print(f'总文件: {len(docx_files)}')

if __name__ == '__main__':
    main()
