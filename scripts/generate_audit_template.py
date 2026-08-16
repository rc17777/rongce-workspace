"""
审计/咨询报告标准模板生成器
基于中银证券研报设计分析，适配政府审计/工程咨询报告格式
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ===== 配色方案（与中银证券研报一致的设计哲学：克制、单色强调） =====
COLOR = {
    'primary':        RGBColor(0x1A, 0x3A, 0x6E),  # 深蓝 #1a3a6e - 主色调（替代中银红色，更适合政府审计）
    'primary_dark':   RGBColor(0x0F, 0x24, 0x48),  # 更深蓝
    'primary_light':  RGBColor(0xD6, 0xE0, 0xF0),  # 浅蓝背景
    'primary_pale':   RGBColor(0xEE, 0xF2, 0xF9),  # 极浅蓝
    'secondary':      RGBColor(0xE8, 0x6A, 0x17),  # 橙色 #e86a17 - 图表强调色
    'secondary_light':RGBColor(0xFD, 0xF0, 0xE6),  # 浅橙
    'text':           RGBColor(0x33, 0x33, 0x33),  # 正文黑 #333333
    'text_light':     RGBColor(0x66, 0x66, 0x66),  # 辅助灰
    'white':          RGBColor(0xFF, 0xFF, 0xFF),
    'black':          RGBColor(0x00, 0x00, 0x00),
    'table_border':   RGBColor(0x99, 0x99, 0x99),  # 表格边框灰
    'table_header':   RGBColor(0x1A, 0x3A, 0x6E),  # 表头底色（与主色一致）
    'row_alt':        RGBColor(0xF5, 0xF7, 0xFA),  # 交替行色
    'risk_high':      RGBColor(0xCC, 0x33, 0x33),  # 高风险红
    'risk_medium':    RGBColor(0xE8, 0x6A, 0x17),  # 中风险橙
    'risk_low':       RGBColor(0x33, 0x99, 0x33),  # 低风险绿
    'green':          RGBColor(0x27, 0xAE, 0x60),
    'green_bg':       RGBColor(0xE8, 0xF8, 0xF0),
    'red_bg':         RGBColor(0xFD, 0xED, 0xED),
    'yellow_bg':      RGBColor(0xFE, 0xF9, 0xE7),
}

# ===== 工具函数 =====
def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表风格）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if spec:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{spec["sz"]}" w:space="0" w:color="{spec["color"]}"/>'
            )
            borders.append(border)
    tcPr.append(borders)

def make_border_spec(sz, color_hex):
    return {'sz': str(sz), 'color': color_hex}

def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def add_heading_styled(doc, text, level=1):
    """添加自定义标题（与政府审计格式一致）"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = COLOR['primary']
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(13)
    set_paragraph_spacing(h, before=12, after=6)
    return h

def add_body(doc, text, bold=False, size=10.5, color=None, alignment=None, indent=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or COLOR['text']
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    set_paragraph_spacing(p, before=2, after=4, line=1.35)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    return p

def add_key_finding(doc, text, symbol='►'):
    """添加核心发现段落（带符号标记，借鉴研报■标记）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    # 符号
    run_sym = p.add_run(f'{symbol}  ')
    run_sym.font.size = Pt(11)
    run_sym.font.color.rgb = COLOR['primary']
    run_sym.bold = True
    # 内容
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = COLOR['text']
    run_text.font.name = '宋体'
    run_text._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_paragraph_spacing(p, before=2, after=4, line=1.35)
    return p

def add_source_note(doc, text):
    """添加数据来源标注（借鉴研报"资料来源：XXX"格式）"""
    p = doc.add_paragraph()
    run = p.add_run(f'数据来源：{text}')
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR['text_light']
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.italic = True
    set_paragraph_spacing(p, before=0, after=8)
    return p

def create_three_line_table(doc, headers, rows, col_widths=None, header_color=None):
    """创建标准三线表（学术/审计标准格式）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_color = header_color or COLOR['table_header']

    # 设置表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR['white']
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, f'{hdr_color[0]:02x}{hdr_color[1]:02x}{hdr_color[2]:02x}')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 设置数据行
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR['text']
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 数字右对齐，文字左对齐
            if isinstance(value, (int, float)) or (isinstance(value, str) and value.replace('.', '').replace(',', '').replace('-', '').replace('%', '').isdigit()):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 交替行底色
            if r % 2 == 1:
                set_cell_shading(cell, f'{COLOR["row_alt"][0]:02x}{COLOR["row_alt"][1]:02x}{COLOR["row_alt"][2]:02x}')

    # 三线表边框：顶线粗、表头下线粗、底线粗，其余无线
    border_color = f'{COLOR["table_border"][0]:02x}{COLOR["table_border"][1]:02x}{COLOR["table_border"][2]:02x}'
    thick = make_border_spec(12, border_color)  # 粗线 1.5pt
    thin = make_border_spec(4, border_color)    # 细线 0.5pt

    n_rows = len(rows) + 1
    for r_idx in range(n_rows):
        for c_idx in range(len(headers)):
            cell = table.rows[r_idx].cells[c_idx]
            top = thick if r_idx == 0 else (thick if r_idx == 1 else None)
            bottom = thick if r_idx == n_rows - 1 else (None)
            set_cell_borders(cell, top=top, bottom=bottom,
                             left=make_border_spec(0, 'FFFFFF'), right=make_border_spec(0, 'FFFFFF'))

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table

def create_finding_summary_table(doc, findings):
    """
    创建问题发现汇总表（审计报告核心表格）
    findings: [(编号, 问题描述, 涉及金额, 风险等级, 建议)]
    """
    table = doc.add_table(rows=1 + len(findings), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['编号', '问题描述', '涉及金额(万元)', '风险等级', '整改建议']
    hdr_color_hex = f'{COLOR["table_header"][0]:02x}{COLOR["table_header"][1]:02x}{COLOR["table_header"][2]:02x}'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR['white']
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, hdr_color_hex)

    risk_colors = {
        '高': (COLOR['red_bg'], COLOR['risk_high']),
        '中': (COLOR['yellow_bg'], COLOR['risk_medium']),
        '低': (COLOR['green_bg'], COLOR['risk_low']),
    }

    for r, (no, desc, amount, risk, suggestion) in enumerate(findings):
        row_data = [no, desc, amount, risk, suggestion]
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if c == 3:  # 风险等级列
                run.bold = True
                bg, fg = risk_colors.get(str(val), (None, COLOR['text']))
                run.font.color.rgb = fg
                if bg:
                    set_cell_shading(cell, f'{bg[0]:02x}{bg[1]:02x}{bg[2]:02x}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c == 2:  # 金额列
                run.font.color.rgb = COLOR['text']
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif c == 0:  # 编号列
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.color.rgb = COLOR['text_light']
            else:
                run.font.color.rgb = COLOR['text']
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 三线表边框
    border_color_hex = f'{COLOR["table_border"][0]:02x}{COLOR["table_border"][1]:02x}{COLOR["table_border"][2]:02x}'
    thick = make_border_spec(12, border_color_hex)
    n_rows = len(findings) + 1
    n_cols = 5
    for r_idx in range(n_rows):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            top = thick if r_idx == 0 else (thick if r_idx == 1 else None)
            bottom = thick if r_idx == n_rows - 1 else None
            set_cell_borders(cell, top=top, bottom=bottom, left=make_border_spec(0, 'FFFFFF'), right=make_border_spec(0, 'FFFFFF'))

    # 列宽
    widths = [1.0, 5.5, 2.5, 1.5, 5.5]
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)

    return table

def add_chart_placeholder(doc, title, chart_type, description):
    """添加图表占位符（含标题、类型说明、数据来源行）"""
    # 图表标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{title}')
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR['primary']
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = True
    set_paragraph_spacing(p, before=10, after=2)

    # 占位框
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ph = placeholder.add_run(f'[{chart_type}：{description}]')
    run_ph.font.size = Pt(9)
    run_ph.font.color.rgb = COLOR['text_light']
    run_ph.font.name = '宋体'
    run_ph._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_ph.italic = True

    # 使用表格模拟图表区域
    chart_bg = doc.add_table(rows=10, cols=1)
    chart_bg.alignment = WD_TABLE_ALIGNMENT.CENTER
    bg_hex = f'{COLOR["row_alt"][0]:02x}{COLOR["row_alt"][1]:02x}{COLOR["row_alt"][2]:02x}'
    for row in chart_bg.rows:
        cell = row.cells[0]
        set_cell_shading(cell, bg_hex)
        cell.height = Cm(0.5)
        p_bg = cell.paragraphs[0]
        p_bg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_bg = p_bg.add_run(f'  {chart_type}区域 - 建议配色：主色{COLOR["primary"]} / 强调色{COLOR["secondary"]} / 灰色系辅助  ')
        run_bg.font.size = Pt(7.5)
        run_bg.font.color.rgb = COLOR['text_light']
        run_bg.font.name = '宋体'
        run_bg._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 数据来源
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_src = p_src.add_run('数据来源：[请填写具体数据来源及取证方式]')
    run_src.font.size = Pt(8)
    run_src.font.color.rgb = COLOR['text_light']
    run_src.font.name = '宋体'
    run_src._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_src.italic = True
    set_paragraph_spacing(p_src, before=2, after=12)

# ===== 主函数：生成审计报告模板 =====
def generate_audit_template():
    doc = Document()

    # ----- 页面设置 -----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ----- 页眉页脚 -----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = hp.add_run('[项目名称]  ｜  审计/咨询报告')
    run_h.font.size = Pt(8)
    run_h.font.color.rgb = COLOR['text_light']
    run_h.font.name = '宋体'
    run_h._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 右侧日期
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h2 = hp2.add_run(datetime.date.today().strftime('%Y年%m月%d日'))
    run_h2.font.size = Pt(8)
    run_h2.font.color.rgb = COLOR['text_light']
    run_h2.font.name = '宋体'
    run_h2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 页脚：页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 插入页码域
    run_f = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_f._r.append(fldChar1)
    run_f2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_f2._r.append(instrText)
    run_f3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_f3._r.append(fldChar2)
    fp.add_run(' / ')
    # 总页数域
    run_f4 = fp.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_f4._r.append(fldChar3)
    run_f5 = fp.add_run()
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run_f5._r.append(instrText2)
    run_f6 = fp.add_run()
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_f6._r.append(fldChar4)
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR['text_light']

    # ========== 封面页 ==========
    # 顶部色带
    p_band = doc.add_paragraph()
    p_band.paragraph_format.space_before = Pt(0)
    p_band.paragraph_format.space_after = Pt(0)
    # 用一条水平线模拟色带
    run_band = p_band.add_run('━' * 60)
    run_band.font.size = Pt(14)
    run_band.font.color.rgb = COLOR['primary']

    # 空行
    doc.add_paragraph()

    # 报告类型
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_type = p_type.add_run('审 计 / 咨 询 报 告')
    run_type.font.size = Pt(14)
    run_type.font.color.rgb = COLOR['text_light']
    run_type.font.name = '宋体'
    run_type._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # 主标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('[请输入报告标题]')
    run_title.font.size = Pt(26)
    run_title.font.color.rgb = COLOR['primary']
    run_title.font.name = '黑体'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run_title.bold = True

    doc.add_paragraph()

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('[请输入报告副标题（可选）]')
    run_sub.font.size = Pt(15)
    run_sub.font.color.rgb = COLOR['text']
    run_sub.font.name = '宋体'
    run_sub._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()
    doc.add_paragraph()

    # 核心发现摘要框（借鉴研报封面"支撑评级的要点"）
    p_finding_title = doc.add_paragraph()
    run_ft = p_finding_title.add_run('核心发现摘要')
    run_ft.font.size = Pt(13)
    run_ft.font.color.rgb = COLOR['primary']
    run_ft.font.name = '黑体'
    run_ft._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run_ft.bold = True

    add_key_finding(doc, '[请填写核心发现1：用一句话概括最重要的审计发现或咨询结论]')
    add_key_finding(doc, '[请填写核心发现2：概括第二重要的发现或建议]')
    add_key_finding(doc, '[请填写核心发现3：概括第三重要的发现或建议]')

    doc.add_paragraph()
    doc.add_paragraph()

    # 项目信息表（封面底部）
    info_table = doc.add_table(rows=6, cols=2)
    info_data = [
        ('委托单位', '[请输入委托单位全称]'),
        ('被审计/咨询单位', '[请输入被审计单位全称]'),
        ('项目类型', '□ 绩效评价  □ 资产清查  □ 专项债审计\n□ 工程结算  □ 预算编制  □ 其他：____'),
        ('报告日期', datetime.date.today().strftime('%Y年%m月%d日')),
        ('编制单位', '四川融策会计师事务所 / 四川融策工程咨询公司'),
        ('风险评级', '□ 高风险    □ 中风险    □ 低风险'),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.rows[i].cells[0]
        cell_r = info_table.rows[i].cells[1]
        cell_l.text = ''
        cell_r.text = ''
        pl = cell_l.paragraphs[0]
        rl = pl.add_run(label)
        rl.font.size = Pt(10)
        rl.font.color.rgb = COLOR['primary']
        rl.font.name = '宋体'
        rl._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        rl.bold = True
        pr = cell_r.paragraphs[0]
        rr = pr.add_run(value)
        rr.font.size = Pt(10)
        rr.font.color.rgb = COLOR['text']
        rr.font.name = '宋体'
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 左侧列底色
        set_cell_shading(cell_l, f'{COLOR["primary_pale"][0]:02x}{COLOR["primary_pale"][1]:02x}{COLOR["primary_pale"][2]:02x}')

    # 封面分页
    doc.add_page_break()

    # ========== 目录页 ==========
    add_heading_styled(doc, '目  录', level=1)

    toc_items = [
        ('一、项目概述', '3'),
        ('    1.1 项目背景', '3'),
        ('    1.2 审计/咨询目标', '3'),
        ('    1.3 审计/咨询范围与方法', '4'),
        ('', ''),
        ('二、项目总体情况', '5'),
        ('    2.1 基本情况概述', '5'),
        ('    2.2 主要指标分析', '5'),
        ('    2.3 资金使用总体评价', '6'),
        ('', ''),
        ('三、重点问题分析', '7'),
        ('    3.1 预算执行方面', '7'),
        ('    3.2 资金管理方面', '8'),
        ('    3.3 项目管理方面', '9'),
        ('    3.4 内部控制方面', '10'),
        ('', ''),
        ('四、问题发现汇总', '11'),
        ('', ''),
        ('五、改进建议', '12'),
        ('', ''),
        ('六、结论与总体评价', '14'),
        ('', ''),
        ('附录一：图表索引', '15'),
        ('附录二：配色与排版规范', '16'),
    ]

    toc_table = doc.add_table(rows=len(toc_items) + 1, cols=2)
    # 表头
    for i, h in enumerate(['章节', '页码']):
        cell = toc_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR['primary']
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, f'{COLOR["primary_pale"][0]:02x}{COLOR["primary_pale"][1]:02x}{COLOR["primary_pale"][2]:02x}')

    for i, (item, page) in enumerate(toc_items):
        cell_l = toc_table.rows[i + 1].cells[0]
        cell_r = toc_table.rows[i + 1].cells[1]
        cell_l.text = ''
        cell_r.text = ''
        pl = cell_l.paragraphs[0]
        rl = pl.add_run(item)
        rl.font.size = Pt(10)
        rl.font.color.rgb = COLOR['primary'] if item and not item.startswith('    ') else COLOR['text']
        rl.font.name = '宋体'
        rl._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        rl.bold = bool(item and not item.startswith('    '))
        pr = cell_r.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pr.add_run(page)
        rr.font.size = Pt(10)
        rr.font.color.rgb = COLOR['text_light']
        rr.font.name = '宋体'
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 清理边框
    for row in toc_table.rows:
        for cell in row.cells:
            set_cell_borders(cell, left=make_border_spec(0, 'FFFFFF'), right=make_border_spec(0, 'FFFFFF'),
                             top=make_border_spec(0, 'FFFFFF'), bottom=make_border_spec(0, 'FFFFFF'))

    toc_table.rows[0].cells[0].width = Cm(12)
    toc_table.rows[0].cells[1].width = Cm(3)

    doc.add_page_break()

    # ========== 一、项目概述 ==========
    add_heading_styled(doc, '一、项目概述', level=1)

    add_heading_styled(doc, '1.1 项目背景', level=2)
    add_body(doc, '根据[委托文件/合同编号]，[委托单位]委托我单位对[被审计单位]开展[项目类型]。本报告旨在全面反映[项目名称]的实施情况、资金使用效果及存在问题，并提出针对性的改进建议。')
    add_body(doc, '本项目于[开始日期]启动，现场工作自[现场日期]至[现场结束日期]，历时[天数]天，累计投入[人数]人次。项目组通过资料审查、现场勘查、访谈调研、数据分析等多种方式开展工作，共查阅凭证[数量]份、合同[数量]份、工程资料[数量]份。')

    add_heading_styled(doc, '1.2 审计/咨询目标', level=2)
    add_body(doc, '本次审计/咨询工作的主要目标如下：')
    goals = [
        '核实项目资金的拨付、使用及结余情况，评价资金使用合规性',
        '检查项目实施进度与质量，评价项目管理规范性',
        '识别项目实施过程中的风险点和薄弱环节',
        '评估项目绩效目标完成情况，提出优化建议',
        '为单位决策和改进管理提供参考依据',
    ]
    for g in goals:
        add_key_finding(doc, g, '①')

    add_heading_styled(doc, '1.3 审计/咨询范围与方法', level=2)
    add_body(doc, '本次审计/咨询工作采取以下方法：')

    # 方法列表表
    method_data = [
        ('资料审查法', '核查财务凭证、合同文件、项目管理档案等原始资料', '全面覆盖'),
        ('现场勘查法', '实地走访项目现场，核查工程进度及质量', '重点抽查'),
        ('数据分析法', '运用统计分析、趋势分析等方法对资金数据进行深度分析', '全面覆盖'),
        ('访谈调研法', '对项目负责人、财务人员、施工单位等进行结构化访谈', '关键人员'),
        ('比对核验法', '将实际执行情况与立项批复、预算方案进行逐项比对', '重点事项'),
    ]
    methods_table = create_three_line_table(doc,
        ['审计方法', '方法说明', '覆盖范围'],
        method_data,
        col_widths=[3, 10, 3]
    )
    doc.add_paragraph()

    doc.add_page_break()

    # ========== 二、项目总体情况 ==========
    add_heading_styled(doc, '二、项目总体情况', level=1)

    add_heading_styled(doc, '2.1 基本情况概述', level=2)
    add_body(doc, '[请填写项目基本情况，包括项目立项背景、总投资额、资金来源构成、建设内容、实施周期等关键信息。建议用精炼的语言在2-3段内完成。]')

    add_heading_styled(doc, '2.2 主要指标分析', level=2)

    # 预算执行对比表
    add_body(doc, '预算执行情况对比表：', bold=True, size=10)
    budget_data = [
        ('项目A', '1,000.00', '856.32', '85.6%', '正常'),
        ('项目B', '500.00', '423.15', '84.6%', '正常'),
        ('项目C', '300.00', '298.70', '99.6%', '正常'),
        ('项目D', '200.00', '67.50', '33.8%', '⚠ 偏低'),
        ('合计', '2,000.00', '1,645.67', '82.3%', '—'),
    ]
    budget_table = create_three_line_table(doc,
        ['项目名称', '预算金额(万元)', '实际支出(万元)', '执行率', '状态'],
        budget_data,
        col_widths=[4, 3.5, 3.5, 3, 2]
    )
    add_source_note(doc, '[数据来源说明，如：根据项目财务台账及银行对账单汇总，统计截至202X年X月X日')
    doc.add_paragraph()

    # 饼图占位
    add_chart_placeholder(doc, '图表1：项目资金构成分布', '饼图/环形图',
                         '主色：深蓝#1a3a6e渐变，建议蓝灰色调色盘，标注占比百分比')
    add_source_note(doc, '[数据来源]')
    doc.add_paragraph()

    # 柱状图占位
    add_chart_placeholder(doc, '图表2：各项目预算执行率对比', '柱状图',
                         '单色渐变柱状图，主色#1a3a6e，达标线用绿色虚线标注')
    add_source_note(doc, '[数据来源]')

    doc.add_page_break()

    # ========== 三、重点问题分析 ==========
    add_heading_styled(doc, '三、重点问题分析', level=1)

    add_heading_styled(doc, '3.1 预算执行方面', level=2)
    add_body(doc, '[请描述预算执行方面发现的问题，包括：预算编制是否合理、预算调整是否合规、执行偏差的原因分析等。]')
    add_body(doc, '示例：经核查，项目D预算执行率仅为33.8%，主要原因为：① 项目前期论证不充分，立项条件不成熟；② 实施过程中遇到[具体困难]，导致进度严重滞后。涉及未执行预算资金132.50万元。')

    add_chart_placeholder(doc, '图表3：月度预算执行进度趋势', '折线图+柱状图组合',
                         '折线图：月度累计执行率（蓝色实线），柱状图：月度支出额（浅蓝色），计划线：红色虚线')
    add_source_note(doc, '[数据来源]')
    doc.add_paragraph()

    add_heading_styled(doc, '3.2 资金管理方面', level=2)
    add_body(doc, '[请描述资金管理方面发现的问题，包括：资金拨付及时性、专款专用情况、资金沉淀问题等。]')

    add_heading_styled(doc, '3.3 项目管理方面', level=2)
    add_body(doc, '[请描述项目管理方面发现的问题，包括：招投标合规性、合同管理、工程变更、验收程序等。]')

    add_heading_styled(doc, '3.4 内部控制方面', level=2)
    add_body(doc, '[请描述内控方面发现的问题，包括：制度建设、岗位设置、审批流程、档案管理等。]')

    doc.add_page_break()

    # ========== 四、问题发现汇总 ==========
    add_heading_styled(doc, '四、问题发现汇总', level=1)
    add_body(doc, '经审计/核查，本次共发现[数量]项问题，按风险等级分类如下：')

    # 风险统计表
    risk_stats = [
        ('高风险', '3', '15.8%', '需立即整改'),
        ('中风险', '8', '42.1%', '限期整改'),
        ('低风险', '8', '42.1%', '持续关注'),
        ('合计', '19', '100%', '—'),
    ]
    risk_table = create_three_line_table(doc,
        ['风险等级', '问题数量', '占比', '处理建议'],
        risk_stats,
        col_widths=[4, 3, 3, 6]
    )
    doc.add_paragraph()

    add_body(doc, '具体问题发现清单如下：', bold=True, size=10)
    sample_findings = [
        ('F-001', '项目D预算执行率仅33.8%，立项审批与实际脱节，项目推进严重滞后', '132.50', '高', '建议重新评估项目可行性，必要时终止或调整方案，收回结余资金'),
        ('F-002', '部分支出未按规定取得合规发票，涉及金额45.30万元', '45.30', '高', '限期补充取得合规票据，对无法补救的按税法规定处理'),
        ('F-003', '项目A未经批准擅自调整部分建设内容，涉及变更金额85.00万元', '85.00', '高', '补充变更审批手续，追究相关人员责任，完善变更管理制度'),
        ('F-004', '部分合同签订不规范，存在要素缺失、条款模糊等问题', '—', '中', '修订完善合同模板，加强合同审核流程建设'),
        ('F-005', '专项资金存在短期挪用现象，虽然已在期末归还但内部控制存在漏洞', '120.00', '中', '强化专项资金专户管理，建立定期对账机制'),
        ('F-006', '项目C虽执行率正常但成果交付存在部分质量缺陷', '—', '中', '督促整改，建立质量验收标准及扣款机制'),
        ('F-007', '档案管理不够规范，部分过程性文件缺失', '—', '低', '建立完整的项目档案管理制度和归档清单'),
        ('F-008', '绩效自评报告数据与财务数据存在口径不一致', '—', '低', '统一数据口径，规范绩效指标填报要求'),
    ]
    finding_table = create_finding_summary_table(doc, sample_findings)
    add_source_note(doc, '以上问题分类依据《政府会计制度》及相关审计准则。风险等级按涉及金额、性质严重程度、整改难度综合判定。')
    doc.add_paragraph()

    doc.add_page_break()

    # ========== 五、改进建议 ==========
    add_heading_styled(doc, '五、改进建议', level=1)
    add_body(doc, '针对上述发现的问题，提出以下改进建议：')

    suggestions = [
        ('加强项目前期论证', '建议在项目立项阶段增加可行性研究深度，引入专家评审机制，确保立项条件充分、预算编制合理，从源头杜绝"先天不足"项目。'),
        ('完善资金管理制度', '建立专项资金全流程监控机制，实行专户管理、专款专用。完善资金拨付与项目进度的联动机制，杜绝资金闲置或挪用。'),
        ('规范合同与变更管理', '制定统一的合同模板体系，加强合同审核流程。建立变更审批的层级管理制度，重大变更须经集体决策并书面审批。'),
        ('强化绩效评价闭环', '建立"目标设定→过程监控→结果评价→反馈应用"的全链条绩效管理体系，将评价结果与下年度预算安排挂钩。'),
    ]
    for i, (title, detail) in enumerate(suggestions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        # 编号
        run_no = p.add_run(f'{i}. ')
        run_no.font.size = Pt(11)
        run_no.font.color.rgb = COLOR['primary']
        run_no.bold = True
        # 标题
        run_t = p.add_run(f'{title}  ')
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = COLOR['primary']
        run_t.font.name = '宋体'
        run_t._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_t.bold = True
        # 详情
        run_d = p.add_run(detail)
        run_d.font.size = Pt(10)
        run_d.font.color.rgb = COLOR['text']
        run_d.font.name = '宋体'
        run_d._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_paragraph_spacing(p, before=4, after=8, line=1.4)

    doc.add_page_break()

    # ========== 六、结论 ==========
    add_heading_styled(doc, '六、结论与总体评价', level=1)
    add_body(doc, '[请撰写总体评价和结论，包括：对项目整体实施效果的评价、主要成绩和亮点、存在的突出问题、改进建议的优先级排序等。]')
    add_body(doc, '综合评价：□ 优秀    □ 良好    □ 一般    □ 较差', bold=True, size=11)

    doc.add_page_break()

    # ========== 附录一：图表索引 ==========
    add_heading_styled(doc, '附录一：图表索引', level=1)
    add_body(doc, '本报告所含图表清单如下（借鉴券商研报图表目录格式）：')

    chart_index = [
        ('图表1', '项目资金构成分布', '饼图/环形图', '第5页'),
        ('图表2', '各项目预算执行率对比', '柱状图', '第5页'),
        ('图表3', '月度预算执行进度趋势', '折线图+柱状图', '第7页'),
        ('表1', '预算执行情况对比表', '三线表', '第5页'),
        ('表2', '问题风险等级统计表', '三线表', '第11页'),
        ('表3', '问题发现清单', '风险分类表', '第11-12页'),
        ('表4', '审计方法清单', '三线表', '第4页'),
    ]
    ci_table = create_three_line_table(doc,
        ['编号', '图表名称', '图表类型', '所在页码'],
        chart_index,
        col_widths=[2, 7, 4, 3]
    )

    doc.add_page_break()

    # ========== 附录二：配色与排版规范 ==========
    add_heading_styled(doc, '附录二：配色与排版规范', level=1)
    add_body(doc, '本模板配色与排版设计参考中银证券行业深度研报，遵循"克制、统一、专业"原则。', bold=True)

    add_heading_styled(doc, '一、配色方案', level=2)

    colors_data = [
        ('主色', f'深蓝 #{COLOR["primary"][0]:02x}{COLOR["primary"][1]:02x}{COLOR["primary"][2]:02x}', '标题、表头、重点文字、图表主色'),
        ('强调色', f'橙色 #{COLOR["secondary"][0]:02x}{COLOR["secondary"][1]:02x}{COLOR["secondary"][2]:02x}', '图表对比色、关键数据标注'),
        ('正文色', f'深灰 #{COLOR["text"][0]:02x}{COLOR["text"][1]:02x}{COLOR["text"][2]:02x}', '正文段落文字'),
        ('辅助色', f'灰色 #{COLOR["text_light"][0]:02x}{COLOR["text_light"][1]:02x}{COLOR["text_light"][2]:02x}', '数据来源、页码、次要信息'),
        ('高风险', f'红色 #{COLOR["risk_high"][0]:02x}{COLOR["risk_high"][1]:02x}{COLOR["risk_high"][2]:02x}', '高风险问题标识'),
        ('中风险', f'橙色 #{COLOR["risk_medium"][0]:02x}{COLOR["risk_medium"][1]:02x}{COLOR["risk_medium"][2]:02x}', '中风险问题标识'),
        ('低风险', f'绿色 #{COLOR["risk_low"][0]:02x}{COLOR["risk_low"][1]:02x}{COLOR["risk_low"][2]:02x}', '低风险问题标识'),
        ('浅蓝背景', f'#{COLOR["primary_pale"][0]:02x}{COLOR["primary_pale"][1]:02x}{COLOR["primary_pale"][2]:02x}', '封面信息表左侧列、目录表头'),
        ('交替行色', f'#{COLOR["row_alt"][0]:02x}{COLOR["row_alt"][1]:02x}{COLOR["row_alt"][2]:02x}', '表格数据行交替底色'),
    ]
    color_table = create_three_line_table(doc,
        ['用途', '颜色值', '说明'],
        colors_data,
        col_widths=[3, 5, 8]
    )

    doc.add_paragraph()
    add_heading_styled(doc, '二、表格规范', level=2)
    add_body(doc, '本模板所有正式表格均采用"三线表"（顶线、表头下线、底线）设计，这是学术论文和政府报告的通用标准：')
    table_rules = [
        '顶线和底线用粗线（1.5pt），表头下线用粗线（1.5pt），其余位置无线条',
        '表头用深蓝色底色+白色文字，与主色调统一',
        '数据行使用交替浅色底色（斑马条纹），提高长表格可读性',
        '数字列右对齐，文字列左对齐，表头居中',
        '每个表格下方标注"数据来源"，格式与券商研报一致',
        '金额列标注单位（万元/亿元），不可省略',
        '负值使用括号表示，如（5.50），比红色更专业',
    ]
    for rule in table_rules:
        add_key_finding(doc, rule, '·')

    doc.add_paragraph()
    add_heading_styled(doc, '三、图表规范', level=2)
    chart_rules = [
        '默认图表配色：主色（深蓝）为主，强调色（橙）用于对比，灰色系辅助',
        '饼图使用深蓝色系渐变调色盘，不宜超过5个扇区',
        '柱状图使用单色渐变，不建议多色柱子（除非必须对比）',
        '折线图主趋势线深蓝色加粗，辅助线灰色虚线',
        '每张图表必须有标题（居中、加粗、深蓝色）和数据来源标注',
        '图表风格全报告保持一致，不混用不同风格的默认配色',
    ]
    for rule in chart_rules:
        add_key_finding(doc, rule, '·')

    doc.add_paragraph()
    add_heading_styled(doc, '四、文字排版规范', level=2)
    text_rules = [
        '正文：宋体 10.5pt，行距1.35倍，段间距4pt',
        '一级标题：黑体 18pt，深蓝色，段前12pt/段后6pt',
        '二级标题：黑体 15pt，深蓝色',
        '三级标题：黑体 13pt，深蓝色',
        '页眉：宋体 8pt，灰色，左侧项目名称+右侧日期',
        '页脚：宋体 8pt，灰色，居中页码（第X页/共Y页）',
        '关键发现：使用 ► 或 · 符号引导，缩进0.5cm',
        '段落首行不缩进（现代排版风格），段间距替代首行缩进',
    ]
    for rule in text_rules:
        add_key_finding(doc, rule, '·')

    # ===== 保存 =====
    output_path = r'D:\openclaw-workspace\审计报告标准模板.docx'
    doc.save(output_path)
    print(f'[OK] 审计报告模板已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    generate_audit_template()
