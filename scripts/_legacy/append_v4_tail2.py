# -*- coding: utf-8 -*-
"""Append tail to rongce_v4_brochure.py"""
from pathlib import Path

tail = r'''
def p9_experience():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_header(d, "代表经验", "REPRESENTATIVE EXPERIENCE")
    draw(d, "长期服务省、市、县多级财政和主管部门，覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。", 160, 240, font(30), MU, 620, 10)
    # Service territory box
    d.rounded_rectangle([880, 130, 1450, 550], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((960, 180), "服务版图", font=font(36, True), fill=rga(TL, 60))
    d.text((960, 230), "SERVICE TERRITORY", font=font(18), fill=rga(TL, 40))
    # Territory items with decorative dots and lines
    terr = [("成都总部", "四川省会，核心枢纽"), ("阿坝州办事处", "川西高原，覆盖藏区"),
            ("西藏办事处", "高原地区拓展"), ("覆盖川、藏、黔", "三省联动，跨区域服务")]
    for i, (name, desc) in enumerate(terr):
        y0 = 280 + i * 60
        d.ellipse([960, y0, 990, y0+30], fill=rga(GD))
        draw(d, name, 1010, y0, font(26, True), IK, 300, 2)
        draw(d, desc, 1010, y0+28, font(18), MU, 300, 2)
        if i < 3:
            d.line([975, y0+30, 975, y0+60], fill=rga(GD, 100), width=2)
    # Clients section
    d.text((160, 390), "代表客户", font=font(32, True), fill=NW)
    d.text((160, 428), "REPRESENTATIVE CLIENTS", font=font(18), fill=rga(GD))
    clients = ["四川省财政厅", "省公安厅交警总队", "省民政厅", "省农业农村厅", "省教育厅",
        "省退役军人事务厅", "省药品监督管理局", "达州市财政局", "绵阳市财政局",
        "宜宾市财政局", "什邡市财政局", "德阳市财政局", "康定市财政局", "九寨沟县财政局"]
    yv = 470
    for i, c in enumerate(clients):
        if i % 5 == 0 and i > 0:
            yv += 40
        cx = 160 + (i % 5) * 280
        d.rounded_rectangle([cx, yv, cx+240, yv+36], radius=6, fill=rga(WH), outline=rga(GD, 80), width=1)
        tw = d.textlength(c, font=font(20))
        d.text((cx+120-tw//2, yv+5), c, font=font(20), fill=NW)
    # Industry coverage note
    d.text((160, 710), "覆盖领域：财政、审计、教育、民政、交通、国资、医保、药监、公安", font=font(22), fill=MU)
    bottom_bar(d)
    return img

# =========== P10: Contact (improved) ===========
def p10_contact():
    img = Image.new("RGB", (W, H), rgb(BG))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, H], fill=rga(NW))
    d.rectangle([0, 0, W, 280], fill=rga(NW))
    d.rectangle([0, 280, W, 290], fill=rga(GD))
    d.text((130, 80), "合作价值", font=font(52, True), fill=WH)
    d.text((130, 150), "VALUE PROPOSITION", font=font(22), fill=rga(GD, 200))
    d.line([130, 190, 500, 190], fill=rga(GD), width=3)
    draw(d, "我们交付的不只是一份报告，更是一套可执行的改进方案。", 130, 260, font(30), WH, 600, 10)
    # Contact card
    d.rounded_rectangle([130, 340, 750, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "联系融策", 180, 370, font(36, True), GD, 500, 8)
    for i, (icon, line) in enumerate([("\u260e", "028-87659276"), ("\u2709", "scrccpa@163.com"),
                                       ("\u25b6", "www.scrccpa.com"), ("\u2302", "成都市金牛区金周路595号")]):
        d.text((180, 440+i*48), icon, font=font(26), fill=GD)
        d.text((215, 438+i*48), line, font=font(26), fill=WH)
    # Commitment card
    d.rounded_rectangle([880, 340, 1530, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "合作承诺", 930, 380, font(36, True), GD, 500, 8)
    for i, item in enumerate(["客观公正  实事求是", "质量优先  证据闭环",
                               "问题有依据  结论可解释", "建议能落地  整改有追踪"]):
        d.ellipse([930, 460+i*48, 946, 476+i*48], fill=rga(GD))
        d.text((965, 458+i*48), item, font=font(26), fill=WH)
    # Bottom decorative elements
    d.ellipse([W-300, H-400, W-50, H-150], fill=rga(GD, 15))
    d.ellipse([W-250, H-350, W-100, H-200], fill=rga(TL, 15))
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text((130, H-40), "四川融策会计师事务所有限公司", font=font(20), fill=rga(GD, 180))
    return img

# =========== PDF Generation ===========
def generate_pdf(images):
    c = canvas.Canvas(str(PDF), pagesize=A4)
    pw, ph = A4
    for img_path in images:
        im = Image.open(img_path)
        im = im.resize((int(pw), int(ph)-20), LANCZOS)
        tmp = WK / "tmp_render.png"
        im.save(tmp)
        c.drawImage(str(tmp), 0, 10, width=pw, height=ph-20, preserveAspectRatio=False)
        tmp.unlink(missing_ok=True)
        c.showPage()
    c.save()

# =========== DOCX Generation ===========
def generate_docx(images):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    lbls = ["封面", "关于融策", "核心方法", "服务体系", "政府审计",
            "预算绩效", "工程咨询", "数字化", "代表经验", "合作价值"]
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v4清新风"
    for idx, (lbl, img_path) in enumerate(zip(lbls, images), 1):
        if idx > 1: doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(lbl)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(22); r.font.bold = True
        doc.add_picture(str(img_path), width=Cm(17.0))
    doc.save(DOCX)

# =========== Main ===========
def main():
    OUT.mkdir(parents=True, exist_ok=True)
    WK.mkdir(parents=True, exist_ok=True)
    makers = [p1_cover, p2_about, p3_method, p4_services,
              p5_gov, p6_perf, p7_eng, p8_digital,
              p9_experience, p10_contact]
    print("Rendering 10 pages...")
    images = []
    for i, maker in enumerate(makers, 1):
        out = WK / f"page_{i:02d}.png"
        if not out.exists():
            print(f"  Page {i}...")
            maker().save(out, quality=95)
        images.append(out)
    print("Generating PDF...")
    generate_pdf(images)
    print("Generating DOCX...")
    generate_docx(images)
    print(f"DONE\n{DOCX}\n{PDF}")

if __name__ == "__main__":
    main()
'''

p = Path(r"C:\Users\scrccpa\.openclaw\workspace\scripts\rongce_v4_brochure.py")
with open(p, "a", encoding="utf-8") as f:
    f.write(tail)
print(f"Appended. Size: {p.stat().st_size} bytes")

import py_compile
try:
    py_compile.compile(str(p), doraise=True)
    print("SYNTAX OK")
except py_compile.PyCompileError as e:
    print(f"ERROR: {e}")