# -*- coding: utf-8 -*-
"""在 v3 文档中补充文字内容，使总字数达到 12000+ """
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, copy
from lxml import etree

doc_path = r'D:\openclaw-workspace\output\模拟案例一_绩效目标编制辅导与审核工作方案.docx'
doc = Document(doc_path)

def bd(text, font='仿宋', size=14):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(size*2)
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    para.paragraph_format.line_spacing = 1.5

def hd(text, level=2):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = '黑体'
    run.font.size = Pt({2:15,3:14}.get(level,14))
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

def find_and_insert_after(doc, keyword, new_paragraphs_fn):
    """找到包含关键词的段落索引，在其后面插入新段落"""
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            target_idx = i
            break
    if target_idx is None:
        print(f'  [未找到] 关键词: {keyword}')
        return
    
    # 收集 target 之后的所有元素
    body = doc.element.body
    target_elem = doc.paragraphs[target_idx]._element
    
    # 获取 target 之后下一个段落的索引（即要在哪里插入）
    insert_after_elem = target_elem
    next_sibling = insert_after_elem.getnext()
    
    # 使用临时document构建要插入的段落
    tmp_doc = Document()
    new_paragraphs_fn(tmp_doc)
    
    # 将临时文档的段落元素插入到 target 之后
    for tmp_para in tmp_doc.paragraphs:
        cloned = copy.deepcopy(tmp_para._element)
        insert_after_elem.addnext(cloned)
        insert_after_elem = cloned
    
    print(f'  已在 "{keyword[:30]}..." 后插入{len(tmp_doc.paragraphs)}段')

# ============================================================
# 1. 在 "第一阶段：前置辅导" 后补充详细内容
# ============================================================
def add_prep_detail(tmp_doc):
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('（1）政策解读与培训：组织区交通运输局财务及业务科室负责人，集中培训1天（约6学时）。培训内容涵盖：预算绩效管理政策演进脉络（从预算法修订到34号文的十年改革历程）、\"管理效率-履职效能\"二维框架的构建方法、三级指标的分类与命名规范、SMART指标值设定原则、常见编制误区及纠正方法（如\"指标值越好看越好\"的错误倾向）、绩效目标编制说明的撰写要点等六个模块。培训采用\"PPT讲解+案例研讨+实操演练\"三段式设计，每个模块均设有互动环节，确保参训人员不仅\"听得懂\"更能\"上手做\"。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('（2）部门职能梳理：结合区交通运输局\"三定\"方案和年度工作计划，对部门核心职能进行结构化梳理。区交通运输局承担公路建设与养护管理、道路运输市场监管、交通安全监督管理、交通基础设施建设管理四大核心职能。通过建立\"职能→任务→指标\"的三级映射关系，确保每一项核心职能在绩效目标体系中都有对应的指标来衡量。例如：\"公路养护管理\"职能映射到\"路面使用性能指数(PQI)\"和\"单位公路养护成本\"两项指标；\"交通建设管理\"职能映射到\"项目完工及时率\"\"工程验收合格率\"\"年度完工项目数量\"等多项指标。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('（3）历史数据分析：收集区交通运输局2023-2025年度部门决算报表、绩效自评报告、审计报告及整改情况等三类历史资料。重点分析四个维度的数据：预算执行偏离度趋势（近三年预算偏离度的均值和方差）、项目完工及时率变化（是否有季节性规律或年度改善趋势）、资金使用效率（人均经费、单位养护成本等的变动方向）、审计发现问题类型及整改率。通过历史数据回溯，为后续指标参考值设定和指标值合理性判断提供数据支撑。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('（4）资料清单交付：向区交通运输局提供分类资料清单，分为A类（必需资料：预算批复文件、决算报表、\"三定\"方案、年度工作计划）、B类（补充资料：政府采购计划、项目实施进度表、内控制度汇编、在编人员名册）、C类（参考资料：上级主管部门考核要求、历史年度绩效评价报告、审计整改台账）三个层级，逐项标注用途和对接科室，在项目启动会上完成交接签收。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

find_and_insert_after(doc, '第一阶段：前置辅导', add_prep_detail)

# ============================================================
# 2. 在 "四维审核法" 后补充审核标准细节
# ============================================================
def add_audit_detail(tmp_doc):
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('四维审核法的审核清单共22项检查要点，逐项评分后汇总形成审核意见。其中：完整性维度5项（是否覆盖全部核心职能、是否包含管理效率与履职效能、三级指标是否完整、是否有指标编制说明、是否有数据来源标注）；科学性维度6项（指标分类是否有误、层级关系是否正确、是否存在交叉重复、指标名称是否准确、指标值设定是否有逻辑依据、定性指标是否有等级标准）；合理性维度6项（指标值是否有历史数据支撑、目标值是否切合实际、参考值选取是否有依据、指标数量是否均衡、权重分配是否合理、定性指标标准是否清晰）；可操作性维度5项（指标是否可量化或可分级、数据来源是否明确可靠、数据采集是否可行、评价标准是否无歧义、是否可年度对比）。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

find_and_insert_after(doc, '图5："四维审核"方法体系', add_audit_detail)

# ============================================================
# 3. 在问题诊断开篇 补充审核思路说明
# ============================================================
def add_diagnosis_intro(tmp_doc):
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('本次审核采用\"由表及里、由整体到局部\"的递进式诊断路径：第一步，通过雷达图做整体评分，快速锁定指标体系的薄弱维度；第二步，按七维诊断框架逐项深入，每个维度均按照\"现象描述→问题定性→政策依据→影响分析\"四个层次展开；第三步，将七个维度的问题汇总为映射关系图，确保修改建议的系统性和针对性。整个诊断过程遵循\"独立、客观、证据导向\"的原则，每一个问题判断均对应到具体的政策条款或行业规范，避免主观臆断。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('此外，需要特别说明的是，本案例分析关注的是指标体系的内在质量——即指标设置本身是否科学合理——而非对区交通运输局实际工作绩效的评价。两者有本质区别：指标体系有缺陷不代表实际工作做得不好，但指标体系有缺陷将导致实际工作的绩效无法被客观、准确地衡量。这也是绩效目标编制审核工作的核心价值所在——在评价开始之前，先把\"尺子\"校准。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

find_and_insert_after(doc, '图9：七大问题→修改建议', add_diagnosis_intro)

# ============================================================
# 4. 在时间保障表后 补充实施风险与应对
# ============================================================
def add_risk_detail(tmp_doc):
    para = tmp_doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run('为确保项目按计划推进，需重点关注以下实施风险：一是资料交接延误风险——如区交通运输局未能按时提供所需历史数据和佐证材料，将影响审核工作的数据比对环节。应对措施：在项目启动阶段一次性提供详细资料清单，设置资料提交截止日（第1周末），逾期未提交的资料在审核意见书中标注\"因数据缺失无法验证\"。二是审核意见争议风险——对于审核发现的部分问题（尤其是涉及业务判断的专业指标），可能存在双方认知差异。应对措施：在沟通会议前，审计团队内部先进行\"争议预判\"，对每个可能引发争议的问题准备2-3套替代方案，提高沟通效率。三是修改反复风险——可能出现\"改了又改\"的情况，导致工作周期失控。应对措施：明确修改轮次上限（原则上不超过两轮），第二轮修改后仍未达成一致的事项，提交区财政局裁决。')
    para.paragraph_format.line_spacing = 1.5
    run.font.name = '仿宋'; run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

find_and_insert_after(doc, '表6：项目时间进度安排表', add_risk_detail)

# ============================================================
# 保存
# ============================================================
doc.save(doc_path)

total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'补充完成，总字数：{total}字')
print(f'文件大小：{os.path.getsize(doc_path)/1024:.1f} KB')
