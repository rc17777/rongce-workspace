# -*- coding: utf-8 -*-
"""再补一波，确保 12000+ 字"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import os, copy

doc_path = r'D:\openclaw-workspace\output\模拟案例一_绩效目标编制辅导与审核工作方案.docx'
doc = Document(doc_path)

def bd(tmp, text):
    para = tmp.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run(text)
    run.font.name = '仿宋'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    para.paragraph_format.line_spacing = 1.5

def insert_after(doc, keyword, fn):
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            target_elem = para._element
            tmp = Document()
            fn(tmp)
            insert_after = target_elem
            for tmp_p in tmp.paragraphs:
                cloned = copy.deepcopy(tmp_p._element)
                insert_after.addnext(cloned)
                insert_after = cloned
            n = len(tmp.paragraphs)
            print(f'  OK {n}段 @ ...{keyword[-20:]}')
            return
    print(f'  MISS: {keyword[:30]}')

# 问题二补充
insert_after(doc, '4项指标占据了全部18项履职效能指标', lambda t: [
    bd(t, '需要特别强调的是，效益指标的量化改造并非要求所有指标都必须变成硬数字，有些公共管理效果确实难以直接用数字衡量。关键在于：即使不能直接量化，也必须建立科学、透明的评价规则，让定性判断变成规则下的判断而非拍脑袋的判断。例如促进农村地区经济发展，虽然不可能精确测算出交通资金贡献了GDP的百分之几，但可以通过农村公路沿线新增市场主体数量、农产品外运成本下降比例、沿线旅游收入增长率等替代指标，从多个侧面间接反映经济发展的促进效果。'),
])

# 问题三补充
insert_after(doc, '结构严重失衡', lambda t: [
    bd(t, '从绩效评价实务经验来看，管理效率指标过少会产生两个直接后果：一是在绩效评价时，评价重心的百分之八十以上集中在履职效能上，对资金使用规范性关注不足，容易形成只要事办成了怎么花钱都行的错误导向；二是在财政审核时，财政部门倾向于关注管理规范性和预算合规性，如果指标体系中这类指标太少，可能导致绩效目标在财政审核环节被退回或要求大幅修改。将管理效率指标从4项增加到12项，本质上是让运行管理和业务产出两个维度回到相对均衡的状态。'),
])

# 问题四补充
insert_after(doc, '直接导致后续绩效监控无法开展', lambda t: [
    bd(t, '此外，一般性支出金额的2023年和2024年数据标注为横线（而非xx），暗示这两年可能没有单独统计过一般性支出，或者统计口径发生过变化。如果确实如此，需要在指标编制说明中予以明确，并说明2026年指标值的测算基础（是参考2025年数据还是参考同级其他区县数据）。数据的断档本身不一定构成问题，但对断档原因不做说明、对补救措施不做交代，就构成了信息不完整的实质性缺陷。'),
])

# 汇总表前补充使用说明
insert_after(doc, '将所有修改建议汇总为以下三张表', lambda t: [
    bd(t, '区交通运输局在使用以下三张汇总表时，建议按照先分类调整、再量化改造、最后指标补充的顺序逐表落实。分类调整是结构性修正，决定了指标体系的骨架是否正确；量化改造是内容性修正，决定了每项指标的血肉是否丰满；指标补充是完整性修正，决定了指标体系的版图是否无遗漏。三张表形成一个递进式的修改工作清单，便于分工落实和进度跟踪。'),
])

# 结语前补充展望
insert_after(doc, '绩效目标是预算绩效管理的第一粒扣子', lambda t: [
    bd(t, '展望后续工作，本次辅导与审核的成果不应止步于一份绩效目标定稿。建议区交通运输局以本次工作为起点，将绩效目标管理融入日常预算管理流程：每季度对照绩效目标开展一次执行进度自查（绩效监控），年中根据实际情况对绩效目标进行一次动态调整（如需），年度终了时以绩效目标为基准开展自评（绩效自评），形成编制-监控-评价-反馈的完整绩效管理闭环。四川融策会计师事务所有限公司愿意在后续的绩效监控和绩效评价环节继续提供专业支持，助力区交通运输局建立长效化的预算绩效管理机制。'),
])

doc.save(doc_path)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
print(f'=== 最终文档 ===')
print(f'总字数: {total}字')
print(f'图片数: {img_count}张')
print(f'表格数: {len(doc.tables)}张')
print(f'文件大小: {os.path.getsize(doc_path)/1024:.1f} KB')
