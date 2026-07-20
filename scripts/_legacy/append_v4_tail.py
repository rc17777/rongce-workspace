# -*- coding: utf-8 -*-
"""追加 v4 脚本尾部"""
from pathlib import Path

tail = r'''
def page8_digital():
    return service_page("数字化审计能力", "DIGITAL AUDIT CAPABILITIES",
        "用数据扩大覆盖面、提高发现率、增强证据质量----把审计经验沉淀为可复用的数据工具。",
        [("数据标准", ["财务、预算、支付、合同", "采购、资产、工程项目字段整理"]),
         ("规则模型", ["重复支付、超预算执行识别", "供应商异常、资金沉淀检测"]),
         ("穿透核查", ["疑点来源、核查路径", "佐证材料、影响金额、整改建议"]),
         ("报告复核", ["金额汇总校验、口径一致性", "附表闭环、结论依据可追溯"])], 5)

def page9_experience():
    img = Image.new("RGB", (W, H), rgb(WHT))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, 18, H], fill=rgba(NAVY))
    d.rectangle([18, 0, 24, H], fill=rgba(GOLD))
    d.text((160, 100), "代表经验", font=fnt(52, True), fill=NAVY)
    d.text((160, 168), "REPRESENTATIVE EXPERIENCE", font=fnt(22), fill=rgba(GOLD))
    d.line([160, 205, 600, 205], fill=rgba(GOLD), width=3)
    wrapdraw(d, "长期服务省、市、县多级财政和主管部门，覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。", 160, 240, fnt(30), MUTED, 620, 10)
    d.rounded_rectangle([880, 130, 1450, 550], radius=20, fill=rgba(TEAL, 12), outline=rgba(TEAL, 30), width=2)
    d.text((960, 180), "服务版图", font=fnt(36, True), fill=rgba(TEAL, 60))
    for i, r in enumerate(["成都总部", "阿坝州办事处", "西藏办事处", "覆盖川、藏、黔"]):
        d.ellipse([960, 270+i*60, 990, 300+i*60], fill=rgba(GOLD))
        wrapdraw(d, r, 1005, 270+i*6, fnt(26), INK, 300, 4)
    d.text((160, 390), "代表客户", font=fnt(32, True), fill=NAVY)
    clients = ["四川省财政厅", "省公安厅交警总队", "省民政厅", "省农业农村厅", "省教育厅",
        "省退役军人事务厅", "省药品监督管理局", "达州市财政局", "绵阳市财政局",
        "宜宾市财政局", "什邡市财政局", "德阳市财政局", "康定市财政局", "九寨沟县财政局"]
    yv = 440
    for i, c in enumerate(clients):
        if i % 5 == 0 and i > 0: yv += 40
        cx = 160 + (i % 5) * 280
        d.rounded_rectangle([cx, yv, cx+240, yv+36], radius=6, fill=rgba(WHT), outline=rgba(GOLD, 80), width=1)
        tw = d.textlength(c, font=fnt(20))
        d.text((cx+120-tw//2, yv+5), c, font=fnt(20), fill=NAVY)
    d.rectangle([0, H-50, W, H], fill=rgba(NAVY))
    tx = "公开公正  用心服务  诚信为本  服务至上  追求卓越"
    d.text(((W-d.textlength(tx, font=fnt(18)))//2, H-42), tx, font=fnt(18), fill=rgba(WHT, 180))
    return img

def page10_contact():
    img = Image.new("RGB", (W, H), rgb(BGE))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, H], fill=rgba(NAVY))
    d.rectangle([0, 0, W, 280], fill=rgba(NAVY))
    d.rectangle([0, 280, W, 290], fill=rgba(GOLD))
    d.text((130, 80), "合作价值", font=fnt(52, True), fill=WHT)
    d.text((130, 150), "VALUE PROPOSITION", font=fnt(22), fill=rgba(GOLD, 200))
    d.line([130, 190, 500, 190], fill=rgba(GOLD), width=3)
    wrapdraw(d, "我们交付的不只是一份报告，更是一套可执行的改进方案。", 130, 260, fnt(30), WHT, 600, 10)
    d.rounded_rectangle([130, 340, 750, 540], radius=16, fill=rgba(WHT, 30), outline=rgba(GOLD, 80), width=2)
    wrapdraw(d, "联系融策", 180, 370, fnt(36, True), GOLD, 500, 8)
    for i, line in enumerate(["028-87659276", "scrccpa@163.com", "www.scrccpa.com", "成都市金牛区金周路595号"]):
        d.ellipse([180, 440+i*48, 196, 456+i*48], fill=rgba(GOLD))
        d.text((215, 438+i*48), line, font=fnt(26), fill=WHT)
    d.rounded_rectangle([880, 340, 1530, 750], radius=16, fill=rgba(WHT, 30), outline=rgba(GOLD, 80), width=2)
    wrapdraw(d, "合作承诺", 930, 380, fnt(36, True), GOLD, 500, 8)
    for i, item in enumerate(["客观公正  实事求是", "质量优先  证据闭环", "问题有依据  结论可解释", "建议能落地  整改有追踪"]):
        d.ellipse([930, 460+i*48, 946, 476+i*48], fill=rgba(GOLD))
        d.text((965, 458+i*48), item, font=fnt(26), fill=WHT)
    d.rectangle([0, H-50, W, H], fill=rgba(NAVY))
    d.text((130, H-40), "四川融策会计师事务所有限公司", font=fnt(20), fill=rgba(GOLD, 180))
    return img

def generate_pdf(images):
    c = canvas.Canvas(str(PDF_), pagesize=A4)
    pw, ph = A4
    from PIL.Image import LANCZOS
    for img_path in images:
        im = Image.open(img_path)
        im = im.resize((int(pw), int(ph)-20), LANCZOS)
        tmp = WK / "tmp_render.png"
        im.save(tmp)
        c.drawImage(str(tmp), 0, 10, width=pw, height=ph-20, preserveAspectRatio=False)
        tmp.unlink(missing_ok=True)
        c.showPage()
    c.save()

def generate_docx(images):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    lbls = ["封面", "关于融策", "核心方法", "服务体系", "政府审计", "预算绩效", "工程咨询", "数字化", "代表经验", "合作价值"]
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v4清新风"
    for idx, (lbl, img_path) in enumerate(zip(lbls, images), 1):
        if idx > 1: doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(lbl)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(22); r.font.bold = True
        doc.add_picture(str(img_path), width=Cm(17.0))
    doc.save(DOCX)

def main():
    OUT.mkdir(parents=True, exist_ok=True)
    WK.mkdir(parents=True, exist_ok=True)
    makers = [page1_cover, page2_about, page3_method, page4_services,
              page5_gov, page6_perf, page7_eng, page8_digital,
              page9_experience, page10_contact]
    print("Rendering 10 pages...")
    images = []
    for i, maker in enumerate(makers, 1):
        out = WK / f"page_{i:02d}.png"
        if not out.exists():
            print(f"  Page {i}...")
            maker().save(out, quality=95)
        images.append(out)
    print("PDF...")
    generate_pdf(images)
    print("DOCX...")
    generate_docx(images)
    print(f"DONE\n{DOCX}\n{PDF_}")

if __name__ == "__main__":
    main()
'''

if __name__ == "__main__":
    p = Path(r"C:\Users\scrccpa\.openclaw\workspace\scripts\generate_rongce_brochure_v4.py")
    with open(p, "a", encoding="utf-8") as f:
        f.write(tail)
    print(f"Appended. File size: {p.stat().st_size} bytes")
    lines = p.read_text(encoding="utf-8").splitlines()
    print(f"Total lines: {len(lines)}")