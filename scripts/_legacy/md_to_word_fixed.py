# -*- coding: utf-8 -*-
"""
修复版：Markdown转Word，解决表格乱码问题
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

input_dir = r'C:\Users\scrccpa\.openclaw\workspace\chayu_analysis'
output_dir = input_dir

md_files = [
    '资产确权情况说明模板.md',
    '四个一批项目资料齐全率统计报告.md',
    '低效闲置资产盘活方案.md',
]

def set_cell_border(cell):
    """设置单元格边框"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def md_to_docx(md_file, output_file):
    """将Markdown转Word（修复版）"""
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # 读取Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 解析Markdown
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 一级标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue
        
        # 二级标题
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            for run in p.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue
        
        # 三级标题
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            for run in p.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue
        
        # 四级标题
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            for run in p.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue
        
        # 分隔线
        if line.strip() == '---':
            p = doc.add_paragraph('─' * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 引用块
        if line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.font.italic = True
            i += 1
            continue
        
        # 表格检测
        if line.startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tl = lines[i].strip()
                # 跳过分隔行
                if not re.match(r'^\|[\s\-:]+\|', tl):
                    table_lines.append(tl)
                i += 1
            
            if table_lines:
                # 解析表格
                rows = []
                for tl in table_lines:
                    # 分割单元格
                    cells = [c.strip() for c in tl.split('|')]
                    # 去掉首尾空元素
                    cells = [c for c in cells if c]
                    if cells:
                        rows.append(cells)
                
                if rows and len(rows) > 0:
                    # 创建Word表格
                    max_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(rows):
                        for c_idx in range(max_cols):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
                            
                            # 设置单元格文本
                            cell.text = cell_text
                            
                            # 设置字体
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                    run.font.size = Pt(10.5)
                                    # 表头加粗
                                    if r_idx == 0:
                                        run.font.bold = True
                            
                            # 设置边框
                            set_cell_border(cell)
                    
                    # 添加表格后的空行
                    doc.add_paragraph()
            
            continue
        
        # 无序列表
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue
        
        # 有序列表
        match = re.match(r'^(\d+)\.\s+(.+)', line)
        if match:
            text = match.group(2)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue
        
        # 普通段落
        p = doc.add_paragraph()
        
        # 处理加粗标记 **text**
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                run = p.add_run(part)
            
            # 设置字体
            if p.runs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        i += 1
    
    # 保存
    doc.save(output_file)
    print(f'✅ 已转换: {os.path.basename(output_file)}')

# 转换三个文件
print('开始转换...\n')
import time
for md_file in md_files:
    md_path = os.path.join(input_dir, md_file)
    # 加时间戳避免文件占用
    docx_file = md_file.replace('.md', f'_{int(time.time())}.docx')
    docx_path = os.path.join(output_dir, docx_file)
    
    if os.path.exists(md_path):
        md_to_docx(md_path, docx_path)
    else:
        print(f'❌ 文件不存在: {md_file}')

print(f'\n全部完成！输出目录: {output_dir}')
