#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Deep extract 顺华 pricing tables - check for images/form fields"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from pathlib import Path

path = Path(r'C:\Users\scrccpa\Desktop\校服\2025年校服采购\校服\2025年\投标文件\投标文件\成都顺华服装有限公司\其他投标文件.doc')
doc = Document(str(path))

for ti, table in enumerate(doc.tables):
    first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
    if any(kw in first_row_text for kw in ['投标报价', '分项报价', '单价', '开标一览']):
        print(f'\n=== Table {ti+1} ===')
        for ri, row in enumerate(table.rows):
            cells_info = []
            for ci, cell in enumerate(row.cells):
                cell_xml = cell._element.xml
                has_image = 'pic:pic' in cell_xml or 'w:drawing' in cell_xml or 'v:imagedata' in cell_xml
                has_sdt = 'w:sdt' in cell_xml  # structured document tag / content control
                text = cell.text.strip()
                flags = []
                if has_image: flags.append('IMG')
                if has_sdt: flags.append('SDT')
                flag_str = f' [{",".join(flags)}]' if flags else ''
                cells_info.append(f'C{ci}:{text}{flag_str}')
            print(f'  Row{ri}: {" | ".join(cells_info)}')
            
        # Also check paragraphs in cells for pricing
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            if any(kw in run.text for kw in ['元', '报价', '价']):
                                print(f'  [Run in T{ti+1}R{ri}C{ci}]: {run.text[:100]}')

# Also check document for SDT (structured document tags - form fields)
print('\n=== SDT/Content Controls ===')
for i, para in enumerate(doc.paragraphs):
    p_xml = para._element.xml
    if 'w:sdt' in p_xml:
        print(f'  Para{i}: {para.text[:100]}')

# Also scan ALL paragraphs for any pricing numbers
print('\n=== All paragraphs with numbers + yuan ===')
import re
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if re.search(r'\d+.*元', t):
        print(f'  P{i}: {t[:150]}')
