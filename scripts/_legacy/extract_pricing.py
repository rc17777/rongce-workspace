#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从校服投标文件DOCX中提取报价和关键信息"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from pathlib import Path

BASE = Path(r"C:\Users\scrccpa\Desktop\校服\2025年校服采购\校服\2025年\投标文件\投标文件")

# Target: other/business bid files
targets = [
    (BASE / "四川乐吉玛帝诺服饰有限公司" / "成都市教育科学研究院附属中学学生校服采购项目其他投标文件.docx", "乐吉玛帝诺"),
    (BASE / "四川牧森服饰有限公司" / "成都市教育科学研究院附属中学其他投标文件-最终版.docx", "牧森"),
    (BASE / "江苏苏美达伊顿纪德品牌管理有限公司" / "20251012【其他投标文件】科教院附中.docx", "苏美达伊顿纪德"),
]

def find_price_tables(doc, company):
    """Extract price-related tables and text"""
    results = []
    
    # First, find price-related paragraphs
    price_keywords = ['报价', '价格', '单价', '金额', '元/套', '分项报价', '报价表', '开标一览表',
                      '投标报价', '报价明细', '总计', '合计', '全套', '夏装', '春秋', '冬装']
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # Check if contains price keywords
        if any(kw in text for kw in price_keywords):
            # Get context (before and after)
            ctx_start = max(0, i-2)
            ctx_end = min(len(doc.paragraphs), i+5)
            ctx_lines = []
            for j in range(ctx_start, ctx_end):
                t = doc.paragraphs[j].text.strip()
                if t:
                    ctx_lines.append(t)
            results.append(('\n'.join(ctx_lines), 'paragraph'))
    
    # Second, extract all tables that might contain pricing
    for ti, table in enumerate(doc.tables):
        # Check if table contains price keywords
        table_text = []
        for row in table.rows[:3]:  # check first 3 rows
            cells_text = [cell.text.strip() for cell in row.cells]
            table_text.append(' | '.join(cells_text))
        
        table_str = '\n'.join(table_text)
        if any(kw in table_str for kw in price_keywords + ['元', '￥', '¥', '金额']):
            # Extract full table
            all_rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                all_rows.append(' | '.join(cells))
            results.append((f"[Table {ti+1}]\n" + '\n'.join(all_rows), 'table'))
    
    return results

for path, company in targets:
    if not path.exists():
        print(f"\n{'='*60}")
        print(f"❌ {company}: File not found - {path}")
        continue
    
    print(f"\n{'='*60}")
    print(f"📄 {company} - 商务标提取")
    print(f"   文件大小: {path.stat().st_size/1024/1024:.1f} MB")
    
    try:
        doc = Document(str(path))
        print(f"   段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
        
        results = find_price_tables(doc, company)
        
        if not results:
            # Scan first 200 paragraphs for structure
            print("   ⚠️ 未找到价格关键词，显示前50个有效段落的结构:")
            count = 0
            for para in doc.paragraphs:
                t = para.text.strip()
                if t and count < 50:
                    print(f"   [{count}] {t[:120]}")
                    count += 1
        else:
            print(f"   找到 {len(results)} 个相关段落/表格:")
            for r, rtype in results[:10]:  # limit to 10
                print(f"\n   --- {rtype} ---")
                for line in r.split('\n')[:30]:
                    print(f"   {line[:150]}")
                    
    except Exception as e:
        print(f"   ❌ 提取失败: {e}")

print("\nDone!")
