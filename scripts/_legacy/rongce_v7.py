# -*- coding: utf-8 -*-
"""四川融策宣传册 v7 —— 独立视觉版
封面重做 + 每页独立设计，不再使用统一模板。
"""
from __future__ import annotations
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from PIL.Image import LANCZOS
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

OUT = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
WK = Path("work/sichuan_rongce_brochure/v7_assets")
DOCX = OUT / "四川融策宣传册_完善稿_v7_独立视觉版.docx"
PDF = OUT / "四川融策宣传册_完善稿_v7_独立视觉版.pdf"

NW = "#1A365D"; TL = "#3A7B8A"; GD = "#D4A574";
MU = "#718096"; IK = "#2D3748"; WH = "#FFFFFF"; BG = "#FAF8F4"
W, H = 1653, 2339
M = 160  # margin

def rgb(s):
    h = s.lstrip("#"); return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
def rga(s, a=255): return (*rgb(s), a)

def font(size, bold=False):
    for p in [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if Path(p).exists(): return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()

def draw(d, text, x, y, f, fill, w=None, gap=6, align="left"):
    mw = w or 9999; lines = []
    for para in text.split("\n"):
        cur = ""
        for ch in para:
            if d.textlength(cur + ch, font=f) <= mw: cur += ch
            else:
                if cur: lines.append(cur); cur = ch
        if cur: lines.append(cur)
    for line in lines:
        xx = x
        if align == "center" and w: xx = x + (w - d.textlength(line, font=f)) // 2
        d.text((xx, y), line, font=f, fill=fill); y += f.size + gap
    return y

def left_bar(d):
    d.rectangle([0, 0, 18, H], fill=rga(NW)); d.rectangle([18, 0, 24, H], fill=rga(GD))

def bottom_bar(d):
    tx = "公开公正  用心服务  诚信为本  服务至上  追求卓越"
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text(((W-d.textlength(tx, font=font(18)))//2, H-42), tx, font=font(18), fill=rga(WH, 180))

def page_title(d, title, en, y=120):
    d.text((M, y), title, font=font(52, True), fill=NW)
    d.text((M, y+68), en, font=font(22), fill=rga(GD))
    d.line([M, y+105, M+440, y+105], fill=rga(GD), width=3)

# ===================== P1: COVER (redesigned) =====================
def p1_cover():
    img = Image.new("RGB", (W, H), rgb(NW))
    d = ImageDraw.Draw(img, "RGBA")
    # Background geometric pattern
    for r in range(10):
        d.ellipse([W-600+r*35, -200+r*35, W+200+r*35, 600+r*35], fill=rga(TL, 3))
    for r in range(6):
        d.ellipse([W-500+r*40, 300+r*40, W+100+r*40, 900+r*40], fill=rga(TL, 2))
    # Diamond pattern top-right
    cx, cy = W-250, 200
    for i in range(4):
        s = 100 - i*20
        pts = [(cx, cy-s*2), (cx+s, cy), (cx, cy+s*2), (cx-s, cy)]
        d.polygon(pts, fill=rga(GD, 15-i*3), outline=rga(GD, 30-i*5))
    # Brand block
    d.rectangle([0, 0, W, 350], fill=rga(NW))
    d.rectangle([0, 350, W, 360], fill=rga(GD))
    d.text((M, 90), "SICHUAN", font=font(72, True), fill=rga(WH, 180))
    d.text((M, 165), "RONGCE", font=font(110, True), fill=WH)
    # Large center statement
    draw(d, "谋专业之策", M+20, 480, font(64, True), WH, 700, 16)
    draw(d, "融品质之精", M+20, 560, font(64, True), WH, 700, 16)
    d.line([M+20, 660, M+420, 660], fill=rga(GD), width=5)
    draw(d, "政府审计与工程咨询综合服务机构", M+20, 700, font(32), rga(WH, 200), 700, 10)
    # Bottom info
    d.rectangle([0, H-70, W, H], fill=rga(NW))
    draw(d, "审计 \u00b7 绩效 \u00b7 财政监督 \u00b7 工程咨询 \u00b7 数字化分析", M, H-52, font(22), rga(WH, 180), 800, 0, "center")
    # Corner accents
    d.line([W-100, 60, W-30, 60], fill=rga(GD, 80), width=3)
    d.line([W-30, 60, W-30, 130], fill=rga(GD, 80), width=3)
    # Bottom-right decorative triangle
    d.polygon([(W-100, H-100), (W-40, H-100), (W-100, H-40)], fill=rga(GD, 25))
    d.polygon([(W-100, H-100), (W-55, H-100), (W-100, H-55)], fill=rga(GD, 40))
    return img

# ===================== P2: WHY RONGCE =====================
def p2_about():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "为什么是融策？", "WHY RONGCE")
    # Faded page number
    d.text((W-250, 0), "02", font=font(240, True), fill=rga(NW, 6))
    # Headline
    y = draw(d, "财政资金的管理，核心是两个问题：", M, 280, font(36), IK, 700, 12)
    y = draw(d, "\u201c资金到底去哪了？\u201d 和 \u201c这钱花得值不值？\u201d", M, y+4, font(36, True), NW, 700, 12)
    y = draw(d, "\n融策做的事情，就是用专业的方法和数据工具，帮您把这两个问题搞清楚。", M, y+20, font(28), MU, 700, 10)
    # Three cards
    for i, (icon, head, desc) in enumerate([
        ("\u2460", "不只查账", "合同 · 项目 · 资产\n内控制度 · 决策程序\n账表只是入口"),
        ("\u2461", "不只找问题", "能不能改 · 怎么改\n改得怎么样\n整改落地才是目标"),
        ("\u2462", "不只靠经验", "数据分析+现场核查\n用数据扩大覆盖面\n用现场核实关键点"),
    ]):
        x = M + i * 380
        d.rounded_rectangle([x, 470, x+340, 690], radius=14, fill=rga("#F0EDE8"), outline=rga(TL, 50), width=2)
        d.rounded_rectangle([x, 470, x+340, 555], radius=14, fill=rga(TL, 20))
        d.text((x+28, 495), icon, font=font(36, True), fill=NW)
        d.text((x+76, 500), head, font=font(30, True), fill=NW)
        draw(d, desc, x+28, 580, font(24), MU, 320, 7)
    # Company brief
    y = 740
    d.rounded_rectangle([M, 730, W-M, 870], radius=14, fill=rga("#E8E0D4", 60))
    for item in [
        "始于2000年，四川省内较早成立的会计师事务所之一",
        "长期服务财政、审计、教育、民政、交通、国资、医保等领域",
        "审计 + 绩效 + 财政监督 + 工程咨询协同发展",
        "覆盖四川、西藏、贵州三省"
    ]:
        d.ellipse([180, y+8, 192, y+20], fill=rga(GD))
        y = draw(d, item, 210, y, font(26), IK, 1250, 10)
    # Decorative diamond cluster bottom-right
    for j in range(3):
        s = 20 - j*5; x0 = W-100; y0 = H-250 + j*40
        d.polygon([(x0, y0-s), (x0+s, y0), (x0, y0+s), (x0-s, y0)], fill=rga(TL, 20-j*5))
    bottom_bar(d)
    return img

# ===================== P3: METHODOLOGY =====================
def p3_method():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "我们怎么做？", "OUR METHODOLOGY")
    d.text((W-250, 0), "03", font=font(240, True), fill=rga(NW, 6))
    # Three-line cards with improved styling
    for i, (num, title, items) in enumerate([
        ("01", "资金线", ["钱从哪来？到哪去？", "花得值不值？", "三流交叉核验"]),
        ("02", "项目线", ["立项·招标·合同", "施工·验收·结算", "全生命周期穿透"]),
        ("03", "责任线", ["决策·岗位·内控", "整改落实追踪", "结论有据·责任可追"]),
    ]):
        x = M + i * 380
        d.rounded_rectangle([x, 280, x+340, 580], radius=16, fill=rga("#F8F6F2"), outline=rga(TL, 60), width=2)
        d.rounded_rectangle([x, 280, x+340, 370], radius=16, fill=rga(TL, 30))
        d.text((x+32, 305), num, font=font(40, True), fill=rga(NW, 80))
        d.text((x+80, 315), title, font=font(32, True), fill=NW)
        yy = 400
        for item in items:
            d.ellipse([x+32, yy+8, x+44, yy+20], fill=rga(GD))
            yy = draw(d, item, x+55, yy, font(24), IK, 340-70, 6)
    # Connecting arrow between cards
    for i in range(2):
        x1 = M + 380*(i+1) - 40
        d.line([x1, 430, x1+80, 430], fill=rga(GD, 60), width=2)
        d.polygon([(x1+70, 425), (x1+80, 430), (x1+70, 435)], fill=rga(GD, 60))
    # Five-step flow
    steps = ["制度审查", "数据核验", "现场核查", "证据闭环", "整改建议"]
    d.line([240, 660, 1410, 660], fill=rga(TL, 50), width=3)
    for i, step in enumerate(steps):
        x = 220 + i * 260
        d.ellipse([x, 620, x+80, 700], fill=rga(NW if i%2==0 else TL))
        tw = d.textlength(str(i+1), font=font(32, True))
        d.text((x+40-tw//2, 638), str(i+1), font=font(32, True), fill=WH)
        draw(d, step, x+100, 638, font(26), IK, 150, 6)
    # Closing statement
    d.rounded_rectangle([M, 810, W-M, 920], radius=14, fill=rga("#E8E0D4", 80))
    draw(d, "我们的目标不是出具一份报告，而是帮助委托方看清问题、厘清责任、找到改进路径。", M+40, 840, font(28), IK, 1250, 10, "center")
    bottom_bar(d)
    return img

# ===================== P4: SERVICE OVERVIEW =====================
def p4_services():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "服务体系", "SERVICE SYSTEM")
    d.text((W-250, 0), "04", font=font(240, True), fill=rga(NW, 6))
    draw(d, "从单项审计到全过程咨询，从结果评价到制度优化——五线并进，协同服务。", M, 240, font(30), MU, 650, 10)
    services = [
        ("政府审计", "经责 · 专项\n财政监督 · 工程决算", rga(NW)),
        ("预算绩效", "事前 · 目标\n监控 · 评价 · 应用", rga(TL)),
        ("工程咨询", "预算 · 评审\n结算 · 全过程", rga(NW)),
        ("采购审计", "程序合规\n围标串标 · 履约", rga(TL)),
        ("管理咨询", "内控 · 资产\n整改 · 流程优化", rga(NW)),
    ]
    for i, (title, desc, clr) in enumerate(services):
        x = M + i * 280
        d.rounded_rectangle([x, 310, x+250, 480], radius=14, fill=rga("#F8F6F2"), outline=rga(TL, 30), width=2)
        d.rounded_rectangle([x, 310, x+250, 370], radius=14, fill=clr)
        d.text((x+125 - d.textlength(title, font=font(32, True))//2, 320), title, font=font(32, True), fill=WH)
        draw(d, desc, x+20, 390, font(22), MU, 210, 6, "center")
    # Connecting line
    d.line([M+20, 490, W-M-20, 490], fill=rga(TL, 30), width=2)
    # Value statement
    d.rounded_rectangle([M, 540, W-M, 660], radius=14, fill=rga("#E8E0D4", 60))
    draw(d, "五条业务线形成三种协同效应：审计发现的问题，绩效评价提供改进方向，工程咨询落实项目执行，管理咨询固化制度成果。", M+40, 570, font(28), IK, 1300, 8, "center")
    bottom_bar(d)
    return img

# ===================== P5: GOVERNMENT AUDIT =====================
def p5_gov():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "政府审计", "GOVERNMENT AUDIT")
    d.text((W-250, 0), "05", font=font(240, True), fill=rga(NW, 6))
    # Three-layer narrative
    draw(d, "从\u201c有没有\u201d到\u201c对不对\u201d再到\u201c值不值\u201d——把账表、合同、项目、资金、资产、责任贯通核查。", M, 240, font(30), MU, 650, 10)
    # Three layers as vertical columns
    layers = [
        ("\u201c有没有\u201d", "合规性审查", ["制度是否完善", "程序是否合规", "记录是否完整"], NW),
        ("\u201c对不对\u201d", "准确性核查", ["数据是否准确", "金额是否一致", "归属是否清晰"], TL),
        ("\u201c值不值\u201d", "绩效性评价", ["投入是否合理", "产出是否达标", "效果是否持续"], GD),
    ]
    for i, (tag, label, items, clr) in enumerate(layers):
        x = M + i * 380
        d.rounded_rectangle([x, 310, x+340, 520], radius=14, fill=WH, outline=rga(clr, 60), width=2)
        d.rounded_rectangle([x, 310, x+340, 370], radius=14, fill=rga(clr, 40))
        d.text((x+30, 325), tag, font=font(28, True), fill=rgb(clr))
        d.text((x+32, 355), label, font=font(22), fill=rga(clr, 180))
        yy = 400
        for item in items:
            d.ellipse([x+30, yy+8, x+42, yy+20], fill=rga(GD))
            yy = draw(d, item, x+52, yy, font(24), IK, 280, 6)
    # Arrow between layers
    for i in range(2):
        x1 = M + 380*(i+1) - 50
        d.line([x1, 415, x1+60, 415], fill=rga(GD, 60), width=2)
        d.polygon([(x1+50, 410), (x1+60, 415), (x1+50, 420)], fill=rga(GD, 60))
    # Service areas grid
    d.text((M, 560), "服务领域", font=font(32, True), fill=NW)
    d.text((M, 600), "SERVICE AREAS", font=font(18), fill=rga(GD))
    areas = [("经济责任审计", "重大决策 · 资金安全 · 廉政风险"),
             ("专项资金审计", "申报分配 · 拨付使用 · 绩效评价"),
             ("财政监督检查", "预算执行 · 财经纪律 · 采购合规"),
             ("工程决算财务审计", "成本归集 · 资金来源 · 资产交付")]
    for i, (name, desc) in enumerate(areas):
        x = M + (i%2)*680
        y0 = 640 + (i//2)*90
        d.rounded_rectangle([x, y0, x+640, y0+72], radius=8, fill=WH, outline=rga(GD, 60), width=1)
        d.rectangle([x, y0, x+10, y0+72], fill=rga(NW))
        draw(d, name, x+24, y0+8, font(28, True), NW, 600, 4)
        draw(d, desc, x+24, y0+40, font(22), MU, 600, 4)
    d.line([M, H-65, W-M, H-65], fill=rga(TL, 20), width=2)
    bottom_bar(d)
    return img

# ===================== P6: BUDGET PERFORMANCE =====================
def p6_perf():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "预算绩效管理", "BUDGET PERFORMANCE")
    d.text((W-250, 0), "06", font=font(240, True), fill=rga(NW, 6))
    draw(d, "您关心的不只是\u201c花了多少钱\u201d，更是\u201c效果怎么样\u201d——让财政资金从\u201c花了没有\u201d走向\u201c花得值不值\u201d。", M, 240, font(30), MU, 650, 10)
    # Cycle layout - 5 nodes in a circle-like arrangement
    steps = [
        ("事前评估", "必要性 · 可行性\n财政承受能力", (M+20, 380)),
        ("目标审核", "完整性 · 可衡量\n绩效责任书", (M+380, 340)),
        ("运行监控", "进度追踪 · 偏差预警\n资金支付监控", (M+730, 380)),
        ("重点评价", "政策 · 部门 · 项目\n专项资金评价", (M+380, 500)),
        ("结果应用", "整改清单 · 预算挂钩\n制度优化", (M+20, 500)),
    ]
    # Center node
    d.ellipse([M+305, 410, M+505, 470], fill=rga(NW))
    draw(d, "绩效管理", M+405 - d.textlength("绩效管理", font=font(28, True))//2, 420, font(28, True), WH, 150, 6, "center")
    draw(d, "闭环", M+405 - d.textlength("闭环", font=font(22))//2, 445, font(22), rga(WH, 180), 150, 0, "center")
    # Nodes
    for i, (name, desc, (nx, ny)) in enumerate(steps):
        d.rounded_rectangle([nx, ny, nx+230, ny+70], radius=10, fill=WH, outline=rga(TL, 50), width=2)
        draw(d, f"0{i+1} {name}", nx+14, ny+6, font(26, True), NW, 200, 4)
        draw(d, desc, nx+14, ny+36, font(18), MU, 200, 2)
    # Bottom: additional details
    d.rounded_rectangle([M, 640, W-M, 800], radius=14, fill=rga("#E8E0D4", 60))
    draw(d, "全周期服务覆盖", M+40, 670, font(32, True), NW, 600, 8)
    yy_dt = 710
    for item in ["事前评估 \u2192 目标审核 \u2192 运行监控 \u2192 重点评价 \u2192 结果应用",
                 "每个环节交付：核查清单 + 数据底稿 + 分析报告 + 整改建议"]:
        yy_dt = draw(d, item, M+40, yy_dt+14, font(24), IK, 1200, 6)
    d.line([M, H-65, W-M, H-65], fill=rga(TL, 20), width=2)
    bottom_bar(d)
    return img

# ===================== P7: ENGINEERING =====================
def p7_eng():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "工程咨询与财政评审", "ENGINEERING CONSULTING")
    d.text((W-250, 0), "07", font=font(240, True), fill=rga(NW, 6))
    draw(d, "从概算到结算，帮您把好每一道关——工程造价、合同履约、资金支付和项目绩效协同管理。", M, 240, font(30), MU, 650, 10)
    # Four phases horizontal
    phases = [
        ("预算编制\n财政评审", "工程量·定额·材料价格\n措施费·取费标准", rga(NW)),
        ("清单及\n招标控制价", "招标文件质量提升\n控制价编制", rga(TL)),
        ("结算审核", "合同条款·变更签证\n隐蔽工程·现场核实", rga(NW)),
        ("全过程\n工程咨询", "前期·招采·实施\n验收·结算·绩效", rga(TL)),
    ]
    for i, (title, desc, clr) in enumerate(phases):
        x = M + i * 330
        y0 = 310
        d.rounded_rectangle([x, y0, x+290, y0+180], radius=14, fill=rga("#F8F6F2"), outline=rga(GD, 40), width=2)
        d.rounded_rectangle([x, y0, x+290, y0+90], radius=14, fill=clr)
        draw(d, title, x+20, y0+8, font(26, True), WH, 250, 4)
        draw(d, desc, x+20, y0+110, font(22), MU, 250, 4)
        if i < 3:
            dx = x+290
            d.line([dx, y0+90, dx+40, y0+90], fill=rga(GD, 60), width=2)
            d.polygon([(dx+30, y0+85), (dx+40, y0+90), (dx+30, y0+95)], fill=rga(GD, 60))
    # Key capabilities card
    d.rounded_rectangle([M, 550, W-M, 760], radius=14, fill=rga("#E8E0D4", 60))
    draw(d, "核心能力", M+40, 580, font(32, True), NW, 600, 8)
    caps = ["工程量清单编制与复核", "材料设备价格数据库", "合同条款审核与风险识别", "现场工程变更与签证管理"]
    for i, cap in enumerate(caps):
        cx = M+40 + (i%2)*680
        cy = 625 + (i//2)*50
        d.ellipse([cx, cy+6, cx+16, cy+22], fill=rga(GD))
        draw(d, cap, cx+24, cy, font(24), IK, 600, 4)
    d.line([M, H-65, W-M, H-65], fill=rga(TL, 20), width=2)
    bottom_bar(d)
    return img

# ===================== P8: DIGITAL =====================
def p8_digital():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "数字化审计能力", "DIGITAL AUDIT CAPABILITIES")
    d.text((W-250, 0), "08", font=font(240, True), fill=rga(NW, 6))
    draw(d, "用数据扩大覆盖面、提高发现率、增强证据质量——把审计经验沉淀为可复用的数据工具。", M, 240, font(30), MU, 650, 10)
    # Four pillars
    pillars = [
        ("数据标准", "01", "财务·预算·支付·合同\n采购·资产·工程字段", rga(NW)),
        ("规则模型", "02", "重复支付·超预算执行\n供应商异常·资金沉淀", rga(TL)),
        ("穿透核查", "03", "疑点来源·核查路径\n佐证材料·整改建议", rga(NW)),
        ("报告复核", "04", "金额校验·口径一致\n附表闭环·依据可溯", rga(TL)),
    ]
    for i, (name, num, desc, clr) in enumerate(pillars):
        x = M + (i%2)*700
        y = 310 + (i//2)*220
        d.rounded_rectangle([x, y, x+640, y+190], radius=14, fill=rga("#F8F6F2"), outline=rga(GD, 40), width=2)
        d.rounded_rectangle([x, y, x+140, y+190], radius=14, fill=clr)
        d.text((x+20, y+20), num, font=font(56, True), fill=rga(WH, 60))
        d.text((x+24, y+90), name, font=font(32, True), fill=WH)
        draw(d, desc, x+160, y+40, font(24), IK, 450, 6)
    # Bottom statement
    d.rounded_rectangle([M, 780, W-M, 900], radius=14, fill=rga("#E8E0D4", 60))
    draw(d, "数字化不是替代专业判断，而是让专业判断覆盖更多数据、锁定更准疑点、生成更强证据。", M+40, 815, font(28), IK, 1300, 8, "center")
    d.line([M, H-65, W-M, H-65], fill=rga(TL, 20), width=2)
    bottom_bar(d)
    return img

# ===================== P9: EXPERIENCE =====================
def p9_experience():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "代表经验", "REPRESENTATIVE EXPERIENCE")
    d.text((W-250, 0), "09", font=font(240, True), fill=rga(NW, 6))
    draw(d, "长期服务省、市、县多级财政和主管部门，覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。", M, 240, font(30), MU, 650, 10)
    # Territory box - right side
    d.rounded_rectangle([920, 120, 1480, 520], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((1000, 170), "服务版图", font=font(36, True), fill=rga(TL, 60))
    d.text((1000, 220), "SERVICE TERRITORY", font=font(18), fill=rga(TL, 40))
    terr = [("成都总部","四川省会，核心枢纽"),("阿坝州","川西高原，覆盖藏区"),
            ("西藏办事处","高原地区拓展"),("覆盖川藏黔","三省联动，跨区域服务")]
    for i, (nm, dc) in enumerate(terr):
        y0 = 270 + i*56
        d.ellipse([1000, y0, 1030, y0+30], fill=rga(GD))
        draw(d, nm, 1050, y0, font(26, True), IK, 300, 2)
        draw(d, dc, 1050, y0+28, font(18), MU, 300, 2)
        if i < 3: d.line([1015, y0+30, 1015, y0+56], fill=rga(GD, 100), width=2)
    # Clients grid
    d.text((M, 380), "代表客户", font=font(32, True), fill=NW)
    d.text((M, 418), "REPRESENTATIVE CLIENTS", font=font(18), fill=rga(GD))
    clients = ["四川省财政厅","省公安厅交警总队","省民政厅","省农业农村厅","省教育厅",
               "省退役军人事务厅","省药品监督管理局","达州市财政局","绵阳市财政局",
               "宜宾市财政局","什邡市财政局","德阳市财政局","康定市财政局","九寨沟县财政局"]
    yv = 460
    for i, c in enumerate(clients):
        if i > 0 and i % 5 == 0: yv += 40
        cx = M + (i%5)*280
        d.rounded_rectangle([cx, yv, cx+250, yv+36], radius=6, fill=WH, outline=rga(GD, 80), width=1)
        tw = d.textlength(c, font=font(20))
        d.text((cx+125-tw//2, yv+5), c, font=font(20), fill=NW)
    d.text((M, 700), "覆盖领域：财政、审计、教育、民政、交通、国资、医保、药监、公安", font=font(22), fill=MU)
    # Decorative diamonds
    for j in range(3):
        s = 16 - j*4; x0 = W-70; y0 = H-220 + j*35
        d.polygon([(x0, y0-s), (x0+s, y0), (x0, y0+s), (x0-s, y0)], fill=rga(TL, 25-j*6))
    d.line([M, H-65, W-M, H-65], fill=rga(TL, 20), width=2)
    bottom_bar(d)
    return img

# ===================== P10: CONTACT =====================
def p10_contact():
    img = Image.new("RGB", (W, H), rgb(NW))
    d = ImageDraw.Draw(img, "RGBA")
    # Geometric background
    for r in range(8):
        d.ellipse([W-500+r*30, -100+r*30, W+100+r*30, 500+r*30], fill=rga(TL, 4))
    d.rectangle([0, 280, W, 290], fill=rga(GD))
    # Title
    d.text((M-30, 80), "合作价值", font=font(52, True), fill=WH)
    d.text((M-30, 150), "VALUE PROPOSITION", font=font(22), fill=rga(GD, 200))
    d.line([M-30, 190, M+370, 190], fill=rga(GD), width=3)
    draw(d, "我们交付的不只是一份报告，更是一套可执行的改进方案。", M-30, 260, font(30), WH, 650, 10)
    # Left: contact
    d.rounded_rectangle([M-30, 340, 780, 680], radius=16, fill=rga(WH, 25), outline=rga(GD, 80), width=2)
    draw(d, "联系融策", 180, 370, font(36, True), GD, 500, 8)
    for i, (icon, line) in enumerate([("\u260e","028-87659276"),("\u2709","scrccpa@163.com"),
                ("\u25b6","www.scrccpa.com"),("\u2302","成都市金牛区金周路 595 号")]):
        d.text((180, 460+i*48), icon, font=font(26), fill=GD)
        d.text((215, 458+i*48), line, font=font(26), fill=WH)
    # Right: commitment
    d.rounded_rectangle([880, 340, 1530, 680], radius=16, fill=rga(WH, 25), outline=rga(GD, 80), width=2)
    draw(d, "合作承诺", 930, 380, font(36, True), GD, 500, 8)
    for i, item in enumerate(["客观公正  实事求是","质量优先  证据闭环",
                               "问题有依据  结论可解释","建议能落地  整改有追踪"]):
        d.ellipse([930, 480+i*48, 946, 496+i*48], fill=rga(GD))
        d.text((965, 478+i*48), item, font=font(26), fill=WH)
    # Decorative
    d.ellipse([W-280, H-350, W-30, H-100], fill=rga(GD, 10))
    d.ellipse([W-220, H-300, W-80, H-160], fill=rga(TL, 12))
    # Diamond cluster bottom-left
    for j in range(4):
        s = 18 - j*4; x0 = 100; y0 = H-200 + j*32
        d.polygon([(x0, y0-s), (x0+s, y0), (x0, y0+s), (x0-s, y0)], fill=rga(GD, 20-j*4))
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text((M-30, H-40), "四川融策会计师事务所有限公司", font=font(20), fill=rga(GD, 180))
    return img

# ===================== PDF & DOCX =====================
def generate_pdf(images):
    pw, ph = 1654, 2338
    c = canvas.Canvas(str(PDF), pagesize=(pw, ph))
    for img_path in images:
        im = Image.open(img_path).convert("RGB")
        tmp = WK / "tmp_f.jpg"
        im.save(tmp, "JPEG", quality=97)
        c.drawImage(str(tmp), 0, 0, width=pw, height=ph, preserveAspectRatio=False)
        tmp.unlink(missing_ok=True)
        c.showPage()
    c.save()

def generate_docx(images):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    from copy import deepcopy
    lbls = ["封面","为什么是融策","方法论","服务体系","政府审计","预算绩效","工程咨询","数字化","代表经验","合作价值"]
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v7独立视觉版"
    for idx, (lbl, img_path) in enumerate(zip(lbls, images), 1):
        if idx > 1: doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0); sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(lbl)
        r.font.name = "微软雅黑"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(22); r.font.bold = True
        doc.add_picture(str(img_path), width=Cm(17.0))
    doc.save(DOCX)

def main():
    OUT.mkdir(parents=True, exist_ok=True); WK.mkdir(parents=True, exist_ok=True)
    makers = [p1_cover, p2_about, p3_method, p4_services, p5_gov,
              p6_perf, p7_eng, p8_digital, p9_experience, p10_contact]
    print("Rendering 10 unique pages...")
    images = []
    for i, maker in enumerate(makers, 1):
        out = WK / f"page_{i:02d}.png"
        if not out.exists():
            print(f"  Page {i}...")
            img = maker()
            img.save(out, quality=95)
        images.append(out)
    print("Generating PDF...")
    generate_pdf(images)
    print("Generating DOCX...")
    generate_docx(images)
    print(f"DONE\n{DOCX}\n{PDF}")

if __name__ == "__main__":
    main()