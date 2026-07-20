"""
融策审计复核：报告↔附表数据勾稽对账
复核引擎 - 精确版
"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')

# 直接指定文件路径（避免编码问题）
BASE = r'C:\Users\scrccpa\Desktop\新建文件夹'

# 手动罗列两个项目的文件
projects = {
    'S220安羌镇至茸安乡段': {
        'dir': os.path.join(BASE, os.listdir(BASE)[0]),
    },
    'S452垮沙乡至柯河乡段': {
        'dir': os.path.join(BASE, os.listdir(BASE)[1]),
    },
}

for k, v in projects.items():
    d = v['dir']
    files = os.listdir(d)
    for f in files:
        fp = os.path.join(d, f)
        if f.endswith('.docx') and '审核报告' in f and '说明' not in f:
            v['报告'] = fp
        elif f.endswith('.docx') and '说明' in f:
            v['编制说明'] = fp
        elif f.endswith('.xls') and '汇总表' in f:
            v['附表1'] = fp
        elif f.endswith('.xls') and '明细表' in f:
            v['附表2'] = fp
    print(f"\n{k}:")
    for key in ['报告','编制说明','附表1','附表2']:
        val = v.get(key, '未找到')
        print(f"  {key}: {os.path.basename(val) if val != '未找到' else val}")

# ========== 读取数据 ==========
from docx import Document
import xlrd

def read_docx_text(fp):
    doc = Document(fp)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for i, table in enumerate(doc.tables):
        lines.append(f"【表格{i+1}】")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append('|'.join(cells))
    return '\n'.join(lines)

def read_xls_text(fp):
    book = xlrd.open_workbook(fp, formatting_info=False)
    lines = []
    for sn in book.sheet_names():
        lines.append(f"【Sheet: {sn}】")
        sh = book.sheet_by_name(sn)
        for r in range(sh.nrows):
            row = []
            for c in range(sh.ncols):
                cell = sh.cell(r,c)
                v = cell.value
                if cell.ctype == 2:  # number
                    row.append(str(v))
                elif v:
                    row.append(str(v).strip())
                else:
                    row.append('')
            lines.append('|'.join(row))
    return '\n'.join(lines)

def find_amounts(text, source_label):
    """从文本中提取所有金额信息"""
    results = []
    
    # 金额模式：数字 + 元/万元
    patterns = [
        # 带单位的金额
        r'(概算[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(送审[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(审定[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(核减[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(审核[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(报审[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(合同[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(已支付[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(建筑安装工程[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(待摊投资[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(项目总投资[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(建安工程投资[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(设备投资[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(其他投资[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(工程费用[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(工程建设其他费[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(预备费[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(合计[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
        r'(总计[^0-9]*?)([\d,]+(?:\.\d+)?)\s*(万元|万|元)',
    ]
    
    for pat in patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            label = m.group(1).strip()
            num_str = m.group(2).replace(',', '')
            unit = m.group(3)
            try:
                num = float(num_str)
                if unit in ('元',):
                    num = round(num / 10000, 4)
                results.append({
                    'label': label,
                    'value': num,
                    'unit': '万元',
                    'raw': m.group(0),
                    'source': source_label
                })
            except:
                pass
    
    # 也找表格中的结构化数据
    # 查看文本中是否包含审定/送审两列的数据表
    lines = text.split('\n')
    for line in lines:
        cells = line.split('|')
        if len(cells) >= 4:
            # 检测是否像表格行：有数字+文字
            # 尝试提取序号和金额
            pass
    
    return results

def find_significant_amounts(amounts, max_per_label=3):
    """提取关键金额，每个标签最多取3个"""
    seen = {}
    filtered = []
    for a in amounts:
        label_key = a['label'][:6]  # 前6字做key
        if label_key not in seen:
            seen[label_key] = 0
        if seen[label_key] < max_per_label:
            filtered.append(a)
            seen[label_key] += 1
    return filtered

def normalize_label(label):
    """标准化标签名称"""
    label = re.sub(r'[\s：:，,。、（）()\-]', '', label)
    # 提取关键词
    for kw in ['概算','送审','审定','核减','审核','报审','合同','已支付',
               '建筑安装工程','待摊投资','项目总投资','建安工程投资',
               '设备投资','其他投资','工程费用','工程建设其他费',
               '预备费','合计','总计']:
        if kw in label:
            return kw
    return label[:8]

def fmt(v):
    return f"{v:,.4f}"

print("\n" + "="*70)
print("  融策审计复核：报告↔附表全面数据勾稽核对")
print("="*70)

# ========== 核心复核逻辑 ==========
all_findings = []  # 用于输出excel

for proj_name, v in projects.items():
    print(f"\n{'='*70}")
    print(f"  项目：{proj_name}")
    print(f"{'='*70}")
    
    report_amts = []
    table1_amts = []
    table2_amts = []
    
    # 读取报告
    if v.get('报告'):
        text = read_docx_text(v['报告'])
        report_amts = find_amounts(text, '审核报告')
        print(f"\n  审核报告: {len(text)}字符, 提取{len(report_amts)}个金额")
    
    # 读取编制说明
    report_text_full = ""
    if v.get('编制说明'):
        text = read_docx_text(v['编制说明'])
        report_text_full = text
        extra = find_amounts(text, '编制说明')
        report_amts.extend(extra)
    
    # 读取附表1
    if v.get('附表1'):
        text = read_xls_text(v['附表1'])
        table1_amts = find_amounts(text, '附表1')
        print(f"  附表1(汇总表): {len(text)}字符, 提取{len(table1_amts)}个金额")
    
    # 读取附表2
    if v.get('附表2'):
        text = read_xls_text(v['附表2'])
        table2_amts = find_amounts(text, '附表2')
        print(f"  附表2(明细表): {len(text)}字符, 提取{len(table2_amts)}个金额")
    
    # ===== 复核维度1: 报告↔附表勾稽 =====
    print(f"\n  ┌{'─'*60}┐")
    print(f"  │ 复核维度1：报告 ↔ 附表数据一致性")
    print(f"  └{'─'*60}┘")
    
    # 构建标准化金额映射
    def build_val_map(amts):
        m = {}
        for a in amts:
            k = normalize_label(a['label'])
            if k not in m:
                m[k] = []
            m[k].append(a['value'])
        return m
    
    report_map = build_val_map(report_amts)
    t1_map = build_val_map(table1_amts)
    t2_map = build_val_map(table2_amts)
    
    # 对比关键金额字段
    check_keys = ['概算','送审','审定','核减','建筑安装工程','待摊投资',
                  '合同','已支付','合计','总计','项目总投资','工程费用']
    
    for key in check_keys:
        r_vals = report_map.get(key, [])
        t1_vals = t1_map.get(key, [])
        t2_vals = t2_map.get(key, [])
        
        if r_vals or t1_vals or t2_vals:
            print(f"\n  📊 {key}:")
            if r_vals:
                print(f"    报告: {', '.join(fmt(v) for v in r_vals[:3])}万")
            if t1_vals:
                print(f"    附表1: {', '.join(fmt(v) for v in t1_vals[:3])}万")
            if t2_vals:
                print(f"    附表2: {', '.join(fmt(v) for v in t2_vals[:3])}万")
            
            # 取第一个值对比
            base_val = None
            base_src = None
            for src_vals, src_name in [(r_vals, '报告'), (t1_vals, '附表1'), (t2_vals, '附表2')]:
                if src_vals:
                    base_val = src_vals[0]
                    base_src = src_name
                    break
            
            if base_val:
                for other_vals, other_name in [(r_vals, '报告'), (t1_vals, '附表1'), (t2_vals, '附表2')]:
                    if other_vals and other_name != base_src:
                        for ov in other_vals[:3]:
                            diff = abs(base_val - ov)
                            if diff < 0.01:
                                print(f"    ✅ {base_src}={other_name}: {fmt(base_val)}万, 一致")
                            elif diff < 1:
                                print(f"    ⚠️ {base_src}={other_name}: {fmt(base_val)}万 vs {fmt(ov)}万, 差{diff:.4f}万")
                            else:
                                print(f"    ❌ {base_src}={other_name}: {fmt(base_val)}万 vs {fmt(ov)}万, 差{diff:.4f}万")
            
            all_findings.append({
                '项目': proj_name,
                '复核维度': '报告↔附表勾稽',
                '字段': key,
                '报告金额(万)': ', '.join(fmt(v) for v in r_vals[:3]) if r_vals else '',
                '附表1金额(万)': ', '.join(fmt(v) for v in t1_vals[:3]) if t1_vals else '',
                '附表2金额(万)': ', '.join(fmt(v) for v in t2_vals[:3]) if t2_vals else '',
                '结论': '',
            })
    
    # ===== 复核维度2: 报告内部数据勾稽 =====
    print(f"\n  ┌{'─'*60}┐")
    print(f"  │ 复核维度2：报告内部数据勾稽关系")
    print(f"  └{'─'*60}┘")
    
    # 检查：概算 ≈ 送审？
    if '概算' in report_map and '送审' in report_map:
        gs = report_map['概算'][0]
        ss = report_map['送审'][0]
        diff = abs(gs - ss)
        conclusion = '一致' if diff < 0.01 else f"差异{diff:.2f}万"
        flag = '✅' if diff < 0.01 else '⚠️' if diff < 10 else '❌'
        print(f"\n  {flag} 概算({fmt(gs)}万) vs 送审({fmt(ss)}万): {conclusion}")
        all_findings.append({
            '项目': proj_name, '复核维度': '报告内部勾稽-概算vs送审',
            '字段': f'概算({fmt(gs)}万)/送审({fmt(ss)}万)',
            '结论': conclusion
        })
    
    # 检查：送审 ≈ 审定 + 核减
    if '送审' in report_map and '审定' in report_map and '核减' in report_map:
        ss = report_map['送审'][0]
        ds = report_map['审定'][0]
        hj = report_map['核减'][0]
        calc = round(ds + hj, 4)
        diff = abs(ss - calc)
        conclusion = '一致' if diff < 0.01 else f"差异{diff:.2f}万"
        flag = '✅' if diff < 0.01 else '⚠️' if diff < 10 else '❌'
        print(f"  {flag} 送审({fmt(ss)}万) = 审定({fmt(ds)}万) + 核减({fmt(hj)}万) = {fmt(calc)}万: {conclusion}")
        all_findings.append({
            '项目': proj_name, '复核维度': '报告内部勾稽-送审=审定+核减',
            '字段': f'送审({fmt(ss)}万)/审定({fmt(ds)}万)/核减({fmt(hj)}万)',
            '结论': conclusion
        })
    
    # 检查：建安工程投资 + 待摊投资 + 其他 ≈ 总投资
    for key in ['建筑安装工程','建安工程投资']:
        if key in report_map:
            ja = report_map[key][0]
            print(f"\n  建安投资: {fmt(ja)}万")
            break
    for key in ['待摊投资']:
        if key in report_map:
            dt = report_map[key][0]
            print(f"  待摊投资: {fmt(dt)}万")
            break
    if '项目总投资' in report_map:
        ztz = report_map['项目总投资'][0]
        print(f"  项目总投资: {fmt(ztz)}万")
    
    # 检查：合同金额 vs 审定金额
    if '合同' in report_map and '审定' in report_map:
        ht = report_map['合同'][0]
        ds = report_map['审定'][0]
        diff = abs(ht - ds)
        conclusion = '一致' if diff < 0.01 else f"差异{diff:.2f}万"
        flag = '✅' if diff < 0.01 else '⚠️'
        print(f"\n  {flag} 合同({fmt(ht)}万) vs 审定({fmt(ds)}万): {conclusion}")
        all_findings.append({
            '项目': proj_name, '复核维度': '报告内部勾稽-合同vs审定',
            '字段': f'合同({fmt(ht)}万)/审定({fmt(ds)}万)',
            '结论': conclusion
        })

print(f"\n{'='*70}")
print(f"  复核完成！共{len(all_findings)}条检查记录")
print(f"{'='*70}")

# ========== 输出Excel ==========
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '勾稽复核对账表'
    
    # 标题
    ws.merge_cells('A1:G1')
    ws['A1'] = '融策会计师事务所 - 竣工财务决算审核报告数据勾稽复核对账表'
    ws['A1'].font = Font(name='微软雅黑', size=14, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 35
    
    # 表头
    headers = ['序号', '项目名称', '复核维度', '表单/字段', '报告金额(万元)', '附表金额(万元)', '复核结论']
    ws.append(headers)
    for c in range(1, 8):
        cell = ws.cell(row=2, column=c)
        cell.font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='1A237E', end_color='1A237E', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[2].height = 25
    
    # 分隔行标题
    proj_rows = {}
    row_num = 3
    seq = 1
    
    for proj_name in list(projects.keys()):
        proj_rows[proj_name] = row_num
        
        # 用颜色标注不同项目
        for f in all_findings:
            if f['项目'] != proj_name:
                continue
            
            ws.append([
                seq,
                proj_name,
                f.get('复核维度', ''),
                f.get('字段', ''),
                f.get('报告金额(万)', ''),
                f.get('附表1金额(万)', ''),
                f.get('结论', '')
            ])
            seq += 1
            row_num += 1
        
        # 合并项目名称单元格（可视化提升）
    
    # 设置列宽
    col_widths = [6, 30, 25, 35, 18, 18, 15]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[chr(64+i)].width = w
    
    # 设置边框
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    for row in ws.iter_rows(min_row=2, max_row=row_num, max_col=7):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # 颜色标记结论
    green_fill = PatternFill(start_color='E8F5E9', fill_type='solid')
    yellow_fill = PatternFill(start_color='FFF3E0', fill_type='solid')
    red_fill = PatternFill(start_color='FFEBEE', fill_type='solid')
    
    for row in ws.iter_rows(min_row=3, max_row=row_num, max_col=7):
        conclusion_cell = row[6]  # 第7列=结论
        if conclusion_cell.value:
            val = str(conclusion_cell.value)
            if '一致' in val or '通过' in val:
                for c in row:
                    c.fill = green_fill
            elif '差异' in val:
                for c in row:
                    c.fill = yellow_fill
            elif '❌' in val:
                for c in row:
                    c.fill = red_fill
    
    # 第二个sheet：详细金额清单
    ws2 = wb.create_sheet('金额明细清单')
    ws2.merge_cells('A1:F1')
    ws2['A1'] = '各文件提取的金额明细'
    ws2['A1'].font = Font(name='微软雅黑', size=12, bold=True)
    ws2['A1'].alignment = Alignment(horizontal='center')
    
    detail_headers = ['序号', '项目', '来源文件', '金额标签', '金额(万元)', '原始文本']
    ws2.append(detail_headers)
    for c in range(1, 7):
        cell = ws2.cell(row=2, column=c)
        cell.font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='283593', fill_type='solid')
    
    dseq = 1
    for proj_name, v in projects.items():
        for src_name in ['报告','编制说明','附表1','附表2']:
            src_key = src_name
            src_amts = []
            if src_key == '报告':
                if v.get('报告'):
                    src_amts = find_amounts(read_docx_text(v['报告']), src_name)
            elif src_key == '编制说明':
                if v.get('编制说明'):
                    src_amts = find_amounts(read_docx_text(v['编制说明']), src_name)
            elif src_key == '附表1':
                if v.get('附表1'):
                    src_amts = find_amounts(read_xls_text(v['附表1']), src_name)
            elif src_key == '附表2':
                if v.get('附表2'):
                    src_amts = find_amounts(read_xls_text(v['附表2']), src_name)
            
            for a in find_significant_amounts(src_amts, max_per_label=5):
                ws2.append([dseq, proj_name, a['source'], a['label'], a['value'], a['raw']])
                dseq += 1
    
    for i, w in enumerate([6, 25, 12, 18, 14, 40], 1):
        ws2.column_dimensions[chr(64+i)].width = w
    for row in ws2.iter_rows(min_row=2, max_row=dseq+1, max_col=6):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
    
    output_path = r'C:\Users\scrccpa\Desktop\融策审计勾稽复核表.xlsx'
    wb.save(output_path)
    print(f"\n✅ 复核表已保存到: {output_path}")
    print(f"   Sheet1: 勾稽对账总表")
    print(f"   Sheet2: 金额明细清单")
    
except ImportError:
    print("\n⚠️ openpyxl未安装，尝试保存CSV...")
    import csv
    output_path = r'C:\Users\scrccpa\Desktop\融策审计勾稽复核表.csv'
    with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.writer(f)
        w.writerow(['序号','项目名称','复核维度','字段','报告金额','附表金额','复核结论'])
        for i, fi in enumerate(all_findings, 1):
            w.writerow([
                i, fi['项目'], fi.get('复核维度',''), fi.get('字段',''),
                fi.get('报告金额(万)',''), fi.get('附表1金额(万)',''), fi.get('结论','')
            ])
    print(f"\n✅ 复核表(CSV)已保存到: {output_path}")

print("\n复核完毕。")
