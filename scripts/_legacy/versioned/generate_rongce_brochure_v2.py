from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

OUT_DIR = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
WORK_DIR = Path("work/sichuan_rongce_brochure/v2_assets")
DOCX_PATH = OUT_DIR / "四川融策宣传册_完善稿_v2_图文版.docx"
PDF_PATH = OUT_DIR / "四川融策宣传册_完善稿_v2_图文版.pdf"

BLUE = "0A1F3F"
TEAL = "1A5C6E"
GOLD = "C5955C"
WARM = "F5F2EC"
INK = "20252B"
MUTED = "667085"

PAGES = [
    {
        "title": "四川融策",
        "subtitle": "政府审计与工程咨询综合服务机构",
        "lead": "以财政资金、公共资源、工程项目和预算绩效为核心场景，提供审计、评价、咨询、监督检查和数据分析一体化服务。",
        "visual_title": "财政治理专业服务",
        "visual_items": ["政府审计", "预算绩效", "财政监督", "工程咨询", "数据分析"],
        "points": ["谋专业之策，融品质之精", "发现风险、厘清责任、推动整改", "让专业报告转化为治理改进"],
    },
    {
        "title": "公司定位",
        "subtitle": "把专业经验转化为治理能力",
        "lead": "四川融策会计师事务所有限公司成立于2000年，总部位于成都，长期服务财政、审计、教育、民政、农业农村、交通、国资、医保等领域。",
        "visual_title": "20+ 年积累",
        "visual_items": ["400+ 客户", "50人团队", "成都总部", "阿坝/西藏服务支点"],
        "points": ["会计、造价、绩效、工程复合团队", "熟悉政府业务流程和基层执行场景", "形成审计、绩效、工程咨询协同格局"],
    },
    {
        "title": "核心优势",
        "subtitle": "专业判断、现场经验、数据能力三位一体",
        "lead": "融策不只看账和表，更把资金、项目、合同、资产、制度和责任链条贯通核查，保证问题有证据、结论有依据、建议能落地。",
        "visual_title": "融策方法",
        "visual_items": ["制度审查", "数据核验", "现场核查", "证据闭环", "整改建议"],
        "points": ["资质团队支撑专业判断", "项目案例库沉淀风险规律", "重大结论复核到数据来源和计算依据"],
    },
    {
        "title": "服务体系",
        "subtitle": "围绕财政管理和公共治理形成六类服务",
        "lead": "从单项审计到全过程咨询，从结果评价到制度优化，从人工核查到数据化识别，形成面向委托方的综合服务体系。",
        "visual_title": "六类服务",
        "visual_items": ["政府审计", "预算绩效", "财政监督", "工程咨询", "采购审计", "管理咨询"],
        "points": ["政府审计：经责、收支、专项资金、资产清查、国企审计、竣工决算财务审计", "预算绩效：事前评估、目标审核、运行监控、重点评价、结果应用", "工程咨询：预算编制、财政评审、结算审核、全过程工程咨询"],
    },
    {
        "title": "政府审计",
        "subtitle": "从合规审查延伸到治理改进",
        "lead": "面向财政、审计、主管部门及国有企事业单位，提供经济责任、财务收支、专项资金、资产清查、财政监督检查和竣工决算财务审计等服务。",
        "visual_title": "审计闭环",
        "visual_items": ["账表", "合同", "项目", "资金", "资产", "责任"],
        "points": ["经济责任审计：重大决策、资金资产、项目建设、内控风险", "专项资金审计：申报、分配、拨付、使用、绩效、结余沉淀", "竣工决算财务审计：建设成本归集、资金来源、资产交付、尾工尾款"],
    },
    {
        "title": "预算绩效管理",
        "subtitle": "让财政资金从“花了没有”走向“花得值不值”",
        "lead": "围绕预算编制、执行监控、绩效评价、结果应用四个环节，帮助财政部门和预算单位建立全过程预算绩效管理链条。",
        "visual_title": "全过程绩效",
        "visual_items": ["事前评估", "目标审核", "运行监控", "绩效评价", "结果应用"],
        "points": ["指标体系贴合项目实际", "评价证据支撑每项结论", "问题建议反馈预算安排和项目管理"],
    },
    {
        "title": "工程咨询与财政评审",
        "subtitle": "把工程语言和财政管理要求对齐",
        "lead": "工程咨询业务覆盖预算编制、工程量清单、招标控制价、财政评审、结算审核和全过程工程咨询，帮助委托方控制投资风险。",
        "visual_title": "工程咨询链条",
        "visual_items": ["预算编制", "清单控制价", "财政评审", "过程管理", "结算审核"],
        "points": ["核工程量、定额套用、材料价格和取费标准", "查合同条款、变更签证、隐蔽工程和支付资料", "协同项目管理、造价控制、验收结算和绩效评价"],
    },
    {
        "title": "数字化改革服务能力",
        "subtitle": "用数据提高覆盖面、发现率和证据质量",
        "lead": "融策将审计经验沉淀为数据标准、识别规则、疑点模型和底稿模板，提升项目启动、疑点筛选、现场核查和报告复核效率。",
        "visual_title": "数据化作业",
        "visual_items": ["数据标准", "规则模型", "疑点清单", "穿透核查", "底稿沉淀"],
        "points": ["覆盖财务、预算、支付、合同、采购、资产、工程项目数据", "识别重复支付、超预算执行、供应商异常、资金沉淀等风险", "形成疑点来源、核查过程、佐证材料、影响金额、整改建议闭环"],
    },
    {
        "title": "代表服务经验",
        "subtitle": "长期服务省、市、县多级财政和主管部门",
        "lead": "服务对象覆盖省级部门、市县财政部门、国有企事业单位和工程建设单位，项目类型涵盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。",
        "visual_title": "服务版图",
        "visual_items": ["省级部门", "市县财政", "主管部门", "国资单位", "工程项目"],
        "points": ["四川省财政厅、省公安厅交警总队、省民政厅、省农业农村厅等绩效服务经验", "达州、绵阳、宜宾、什邡、德阳、康定、九寨沟等市县项目经验", "交通、教育、民政、农业农村、医保等重点民生领域审计评价"],
    },
    {
        "title": "合作价值",
        "subtitle": "交付报告，更交付可执行的改进方案",
        "lead": "融策坚持客观公正、实事求是、质量优先。每个项目都力求沉淀风险清单、指标体系、数据规则和整改建议，帮助委托方持续改进管理。",
        "visual_title": "联系我们",
        "visual_items": ["028-87659276", "scrccpa@163.com", "www.scrccpa.com", "四川融策会计师事务所有限公司"],
        "points": ["公开公正  用心服务", "诚信为本  服务至上", "追求卓越  持续改进"],
    },
]


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], width: int, fnt, fill, line_gap=8) -> int:
    x, y = xy
    current = ""
    lines = []
    for char in text:
        trial = current + char
        if draw.textlength(trial, font=fnt) <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def make_visual(page: dict, index: int) -> Path:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 760), hex_to_rgb(WARM))
    draw = ImageDraw.Draw(img)
    blue = hex_to_rgb(BLUE)
    teal = hex_to_rgb(TEAL)
    gold = hex_to_rgb(GOLD)
    ink = hex_to_rgb(INK)
    muted = hex_to_rgb(MUTED)

    draw.rectangle((0, 0, 1600, 760), fill=hex_to_rgb(WARM))
    draw.rectangle((0, 0, 1600, 120), fill=blue)
    draw.rectangle((0, 120, 1600, 132), fill=gold)
    draw.text((70, 34), page["visual_title"], font=font(46, True), fill=(255, 255, 255))
    draw.text((1280, 38), f"RONGCE  {index:02d}", font=font(26, True), fill=gold)

    # Abstract layered waves, matching the existing brochure's restrained Chinese pattern.
    for row in range(5):
        y = 610 + row * 26
        for x in range(-100, 1700, 160):
            draw.arc((x, y - 70, x + 210, y + 80), 190, 350, fill=(205, 149, 92), width=3)
            draw.arc((x + 30, y - 42, x + 185, y + 58), 190, 350, fill=(26, 92, 110), width=2)

    items = page["visual_items"]
    if index in {3, 5, 6, 7, 8}:
        start_x = 120
        y = 330
        gap = 245 if len(items) <= 5 else 205
        for i, item in enumerate(items):
            x = start_x + i * gap
            draw.ellipse((x, y, x + 112, y + 112), fill=blue if i % 2 == 0 else teal, outline=gold, width=5)
            draw.text((x + 38, y + 24), str(i + 1).zfill(2), font=font(34, True), fill=(255, 255, 255))
            draw.line((x + 112, y + 56, x + gap - 18, y + 56), fill=gold, width=4)
            draw_wrapped(draw, item, (x - 12, y + 135), 150, font(25, True), ink, line_gap=4)
    elif index in {4, 9}:
        cols = 3
        for i, item in enumerate(items):
            col = i % cols
            row = i // cols
            x = 120 + col * 470
            y = 210 + row * 170
            draw.rounded_rectangle((x, y, x + 380, y + 120), radius=18, fill=(255, 255, 255), outline=gold, width=4)
            draw.rectangle((x, y, x + 28, y + 120), fill=teal if i % 2 else blue)
            draw_wrapped(draw, item, (x + 55, y + 36), 290, font(30, True), ink, line_gap=4)
    else:
        for i, item in enumerate(items):
            x = 120 + (i % 3) * 455
            y = 210 + (i // 3) * 160
            color = [blue, teal, gold][i % 3]
            draw.rounded_rectangle((x, y, x + 360, y + 110), radius=14, fill=(255, 255, 255), outline=color, width=4)
            draw.ellipse((x + 26, y + 27, x + 82, y + 83), fill=color)
            draw.text((x + 44, y + 38), "·", font=font(38, True), fill=(255, 255, 255))
            draw_wrapped(draw, item, (x + 105, y + 34), 220, font(28, True), ink, line_gap=4)

    draw_wrapped(draw, page["lead"], (90, 152), 1420, font(30), muted, line_gap=10)
    path = WORK_DIR / f"visual_{index:02d}.png"
    img.save(path, quality=95)
    return path


def set_run(run, size=11, bold=False, color=INK, east_asia="微软雅黑") -> None:
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_docx_page(doc: Document, page: dict, image_path: Path, first: bool) -> None:
    if not first:
        doc.add_section(WD_SECTION.NEW_PAGE)
    sec = doc.sections[-1]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.3)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    r = p.add_run(page["title"])
    set_run(r, size=24, bold=True, color=BLUE)
    p = doc.add_paragraph()
    r = p.add_run(page["subtitle"])
    set_run(r, size=12.5, bold=True, color=GOLD)
    p = doc.add_paragraph()
    r = p.add_run(page["lead"])
    set_run(r, size=10.5, color=INK, east_asia="宋体")

    doc.add_picture(str(image_path), width=Cm(16.8))

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    hdr = table.cell(0, 0)
    set_cell_shading(hdr, BLUE)
    r = hdr.paragraphs[0].add_run("服务要点")
    set_run(r, size=10.5, bold=True, color="FFFFFF")
    for item in page["points"]:
        cell = table.add_row().cells[0]
        r = cell.paragraphs[0].add_run("• " + item)
        set_run(r, size=10, color=INK, east_asia="宋体")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = foot.add_run("公开公正  用心服务  诚信为本  服务至上  追求卓越")
    set_run(r, size=8, color=GOLD)


def register_pdf_font() -> str:
    for path, name in [
        (r"C:\Windows\Fonts\msyh.ttc", "MSYH"),
        (r"C:\Windows\Fonts\simhei.ttf", "SIMHEI"),
        (r"C:\Windows\Fonts\simsun.ttc", "SIMSUN"),
    ]:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont(name, path))
            return name
    return "Helvetica"


def generate_docx(images: list[Path]) -> None:
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v2图文版"
    for i, page in enumerate(PAGES):
        add_docx_page(doc, page, images[i], first=i == 0)
    doc.save(DOCX_PATH)


def generate_pdf(images: list[Path]) -> None:
    pdf_font = register_pdf_font()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.2 * cm,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Title"], fontName=pdf_font, fontSize=25, leading=30, textColor=colors.HexColor("#" + BLUE), alignment=TA_LEFT, spaceAfter=4)
    subtitle = ParagraphStyle("subtitle", parent=styles["Heading2"], fontName=pdf_font, fontSize=12, leading=16, textColor=colors.HexColor("#" + GOLD), alignment=TA_LEFT, spaceAfter=8)
    lead = ParagraphStyle("lead", parent=styles["BodyText"], fontName=pdf_font, fontSize=9.6, leading=15, textColor=colors.HexColor("#" + INK), alignment=TA_LEFT, spaceAfter=10)
    point = ParagraphStyle("point", parent=styles["BodyText"], fontName=pdf_font, fontSize=9.3, leading=13.5, textColor=colors.HexColor("#" + INK))
    head = ParagraphStyle("head", parent=point, textColor=colors.white)
    story = []
    usable_width = A4[0] - 3.2 * cm
    for i, page in enumerate(PAGES):
        story.append(Paragraph(page["title"], title))
        story.append(Paragraph(page["subtitle"], subtitle))
        story.append(Paragraph(page["lead"], lead))
        story.append(PdfImage(str(images[i]), width=usable_width, height=usable_width * 760 / 1600))
        story.append(Spacer(1, 0.25 * cm))
        data = [[Paragraph("服务要点", head)]] + [[Paragraph("• " + item, point)] for item in page["points"]]
        table = Table(data, colWidths=[usable_width])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BLUE)),
            ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#" + WARM)),
            ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#" + GOLD)),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D8C7A9")),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(table)
        if i != len(PAGES) - 1:
            story.append(PageBreak())
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    images = [make_visual(page, i + 1) for i, page in enumerate(PAGES)]
    generate_docx(images)
    generate_pdf(images)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
