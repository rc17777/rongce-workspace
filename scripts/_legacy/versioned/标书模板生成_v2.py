#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""融策标书模板 v2.0 — 高级版"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ========== 高级配色 ==========
C_DARK_BLUE = RGBColor(0x0A, 0x1F, 0x3F)   # #0A1F3F 深藏蓝
C_TEAL = RGBColor(0x1A, 0x5C, 0x6E)         # #1A5C6E 深青绿
C_GOLD = RGBColor(0xC5, 0x95, 0x5C)         # #C5955C 铜金色
C_GOLD_LIGHT = RGBColor(0xE8, 0xD5, 0xB5)   # 浅金色
C_BG_WARM = RGBColor(0xFA, 0xFA, 0xF8)      # 暖白底
C_TEXT = RGBColor(0x2D, 0x2D, 0x2D)         # 深灰正文
C_TEXT_LIGHT = RGBColor(0x8B, 0x9D, 0xAF)   # 浅灰辅助
C_TABLE_ALT = RGBColor(0xF8, 0xF6, 0xF0)    # 隔行暖灰
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_SECTION_BG = RGBColor(0xF0, 0xEC, 0xE3)   # 章节装饰底

# ========== 页面设置 ==========
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

# 设置默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.font.color.rgb = C_TEXT
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 辅助函数 ==========

def add_horizontal_line(before_para=True, after_para=True):
    """加入分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # 使用边框作为线条
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="C5955C"/>'
        f'</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)

def add_section_heading(text, number=""):
    """添加带装饰线的章节标题"""
    # 装饰条（金色）
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(16)
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run('')
    # 用金色方框作为装饰
    p1Bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="2" w:color="0A1F3F"/>'
        f'</w:pBdr>'
    )
    p1._element.get_or_add_pPr().append(p1Bdr)
    
    # 标题
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(12)
    
    full_text = f'{number}  {text}' if number else text
    run = p2.add_run(full_text)
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = C_DARK_BLUE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_sub_heading(text):
    """添加二级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    # 金色小方块 + 标题
    run_marker = p.add_run('■ ')
    run_marker.font.size = Pt(11)
    run_marker.font.color.rgb = C_GOLD
    run_marker.font.name = '微软雅黑'
    run_marker._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = C_TEAL
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_body(text, size=12):
    """添加正文"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = C_TEXT
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Pt(size * 2)  # 首行缩进
    p.paragraph_format.line_spacing = 1.5
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color_hex='CCCCCC'):
    """设置单元格边框"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_cell(cell, text, bold=False, color=None, size=10, align=None, fn='微软雅黑', bg=None):
    """设置单元格文字和样式"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.color.rgb = color or C_TEXT
    run.font.size = Pt(size)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    p.alignment = align or WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    if bg:
        set_cell_shading(cell, bg)
    set_cell_border(cell, 'D0C8B8')

def make_premium_table(headers, rows, first_col_bold=True):
    """制作高级表格"""
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头：深藏蓝底+金字
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=C_GOLD, size=10, bg='0A1F3F')
    
    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            is_first = (ci == 0 and first_col_bold)
            bg = 'FAFAF6' if ri % 2 == 0 else 'F5F2EC'
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cell, str(val), bold=is_first, color=C_TEXT, size=10, 
                     align=align, bg=bg)
    
    # 设置表格宽度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    
    # 设置表格边框
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="C5955C"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0C8B8"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="C5955C"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0C8B8"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0D8C8"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0D8C8"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    doc.add_paragraph()  # 表后空行
    return table

def set_page_border(section, color='F0ECDF'):
    """设置页面边框"""
    section._sectPr

# ================================================================
# ===== 正文内容 =====
# ================================================================

# ===== 第1页：封面（插入封面图） =====
cover_path = r'D:\openclaw-workspace\bid_aba\封面-投标文件.png'
if os.path.exists(cover_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Cm(16.0))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

doc.add_page_break()

# ===== 第2页：目录 =====
add_section_heading('目  录')

toc_items = [
    ('第一章  公司简介', 1),
    ('第二章  项目实施方案', 3),
    ('  一、项目概况与理解', 3),
    ('  二、工作内容与流程', 4),
    ('  三、审计重点与方法', 7),
    ('  四、人员配置计划', 10),
    ('第三章  类似业绩', 12),
    ('第四章  服务承诺', 15),
    ('第五章  附件', 17),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    
    is_sub = title.startswith('  ')
    run = p.add_run(title.strip())
    run.font.size = Pt(12 if not is_sub else 11)
    run.bold = not is_sub
    run.font.color.rgb = C_DARK_BLUE if not is_sub else C_TEXT
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 虚线 + 页码
    dots = '.' * (60 - len(title) + (15 if is_sub else 0))
    run2 = p.add_run(f' {dots} {page}')
    run2.font.size = Pt(12 if not is_sub else 11)
    run2.font.color.rgb = C_TEXT_LIGHT
    run2.font.name = '微软雅黑'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ===== 第一章：公司简介 =====
add_section_heading('第一章  公司简介')

add_sub_heading('1.1 公司概况')

add_body('四川融策会计师事务所（普通合伙）（以下简称"我公司"）成立于20XX年，是经四川省财政厅批准设立的综合性会计师事务所。公司注册资本XXX万元，现有从业人员XX人，其中注册会计师XX人、中级以上职称XX人。')

add_body('我公司业务涵盖政府审计、绩效评价、资产清查、专项债申报、监督检查、工程预算编制、财政评审、全过程工程咨询、工程结算审核等专业领域。多年来，我公司深耕政府审计和工程咨询服务领域，积累了丰富的项目经验，具备为各级政府部门和企事业单位提供高质量专业服务的能力。')

add_body('公司秉承"诚信为本、专业立身"的服务理念，建立了完善的三级复核质量控制体系，确保每一份报告经得起推敲、每一个数据经得起检验。')

add_sub_heading('1.2 资质荣誉')

make_premium_table(
    headers=['序号', '资质/荣誉名称', '发证机关', '有效期'],
    rows=[
        ('1', '会计师事务所执业证书', '四川省财政厅', '长期有效'),
        ('2', 'ISO9001质量管理体系认证', '中国质量认证中心', '2024-2027'),
        ('3', 'AAA级信用等级证书', 'XX信用评估机构', '2025-2028'),
        ('4', '工程咨询单位乙级资信', '四川省工程咨询协会', '2024-2026'),
    ]
)

add_sub_heading('1.3 组织架构')

# 插入组织架构图
flow_path = r'D:\openclaw-workspace\bid_aba\流程图-审计工作流程.png'
if os.path.exists(flow_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(flow_path, width=Cm(14.0))
    p.paragraph_format.space_before = Pt(6)

doc.add_page_break()

# ===== 第二章：项目实施方案 =====
add_section_heading('第二章  项目实施方案')

add_sub_heading('一、项目概况与理解')

add_body('本项目为XXX项目竣工财务决算审核。主要工作内容包括：对项目建设资金到位和使用情况进行审核、对建筑安装工程投资进行核定、对设备投资和待摊投资进行逐项复核、编制竣工财务决算报表及审核报告。')

add_body('我公司将在充分理解项目背景和需求的基础上，组织专业团队严格按照《基本建设项目竣工财务决算管理暂行办法》（财建〔2016〕503号）、《会计师事务所从事基本建设工程预算结算决算审核暂行办法》（财协字〔1999〕103号）等政策法规的要求，提供专业、规范的审计服务。')

add_sub_heading('二、工作内容与流程')

add_body('本项目审核工作分三个阶段实施：')

# 流程说明表
make_premium_table(
    headers=['阶段', '工作内容', '工作标准', '输出成果'],
    rows=[
        ('准备阶段', '接受委托→签订业务约定书\n→成立项目组→制定审计方案\n→下发资料清单', '方案经三级审核\n资料清单完整', '审计方案\n资料清单'),
        ('实施阶段', '进驻现场→资料核对→盘点核实\n→资金使用审核→二类费用复核\n→底稿编制', '账实相符\n依据充分\n底稿规范', '盘点记录\n复核意见表\n审计底稿'),
        ('报告阶段', '编制报告→三级复核→\n出具正式报告→归档', '格式规范\n数据准确\n结论客观', '审核报告\n附表'),
    ]
)

add_sub_heading('三、审计重点与方法')

add_body('本项目审核重点包括但不限于：')

make_premium_table(
    headers=['序号', '审核重点', '审核方法', '风险关注点'],
    rows=[
        ('1', '资金到位及使用情况', '核对拨款文件、银行流水\n与账面数逐一勾稽', '截留、挪用、超范围使用'),
        ('2', '建筑安装工程投资', '核实施工合同、结算书\n工程量清单、变更签证', '虚列工程量\n高套定额'),
        ('3', '待摊投资（二类费用）', '逐项对照国家计费标准\n复核计费基数与费率', '超标准计费\n无依据收费'),
        ('4', '设备投资', '核实采购合同、验收记录\n与账面一致', '虚列设备\n以次充好'),
        ('5', '交付使用资产', '核实资产清单与实物\n确保账实相符', '资产流失\n账实不符'),
    ]
)

add_sub_heading('四、人员配置计划')

make_premium_table(
    headers=['序号', '姓名', '职务', '执业资格', '本项目职责'],
    rows=[
        ('1', '李开', '合伙人/总经理', '注册会计师', '项目总负责人'),
        ('2', 'XXX', '部门经理', '注册会计师', '现场负责人'),
        ('3', 'XXX', '项目经理', '中级会计师', '审计实施'),
        ('4', 'XXX', '项目经理', '造价工程师', '造价审核'),
        ('5', 'XXX', '审计助理', '初级会计师', '底稿编制/数据录入'),
    ]
)

add_body('以上人员均具备丰富的政府审计和竣工财务决算审核经验，项目负责人李开拥有XX年以上从业经历，主持完成过多个类似项目。')

doc.add_page_break()

# ===== 第三章：业绩 =====
add_section_heading('第三章  类似业绩')

add_body('近三年，我公司承接并完成的同类项目业绩如下：')

make_premium_table(
    headers=['序号', '项目名称', '委托单位', '合同金额\n（万元）', '完成时间', '服务内容'],
    rows=[
        ('1', '阿坝州S220线安羌镇至茸安乡段\n灾害修复整治工程竣工财务决算审核', '阿坝州财政局', 'XX', '2026年', '竣工财务决算审核'),
        ('2', '阿坝州S452线垮沙乡至柯河乡段\n灾害修复整治工程竣工财务决算审核', '阿坝州财政局', 'XX', '2026年', '竣工财务决算审核'),
        ('3', 'XX县20XX年度财政预算执行\n及其他财政收支情况审计', 'XX县审计局', 'XX', '2025年', '预算执行审计'),
        ('4', 'XX市教育局营养改善计划\n专项资金审计', 'XX市教育局', 'XX', '2025年', '专项资金审计'),
        ('5', 'XX区20XX年财政绩效评价', 'XX区财政局', 'XX', '2024年', '绩效评价'),
        ('6', 'XX县国有资产清查服务', 'XX县财政局', 'XX', '2024年', '资产清查'),
        ('7', 'XX镇乡村振兴专项债\n项目申报服务', 'XX镇人民政府', 'XX', '2024年', '专项债申报'),
    ]
)

add_body('注：以上为部分代表业绩，可根据项目需要提供完整的业绩清单及证明材料。', size=10)

doc.add_page_break()

# ===== 第四章：服务承诺 =====
add_section_heading('第四章  服务承诺')

promises = [
    ('质量承诺', '严格执行三级复核制度，确保审核报告数据准确、依据充分、结论客观。如因我公司原因导致报告质量问题，我公司承担相应责任。'),
    ('时间承诺', '自收到全部资料之日起XX个工作日内出具正式审核报告。如需加急，可安排专人加班，最快XX个工作日交付。'),
    ('保密承诺', '对项目实施过程中知悉的所有信息严格保密，不向任何第三方泄露，项目结束后全部资料归档封存。'),
    ('人员承诺', '项目主要人员一旦确定，未经甲方书面同意不随意更换。如需调整，提前报甲方审批。'),
]

for title, content in promises:
    add_sub_heading(title)
    add_body(content)

doc.add_page_break()

# ===== 第五章：排版规范附录 =====
add_section_heading('附录  融策投标文件排版规范')

add_body('为确保投标文件质量，特制定以下排版规范，公司所有投标文件均按此标准执行：')

add_horizontal_line()

add_sub_heading('1. 配色方案')
make_premium_table(
    headers=['用途', '色值', '色样说明'],
    rows=[
        ('主色  ——  标题/表头/装饰', '#0A1F3F', '深藏蓝色，沉稳大气'),
        ('辅色  ——  二级标题', '#1A5C6E', '深青绿色，专业内敛'),
        ('点缀色 ——  装饰线条/标记', '#C5955C', '铜金色，提升质感'),
        ('背景色 ——  隔行/底纹', '#F5F2EC', '暖灰底，柔和舒适'),
        ('正文色 ——  正文内容', '#2D2D2D', '深灰色，护眼专业'),
    ]
)

add_sub_heading('2. 字体规范')
make_premium_table(
    headers=['用途', '字体', '字号', '备注'],
    rows=[
        ('封面标题', '微软雅黑', '36-42pt', '加粗深蓝'),
        ('一级标题', '微软雅黑', '18pt', '加粗深蓝'),
        ('二级标题', '微软雅黑', '14pt', '加粗深青绿'),
        ('正文', '宋体', '12pt（小四）', '首行缩进2字符'),
        ('表格表头', '微软雅黑', '10pt', '加粗金色字'),
        ('表格正文', '微软雅黑', '10pt', '深灰色'),
        ('页眉/页脚', '微软雅黑', '9pt', '浅灰色'),
    ]
)

add_sub_heading('3. 页边距与间距')
make_premium_table(
    headers=['项目', '标准值', '说明'],
    rows=[
        ('上边距', '2.5cm', '封面统一'),
        ('下边距', '2.0cm', ''),
        ('左边距', '2.8cm', '预留装订空间'),
        ('右边距', '2.8cm', ''),
        ('行距', '1.5倍', '正文'),
        ('段前距', '12pt', '标题前'),
        ('段后距', '6pt', '标题后'),
    ]
)

add_sub_heading('4. 表格规范')
add_body('① 表头：深蓝色(#0A1F3F)底 + 金色(#C5955C)字 + 加粗')
add_body('② 数据行：隔行用暖灰色(#F5F2EC)底')
add_body('③ 首列：加粗 + 左对齐')
add_body('④ 边框：上下边框用金色(#C5955C)，内部用浅褐色(#E0D8C8)')
add_body('⑤ 金额：保留两位小数，每三位用逗号分隔')

add_sub_heading('5. 封面规范')
add_body('① 封面使用统一模板（深蓝底+金色装饰线条+徽标）')
add_body('② 项目名称、招标编号、投标单位、日期为必填项')
add_body('③ 封面不编页码')

# ===== 保存 =====
output_path = r'D:\openclaw-workspace\bid_aba\融策标书模板_v2_高级版.docx'
doc.save(output_path)
print(f"✅ 高级标书模板 v2.0 已生成: {output_path}")

sz = os.path.getsize(output_path)
print(f"   文件大小: {sz:,}B ({sz//1024}KB)")
print(f"   包含页面：封面(插入图) → 目录 → 公司简介+资质表+流程图 → 实施方案+5张表 → 业绩表 → 服务承诺 → 排版规范说明")
