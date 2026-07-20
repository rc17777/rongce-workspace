"""
融策审计复核引擎 v3 - 精确提取xls表格数据
"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')

BASE = r'C:\Users\scrccpa\Desktop\新建文件夹'

projects = {}
for d in os.listdir(BASE):
    full_dir = os.path.join(BASE, d)
    if not os.path.isdir(full_dir):
        continue
    proj = {'dir': full_dir, 'files': {}}
    for f in os.listdir(full_dir):
        fp = os.path.join(full_dir, f)
        if f.endswith('.docx'):
            if '审核报告' in f and '说明' not in f:
                proj['files']['报告'] = fp
            elif '说明' in f:
                proj['files']['编制说明'] = fp
        elif f.endswith('.xls'):
            if '附件1' in f or '报表' in f:
                proj['files']['附表1'] = fp
            elif '附件2' in f or '审核表' in f:
                proj['files']['附表2'] = fp
    key = d[:20]
    projects[key] = proj

from docx import Document
import xlrd

def read_docx_text(fp):
    doc = Document(fp)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: lines.append(t)
    for i, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            lines.append('|'.join(cells))
    return '\n'.join(lines)

def extract_xls_data(fp, sheet_name=None):
    """提取xls中具体sheet的数据为行记录"""
    book = xlrd.open_workbook(fp, formatting_info=False)
    
    if sheet_name:
        sheets_to_check = [s for s in book.sheet_names() if sheet_name in s]
        if not sheets_to_check:
            sheets_to_check = [book.sheet_names()[0]]
    else:
        sheets_to_check = book.sheet_names()
    
    rows = []
    for sn in sheets_to_check:
        sh = book.sheet_by_name(sn)
        for r in range(sh.nrows):
            row = []
            for c in range(sh.ncols):
                cell = sh.cell(r,c)
                v = cell.value
                if cell.ctype == 2:  # number
                    row.append(('num', v))
                elif v:
                    row.append(('txt', str(v).strip()))
                else:
                    row.append(('emp', ''))
            rows.append((sn, r, row))
    return rows

# 专门提取财务决算表(02表)和审核汇总表(01表)的数据
def extract_finance_data(fp, sheet_prefix):
    """提取财务相关sheet的数据"""
    book = xlrd.open_workbook(fp, formatting_info=False)
    
    for sn in book.sheet_names():
        if not sn.startswith(sheet_prefix):
            continue
        sh = book.sheet_by_name(sn)
        print(f"\n  Sheet: {sn} ({sh.nrows}行×{sh.ncols}列)")
        yield sn, sh

print("="*70)
print("  融策审计复核：报告↔附表精确勾稽对账")
print("="*70)

for proj_key, proj in projects.items():
    print(f"\n{'='*70}")
    proj_name = os.path.basename(proj['dir'])
    print(f"  项目：{proj_name}")
    print(f"{'='*70}")
    
    files = proj['files']
    for k in ['报告','编制说明','附表1','附表2']:
        fn = os.path.basename(files[k]) if k in files else '未找到'
        print(f"  {k}: {fn}")
    
    # ===== 1. 从审核报告提取关键数据 =====
    report_text = ''
    if '报告' in files:
        report_text = read_docx_text(files['报告'])
    
    if '编制说明' in files:
        note_text = read_docx_text(files['编制说明'])
        report_text += '\n' + note_text
    
    # 用更精确的正则提取金额
    # 报告中的典型表述：
    # "送审金额为29,788,228.07元"
    # "审定金额为24,841,922.02元"
    # "核减金额为4,946,306.05元"
    # "项目概算总投资2026万元"
    
    report_data = {}
    
    # 提取概算
    m = re.search(r'概算[^0-9]*?(\d[\d,.]*\d)\s*(?:万?元?)', report_text)
    if m:
        val = float(m.group(1).replace(',',''))
        report_data['概算总投资'] = round(val, 4)
    
    # 提取项目总投资
    m = re.search(r'项目总投资[^0-9]*?(\d[\d,.]*\d)\s*(万?元?)', report_text)
    if m:
        val = float(m.group(1).replace(',',''))
        if m.group(2) == '元':
            val = round(val/10000, 4)
        report_data['项目总投资'] = round(val, 4)
    
    # 提取送审
    m = re.search(r'送审[^0-9]*?(\d[\d,.]*\d)\s*(元)', report_text)
    if m:
        val = float(m.group(1).replace(',','')) / 10000
        report_data['送审金额'] = round(val, 4)
    
    # 提取审定（多个审定金额，取第一个主要审定）
    # 模式: "审定金额为..." 或 "审定价"
    m = re.search(r'审定[^0-9]*?(\d[\d,.]*\d)\s*(元)', report_text)
    if m:
        val = float(m.group(1).replace(',','')) / 10000
        report_data['审定金额'] = round(val, 4)
    
    # 提取核减
    m = re.search(r'核减[^0-9]*?(\d[\d,.]*\d)\s*(元)', report_text)
    if m:
        val = float(m.group(1).replace(',','')) / 10000
        report_data['核减金额'] = round(val, 4)
    
    # 建安工程投资
    for kw in ['建筑安装工程','建安工程投资','建安投资']:
        m = re.search(rf'{kw}[^0-9]*?(\d[\d,.]*\d)\s*(元)', report_text)
        if m:
            val = float(m.group(1).replace(',','')) / 10000
            report_data['建安工程投资'] = round(val, 4)
            break
    
    # 待摊投资
    m = re.search(r'待摊投资[^0-9]*?(\d[\d,.]*\d)\s*(元)', report_text)
    if m:
        val = float(m.group(1).replace(',','')) / 10000
        report_data['待摊投资'] = round(val, 4)
    
    # 已支付
    m = re.search(r'已支付[^0-9]*?(\d[\d,.]*\d)\s*(元)', report_text)
    if m:
        val = float(m.group(1).replace(',','')) / 10000
        report_data['已支付'] = round(val, 4)
    
    print(f"\n  【报告提取数据】")
    for k, v in report_data.items():
        print(f"    {k}: {v:.4f}万元")
    
    # ===== 2. 从附表1(财务决算表)提取关键数据 =====
    table1_data = {}
    if '附表1' in files:
        print(f"\n  【附表1 - 财务决算表（02表）】")
        for sn, sh in extract_finance_data(files['附表1'], '02'):
            # 查找：工程名称 / 总投资 / 建安投资 / 待摊投资
            cols_map = {}
            header_row = -1
            for r in range(min(5, sh.nrows)):
                for c in range(sh.ncols):
                    v = str(sh.cell(r,c).value).strip()
                    if '项目' in v or '工程' in v or '概算' in v or '决算' in v:
                        header_row = r
                    if v in ['项目','工程项目','费用名称']:
                        cols_map['name'] = c
                    elif '概算' in v:
                        cols_map['概算'] = c
                    elif '决算' in v or '实际' in v:
                        cols_map['决算'] = c
            
            # 如果没找到表头，直接尝试按行读取
            for r in range(3, sh.nrows):
                # 读取第1列作为项目名称
                name = str(sh.cell(r, 0).value).strip()
                if not name:
                    # 尝试第2列
                    name = str(sh.cell(r, 1).value).strip()
                
                # 找数字
                vals = {}
                for c in range(sh.ncols):
                    cell = sh.cell(r, c)
                    if cell.ctype == 2 and cell.value > 0:
                        # 判断是哪一列
                        header = str(sh.cell(0, c).value).strip() if sh.nrows > 0 else ''
                        if '概算' in header:
                            vals['概算'] = cell.value
                        elif '决算' in header or '实际' in header:
                            vals['决算'] = cell.value
                        else:
                            vals[f'列{c}'] = cell.value
                
                # 输出有价值的数据行
                if name and vals:
                    print(f"    {name}: {vals}")
                    for k, v in vals.items():
                        table1_data[f'{name}_{k}'] = v
            
            # 直接输出所有行
            print(f"\n    【02表原始数据】")
            for r in range(sh.nrows):
                row = []
                for c in range(min(9, sh.ncols)):
                    cell = sh.cell(r,c)
                    v = cell.value
                    if cell.ctype == 2 and v > 0:
                        row.append(f"[C{c}]{v:.2f}")
                    elif v:
                        row.append(f"[C{c}]{str(v)[:20]}")
                if any(x for x in row):
                    print(f"    R{r}: {' | '.join(row)}")
    
    # ===== 3. 从附表2(审核汇总表)提取关键数据 =====
    table2_data = {}
    if '附表2' in files:
        print(f"\n  【附表2 - 审核汇总表（01表）】")
        for sn, sh in extract_finance_data(files['附表2'], '01'):
            for r in range(sh.nrows):
                row = []
                for c in range(min(8, sh.ncols)):
                    cell = sh.cell(r,c)
                    v = cell.value
                    if cell.ctype == 2 and v > 0:
                        row.append(f"[C{c}]{v:.2f}")
                    elif v:
                        row.append(f"[C{c}]{str(v)[:20]}")
                if any(x for x in row):
                    print(f"    R{r}: {' | '.join(row)}")
            
            # 提取关键数据项
            for r in range(3, sh.nrows):
                col0 = str(sh.cell(r, 0).value).strip()
                col1 = str(sh.cell(r, 1).value).strip()
                
                # 找送审审定数据行
                for c in range(sh.ncols):
                    cell = sh.cell(r,c)
                    if cell.ctype == 2 and cell.value > 1000:
                        header = str(sh.cell(0,c).value).strip() if sh.nrows > 0 else str(sh.cell(1,c).value).strip()
                        key = f"{col0}_{col1}_{header}" if header else f"{col0}_{col1}_C{c}"
                        table2_data[key] = cell.value
                        
                        if col0:
                            tag = col0[:12]
                            print(f"    {tag} | C{c}={cell.value:.2f}")

    # ===== 4. 勾稽复核 =====
    print(f"\n  {'='*60}")
    print(f"  【勾稽复核结果】")
    print(f"  {'='*60}")
    
    findings = []
    
    # 重点复核1：报告送审 vs 附表审定数据
    if '送审金额' in report_data and '审定金额' in report_data:
        ss = report_data['送审金额']
        ds = report_data['审定金额']
        if ss and ds:
            diff = ss - ds
            print(f"  报告送审({ss:.4f}万) - 报告审定({ds:.4f}万) = {diff:.4f}万（核减额）")
            findings.append(('报告送审-审定差', diff, ss, ds))
            
            if '核减金额' in report_data:
                hj = report_data['核减金额']
                if abs(diff - hj) < 0.01:
                    print(f"    ✅ 核减金额{hj:.4f}万 与计算一致")
                else:
                    print(f"    ⚠️ 核减金额{hj:.4f}万 与计算{diff:.4f}万 差异{abs(diff-hj):.4f}万")
    
    # 复核2：建安+待摊 vs 总投资
    if '建安工程投资' in report_data and '待摊投资' in report_data:
        ja = report_data['建安工程投资']
        dt = report_data['待摊投资']
        sub = ja + dt
        print(f"  建安({ja:.4f}万) + 待摊({dt:.4f}万) = {sub:.4f}万")
        if '项目总投资' in report_data:
            ztz = report_data['项目总投资']
            diff = ztz - sub
            if abs(diff) < 0.01:
                print(f"    ✅ 与项目总投资{ztz:.4f}万一致")
            elif abs(diff) < 100:
                print(f"    ⚠️ 与项目总投资{ztz:.4f}万差{diff:.4f}万（可能含设备/其他投资）")
            else:
                print(f"    ❌ 与项目总投资{ztz:.4f}万差{diff:.4f}万")
    
    # 复核3：概算 vs 送审
    if '概算总投资' in report_data and '送审金额' in report_data:
        gs = report_data['概算总投资']
        ss = report_data['送审金额']
        diff = ss - gs
        print(f"  概算({gs:.4f}万) vs 送审({ss:.4f}万): {'一致' if abs(diff)<0.01 else f'送审超概{diff:.4f}万'}")
    
    print()

# 保存结果
print(f"\n✅ 复核完成！详细数据已输出如上。")