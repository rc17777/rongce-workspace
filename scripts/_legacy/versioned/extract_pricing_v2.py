#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""精确提取校服投标报价表和公司信息"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from pathlib import Path

BASE = Path(r"C:\Users\scrccpa\Desktop\校服\2025年校服采购\校服\2025年\投标文件\投标文件")

targets = [
    (BASE / "四川乐吉玛帝诺服饰有限公司" / "成都市教育科学研究院附属中学学生校服采购项目其他投标文件.docx", "乐吉玛帝诺"),
    (BASE / "四川牧森服饰有限公司" / "成都市教育科学研究院附属中学其他投标文件-最终版.docx", "牧森"),
    (BASE / "江苏苏美达伊顿纪德品牌管理有限公司" / "20251012【其他投标文件】科教院附中.docx", "苏美达伊顿纪德"),
]

def extract_tables_near_keyword(doc, keywords, max_distance=5):
    """Extract tables that appear near keyword paragraphs"""
    results = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(kw in text for kw in keywords):
            # Look for nearest table after this paragraph
            # In docx, tables are separate from paragraphs
            pass
    
    # Alternative: extract all tables and label them
    for ti, table in enumerate(doc.tables):
        all_rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # Skip fully empty rows
            if any(c for c in cells):
                all_rows.append(' | '.join(cells))
        
        table_text = '\n'.join(all_rows)
        
        # Check if it's a pricing table
        if any(kw in table_text for kw in ['开标一览表', '分项报价', '投标报价', '单价', '元/套']):
            results.append((ti, all_rows))
    
    return results

def extract_company_info(doc):
    """Extract company registration info"""
    info = {}
    patterns = {
        'company_name': r'(?:投标人名称|单位名称|供应商名称)[：:]\s*(.+)',
        'address': r'(?:地\s*址|通讯地址)[：:]\s*(.+)',
        'legal_rep': r'(?:法定代表人|单位负责人)[：:]?\s*姓名[：:]?\s*(.+?)(?:\s|性别|年龄)',
        'contact': r'(?:被授权人|授权代表)[：:]?\s*(.+?)(?:\s|职务)',
        'phone': r'(?:电话|联系电话)[：:]\s*(\S+)',
        'registered_capital': r'注册资本[：:]\s*(.+)',
    }
    
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    
    for key, pattern in patterns.items():
        m = re.search(pattern, all_text)
        if m:
            info[key] = m.group(1).strip()[:80]
    
    return info

for path, company in targets:
    if not path.exists():
        continue
    
    print(f"\n{'='*70}")
    print(f"📊 {company}")
    print(f"{'='*70}")
    
    doc = Document(str(path))
    
    # Extract price tables
    tables = extract_tables_near_keyword(doc, ['开标一览表', '分项报价'])
    for ti, rows in tables:
        print(f"\n[Table {ti+1}]")
        for row in rows:
            print(f"  {row[:200]}")
    
    # Extract company info
    info = extract_company_info(doc)
    if info:
        print(f"\n[公司信息]")
        for k, v in info.items():
            print(f"  {k}: {v}")

print("\nDone!")
