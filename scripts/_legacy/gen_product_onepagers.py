#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成两个产品一页纸"""
import sys,os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DARK = RGBColor(0x0A, 0x1F, 0x3F)
GOLD = RGBColor(0xC5, 0x95, 0x5C)
TEAL = RGBColor(0x1A, 0x5C, 0x6E)

def make_header_cell(cell, text):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell._element.get_or_add_tcPr().append(
        cell._element.makeelement(qn('w:shd'), {qn('w:fill'): '0A1F3F', qn('w:val'): 'clear'})
    )

# ===== 产品一：围标检测 =====
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

h = doc.add_heading('融策智能围标串标检测系统', level=0)
for run in h.runs:
    run.font.color.rgb = DARK

doc.add_paragraph('产品一页纸  |  四川融策会计师事务所  |  2026年6月').runs[0].font.color.rgb = TEAL

doc.add_heading('一句话定位', level=2)
doc.add_paragraph('11层AI检测，让围标串标无所遁形——不需要代理机构配合，3层证据即可定案。')

doc.add_heading('技术架构：11层递进式检测体系', level=2)
t = doc.add_table(12, 3, style='Table Grid')
t.cell(0,0).text = '层级'
t.cell(0,1).text = '检测维度'
t.cell(0,2).text = '数据源'
for j in range(3):
    make_header_cell(t.cell(0,j), t.cell(0,j).text)
layers = [
    ('L1', '报价规律异常', '开标一览表'),
    ('L2', '投标IP/MAC地址', '代理后台日志'),
    ('L3', '文本雷同（TF-IDF）', '投标文件.docx'),
    ('L4', '图片哈希值重合', '.docx解压word/media/'),
    ('L5', '元数据交叉比对', 'core.xml/WPS GUID'),
    ('L6', '文档结构比对', '段落/表格/图片位置'),
    ('L7', '打印机/扫描仪型号', 'PDF Producer字段'),
    ('L8', '工商关联关系', '天眼查/企查查'),
    ('L9', '保证金资金链', '银行汇款凭证'),
    ('L10', '标书格式内容重合率', '全量文本/图片/字体'),
    ('L11', '意思联络证据', '微信/通话/协议（司法专用）'),
]
for i, (l, v, s) in enumerate(layers):
    t.cell(i+1,0).text = l
    t.cell(i+1,1).text = v
    t.cell(i+1,2).text = s

doc.add_heading('核心优势', level=2)
for item in [
    '🔑 无需代理配合：L3+L4+L5三杀即可定案——破解"代理不给IP"的行业顽疾',
    '📊 全量自动化：一份投标文件解压→文本提取→哈希计算→元数据读取→交叉比对，全自动',
    '🎯 铁证级输出：元数据指纹、图片哈希、TF-IDF相似度——数字证据，不可抵赖',
    '📋 取数函集成：配备7层破解法取数方案，代理不配合时有替代路径',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('典型案例', level=2)
doc.add_paragraph(
    '某县中小学厕所改扩建项目招标审计：8家投标方，6家公司PDF元数据指纹完全一致'
    '（同一扫描仪同一天同一设置）、声明函100%字体重合、所有投标书页数全同。'
    'L3+L4+L5+L10四层证据锁定围标。'
)

doc.add_heading('服务方式', level=2)
for item in [
    '按项目收费：5,000-20,000元/项目（视投标方数量和数据完整度）',
    '按年订阅：30,000-80,000元/年（全年招标项目不限量检测）',
    '交付物：《围标串标AI检测报告》含证据链、相似度矩阵、结论建议',
]:
    doc.add_paragraph(item, style='List Bullet')

out1 = r'D:\openclaw-workspace\bid_aba\产品一页纸-围标检测.docx'
doc.save(out1)
print(f'Done: {out1}')

# ===== 产品二：经责量化评价 =====
doc2 = Document()
doc2.styles['Normal'].font.name = '微软雅黑'
doc2.styles['Normal'].font.size = Pt(11)

h = doc2.add_heading('融策经责审计量化评价系统', level=0)
for run in h.runs:
    run.font.color.rgb = DARK

doc2.add_paragraph('产品一页纸  |  四川融策会计师事务所  |  2026年6月').runs[0].font.color.rgb = TEAL

doc2.add_heading('一句话定位', level=2)
doc2.add_paragraph('从"拍脑袋打分"到"量化可追溯"——7+1套指标体系，让经责审计有据可依、有据可查。')

doc2.add_heading('指标体系', level=2)
t2 = doc2.add_table(9, 2, style='Table Grid')
t2.cell(0,0).text = '体系'
t2.cell(0,1).text = '适用对象'
for j in range(2):
    make_header_cell(t2.cell(0,j), t2.cell(0,j).text)
indices = [
    ('工程项目', '建设/交通/水利/市政类'),
    ('行政事业', '机关/事业单位'),
    ('商业竞争', '市场化国企'),
    ('公益功能', '供水/供气/公交/环卫'),
    ('特定功能', '融资平台/金控/担保'),
    ('金融', '银行/保险/证券/基金'),
    ('科创', '科技/研发/孵化器'),
    ('自定义', '客户定制指标'),
]
for i, (a,b) in enumerate(indices):
    t2.cell(i+1,0).text = a
    t2.cell(i+1,1).text = b

doc2.add_heading('核心机制', level=2)
for item in [
    '⚖️ 6种一票否决：重大损失·严重违纪·重大舆情·安全事故·环保事件·系统性风险',
    '📐 四类调整因素：任期内外划断·宏观经济影响·不可抗力·政策变更',
    '📊 三级损失标准：一般<100万·较大100万~1000万·重大>1000万',
    '🔴 终身问责原则：退休≠安全，离职≠免责，调任≠清零',
]:
    doc2.add_paragraph(item, style='List Bullet')

doc2.add_heading('交付物', level=2)
doc2.add_paragraph('Excel量化评价工作表（含评分明细、证据链索引、调整因素说明）+ 评价结论摘要')

out2 = r'D:\openclaw-workspace\bid_aba\产品一页纸-经责评价.docx'
doc2.save(out2)
print(f'Done: {out2}')
