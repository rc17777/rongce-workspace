# -*- coding: utf-8 -*-
"""
在已生成的文档中插入配图（图文并茂增强版）
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc_path = r'D:\openclaw-workspace\output\模拟案例一_绩效目标编制辅导与审核工作方案.docx'
chart_dir = r'D:\openclaw-workspace\output\charts'

doc = Document(doc_path)

def insert_image_at_keyword(doc, keyword, img_path, width_inches=5.8, caption=None):
    """在包含关键词的段落后插入图片"""
    inserted = False
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text and not inserted:
            # 在找到的段落后面插入图片
            # 创建新的图片段落
            next_para = para
            
            # 空行
            spacer = doc.paragraphs[i]
            # 在段落之后插入空行和图片
            # 由于python-docx的限制，我们采用另一种方式：
            # 在实际段落对象后面插入
            img_para = doc.add_paragraph()
            # 需要把段落移到正确位置
            # 简化处理：在文档末尾附近的特定位置插入
            
            inserted = True
            break
    return inserted

# 由于python-docx的段落插入限制，采用按索引插入的方式
# 搜索关键段落并记录索引位置

def find_para_index(doc, keyword):
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            return i
    return None

def add_image_paragraph(doc, img_path, width_inches=5.8, caption=''):
    """添加居中图片段落"""
    if not os.path.exists(img_path):
        print(f'  警告：图片不存在 {img_path}')
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.font.size = Pt(10)
        run_cap.font.name = '楷体'
        run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()  # 空行

# 由于python-docx无法在段落中间插入，我们在文档末尾之前插入图片
# 采用策略：找到合适的插入点，使用doc.paragraphs移除后重新构建
# 更实用的方法：在特定段落后添加临时标记，然后重新生成

# 简单方案：识别关键位置附近，利用 doc.add_paragraph 追加到文档末尾，
# 然后移动元素...这太复杂了
# 更好的方案：直接在文档末尾之前（结语之前）添加配图章节

# 找到"结语"段落
jieyu_idx = None
for i, para in enumerate(doc.paragraphs):
    if '结  语' in para.text or '结语' in para.text:
        jieyu_idx = i
        break

# 找到"五、保障措施"之后的位置插入图表汇总
baozhang_idx = None
for i, para in enumerate(doc.paragraphs):
    if '五、保障措施与质量管控' in para.text:
        baozhang_idx = i
        break

# 更简单的方法：在文档末尾追加图片附录
# 找到结尾的落款
luokuan_idx = None
for i, para in enumerate(doc.paragraphs):
    if '四川融策会计师事务所有限公司' in para.text:
        luokuan_idx = i
        break

# 在落款前插入图表
if luokuan_idx:
    # 找到落款前的空行
    insert_idx = luokuan_idx - 1
    while insert_idx > 0 and doc.paragraphs[insert_idx].text.strip() == '':
        insert_idx -= 1
    insert_idx += 1  # 在第一个空行后

# 采用更直接的方法：在保障措施章节部分，通过重新排列段落来插入图片
# 但由于python-docx的限制，最简单的是：
# 1. 保存当前文档
# 2. 重新打开并通过底层XML插入图片

# 实际可行方案：使用底层lxml操作
from docx.oxml.ns import qn as nqn
from lxml import etree

# 找到"五、保障措施与质量管控"下的"（六）时间保障"表格后的位置
target_keywords = ['图1：绩效目标编制辅导与审核"五阶段闭环"工作流程',
                   '图2：绩效指标体系结构对比分析',
                   '图3：绩效指标分类分布对比',
                   '图4：现行绩效指标体系七维雷达诊断',
                   '图5：三级复核质量控制金字塔']

# 找到"工作流程总览"后插入图1
flow_idx = None
for i, para in enumerate(doc.paragraphs):
    if '工作流程总览' in para.text:
        flow_idx = i
        break

# 找到流程表格后的段落
if flow_idx:
    # 表格在段落之后，我们需要找到"图1："的那段
    for i in range(flow_idx, len(doc.paragraphs)):
        if '图1' in doc.paragraphs[i].text and '工作流程' in doc.paragraphs[i].text:
            # 在这段之后插入图片
            # 使用底层XML操作
            para_elem = doc.paragraphs[i]._element
            parent = para_elem.getparent()
            
            # 创建图片段落
            img_para_elem = etree.SubElement(parent, nqn('w:p'))
            # 创建空段落
            spacer_elem = etree.SubElement(parent, nqn('w:p'))
            
            # 在img_para中插入图片运行
            img_run_elem = etree.SubElement(img_para_elem, nqn('w:r'))
            drawing_elem = etree.SubElement(img_run_elem, nqn('w:drawing'))
            
            # 读取模板图片的XML
            # 需要创建inline drawing... 太复杂
            
            # 换个思路，用新建段落的方式
            # 在parent中移动元素
            parent.remove(img_para_elem)
            parent.remove(spacer_elem)
            
            # 在para_elem之后插入
            idx = list(parent).index(para_elem)
            parent.insert(idx + 1, img_para_elem)
            parent.insert(idx + 2, spacer_elem)
            break

# 上面的方法太复杂了。让我改用更简单的方案：
# 直接重新生成整个文档，在正确的位置插入图片

print("采用重新生成方式，整合图片到文档...")

# 重新生成一个整合版
exec(open(r'D:\openclaw-workspace\scripts\generate_simulation_report.py', encoding='utf-8').read().replace(
    "doc.save(output_path)",
    "# 暂不保存"
))

# 但这样也不行... 最简单的方法：
# 重新写一个完整的脚本，包含图片插入
print("请运行整合脚本...")
