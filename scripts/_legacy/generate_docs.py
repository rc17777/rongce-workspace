# -*- coding: utf-8 -*-
"""生成监督检查(D)和财政评审(F)的Word+Excel文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
import os

OUTPUT_DIR = r"D:\openclaw-workspace\references"

# ============================================================
# 通用样式工具
# ============================================================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table_to_doc(doc, headers, rows, col_widths=None):
    """便捷添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        from docx.shared import RGBColor as RC
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E79')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

# ============================================================
# WORD: 监督检查(D)
# ============================================================

def create_word_D():
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading('业务线D：监督检查 — 案例深度分析与操作卡', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph('覆盖案例：案例1（违规奖金1837万）、案例3（小金库2800万）、案例5（采购回扣370万）、案例6（违规担保8.9亿）、案例8（账外收费6800万）')
    doc.add_paragraph('编制日期：2026-05-09')
    doc.add_paragraph()

    # ---- 案例1 ----
    add_heading_styled(doc, '案例1：某国有能源集团 · 违规发放奖金1837万元', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('在工资总额之外，以"专项奖励""考核兑现"等名义，向集团本部及3家子公司班子成员12人发放奖金1837万元。资金来源：管理费用/营业外支出科目列支。审批路径：分管领导直接签批，绕过工资总额审核流程。')

    add_heading_styled(doc, '违规链条拆解', 2)
    add_table_to_doc(doc,
        ['环节', '正常做法', '实际操作', '漏洞'],
        [
            ['制度层面', '工资总额由国资委核定批复', '在总额外另设名目，绕开批复', '子公司资金支出缺乏集团层面总额监控'],
            ['审批层面', '薪酬委员会审议→董事会批准', '分管领导签字即发放', '"专项奖励"名目无明确定义，自由裁量权过大'],
            ['核算层面', '奖金通过"应付职工薪酬"核算', '混入管理费用/营业外支出', '科目滥用掩盖薪酬性质'],
            ['发放层面', '纳入工资总额台账', '单独制表、发放，不计入总额', '两套工资表并行'],
        ],
        [3, 4, 4, 4.5]
    )

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 获取国资委批复的工资总额，与应付职工薪酬贷方发生额比对，发现差异1837万')
    doc.add_paragraph('2. 从管理费用中提取摘要含"奖励""兑现""津贴"的凭证，逐笔穿透')
    doc.add_paragraph('3. 将12名班子成员全部收入加总，与个税申报比对，发现"第二张工资表"')

    add_heading_styled(doc, '关键证据', 2)
    for e in ['国资委工资总额批复文件', '"专项奖励"发放明细表（含签收记录）', '管理费用/营业外支出明细账（奖励类摘要）', '12名班子成员个税申报记录']:
        doc.add_paragraph(f'• {e}', style='List Bullet')

    # ---- 案例3 ----
    add_heading_styled(doc, '案例3：某省属建筑企业 · 挂靠项目私设小金库2800万', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('项目经理挂靠施工队，以"管理费""协调费"名义截留项目资金，在账外设立小金库2800万元。资金存在形式：个人银行卡 + 现金 + 第三方账户。用途：项目部日常开支和关系维护。')

    add_heading_styled(doc, '违规链条拆解', 2)
    add_table_to_doc(doc,
        ['环节', '正常做法', '实际操作', '漏洞'],
        [
            ['项目承接', '企业自有资质投标', '挂靠方以企业名义投标，缴纳"管理费"', '挂靠本身违规，管理费成为小金库源头'],
            ['资金拨付', '总包按进度拨付，全部入账', '拨付时分两笔：入账款 + 截留款', '总包未监控分包方资金使用'],
            ['成本归集', '所有支出凭发票入账', '截留部分无发票、无凭证', '以"管理费""协调费"代替真实票据'],
            ['资金存放', '对公账户统一管理', '个人卡、现金、第三方账户', '项目经理个人控制资金，无人制衡'],
        ],
        [3, 4, 4, 4.5]
    )

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 对比业主付款→总包→分包的完整链条，发现分包实际收款 < 合同约定')
    doc.add_paragraph('2. 大量付款无对应发票，摘要均为"管理费""协调费"')
    doc.add_paragraph('3. 项目部基层财务人员匿名举报')
    doc.add_paragraph('4. 项目经理名下发现大额异常资金往来')

    # ---- 案例5 ----
    add_heading_styled(doc, '案例5：某市属医院 · 耗材采购吃回扣370万', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('设备科科长与耗材供应商约定，按采购额的3%-5%收取回扣，5年累计370万元。回扣形式：现金 + 转账至亲属账户（约280万）+ 科室聚餐消费（约90万）。')

    add_heading_styled(doc, '违规链条拆解', 2)
    add_table_to_doc(doc,
        ['环节', '正常做法', '实际操作', '漏洞'],
        [
            ['供应商选择', '公开招标，采购委员会集体决策', '科室推荐供应商，招标流于形式', '设备科同时负责需求提出和供应商推荐'],
            ['价格谈判', '多家比价，公开竞价', '选定供应商后谈"返点"，价格虚高', '返点前置导致中标价已含回扣空间'],
            ['验收环节', '使用科室+设备科+纪委三方验收', '设备科单独验收', '缺乏验收制衡'],
            ['付款环节', '按合同进度付款', '回扣按采购额同步结算', '付款与回扣形成闭环节奏'],
        ],
        [3, 4, 4, 4.5]
    )

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 采购单价与同期市场公开报价对比，发现系统性偏高8%-15%')
    doc.add_paragraph('2. 3家供应商占全部耗材采购额85%，且长期未轮换')
    doc.add_paragraph('3. 设备科科长及其亲属银行账户大额资金流入时间与采购付款高度吻合')
    doc.add_paragraph('4. 科室聚餐频次、人均消费远超同院其他科室')

    # ---- 案例6 ----
    add_heading_styled(doc, '案例6：某市属国有企业 · 违规对外担保8.9亿元', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('集团在未经董事会审议、未上报国资监管部门的情况下，为3家民营关联企业提供连带责任担保，担保金额合计8.9亿元。其中1笔已代偿2.3亿元无法追回，担保余额占净资产比高达65%（远超国资委50%红线），且无任何反担保措施。')

    add_heading_styled(doc, '违规链条拆解', 2)
    add_table_to_doc(doc,
        ['环节', '正常做法', '实际操作', '漏洞'],
        [
            ['决策审批', '董事会审议→国资监管部门审批/备案', '主要领导个人拍板', '印章管理失控，担保合同未经法务审核即用印'],
            ['信息披露', '财务报告附注披露担保事项', '担保事项不入账、不披露', '或有负债未在报表反映，财务监督失效'],
            ['风险评估', '对被担保方尽职调查', '未做实质尽调，基于私人关系担保', '3家被担保企业均为民企，与领导存在关联'],
            ['事后管理', '跟踪被担保方经营、定期风险评估', '无人跟踪，代偿发生后才发现风险', '缺乏担保后管理机制'],
        ],
        [3, 4, 4, 4.5]
    )

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 向所有合作银行发函询证，银行回函揭示担保信息')
    doc.add_paragraph('2. 人民银行征信中心查询，发现对外担保记录')
    doc.add_paragraph('3. 调取印章使用登记簿，发现未经审批的担保合同用印')
    doc.add_paragraph('4. 8.9亿担保无任何反担保措施，明显异常')

    # ---- 案例8 ----
    add_heading_styled(doc, '案例8：某教育集团 · 违规收费账外管理6800万', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('集团旗下5所民办学校以"赞助费""培训费""校服费"等名义额外收取费用6800万元，存入个人账户或第三方账户，未纳入学校财务统一核算。')

    add_heading_styled(doc, '违规链条拆解', 2)
    add_table_to_doc(doc,
        ['环节', '正常做法', '实际操作', '漏洞'],
        [
            ['收费标准', '经物价/教育部门批准/备案', '批准项目之外另立名目收费', '监管只查批准项目不查"赞助""捐赠"类'],
            ['收费渠道', '学校公账统一收取', '指定个人账户或第三方账户收取', '家长缺乏辨别能力'],
            ['票据管理', '开具财政票据/税务发票', '不开票或开自制收据', '无票据即无监管痕迹'],
            ['资金核算', '全额纳入学校财务统一核算', '公账之外另立账册', '两套账并行，法定账套不反映真实状况'],
        ],
        [3, 4, 4, 4.5]
    )

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 根据学籍人数和批准收费标准测算应收学费，与账面入账数对比，发现缺口')
    doc.add_paragraph('2. 抽查学生家长实际缴费金额和收费项目，发现大量未入账项目')
    doc.add_paragraph('3. 学校负责人/财务人员个人账户发现大额、高频收款')
    doc.add_paragraph('4. 学校新建设施金额远超公账可支付范围，资金必有其他来源')

    # ---- 风险矩阵 ----
    add_heading_styled(doc, '风险矩阵', 1)
    add_table_to_doc(doc,
        ['编号', '风险描述', '案例', '可能性', '影响', '等级', '关键控制'],
        [
            ['D-R01', '工资总额外违规发放薪酬福利', '1', '中', '高', '🔴高', '工资总额预算→发放→个税三线比对'],
            ['D-R02', '项目部/分公司账外资金循环', '3', '高', '高', '🔴高', '全账户银行函证+资金流向穿透'],
            ['D-R03', '大额采购供应商围标+回扣', '5', '高', '高', '🔴高', '采购比价+供应商轮换+关联排查'],
            ['D-R04', '未经审批对外提供担保', '6', '中', '致命', '🔴高', '银行函证+征信查询+用印审查'],
            ['D-R05', '收费不入法定账簿', '8', '中', '高', '🔴高', '业务量×单价 vs 账面收入比对'],
            ['D-R06', '津补贴名实不符', '1', '中', '中', '🟡中', '津贴发放依据核实'],
            ['D-R07', '挂靠方截留工程款', '3', '高', '中', '🟡中', '付款比例 vs 合同约定比对'],
            ['D-R08', '供应商长期不轮换', '5', '高', '中', '🟡中', '供应商清单年度变动分析'],
            ['D-R09', '印章管理失控', '6', '低', '致命', '🟡中', '用印审批+登记+抽查'],
            ['D-R10', '票据管理混乱', '8', '中', '中', '🟡中', '票据领用/核销台账'],
        ],
        [2, 5, 1, 1.5, 1.5, 1.5, 5.5]
    )

    # ---- 操作卡 ----
    add_heading_styled(doc, '现场操作卡', 1)

    add_heading_styled(doc, '操作卡D-A：48小时进场快速扫描清单', 2)
    doc.add_paragraph('Day 1 上午：')
    for item in ['调取：工资总额批复文件、应付职工薪酬科目余额表', '调取：全部银行账户清单（含已销户）', '调取：对外担保台账（如无则要求书面确认零担保）']:
        doc.add_paragraph(f'  □ {item}')
    doc.add_paragraph('Day 1 下午：')
    for item in ['调取：收费许可证/批复、收费标准备案文件', '调取：近3年采购合同清单（按金额排序，取前20大）', '调取：近3年薪酬发放汇总表（分月/分项目）', '向主要合作银行发出询证函（担保+账户+存款）']:
        doc.add_paragraph(f'  □ {item}')
    doc.add_paragraph('Day 2 上午：')
    for item in ['数据比对①：工资总额 vs 实际发放', '数据比对②：业务量×单价 vs 账面收入', '数据比对③：采购中标价 vs 市场价']:
        doc.add_paragraph(f'  □ {item}')
    doc.add_paragraph('Day 2 下午：')
    for item in ['访谈财务负责人（重点：薪酬、担保、收费）', '访谈采购/后勤负责人（重点：供应商选择）', '突击保险柜盘查（现金+个人卡+账外凭证）', '形成初步风险判断，调整后续审计重点']:
        doc.add_paragraph(f'  □ {item}')

    add_heading_styled(doc, '操作卡D-B：小金库核查七步法', 2)
    steps = ['全账户函证 → 找出账外银行账户', '收入完整性测算 → 业务量×单价 vs 账面收入', '支出真实性检查 → 大额无票支出追踪', '个人卡排查 → 关键岗位人员银行流水', '突击盘点 → 保险柜现金+存放物品', '供应商/分包方延伸 → 向对方核实实际收款金额', '员工访谈/举报 → 财务人员+离职人员']
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    add_heading_styled(doc, '操作卡D-C：回扣发现五线索', 2)
    clues = ['价格异常 → 采购价持续高于市场价10%以上', '供应商锁定 → 同一供应商多年未轮换，市占率>50%', '采购量异常 → 采购量不随业务量变化（刚需采购）', '资金回流 → 供应商→采购人员/亲属账户资金往来', '生活消费异常 → 个人消费水平与收入明显不匹配']
    for i, c in enumerate(clues, 1):
        doc.add_paragraph(f'线索{i}：{c}')

    filepath = os.path.join(OUTPUT_DIR, '监督检查-案例分析与操作卡.docx')
    doc.save(filepath)
    print(f'[OK] {filepath}')
    return filepath


# ============================================================
# WORD: 财政评审(F)
# ============================================================

def create_word_F():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    title = doc.add_heading('业务线F：财政评审 — 案例深度分析与操作卡', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph('覆盖案例：案例1（违规奖金1837万）、案例4（研发造假套补1.2亿）、案例10（虚报数据套燃油补3100万）')
    doc.add_paragraph('编制日期：2026-05-09')
    doc.add_paragraph()

    # ---- 案例4 ----
    add_heading_styled(doc, '案例4：某科技集团 · 研发费用造假套取财政补贴1.2亿元', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('将已量产3个产品包装成"在研项目"，伪造研发立项文件、研发辅助账和虚假人员名单，累计虚报研发费用1.2亿元，套取高新技术补贴及税收优惠。')

    add_heading_styled(doc, '造假链条拆解', 2)
    add_table_to_doc(doc,
        ['造假环节', '正常做法', '具体造假手法', '破绽'],
        [
            ['立项文件', '真实形成立项报告、可研报告、预算审批', '修改已量产产品的立项日期，将投产日期改为"预计投产"', '立项文件中的技术指标与已量产产品完全一致'],
            ['人员名单', '实际参与研发的人员登记', '将生产车间工人、销售人员列入研发名单，并伪造工时', '社保记录显示"研发人员"实际在生产/销售岗位'],
            ['辅助账', '按项目真实归集研发费用', '将生产材料费、差旅费、招待费挪入研发辅助账', '同一发票在辅助账和生产账中重复列支'],
            ['成果证据', '真实测试报告、专利申请', '修改已有产品的测试报告日期、伪造专家验收意见', '专利与"在研项目"技术方案无关'],
        ],
        [2.5, 3.5, 4.5, 4.5]
    )

    add_heading_styled(doc, '套取路径与金额', 2)
    doc.add_paragraph('研发费用虚增 → 满足高新认定门槛 → 所得税优惠（15% vs 25%）：约7000万')
    doc.add_paragraph('研发费用加计扣除 → 应纳税所得额虚减 → 少缴所得税')
    doc.add_paragraph('申报各类科技补贴 → 满足研发投入条件 → 直接套取：约5000万')
    doc.add_paragraph('合计套取：1.2亿元')

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 查阅产品目录，发现3个"在研项目"对应的产品已在3年前上市')
    doc.add_paragraph('2. 比对个税/社保申报，发现"研发团队"80%人员实际岗位为生产和销售')
    doc.add_paragraph('3. 研发辅助账中的材料费、差旅费已在总账其他科目列支')
    doc.add_paragraph('4. 现场查看研发场地，"实验室"实为生产车间')

    # ---- 案例10 ----
    add_heading_styled(doc, '案例10：某市公交集团 · 套取燃油补贴3100万元', 1)
    add_heading_styled(doc, '违规手段', 2)
    doc.add_paragraph('通过虚报运营里程（篡改GPS数据，虚增35%）、伪造车辆油耗数据（虚增20%），累计套取国家燃油补贴3100万元。')
    doc.add_paragraph('套取公式：燃油补贴 = 运营里程 × 百公里油耗 × 补贴系数')
    doc.add_paragraph('真实应补：2200万km × 28L × 系数 ≈ 1848万')
    doc.add_paragraph('申报金额：2970万km × 33.6L × 系数 ≈ 4948万')
    doc.add_paragraph('套取差额：4948 - 1848 = 3100万元')

    add_heading_styled(doc, '造假链条拆解', 2)
    add_table_to_doc(doc,
        ['造假环节', '正常做法', '具体手法', '破绽'],
        [
            ['GPS数据', '直接从车辆监控平台导出，不可修改', '导出后手工篡改Excel，或让服务商生成"定制"报表', 'GPS原始数据库日志与导出报表不一致'],
            ['加油记录', '加油卡系统数据 + 纸质加油单', '伪造加油小票、虚增加油量', '加油卡系统数据可与加油站交叉比对'],
            ['车型油耗', '工信部公告油耗数据', '申报时使用偏高数据', '同型号车辆油耗应为固定值'],
            ['车辆范围', '仅限营运公交车辆', '将维修/通勤/已报废车辆混入', '车辆GPS轨迹可验证是否实际营运'],
        ],
        [2.5, 3.5, 4.5, 4.5]
    )

    add_heading_styled(doc, '审计发现路径', 2)
    doc.add_paragraph('1. 直接登录GPS监控平台数据库导出原始数据，与申报报表对比，差异35%')
    doc.add_paragraph('2. 调取加油站加油卡系统数据，与申报油耗数据比对')
    doc.add_paragraph('3. 查询工信部公告的同型号公交车百公里油耗，申报数据偏差巨大')
    doc.add_paragraph('4. 维修记录显示大量车辆长期停驶，但GPS数据仍显示"运营"')

    # ---- 案例1（财政评审视角）----
    add_heading_styled(doc, '案例1在财政评审中的视角', 1)
    doc.add_paragraph('在财政评审中，案例1（违规发放奖金1837万）的评审重点聚焦于预算编制准确性和执行合规性：')
    add_table_to_doc(doc,
        ['评审维度', '正常标准', '实际执行', '问题'],
        [
            ['预算编制准确性', '工资总额预算覆盖全部薪酬', '预算编制未纳入"专项奖励"', '偏差18.37%'],
            ['预算执行合规性', '工资总额不突破批复', '突破1837万', '违规发放'],
            ['科目使用规范性', '薪酬通过应付职工薪酬核算', '混入管理费用/营业外支出', '科目滥用'],
            ['资金使用效率', '薪酬与绩效挂钩', '名实不符，缺乏绩效依据', '国有资本收益被侵蚀'],
        ],
        [3, 4, 4, 4]
    )

    # ---- 风险矩阵 ----
    add_heading_styled(doc, '风险矩阵', 1)
    add_table_to_doc(doc,
        ['编号', '风险描述', '案例', '可能性', '影响', '等级', '关键控制'],
        [
            ['F-R01', '研发项目虚构，套取高新补贴及税收优惠', '4', '高', '致命', '🔴高', '研发项目真实性核查'],
            ['F-R02', '申报数据造假套取专项补贴', '10', '高', '高', '🔴高', '原始系统数据 vs 申报数据比对'],
            ['F-R03', '研发费用归集不实，虚增多列', '4', '高', '高', '🔴高', '辅助账 vs 总账交叉比对'],
            ['F-R04', '工资总额预算与实际执行严重偏离', '1', '中', '高', '🔴高', '预算批复→执行→决算全链条比对'],
            ['F-R05', '非研发人员混入研发队伍', '4', '高', '中', '🟡中', '研发人员名单 vs 社保/个税岗位信息'],
            ['F-R06', '科目滥用掩盖真实经济业务', '1', '中', '中', '🟡中', '管理费用/营业外支出摘要关键字筛查'],
            ['F-R07', '补贴申报与税收优惠共享同一套虚假材料', '4', '中', '高', '🟡中', '各口申报材料横向比对'],
            ['F-R08', 'GPS/加油等第三方系统数据被篡改', '10', '中', '高', '🟡中', '直接从源系统数据库提取'],
        ],
        [2, 5, 1, 1.5, 1.5, 1.5, 5.5]
    )

    # ---- 操作卡 ----
    add_heading_styled(doc, '现场操作卡', 1)

    add_heading_styled(doc, '操作卡F-A：研发费用造假识别六步法', 2)
    steps = [
        '产品上市时间 vs 研发立项时间 → 产品已上市却还在"研发"= 项目虚构',
        '研发人员名册 vs 社保/个税岗位 → 大量非研发岗位人员列入 = 人员造假',
        '研发辅助账 vs 财务总账 → 同一笔费用在两套账中出现 = 重复列支',
        '研发场地现场查看 → "研发实验室"实为车间/办公室 = 项目虚构',
        '研发成果核验 → 无专利/无样品/无测试报告 = 可能存在虚假',
        '多口申报数据横向比对 → 同一指标在不同申报中数值不同 = 至少一处造假',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    add_heading_styled(doc, '操作卡F-B：补贴数据真实性核查五步骤', 2)
    steps2 = [
        '确定原始数据源（GPS数据库/加油卡系统/车管系统等）',
        '直接从源系统导出原始数据（不走申报报表）',
        '源数据 vs 申报数据逐项比对',
        '偏差＞5% → 深入追溯原因',
        '无法从源系统导出 → 数据可靠性存疑，标注高风险',
    ]
    for i, s in enumerate(steps2, 1):
        doc.add_paragraph(f'{i}. {s}')

    add_heading_styled(doc, '操作卡F-C：财政评审进场48小时清单', 2)
    doc.add_paragraph('Day 1 上午：')
    for item in ['获取近3年工资总额批复及调整文件', '获取近3年所有财政补贴申报及批复清单', '获取高新认定全套材料（如适用）']:
        doc.add_paragraph(f'  □ {item}')
    doc.add_paragraph('Day 1 下午：')
    for item in ['获取应付职工薪酬科目余额表（分月）', '获取研发费用辅助账（如适用）', '获取个税全员全额申报数据', '获取GPS/加油卡等业务系统数据库访问权限']:
        doc.add_paragraph(f'  □ {item}')
    doc.add_paragraph('Day 2 上午：')
    for item in ['数据比对①：工资总额批复 vs 实际执行', '数据比对②：研发人员名单 vs 社保岗位信息', '数据比对③：补贴申报数据 vs 源系统数据']:
        doc.add_paragraph(f'  □ {item}')
    doc.add_paragraph('Day 2 下午：')
    for item in ['现场查看：研发场所、设备、人员', '访谈：研发负责人、财务负责人、人力资源负责人', '形成初步风险判断，调整后续审计重点']:
        doc.add_paragraph(f'  □ {item}')

    add_heading_styled(doc, '操作卡F-D：财政资金"一鱼多吃"快速筛查', 2)
    doc.add_paragraph('取一笔大额支出 → 追问资金去向：')
    for q in ['在研发辅助账中列支了吗？', '在加计扣除中申报了吗？', '在补贴申报中使用了吗？', '在常规成本中列支了吗？']:
        doc.add_paragraph(f'  □ {q}')
    doc.add_paragraph('→ 同一笔支出出现在≥2个申报渠道 = "一鱼多吃"嫌疑')

    filepath = os.path.join(OUTPUT_DIR, '财政评审-案例分析与操作卡.docx')
    doc.save(filepath)
    print(f'[OK] {filepath}')
    return filepath


# ============================================================
# EXCEL: 监督检查(D) 五套工作底稿
# ============================================================

def create_excel_D():
    wb = openpyxl.Workbook()
    # 删除默认sheet
    wb.remove(wb.active)

    # 通用样式
    header_font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell_font = Font(name='微软雅黑', size=9)
    cell_align = Alignment(vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    title_font = Font(name='微软雅黑', size=14, bold=True, color='1F4E79')
    subtitle_font = Font(name='微软雅黑', size=9, color='666666')

    def create_sheet(ws, title_text, headers, row_labels=None, col_count=6, row_count=6):
        """创建标准底稿sheet"""
        # 标题
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
        ws['A1'] = title_text
        ws['A1'].font = title_font
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30

        # 信息行
        info_items = ['被检查单位：_______________', '检查期间：20__年__月至20__年__月',
                      '编制人/日期：___ / 20__.__.__', '复核人/日期：___ / 20__.__.__']
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=int(col_count/2))
        ws.merge_cells(start_row=2, start_column=int(col_count/2)+1, end_row=2, end_column=col_count)
        ws['A2'] = f'项目编号：[  ]-___-001    {info_items[0]}    {info_items[1]}'
        ws['A2'].font = subtitle_font
        ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=col_count)
        ws['A3'] = f'{info_items[2]}    {info_items[3]}'
        ws['A3'].font = subtitle_font

        # 表头
        start_row = 5
        for i, h in enumerate(headers):
            cell = ws.cell(row=start_row, column=i+1, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
        ws.row_dimensions[start_row].height = 25

        # 数据行
        for r in range(row_count):
            row_num = start_row + 1 + r
            ws.row_dimensions[row_num].height = 22
            for c in range(len(headers)):
                cell = ws.cell(row=row_num, column=c+1, value='')
                cell.font = cell_font
                cell.alignment = cell_align
                cell.border = thin_border

        # 列宽
        for i in range(len(headers)):
            ws.column_dimensions[get_column_letter(i+1)].width = 16

        return start_row + 1 + row_count + 1  # 返回下一个可用行

    # ---- D-01 薪酬合规检查表 ----
    ws1 = wb.create_sheet('D-01 薪酬合规检查')
    ws1.sheet_properties.tabColor = '1F4E79'

    ws1.merge_cells('A1:F1')
    ws1['A1'] = '底稿D-01：薪酬合规检查表'
    ws1['A1'].font = title_font
    ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws1.row_dimensions[1].height = 30
    ws1.merge_cells('A2:F2')
    ws1['A2'] = '项目编号：[  ]-薪酬-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws1['A2'].font = subtitle_font

    # 一、工资总额批复与发放比对
    r = 4
    ws1.merge_cells(f'A{r}:F{r}')
    ws1[f'A{r}'] = '一、工资总额批复与发放比对'
    ws1[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)

    headers1 = ['年度', '国资委/主管部门批复总额(A)', '应付职工薪酬贷方发生额(B)', '差异(C=B-A)', '差异率', '说明']
    for i, h in enumerate(headers1):
        cell = ws1.cell(row=5, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(6, 10):
        for c in range(1, 7):
            cell = ws1.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
    ws1.merge_cells('A9:F9')
    ws1['A9'] = '合计'
    ws1['A9'].font = Font(name='微软雅黑', size=9, bold=True)
    ws1['A9'].alignment = cell_align

    # 二、工资表与银行代发比对
    r = 11
    ws1.merge_cells(f'A{r}:F{r}')
    ws1[f'A{r}'] = '二、工资表与银行代发比对'
    ws1[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    headers_b = ['月份', '工资表人数', '工资表总额', '银行代发人数', '银行代发总额', '差异说明']
    for i, h in enumerate(headers_b):
        cell = ws1.cell(row=12, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(13, 17):
        for c in range(1, 7):
            cell = ws1.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 三、非常规发放项目排查
    r = 18
    ws1.merge_cells(f'A{r}:F{r}')
    ws1[f'A{r}'] = '三、非常规发放项目排查'
    ws1[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    headers_c = ['发放名目', '发放期间', '发放金额', '审批文件', '是否在工资总额内', '结论']
    for i, h in enumerate(headers_c):
        cell = ws1.cell(row=19, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    items_c = ['专项奖励', '考核兑现', '特别贡献', '节日慰问', '(其他)']
    for idx, item in enumerate(items_c):
        row_n = 20 + idx
        for c in range(1, 7):
            cell = ws1.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws1.cell(row=row_n, column=1, value=item).font = cell_font

    # 四、班子成员收入汇总
    r = 26
    ws1.merge_cells(f'A{r}:F{r}')
    ws1[f'A{r}'] = '四、班子成员收入汇总'
    ws1[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    headers_d = ['姓名', '职务', '工资总额内收入', '工资总额外收入', '个税申报收入', '差异']
    for i, h in enumerate(headers_d):
        cell = ws1.cell(row=27, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(28, 34):
        for c in range(1, 7):
            cell = ws1.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 五、检查结论
    r = 35
    ws1.merge_cells(f'A{r}:F{r}')
    ws1[f'A{r}'] = '五、检查结论'
    ws1[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    ws1.merge_cells('A36:F36')
    ws1['A36'] = '□ 未发现异常'
    ws1['A36'].font = cell_font
    ws1.merge_cells('A37:F38')
    ws1['A37'] = '□ 发现以下问题：'
    ws1['A37'].font = cell_font

    # 列宽
    widths1 = [14, 22, 18, 14, 10, 20]
    for i, w in enumerate(widths1):
        ws1.column_dimensions[get_column_letter(i+1)].width = w

    # ---- D-02 小金库专项排查表 ----
    ws2 = wb.create_sheet('D-02 小金库排查')
    ws2.sheet_properties.tabColor = 'C00000'

    ws2.merge_cells('A1:G1')
    ws2['A1'] = '底稿D-02：小金库专项排查表'
    ws2['A1'].font = title_font; ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws2.row_dimensions[1].height = 30
    ws2.merge_cells('A2:G2')
    ws2['A2'] = '项目编号：[  ]-小金库-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws2['A2'].font = subtitle_font

    # 一、银行账户完整性核查
    r = 4
    ws2.merge_cells(f'A{r}:G{r}')
    ws2[f'A{r}'] = '一、银行账户完整性核查'
    ws2[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h1 = ['序号', '开户行', '账号', '账户性质', '财务账是否反映', '函证结果', '备注']
    for i, h in enumerate(h1):
        cell = ws2.cell(row=5, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(6, 14):
        for c in range(1, 8):
            cell = ws2.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 二、项目部资金流向
    r = 15
    ws2.merge_cells(f'A{r}:G{r}')
    ws2[f'A{r}'] = '二、项目部/分公司资金流向检查'
    ws2[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h2 = ['项目部/分公司', '合同收入(A)', '对公账户收款(B)', '差异(A-B)', '资金去向', '', '']
    for i, h in enumerate(h2):
        cell = ws2.cell(row=16, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    # merge last three cells
    ws2.merge_cells('E16:G16')
    ws2['E16'] = '资金去向'
    ws2['E16'].font = header_font; ws2['E16'].fill = header_fill; ws2['E16'].alignment = header_align; ws2['E16'].border = thin_border
    for row_n in range(17, 22):
        for c in range(1, 8):
            cell = ws2.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 三、个人账户排查
    r = 23
    ws2.merge_cells(f'A{r}:G{r}')
    ws2[f'A{r}'] = '三、个人账户排查（高风险人员）'
    ws2[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h3 = ['姓名', '职务', '账户数量', '检查期间流入', '检查期间流出', '可疑交易描述', '']
    for i, h in enumerate(h3):
        cell = ws2.cell(row=24, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws2.merge_cells('F24:G24')
    ws2['F24'] = '可疑交易描述'
    ws2['F24'].font = header_font; ws2['F24'].fill = header_fill; ws2['F24'].alignment = header_align; ws2['F24'].border = thin_border
    for row_n in range(25, 30):
        for c in range(1, 8):
            cell = ws2.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 四、大额无票支出
    r = 31
    ws2.merge_cells(f'A{r}:G{r}')
    ws2[f'A{r}'] = '四、大额无票支出检查（摘要含"管理费""协调费"等）'
    ws2[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h4 = ['日期', '凭证号', '摘要', '金额', '收款方', '是否有发票', '是否有审批']
    for i, h in enumerate(h4):
        cell = ws2.cell(row=32, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(33, 39):
        for c in range(1, 8):
            cell = ws2.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 五、现金盘点
    r = 40
    ws2.merge_cells(f'A{r}:G{r}')
    ws2[f'A{r}'] = '五、现金盘点与突击检查'
    ws2[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    check_items = [
        ('保险柜现金盘点', '账面___ / 实盘___ / 差异___'),
        ('是否存在个人银行卡存放于单位', '□是 □否'),
        ('是否存在账外账/账外凭证', '□是 □否'),
        ('是否存在未入账收款收据', '□是 □否'),
    ]
    for idx, (item, val) in enumerate(check_items):
        ws2.merge_cells(f'A{41+idx}:C{41+idx}')
        ws2.merge_cells(f'D{41+idx}:G{41+idx}')
        ws2[f'A{41+idx}'] = item
        ws2[f'D{41+idx}'] = val
        for c in range(1, 8):
            ws2.cell(row=41+idx, column=c).font = cell_font
            ws2.cell(row=41+idx, column=c).border = thin_border

    # 六、检查结论
    r = 46
    ws2.merge_cells(f'A{r}:G{r}')
    ws2[f'A{r}'] = '六、检查结论'
    ws2[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    ws2.merge_cells('A47:G47')
    ws2['A47'] = '□ 未发现小金库'
    ws2['A47'].font = cell_font
    ws2.merge_cells('A48:G49')
    ws2['A48'] = '□ 发现小金库___个，金额合计___元'
    ws2['A48'].font = cell_font

    widths2 = [6, 14, 18, 12, 12, 12, 20]
    for i, w in enumerate(widths2):
        ws2.column_dimensions[get_column_letter(i+1)].width = w

    # ---- D-03 采购合规检查表 ----
    ws3 = wb.create_sheet('D-03 采购合规检查')
    ws3.sheet_properties.tabColor = 'ED7D31'

    ws3.merge_cells('A1:G1')
    ws3['A1'] = '底稿D-03：采购合规检查表'
    ws3['A1'].font = title_font; ws3['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws3.row_dimensions[1].height = 30
    ws3.merge_cells('A2:G2')
    ws3['A2'] = '项目编号：[  ]-采购-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws3['A2'].font = subtitle_font

    # 一、供应商集中度分析
    r = 4; ws3.merge_cells(f'A{r}:G{r}'); ws3[f'A{r}'] = '一、供应商集中度分析'; ws3[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h1 = ['供应商名称', '采购品类', '近3年采购总额', '占同类采购比', '合作起始年份', '是否轮换过', '备注']
    for i, h in enumerate(h1):
        cell = ws3.cell(row=5, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(6, 12):
        for c in range(1, 8):
            cell = ws3.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 二、采购价格对比
    r = 13; ws3.merge_cells(f'A{r}:G{r}'); ws3[f'A{r}'] = '二、采购价格对比分析'; ws3[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h2 = ['品名', '规格', '中标单价', '同期市场价', '同期其他单位采购价', '偏离率', '异常判定']
    for i, h in enumerate(h2):
        cell = ws3.cell(row=14, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(15, 20):
        for c in range(1, 8):
            cell = ws3.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    # 三、供应商关联关系
    r = 21; ws3.merge_cells(f'A{r}:G{r}'); ws3[f'A{r}'] = '三、供应商关联关系排查'; ws3[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h3 = ['供应商', '股东/实际控制人', '是否与被检查单位人员关联', '关联关系描述', '', '', '']
    for i, h in enumerate(h3):
        cell = ws3.cell(row=22, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws3.merge_cells('D22:G22')
    ws3['D22'] = '关联关系描述'
    ws3['D22'].font = header_font; ws3['D22'].fill = header_fill; ws3['D22'].alignment = header_align; ws3['D22'].border = thin_border
    for row_n in range(23, 28):
        for c in [1, 2, 3]:  # write A-C first
            cell = ws3.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws3.merge_cells(f'D{row_n}:G{row_n}')
        cell = ws3.cell(row=row_n, column=4, value='')
        cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        for c in range(5, 8):
            ws3.cell(row=row_n, column=c).border = thin_border

    # 四、关键岗位人员账户
    r = 29; ws3.merge_cells(f'A{r}:G{r}'); ws3[f'A{r}'] = '四、关键岗位人员账户排查'; ws3[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h4 = ['姓名', '职务', '与供应商资金往来', '往来金额', '往来说明', '', '']
    for i, h in enumerate(h4):
        cell = ws3.cell(row=30, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws3.merge_cells('E30:G30'); ws3['E30'] = '往来说明'; ws3['E30'].font = header_font; ws3['E30'].fill = header_fill; ws3['E30'].alignment = header_align; ws3['E30'].border = thin_border
    for row_n in range(31, 36):
        for c in range(1, 5):
            ws3.cell(row=row_n, column=c, value='').font = cell_font
            ws3.cell(row=row_n, column=c).border = thin_border
        ws3.merge_cells(f'E{row_n}:G{row_n}')
        ws3.cell(row=row_n, column=5, value='').font = cell_font
        ws3.cell(row=row_n, column=5).border = thin_border
        for c in range(6, 8):
            ws3.cell(row=row_n, column=c).border = thin_border

    # 五、采购程序
    r = 37; ws3.merge_cells(f'A{r}:G{r}'); ws3[f'A{r}'] = '五、采购程序检查'; ws3[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h5 = ['采购项目', '金额', '采购方式', '招标文件是否合规', '评标过程是否合规', '中标结果是否合理', '']
    for i, h in enumerate(h5):
        cell = ws3.cell(row=38, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws3.merge_cells('F38:G38'); ws3['F38'] = '中标结果是否合理'; ws3['F38'].font = header_font; ws3['F38'].fill = header_fill; ws3['F38'].alignment = header_align; ws3['F38'].border = thin_border
    for row_n in range(39, 44):
        for c in range(1, 6):
            ws3.cell(row=row_n, column=c, value='').font = cell_font
            ws3.cell(row=row_n, column=c).border = thin_border
        ws3.merge_cells(f'F{row_n}:G{row_n}')
        ws3.cell(row=row_n, column=6, value='').font = cell_font
        ws3.cell(row=row_n, column=6).border = thin_border
        ws3.cell(row=row_n, column=7).border = thin_border

    # 六、结论
    r = 45; ws3.merge_cells(f'A{r}:G{r}'); ws3[f'A{r}'] = '六、检查结论'; ws3[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    ws3.merge_cells('A46:G46'); ws3['A46'] = '□ 未发现异常'; ws3['A46'].font = cell_font
    ws3.merge_cells('A47:G48'); ws3['A47'] = '□ 发现以下问题：'; ws3['A47'].font = cell_font

    for i, w in enumerate([14, 12, 14, 12, 14, 12, 20]):
        ws3.column_dimensions[get_column_letter(i+1)].width = w

    # ---- D-04 对外担保检查表 ----
    ws4 = wb.create_sheet('D-04 对外担保检查')
    ws4.sheet_properties.tabColor = 'FFC000'

    ws4.merge_cells('A1:F1')
    ws4['A1'] = '底稿D-04：对外担保检查表'
    ws4['A1'].font = title_font; ws4['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws4.row_dimensions[1].height = 30
    ws4.merge_cells('A2:F2')
    ws4['A2'] = '项目编号：[  ]-担保-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws4['A2'].font = subtitle_font

    sections_4 = [
        ('一、银行函证汇总', ['开户行', '函证发出日', '回函日', '是否披露担保', '担保金额', '被担保方']),
        ('二、企业征信查询结果', ['查询日期', '对外担保笔数', '担保总额', '其中关注/不良类', '', '']),
        ('三、担保合同与审批比对', ['担保合同编号', '合同签订日', '担保金额', '被担保方', '董事会决议', '国资审批']),
        ('三续、反担保与披露', ['合同编号', '反担保措施', '财务披露', '是否关联方', '', '']),
        ('四、被担保方经营状况', ['被担保方', '注册资本', '近1年营收', '资产负债率', '偿债能力评估', '']),
        ('五、担保风险指标', ['指标', '计算公式', '实际值', '国资委红线', '是否超标', '']),
    ]

    current_row = 4
    for title_text, headers in sections_4:
        n_cols = len([h for h in headers if h])
        ws4.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=max(n_cols, 6))
        ws4[f'A{current_row}'] = title_text
        ws4[f'A{current_row}'].font = Font(name='微软雅黑', size=11, bold=True)
        current_row += 1

        for i, h in enumerate(headers):
            if h:
                cell = ws4.cell(row=current_row, column=i+1, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
        current_row += 1

        for row_n in range(current_row, current_row + 4):
            for c in range(1, max(n_cols, 6) + 1):
                cell = ws4.cell(row=row_n, column=c, value='')
                cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        current_row += 4

        # special data for 五
        if title_text.startswith('五、担保风险指标'):
            indicators = ['担保余额/净资产', '已代偿/担保余额']
            for idx, ind in enumerate(indicators):
                for c in range(1, max(n_cols, 6) + 1):
                    ws4.cell(row=current_row - 4 + idx, column=c).font = cell_font
                ws4.cell(row=current_row - 4 + idx, column=1).value = ind

    # 六、检查结论
    ws4.merge_cells(f'A{current_row}:F{current_row}')
    ws4[f'A{current_row}'] = '六、检查结论'; ws4[f'A{current_row}'].font = Font(name='微软雅黑', size=11, bold=True)
    current_row += 1
    ws4.merge_cells(f'A{current_row}:F{current_row}'); ws4[f'A{current_row}'] = '□ 未发现异常'; ws4[f'A{current_row}'].font = cell_font
    current_row += 1
    ws4.merge_cells(f'A{current_row}:F{current_row}'); ws4[f'A{current_row}'] = '□ 发现违规担保___笔，金额___元'; ws4[f'A{current_row}'].font = cell_font

    for i, w in enumerate([16, 14, 14, 14, 16, 16]):
        ws4.column_dimensions[get_column_letter(i+1)].width = w

    # ---- D-05 收费合规检查表 ----
    ws5 = wb.create_sheet('D-05 收费合规检查')
    ws5.sheet_properties.tabColor = '70AD47'

    ws5.merge_cells('A1:G1')
    ws5['A1'] = '底稿D-05：收费合规检查表'
    ws5['A1'].font = title_font; ws5['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws5.row_dimensions[1].height = 30
    ws5.merge_cells('A2:G2')
    ws5['A2'] = '项目编号：[  ]-收费-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws5['A2'].font = subtitle_font

    r = 4
    ws5.merge_cells(f'A{r}:G{r}'); ws5[f'A{r}'] = '一、收费标准核查'; ws5[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h1 = ['收费项目', '批准标准', '实际收取标准', '差异', '是否公示', '', '']
    for i, h in enumerate(h1):
        cell = ws5.cell(row=5, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws5.merge_cells('E5:G5'); ws5['E5'] = '是否公示'; ws5['E5'].font = header_font; ws5['E5'].fill = header_fill; ws5['E5'].alignment = header_align; ws5['E5'].border = thin_border
    for row_n in range(6, 11):
        for c in range(1, 5):
            cell = ws5.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws5.merge_cells(f'E{row_n}:G{row_n}')
        cell = ws5.cell(row=row_n, column=5, value='')
        cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        for c in range(6, 8):
            ws5.cell(row=row_n, column=c).border = thin_border

    r = 12
    ws5.merge_cells(f'A{r}:G{r}'); ws5[f'A{r}'] = '二、收入测算与账面比对'; ws5[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h2 = ['收费项目', '业务量(A)', '单价(B)', '测算收入(C=A×B)', '账面收入(D)', '差异(C-D)', '差异率']
    for i, h in enumerate(h2):
        cell = ws5.cell(row=13, column=i+1, value=h)
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    for row_n in range(14, 18):
        for c in range(1, 8):
            cell = ws5.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border

    r = 19
    ws5.merge_cells(f'A{r}:G{r}'); ws5[f'A{r}'] = '三、非批准收费项目排查'; ws5[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h3 = ['收费名目', '收取期间', '收取对象', '收费标准', '收款账户', '是否入账', '']
    for i, h in enumerate(h3):
        cell = ws5.cell(row=20, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws5.merge_cells('F20:G20'); ws5['F20'] = '是否入账'; ws5['F20'].font = header_font; ws5['F20'].fill = header_fill; ws5['F20'].alignment = header_align; ws5['F20'].border = thin_border
    items_3 = ['赞助费', '培训费', '材料费', '(其他)', '(其他)']
    for idx, item in enumerate(items_3):
        row_n = 21 + idx
        for c in range(1, 6):
            cell = ws5.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws5.merge_cells(f'F{row_n}:G{row_n}')
        cell = ws5.cell(row=row_n, column=6, value='')
        cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws5.cell(row=row_n, column=7).border = thin_border
        ws5.cell(row=row_n, column=1).value = item

    r = 27
    ws5.merge_cells(f'A{r}:G{r}'); ws5[f'A{r}'] = '四、收款账户排查'; ws5[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    h4 = ['账户名', '账号', '开户行', '账户性质', '收款金额', '是否纳入财务核算', '']
    for i, h in enumerate(h4):
        cell = ws5.cell(row=28, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws5.merge_cells('F28:G28'); ws5['F28'] = '是否纳入财务核算'; ws5['F28'].font = header_font; ws5['F28'].fill = header_fill; ws5['F28'].alignment = header_align; ws5['F28'].border = thin_border
    for row_n in range(29, 34):
        for c in range(1, 6):
            cell = ws5.cell(row=row_n, column=c, value='')
            cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws5.merge_cells(f'F{row_n}:G{row_n}')
        cell = ws5.cell(row=row_n, column=6, value='')
        cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        ws5.cell(row=row_n, column=7).border = thin_border

    r = 35
    ws5.merge_cells(f'A{r}:G{r}'); ws5[f'A{r}'] = '五、检查结论'; ws5[f'A{r}'].font = Font(name='微软雅黑', size=11, bold=True)
    ws5.merge_cells('A36:G36'); ws5['A36'] = '□ 未发现异常'; ws5['A36'].font = cell_font
    ws5.merge_cells('A37:G38'); ws5['A37'] = '□ 发现违规收费___项，金额___元'; ws5['A37'].font = cell_font

    for i, w in enumerate([14, 12, 12, 12, 16, 14, 10]):
        ws5.column_dimensions[get_column_letter(i+1)].width = w

    filepath = os.path.join(OUTPUT_DIR, '监督检查-工作底稿模板.xlsx')
    wb.save(filepath)
    print(f'[OK] {filepath}')
    return filepath


# ============================================================
# EXCEL: 财政评审(F) 四套工作底稿
# ============================================================

def create_excel_F():
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    header_font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell_font = Font(name='微软雅黑', size=9)
    cell_align = Alignment(vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    title_font = Font(name='微软雅黑', size=14, bold=True, color='1F4E79')
    subtitle_font = Font(name='微软雅黑', size=9, color='666666')
    section_font = Font(name='微软雅黑', size=11, bold=True)

    # ---- F-01 研发项目真实性核查表 ----
    ws1 = wb.create_sheet('F-01 研发项目真实性')
    ws1.sheet_properties.tabColor = '1F4E79'

    ws1.merge_cells('A1:H1')
    ws1['A1'] = '底稿F-01：研发项目真实性核查表'
    ws1['A1'].font = title_font; ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws1.row_dimensions[1].height = 30
    ws1.merge_cells('A2:H2')
    ws1['A2'] = '项目编号：[  ]-研发-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws1['A2'].font = subtitle_font

    sections_f1 = [
        ('一、研发项目清单', ['项目编号', '项目名称', '立项日期', '预算金额', '实际支出', '状态(在研/结题/中止)', '', '']),
        ('二、项目真实性核查（□是 □否）', ['核查项', '项目1', '项目2', '项目3', '项目4', '项目5', '', '']),
        ('二核查项列表', ['立项文件是否完整', '研发人员是否在册（社保/个税）', '研发场所是否真实存在', '研发设备是否到位',
                          '研发日志是否连续记录', '是否存在对应产品已上市', '研发成果是否可验证']),
        ('三、研发人员穿透核查', ['姓名', '申报岗位', '社保登记岗位', '个税扣缴单位', '实际工作部门', '差异(一致/不一致)', '', '']),
        ('四、研发费用辅助账检查', ['项目', '辅助账金额(A)', '总账科目金额(B)', '其中重复列支', '核实后金额', '核实后占比', '', '']),
        ('四费用类别', ['材料费', '人工费', '折旧费', '差旅费', '其他']),
        ('五、高新认定指标复核', ['指标', '高新认定要求', '申报值', '核实值', '是否达标', '', '', '']),
        ('五指标', ['研发费用占销售收入比例', '研发人员占职工总数比例', '高新技术产品收入占比', '核心知识产权数量']),
        ('六、套取补贴/税收优惠计算', ['项目', '申报依据', '核实后依据', '套取金额', '', '', '', '']),
        ('六项目', ['所得税优惠（15% vs 25%）', '研发加计扣除少缴税', '科技补贴A', '科技补贴B']),
    ]

    current_row = 4
    for title_text, headers in sections_f1:
        n_cols = len([h for h in headers if h])
        ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=max(n_cols, 6))
        ws1[f'A{current_row}'] = title_text
        ws1[f'A{current_row}'].font = section_font
        current_row += 1

        if '核查项列表' in title_text or '费用类别' in title_text or '五指标' in title_text or '六项目' in title_text:
            for idx, item in enumerate(headers):
                ws1.merge_cells(f'A{current_row + idx}:E{current_row + idx}')
                ws1.cell(row=current_row + idx, column=1, value=item).font = cell_font
                for c in range(1, 9):
                    ws1.cell(row=current_row + idx, column=c).border = thin_border
            current_row += len(headers)
            continue

        for i, h in enumerate(headers):
            if h:
                cell = ws1.cell(row=current_row, column=i+1, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
        current_row += 1

        row_count = 6 if '清单' in title_text or '核查' in title_text else 4
        for row_n in range(current_row, current_row + row_count):
            for c in range(1, max(n_cols, 6) + 1):
                cell = ws1.cell(row=row_n, column=c, value='')
                cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        current_row += row_count

    # 七、检查结论
    ws1.merge_cells(f'A{current_row}:H{current_row}')
    ws1[f'A{current_row}'] = '七、检查结论'; ws1[f'A{current_row}'].font = section_font
    current_row += 1
    ws1.merge_cells(f'A{current_row}:H{current_row}'); ws1[f'A{current_row}'] = '□ 研发费用归集合规'; ws1[f'A{current_row}'].font = cell_font
    current_row += 1
    ws1.merge_cells(f'A{current_row}:H{current_row + 1}'); ws1[f'A{current_row}'] = '□ 发现以下问题：'; ws1[f'A{current_row}'].font = cell_font

    for i, w in enumerate([14, 14, 14, 14, 14, 12, 12, 20]):
        ws1.column_dimensions[get_column_letter(i+1)].width = w

    # ---- F-02 专项补贴申报核查表 ----
    ws2 = wb.create_sheet('F-02 专项补贴申报核查')
    ws2.sheet_properties.tabColor = 'ED7D31'

    ws2.merge_cells('A1:G1')
    ws2['A1'] = '底稿F-02：专项补贴申报核查表'
    ws2['A1'].font = title_font; ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws2.row_dimensions[1].height = 30
    ws2.merge_cells('A2:G2')
    ws2['A2'] = '项目编号：[  ]-补贴-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws2['A2'].font = subtitle_font

    sections_f2 = [
        ('一、补贴申报清单', ['补贴名称', '申报年度', '申报金额', '批复金额', '到账金额', '申报依据文件', '']),
        ('二、核心申报数据核查', ['补贴项目', '关键申报数据(A)', '数据来源', '系统原始数据(B)', '差异(A-B)', '偏差率', '']),
        ('三、数据源追溯', ['申报数据项', '申报表中数值', '最原始数据来源', '是否从源系统直接导出', '中间有无加工', '', '']),
        ('四、申报范围合规性', ['核查项', '申报范围', '核实后范围', '差异', '', '', '']),
        ('四核查项', ['营运车辆数', '非营运车辆是否混入', '已报废车辆是否剔除', '停驶期间里程是否剔除']),
        ('五、补贴资金使用跟踪', ['补贴资金用途', '申报承诺', '实际用途', '是否偏离', '', '', '']),
    ]

    current_row = 4
    for title_text, headers in sections_f2:
        n_cols = len([h for h in headers if h])
        ws2.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=max(n_cols, 6))
        ws2[f'A{current_row}'] = title_text; ws2[f'A{current_row}'].font = section_font
        current_row += 1

        if '核查项' in title_text:
            for idx, item in enumerate(headers):
                ws2.merge_cells(f'A{current_row + idx}:D{current_row + idx}')
                ws2.cell(row=current_row + idx, column=1, value=item).font = cell_font
                for c in range(1, 8):
                    ws2.cell(row=current_row + idx, column=c).border = thin_border
            current_row += len(headers)
            continue

        for i, h in enumerate(headers):
            if h:
                cell = ws2.cell(row=current_row, column=i+1, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
        current_row += 1

        for row_n in range(current_row, current_row + 4):
            for c in range(1, max(n_cols, 6) + 1):
                cell = ws2.cell(row=row_n, column=c, value='')
                cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        current_row += 4

    # 六、结论
    ws2.merge_cells(f'A{current_row}:G{current_row}')
    ws2[f'A{current_row}'] = '六、检查结论'; ws2[f'A{current_row}'].font = section_font
    current_row += 1
    ws2.merge_cells(f'A{current_row}:G{current_row}'); ws2[f'A{current_row}'] = '□ 申报数据真实合规'; ws2[f'A{current_row}'].font = cell_font
    current_row += 1
    ws2.merge_cells(f'A{current_row}:G{current_row + 1}'); ws2[f'A{current_row}'] = '□ 发现以下问题：'; ws2[f'A{current_row}'].font = cell_font

    for i, w in enumerate([16, 14, 16, 14, 14, 16, 14]):
        ws2.column_dimensions[get_column_letter(i+1)].width = w

    # ---- F-03 工资总额预算评审表 ----
    ws3 = wb.create_sheet('F-03 工资总额预算评审')
    ws3.sheet_properties.tabColor = 'C00000'

    ws3.merge_cells('A1:G1')
    ws3['A1'] = '底稿F-03：工资总额预算评审表'
    ws3['A1'].font = title_font; ws3['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws3.row_dimensions[1].height = 30
    ws3.merge_cells('A2:G2')
    ws3['A2'] = '项目编号：[  ]-工资-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws3['A2'].font = subtitle_font

    sections_f3 = [
        ('一、工资总额预算与执行比对', ['年度', '年初预算批复(A)', '年中调整(B)', '调整后预算(C=A+B)', '实际执行(D)', '超支额(D-C)', '超支率']),
        ('二、工资总额内/外发放结构分析', ['发放项目', '是否在工资总额内', '发放金额', '审批程序', '列支科目', '', '']),
        ('二项目', ['基本工资', '绩效工资', '专项奖励', '考核兑现', '节日慰问', '(其他)']),
        ('三、科目列支合规性检查', ['科目', '含奖励/津贴/补贴/兑现摘要的金额', '是否属于薪酬性质', '是否纳入工资总额', '', '', '']),
        ('三科目', ['管理费用', '营业外支出', '其他应付款', '在建工程']),
        ('四、关键岗位收入汇总', ['姓名', '职务', '工资总额内', '管理费列支', '营业外列支', '合计', '个税申报']),
    ]

    current_row = 4
    for title_text, headers in sections_f3:
        n_cols = len([h for h in headers if h])
        ws3.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=max(n_cols, 6))
        ws3[f'A{current_row}'] = title_text; ws3[f'A{current_row}'].font = section_font
        current_row += 1

        if '项目' in title_text or '科目' in title_text:
            for idx, item in enumerate(headers):
                ws3.merge_cells(f'A{current_row + idx}:C{current_row + idx}')
                ws3.cell(row=current_row + idx, column=1, value=item).font = cell_font
                for c in range(1, 8):
                    ws3.cell(row=current_row + idx, column=c).border = thin_border
            current_row += len(headers)
            continue

        for i, h in enumerate(headers):
            if h:
                cell = ws3.cell(row=current_row, column=i+1, value=h)
                cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
        current_row += 1

        row_count = 4 if '比对' in title_text else 4
        for row_n in range(current_row, current_row + row_count):
            for c in range(1, max(n_cols, 6) + 1):
                cell = ws3.cell(row=row_n, column=c, value='')
                cell.font = cell_font; cell.alignment = cell_align; cell.border = thin_border
        current_row += row_count

    # 五、结论
    ws3.merge_cells(f'A{current_row}:G{current_row}')
    ws3[f'A{current_row}'] = '五、检查结论'; ws3[f'A{current_row}'].font = section_font
    current_row += 1
    ws3.merge_cells(f'A{current_row}:G{current_row}'); ws3[f'A{current_row}'] = '□ 工资总额管理合规'; ws3[f'A{current_row}'].font = cell_font
    current_row += 1
    ws3.merge_cells(f'A{current_row}:G{current_row + 1}'); ws3[f'A{current_row}'] = '□ 发现以下问题：'; ws3[f'A{current_row}'].font = cell_font

    for i, w in enumerate([14, 16, 14, 14, 14, 14, 10]):
        ws3.column_dimensions[get_column_letter(i+1)].width = w

    # ---- F-04 税收优惠与补贴联动检查表 ----
    ws4 = wb.create_sheet('F-04 税收优惠与补贴联动检查')
    ws4.sheet_properties.tabColor = '70AD47'

    ws4.merge_cells('A1:F1')
    ws4['A1'] = '底稿F-04：税收优惠与补贴联动检查表'
    ws4['A1'].font = title_font; ws4['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws4.row_dimensions[1].height = 30
    ws4.merge_cells('A2:F2')
    ws4['A2'] = '项目编号：[  ]-联动-001    被检查单位：_______________    编制人/日期：___ / 20__.__.__    复核人/日期：___ / 20__.__.__'
    ws4['A2'].font = subtitle_font

    r = 4
    ws4.merge_cells(f'A{r}:F{r}'); ws4[f'A{r}'] = '一、同一套材料多口申报检查'; ws4[f'A{r}'].font = section_font
    h1 = ['申报事项', '申报部门', '使用关键数据', '与其他申报是否一致', '差异说明', '']
    for i, h in enumerate(h1):
        cell = ws4.cell(row=5, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws4.merge_cells('E5:F5'); ws4['E5'] = '差异说明'; ws4['E5'].font = header_font; ws4['E5'].fill = header_fill; ws4['E5'].alignment = header_align; ws4['E5'].border = thin_border
    items_1 = ['高新认定(科技厅)', '研发加计扣除(税务局)', '科技补贴A(科技厅)', '科技补贴B(经信委)']
    for idx, item in enumerate(items_1):
        row_n = 6 + idx
        ws4.merge_cells(f'E{row_n}:F{row_n}')
        ws4.cell(row=row_n, column=1, value=item).font = cell_font
        for c in range(1, 7):
            ws4.cell(row=row_n, column=c).border = thin_border

    r = 11
    ws4.merge_cells(f'A{r}:F{r}'); ws4[f'A{r}'] = '二、财政资金"一鱼多吃"排查'; ws4[f'A{r}'].font = section_font
    h2 = ['同一笔支出', '加计扣除中申报', '补贴中申报', '常规成本中列支', '重复使用次数', '']
    for i, h in enumerate(h2):
        cell = ws4.cell(row=12, column=i+1, value=h if h else '')
        cell.font = header_font; cell.fill = header_fill; cell.alignment = header_align; cell.border = thin_border
    ws4.merge_cells('E12:F12'); ws4['E12'] = '重复使用次数'; ws4['E12'].font = header_font; ws4['E12'].fill = header_fill; ws4['E12'].alignment = header_align; ws4['E12'].border = thin_border
    items_2 = ['材料费___', '人工费___', '折旧费___', '检测费___', '(其他)___']
    for idx, item in enumerate(items_2):
        row_n = 13 + idx
        ws4.merge_cells(f'E{row_n}:F{row_n}')
        ws4.cell(row=row_n, column=1, value=item).font = cell_font
        for c in range(1, 7):
            ws4.cell(row=row_n, column=c).border = thin_border

    r = 19
    ws4.merge_cells(f'A{r}:F{r}'); ws4[f'A{r}'] = '三、检查结论'; ws4[f'A{r}'].font = section_font
    ws4.merge_cells('A20:F20'); ws4['A20'] = '□ 未发现重复申报'; ws4['A20'].font = cell_font
    ws4.merge_cells('A21:F22'); ws4['A21'] = '□ 发现重复申报问题：'; ws4['A21'].font = cell_font

    for i, w in enumerate([16, 14, 16, 16, 16, 16]):
        ws4.column_dimensions[get_column_letter(i+1)].width = w

    filepath = os.path.join(OUTPUT_DIR, '财政评审-工作底稿模板.xlsx')
    wb.save(filepath)
    print(f'[OK] {filepath}')
    return filepath


# ============================================================
# 主流程
# ============================================================

if __name__ == '__main__':
    print("开始生成文档...")
    create_word_D()
    create_word_F()
    create_excel_D()
    create_excel_F()
    print("\n全部完成！生成4个文件：")
    print("  Word: 监督检查-案例分析与操作卡.docx")
    print("  Word: 财政评审-案例分析与操作卡.docx")
    print("  Excel: 监督检查-工作底稿模板.xlsx (5个Sheet)")
    print("  Excel: 财政评审-工作底稿模板.xlsx (4个Sheet)")
