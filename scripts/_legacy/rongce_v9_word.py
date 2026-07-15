# -*- coding: utf-8 -*-
"""
四川融策宣传册 v9 —— 可编辑 Word 原生版
每个页面用独立节(section)，文字/表格/颜色均可直接在 Word 中编辑。
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
DOCX = OUT / "四川融策宣传册_v9_可编辑Word版.docx"

# Colors
NW = RGBColor(0x1A, 0x36, 0x5D)
TL = RGBColor(0x3A, 0x7B, 0x8A)
GD = RGBColor(0xD4, 0xA5, 0x74)
MU = RGBColor(0x71, 0x80, 0x96)
DK = RGBColor(0x2D, 0x37, 0x48)
WH = RGBColor(0xFF, 0xFF, 0xFF)
GRY = RGBColor(0x99, 0x99, 0x99)

doc = Document()

# ==================== HELPERS ====================
def p(doc_or_cell, text, size=11, bold=False, color=DK, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.2):
    """Add a formatted paragraph. Returns the paragraph."""
    para = doc_or_cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    run = para.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return para

def heading(text, en="", size=28):
    """Page title with English subtitle"""
    p(doc, text, size=size, bold=True, color=NW, before=0, after=0)
    if en:
        p(doc, en, size=12, color=GD, before=0, after=0)
    p(doc, "\u2500" * 40, size=6, color=GD, before=2, after=8)

def bullet(text, size=11, color=DK, indent=0):
    """Bullet point"""
    p(doc, ("    " * indent) + "\u2022 " + text, size=size, color=color, before=1, after=1)

def spacer(h=6):
    p(doc, "", size=2, before=h, after=0)

def section_break(margins=(2.5, 2.0, 2.8, 2.8)):
    """Start new page/section"""
    sec = doc.add_section()
    sec.top_margin = Cm(margins[0])
    sec.bottom_margin = Cm(margins[1])
    sec.left_margin = Cm(margins[2])
    sec.right_margin = Cm(margins[3])
    return sec

def make_table(rows, cols, col_widths=None):
    """Create a clean table"""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

def header_row(tbl, texts, bg="1A365D"):
    """Style the header row of a table"""
    row = tbl.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'))
        p(cell, text, size=11, bold=True, color=WH, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)

def data_row(tbl, row_idx, texts, bg=None):
    """Populate a data row"""
    row = tbl.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        if bg:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'))
        p(cell, text, size=10, color=DK, before=1, after=1)

# ==================== PAGE 1: COVER ====================
sec = doc.sections[0]
sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)

spacer(40)
p(doc, "SICHUAN  RONGCE", size=20, bold=True, color=GD, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
spacer(60)
p(doc, "\u56db\u5ddd\u878d\u7b56\u4f1a\u8ba1\u5e08\u4e8b\u52a1\u6240", size=20, color=MU, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
p(doc, "\u8c0b\u4e13\u4e1a\u4e4b\u7b56  \u878d\u54c1\u8d28\u4e4b\u7cbe", size=32, bold=True, color=NW, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=8)
p(doc, "\u2500" * 50, size=6, color=GD, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
p(doc, "\u653f\u5e9c\u5ba1\u8ba1\u4e0e\u5de5\u7a0b\u54a8\u8be2\u7efc\u5408\u670d\u52a1\u673a\u6784", size=16, color=MU, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
spacer(80)
p(doc, "\u5ba1\u8ba1 \u00b7 \u7ee9\u6548 \u00b7 \u8d22\u653f\u76d1\u7763 \u00b7 \u5de5\u7a0b\u54a8\u8be2 \u00b7 \u6570\u5b57\u5316\u5206\u6790", size=12, color=GRY, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

# ==================== PAGE 2: WHY RONGCE ====================
section_break()
heading("\u4e3a\u4ec0\u4e48\u662f\u878d\u7b56\uff1f", "WHY RONGCE")
spacer(4)
p(doc, "\u8d22\u653f\u8d44\u91d1\u7684\u7ba1\u7406\uff0c\u6838\u5fc3\u662f\u4e24\u4e2a\u95ee\u9898\uff1a", size=16, color=MU, before=4, after=2)
p(doc, "\u201c\u8d44\u91d1\u5230\u5e95\u53bb\u54ea\u4e86\uff1f\u201d \u548c \u201c\u8fd9\u94b1\u82b1\u5f97\u503c\u4e0d\u503c\uff1f\u201d", size=18, bold=True, color=NW, before=0, after=4)
p(doc, "\u878d\u7b56\u505a\u7684\u4e8b\u60c5\uff0c\u5c31\u662f\u7528\u4e13\u4e1a\u7684\u65b9\u6cd5\u548c\u6570\u636e\u5de5\u5177\uff0c\u5e2e\u60a8\u628a\u8fd9\u4e24\u4e2a\u95ee\u9898\u641e\u6e05\u695a\u3002", size=13, color=MU, before=4, after=8)
spacer(2)

# Three differentiators
for num, head, desc in [
    ("\u2460", "\u4e0d\u53ea\u67e5\u8d26", "\u6211\u4eec\u770b\u7684\u662f\u5408\u540c\u3001\u9879\u76ee\u3001\u8d44\u4ea7\u3001\u5185\u63a7\u5236\u5ea6\u548c\u51b3\u7b56\u7a0b\u5e8f\u3002\u8d26\u8868\u53ea\u662f\u5165\u53e3\u3002"),
    ("\u2461", "\u4e0d\u53ea\u627e\u95ee\u9898", "\u6211\u4eec\u66f4\u5173\u5fc3\u95ee\u9898\u80fd\u4e0d\u80fd\u6539\u3001\u600e\u4e48\u6539\u3001\u6539\u5f97\u600e\u4e48\u6837\u3002\u6574\u6539\u843d\u5730\u624d\u662f\u76ee\u6807\u3002"),
    ("\u2462", "\u4e0d\u53ea\u9760\u7ecf\u9a8c", "\u6570\u636e\u5206\u6790\u548c\u73b0\u573a\u6838\u67e5\u7ed3\u5408\uff0c\u7528\u6570\u636e\u6269\u5927\u8986\u76d6\u9762\uff0c\u7528\u73b0\u573a\u6838\u5b9e\u5173\u952e\u70b9\u3002"),
]:
    p(doc, f"{num}  {head}", size=16, bold=True, color=NW, before=6, after=2)
    p(doc, f"    {desc}", size=12, color=MU, before=0, after=4)
spacer(4)
p(doc, "\u2500" * 60, size=4, color=GD, before=4, after=4)

# Company info
for item in [
    "\u59cb\u4e8e 2000 \u5e74\uff0c\u56db\u5ddd\u7701\u5185\u8f83\u65e9\u6210\u7acb\u7684\u4f1a\u8ba1\u5e08\u4e8b\u52a1\u6240\u4e4b\u4e00",
    "\u957f\u671f\u670d\u52a1\u8d22\u653f\u3001\u5ba1\u8ba1\u3001\u6559\u80b2\u3001\u6c11\u653f\u3001\u4ea4\u901a\u3001\u56fd\u8d44\u3001\u533b\u4fdd\u7b49\u9886\u57df",
    "\u5f62\u6210\u5ba1\u8ba1 + \u7ee9\u6548 + \u8d22\u653f\u76d1\u7763 + \u5de5\u7a0b\u54a8\u8be2\u534f\u540c\u53d1\u5c55\u7684\u4e1a\u52a1\u683c\u5c40",
    "\u8986\u76d6\u56db\u5ddd\u3001\u897f\u85cf\u3001\u8d35\u5dde\u4e09\u7701"
]:
    bullet(item, size=11, color=DK, indent=0)
spacer(8)
p(doc, "\u516c\u5f00\u516c\u6b63  \u7528\u5fc3\u670d\u52a1  \u8bda\u4fe1\u4e3a\u672c  \u670d\u52a1\u81f3\u4e0a  \u8ffd\u6c42\u5353\u8d8a", size=10, color=MU, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

# ==================== PAGE 3: METHODOLOGY ====================
section_break()
heading("\u6211\u4eec\u600e\u4e48\u505a\uff1f", "OUR METHODOLOGY")
spacer(2)

# Three lines
for num, title, items in [
    ("01", "\u8d44\u91d1\u7ebf", ["\u94b1\u4ece\u54ea\u6765\uff1f\u5230\u54ea\u53bb\uff1f\u82b1\u5f97\u503c\u4e0d\u503c\uff1f", "\u8d44\u91d1\u6d41\u3001\u652f\u4ed8\u6d41\u3001\u7968\u636e\u6d41\u4e09\u6d41\u4ea4\u53c9\u6838\u9a8c"]),
    ("02", "\u9879\u76ee\u7ebf", ["\u7acb\u9879\u3001\u62db\u6807\u3001\u5408\u540c\u3001\u65bd\u5de5\u3001\u9a8c\u6536\u3001\u7ed3\u7b97", "\u5168\u751f\u547d\u5468\u671f\u7a7f\u900f\u6838\u67e5\uff0c\u4e0d\u7559\u6b7b\u89d2"]),
    ("03", "\u8d23\u4efb\u7ebf", ["\u51b3\u7b56\u7a0b\u5e8f\u3001\u5c97\u4f4d\u804c\u8d23\u3001\u5185\u63a7\u5236\u5ea6\u3001\u6574\u6539\u843d\u5b9e", "\u628a\u95ee\u9898\u8ffd\u6eaf\u5230\u4eba\uff0c\u8ba9\u7ed3\u8bba\u6709\u4f9d\u636e\u3001\u8d23\u4efb\u53ef\u8ffd\u6eaf"]),
]:
    p(doc, f"[{num}]  {title}", size=16, bold=True, color=NW, before=6, after=2)
    for item in items:
        bullet(item, size=11, color=DK, indent=1)
spacer(4)

# Five steps
p(doc, "\u5de5\u4f5c\u6d41\u7a0b\uff1a", size=14, bold=True, color=NW, before=8, after=4)
p(doc, "\u2460 \u5236\u5ea6\u5ba1\u67e5  \u2192  \u2461 \u6570\u636e\u6838\u9a8c  \u2192  \u2462 \u73b0\u573a\u6838\u67e5  \u2192  \u2463 \u8bc1\u636e\u95ed\u73af  \u2192  \u2464 \u6574\u6539\u5efa\u8bae", size=14, bold=True, color=TL, before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
spacer(2)
p(doc, "\u2500" * 60, size=4, color=GD, before=4, after=4)
p(doc, "\u6211\u4eec\u7684\u76ee\u6807\u4e0d\u662f\u51fa\u5177\u4e00\u4efd\u62a5\u544a\uff0c\u800c\u662f\u5e2e\u52a9\u59d4\u6258\u65b9\u770b\u6e05\u95ee\u9898\u3001\u5398\u6e05\u8d23\u4efb\u3001\u627e\u5230\u6539\u8fdb\u8def\u5f84\u3002", size=13, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=0)

# ==================== PAGE 4: SERVICES OVERVIEW ====================
section_break()
heading("\u670d\u52a1\u4f53\u7cfb", "SERVICE SYSTEM")
spacer(2)
p(doc, "\u4ece\u5355\u9879\u5ba1\u8ba1\u5230\u5168\u8fc7\u7a0b\u54a8\u8be2\uff0c\u4ece\u7ed3\u679c\u8bc4\u4ef7\u5230\u5236\u5ea6\u4f18\u5316\uff0c\u4ece\u4eba\u5de5\u6838\u67e5\u5230\u6570\u636e\u5316\u8bc6\u522b\u3002", size=14, color=MU, before=4, after=8)

tbl = make_table(7, 2, [4.0, 12.0])
header_row(tbl, ["\u4e1a\u52a1\u7ebf", "\u670d\u52a1\u5185\u5bb9"])
for i, (name, items) in enumerate([
    ("\u653f\u5e9c\u5ba1\u8ba1", "\u7ecf\u6d4e\u8d23\u4efb\u5ba1\u8ba1 \u00b7 \u4e13\u9879\u8d44\u91d1\u5ba1\u8ba1 \u00b7 \u8d22\u653f\u76d1\u7763\u68c0\u67e5 \u00b7 \u5de5\u7a0b\u51b3\u7b97\u8d22\u52a1\u5ba1\u8ba1"),
    ("\u9884\u7b97\u7ee9\u6548\u7ba1\u7406", "\u4e8b\u524d\u8bc4\u4f30 \u00b7 \u76ee\u6807\u5ba1\u6838 \u00b7 \u8fd0\u884c\u76d1\u63a7 \u00b7 \u91cd\u70b9\u8bc4\u4ef7 \u00b7 \u7ed3\u679c\u5e94\u7528"),
    ("\u5de5\u7a0b\u54a8\u8be2", "\u9884\u7b97\u7f16\u5236 \u00b7 \u8d22\u653f\u8bc4\u5ba1 \u00b7 \u7ed3\u7b97\u5ba1\u6838 \u00b7 \u5168\u8fc7\u7a0b\u5de5\u7a0b\u54a8\u8be2"),
    ("\u91c7\u8d2d\u5ba1\u8ba1", "\u91c7\u8d2d\u7a0b\u5e8f\u5408\u89c4 \u00b7 \u56f4\u6807\u4e32\u6807\u7ebf\u7d22 \u00b7 \u5408\u540c\u5c65\u7ea6\u6838\u67e5"),
    ("\u7ba1\u7406\u54a8\u8be2", "\u5185\u63a7\u5efa\u8bbe \u00b7 \u8d44\u4ea7\u7ba1\u7406 \u00b7 \u6574\u6539\u63d0\u5347 \u00b7 \u6d41\u7a0b\u4f18\u5316"),
    ("\u6570\u5b57\u5316\u5ba1\u8ba1", "\u6570\u636e\u6807\u51c6 \u00b7 \u89c4\u5219\u6a21\u578b \u00b7 \u7a7f\u900f\u6838\u67e5 \u00b7 \u62a5\u544a\u590d\u6838"),
], 1):
    data_row(tbl, i, [name, items], bg="F5F2EC" if i%2==0 else None)
spacer(6)
p(doc, "\u4e94\u6761\u4e1a\u52a1\u7ebf\u5f62\u6210\u534f\u540c\u6548\u5e94\uff1a\u5ba1\u8ba1\u53d1\u73b0\u95ee\u9898\uff0c\u7ee9\u6548\u8bc4\u4ef7\u63d0\u4f9b\u6539\u8fdb\u65b9\u5411\uff0c\u5de5\u7a0b\u54a8\u8be2\u843d\u5b9e\u9879\u76ee\u6267\u884c\uff0c\u7ba1\u7406\u54a8\u8be2\u56fa\u5316\u5236\u5ea6\u6210\u679c\u3002", size=12, color=MU, before=4, after=0)

# ==================== PAGE 5: GOVERNMENT AUDIT ====================
section_break()
heading("\u653f\u5e9c\u5ba1\u8ba1", "GOVERNMENT AUDIT")
p(doc, "\u4ece\u201c\u6709\u6ca1\u6709\u201d\u5230\u201c\u5bf9\u4e0d\u5bf9\u201d\u518d\u5230\u201c\u503c\u4e0d\u503c\u201d\u2014\u2014\u628a\u8d26\u8868\u3001\u5408\u540c\u3001\u9879\u76ee\u3001\u8d44\u91d1\u3001\u8d44\u4ea7\u3001\u8d23\u4efb\u8d2f\u901a\u6838\u67e5\u3002", size=14, color=MU, before=4, after=8)

# Three layers
for layer, label, desc in [
    ("\u201c\u6709\u6ca1\u6709\u201d", "\u5408\u89c4\u6027\u5ba1\u67e5", "\u5236\u5ea6\u662f\u5426\u5b8c\u5584 \u00b7 \u7a0b\u5e8f\u662f\u5426\u5408\u89c4 \u00b7 \u8bb0\u5f55\u662f\u5426\u5b8c\u6574"),
    ("\u201c\u5bf9\u4e0d\u5bf9\u201d", "\u51c6\u786e\u6027\u6838\u67e5", "\u6570\u636e\u662f\u5426\u51c6\u786e \u00b7 \u91d1\u989d\u662f\u5426\u4e00\u81f4 \u00b7 \u5f52\u5c5e\u662f\u5426\u6e05\u6670"),
    ("\u201c\u503c\u4e0d\u503c\u201d", "\u7ee9\u6548\u6027\u8bc4\u4ef7", "\u6295\u5165\u662f\u5426\u5408\u7406 \u00b7 \u4ea7\u51fa\u662f\u5426\u8fbe\u6807 \u00b7 \u6548\u679c\u662f\u5426\u6301\u7eed"),
]:
    p(doc, f"{layer}  {label}", size=15, bold=True, color=NW, before=6, after=2)
    p(doc, f"    {desc}", size=12, color=MU, before=0, after=4)
spacer(4)

# Areas table
p(doc, "\u670d\u52a1\u9886\u57df", size=14, bold=True, color=NW, before=8, after=4)
tbl2 = make_table(5, 2, [4.0, 12.0])
header_row(tbl2, ["\u9886\u57df", "\u91cd\u70b9\u5185\u5bb9"])
for i, (name, items) in enumerate([
    ("\u7ecf\u6d4e\u8d23\u4efb\u5ba1\u8ba1", "\u91cd\u5927\u51b3\u7b56\u4e0e\u8d44\u91d1\u8d44\u4ea7\u5b89\u5168 \u00b7 \u9879\u76ee\u5efa\u8bbe\u4e0e\u5185\u63a7\u98ce\u9669\u8bc6\u522b \u00b7 \u5ec9\u653f\u98ce\u9669\u6392\u67e5\u4e0e\u8d23\u4efb\u8ffd\u6eaf"),
    ("\u4e13\u9879\u8d44\u91d1\u5ba1\u8ba1", "\u7533\u62a5\u5206\u914d\u3001\u62e8\u4ed8\u4f7f\u7528\u5168\u94fe\u6761\u6838\u67e5 \u00b7 \u7ee9\u6548\u8bc4\u4ef7\u4e0e\u7ed3\u4f59\u6c89\u6dc0\u5206\u6790"),
    ("\u8d22\u653f\u76d1\u7763\u68c0\u67e5", "\u9884\u7b97\u6267\u884c\u4e0e\u8d22\u7ecf\u7eaa\u5f8b\u5408\u89c4 \u00b7 \u653f\u5e9c\u91c7\u8d2d\u4e0e\u4f1a\u8ba1\u4fe1\u606f\u8d28\u91cf\u68c0\u67e5"),
    ("\u5de5\u7a0b\u51b3\u7b97\u8d22\u52a1\u5ba1\u8ba1", "\u5efa\u8bbe\u6210\u672c\u5f52\u96c6\u4e0e\u8d44\u91d1\u6765\u6e90\u6838\u5b9e \u00b7 \u8d44\u4ea7\u4ea4\u4ed8\u4e0e\u5c3e\u5de5\u5c3e\u6b3e\u7ba1\u7406"),
], 1):
    data_row(tbl2, i, [name, items], bg="F5F2EC" if i%2==0 else None)

# ==================== PAGE 6: BUDGET PERFORMANCE ====================
section_break()
heading("\u9884\u7b97\u7ee9\u6548\u7ba1\u7406", "BUDGET PERFORMANCE")
p(doc, "\u60a8\u5173\u5fc3\u7684\u4e0d\u53ea\u662f\u201c\u82b1\u4e86\u591a\u5c11\u94b1\u201d\uff0c\u66f4\u662f\u201c\u6548\u679c\u600e\u4e48\u6837\u201d\u2014\u2014\u8ba9\u8d22\u653f\u8d44\u91d1\u4ece\u201c\u82b1\u4e86\u6ca1\u6709\u201d\u8d70\u5411\u201c\u82b1\u5f97\u503c\u4e0d\u503c\u201d\u3002", size=14, color=MU, before=4, after=10)

# Five-step cycle (simple list with arrows)
p(doc, "\u5168\u5468\u671f\u670d\u52a1\u8986\u76d6\uff1a", size=14, bold=True, color=NW, before=0, after=6)
p(doc, "\u4e8b\u524d\u8bc4\u4f30  \u2192  \u76ee\u6807\u5ba1\u6838  \u2192  \u8fd0\u884c\u76d1\u63a7  \u2192  \u91cd\u70b9\u8bc4\u4ef7  \u2192  \u7ed3\u679c\u5e94\u7528", size=14, bold=True, color=TL, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10)

tbl3 = make_table(6, 2, [3.5, 12.5])
header_row(tbl3, ["\u73af\u8282", "\u670d\u52a1\u5185\u5bb9"])
data = [
    ("\u4e8b\u524d\u7ee9\u6548\u8bc4\u4f30", "\u5fc5\u8981\u6027\u3001\u53ef\u884c\u6027\u4e0e\u8d22\u653f\u627f\u53d7\u80fd\u529b\u5206\u6790 \u00b7 \u9884\u671f\u7ee9\u6548\u4e0e\u6295\u5165\u6210\u672c\u7efc\u5408\u8bc4\u4f30"),
    ("\u7ee9\u6548\u76ee\u6807\u5ba1\u6838", "\u76ee\u6807\u5b8c\u6574\u6027\u4e0e\u6307\u6807\u53ef\u8861\u91cf\u6027\u5ba1\u67e5 \u00b7 \u9884\u7b97\u5339\u914d\u6027\u4e0e\u7ee9\u6548\u8d23\u4efb\u4e66\u5ba1\u6838"),
    ("\u7ee9\u6548\u8fd0\u884c\u76d1\u63a7", "\u6267\u884c\u8fdb\u5ea6\u4e0e\u8d44\u91d1\u652f\u4ed8\u8ffd\u8e2a \u00b7 \u4ea7\u51fa\u504f\u5dee\u5206\u6790\u4e0e\u98ce\u9669\u9884\u8b66"),
    ("\u91cd\u70b9\u7ee9\u6548\u8bc4\u4ef7", "\u653f\u7b56\u8bc4\u4ef7\u3001\u90e8\u95e8\u6574\u4f53\u8bc4\u4ef7 \u00b7 \u9879\u76ee\u652f\u51fa\u4e0e\u4e13\u9879\u8d44\u91d1\u8bc4\u4ef7"),
    ("\u7ed3\u679c\u5e94\u7528", "\u6574\u6539\u6e05\u5355\u4e0e\u9884\u7b97\u6302\u94a9 \u00b7 \u7ba1\u7406\u5236\u5ea6\u4f18\u5316\u5efa\u8bae"),
]
for i, (name, items) in enumerate(data, 1):
    data_row(tbl3, i, [name, items], bg="F5F2EC" if i%2==0 else None)
spacer(6)
p(doc, "\u6bcf\u4e2a\u73af\u8282\u4ea4\u4ed8\uff1a\u6838\u67e5\u6e05\u5355 + \u6570\u636e\u5e95\u7a3f + \u5206\u6790\u62a5\u544a + \u6574\u6539\u5efa\u8bae", size=12, color=MU, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=0)

# ==================== PAGE 7: ENGINEERING ====================
section_break()
heading("\u5de5\u7a0b\u54a8\u8be2\u4e0e\u8d22\u653f\u8bc4\u5ba1", "ENGINEERING CONSULTING")
p(doc, "\u4ece\u6982\u7b97\u5230\u7ed3\u7b97\uff0c\u5e2e\u60a8\u628a\u597d\u6bcf\u4e00\u9053\u5173\u2014\u2014\u5de5\u7a0b\u9020\u4ef7\u3001\u5408\u540c\u5c65\u7ea6\u3001\u8d44\u91d1\u652f\u4ed8\u548c\u9879\u76ee\u7ee9\u6548\u534f\u540c\u7ba1\u7406\u3002", size=14, color=MU, before=4, after=8)

# Four phases horizontal
p(doc, "\u2630  \u9884\u7b97\u7f16\u5236\u4e0e\u8d22\u653f\u8bc4\u5ba1  \u2192   \u2630  \u6e05\u5355\u53ca\u62db\u6807\u63a7\u5236\u4ef7  \u2192   \u2630  \u7ed3\u7b97\u5ba1\u6838  \u2192   \u2630  \u5168\u8fc7\u7a0b\u5de5\u7a0b\u54a8\u8be2", size=13, bold=True, color=NW, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)

tbl4 = make_table(5, 2, [4.5, 11.5])
header_row(tbl4, ["\u670d\u52a1", "\u91cd\u70b9\u5185\u5bb9"])
for i, (name, items) in enumerate([
    ("\u9884\u7b97\u7f16\u5236\u4e0e\u8d22\u653f\u8bc4\u5ba1", "\u5de5\u7a0b\u91cf\u3001\u5b9a\u989d\u5957\u7528\u6838\u9a8c \u00b7 \u6750\u6599\u4ef7\u683c\u3001\u63aa\u65bd\u8d39\u3001\u53d6\u8d39\u6807\u51c6"),
    ("\u6e05\u5355\u53ca\u62db\u6807\u63a7\u5236\u4ef7", "\u63d0\u5347\u62db\u6807\u6587\u4ef6\u548c\u63a7\u5236\u4ef7\u7f16\u5236\u8d28\u91cf"),
    ("\u7ed3\u7b97\u5ba1\u6838", "\u5408\u540c\u6761\u6b3e\u3001\u53d8\u66f4\u7b7e\u8bc1\u6838\u67e5 \u00b7 \u9690\u853d\u5de5\u7a0b\u3001\u73b0\u573a\u5de5\u7a0b\u91cf\u6838\u9a8c"),
    ("\u5168\u8fc7\u7a0b\u5de5\u7a0b\u54a8\u8be2", "\u9879\u76ee\u524d\u671f\u3001\u62db\u91c7\u3001\u5b9e\u65bd \u00b7 \u9a8c\u6536\u3001\u7ed3\u7b97\u3001\u7ee9\u6548\u8bc4\u4ef7\u534f\u540c"),
], 1):
    data_row(tbl4, i, [name, items], bg="F5F2EC" if i%2==0 else None)

# ==================== PAGE 8: DIGITAL ====================
section_break()
heading("\u6570\u5b57\u5316\u5ba1\u8ba1\u80fd\u529b", "DIGITAL AUDIT CAPABILITIES")
p(doc, "\u7528\u6570\u636e\u6269\u5927\u8986\u76d6\u9762\u3001\u63d0\u9ad8\u53d1\u73b0\u7387\u3001\u589e\u5f3a\u8bc1\u636e\u8d28\u91cf\u2014\u2014\u628a\u5ba1\u8ba1\u7ecf\u9a8c\u6c89\u6dc0\u4e3a\u53ef\u590d\u7528\u7684\u6570\u636e\u5de5\u5177\u3002", size=14, color=MU, before=4, after=8)

for num, name, items in [
    ("01", "\u6570\u636e\u6807\u51c6", "\u8d22\u52a1\u3001\u9884\u7b97\u3001\u652f\u4ed8\u3001\u5408\u540c\u3001\u91c7\u8d2d\u3001\u8d44\u4ea7\u3001\u5de5\u7a0b\u9879\u76ee\u5b57\u6bb5\u6574\u7406"),
    ("02", "\u89c4\u5219\u6a21\u578b", "\u91cd\u590d\u652f\u4ed8\u3001\u8d85\u9884\u7b97\u6267\u884c\u8bc6\u522b \u00b7 \u4f9b\u5e94\u5546\u5f02\u5e38\u3001\u8d44\u91d1\u6c89\u6dc0\u68c0\u6d4b"),
    ("03", "\u7a7f\u900f\u6838\u67e5", "\u7591\u70b9\u6765\u6e90\u3001\u6838\u67e5\u8def\u5f84\u3001\u4f50\u8bc1\u6750\u6599\u3001\u5f71\u54cd\u91d1\u989d\u3001\u6574\u6539\u5efa\u8bae"),
    ("04", "\u62a5\u544a\u590d\u6838", "\u91d1\u989d\u6c47\u603b\u6821\u9a8c\u3001\u53e3\u5f84\u4e00\u81f4\u6027 \u00b7 \u9644\u8868\u95ed\u73af\u3001\u7ed3\u8bba\u4f9d\u636e\u53ef\u8ffd\u6eaf"),
]:
    p(doc, f"[{num}]  {name}", size=15, bold=True, color=NW, before=6, after=2)
    p(doc, f"    {items}", size=12, color=MU, before=0, after=4)
spacer(4)
p(doc, "\u2500" * 60, size=4, color=GD, before=4, after=4)
p(doc, "\u6570\u5b57\u5316\u4e0d\u662f\u66ff\u4ee3\u4e13\u4e1a\u5224\u65ad\uff0c\u800c\u662f\u8ba9\u4e13\u4e1a\u5224\u65ad\u8986\u76d6\u66f4\u591a\u6570\u636e\u3001\u9501\u5b9a\u66f4\u51c6\u7591\u70b9\u3001\u751f\u6210\u66f4\u5f3a\u8bc1\u636e\u3002", size=13, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=0)

# ==================== PAGE 9: EXPERIENCE ====================
section_break()
heading("\u4ee3\u8868\u7ecf\u9a8c", "REPRESENTATIVE EXPERIENCE")
p(doc, "\u957f\u671f\u670d\u52a1\u7701\u3001\u5e02\u3001\u53bf\u591a\u7ea7\u8d22\u653f\u548c\u4e3b\u7ba1\u90e8\u95e8\uff0c\u8986\u76d6\u9884\u7b97\u7ee9\u6548\u3001\u8d22\u653f\u76d1\u7763\u3001\u4e13\u9879\u8d44\u91d1\u8bc4\u4ef7\u3001\u5de5\u7a0b\u51b3\u7b97\u548c\u7ba1\u7406\u54a8\u8be2\u3002", size=14, color=MU, before=4, after=8)

# Territory
p(doc, "\u670d\u52a1\u7248\u56fe", size=14, bold=True, color=NW, before=4, after=4)
for name, desc in [("\u6210\u90fd\u603b\u90e8", "\u56db\u5ddd\u7701\u4f1a\uff0c\u6838\u5fc3\u67a2\u7ebd"), ("\u963f\u575d\u5dde\u529e\u4e8b\u5904", "\u5ddd\u897f\u9ad8\u539f\uff0c\u8986\u76d6\u85cf\u533a"), ("\u897f\u85cf\u529e\u4e8b\u5904", "\u9ad8\u539f\u5730\u533a\u62d3\u5c55"), ("\u8986\u76d6\u5ddd\u3001\u85cf\u3001\u9ed4", "\u4e09\u7701\u8054\u52a8\uff0c\u8de8\u533a\u57df\u670d\u52a1")]:
    bullet(f"{name}\uff1a{desc}", size=11, color=DK, indent=0)
spacer(4)

# Clients
p(doc, "\u4ee3\u8868\u5ba2\u6237", size=14, bold=True, color=NW, before=8, after=4)
clients = [
    "\u56db\u5ddd\u7701\u8d22\u653f\u5385  \u00b7  \u7701\u516c\u5b89\u5385\u4ea4\u8b66\u603b\u961f  \u00b7  \u7701\u6c11\u653f\u5385",
    "\u7701\u519c\u4e1a\u519c\u6751\u5385  \u00b7  \u7701\u6559\u80b2\u5385  \u00b7  \u7701\u9000\u5f79\u519b\u4eba\u4e8b\u52a1\u5385",
    "\u7701\u836f\u54c1\u76d1\u7763\u7ba1\u7406\u5c40  \u00b7  \u8fbe\u5dde\u5e02\u8d22\u653f\u5c40  \u00b7  \u7ef5\u9633\u5e02\u8d22\u653f\u5c40",
    "\u5b9c\u5bbe\u5e02\u8d22\u653f\u5c40  \u00b7  \u4ec0\u90a1\u5e02\u8d22\u653f\u5c40  \u00b7  \u5fb7\u9633\u5e02\u8d22\u653f\u5c40",
    "\u5eb7\u5b9a\u5e02\u8d22\u653f\u5c40  \u00b7  \u4e5d\u5be8\u6c9f\u53bf\u8d22\u653f\u5c40",
]
for c in clients:
    bullet(c, size=11, color=DK, indent=0)
spacer(4)
p(doc, "\u8986\u76d6\u9886\u57df\uff1a\u8d22\u653f\u3001\u5ba1\u8ba1\u3001\u6559\u80b2\u3001\u6c11\u653f\u3001\u4ea4\u901a\u3001\u56fd\u8d44\u3001\u533b\u4fdd\u3001\u836f\u76d1\u3001\u516c\u5b89", size=12, color=MU, before=6, after=0)

# ==================== PAGE 10: CONTACT ====================
section_break()
heading("\u5408\u4f5c\u4ef7\u503c", "VALUE PROPOSITION")
spacer(6)

p(doc, "\u6211\u4eec\u4ea4\u4ed8\u7684\u4e0d\u53ea\u662f\u4e00\u4efd\u62a5\u544a\uff0c\u66f4\u662f\u4e00\u5957\u53ef\u6267\u884c\u7684\u6539\u8fdb\u65b9\u6848\u3002", size=14, color=NW, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10)
spacer(8)

# Contact info
p(doc, "\u8054\u7cfb\u878d\u7b56", size=18, bold=True, color=GD, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
p(doc, "\u260e  028-87659276", size=14, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)
p(doc, "\u2709  scrccpa@163.com", size=14, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)
p(doc, "\u25b6  www.scrccpa.com", size=14, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)
p(doc, "\u2302  \u6210\u90fd\u5e02\u91d1\u725b\u533a\u91d1\u5468\u8def 595 \u53f7", size=14, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)
spacer(12)

# Values
p(doc, "\u5408\u4f5c\u627f\u8bfa", size=18, bold=True, color=GD, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
for item in ["\u5ba2\u89c2\u516c\u6b63  \u5b9e\u4e8b\u6c42\u662f", "\u8d28\u91cf\u4f18\u5148  \u8bc1\u636e\u95ed\u73af", "\u95ee\u9898\u6709\u4f9d\u636e  \u7ed3\u8bba\u53ef\u89e3\u91ca", "\u5efa\u8bae\u80fd\u843d\u5730  \u6574\u6539\u6709\u8ffd\u8e2a"]:
    p(doc, "\u25c6  " + item, size=13, color=DK, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)

spacer(24)
p(doc, "\u56db\u5ddd\u878d\u7b56\u4f1a\u8ba1\u5e08\u4e8b\u52a1\u6240\u6709\u9650\u516c\u53f8", size=12, color=GRY, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

# ==================== SAVE ====================
doc.core_properties.title = "四川融策宣传册 v9 可编辑Word版"
doc.save(str(DOCX))
print(f"DONE: {DOCX}")
print(f"Size: {DOCX.stat().st_size / 1024:.1f} KB")
