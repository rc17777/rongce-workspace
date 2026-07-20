# -*- coding: utf-8 -*-
"""
将三个Markdown文档转换为Word格式
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

input_dir = r'C:\Users\scrccpa\.openclaw\workspace\chayu_analysis'
output_dir = input_dir

md_files = [
    '资产确权情况说明模板.md',
    '四个一批项目资料齐全率统计报告.md',
    '低效闲置资产盘活方案.md',
]

def md_to_docx(md_file, output_file):
    """将Markdown转Word"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 读取Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 解析Markdown
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 一级标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 二级标题
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        
        # 三级标题
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        
        # 四级标题
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            i += 1
            continue
        
        # 分隔线
        if line.startswith('---'):
            doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # 引用块
        if line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Cm(1)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue
        
        # 表格 (简化处理：检测 | 开头)
        if '|' in line and not line.startswith('|---'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            
            # 过滤分隔行
            table_lines = [l for l in table_lines if not l.startswith('|---') and not l.startswith('| ---')]
            
            if table_lines:
                # 解析表格
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')]
                    # 去掉首尾空单元格
                    if cells and not cells[0]:
                        cells = cells[1:]
                    if cells and not cells[-1]:
                        cells = cells[:-1]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # 创建Word表格
                    max_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_text in enumerate(row):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                            # 表头加粗
                            if r_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            continue
        
        # 无序列表
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # 有序列表
        match = re.match(r'^(\d+)\.\s+(.+)', line)
        if match:
            text = match.group(2)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # 普通段落
        # 处理加粗 **text**
        para_text = line
        p = doc.add_paragraph()
        
        # 简单替换加粗标记
        parts = re.split(r'(\*\*[^*]+\*\*)', para_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        
        i += 1
    
    # 保存
    doc.save(output_file)
    print(f'✅ 已转换: {output_file}')

# 转换三个文件
for md_file in md_files:
    md_path = os.path.join(input_dir, md_file)
    docx_file = md_file.replace('.md', '.docx')
    docx_path = os.path.join(output_dir, docx_file)
    
    if os.path.exists(md_path):
        md_to_docx(md_path, docx_path)
    else:
        print(f'❌ 文件不存在: {md_path}')

print(f'\n全部完成！输出目录: {output_dir}')
