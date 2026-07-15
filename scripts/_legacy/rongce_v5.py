# -*- coding: utf-8 -*-
"""四川融策宣传册 v5 —— 第一性原理深挖版（第一部分：工具函数 + P1-P3）"""
from __future__ import annotations
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from PIL.Image import LANCZOS
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

OUT = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
WK = Path("work/sichuan_rongce_brochure/v5_assets")
DOCX = OUT / "四川融策宣传册_完善稿_v5_第一性原理版.docx"
PDF = OUT / "四川融策宣传册_完善稿_v5_第一性原理版.pdf"

NW = "#1A365D"
TL = "#3A7B8A"
GD = "#D4A574"
MU = "#718096"
IK = "#2D3748"
WH = "#FFFFFF"
BG = "#FAF8F4"
W, H = 1653, 2339

def rgb(s):
    h = s.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
def rga(s, a=255):
    return (*rgb(s), a)

def font(size, bold=False):
    for p in [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()

def draw(d, text, x, y, f, fill, w=None, gap=6, align="left"):
    mw = w or 9999
    lines = []
    for para in text.split("\n"):
        cur = ""
        for ch in para:
            if d.textlength(cur + ch, font=f) <= mw:
                cur += ch
            else:
                if cur: lines.append(cur)
                cur = ch
        if cur: lines.append(cur)
    for line in lines:
        xx = x
        if align == "center" and w:
            xx = x + (w - d.textlength(line, font=f)) // 2
        d.text((xx, y), line, font=f, fill=fill)
        y += f.size + gap
    return y

def left_bar(d):
    d.rectangle([0, 0, 18, H], fill=rga(NW))
    d.rectangle([18, 0, 24, H], fill=rga(GD))

def bottom_bar(d):
    tx = "公开公正  用心服务  诚信为本  服务至上  追求卓越"
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text(((W-d.textlength(tx, font=font(18)))//2, H-42), tx, font=font(18), fill=rga(WH, 180))

def page_title(d, title, en, y=120):
    d.text((160, y), title, font=font(52, True), fill=NW)
    d.text((160, y+68), en, font=font(22), fill=rga(GD))
    d.line([160, y+105, 600, y+105], fill=rga(GD), width=3)

# ======= P1: Cover =======
def p1_cover():
    img = Image.new("RGB", (W, H), rgb(BG))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, 260], fill=rga(NW))
    d.rectangle([0, 260, W, 270], fill=rga(GD))
    d.ellipse([W-500, -150, W-50, 260], fill=rga(TL, 35))
    d.ellipse([W-400, -80, W-150, 180], fill=rga(GD, 20))
    d.text((130, 60), "SICHUAN", font=font(72, True), fill=rga(WH, 200))
    d.text((130, 135), "RONGCE", font=font(100, True), fill=WH)
    d.text((130, 320), "谋专业之策  融品质之精", font=font(48, True), fill=NW)
    d.text((130, 400), "政府审计与工程咨询综合服务机构", font=font(30), fill=MU)
    d.line([130, 480, 500, 480], fill=rga(GD), width=5)
    d.rectangle([0, H-60, W, H], fill=rga(NW))
    d.text((130, H-45), "审计 · 绩效 · 财政监督 · 工程咨询 · 数字化分析", font=font(20), fill=rga(WH, 200))
    d.line([W-80, 80, W-30, 80], fill=rga(GD, 60), width=3)
    d.line([W-30, 80, W-30, 130], fill=rga(GD, 60), width=3)
    return img

# ======= P2: Why Rongce =======
def p2_about():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "为什么是融策？", "WHY RONGCE")
    y = draw(d, "财政资金的管理，核心是两个问题：", 160, 280, font(36), IK, 650, 12)
    y = draw(d, "\u201c资金到底去哪了？\u201d 和 \u201c这钱花得值不值？\u201d", 160, y, font(36, True), NW, 650, 12)
    y = draw(d, "\n融策做的事情，就是用专业的方法和数据工具，帮您把这两个问题搞清楚。", 160, y+20, font(28), MU, 650, 10)
    for i, (icon, head, desc) in enumerate([
        ("\u2460", "不只查账", "我们看的是你的合同、项目、资产、\n内部控制制度和决策程序。账表只是入口。"),
        ("\u2461", "不只找问题", "我们更关心问题能不能改、怎么改、\n改得怎么样。整改落地才是目标。"),
        ("\u2462", "不只靠经验", "数据分析和现场核查结合，\n用数据扩大覆盖面，用现场核实关键点。"),
    ]):
        x = 160 + i * 380
        d.rounded_rectangle([x, 470, x+340, 680], radius=14, fill=rga(WH), outline=rga(TL, 50), width=2)
        d.rounded_rectangle([x, 470, x+340, 560], radius=14, fill=rga(TL, 25))
        d.text((x+32, 495), icon, font=font(36, True), fill=NW)
        d.text((x+80, 500), head, font=font(30, True), fill=NW)
        draw(d, desc, x+32, 580, font(24), MU, 340-64, 8)
    y = 740
    d.rounded_rectangle([160, 720, 1490, 880], radius=14, fill=rga("#E8E0D4", 60))
    for item in [
        "始于 2000 年，四川省内较早成立的会计师事务所之一",
        "长期服务财政、审计、教育、民政、交通、国资、医保等领域",
        "形成审计、绩效、财政监督、工程咨询协同发展的业务格局",
        "覆盖四川、西藏、贵州三省"
    ]:
        d.ellipse([180, y+8, 192, y+20], fill=rga(GD))
        y = draw(d, item, 210, y, font(26), IK, 1250, 12)
    bottom_bar(d)
    return img

# ======= P3: Methodology =======
def p3_method():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "我们怎么做？", "OUR METHODOLOGY")
    for i, (num, title, items) in enumerate([
        ("01", "资金线", ["钱从哪来？到哪去？花得值不值？", "资金流、支付流、票据流三流交叉核验"]),
        ("02", "项目线", ["立项、招标、合同、施工、验收、结算", "全生命周期穿透核查，不留死角"]),
        ("03", "责任线", ["决策程序、岗位职责、内控制度、整改落实", "把问题追溯到人，让结论有依据、责任可追溯"]),
    ]):
        x = 160 + i * 380
        d.rounded_rectangle([x, 280, x+340, 560], radius=16, fill=rga(WH), outline=rga(TL, 60), width=2)
        d.rounded_rectangle([x, 280, x+340, 370], radius=16, fill=rga(TL, 30))
        d.text((x+32, 305), num, font=font(40, True), fill=rga(NW, 80))
        d.text((x+80, 315), title, font=font(32, True), fill=NW)
        y0 = 400
        for item in items:
            d.ellipse([x+32, y0+6, x+44, y0+18], fill=rga(GD))
            y0 = draw(d, item, x+55, y0, font(22), MU, 340-70, 8)
        y0 += 10
    steps = ["制度审查", "数据核验", "现场核查", "证据闭环", "整改建议"]
    for i, step in enumerate(steps):
        x = 200 + i * 270
        d.ellipse([x, 620, x+80, 700], fill=rga(NW if i%2==0 else TL))
        tw = d.textlength(str(i+1), font=font(32, True))
        d.text((x+40-tw//2, 638), str(i+1), font=font(32, True), fill=WH)
        draw(d, step, x+100, 638, font(26), IK, 150, 6)
        if i < 4:
            d.line([x+80, 660, x+190, 660], fill=rga(GD), width=3)
    d.rounded_rectangle([160, 800, 1490, 930], radius=14, fill=rga("#E8E0D4", 80))
    draw(d, "我们的目标不是出具一份报告，而是帮助委托方看清问题、厘清责任、找到改进路径。", 200, 835, font(28), IK, 1250, 10, "center")
    bottom_bar(d)
    return img

# ======= Service page template =======
def service_page(title, en, sub, items, n):
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, title, en, 100)
    draw(d, sub, 160, 240, font(30), MU, 620, 10)
    d.rounded_rectangle([880, 120, 1450, 580], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((960, 170), f"0{n}", font=font(80, True), fill=rga(TL, 60))
    d.text((960, 270), "RONGCE", font=font(40, True), fill=rga(TL, 40))
    tags = {
        1: ['从单项到全过程', '从结果到制度'],
        2: ['从有没有', '到值不值'],
        3: ['让资金从花了', '走向花得值'],
        4: ['从概算到结算', '全关贯通'],
        5: ['数据扩大覆盖面', '经验沉淀为工具'],
    }.get(n, [])
    for i, t in enumerate(tags):
        draw(d, t, 1020, 360+i*60, font(28, True), rga(TL, 80), 400, 6)
    y = 330
    for i, item in enumerate(items):
        clr = rga(NW if i%2==0 else TL)
        d.rounded_rectangle([160, y, 1580, y+48], radius=8, fill=rga(WH), outline=rga(GD, 80), width=1)
        d.rectangle([160, y, 168, y+48], fill=clr)
        y = draw(d, item[0], 185, y+8, font(26, True), IK, 1350, 6) + 10
        for sub in item[1]:
            d.ellipse([185, y+8, 193, y+16], fill=rga(GD))
            y = draw(d, sub, 205, y, font(22), MU, 1300, 6) + 10
    bottom_bar(d)
    return img

def p4_services():
    return service_page("服务体系", "SERVICE SYSTEM",
        "从单项审计到全过程咨询，从结果评价到制度优化，从人工核查到数据化识别。",
        [("政府审计", ["经济责任审计", "专项资金审计", "财政监督检查", "工程决算财务审计"]),
         ("预算绩效管理", ["事前评估 · 目标审核", "运行监控 · 重点评价", "结果应用"]),
         ("工程咨询", ["预算编制 · 财政评审", "结算审核 · 全过程工程咨询"]),
         ("采购审计", ["采购程序合规", "围标串标线索", "合同履约核查"]),
         ("管理咨询", ["内控建设 · 资产管理", "整改提升 · 流程优化"])], 1)

def p5_gov():
    return service_page("政府审计", "GOVERNMENT AUDIT",
        "从“有没有”到“对不对”再到“值不值”——把账表、合同、项目、资金、资产、责任贯通核查。",
        [("经济责任审计", ["重大决策与资金资产安全", "项目建设与内控风险识别", "廉政风险排查与责任追溯"]),
         ("专项资金审计", ["申报分配、拨付使用全链条核查", "绩效评价与结余沉淀分析"]),
         ("财政监督检查", ["预算执行与财经纪律合规", "政府采购与会计信息质量检查"]),
         ("工程决算财务审计", ["建设成本归集与资金来源核实", "资产交付与尾工尾款管理"])], 2)

def p6_perf():
    return service_page("预算绩效管理", "BUDGET PERFORMANCE",
        "您关心的不只是“花了多少钱”，更是“效果怎么样”——让财政资金从“花了没有”走向“花得值不值”。",
        [("事前绩效评估", ["必要性、可行性与财政承受能力分析", "预期绩效与投入成本综合评估"]),
         ("绩效目标审核", ["目标完整性与指标可衡量性审查", "预算匹配性与绩效责任书审核"]),
         ("绩效运行监控", ["执行进度与资金支付追踪", "产出偏差分析与风险预警"]),
         ("重点绩效评价", ["政策评价、部门整体评价", "项目支出与专项资金评价"]),
         ("结果应用", ["整改清单与预算挂钩", "管理制度优化建议"])], 3)


def p7_eng():
    return service_page("工程咨询与财政评审", "ENGINEERING CONSULTING",
        "从概算到结算，帮您把好每一道关——工程造价、合同履约、资金支付和项目绩效协同管理。",
        [("预算编制与财政评审", ["工程量、定额套用核验", "材料价格、措施费、取费标准"]),
         ("清单及招标控制价", ["提升招标文件和控制价编制质量"]),
         ("结算审核", ["合同条款、变更签证核查", "隐蔽工程、现场工程量核验"]),
         ("全过程工程咨询", ["项目前期、招采、实施", "验收、结算、绩效评价协同"])], 4)

# ======= P8: Digital Audit =======
def p8_digital():
    return service_page("数字化审计能力", "DIGITAL AUDIT CAPABILITIES",
        "用数据扩大覆盖面、提高发现率、增强证据质量——把审计经验沉淀为可复用的数据工具。",
        [("数据标准", ["财务、预算、支付、合同", "采购、资产、工程项目字段整理"]),
         ("规则模型", ["重复支付、超预算执行识别", "供应商异常、资金沉淀检测"]),
         ("穿透核查", ["疑点来源、核查路径", "佐证材料、影响金额、整改建议"]),
         ("报告复核", ["金额汇总校验、口径一致性", "附表闭环、结论依据可追溯"])], 5)

# ======= P9: Experience =======
def p9_experience():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "代表经验", "REPRESENTATIVE EXPERIENCE")
    draw(d, "长期服务省、市、县多级财政和主管部门，覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。", 160, 240, font(30), MU, 620, 10)
    d.rounded_rectangle([880, 130, 1450, 550], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((960, 180), "服务版图", font=font(36, True), fill=rga(TL, 60))
    d.text((960, 230), "SERVICE TERRITORY", font=font(18), fill=rga(TL, 40))
    terr = [("成都总部","四川省会，核心枢纽"),("阿坝州办事处","川西高原，覆盖藏区"),
            ("西藏办事处","高原地区拓展"),("覆盖川、藏、黔","三省联动，跨区域服务")]
    for i, (nm, dc) in enumerate(terr):
        y0 = 280 + i*60
        d.ellipse([960, y0, 990, y0+30], fill=rga(GD))
        draw(d, nm, 1010, y0, font(26, True), IK, 300, 2)
        draw(d, dc, 1010, y0+28, font(18), MU, 300, 2)
        if i < 3: d.line([975, y0+30, 975, y0+60], fill=rga(GD, 100), width=2)
    d.text((160, 390), "代表客户", font=font(32, True), fill=NW)
    d.text((160, 428), "REPRESENTATIVE CLIENTS", font=font(18), fill=rga(GD))
    clients = ["四川省财政厅","省公安厅交警总队","省民政厅","省农业农村厅","省教育厅",
        "省退役军人事务厅","省药品监督管理局","达州市财政局","绵阳市财政局",
        "宜宾市财政局","什邡市财政局","德阳市财政局","康定市财政局","九寨沟县财政局"]
    yv = 470
    for i, c in enumerate(clients):
        if i % 5 == 0 and i > 0: yv += 40
        cx = 160 + (i%5)*280
        d.rounded_rectangle([cx, yv, cx+240, yv+36], radius=6, fill=rga(WH), outline=rga(GD, 80), width=1)
        tw = d.textlength(c, font=font(20))
        d.text((cx+120-tw//2, yv+5), c, font=font(20), fill=NW)
    d.text((160, 710), "覆盖领域：财政、审计、教育、民政、交通、国资、医保、药监、公安", font=font(22), fill=MU)
    bottom_bar(d)
    return img

# ======= P10: Contact =======
def p10_contact():
    img = Image.new("RGB", (W, H), rgb(BG))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, H], fill=rga(NW))
    d.rectangle([0, 0, W, 280], fill=rga(NW))
    d.rectangle([0, 280, W, 290], fill=rga(GD))
    d.text((130, 80), "合作价值", font=font(52, True), fill=WH)
    d.text((130, 150), "VALUE PROPOSITION", font=font(22), fill=rga(GD, 200))
    d.line([130, 190, 500, 190], fill=rga(GD), width=3)
    draw(d, "我们交付的不只是一份报告，更是一套可执行的改进方案。", 130, 260, font(30), WH, 600, 10)
    d.rounded_rectangle([130, 340, 750, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "联系融策", 180, 370, font(36, True), GD, 500, 8)
    for i, (icon, line) in enumerate([("\u260e","028-87659276"),("\u2709","scrccpa@163.com"),
                                       ("\u25b6","www.scrccpa.com"),("\u2302","成都市金牛区金周路595号")]):
        d.text((180, 440+i*48), icon, font=font(26), fill=GD)
        d.text((215, 438+i*48), line, font=font(26), fill=WH)
    d.rounded_rectangle([880, 340, 1530, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "合作承诺", 930, 380, font(36, True), GD, 500, 8)
    for i, item in enumerate(["客观公正  实事求是","质量优先  证据闭环",
                               "问题有依据  结论可解释","建议能落地  整改有追踪"]):
        d.ellipse([930, 460+i*48, 946, 476+i*48], fill=rga(GD))
        d.text((965, 458+i*48), item, font=font(26), fill=WH)
    d.ellipse([W-300, H-400, W-50, H-150], fill=rga(GD, 15))
    d.ellipse([W-250, H-350, W-100, H-200], fill=rga(TL, 15))
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text((130, H-40), "四川融策会计师事务所有限公司", font=font(20), fill=rga(GD, 180))
    return img

# ======= PDF =======
def generate_pdf(images):
    c = canvas.Canvas(str(PDF), pagesize=A4)
    pw, ph = A4
    for img_path in images:
        im = Image.open(img_path).resize((int(pw), int(ph)-20), LANCZOS)
        tmp = WK / "tmp_render.png"
        im.save(tmp)
        c.drawImage(str(tmp), 0, 10, width=pw, height=ph-20, preserveAspectRatio=False)
        tmp.unlink(missing_ok=True)
        c.showPage()
    c.save()

# ======= DOCX =======
def generate_docx(images):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    lbls = ["封面","为什么是融策","方法论","服务体系","政府审计","预算绩效","工程咨询","数字化","代表经验","合作价值"]
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v5第一性原理版"
    for idx, (lbl, img_path) in enumerate(zip(lbls, images), 1):
        if idx > 1: doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(lbl)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(22); r.font.bold = True
        doc.add_picture(str(img_path), width=Cm(17.0))
    doc.save(DOCX)

# ======= Main =======
def main():
    OUT.mkdir(parents=True, exist_ok=True)
    WK.mkdir(parents=True, exist_ok=True)
    # Import all page functions from part1
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    from rongce_v5_part1 import p1_cover, p2_about, p3_method, p4_services, p5_gov, p6_perf
    makers = [p1_cover, p2_about, p3_method, p4_services, p5_gov, p6_perf,
              p7_eng, p8_digital, p9_experience, p10_contact]
    print("Rendering 10 pages...")
    images = []
    for i, maker in enumerate(makers, 1):
        out = WK / f"page_{i:02d}.png"
        if not out.exists():
            print(f"  Page {i}...")
            maker().save(out, quality=95)
        images.append(out)
    print("Generating PDF...")
    generate_pdf(images)
    print("Generating DOCX...")
    generate_docx(images)
    print(f"DONE\n{DOCX}\n{PDF}")

if __name__ == "__main__":
    main()