#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""提取审计报告和附表内容"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

import docx
import openpyxl

# ===== 1. 提取报告 =====
path = r'C:\Users\scrccpa\Desktop\四川食在攒劲餐饮服务有限公司2024年7月—2025年12月财务收支专项审计报告\四川食在攒劲餐饮服务有限公司2024年7月—2025年12月财务收支专项审计报告-（6月1日）.docx'
doc = docx.Document(path)
print('===== 报告全文 =====')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[P{i}] {p.text}')

print('\n\n===== 报告内表格 =====')
for ti, tbl in enumerate(doc.tables):
    print(f'\n--- 表格{ti+1} ---')
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f'  R{ri}: {" | ".join(cells)}')

# ===== 2. 提取附表 =====
xlsx_path = r'C:\Users\scrccpa\Desktop\四川食在攒劲餐饮服务有限公司2024年7月—2025年12月财务收支专项审计报告\附件1：截止2025年12月31日往来余额比对明细表.xlsx'
wb = openpyxl.load_workbook(xlsx_path, data_only=True)
print('\n\n===== 附表Sheet列表 =====')
for name in wb.sheetnames:
    print(f'  Sheet: {name}')

print('\n\n===== 附表内容 =====')
for name in wb.sheetnames:
    ws = wb[name]
    print(f'\n--- Sheet: {name} (rows={ws.max_row}, cols={ws.max_column}) ---')
    for ri, row in enumerate(ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True)):
        vals = [str(v) if v is not None else '' for v in row]
        print(f'  R{ri}: {" | ".join(vals)}')
