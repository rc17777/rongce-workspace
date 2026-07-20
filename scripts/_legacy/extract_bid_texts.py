#!/usr/bin/env python3
"""提取校服采购项目所有文件文本"""
import pdfplumber
import os
import sys

BASE = r"C:\Users\scrccpa\Desktop\校服\2025年校服采购\校服\2025年"
OUT = r"D:\openclaw-workspace\output\校服分析\txt"
os.makedirs(OUT, exist_ok=True)

def extract_pdf(pdf_path, out_name):
    """提取PDF文本"""
    out_path = os.path.join(OUT, out_name)
    try:
        with pdfplumber.open(pdf_path) as pdf:
            texts = []
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t:
                    texts.append(f"--- Page {i+1} ---\n{t}")
            full = "\n\n".join(texts)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(full)
            print(f"✅ PDF: {out_name} ({len(full)} chars, {len(pdf.pages)} pages)")
            return len(full)
    except Exception as e:
        print(f"❌ PDF {out_name}: {e}")
        return 0

def extract_docx(docx_path, out_name):
    """提取DOCX文本"""
    out_path = os.path.join(OUT, out_name)
    try:
        from docx import Document
        doc = Document(docx_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        # Also extract tables
        for i, table in enumerate(doc.tables):
            texts.append(f"\n--- Table {i+1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                texts.append(" | ".join(cells))
        full = "\n".join(texts)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(full)
        print(f"✅ DOCX: {out_name} ({len(full)} chars)")
        return len(full)
    except Exception as e:
        print(f"❌ DOCX {out_name}: {e}")
        return 0

# Step 1: Extract key procurement documents
print("=" * 60)
print("Step 1: 招标文件 & 采购公告")
print("=" * 60)

extract_pdf(
    os.path.join(BASE, "1.工会及招标采购", "招标文件-成都市教育科学研究院附属中学学校校服采购项目（9.25）.pdf"),
    "01-招标文件.txt"
)

extract_pdf(
    os.path.join(BASE, "1.工会及招标采购", "评审结果和采购结果文.pdf"),
    "02-评审结果.txt"
)

extract_pdf(
    os.path.join(BASE, "1.工会及招标采购", "成都市教育科学研究院附属中学学校校服采购项目招标通知书.pdf"),
    "03-招标通知书.txt"
)

# 采购公告 & 前期文件
proc_docs = [
    ("1.校服采购前期纪要.pdf", "04-采购前期纪要.txt"),
    ("2.需求公告前公示.pdf", "05-需求公告前公示.txt"),
    ("3.校服选用需求统计.pdf", "06-需求统计.txt"),
    ("4.校服采购公告.pdf", "07-采购公告.txt"),
    ("5.选用方式及代理机构.pdf", "08-选用方式及代理机构.txt"),
    ("6.校服采购承诺书.pdf", "09-采购承诺书.txt"),
    ("7.选用小组签到表.pdf", "10-选用小组签到表.txt"),
]
for fname, oname in proc_docs:
    extract_pdf(os.path.join(BASE, "1.工会及招标采购", "采购需求公告", fname), oname)

print()
print("=" * 60)
print("Step 2: 投标文件 - 资格标")
print("=" * 60)

# 资格投标文件 (all are docx or doc)
bidders = [
    ("四川锦鸿德凯实业有限公司", 
     "成都市教育科学研究院附属中学学校校服采购项目资格投标文件.docx", "B1-锦鸿德凯-资格标.txt"),
    ("四川琳耀商贸有限公司",
     "成都市教育科学研究院附属中学资格投标文件-终稿.docx", "B2-琳耀-资格标.txt"),
    ("成都顺华服装有限公司",
     "资格投标文件.doc", "B4-顺华-资格标.txt"),
    ("鼎新邑和锦富康汇品牌管理有限公司",
     "20251012资格投标文件终稿(1).docx", "B5-鼎新邑和-资格标.txt"),
]

BID_DIR = os.path.join(BASE, "投标文件", "投标文件")
for company, fname, oname in bidders:
    path = os.path.join(BID_DIR, company, fname)
    if os.path.exists(path):
        if fname.endswith('.docx'):
            extract_docx(path, oname)
        elif fname.endswith('.doc'):
            print(f"⚠️  .doc file: {oname} - need conversion, skipping for now")
    else:
        print(f"❌ Not found: {path}")

# 博博士 is .doc
print("⚠️  B3-博博士-资格标.doc - binary format, need special handling")

print()
print("=" * 60)
print("Step 3: 投标文件 - 商务标")
print("=" * 60)

biz_bids = [
    ("四川锦鸿德凯实业有限公司", "成都市教育科学研究院附属中学学校校服采购项目商务投标文件.docx", "B1-锦鸿德凯-商务标.txt"),
    ("四川琳耀商贸有限公司", "成都市教育科学研究院附属中学商务投标文件-终稿.docx", "B2-琳耀-商务标.txt"),
    ("鼎新邑和锦富康汇品牌管理有限公司", "20251012商务投标文件终稿(1).docx", "B5-鼎新邑和-商务标.txt"),
]

for company, fname, oname in biz_bids:
    path = os.path.join(BID_DIR, company, fname)
    if os.path.exists(path):
        if fname.endswith('.docx'):
            extract_docx(path, oname)
    else:
        print(f"❌ Not found: {path}")

print("⚠️  B3-博博士-商务标.doc & B4-顺华-商务标.doc - binary format, need special handling")

print()
print("=" * 60)
print("Step 4: 合同文件")
print("=" * 60)

extract_pdf(
    os.path.join(BASE, "2.合同签订", "合同审计稿.pdf"),
    "11-合同审计稿.txt"
)
extract_pdf(
    os.path.join(BASE, "2.合同签订", "成都市教育科学研究院附属中学学校校服采购项目合同.pdf"),
    "12-合同.txt"
)

print()
print("Done!")
