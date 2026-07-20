#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成《2026年阿坝县财政局采购政府性投资项目竣工财务决算机构》
服务方案 Word 文档
"""
import os
import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\openclaw-workspace\2026年阿坝县财政局竣工财务决算审核服务方案_v3.docx"
SVG_DIR = r"D:\openclaw-workspace\svgs"
os.makedirs(SVG_DIR, exist_ok=True)

# ============================================================
# SVG 图表生成
# ============================================================

def generate_flowchart_svg():
    """审核工作流程图"""
    svg = '''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 1100" width="800" height="1100">
  <defs>
    <marker id="arrowhead" markerWidth="10" markerHeight="7" refX="10" refY="3.5" orient="auto">
      <polygon points="0 0, 10 3.5, 0 7" fill="#2c3e50"/>
    </marker>
    <linearGradient id="headerGrad" x1="0%" y1="0%" x2="100%" y2="0%">
      <stop offset="0%" style="stop-color:#1a5276;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#2980b9;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="stepGrad" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#ecf0f1;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#d5dbdb;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="reviewGrad" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#eaf2f8;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#d4e6f1;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="outputGrad" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#e8f8f5;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#d1f2eb;stop-opacity:1"/>
    </linearGradient>
    <filter id="shadow" x="-5%" y="-5%" width="110%" height="110%">
      <feDropShadow dx="2" dy="3" stdDeviation="3" flood-opacity="0.15"/>
    </filter>
  </defs>

  <!-- 标题 -->
  <rect x="50" y="20" width="700" height="50" rx="8" fill="url(#headerGrad)" filter="url(#shadow)"/>
  <text x="400" y="52" text-anchor="middle" fill="white" font-size="20" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阿坝县工程竣工财务决算审核工作流程图</text>

  <!-- ====== 阶段一：接受委托 ====== -->
  <rect x="50" y="90" width="700" height="130" rx="6" fill="none" stroke="#1a5276" stroke-width="2"/>
  <text x="70" y="115" fill="#1a5276" font-size="16" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阶段一：接受委托与前期准备</text>

  <rect x="100" y="130" width="180" height="36" rx="4" fill="url(#stepGrad)" stroke="#7f8c8d" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="190" y="153" text-anchor="middle" fill="#2c3e50" font-size="13" font-family="Microsoft YaHei, SimHei, sans-serif">签订委托协议</text>

  <line x1="280" y1="148" x2="330" y2="148" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="335" y="130" width="180" height="36" rx="4" fill="url(#stepGrad)" stroke="#7f8c8d" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="425" y="153" text-anchor="middle" fill="#2c3e50" font-size="13" font-family="Microsoft YaHei, SimHei, sans-serif">组建审核工作组</text>

  <line x1="515" y1="148" x2="565" y2="148" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="570" y="130" width="150" height="36" rx="4" fill="url(#stepGrad)" stroke="#7f8c8d" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="645" y="153" text-anchor="middle" fill="#2c3e50" font-size="13" font-family="Microsoft YaHei, SimHei, sans-serif">收集项目资料</text>

  <!-- 连接线 -->
  <line x1="400" y1="220" x2="400" y2="255" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <!-- ====== 阶段二：资料审核 ====== -->
  <rect x="50" y="260" width="700" height="130" rx="6" fill="none" stroke="#1a5276" stroke-width="2"/>
  <text x="70" y="285" fill="#1a5276" font-size="16" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阶段二：资料初审与现场核查</text>

  <rect x="100" y="300" width="180" height="36" rx="4" fill="url(#stepGrad)" stroke="#7f8c8d" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="190" y="323" text-anchor="middle" fill="#2c3e50" font-size="13" font-family="Microsoft YaHei, SimHei, sans-serif">资料完整性审查</text>

  <line x1="280" y1="318" x2="330" y2="318" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="335" y="300" width="180" height="36" rx="4" fill="url(#stepGrad)" stroke="#7f8c8d" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="425" y="323" text-anchor="middle" fill="#2c3e50" font-size="13" font-family="Microsoft YaHei, SimHei, sans-serif">现场实地踏勘核查</text>

  <line x1="515" y1="318" x2="565" y2="318" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="570" y="300" width="150" height="36" rx="4" fill="url(#stepGrad)" stroke="#7f8c8d" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="645" y="323" text-anchor="middle" fill="#2c3e50" font-size="13" font-family="Microsoft YaHei, SimHei, sans-serif">编制资料清单</text>

  <!-- 连接线 -->
  <line x1="400" y1="390" x2="400" y2="425" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <!-- ====== 阶段三：实质审核 ====== -->
  <rect x="50" y="430" width="700" height="200" rx="6" fill="none" stroke="#1a5276" stroke-width="2"/>
  <text x="70" y="455" fill="#1a5276" font-size="16" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阶段三：实质性审核</text>

  <rect x="80" y="475" width="155" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="157" y="498" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">工程价款审核</text>
  <text x="157" y="518" text-anchor="middle" fill="#5d6d7e" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">结算/变更/索赔</text>

  <line x1="235" y1="502" x2="270" y2="502" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="275" y="475" width="155" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="352" y="498" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">资金来源审核</text>
  <text x="352" y="518" text-anchor="middle" fill="#5d6d7e" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">拨付/配套/自筹</text>

  <line x1="430" y1="502" x2="465" y2="502" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="470" y="475" width="155" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="547" y="498" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">资金使用审核</text>
  <text x="547" y="518" text-anchor="middle" fill="#5d6d7e" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">合规性/真实性</text>

  <line x1="625" y1="502" x2="660" y2="502" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="665" y="475" width="70" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="700" y="508" text-anchor="middle" fill="#1a5276" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">资产</text>
  <text x="700" y="523" text-anchor="middle" fill="#1a5276" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">交付</text>

  <!-- 第二行 -->
  <rect x="125" y="545" width="155" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="202" y="568" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">待摊投资审核</text>
  <text x="202" y="588" text-anchor="middle" fill="#5d6d7e" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">管理费/监理费等</text>

  <line x1="280" y1="572" x2="335" y2="572" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="340" y="545" width="155" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="417" y="568" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">其他投资审核</text>
  <text x="417" y="588" text-anchor="middle" fill="#5d6d7e" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">设备/工器具等</text>

  <line x1="495" y1="572" x2="550" y2="572" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="555" y="545" width="155" height="55" rx="4" fill="url(#reviewGrad)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="632" y="568" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">尾工工程审核</text>
  <text x="632" y="588" text-anchor="middle" fill="#5d6d7e" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">预留/未完工程</text>

  <!-- 连接线 -->
  <line x1="400" y1="630" x2="400" y2="665" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <!-- ====== 阶段四：三级复核 ====== -->
  <rect x="50" y="670" width="700" height="130" rx="6" fill="none" stroke="#1a5276" stroke-width="2"/>
  <text x="70" y="695" fill="#1a5276" font-size="16" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阶段四：三级复核与报告出具</text>

  <rect x="70" y="715" width="170" height="50" rx="4" fill="#fdebd0" stroke="#e67e22" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="155" y="738" text-anchor="middle" fill="#a04000" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">一级复核</text>
  <text x="155" y="756" text-anchor="middle" fill="#7d6608" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">项目负责人</text>

  <line x1="240" y1="740" x2="285" y2="740" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="290" y="715" width="170" height="50" rx="4" fill="#d5f5e3" stroke="#27ae60" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="375" y="738" text-anchor="middle" fill="#1e8449" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">二级复核</text>
  <text x="375" y="756" text-anchor="middle" fill="#145a32" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">部门负责人</text>

  <line x1="460" y1="740" x2="505" y2="740" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="510" y="715" width="210" height="50" rx="4" fill="#d6eaf8" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="615" y="738" text-anchor="middle" fill="#1a5276" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">三级复核（终审）</text>
  <text x="615" y="756" text-anchor="middle" fill="#1b4f72" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">主任会计师/质控部</text>

  <!-- 连接线 -->
  <line x1="400" y1="800" x2="400" y2="835" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <!-- ====== 阶段五：成果交付 ====== -->
  <rect x="50" y="840" width="700" height="120" rx="6" fill="none" stroke="#1a5276" stroke-width="2"/>
  <text x="70" y="865" fill="#1a5276" font-size="16" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阶段五：成果交付与后续服务</text>

  <rect x="100" y="880" width="180" height="50" rx="4" fill="url(#outputGrad)" stroke="#1abc9c" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="190" y="903" text-anchor="middle" fill="#0e6655" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">出具审核报告</text>
  <text x="190" y="921" text-anchor="middle" fill="#117a65" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">征求意见→正式出具</text>

  <line x1="280" y1="905" x2="330" y2="905" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="335" y="880" width="180" height="50" rx="4" fill="url(#outputGrad)" stroke="#1abc9c" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="425" y="903" text-anchor="middle" fill="#0e6655" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">归档与备案</text>
  <text x="425" y="921" text-anchor="middle" fill="#117a65" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">电子+纸质归档</text>

  <line x1="515" y1="905" x2="565" y2="905" stroke="#2c3e50" stroke-width="2" marker-end="url(#arrowhead)"/>

  <rect x="570" y="880" width="150" height="50" rx="4" fill="url(#outputGrad)" stroke="#1abc9c" stroke-width="1.5" filter="url(#shadow)"/>
  <text x="645" y="903" text-anchor="middle" fill="#0e6655" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">跟踪回访</text>
  <text x="645" y="921" text-anchor="middle" fill="#117a65" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">整改落实情况</text>

  <!-- 底部标注 -->
  <text x="400" y="1000" text-anchor="middle" fill="#95a5a6" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">注：各阶段均设置质量控制节点，确保审核质量符合《基本建设财务规则》及川发改价格〔2013〕901号等要求</text>

  <!-- 左侧时间轴 -->
  <line x1="30" y1="110" x2="30" y2="930" stroke="#bdc3c7" stroke-width="3" stroke-dasharray="8,4"/>
  <circle cx="30" cy="155" r="8" fill="#1a5276"/>
  <circle cx="30" cy="325" r="8" fill="#1a5276"/>
  <circle cx="30" cy="530" r="8" fill="#1a5276"/>
  <circle cx="30" cy="740" r="8" fill="#1a5276"/>
  <circle cx="30" cy="905" r="8" fill="#1a5276"/>
</svg>'''
    path = os.path.join(SVG_DIR, "flowchart.svg")
    with open(path, 'w', encoding='utf-8') as f:
        f.write(svg)
    return path


def generate_quality_arch_svg():
    """质量保证体系架构图"""
    svg = '''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 750" width="800" height="750">
  <defs>
    <linearGradient id="qc1" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" style="stop-color:#1a5276;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#2e86c1;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="qc2" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#d4efdf;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#a9dfbf;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="qc3" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#d6eaf8;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#aed6f1;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="qc4" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#fdebd0;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#fad7a1;stop-opacity:1"/>
    </linearGradient>
    <linearGradient id="qc5" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#f5b7b1;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#f1948a;stop-opacity:1"/>
    </linearGradient>
    <filter id="shadow2" x="-3%" y="-3%" width="106%" height="106%">
      <feDropShadow dx="2" dy="2" stdDeviation="2" flood-opacity="0.12"/>
    </filter>
  </defs>

  <!-- 标题 -->
  <rect x="50" y="15" width="700" height="45" rx="8" fill="url(#qc1)" filter="url(#shadow2)"/>
  <text x="400" y="44" text-anchor="middle" fill="white" font-size="18" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">竣工财务决算审核质量保证体系架构图</text>

  <!-- 顶层：质量管理目标 -->
  <rect x="200" y="80" width="400" height="40" rx="20" fill="url(#qc1)" filter="url(#shadow2)"/>
  <text x="400" y="106" text-anchor="middle" fill="white" font-size="15" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">质量管理目标：审核结论客观公正、数据准确无误</text>

  <!-- 连接线 -->
  <line x1="400" y1="120" x2="400" y2="145" stroke="#1a5276" stroke-width="3"/>
  <line x1="120" y1="145" x2="680" y2="145" stroke="#1a5276" stroke-width="2.5"/>
  <line x1="120" y1="145" x2="120" y2="170" stroke="#1a5276" stroke-width="2.5"/>
  <line x1="300" y1="145" x2="300" y2="170" stroke="#1a5276" stroke-width="2.5"/>
  <line x1="500" y1="145" x2="500" y2="170" stroke="#1a5276" stroke-width="2.5"/>
  <line x1="680" y1="145" x2="680" y2="170" stroke="#1a5276" stroke-width="2.5"/>

  <!-- 第二层：四大支柱 -->
  <rect x="40" y="175" width="160" height="80" rx="6" fill="url(#qc2)" stroke="#27ae60" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="120" y="205" text-anchor="middle" fill="#1e8449" font-size="14" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">制度保障</text>
  <text x="120" y="225" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">执业质量控制制度</text>
  <text x="120" y="242" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">三级复核制度</text>

  <rect x="230" y="175" width="160" height="80" rx="6" fill="url(#qc3)" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="310" y="205" text-anchor="middle" fill="#1a5276" font-size="14" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">人员保障</text>
  <text x="310" y="225" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">CPA持证专业团队</text>
  <text x="310" y="242" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">持续专业培训</text>

  <rect x="420" y="175" width="160" height="80" rx="6" fill="url(#qc4)" stroke="#e67e22" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="500" y="205" text-anchor="middle" fill="#a04000" font-size="14" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">流程保障</text>
  <text x="500" y="225" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">标准化工作底稿</text>
  <text x="500" y="242" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">关键节点控制</text>

  <rect x="610" y="175" width="160" height="80" rx="6" fill="url(#qc5)" stroke="#c0392b" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="690" y="205" text-anchor="middle" fill="#78281f" font-size="14" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">风险管控</text>
  <text x="690" y="225" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">风险识别与评估</text>
  <text x="690" y="242" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">廉洁保密纪律</text>

  <!-- 连接线 -->
  <line x1="120" y1="255" x2="120" y2="300" stroke="#27ae60" stroke-width="2"/>
  <line x1="310" y1="255" x2="310" y2="300" stroke="#2980b9" stroke-width="2"/>
  <line x1="500" y1="255" x2="500" y2="300" stroke="#e67e22" stroke-width="2"/>
  <line x1="690" y1="255" x2="690" y2="300" stroke="#c0392b" stroke-width="2"/>

  <!-- 第三层：具体措施 -->
  <rect x="30" y="305" width="180" height="175" rx="5" fill="white" stroke="#27ae60" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="120" y="325" text-anchor="middle" fill="#1e8449" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">制度保障措施</text>
  <line x1="50" y1="332" x2="190" y2="332" stroke="#27ae60" stroke-width="1"/>
  <text x="40" y="352" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 《执业质量控制制度》</text>
  <text x="40" y="370" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 三级复核实施细则</text>
  <text x="40" y="388" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 审计工作底稿规范</text>
  <text x="40" y="406" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 档案管理制度</text>
  <text x="40" y="424" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 廉政保密制度</text>
  <text x="40" y="442" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 重大问题请示报告制度</text>
  <text x="40" y="460" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 责任追究制度</text>

  <rect x="220" y="305" width="180" height="175" rx="5" fill="white" stroke="#2980b9" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="310" y="325" text-anchor="middle" fill="#1a5276" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">人员保障措施</text>
  <line x1="240" y1="332" x2="380" y2="332" stroke="#2980b9" stroke-width="1"/>
  <text x="230" y="352" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• CPA持证成员≥6人</text>
  <text x="230" y="370" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 高级职称项目负责人</text>
  <text x="230" y="388" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 工程+财务复合背景</text>
  <text x="230" y="406" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 年度专业培训≥40学时</text>
  <text x="230" y="424" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 政府审计项目经验</text>
  <text x="230" y="442" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 人员变更审批制度</text>
  <text x="230" y="460" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 胜任能力动态评估</text>

  <rect x="410" y="305" width="180" height="175" rx="5" fill="white" stroke="#e67e22" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="500" y="325" text-anchor="middle" fill="#a04000" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">流程保障措施</text>
  <line x1="430" y1="332" x2="570" y2="332" stroke="#e67e22" stroke-width="1"/>
  <text x="420" y="352" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 标准化工作底稿模板</text>
  <text x="420" y="370" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 关键节点质量检查表</text>
  <text x="420" y="388" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 审核日志全程记录</text>
  <text x="420" y="406" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 双人交叉复核机制</text>
  <text x="420" y="424" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 时限预警提醒机制</text>
  <text x="420" y="442" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 与采购人沟通反馈</text>
  <text x="420" y="460" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 质量考核与奖惩</text>

  <rect x="600" y="305" width="180" height="175" rx="5" fill="white" stroke="#c0392b" stroke-width="1.5" filter="url(#shadow2)"/>
  <text x="690" y="325" text-anchor="middle" fill="#78281f" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">风险管控措施</text>
  <line x1="620" y1="332" x2="760" y2="332" stroke="#c0392b" stroke-width="1"/>
  <text x="610" y="352" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 项目风险分级评估</text>
  <text x="610" y="370" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 重大风险预警机制</text>
  <text x="610" y="388" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 廉洁从业承诺书</text>
  <text x="610" y="406" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 保密协议签署</text>
  <text x="610" y="424" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 利益冲突申报</text>
  <text x="610" y="442" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 重大事项报告</text>
  <text x="610" y="460" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">• 质量事故应急预案</text>

  <!-- 连接线汇聚 -->
  <line x1="120" y1="480" x2="120" y2="520" stroke="#27ae60" stroke-width="1.5"/>
  <line x1="310" y1="480" x2="310" y2="520" stroke="#2980b9" stroke-width="1.5"/>
  <line x1="500" y1="480" x2="500" y2="520" stroke="#e67e22" stroke-width="1.5"/>
  <line x1="690" y1="480" x2="690" y2="520" stroke="#c0392b" stroke-width="1.5"/>
  <line x1="120" y1="520" x2="690" y2="520" stroke="#1a5276" stroke-width="2.5"/>

  <line x1="400" y1="520" x2="400" y2="550" stroke="#1a5276" stroke-width="3"/>

  <!-- 底层：持续改进 -->
  <rect x="150" y="555" width="500" height="95" rx="8" fill="url(#qc1)" filter="url(#shadow2)"/>
  <text x="400" y="585" text-anchor="middle" fill="#f9e79f" font-size="15" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">质量持续改进机制</text>
  <text x="400" y="610" text-anchor="middle" fill="white" font-size="12" font-family="Microsoft YaHei, SimHei, sans-serif">项目后评价 → 经验总结 → 制度修订 → 培训提升 → 质量考核 → 持续改进</text>
  <text x="400" y="635" text-anchor="middle" fill="#d5dbdb" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">PDCA循环：Plan（策划）→ Do（实施）→ Check（检查）→ Act（改进）</text>

  <!-- 底部 -->
  <text x="400" y="690" text-anchor="middle" fill="#7f8c8d" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">遵循标准：中国注册会计师审计准则、基本建设财务规则、川发改价格〔2013〕901号</text>
  <text x="400" y="715" text-anchor="middle" fill="#7f8c8d" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">《会计师事务所质量管理准则第5101号——业务质量管理》（2024年修订）</text>
</svg>'''
    path = os.path.join(SVG_DIR, "quality_arch.svg")
    with open(path, 'w', encoding='utf-8') as f:
        f.write(svg)
    return path


def generate_org_chart_svg():
    """项目组织架构图"""
    svg = '''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 420" width="800" height="420">
  <defs>
    <linearGradient id="topGrad" x1="0%" y1="0%" x2="100%" y2="0%">
      <stop offset="0%" style="stop-color:#1a5276;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#2e86c1;stop-opacity:1"/>
    </linearGradient>
    <filter id="sh" x="-5%" y="-5%" width="110%" height="110%">
      <feDropShadow dx="2" dy="2" stdDeviation="3" flood-opacity="0.15"/>
    </filter>
  </defs>

  <!-- 标题 -->
  <rect x="50" y="10" width="700" height="40" rx="6" fill="url(#topGrad)" filter="url(#sh)"/>
  <text x="400" y="37" text-anchor="middle" fill="white" font-size="16" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">阿坝县竣工财务决算审核项目组织架构图</text>

  <!-- 第一层：项目总负责人 -->
  <rect x="275" y="70" width="250" height="50" rx="8" fill="url(#topGrad)" filter="url(#sh)"/>
  <text x="400" y="93" text-anchor="middle" fill="white" font-size="14" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">项目总负责人（主任会计师）</text>
  <text x="400" y="110" text-anchor="middle" fill="#d5dbdb" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">全面负责、三级复核终审</text>

  <line x1="400" y1="120" x2="400" y2="145" stroke="#1a5276" stroke-width="2.5"/>
  <line x1="130" y1="145" x2="670" y2="145" stroke="#1a5276" stroke-width="2"/>
  <line x1="130" y1="145" x2="130" y2="165" stroke="#1a5276" stroke-width="2"/>
  <line x1="400" y1="145" x2="400" y2="165" stroke="#1a5276" stroke-width="2"/>
  <line x1="670" y1="145" x2="670" y2="165" stroke="#1a5276" stroke-width="2"/>

  <!-- 第二层 -->
  <rect x="40" y="170" width="180" height="65" rx="6" fill="#d4efdf" stroke="#27ae60" stroke-width="1.5" filter="url(#sh)"/>
  <text x="130" y="195" text-anchor="middle" fill="#1e8449" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">项目负责人</text>
  <text x="130" y="215" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">CPA+高级职称</text>
  <text x="130" y="228" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">一级复核、现场统筹</text>

  <rect x="310" y="170" width="180" height="65" rx="6" fill="#d6eaf8" stroke="#2980b9" stroke-width="1.5" filter="url(#sh)"/>
  <text x="400" y="195" text-anchor="middle" fill="#1a5276" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">质量复核负责人</text>
  <text x="400" y="215" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">质控部门经理</text>
  <text x="400" y="228" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">二级复核、质量监控</text>

  <rect x="580" y="170" width="180" height="65" rx="6" fill="#fdebd0" stroke="#e67e22" stroke-width="1.5" filter="url(#sh)"/>
  <text x="670" y="195" text-anchor="middle" fill="#a04000" font-size="13" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">后勤保障负责人</text>
  <text x="670" y="215" text-anchor="middle" fill="#2c3e50" font-size="10" font-family="Microsoft YaHei, SimHei, sans-serif">综合管理部</text>
  <text x="670" y="228" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">档案/联络/交通保障</text>

  <!-- 连接线 -->
  <line x1="130" y1="235" x2="130" y2="265" stroke="#27ae60" stroke-width="1.5"/>
  <line x1="400" y1="235" x2="400" y2="265" stroke="#2980b9" stroke-width="1.5"/>

  <!-- 第三层 -->
  <line x1="40" y1="265" x2="220" y2="265" stroke="#27ae60" stroke-width="1"/>

  <rect x="15" y="270" width="115" height="55" rx="4" fill="white" stroke="#27ae60" stroke-width="1" filter="url(#sh)"/>
  <text x="72" y="293" text-anchor="middle" fill="#1e8449" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">CPA审核员</text>
  <text x="72" y="312" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">工程财务审核</text>

  <rect x="145" y="270" width="120" height="55" rx="4" fill="white" stroke="#27ae60" stroke-width="1" filter="url(#sh)"/>
  <text x="205" y="293" text-anchor="middle" fill="#1e8449" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">高级审计员</text>
  <text x="205" y="312" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">待摊投资/资金审核</text>

  <rect x="280" y="270" width="115" height="55" rx="4" fill="white" stroke="#27ae60" stroke-width="1" filter="url(#sh)"/>
  <text x="337" y="293" text-anchor="middle" fill="#1e8449" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">审计助理</text>
  <text x="337" y="312" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">资料整理/底稿编制</text>

  <line x1="400" y1="235" x2="400" y2="265" stroke="#2980b9" stroke-width="1.5"/>
  <rect x="325" y="270" width="150" height="55" rx="4" fill="white" stroke="#2980b9" stroke-width="1" filter="url(#sh)"/>
  <text x="400" y="293" text-anchor="middle" fill="#1a5276" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">质控复核组</text>
  <text x="400" y="312" text-anchor="middle" fill="#2c3e50" font-size="9" font-family="Microsoft YaHei, SimHei, sans-serif">独立复核/底稿检查</text>

  <text x="400" y="365" text-anchor="middle" fill="#7f8c8d" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">项目组实行组长负责制，各岗位职责明确、相互配合、相互制约</text>
  <text x="400" y="385" text-anchor="middle" fill="#7f8c8d" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">人员变动需经采购人书面同意，接替人员资质不低于原人员</text>
</svg>'''
    path = os.path.join(SVG_DIR, "org_chart.svg")
    with open(path, 'w', encoding='utf-8') as f:
        f.write(svg)
    return path


def generate_risk_svg():
    """重点难点风险矩阵图"""
    svg = '''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 650" width="800" height="650">
  <defs>
    <linearGradient id="bg" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#fef9e7;stop-opacity:1"/>
      <stop offset="100%" style="stop-color:#f9e79f;stop-opacity:1"/>
    </linearGradient>
    <filter id="sh3"><feDropShadow dx="1" dy="1" stdDeviation="2" flood-opacity="0.1"/></filter>
  </defs>

  <rect x="0" y="0" width="800" height="650" fill="white" rx="10"/>

  <!-- 标题 -->
  <text x="400" y="35" text-anchor="middle" fill="#1a5276" font-size="18" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">竣工财务决算审核重点难点分析矩阵</text>

  <!-- 坐标轴 -->
  <line x1="80" y1="530" x2="750" y2="530" stroke="#2c3e50" stroke-width="2"/>
  <line x1="80" y1="530" x2="80" y2="60" stroke="#2c3e50" stroke-width="2"/>
  <text x="400" y="565" text-anchor="middle" fill="#2c3e50" font-size="12" font-family="Microsoft YaHei, SimHei, sans-serif">风险发生概率 →</text>
  <text x="25" y="300" text-anchor="middle" fill="#2c3e50" font-size="12" font-family="Microsoft YaHei, SimHei, sans-serif" transform="rotate(-90, 25, 300)">影响程度 →</text>

  <!-- 象限标签 -->
  <text x="250" y="210" text-anchor="middle" fill="#27ae60" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">低概率-高影响（重点关注）</text>
  <text x="600" y="210" text-anchor="middle" fill="#c0392b" font-size="12" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">高概率-高影响（优先处置）</text>
  <text x="250" y="500" text-anchor="middle" fill="#7f8c8d" font-size="12" font-family="Microsoft YaHei, SimHei, sans-serif">低概率-低影响（常规管理）</text>
  <text x="600" y="500" text-anchor="middle" fill="#e67e22" font-size="12" font-family="Microsoft YaHei, SimHei, sans-serif">高概率-低影响（持续监控）</text>

  <!-- 象限分隔线 -->
  <line x1="415" y1="60" x2="415" y2="530" stroke="#bdc3c7" stroke-width="1" stroke-dasharray="5,5"/>
  <line x1="80" y1="295" x2="750" y2="295" stroke="#bdc3c7" stroke-width="1" stroke-dasharray="5,5"/>

  <!-- 风险气泡 -->
  <circle cx="620" cy="150" r="45" fill="#e74c3c" opacity="0.7" filter="url(#sh3)"/>
  <text x="620" y="145" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">资金</text>
  <text x="620" y="162" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">管理</text>

  <circle cx="580" cy="200" r="40" fill="#e74c3c" opacity="0.6" filter="url(#sh3)"/>
  <text x="580" y="196" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">工程</text>
  <text x="580" y="213" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">变更</text>

  <circle cx="540" cy="120" r="35" fill="#e67e22" opacity="0.7" filter="url(#sh3)"/>
  <text x="540" y="116" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">待摊</text>
  <text x="540" y="133" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">投资</text>

  <circle cx="300" cy="150" r="38" fill="#2980b9" opacity="0.65" filter="url(#sh3)"/>
  <text x="300" y="146" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">资产</text>
  <text x="300" y="163" text-anchor="middle" fill="white" font-size="11" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">交付</text>

  <circle cx="200" cy="180" r="32" fill="#2980b9" opacity="0.55" filter="url(#sh3)"/>
  <text x="200" y="177" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">尾工</text>
  <text x="200" y="192" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">工程</text>

  <circle cx="560" cy="400" r="32" fill="#27ae60" opacity="0.6" filter="url(#sh3)"/>
  <text x="560" y="397" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">资料</text>
  <text x="560" y="412" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">完整性</text>

  <circle cx="640" cy="370" r="30" fill="#27ae60" opacity="0.5" filter="url(#sh3)"/>
  <text x="640" y="367" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">时限</text>
  <text x="640" y="382" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">压力</text>

  <circle cx="180" cy="420" r="30" fill="#95a5a6" opacity="0.5" filter="url(#sh3)"/>
  <text x="180" y="425" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">天气</text>

  <circle cx="300" cy="380" r="28" fill="#95a5a6" opacity="0.5" filter="url(#sh3)"/>
  <text x="300" y="385" text-anchor="middle" fill="white" font-size="10" font-weight="bold" font-family="Microsoft YaHei, SimHei, sans-serif">交通</text>

  <!-- 图例 -->
  <rect x="80" y="580" width="15" height="15" fill="#e74c3c" opacity="0.7"/>
  <text x="100" y="593" fill="#2c3e50" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">高风险</text>
  <rect x="180" y="580" width="15" height="15" fill="#e67e22" opacity="0.7"/>
  <text x="200" y="593" fill="#2c3e50" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">中高风险</text>
  <rect x="300" y="580" width="15" height="15" fill="#2980b9" opacity="0.65"/>
  <text x="320" y="593" fill="#2c3e50" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">中等风险</text>
  <rect x="420" y="580" width="15" height="15" fill="#27ae60" opacity="0.6"/>
  <text x="440" y="593" fill="#2c3e50" font-size="11" font-family="Microsoft YaHei, SimHei, sans-serif">低风险</text>
</svg>'''
    path = os.path.join(SVG_DIR, "risk_matrix.svg")
    with open(path, 'w', encoding='utf-8') as f:
        f.write(svg)
    return path


# ============================================================
# Word 文档生成
# ============================================================

def set_cell_shading(cell, color):
    """设置单元格底色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    tcPr.append(shading_elm)


def add_table_row(table, cells_data, bold=False, color=None, font_size=10):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(font_size)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if bold:
            run.bold = True
        if color:
            set_cell_shading(cell, color)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return row


def add_heading_styled(doc, text, level=1):
    """添加样式标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def add_para(doc, text, bold=False, indent=False, font_size=12, alignment=None, font_name='宋体'):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True
    return p


def add_image(doc, png_path, width_inches=6.0):
    """插入PNG图片到文档"""
    if os.path.exists(png_path):
        doc.add_picture(png_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def create_document():
    """创建完整的Word文档"""
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ================================================================
    # 封面
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('2026年阿坝县财政局采购政府性投资项目')
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title_p2.add_run('竣工财务决算审核服务方案')
    run2.font.size = Pt(22)
    run2.font.name = '黑体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run2.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ('项目编号：', 'K5132312026000002'),
        ('征集人：', '阿坝县财政局'),
        ('项目地点：', '阿坝藏族羌族自治州阿坝县'),
        ('服务类型：', '工程竣工财务决算审计'),
        ('编制日期：', '2026年5月'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label)
        r1.font.size = Pt(14)
        r1.font.name = '宋体'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r2 = p.add_run(value)
        r2.font.size = Pt(14)
        r2.font.name = '宋体'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r2.bold = True

    doc.add_page_break()

    # ================================================================
    # 目录页（简要）
    # ================================================================
    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '一、审核服务目标',
        '二、审核服务范围',
        '三、审核依据、内容及要点',
        '四、审核工作流程',
        '五、审核过程重点和难点分析',
        '六、审核质量保证措施',
        '七、审核后续支持服务',
    ]
    for item in toc_items:
        add_para(doc, item, font_size=14)
    doc.add_page_break()

    # ================================================================
    # 一、审核服务目标
    # ================================================================
    add_heading_styled(doc, '一、审核服务目标', level=1)

    add_para(doc, '1.1 总体目标', bold=True, font_size=14)
    add_para(doc, '本项目旨在为阿坝县财政局提供专业、规范、高效的政府性投资项目竣工财务决算审核服务。通过对阿坝县范围内政府性投资项目的竣工财务决算进行全面、系统的审核，客观、公正地反映项目建设成果和财务状况，核实项目竣工财务决算的真实性、合法性和完整性，为财政部门批复项目竣工财务决算、办理资产交付手续提供可靠的依据，切实保障财政资金安全和使用效益。', indent=True)

    add_para(doc, '1.2 具体目标', bold=True, font_size=14)
    goals = [
        ('真实性审核目标', '核实项目竣工财务决算报告所反映的各项数据是否真实、准确，确认项目建设成本、投资完成额、资产交付价值等关键数据的真实性，杜绝虚列支出、虚报投资完成额等违规行为。'),
        ('合法性审核目标', '审查项目建设全过程的资金收支活动是否符合国家相关法律法规、财经纪律和基本建设财务管理规定，确认各项支出的合规性，重点审查工程款支付、征地拆迁补偿、设备材料采购等环节的合法合规性。'),
        ('完整性审核目标', '审核项目竣工财务决算是否涵盖了项目建设全过程的全部经济业务，各项资金来源和使用是否完整反映，是否存在应纳入决算而未纳入的事项，确保决算信息全面完整。'),
        ('效益性审核目标', '评价项目建设投资的经济效益和社会效益，分析项目概算执行情况，查找超概算或节余的原因，为财政部门加强政府投资项目管理、提高财政资金使用效益提供决策参考。'),
        ('规范性审核目标', '确保竣工财务决算的编制符合《基本建设财务规则》（财政部令第81号）、《基本建设项目竣工财务决算管理暂行办法》（财建〔2016〕503号）以及四川省和阿坝州相关管理规定的要求，决算报表格式规范、内容完整、勾稽关系正确。'),
    ]
    for title, desc in goals:
        add_para(doc, f'（{goals.index((title, desc))+1}）{title}', bold=True)
        add_para(doc, desc, indent=True)

    add_para(doc, '1.3 服务承诺', bold=True, font_size=14)
    promises = [
        '严格遵守国家法律、法规和行业规范，恪守独立、客观、公正的执业原则，确保审核结论经得起检验。',
        '接到阿坝县财政局通知后，1个工作日内到达采购人办公地点，及时提交人员配置方案。',
        '严格按照合同约定的审核时限要求完成审核工作并出具审核报告（500万元以下5个工作日，500-1000万元10个工作日，1000-5000万元15个工作日，5000万元以上20个工作日）。',
        '不以任何形式将受托审核任务分包、转包或支解给其他中介机构。',
        '严格遵守廉政纪律和保密规定，未经采购人书面许可，不向任何第三方泄露项目信息。',
        '派出完成项目的团队成员均为本单位注册在职人员，且与响应文件承诺的参审人员一致。'
    ]
    for p_text in promises:
        add_para(doc, f'（{promises.index(p_text)+1}）{p_text}', indent=True)

    # ================================================================
    # 二、审核服务范围
    # ================================================================
    add_heading_styled(doc, '二、审核服务范围', level=1)

    add_para(doc, '2.1 服务地域范围', bold=True, font_size=14)
    add_para(doc, '本项目服务范围为阿坝藏族羌族自治州阿坝县行政区域内，由阿坝县财政局指定的政府性投资项目竣工财务决算审核。', indent=True)

    add_para(doc, '2.2 项目类型范围', bold=True, font_size=14)
    project_types = [
        '阿坝县政府投资的新建、改建、扩建、技术改造等各类基本建设项目；',
        '阿坝县政府投资的民生工程、基础设施建设项目、公共服务设施项目；',
        '上级财政转移支付资金安排的各类政府性投资项目；',
        '地方政府专项债券资金安排的建设项目；',
        '其他由阿坝县财政局认定需要纳入竣工财务决算审核范围的政府性投资项目。'
    ]
    for pt in project_types:
        add_para(doc, f'（{project_types.index(pt)+1}）{pt}', indent=True)

    add_para(doc, '2.3 审核内容范围', bold=True, font_size=14)
    add_para(doc, '根据项目竣工财务决算审核的业务特点和相关法规要求，审核内容涵盖以下方面：', indent=True)

    audit_scopes = [
        ('工程价款结算审核', '审核工程竣工结算是否经相关部门审定，结算依据是否充分，结算价款是否准确，工程变更、签证及索赔是否合规合理。'),
        ('资金来源及到位情况审核', '审核项目建设资金来源是否合规，中央、省、州、县各级财政资金到位情况，项目单位自筹资金到位情况，是否存在资金缺口或资金来源不明确的情形。'),
        ('资金使用及支出审核', '审核项目建设资金的使用是否符合预算和概算范围，各项支出是否真实、合规，是否存在挤占、挪用、截留、虚列建设资金等行为。'),
        ('建筑安装工程投资审核', '审核建筑安装工程投资完成额的真实性、准确性，核实工程计量和计价依据，审查施工单位是否按照规定编制竣工结算。'),
        ('设备投资审核', '审核设备购置的真实性和合理性，核实设备数量、规格型号、单价与合同的符合性，审查设备采购程序是否规范。'),
        ('待摊投资审核', '审核建设单位管理费、勘察设计费、监理费、研究试验费、土地使用费、临时设施费等各项待摊投资的合规性和合理性，审查费用支出的标准和依据。'),
        ('其他投资审核', '审核为项目配套的专用设施、生产准备费、办公及生活家具购置费等是否真实、合理。'),
        ('资产交付使用审核', '审核交付使用资产的真实性、准确性，核实资产分类是否正确，资产价值的计算是否合理，交付手续是否完备。'),
        ('尾工工程审核', '审核尾工工程的内容、工程量及预留费用的合理性，审查是否存在将已完工程人为列入尾工以规避审核的情况。'),
        ('概算执行及投资效益分析', '对比分析项目概算与实际执行情况，查找超概原因，评价项目投资效益。'),
    ]
    for title, desc in audit_scopes:
        add_para(doc, f'（{audit_scopes.index((title, desc))+1}）{title}', bold=True)
        add_para(doc, desc, indent=True)

    add_para(doc, '2.4 服务期限范围', bold=True, font_size=14)
    add_para(doc, '本项目服务期限为2年，合同一年一签。在框架协议有效期内，按照阿坝县财政局的委托和项目分配，持续提供竣工财务决算审核服务。如因新一期框架协议公开征集过程中出现废标、质疑投诉等特殊情形，导致新一期框架协议不能及时订立的，征集人可适当延长原框架协议有效期。', indent=True)

    # ================================================================
    # 三、审核依据、内容及要点
    # ================================================================
    add_heading_styled(doc, '三、审核依据、内容及要点', level=1)

    add_para(doc, '3.1 法律法规依据', bold=True, font_size=14)

    # 表格：审核依据
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table1.rows[0]
    hdr.cells[0].text = '类别'
    hdr.cells[1].text = '依据名称'
    hdr.cells[2].text = '文号/发布机关'
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, '1a5276')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    laws = [
        ('法律', '中华人民共和国审计法', '全国人大常委会'),
        ('法律', '中华人民共和国预算法', '全国人大常委会'),
        ('法律', '中华人民共和国会计法', '全国人大常委会'),
        ('法律', '中华人民共和国政府采购法', '全国人大常委会'),
        ('法律', '中华人民共和国招标投标法', '全国人大常委会'),
        ('行政法规', '中华人民共和国审计法实施条例', '国务院'),
        ('行政法规', '中华人民共和国预算法实施条例', '国务院'),
        ('部门规章', '基本建设财务规则', '财政部令第81号'),
        ('部门规章', '基本建设项目竣工财务决算管理暂行办法', '财建〔2016〕503号'),
        ('部门规章', '基本建设项目建设成本管理规定', '财建〔2016〕504号'),
        ('部门规章', '会计师事务所执业许可和监督管理办法', '财政部令第97号'),
        ('地方规章', '四川省基本建设项目竣工财务决算审核规程', '四川省财政厅'),
        ('收费标准', '四川省会计师事务所服务收费管理办法', '川发改价格〔2013〕901号'),
        ('执业准则', '中国注册会计师审计准则', '中国注册会计师协会'),
        ('执业准则', '中国注册会计师职业道德守则', '中国注册会计师协会'),
        ('执业准则', '会计师事务所质量管理准则第5101号', '中国注册会计师协会'),
    ]
    for item in laws:
        add_table_row(table1, item, font_size=9)

    doc.add_paragraph()
    add_para(doc, '3.2 项目具体依据', bold=True, font_size=14)
    proj_basis = [
        '本项目征集文件（项目编号：K5132312026000002）及入围供应商响应文件；',
        '阿坝县财政局与入围供应商签订的框架协议及具体项目委托合同；',
        '项目立项批复文件（项目建议书批复、可行性研究报告批复）；',
        '项目初步设计及概算批复文件；',
        '项目招标文件、中标通知书及施工合同、监理合同、勘察设计合同等各类合同文件；',
        '工程竣工结算审核报告及审计报告；',
        '项目财务核算资料（会计凭证、账簿、报表等）；',
        '项目竣工验收报告及相关验收资料；',
        '其他与项目竣工财务决算相关的文件资料。',
    ]
    for b in proj_basis:
        add_para(doc, f'（{proj_basis.index(b)+1}）{b}', indent=True)

    add_para(doc, '3.3 审核内容及要点', bold=True, font_size=14)
    add_para(doc, '根据《基本建设财务规则》和《基本建设项目竣工财务决算管理暂行办法》的要求，竣工财务决算审核的核心内容包括以下方面，每项内容均需按照相应审核要点逐项落实：', indent=True)

    # 核心审核要点表
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0]
    for i, text in enumerate(['序号', '审核内容', '审核要点']):
        hdr2.cells[i].text = text
        for p in hdr2.cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(hdr2.cells[i], '1a5276')
        for p in hdr2.cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    audit_points = [
        ('1', '竣工财务决算\n报表审核', '①决算报表的完整性、规范性；②报表数据与账簿的一致性；③表内及表间勾稽关系的正确性；④数据计算准确性；⑤报表附注的充分性和适当性'),
        ('2', '项目立项及\n审批程序审核', '①项目立项依据是否充分；②可研、初设审批程序是否合规；③是否存"边勘测边设计边施工"等问题；④项目调整是否按规定报批'),
        ('3', '工程概算执行\n情况审核', '①实际投资与概算的对比分析；②超概原因分析及审批手续；③概算调整是否合规；④是否存在擅自扩大建设规模、提高建设标准等情况'),
        ('4', '资金来源及\n到位审核', '①资金来源渠道是否合规合法；②各级财政资金到位时间和金额；③自筹资金落实情况；④资金缺口分析及原因说明'),
        ('5', '建筑安装工程\n投资审核', '①工程结算是否经审定；②工程量及单价核算准确性；③工程变更、签证手续是否完备；④工程款支付与合同约定的符合性'),
        ('6', '设备投资审核', '①设备采购程序的合规性；②设备数量、规格型号与合同一致性；③设备价格合理性；④不需安装设备及工器具的移交情况'),
        ('7', '待摊投资审核', '①建设单位管理费是否超支；②各项费用的计提标准和依据；③费用支出与项目的相关性；④应分摊费用的分摊方法是否合理'),
        ('8', '其他投资审核', '①建设用地费是否合规；②配套专用设施投资核实；③生产准备费及办公家具购置费的真实性、合理性'),
        ('9', '资产交付\n使用审核', '①交付资产清单的完整性；②资产分类的准确性；③资产价值的合理性；④交付手续和权属证明的完备性'),
        ('10', '尾工工程\n及结余资金审核', '①尾工工程内容和工程量的合理性；②尾工工程预留费用是否适当；③结余资金的真实性及处置合规性；④债权债务清理情况'),
    ]
    for item in audit_points:
        row = add_table_row(table2, item, font_size=9)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    add_para(doc, '3.4 竣工财务决算审核时限要求', bold=True, font_size=14)
    add_para(doc, '根据本项目征集文件要求，竣工财务决算审核时限（自收到完整资料之日起计算）如下：', indent=True)

    table_deadline = doc.add_table(rows=1, cols=3)
    table_deadline.style = 'Table Grid'
    table_deadline.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_d = table_deadline.rows[0]
    for i, text in enumerate(['序号', '项目报送金额', '审核时限']):
        hdr_d.cells[i].text = text
        for p in hdr_d.cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(hdr_d.cells[i], '1a5276')
        for p in hdr_d.cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    deadlines = [
        ('1', '500万元以下（含500万元）', '5个工作日'),
        ('2', '500万元—1000万元（含1000万元）', '10个工作日'),
        ('3', '1000万元—5000万元（含5000万元）', '15个工作日'),
        ('4', '5000万元以上', '20个工作日'),
    ]
    for d in deadlines:
        add_table_row(table_deadline, d, font_size=10)

    doc.add_page_break()

    # ================================================================
    # 四、审核工作流程
    # ================================================================
    add_heading_styled(doc, '四、审核工作流程', level=1)

    add_para(doc, '4.1 流程概述', bold=True, font_size=14)
    add_para(doc, '为确保阿坝县财政局政府性投资项目竣工财务决算审核工作的规范性和高效性，本项目建立了一套完整的"五阶段审核工作流程"，涵盖从接受委托到成果交付的全过程。各阶段均设置明确的质量控制节点和时限要求，确保审核工作有序推进、质量可控。', indent=True)

    add_para(doc, '4.2 审核工作流程图', bold=True, font_size=14)
    add_image(doc, os.path.join(SVG_DIR, 'flowchart.png'), width_inches=5.8)

    add_para(doc, '图1：阿坝县工程竣工财务决算审核工作流程图', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

    add_para(doc, '4.3 各阶段详细说明', bold=True, font_size=14)

    stages = [
        ('阶段一：接受委托与前期准备',
         [
             '签订委托协议：收到阿坝县财政局委托通知后，1个工作日内到达采购人办公地点，与财政局签订具体项目委托协议，明确审核范围、时限、费用及双方权利义务。',
             '组建审核工作组：根据项目规模和特点，按照"专业匹配、经验优先"的原则，组建由注册会计师领衔的审核工作组，明确项目负责人和各级审核人员，制定人员配置方案报采购人备案。',
             '前期调查与资料收集：了解项目基本情况，收集项目立项文件、概算文件、招投标文件、合同文件、工程结算资料、财务核算资料等，编制资料收集清单。',
             '编制审核工作方案：根据项目具体情况，编制包含审核目标、范围、内容、方法、时间安排、人员分工、质量控制措施等内容的工作方案，报采购人审核确认。',
         ]),
        ('阶段二：资料初审与现场核查',
         [
             '资料完整性审查：对收集到的资料进行全面梳理和完整性审查，核实资料的齐全性和真实性，对缺失或不符合要求的资料提出补充清单。',
             '财务资料初步审核：对项目建设期间的会计凭证、账簿、报表等财务资料进行初步审核，核实账务处理的规范性，确认项目独立核算情况。',
             '现场实地踏勘：赴阿坝县项目现场进行实地踏勘，核实工程实际完成情况，确认项目是否按设计文件和合同约定完成建设，了解项目实际运行状况。',
             '沟通与反馈：就资料审核和现场核查发现的问题，及时与项目建设单位和阿坝县财政局进行沟通反馈，形成书面沟通记录。',
         ]),
        ('阶段三：实质性审核',
         [
             '工程价款结算审核：审核工程竣工结算的合理性和准确性，核实工程变更、签证和索赔的真实性和合规性，确认结算审定价款。',
             '资金来源及到位情况审核：逐项核实各类资金来源及实际到位情况，与批复概算进行比对，分析资金到位率及对项目建设的影响。',
             '资金使用及支出审核：审核各项支出的真实性、合规性和合理性，重点关注是否存在虚列支出、挤占挪用等问题，对大额支出进行详查。',
             '建筑安装工程投资审核：核实建筑安装工程投资完成额，审查工程计量和计价依据，确认工程款支付与合同约定的符合性。',
             '设备投资审核：核实设备购置的真实性，审查设备采购程序的合规性，确认设备数量规格与合同的一致性。',
             '待摊投资和其他投资审核：逐项审核各项待摊投资的合规性和合理性，审查费用标准和依据，确认分摊方法的合理性。',
             '资产交付使用审核：核实交付使用资产的真实性和准确性，审查资产分类和计价，确认交付手续的完备性。',
             '尾工工程审核：核实尾工工程内容和工程量，审查预留费用的合理性，确保不存在人为扩大尾工范围以规避审核的情况。',
         ]),
        ('阶段四：三级复核与报告出具',
         [
             '一级复核（项目负责人复核）：项目负责人对审核人员的工作底稿和初步结论进行全面复核，确保审核程序执行到位、取证充分、判断准确，重点复核重大事项和关键数据。',
             '二级复核（部门负责人复核）：部门负责人对一级复核后的工作底稿和审核结论进行独立复核，重点关注审核标准的适用性、重大问题的处理、审核结论的合理性等，对分歧事项进行协调处理。',
             '三级复核（主任会计师/质控部终审）：对审核报告进行全面终审，确保审核结论客观公正、依据充分、表述准确，重点关注重大风险事项、报告的规范性和结论的恰当性。',
             '征求意见与沟通：将审核报告征求意见稿提交被审核项目建设单位征求意见，充分沟通审核发现的问题和审核结论，必要时组织三方（财政局、建设单位、审核机构）会议进行讨论。',
             '出具正式审核报告：在充分征求意见和完成三级复核的基础上，出具正式审核报告。报告应包含项目基本情况、审核依据、审核范围、审核发现、审核结论、建议意见等内容，格式规范、表述清晰。',
         ]),
        ('阶段五：成果交付与后续服务',
         [
             '报告提交与归档：向阿坝县财政局提交正式审核报告（纸质版和电子版），按照档案管理规定完成项目档案的整理和归档。',
             '后续跟踪服务：对审核报告中提出的问题和建议的整改落实情况进行跟踪回访，为阿坝县财政局提供持续的咨询服务。',
             '年度总结与改进：定期对完成的审核项目进行总结分析，归纳共性问题，提出改进建议，不断完善审核方法和质量控制体系。',
         ]),
    ]
    for stage_title, items in stages:
        add_para(doc, stage_title, bold=True)
        for item in items:
            add_para(doc, item, indent=True)

    add_para(doc, '4.4 项目组织架构', bold=True, font_size=14)
    add_image(doc, os.path.join(SVG_DIR, 'org_chart.png'), width_inches=5.8)
    add_para(doc, '图2：阿坝县竣工财务决算审核项目组织架构图', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

    add_para(doc, '4.5 时限管控流程', bold=True, font_size=14)
    add_para(doc, '为确保各项目在征集文件规定的审核时限内完成，建立严格的时限管控机制：', indent=True)
    time_controls = [
        '接收资料当日启动审核时限计时，在审核系统中建立项目时限台账，自动计算审核截止日期。',
        '将审核时限分解为各阶段时间节点：资料初审阶段（占总时限的15%）、现场核查阶段（占总时限的20%）、实质性审核阶段（占总时限的40%）、复核报告阶段（占总时限的25%）。',
        '设置三级时限预警：黄色预警（到期前2个工作日）、橙色预警（到期前1个工作日）、红色预警（到期当日），预警信息自动推送至项目负责人和部门负责人。',
        '建立时限周报制度，每周向采购人报告在审项目进度情况，对可能超期的项目提前说明原因并提出应对措施。',
    ]
    for tc in time_controls:
        add_para(doc, f'（{time_controls.index(tc)+1}）{tc}', indent=True)

    doc.add_page_break()

    # ================================================================
    # 五、审核过程重点和难点分析
    # ================================================================
    add_heading_styled(doc, '五、审核过程重点和难点分析', level=1)

    add_para(doc, '5.1 重点难点分析矩阵', bold=True, font_size=14)
    add_para(doc, '通过对阿坝县近年来政府性投资项目的特征分析，结合竣工财务决算审核的专业特点，本方案运用风险矩阵方法对审核重点和难点进行系统梳理：', indent=True)

    add_image(doc, os.path.join(SVG_DIR, 'risk_matrix.png'), width_inches=5.8)
    add_para(doc, '图3：竣工财务决算审核重点难点分析矩阵', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

    add_para(doc, '5.2 重点分析', bold=True, font_size=14)
    add_para(doc, '根据阿坝县政府性投资项目的特点和竣工财务决算审核的实际需求，以下五个方面为本项目的审核重点：', indent=True)

    focus_points = [
        ('重点一：工程价款结算的审核',
         '阿坝县地处川西高原，因地质条件复杂、气候条件特殊（冬季施工期短、交通运输不便），工程变更和签证较为常见。工程价款结算审核是本项目的最大重点。我们将重点关注：①工程结算是否经财政评审或有资质的造价咨询机构审定；②工程量计算的准确性，特别是隐蔽工程量的核实；③变更和签证是否按规定程序审批，是否存在先施工后审批的情况；④高原地区施工增加费、冬季施工增加费等特殊费用的计取是否合规；⑤工程款支付台账与结算报告的匹配性。'),
        ('重点二：资金来源及到位情况的审核',
         '阿坝县属于民族地区和欠发达地区，政府投资项目资金来源多元化，通常包含中央预算内投资、省级配套资金、州级配套资金、县级配套资金、专项债券资金等多种渠道。审核重点包括：①各级财政资金下达文件与到位金额的逐一核对；②县级配套资金是否足额到位，是否存在因财力不足导致的资金缺口；③专项债券资金的使用是否符合债券管理规定，是否存在挪用或闲置情况；④各类资金是否按规定在银行开设专户、专款专用。'),
        ('重点三：待摊投资的审核',
         '待摊投资审核是本项目的核心重点。针对阿坝县项目的实际情况，重点关注：①建设单位管理费是否按照规定标准计取，是否存在超标准列支；②勘察设计费、监理费等是否按合同约定支付，与工程进度是否匹配；③土地征用及迁移补偿费的真实性和合规性，征地拆迁程序的规范性；④因高原地区特殊地理条件导致的额外费用（如施工降水费、特殊地基处理费等）是否合理；⑤待摊投资的分摊方法和分摊结果是否准确、合理。'),
        ('重点四：资产交付使用审核',
         '资产交付是项目竣工财务决算的最终环节，也是财政资金转化为国有资产的关键节点。重点包括：①交付资产清单是否完整，是否涵盖了项目建设形成的全部资产；②资产分类是否准确（固定资产、流动资产、无形资产等）；③资产价值构成是否合理，是否按规定剔除了不应计入资产价值的费用；④交付资产是否具备交付使用条件，验收手续是否完备；⑤资产权属证明文件是否齐全。'),
        ('重点五：概算执行情况的对比分析',
         '通过概算与实际投资的全面对比分析，查找偏差原因，是评价项目管理水平和投资效益的重要手段。重点包括：①按工程费用、工程建设其他费用、预备费等分类逐项对比；②超概算项目的超概原因分析和审批手续核实；③节余资金的真实性确认和处置合规性审查；④是否存在未经批准擅自扩大建设规模或提高建设标准的情况；⑤将分析结果汇总形成概算执行情况对比分析表，为财政部门提供决策参考。'),
    ]
    for title, desc in focus_points:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    add_para(doc, '5.3 难点分析', bold=True, font_size=14)
    add_para(doc, '结合阿坝县的区域特点及竣工财务决算审核的复杂性，本项目的主要难点和应对策略如下：', indent=True)

    difficulties = [
        ('难点一：高原地区特殊施工条件导致的工程计量和计价审核难度大',
         '阿坝县平均海拔3300米以上，高寒缺氧、施工期短（每年仅约7个月），地质条件复杂。此类特殊条件导致工程变更频繁，部分变更的合理性判断需要较高的专业技术能力。',
         '应对措施：①配备具有高原地区工程项目经验的专业审核人员；②对重大工程变更事项，必要时聘请工程技术专家进行专业判断；③建立工程变更审核标准清单，提高审核标准化程度；④加强与施工单位和监理单位的沟通，充分了解变更的必要性和合理性。'),
        ('难点二：项目资料不完整或不规范',
         '部分项目由于建设周期长、管理人员变动频繁等原因，建设过程资料可能存在缺失或不够规范的情况，特别是早期项目的资料问题更为突出。',
         '应对措施：①在资料收集阶段即进行完整性评估，对缺失资料逐项登记并提出补正要求；②对于确实无法补充的资料，在审核报告中如实披露并评估其对审核结论的影响；③协助建设单位建立和完善项目资料管理台账；④充分利用财政、审计等部门已有档案进行交叉验证。'),
        ('难点三：多资金来源渠道的核实难度大',
         '阿坝县项目资金来源复杂，涉及中央、省、州、县多级财政资金及专项债券资金，各类资金的管理要求和核算方法不尽相同，核实工作量大。',
         '应对措施：①建立各类资金来源明细台账，逐笔核实资金下达文件和到位情况；②请求阿坝县财政局协助提供各级财政资金的下达文件和拨款凭证；③对专项债券资金使用情况进行专项审核，确保符合债券管理规定；④充分利用银行对账单等外部证据进行资金流水的交叉验证。'),
        ('难点四：项目分布广、交通不便带来的现场核查困难',
         '阿坝县地域面积广（10435平方公里），各乡镇之间的交通条件差异较大，部分项目所在区域道路条件较差，尤其是冬季和雨季。',
         '应对措施：①合理规划现场核查路线和时间，集中安排同一方向的多个项目进行现场核查，提高效率；②对交通极为不便的项目，在确保审核质量的前提下，充分利用无人机航拍、远程视频等现代化技术手段辅助现场核查；③提前与项目建设单位协调现场核查时间，做好充足的行前准备和安全保障；④配置越野车辆等适合高原地区使用的交通工具。'),
        ('难点五：审核时限紧张与审核质量保障的矛盾',
         '根据征集文件规定，项目审核时限根据报送金额从5日至20日不等，时间较为紧张，尤其在多项目并行的情况下，质量与效率的平衡难度大。',
         '应对措施：①根据项目规模和复杂程度合理配备人力资源，重大项目增派经验丰富的注册会计师；②建立标准化工作底稿模板和审核程序清单，提高审核效率；③采用"预审+重点详查"相结合的审核策略，对高风险领域实施详查，低风险领域实施抽审；④利用信息化手段辅助数据分析，提高审核效率。'),
        ('难点六：被审核单位的配合度问题',
         '竣工财务决算审核涉及对项目建设单位财务管理和工程管理工作的全面检查，部分单位可能存在配合度不高或抵触情绪。',
         '应对措施：①在审核进场前与项目建设单位召开沟通会，说明审核目的和意义，争取理解与配合；②建立良好的沟通机制，对审核发现的问题及时沟通，给予充分的解释和说明机会；③保持客观公正的职业态度，避免主观臆断和情绪化表达；④对拒不配合或阻挠审核的情况，及时向阿坝县财政局报告，请求协调处理。'),
    ]
    for title, desc, solution in difficulties:
        add_para(doc, title, bold=True)
        add_para(doc, f'问题描述：{desc}', indent=True)
        add_para(doc, f'应对措施：{solution}', indent=True)

    doc.add_page_break()

    # ================================================================
    # 六、审核质量保证措施
    # ================================================================
    add_heading_styled(doc, '六、审核质量保证措施', level=1)

    add_para(doc, '6.1 质量保证体系概述', bold=True, font_size=14)
    add_para(doc, '我们将依据《会计师事务所质量管理准则第5101号——业务质量管理》（2024年修订）的要求，建立覆盖项目全生命周期的质量管理体系。该体系以"制度保障、人员保障、流程保障、风险管控"为四大支柱，通过PDCA循环实现质量管理的持续改进，确保每一个审核项目的质量都达到或超过行业标准和采购人的要求。', indent=True)

    add_para(doc, '6.2 质量保证体系架构图', bold=True, font_size=14)
    add_image(doc, os.path.join(SVG_DIR, 'quality_arch.png'), width_inches=5.8)
    add_para(doc, '图4：竣工财务决算审核质量保证体系架构图', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

    add_para(doc, '6.3 制度保障措施', bold=True, font_size=14)
    add_para(doc, '建立完善的内部质量控制制度体系，主要包括：', indent=True)
    system_measures = [
        ('执业质量控制制度', '明确执业质量标准、质量控制程序和质量责任，从制度层面确保执业质量。该制度涵盖项目承接与保持、人员委派、业务执行、复核、报告出具等全部环节的质量控制要求。'),
        ('三级复核制度', '建立项目负责人一级复核、部门负责人二级复核、主任会计师/质控部三级复核的逐级质量把关机制。各级复核均有明确的复核重点和复核标准，复核过程形成书面记录。上一级复核未通过的，不得进入下一级复核程序。'),
        ('审核工作底稿规范', '制定标准化的审核工作底稿模板，明确底稿的内容、格式、编制要求和归档标准。底稿应充分记录审核程序、获取的审核证据和形成的审核结论，确保审核过程的可追溯性。'),
        ('重大问题请示报告制度', '审核过程中发现的重大问题和疑难事项，应及时向上一级负责人报告，必要时提请合伙人会议或质量控制委员会讨论决定，确保重大事项的处理恰当、谨慎。'),
        ('质量考核与责任追究制度', '建立项目质量考核评价体系，将审核质量纳入员工绩效考核。对因工作疏忽、失职等导致的质量问题，按照责任追究制度追究相关人员的责任。'),
    ]
    for title, desc in system_measures:
        add_para(doc, f'（{system_measures.index((title, desc))+1}）{title}', bold=True)
        add_para(doc, desc, indent=True)

    add_para(doc, '6.4 人员保障措施', bold=True, font_size=14)
    add_para(doc, '人力资源是审核质量的根本保障。我们将从以下方面确保人员保障到位：', indent=True)
    personnel_measures = [
        '严格按照响应文件承诺配置项目团队成员，项目负责人具备注册会计师资格和高级职称，且具有丰富的政府投资项目竣工财务决算审核经验。',
        '项目团队中注册会计师（CPA）持证人员不少于6人，确保每个审核项目均有CPA持证人员主导和复核。',
        '配备具有工程管理、工程造价、财务管理等复合背景的专业人才，形成专业互补的团队结构。',
        '建立年度专业培训制度，每位审核人员每年参加专业培训不少于40学时，内容涵盖审计准则更新、政府投资项目管理政策、竣工财务决算审核实务等。',
        '人员变更须经阿坝县财政局书面同意，接替人员的资质不低于原人员，且不以此作为审核成果延迟交付的理由。',
        '建立人员胜任能力动态评估机制，定期对项目团队成员的专业能力和执业表现进行评估，对不胜任人员及时调整。',
        '实行项目组内部交叉复核制度，不同审核人员之间相互复核工作底稿，互为质量"守门人"。',
    ]
    for pm in personnel_measures:
        add_para(doc, f'（{personnel_measures.index(pm)+1}）{pm}', indent=True)

    add_para(doc, '6.5 流程保障措施', bold=True, font_size=14)
    add_para(doc, '通过标准化的审核流程和关键节点的质量控制，确保审核工作规范有序：', indent=True)
    process_measures = [
        '审核工作方案审批制度：每个项目的审核工作方案须经部门负责人审核、分管领导批准后实施，确保方案的科学性和可行性。',
        '关键节点质量检查表：在工作流程的五个阶段分别设置质量检查节点，配置标准化检查表，确保各阶段工作质量满足要求后再进入下一阶段。',
        '审核日志全程记录：详细记录审核过程中的重要事项、重大判断和关键决策，形成完整的审核日志，作为三级复核的重要依据。',
        '争议事项集体讨论制度：对审核过程中出现的争议事项，组织项目组内讨论，必要时提请技术委员会审议，确保判断的准确性和一致性。',
        '与被审核单位的沟通记录制度：与被审核单位的所有重要沟通均形成书面记录，双方签字确认，确保沟通内容的可追溯性。',
        '项目进度周报制度：每周编制项目进度报告，向采购人报告审核进展、发现的问题和下一步计划，保持信息透明。',
    ]
    for pm in process_measures:
        add_para(doc, f'（{process_measures.index(pm)+1}）{pm}', indent=True)

    add_para(doc, '6.6 风险管控措施', bold=True, font_size=14)
    add_para(doc, '建立全面的风险识别、评估和应对机制：', indent=True)
    risk_measures = [
        '项目风险分级评估：根据项目投资规模、复杂程度、资金来源等因素，对每个项目进行风险等级评估（高/中/低风险），针对不同风险等级配置不同的审核资源。高风险项目实行"双主审"制度。',
        '重大风险预警机制：建立重大风险事项清单和预警指标，审核过程中一旦触发预警阈值，立即启动应急响应程序，报告项目负责人和部门负责人。',
        '廉洁从业承诺书签署：所有参与项目的人员在进场前签署廉洁从业承诺书，承诺不索取或收受被审核单位的财物，不利用职务之便谋取私利。',
        '保密协议签署：全体项目人员签署保密协议，未经采购人书面许可，不向任何第三方披露项目信息、审核发现和审核结论。',
        '利益冲突审查：项目组组建前进行利益冲突审查，确保项目组成员与被审核单位之间不存在可能影响独立性的利害关系。',
        '质量事故应急预案：制定质量事故应急预案，明确质量事故的认定标准、报告程序、处置流程和善后措施，确保在质量事故发生时能够及时、有效应对。',
    ]
    for rm in risk_measures:
        add_para(doc, f'（{risk_measures.index(rm)+1}）{rm}', indent=True)

    add_para(doc, '6.7 三级复核制度详解', bold=True, font_size=14)
    add_para(doc, '三级复核是本项目最核心的质量控制措施，具体实施要求如下：', indent=True)

    # 三级复核表
    table_review = doc.add_table(rows=1, cols=4)
    table_review.style = 'Table Grid'
    table_review.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_r = table_review.rows[0]
    for i, text in enumerate(['复核级别', '复核人', '复核重点', '复核要求']):
        hdr_r.cells[i].text = text
        for p in hdr_r.cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(hdr_r.cells[i], '1a5276')
        for p in hdr_r.cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    reviews = [
        ('一级复核', '项目负责人\n（CPA+高级职称）', '①审核程序是否充分执行\n②工作底稿是否完整、规范\n③取证是否充分、适当\n④审核发现是否准确\n⑤数据计算是否正确\n⑥底稿之间勾稽关系是否一致', '逐项全面复核，\n复核率100%，\n在每份工作底稿上\n签署复核意见'),
        ('二级复核', '部门负责人\n（CPA）', '①重大事项处理是否恰当\n②审核标准适用是否准确\n③审核结论是否合理\n④一级复核是否到位\n⑤与采购人沟通是否充分\n⑥报告征求意见稿质量', '重点复核+抽查，\n对重大事项逐项复核，\n一般事项抽查比例\n不低于30%，\n签署书面复核意见'),
        ('三级复核\n（终审）', '主任会计师\n或质控部', '①审核报告的整体质量\n②重大风险事项的结论\n③报告表述的规范性\n④法律法规遵守情况\n⑤是否存在应披露未\n　披露事项\n⑥审核意见的恰当性', '全面终审，\n对报告全文逐字复核，\n评估审核结论的\n恰当性和充分性，\n签署终审意见'),
    ]
    for rv in reviews:
        row = add_table_row(table_review, rv, font_size=9)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    add_para(doc, '6.8 质量持续改进机制', bold=True, font_size=14)
    add_para(doc, '建立基于PDCA循环的质量持续改进机制：', indent=True)
    pdca = [
        'Plan（策划）：根据上一年度质量检查发现的问题和采购人的反馈意见，制定年度质量改进计划，修订完善质量管理制度和操作规程。',
        'Do（实施）：按照改进计划落实各项改进措施，组织专题培训，更新工作底稿模板和审核程序清单，在项目中应用改进后的方法和工具。',
        'Check（检查）：每半年进行一次项目质量抽查，每年进行一次全面质量检查，邀请外部专家进行同行评审，对照质量标准检查执行情况，查找问题和不足。',
        'Act（改进）：对检查和评审发现的问题进行分析总结，查找根本原因，修订完善相关制度，制定整改措施并跟踪落实，实现质量管理的螺旋式上升。',
    ]
    for pd in pdca:
        add_para(doc, f'（{pdca.index(pd)+1}）{pd}', indent=True)

    doc.add_page_break()

    # ================================================================
    # 七、审核后续支持服务
    # ================================================================
    add_heading_styled(doc, '七、审核后续支持服务', level=1)

    add_para(doc, '7.1 后续支持服务理念', bold=True, font_size=14)
    add_para(doc, '竣工财务决算审核工作的完成不是服务的终点，而是持续服务的起点。我们秉承"专业服务、持续跟进、深度赋能"的服务理念，在出具正式审核报告后，为阿坝县财政局提供全方位、多层次的后续支持服务，确保审核成果的有效运用和项目管理的持续改进。', indent=True)

    add_para(doc, '7.2 后续支持服务内容', bold=True, font_size=14)

    after_services = [
        ('7.2.1 审核报告解读与答疑服务',
         '在提交审核报告后，安排项目负责人为阿坝县财政局提供面对面的审核报告解读服务，对报告中的审核发现、审核结论和建议意见进行详细说明。同时设立长期咨询热线和专属服务微信群，由项目负责人和专业技术人员提供7×24小时的在线答疑服务，及时解答财政局在日常管理中遇到的竣工财务决算相关问题。'),
        ('7.2.2 审核问题整改跟踪服务',
         '对审核报告中提出的问题和建议，协助阿坝县财政局制定整改方案和整改台账，明确整改责任单位和整改时限。在整改期间定期跟踪整改进度，对整改过程中遇到的疑难问题提供专业指导。整改完成后，协助财政局对整改结果进行验收评估，形成整改情况报告，确保问题得到有效解决。'),
        ('7.2.3 项目决算批复支持服务',
         '根据《基本建设财务规则》和《基本建设项目竣工财务决算管理暂行办法》的要求，协助阿坝县财政局完成项目竣工财务决算的批复工作。包括：协助准备决算批复所需的材料和文件，辅助审核被批复项目的决算数据，协助起草决算批复文件，确保批复程序的规范性和批复内容的准确性。'),
        ('7.2.4 资产管理衔接服务',
         '协助阿坝县财政局做好项目竣工财务决算与国有资产管理的衔接工作。具体包括：协助核对资产交付清单与决算报告的一致性，辅助指导项目建设单位规范办理资产交付手续，协助财政部门完成资产管理信息系统中的资产登记和入账工作，确保"项目竣工一项、资产交付一项、管理到位一项"。'),
        ('7.2.5 制度建设咨询服务',
         '基于竣工财务决算审核中发现的共性问题和典型案例，为阿坝县财政局提供项目管理制度建设的咨询服务。包括：协助梳理和完善阿坝县政府投资项目财务管理制度、竣工财务决算编制规程、项目资金管理办法等制度文件，从制度建设层面提升政府投资项目的管理水平。'),
        ('7.2.6 培训与知识转移服务',
         '为阿坝县财政局及项目建设单位的相关管理人员提供专业的培训服务。培训内容可根据实际需求定制，包括但不限于：基本建设财务管理实务、竣工财务决算编制方法与技巧、政府投资项目资金管理要点、常见问题及案例分析等。每年至少组织2次集中培训，培训资料免费提供。'),
        ('7.2.7 政策法规动态更新服务',
         '持续跟踪国家及四川省关于政府投资项目财务管理、竣工财务决算的最新政策法规和行业标准的动态变化，及时向阿坝县财政局推送政策解读和实务操作指引。每季度编制一期《政府投资项目管理政策动态》，以电子版形式发送至财政局指定联系人。'),
        ('7.2.8 年度审核工作总结与分析服务',
         '在每个服务年度结束时，向阿坝县财政局提交年度审核工作总结报告，内容包括：年度审核项目概况、审核发现的共性问题和典型案例、问题成因分析和改进建议、下一年度审核工作思路。通过对年度审核工作的系统总结和分析，为财政局优化政府投资项目管理提供数据支撑和决策参考。'),
        ('7.2.9 应急响应与重大事项报告服务',
         '对阿坝县财政局在政府投资项目管理中遇到的紧急事项或重大疑难问题，提供应急响应服务：2小时内电话响应、24小时内书面回复、48小时内（视情况）到达现场。重大事项包括但不限于：上级审计或检查中发现的问题、项目资金使用异常、重大工程变更争议等。'),
        ('7.2.10 档案管理与数据查询服务',
         '对已完成的审核项目档案实行统一管理，档案保存期限不少于10年。建立电子档案检索系统，为阿坝县财政局提供便捷的档案查询和历史数据统计服务。在框架协议有效期内及期满后2年内，财政局可随时调阅已完成项目的审核工作底稿和审核报告。'),
    ]
    for title, desc in after_services:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    add_para(doc, '7.3 后续支持服务保障措施', bold=True, font_size=14)

    support_measures = [
        ('指定服务联系人', '安排专人作为阿坝县财政局后续服务的固定联系人，确保沟通渠道畅通、信息传递及时。联系人具备注册会计师资格，具有5年以上政府审计服务经验。'),
        ('建立服务台账', '对每一次后续服务建立服务台账，记录服务时间、服务内容、服务结果和财政局反馈意见，确保服务可追溯、可评价。'),
        ('定期服务回访', '每季度安排一次上门回访，主动了解阿坝县财政局在竣工财务决算管理中的新需求和新问题，及时调整和优化后续服务内容。'),
        ('服务质量评价', '每次后续服务结束后，请财政局对服务质量进行评价，评价结果作为项目团队绩效考核和服务改进的重要依据。'),
        ('知识库建设', '建立阿坝县政府投资项目竣工财务决算审核知识库，积累审核经验和典型案例，为后续服务提供知识支撑。'),
    ]
    for title, desc in support_measures:
        add_para(doc, f'（{support_measures.index((title, desc))+1}）{title}', bold=True)
        add_para(doc, desc, indent=True)

    # ================================================================
    # 结束语
    # ================================================================
    doc.add_paragraph()
    add_para(doc, '结束语', bold=True, font_size=14)
    add_para(doc, '本公司（拟投入本项目团队）深耕工程竣工财务决算审计领域多年，积累了丰富的政府投资项目审核经验，对四川省阿坝州地区的政府投资项目管理特点有深入的了解。我们将以高度的责任感和专业的执业精神，严格遵循各项法律法规和执业准则，充分发挥自身的人才优势、技术优势和经验优势，为阿坝县财政局提供优质、高效、全面的竣工财务决算审核服务，为阿坝县政府投资项目的规范化管理和财政资金的安全高效使用贡献力量。', indent=True)
    add_para(doc, '我们郑重承诺：严格按照本服务方案的内容和标准执行，确保每一项服务承诺落到实处，每一个审核项目经得起检验，以专业能力和优质服务赢得阿坝县财政局的信任和认可。', indent=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # 落款
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = company_para.add_run('（供应商名称并加盖公章）')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = date_para.add_run('2026年    月    日')
    run2.font.size = Pt(12)
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ================================================================
    # 保存文档
    # ================================================================
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == '__main__':
    path = create_document()
    print(f'文档已生成：{path}')
