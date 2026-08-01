# -*- coding: utf-8 -*-
"""
口径一致性检查器 v0.2
加载项目数据档案，对报告/底稿进行口径一致性扫描
用法: python caliber_checker.py --profile profiles/pidou_2026/ --report "报告.docx"
"""
import sys, os, json, re, argparse, datetime
from pathlib import Path
sys.stdout.reconfigure(encoding='utf-8')

HERE = Path(__file__).parent

# ─── Extract numbers & claims from report ──────────────
def extract_numeric_claims(text):
    """从文本中提取数值声明,如: '项目评审金额51.2亿元'"""
    claims = []
    # 过滤纯编号/代码行（如 "7.1" "7.2" "C01" "P01"）
    noise_patterns = [
        r'^[\d.]+\s*$', r'^[A-Z]\d{2},?\s*$', r'^[A-Z]\d{2}$',
        r'^[一二三四五六七八九十]\s*$', r'^[\d.]+[）)]',
        r'^[\d]+[、.．]\s*$',  # "1、" "2." "3．"
    ]
    def _is_noise(indicator, raw):
        s = raw.strip()
        for np in noise_patterns:
            if re.fullmatch(np, s):
                return True
        # 纯标点+数字组合
        if len(s) <= 5 and re.match(r'^[\d.（）()\-\s、．]+$', s):
            return True
        # indicator 本身是章节编号（如 "一、编制说明" "2.1"）
        if indicator and re.match(r'^[一二三四五六七八九十\d]+[、.．]\s*\S{0,10}$', indicator):
            return True
        # indicator 太短且无业务含义（如 "C" "N" "Q" 单字母）
        if indicator and len(indicator) <= 2 and re.match(r'^[A-Z]\d?$', indicator):
            return True
        # raw 包含换行符（章节标题跨行，如 "评审业务费（A类）\n  7.2"）
        if '\n' in raw:
            return True
        # indicator 以 "（" 或 "(" 结尾，说明匹配不完整（如 "中心运行经费（27"）
        if indicator and indicator.endswith(('（', '(')):
            return True
        return False
    
    # 模式: 指标名称 + 数字 + 单位
    patterns = [
        # "部门预算1,123万元" / "项目评审金额51.2亿元" / "核差率8.43%"
        # 允许数字和单位之间有括号、空格、连字符
        r'([\u4e00-\u9fff][\u4e00-\u9fff\u3000-\u303f\uff00-\uffef（）\w]{1,30}?)\s*[≥≤>＜=：:]?\s*(\d[\d,.]+)[\s\-（(]*?(亿[元]?|万[元]?%?|[元]?%?|[人次个项件天])(?=[^\w]|$)',
        # "评审业务费（680-1,348万元）"中的范围值，返回 [min, max]
        r'([\u4e00-\u9fff][\u4e00-\u9fff（）]{2,20}?)[（(]\s*(\d[\d,.]+)[\s\-]+(\d[\d,.]+)[\s\-]*?(亿[元]?|万[元]?)',
        # 送审79.74亿元 / 审减8.91亿元
        r'(送审|审减|审减率|核差率|核减)[：:\s]*(\d[\d,.]+)[\s\-]*?(亿[元]?|万[元]?|%)',
        # 预算\d+万元 / 经费\d+万元
        r'(预算|经费|支出|收入)[：:\s]*(\d[\d,.]+)[\s\-]*?(亿[元]?|万[元]?)',
        # 核差率8.43% （数字后直接跟%）
        r'([\u4e00-\u9fff]{2,20}?)\s*(\d[\d,.]+)\s*%',
    ]
    seen_raw = set()
    for pat_idx, pat in enumerate(patterns):
        for m in re.finditer(pat, text):
            raw_text = m.group(0).strip()
            if raw_text in seen_raw:
                continue
            try:
                val_str = m.group(2).replace(',', '')
                val = float(val_str)
                # 过滤单数字纯编号 (如 "1" "2")
                if val < 0.1 and len(raw_text) <= 3:
                    continue
            except (ValueError, AttributeError):
                continue
            indicator = m.group(1).strip() if m.lastindex and m.lastindex >= 1 else ''
            if _is_noise(indicator, raw_text):
                continue
            # 处理不同模式的 group 索引
            if pat_idx == 1:  # 范围值模式
                unit = m.group(4) if m.lastindex and m.lastindex >= 4 else ''
                val2_str = m.group(3).replace(',', '')
                try:
                    val2 = float(val2_str)
                    val_range = [min(val, val2), max(val, val2)]
                except:
                    val_range = None
            elif pat_idx == 4:  # % 模式
                unit = '%'
                val_range = None
            else:
                unit = m.group(3) if m.lastindex and m.lastindex >= 3 else ''
                val_range = None
            seen_raw.add(raw_text)
            claim = {
                'raw': raw_text,
                'indicator': indicator,
                'value': val,
                'unit': unit,
                'position': m.start()
            }
            if val_range:
                claim['value_range'] = val_range
            claims.append(claim)
    return claims

def extract_from_docx(docx_path):
    """从Word文档提取正文"""
    from docx import Document
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格
    tables = []
    for table in doc.tables:
        for row in table.rows:
            tables.append(' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return '\n'.join(paras) + '\n' + '\n'.join(tables)

# ─── Caliber matching ──────────────────────────────────
def normalize_caliber_name(name):
    """口径名称归一化"""
    return re.sub(r'[\s（(）)]', '', name).lower()

def check_consistency(claims, caliber_definitions):
    """
    比对报告中的数值声明与档案里的口径定义
    返回: findings (P0/P1/P2 分级)
    """
    findings = []
    caliber_index = {normalize_caliber_name(c['name']): c for c in caliber_definitions}
    
    for claim in claims:
        cn = normalize_caliber_name(claim['indicator'])
        
        # 模糊匹配
        best_match = None
        for key, cal in caliber_index.items():
            if key in cn or cn in key:
                best_match = cal
                break
        
        if not best_match:
            # 尝试部分匹配
            for key, cal in caliber_index.items():
                overlap = set(key) & set(cn)
                if len(overlap) / max(len(key), len(cn), 1) > 0.4:
                    best_match = cal
                    break
        
        if best_match:
            # 比对单位
            if best_match.get('unit') and claim['unit'] != best_match['unit']:
                findings.append({
                    'level': 'P1',
                    'type': 'unit_mismatch',
                    'claim': claim['raw'],
                    'caliber': best_match['name'],
                    'detail': f"报告用'{claim['unit']}'，档案定义为'{best_match['unit']}'",
                    'suggestion': f"确认口径单位，统一为'{best_match['unit']}'或更新档案"
                })
            
            # 比对数值范围
            expected = best_match.get('expected_range', {})
            if expected:
                # 使用 value_range（如果有）或单值
                if 'value_range' in claim:
                    val_min, val_max = claim['value_range']
                else:
                    val_min = val_max = claim['value']
                
                if 'min' in expected and expected['min'] is not None and val_max < expected['min']:
                    findings.append({
                        'level': 'P1',
                        'type': 'value_out_of_range',
                        'claim': claim['raw'],
                        'caliber': best_match['name'],
                        'detail': f"报告值{val_min}-{val_max}{claim['unit']} < 档案下限{expected['min']}",
                        'suggestion': '核实数据来源，确认是否有统计口径差异'
                    })
                if 'max' in expected and expected['max'] is not None and val_min > expected['max']:
                    findings.append({
                        'level': 'P2',
                        'type': 'value_out_of_range',
                        'claim': claim['raw'],
                        'caliber': best_match['name'],
                        'detail': f"报告值{val_min}-{val_max}{claim['unit']} > 档案上限{expected['max']}",
                        'suggestion': '核实是否数据更新或口径变化'
                    })
        else:
            # 报告中有声明但档案里没有对应口径定义
            findings.append({
                'level': 'P2',
                'type': 'unregistered_caliber',
                'claim': claim['raw'],
                'caliber': '',
                'detail': f"报告指标'{claim['indicator']}'在数据档案中无对应口径定义",
                'suggestion': '在数据档案中补充该指标的口径定义，或确认是否应删除'
            })
    
    return findings

# ─── Main ───────────────────────────────────────────────
def run_caliber_check(profile_dir, report_path=None, report_text=None, project_label=''):
    """执行口径一致性检查"""
    # 加载数据档案
    profiles = []
    pd = Path(profile_dir)
    for pf in sorted(pd.glob('*_profile.json')):
        with open(pf, encoding='utf-8') as f:
            profiles.append(json.load(f))
    
    if not profiles:
        return {'error': f'未找到数据档案 ({profile_dir}/*_profile.json)', 'findings': []}
    
    # 提取报告文本
    if report_path:
        if report_path.endswith('.docx'):
            text = extract_from_docx(report_path)
        elif report_path.endswith('.txt') or report_path.endswith('.md'):
            text = open(report_path, encoding='utf-8').read()
        else:
            text = report_path  # plain text
    elif report_text:
        text = report_text
    else:
        return {'error': '请提供 --report 或 --text', 'findings': []}
    
    # 汇总所有口径
    all_calibers = []
    for p in profiles:
        all_calibers.extend(p.get('key_calibers', []))
    
    # 提取声明 & 检查
    claims = extract_numeric_claims(text)
    findings = check_consistency(claims, all_calibers)
    
    # 统计
    p0 = [f for f in findings if f['level'] == 'P0']
    p1 = [f for f in findings if f['level'] == 'P1']
    p2 = [f for f in findings if f['level'] == 'P2']
    
    result = {
        'project': project_label,
        'checked_at': datetime.datetime.now().isoformat(),
        'profiles_loaded': len(profiles),
        'calibers_defined': len(all_calibers),
        'claims_found': len(claims),
        'findings_count': {'P0': len(p0), 'P1': len(p1), 'P2': len(p2)},
        'findings': findings
    }
    
    return result

def print_report(result):
    """打印检查报告"""
    if 'error' in result:
        print(f"❌ {result['error']}")
        return
    
    c = result['findings_count']
    print(f"\n{'='*60}")
    print(f"  📋 口径一致性检查报告")
    print(f"  项目: {result['project']}")
    print(f"  档案: {result['profiles_loaded']} 份 | 口径: {result['calibers_defined']} 个")
    print(f"  报告声明: {result['claims_found']} 条")
    print(f"{'='*60}")
    
    if c['P0'] == 0 and c['P1'] == 0 and c['P2'] == 0:
        print("  ✅ 无口径不一致问题")
        return
    
    for f in result['findings']:
        emoji = {'P0': '🔴', 'P1': '🟡', 'P2': '⚪'}.get(f['level'], '❓')
        print(f"\n  {emoji} [{f['level']}] {f['type']}")
        print(f"     报告: '{f['claim']}'")
        print(f"     口径: {f['caliber']}" if f['caliber'] else "     口径: (档案无定义)")
        print(f"     说明: {f['detail']}")
        print(f"     建议: {f['suggestion']}")
    
    print(f"\n{'='*60}")
    print(f"  P0={c['P0']} | P1={c['P1']} | P2={c['P2']}")
    print(f"{'='*60}")

if __name__ == '__main__':
    p = argparse.ArgumentParser(description='口径一致性检查器 v0.1')
    p.add_argument('--profile', required=True, help='数据档案目录 (如 profiles/pidou_2026/)')
    p.add_argument('--report', help='报告 docx/txt/md 文件路径')
    p.add_argument('--text', help='直接输入报告文本')
    p.add_argument('--json', action='store_true', help='JSON 输出')
    p.add_argument('--use-llm', action='store_true', help='启用大模型语义匹配（可选，增加成本）')
    
    args = p.parse_args()
    result = run_caliber_check(args.profile, args.report, args.text, 
                               os.path.basename(args.profile.rstrip('/\\')))
    
    # LLM 增强（可选）
    if args.use_llm and result.get('findings'):
        print('\n🤖 启用 LLM 语义匹配...')
        try:
            from caliber_llm_matcher import batch_match_claims
            # 重新提取 claims 并增强
            if args.report:
                if args.report.endswith('.docx'):
                    text = extract_from_docx(args.report)
                elif args.report.endswith('.txt') or args.report.endswith('.md'):
                    text = open(args.report, encoding='utf-8').read()
                else:
                    text = args.report
            elif args.text:
                text = args.text
            else:
                text = ''
            
            claims = extract_numeric_claims(text)
            # 加载口径定义
            profiles = []
            pd = Path(args.profile)
            for pf in sorted(pd.glob('*_profile.json')):
                with open(pf, encoding='utf-8') as f:
                    profiles.append(json.load(f))
            all_calibers = []
            for p in profiles:
                all_calibers.extend(p.get('key_calibers', []))
            
            enhanced_claims = batch_match_claims(claims, all_calibers)
            result['claims'] = enhanced_claims
            result['llm_enabled'] = True
        except Exception as e:
            print(f'  ⚠️ LLM 增强失败: {e}')
            result['llm_enabled'] = False
    
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print_report(result)
