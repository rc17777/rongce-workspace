# -*- coding: utf-8 -*-
"""
融策·审盾 参数溯源校验器 v1.0
=================================
功能：扫描.docx格式的绩效评估报告，自动：
  1. 提取所有数值型参数
  2. 检查参数是否有出处标注
  3. 标记"疑似无出处参数"
  4. 检查报告中是否存在"拍脑袋数字"模式（如55:45）
  5. 给出风险等级

用法：python param_checker.py "报告.docx" [--strict]

输出：JSON格式检查结果 + 控制台简要报告
"""
import sys, re, json, os
from docx import Document

# 无出处信号的文字模式
NO_SOURCE_PATTERNS = [
    r'取整',
    r'综合判断',
    r'经验值',
    r'行业惯例',
    r'（假设）',
    r'（暂定）',
    r'建议取值',
    r'按.*估算',
    r'参考',
]

# 出处信号的文字模式（正面）
HAS_SOURCE_PATTERNS = [
    r'根据《',
    r'依据《',
    r'序号.*文件',
    r'详见',
    r'来源',
    r'数据引自',
    r'《.*》.*第.*页',
    r'表.*显示',
    r'测算底稿',
    r'.xlsx',
    r'.docx',
]

# 敏感参数关键词（这些参数必须有出处）
SENSITIVE_PARAMS = [
    '票价', '成本', '客座率', '比例', '吞吐量', '消费', '补贴', '补助',
    '单价', '人次', '班次', '座位', '比例', '占比', '收费', '收入',
    '支出', '缺口', '标准', '费率', '利息', '折旧',
]

def scan_document(filepath, strict=False):
    """扫描文档，返回检查结果"""
    if not os.path.exists(filepath):
        return {'error': f'文件不存在: {filepath}'}
    
    doc = Document(filepath)
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    findings = []
    issues = []
    
    # 1. 提取带数字的段落
    number_pattern = re.compile(r'[\d,]+\.?\d*')
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text or len(text) < 10:
            continue
        
        nums = number_pattern.findall(text)
        if not nums:
            continue
        
        # 取第一个三位数以上的数字（通常是参数）
        for n in nums:
            try:
                val = float(n.replace(',', ''))
                if 10 < val < 10000000:  # 过滤页码等
                    break
            except:
                continue
        else:
            continue
        
        # 2. 检查这段文字中是否有关键参数词
        has_sensitive = any(kw in text for kw in SENSITIVE_PARAMS)
        
        # 3. 检查是否有出处信号
        has_source = any(re.search(p, text) for p in HAS_SOURCE_PATTERNS)
        has_no_source_signal = any(re.search(p, text) for p in NO_SOURCE_PATTERNS)
        
        # 4. 检查经典比例模式（XX:YY）
        ratio_match = re.search(r'(\d{2})\s*[:：比]\s*(\d{2})', text)
        if ratio_match:
            r1, r2 = ratio_match.group(1), ratio_match.group(2)
            if f'{r1}:{r2}' not in findings:
                findings.append({
                    'type': 'ratio',
                    'value': f'{r1}:{r2}',
                    'paragraph': i+1,
                    'text_snippet': text[:100],
                    'has_source': has_source,
                })
                if not has_source:
                    issues.append({
                        'severity': 'HIGH' if has_sensitive else 'MEDIUM',
                        'type': '无出处的比例假设',
                        'param': f'{r1}:{r2}',
                        'paragraph': i+1,
                        'text': text[:120],
                        'suggestion': '进出港比例等比例参数必须有出处。无出处则改用区间表达。',
                    })
        
        # 5. 检查疑似无出处参数
        if has_no_source_signal and has_sensitive and not has_source:
            issues.append({
                'severity': 'MEDIUM',
                'type': '疑似无出处的敏感参数',
                'param': text[:80],
                'paragraph': i+1,
                'text': text[:120],
                'suggestion': '请在参数依据表中登记或标注出处。',
            })
    
    # 6. 检查是否存在"关键参数依据表"
    has_param_table = any('参数依据表' in p.text or '关键参数' in p.text for p in doc.paragraphs)
    
    # 7. 检查计算链是否闭合（数字相等 = 可疑信号）
    equal_patterns = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        # 查找 "X = Y = Z" 型推导链
        chain_matches = re.findall(r'(\d{3,4})\s*[＝=]\s*(\d{3,4})\s*[＝=]\s*(\d{3,4})', text)
        if chain_matches:
            for a, b, c in chain_matches:
                if a == b or b == c:
                    equal_patterns.append({
                        'type': 'number_chain_equality',
                        'values': f'{a}={b}={c}',
                        'paragraph': i+1,
                        'text_snippet': text[:100],
                    })
                    if not any(re.search(p, text) for p in HAS_SOURCE_PATTERNS):
                        issues.append({
                            'severity': 'LOW',
                            'type': '数字链闭合但无出处验证',
                            'param': f'{a}={b}={c}',
                            'paragraph': i+1,
                            'text': text[:120],
                            'suggestion': '多层相等数字链需每层验证出处，防止自我验证循环。',
                        })
    
    # 8. 汇总评分
    high_issues = [i for i in issues if i['severity'] == 'HIGH']
    medium_issues = [i for i in issues if i['severity'] == 'MEDIUM']
    low_issues = [i for i in issues if i['severity'] == 'LOW']
    
    if not has_param_table:
        issues.insert(0, {
            'severity': 'HIGH',
            'type': '关键参数依据表缺失',
            'param': 'N/A',
            'paragraph': 0,
            'text': '报告中未找到"关键参数依据表"',
            'suggestion': '请使用融策报告模板v2.0，该模板内置参数依据表。',
        })
        high_issues = [issues[0]] + high_issues
    
    # 判定等级
    if len(high_issues) > 0:
        risk = '🔴 高风险 — 建议退回修改'
    elif len(medium_issues) > 2:
        risk = '🟡 中等风险 — 需补充说明'
    elif len(medium_issues) > 0 or len(low_issues) > 3:
        risk = '🟢 低风险 — 可提交，建议复核人关注'
    else:
        risk = '✅ 通过 — 无明显参数风险'
    
    return {
        'file': filepath,
        'risk_level': risk,
        'has_param_table': has_param_table,
        'ratios_found': len(findings),
        'findings': findings,
        'issues': issues,
        'summary': {
            'HIGH': len(high_issues),
            'MEDIUM': len(medium_issues),
            'LOW': len(low_issues),
        }
    }


def print_report(result):
    """控制台友好输出"""
    print('=' * 70)
    print('  融策·审盾 参数溯源校验 v1.0')
    print('=' * 70)
    
    if 'error' in result:
        print(f'\n  ❌ {result["error"]}')
        return
    
    print(f'\n  文件: {os.path.basename(result["file"])}')
    print(f'  参数依据表: {"✅ 已找到" if result["has_param_table"] else "❌ 缺失"}')
    print(f'  发现比例参数: {result["ratios_found"]} 个')
    print(f'\n  风险评级: {result["risk_level"]}')
    print(f'  问题汇总: 🔴{result["summary"]["HIGH"]} 🟡{result["summary"]["MEDIUM"]} 🟢{result["summary"]["LOW"]}')
    
    if result['issues']:
        print(f'\n  {"─" * 60}')
        for issue in result['issues']:
            icon = {'HIGH': '🔴', 'MEDIUM': '🟡', 'LOW': '🟢'}.get(issue['severity'], '•')
            print(f'  {icon} [{issue["type"]}] (P{issue["paragraph"]})')
            print(f'     {issue["suggestion"]}')
    
    if result['ratios_found'] > 0:
        print(f'\n  📊 发现的比例参数:')
        for f in result['findings']:
            src = '有出处' if f['has_source'] else '⚠ 无出处'
            print(f'     {f["value"]}  {src}  → P{f["paragraph"]}: {f["text_snippet"][:60]}...')
    
    print(f'\n{"=" * 70}\n')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python param_checker.py "报告.docx" [--json]')
        print('示例: python param_checker.py "郑州=九寨=武汉...报告.docx"')
        sys.exit(1)
    
    filepath = sys.argv[1]
    strict = '--strict' in sys.argv
    output_json = '--json' in sys.argv
    
    result = scan_document(filepath, strict=strict)
    
    if output_json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print_report(result)
