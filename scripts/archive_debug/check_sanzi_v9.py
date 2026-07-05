from pathlib import Path
from docx import Document

p = Path(r'D:\openclaw-workspace\projects\西昌三资三化投标\西昌三资三化投标服务方案-完整版-V9-课件资料增强版.docx')
d = Document(p)
markers = [
    '1.13 最新讲义资料补充',
    '3.16 最新课件转化',
    '3.17 数字化赋能',
    '5.6 “四张清单”成果体系升级',
    '6.17 最新资料补充',
    '10.7 后续服务升级',
    '附录：本次最新课件资料整合说明',
]
print('file', p)
print('paragraphs', len(d.paragraphs), 'tables', len(d.tables), 'sections', len(d.sections))
for marker in markers:
    found = [(i, para.style.name, para.text[:120]) for i, para in enumerate(d.paragraphs) if marker in para.text]
    print(marker, 'FOUND' if found else 'MISSING', found[:2])
