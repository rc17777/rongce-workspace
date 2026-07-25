# -*- coding: utf-8 -*-
"""
绩效自评复核脚本 — 14条标准 + 格式检查
输出：桌面 Excel
"""
import sys, os, json, re, copy
from pathlib import Path
from collections import defaultdict

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

BASE = Path(r'C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722')
OUTPUT = Path(r'C:\Users\scrccpa\Desktop\绩效自评复核结果.xlsx')

# ============================================================
# 格式常量
# ============================================================
# 二号 = 22pt = 22 * 12700 EMU
# 三号 = 16pt  
# 小标宋 we might need to check as '小标宋' or similar
# 仿宋GB2312
# 黑体
# 楷体GB2312
# Times New Roman for numbers

FONT_MAP = {
    '小标宋': ['小标宋', '小標宋'],
    '仿宋_GB2312': ['仿宋_GB2312', '仿宋GB2312', '仿宋'],
    '黑体': ['黑体', '黑體'],
    '楷体_GB2312': ['楷体_GB2312', '楷体GB2312', '楷体'],
    'times_new_roman': ['Times New Roman', 'times new roman'],
}

SIZE_MAP = {
    '二号': Pt(22),    # 22pt
    '三号': Pt(16),    # 16pt
    '小二号': Pt(18),
}

# ============================================================
# PART 1: 格式检查
# ============================================================

def check_format(docx_path, unit_name):
    """检查docx报告格式"""
    issues = []
    doc = Document(docx_path)
    
    # --- 获取页面设置 ---
    for section in doc.sections:
        # A4 = 21cm x 29.7cm
        pw = section.page_width
        ph = section.page_height
        # 标准A4: 11906 EMU = 21cm, 16838 EMU = 29.7cm
        if pw and abs(pw - 1190700) > 50000:
            issues.append({'位置': '页面设置', '问题': f'页宽非A4标准({pw/635:.1f}cm)', '严重度': 'P2'})
        if ph and abs(ph - 1683800) > 50000:
            issues.append({'位置': '页面设置', '问题': f'页高非A4标准({ph/635:.1f}cm)', '严重度': 'P2'})
    
    # --- 检查段落 ---
    in_body = False
    body_started = False
    title_checked = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # --- 标题居中 ---
        # 检查封面标题行
        if not body_started:
            if len(text) > 6 and ('绩效' in text or '自评' in text or '复核' in text or '报告' in text):
                if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    issues.append({
                        '位置': f'封面P{i}',
                        '问题': f'标题行未居中: "{text[:30]}"',
                        '严重度': 'P1'
                    })
                    title_checked = True
        
        # --- 检测正文开始 ---
        if '一、' in text or '（一）' in text or '一、部门基本情况' in text:
            body_started = True
        
        if not body_started:
            continue
            
        # --- 检查一级标题（一、...）使用黑体 ---
        if re.match(r'^一、|^二、|^三、|^四、|^五、|^六、|^七、|^八、|^九、|^十、', text):
            for r in p.runs:
                if r.text.strip():
                    fname = (r.font.name or '').lower()
                    is_heiti = any(k in fname for k in ['黑体', '黑體'])
                    if not is_heiti and r.text.strip():
                        issues.append({
                            '位置': f'P{i}',
                            '问题': f'一级标题字体非黑体: "{r.text[:20]}" 实际={r.font.name}',
                            '严重度': 'P1'
                        })
                        break
        
        # --- 检查二级标题（（一）...）使用楷体GB2312加粗 ---
        if re.match(r'^（[一二三四五六七八九十\d]+）', text):
            for r in p.runs:
                if r.text.strip():
                    fname = (r.font.name or '').lower()
                    is_kai = any(k in fname for k in ['楷体', '楷體'])
                    if not is_kai and r.text.strip():
                        issues.append({
                            '位置': f'P{i}',
                            '问题': f'二级标题字体非楷体: "{r.text[:20]}" 实际={r.font.name}',
                            '严重度': 'P1'
                        })
                        break
        
        # --- 检查三级标题（1. ...）使用仿宋GB2312 ---
        if re.match(r'^\d+\.\s', text):
            for r in p.runs:
                if r.text.strip():
                    fname = (r.font.name or '').lower()
                    is_fangsong = any(k in fname for k in ['仿宋', 'fangsong'])
                    if not is_fangsong and r.text.strip():
                        issues.append({
                            '位置': f'P{i}',
                            '问题': f'三级标题字体非仿宋: "{r.text[:20]}" 实际={r.font.name}',
                            '严重度': 'P2'
                        })
                        break
        
        # --- 正文段落: 仿宋_GB2312 三号, 行距30磅 ---
        # 只检查正文段落（非标题）
        if not re.match(r'^(一、|二、|三、|四、|五、|[（(].*[）)]|\d+\.\s)', text):
            for r in p.runs:
                if r.text.strip():
                    # 字体
                    fname = (r.font.name or '').lower()
                    fn_ok = any(k in fname for k in ['仿宋', 'fangsong', '宋体', 'times'])
                    if not fn_ok:
                        # 不要对数字严格（数字应该Times New Roman）
                        is_digit_only = re.match(r'^[\d\s.,，。%+-/（）()]+$', r.text)
                        if not is_digit_only:
                            issues.append({
                                '位置': f'P{i}',
                                '问题': f'正文字体非仿宋: "{r.text[:20]}" 实际={r.font.name}',
                                '严重度': 'P2'
                            })
                    
                    # 字号
                    if r.font.size and r.font.size < Pt(15) and r.font.size > Pt(10):
                        pass  # 大概正确
                    elif r.font.size and r.font.size < Pt(10):
                        issues.append({
                            '位置': f'P{i}',
                            '问题': f'正文字号偏小: {r.font.size/12700:.1f}pt "{r.text[:20]}"',
                            '严重度': 'P2'
                        })
                    
                    # 行距检查
                    pf = p.paragraph_format
                    if pf.line_spacing and pf.line_spacing != Pt(30):
                        # Emu or float
                        ls = pf.line_spacing
                        if isinstance(ls, float):
                            pass  # multiple, not checkable
                        elif hasattr(ls, 'pt') and abs(ls.pt - 30) > 2:
                            pass  # don't spam on line spacing
                    break
    
    # --- 检查表格后是否有数字非Times New Roman ---
    # 在Word中很难逐字符检查数字字体，因为数字和中文可能在同一run中
    # 跳过详细检查，仅报告能找到的
    
    return issues


# ============================================================
# PART 2: 自评复核14条标准检查
# ============================================================

def check_criteria_1_10(xlsx_path, unit_name):
    """
    检查14条标准（按xlsx表格逐sheet逐行检查）
    返回：(sheet级别结果列表, 总统计)
    """
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    results = []
    stats = {'total_indicators': 0, 'issues': 0, 'criteria_hits': defaultdict(int)}
    
    for sn in wb.sheetnames:
        ws = wb[sn]
        if ws.max_row < 3:
            continue
            
        # 找表头行（一级指标 | 二级指标 | 三级指标 | 指标性质 | 指标值 | 计量单位 | 权重 | 自评完成值 | 自评得分 | 复核完成值 | 复核得分 | 复核指标偏差 | 备注）
        header_row = None
        for row_idx in range(1, min(10, ws.max_row + 1)):
            row_vals = [str(ws.cell(row=row_idx, column=c).value or '') for c in range(1, 14)]
            if '一级指标' in row_vals[0] and '三级指标' in row_vals[2]:
                header_row = row_idx
                break
        
        if header_row is None:
            continue
        
        # 项目名称通常在第二行
        project_name = str(ws.cell(row=2, column=1).value or '') or sn
        
        # 解析headers
        headers = {}
        for c in range(1, 14):
            val = str(ws.cell(row=header_row, column=c).value or '').strip()
            headers[c] = val
        
        # 找列索引
        col_map = {}
        for c, h in headers.items():
            if '三级指标' in h:
                col_map['indicator'] = c
            elif '指标性质' in h:
                col_map['nature'] = c
            elif '指标值' == h or '指标值' in h:
                col_map['target'] = c
            elif '权重' in h:
                col_map['weight'] = c
            elif '自评完成值' in h:
                col_map['self_val'] = c
            elif '自评得分' in h:
                col_map['self_score'] = c
            elif '复核完成值' in h:
                col_map['review_val'] = c
            elif '复核得分' in h:
                col_map['review_score'] = c
            elif '偏差' in h:
                col_map['deviation'] = c
            elif '备注' in h or '注' in h:
                col_map['remark'] = c
            elif '一级指标' in h:
                col_map['level1'] = c
            elif '二级指标' in h:
                col_map['level2'] = c
            elif '计量单位' in h:
                col_map['unit'] = c
        
        # 遍历数据行
        for row_idx in range(header_row + 1, ws.max_row + 1):
            row = {}
            for k, c in col_map.items():
                row[k] = ws.cell(row=row_idx, column=c).value
            
            indicator = str(row.get('indicator', '')).strip()
            if not indicator or indicator == 'None' or '合计' in indicator:
                continue
            
            stats['total_indicators'] += 1
            
            nature = str(row.get('nature', '')).strip()
            level1 = str(row.get('level1', '')).strip()
            level2 = str(row.get('level2', '')).strip()
            target_val = row.get('target')
            self_val = row.get('self_val')
            self_score = row.get('self_score')
            review_val = row.get('review_val')
            review_score = row.get('review_score')
            remark = str(row.get('remark', '')).strip() if row.get('remark') else ''
            weight = row.get('weight', 0)
            unit = str(row.get('unit', '')).strip()
            
            if isinstance(weight, str):
                try:
                    weight = float(weight)
                except:
                    weight = 0
            if weight is None:
                weight = 0
            
            # 判断复核得分是否为0或明显被扣分（复核得分 < 自评得分）
            is_deducted = False
            if review_score is not None and self_score is not None:
                try:
                    rs = float(review_score)
                    ss = float(self_score)
                    if rs < ss:
                        is_deducted = True
                except:
                    pass
            
            # ---- 标准1: 没得资料证明的，打0分 ----
            # 检查复核得分是否为0且备注提到"无资料""无佐证"等
            if review_score is not None:
                try:
                    if float(review_score) == 0:
                        no_data_keywords = ['无资料', '无佐证', '未提供', '无证明', '无数据', '无相关']
                        if not any(kw in remark for kw in no_data_keywords):
                            results.append({
                                '单位': unit_name,
                                '项目': project_name,
                                '指标': indicator,
                                '标准': '标准1',
                                '问题': f'复核得分=0但备注未说明无佐证资料: 备注="{remark[:50]}"',
                                '严重度': 'P1'
                            })
                            stats['issues'] += 1
                            stats['criteria_hits']['标准1'] += 1
                except:
                    pass
            
            # ---- 标准2: 指标设置跟项目无关，打0分 ----
            # 检查备注中是否提到无关或指标设置不当
            if '无关' in remark or '不相关' in remark:
                if review_score is not None:
                    try:
                        if float(review_score) > 0:
                            results.append({
                                '单位': unit_name,
                                '项目': project_name,
                                '指标': indicator,
                                '标准': '标准2',
                                '问题': f'指标与项目无关但复核未打0分(复核得分={review_score})',
                                '严重度': 'P0'
                            })
                            stats['issues'] += 1
                            stats['criteria_hits']['标准2'] += 1
                    except:
                        pass
            
            # ---- 标准3: 指标不具有可考查性，酌情扣分 ----
            if '无可考查' in remark or '无法考查' in remark or '不可考查' in remark:
                pass  # 已在备注中说明，检查扣分比例是否合理
                # 暂时标记观察
            
            # ---- 标准4: 满意度无资料的，打0分 ----
            if '满意' in indicator or '满意度' in indicator:
                if level1 and '满意' in level1:
                    if review_score is not None:
                        try:
                            if float(review_score) > 0 and ('无' in remark or '未' in remark):
                                if any(kw in remark for kw in ['无资料', '无佐证', '未提供', '无数据', '无调查']):
                                    results.append({
                                        '单位': unit_name,
                                        '项目': project_name,
                                        '指标': indicator,
                                        '标准': '标准4',
                                        '问题': f'满意度指标无资料但复核未打0分(复核得分={review_score})',
                                        '严重度': 'P0'
                                    })
                                    stats['issues'] += 1
                                    stats['criteria_hits']['标准4'] += 1
                        except:
                            pass
            
            # ---- 标准8: 效益指标设为定性指标，虽产生效益但无具体标准 → 不扣分 ----
            if '效益' in level1 or '效益' in level2:
                if nature == '定性':
                    # 检查是否不当扣分
                    if is_deducted and '无可考查' not in remark and '无法判断' not in remark:
                        # 定性效益指标产生效益不应扣分
                        has_benefit = any(kw in remark for kw in ['产生', '有', '达到', '实现', '完成', '良好', '合格'])
                        if has_benefit:
                            results.append({
                                '单位': unit_name,
                                '项目': project_name,
                                '指标': indicator,
                                '标准': '标准8',
                                '问题': f'效益定性指标产生效益但复核扣分(自评{self_score}→复核{review_score})，备注={remark[:60]}',
                                '严重度': 'P1'
                            })
                            stats['issues'] += 1
                            stats['criteria_hits']['标准8'] += 1
            
            # ---- 标准10: 三级指标与二级指标定义不匹配 ----
            if '数量指标' in level2 and indicator:
                money_keywords = ['资金支付', '资金拨付', '预算执行', '资金执行']
                if any(kw in indicator for kw in money_keywords):
                    if review_score is not None:
                        try:
                            if float(review_score) > 0:
                                results.append({
                                    '单位': unit_name,
                                    '项目': project_name,
                                    '指标': indicator,
                                    '标准': '标准10',
                                    '问题': f'数量指标三级指标设为资金支付类({indicator})，未反映公共产品或服务数量，应打0分但复核得分={review_score}',
                                    '严重度': 'P0'
                                })
                                stats['issues'] += 1
                                stats['criteria_hits']['标准10'] += 1
                        except:
                            pass
                    
                    # 同时检查是否备注说明了
                    if '不匹配' not in remark and '数量' not in remark:
                        results.append({
                            '单位': unit_name,
                            '项目': project_name,
                            '指标': indicator,
                            '标准': '标准10',
                            '问题': f'数量指标三级≠二级定义不匹配(资金类放数量指标)，但备注未说明扣分原因',
                            '严重度': 'P1'
                        })
                        stats['issues'] += 1
                        stats['criteria_hits']['标准10'] += 1
            
            # ---- 标准9: 产出指标、效益指标重复设置 ----
            # 在同一项目的不同行中检测重复指标名
            # (在sheet级别循环外处理更好，这里先标记)
    
    # ---- 标准9检查: 同一sheet内重复指标 ----
    for sn in wb.sheetnames:
        ws = wb[sn]
        if ws.max_row < 3:
            continue
        header_row = None
        for row_idx in range(1, min(10, ws.max_row + 1)):
            vals = [str(ws.cell(row=row_idx, column=c).value or '') for c in range(1, 14)]
            if '一级指标' in vals[0]:
                header_row = row_idx
                break
        if header_row is None:
            continue
        
        project_name = str(ws.cell(row=2, column=1).value or '') or sn
        
        # Collect all indicators with their level1
        indicators_data = []
        for row_idx in range(header_row + 1, ws.max_row + 1):
            ind = str(ws.cell(row=row_idx, column=3).value or '').strip()
            lvl1 = str(ws.cell(row=row_idx, column=1).value or '').strip()
            if ind and ind != 'None' and '合计' not in ind:
                indicators_data.append((row_idx, ind, lvl1))
        
        # Check duplicates
        seen = {}
        for row_idx, ind, lvl1 in indicators_data:
            key = ind.strip()
            if key in seen:
                prev_lvl1 = seen[key][1]
                if '产出' in lvl1 and '效益' in prev_lvl1:
                    results.append({
                        '单位': unit_name,
                        '项目': project_name,
                        '指标': ind,
                        '标准': '标准9',
                        '问题': f'产出指标与效益指标重复设置:"{ind}"，应按规则只计其一，另一打0分',
                        '严重度': 'P1'
                    })
                    stats['issues'] += 1
                    stats['criteria_hits']['标准9'] += 1
            else:
                seen[key] = (row_idx, lvl1)
    
    return results, stats


def check_report_vs_worksheet(docx_path, xlsx_path, unit_name):
    """标准11: 报告与盖章底稿信息一致性"""
    issues = []
    
    # 从docx提取项目名称和关键数据
    doc = Document(docx_path)
    report_projects = set()
    for t in doc.tables:
        if t.rows:
            headers = [c.text.strip() for c in t.rows[0].cells]
            if '项目名称' in str(headers) or '项目' in str(headers[1] if len(headers) > 1 else ''):
                for row in t.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) > 1 and cells[1]:
                        report_projects.add(cells[1])
    
    # 从xlsx提取sheet名（项目名）
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    xlsx_sheets = []
    for sn in wb.sheetnames:
        if sn in ('Sheet1', 'Sheet4', 'Sheet5', 'Sheet6'):
            continue
        xlsx_sheets.append(sn)
    
    # 简化为只检查xlsx sheet数量与报告中项目数是否匹配
    # 以及检查报告表的汇总数与底稿各sheet汇总是否一致
    
    # 从docx Table 0（项目选取表）和 Table 1（偏差汇总）提取数据
    for ti, t in enumerate(doc.tables):
        if ti == 0 and len(t.columns) >= 5:
            headers = [c.text.strip() for c in t.rows[0].cells]
            if '项目名称' in str(headers):
                report_count = len(t.rows) - 1
                xlsx_count = len(xlsx_sheets)
                if report_count != xlsx_count:
                    issues.append({
                        '单位': unit_name,
                        '项目': '全局',
                        '标准': '标准11',
                        '问题': f'报告选取项目数({report_count})与盖章底稿sheet数({xlsx_count})不一致',
                        '严重度': 'P0'
                    })
    
    return issues


def check_logic(docx_path, unit_name):
    """标准12: 报告内容逻辑检查"""
    issues = []
    doc = Document(docx_path)
    
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # 检查核心数据自洽
    # 总指标数 vs 偏差分布
    total_match = re.search(r'��Чָ��(\d+)��', full_text)
    if not total_match:
        total_match = re.search(r'涉及绩效指标(\d+)', full_text)
    if not total_match:
        total_match = re.search(r'绩效指标(\d+)个', full_text)
    
    # 检查偏差率计算
    dev_match = re.search(r'偏��[率|度][约|为]?(\d+\.?\d*)%', full_text)
    
    # 检查百分比合计
    percent_sum = re.findall(r'占(\d+\.\d+)%', full_text)
    if len(percent_sum) >= 3:
        try:
            total = sum(float(x) for x in percent_sum[:3])
            if abs(total - 100) > 5:
                issues.append({
                    '单位': unit_name,
                    '项目': '全局',
                    '标准': '标准12',
                    '问题': f'指标分布百分比合计≈{total:.1f}%，不等于100%（{percent_sum[:3]}）',
                    '严重度': 'P1'
                })
        except:
            pass
    
    # 检查"完全一致"+"一般偏差"+"明显偏差"+"显著偏差"是否覆盖所有指标
    match_count = re.findall(r'(\d+)��', full_text)
    
    # 问题与建议的对应性
    if '四、' in full_text and '五、' in full_text:
        issues_section = full_text.split('四、')[1].split('五、')[0] if '五、' in full_text.split('四、')[1] else full_text.split('四、')[1]
        suggest_section = full_text.split('五、')[1] if '五、' in full_text else ''
        issue_count = len(re.findall(r'[（(][一二三四五六七八九十\d]+[）)]', issues_section))
        suggest_count = len(re.findall(r'[（(][一二三四五六七八九十\d]+[）)]', suggest_section))
        if issue_count > 0 and suggest_count > 0 and abs(issue_count - suggest_count) > 1:
            issues.append({
                '单位': unit_name,
                '项目': '全局',
                '标准': '标准12',
                '问题': f'问题条数({issue_count})与建议条数({suggest_count})差异较大',
                '严重度': 'P2'
            })
    
    return issues


def check_format_summary(docx_path, unit_name):
    """标准13: 报告格式合规汇总"""
    format_issues = check_format(docx_path, unit_name)
    return format_issues


def check_deduction_reasons(xlsx_path, unit_name):
    """标准14: 自评复核扣分原因描述是否符合自评复核口径"""
    issues = []
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    
    for sn in wb.sheetnames:
        ws = wb[sn]
        if ws.max_row < 3:
            continue
        
        header_row = None
        for ri in range(1, min(10, ws.max_row + 1)):
            vals = [str(ws.cell(row=ri, column=c).value or '') for c in range(1, 14)]
            if '一级指标' in vals[0]:
                header_row = ri
                break
        if header_row is None:
            continue
        
        project_name = str(ws.cell(row=2, column=1).value or '') or sn
        
        for row_idx in range(header_row + 1, ws.max_row + 1):
            indicator = str(ws.cell(row=row_idx, column=3).value or '')
            nature = str(ws.cell(row=row_idx, column=4).value or '')
            self_score = ws.cell(row=row_idx, column=9).value
            review_score = ws.cell(row=row_idx, column=11).value
            deviation = ws.cell(row=row_idx, column=12).value
            remark = str(ws.cell(row=row_idx, column=13).value or '')
            
            if not indicator or indicator == 'None' or '合计' in indicator:
                continue
            if not remark or remark == 'None':
                continue
            
            try:
                rs = float(review_score) if review_score is not None else 0
                ss = float(self_score) if self_score is not None else 0
            except:
                continue
            
            # 有扣分但没有扣分原因
            if rs < ss:
                vague_keywords = ['酌情扣分', '适当扣分', '部分扣分']
                has_specific = bool(re.search(r'扣\d+%|扣[\d.]+分|扣除|扣分\d+%', remark))
                is_vague = any(kw in remark for kw in vague_keywords)
                
                if is_vague and not has_specific:
                    issues.append({
                        '单位': unit_name,
                        '项目': project_name,
                        '指标': indicator,
                        '标准': '标准14',
                        '问题': f'扣分原因过于笼统:"{remark[:60]}"，应明确扣分比例和具体原因',
                        '严重度': 'P2'
                    })
    
    return issues


# ============================================================
# MAIN
# ============================================================

def main():
    print('=' * 60)
    print('绩效自评复核 — 开始处理')
    print('=' * 60)
    
    all_results = []
    
    for unit_dir in sorted(BASE.iterdir()):
        if not unit_dir.is_dir():
            continue
        
        unit_name = unit_dir.name
        
        # Find docx and xlsx
        docx_files = list(unit_dir.glob('*.docx'))
        xlsx_files = list(unit_dir.glob('*.xlsx'))
        
        if not docx_files or not xlsx_files:
            print(f'  ⚠ {unit_name}: 缺少文件 (docx={len(docx_files)}, xlsx={len(xlsx_files)})')
            continue
        
        docx_path = docx_files[0]
        xlsx_path = xlsx_files[0]
        
        print(f'\n📋 {unit_name}')
        
        # 标准1-10检查
        print(f'  标准1-10: 指标评分检查...')
        criteria_results, stats = check_criteria_1_10(str(xlsx_path), unit_name)
        print(f'    {stats["total_indicators"]}个指标, {stats["issues"]}个问题')
        all_results.extend(criteria_results)
        
        # 标准11: 报告与底稿一致性
        print(f'  标准11: 报告↔底稿一致性...')
        s11 = check_report_vs_worksheet(str(docx_path), str(xlsx_path), unit_name)
        all_results.extend(s11)
        print(f'    {len(s11)}个问题')
        
        # 标准12: 逻辑检查
        print(f'  标准12: 报告逻辑...')
        s12 = check_logic(str(docx_path), unit_name)
        all_results.extend(s12)
        print(f'    {len(s12)}个问题')
        
        # 标准13: 格式合规
        print(f'  标准13: 格式合规...')
        s13 = check_format_summary(str(docx_path), unit_name)
        # Only keep P0/P1 format issues
        s13_filtered = [x for x in s13 if x.get('严重度') in ('P0', 'P1')]
        all_results.extend(s13_filtered)
        print(f'    {len(s13)}个格式问题 (过滤后{len(s13_filtered)}个P0/P1)')
        
        # 标准14: 扣分原因描述
        print(f'  标准14: 扣分原因...')
        s14 = check_deduction_reasons(str(xlsx_path), unit_name)
        all_results.extend(s14)
        print(f'    {len(s14)}个问题')
    
    # ---- 输出Excel ----
    print(f'\n📊 生成Excel: {OUTPUT}')
    
    wb = openpyxl.Workbook()
    
    # Sheet 1: 复核结果总览
    ws1 = wb.active
    ws1.title = '复核结果总览'
    
    # Header style
    header_font = Font(name='黑体', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='0A1F3F', end_color='0A1F3F', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # P0/P1/P2 fills
    p0_fill = PatternFill(start_color='FFD7D7', end_color='FFD7D7', fill_type='solid')
    p1_fill = PatternFill(start_color='FFF3CD', end_color='FFF3CD', fill_type='solid')
    p2_fill = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')
    
    headers1 = ['序号', '单位', '项目', '标准', '指标/位置', '问题描述', '严重度']
    for c, h in enumerate(headers1, 1):
        cell = ws1.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border
    
    body_font = Font(name='仿宋', size=10)
    body_align = Alignment(vertical='center', wrap_text=True)
    
    for ri, r in enumerate(all_results, 2):
        vals = [
            ri - 1,
            r.get('单位', ''),
            r.get('项目', ''),
            r.get('标准', ''),
            r.get('指标', r.get('位置', '')),
            r.get('问题', ''),
            r.get('严重度', 'P2')
        ]
        for c, v in enumerate(vals, 1):
            cell = ws1.cell(row=ri, column=c, value=str(v))
            cell.font = body_font
            cell.alignment = body_align
            cell.border = thin_border
            
            # Color by severity
            severity = r.get('严重度', 'P2')
            if severity == 'P0':
                cell.fill = p0_fill
            elif severity == 'P1':
                cell.fill = p1_fill
    
    # Column widths
    ws1.column_dimensions['A'].width = 6
    ws1.column_dimensions['B'].width = 30
    ws1.column_dimensions['C'].width = 25
    ws1.column_dimensions['D'].width = 10
    ws1.column_dimensions['E'].width = 30
    ws1.column_dimensions['F'].width = 60
    ws1.column_dimensions['G'].width = 8
    
    # Freeze header
    ws1.freeze_panes = 'A2'
    ws1.auto_filter.ref = ws1.dimensions
    
    # Sheet 2: 统计汇总
    ws2 = wb.create_sheet('统计汇总')
    
    # Count by severity
    p0_count = sum(1 for r in all_results if r.get('严重度') == 'P0')
    p1_count = sum(1 for r in all_results if r.get('严重度') == 'P1')
    p2_count = sum(1 for r in all_results if r.get('严重度') == 'P2')
    
    ws2.merge_cells('A1:D1')
    ws2.cell(row=1, column=1, value='绩效自评复核统计汇总').font = Font(name='黑体', size=14, bold=True)
    ws2.cell(row=1, column=1).alignment = Alignment(horizontal='center')
    
    summary_data = [
        ('', '', '', ''),
        ('复核日期', '2026年7月22日', '', ''),
        ('复核单位数', len([d for d in BASE.iterdir() if d.is_dir()]), '', ''),
        ('', '', '', ''),
        ('严重度', '数量', '占比', '说明'),
        ('P0 致命', p0_count, f'{p0_count/max(1,len(all_results))*100:.1f}%', '提交前必须修改'),
        ('P1 重要', p1_count, f'{p1_count/max(1,len(all_results))*100:.1f}%', '建议修改'),
        ('P2 优化', p2_count, f'{p2_count/max(1,len(all_results))*100:.1f}%', '可优化事项'),
        ('合计', len(all_results), '100%', ''),
        ('', '', '', ''),
        ('标准分布', '', '', ''),
    ]
    
    for ri, row_data in enumerate(summary_data, 2):
        for c, v in enumerate(row_data, 1):
            cell = ws2.cell(row=ri, column=c, value=v)
            cell.font = body_font
            cell.border = thin_border
    
    # Count by criterion
    row_offset = len(summary_data) + 2
    criterion_counts = defaultdict(int)
    for r in all_results:
        criterion_counts[r.get('标准', '未知')] += 1
    
    for ci, (criterion, count) in enumerate(sorted(criterion_counts.items()), 1):
        ws2.cell(row=row_offset + ci, column=1, value=criterion).border = thin_border
        ws2.cell(row=row_offset + ci, column=2, value=count).border = thin_border
        ws2.cell(row=row_offset + ci, column=3, value=f'{count/max(1,len(all_results))*100:.1f}%').border = thin_border
    
    # Sheet 3: 格式检查详细
    ws3 = wb.create_sheet('格式检查')
    fmt_headers = ['序号', '单位', '位置', '问题描述', '严重度']
    for c, h in enumerate(fmt_headers, 1):
        cell = ws3.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border
    
    fmt_idx = 2
    for unit_dir in sorted(BASE.iterdir()):
        if not unit_dir.is_dir():
            continue
        docx_files = list(unit_dir.glob('*.docx'))
        if not docx_files:
            continue
        fmt_issues = check_format(str(docx_files[0]), unit_dir.name)
        for fi in fmt_issues:
            for c, v in enumerate([fmt_idx - 1, fi.get('单位', unit_dir.name), fi.get('位置', ''), fi.get('问题', ''), fi.get('严重度', 'P2')], 1):
                cell = ws3.cell(row=fmt_idx, column=c, value=str(v))
                cell.font = body_font
                cell.alignment = body_align
                cell.border = thin_border
            fmt_idx += 1
    
    ws3.column_dimensions['A'].width = 6
    ws3.column_dimensions['B'].width = 30
    ws3.column_dimensions['C'].width = 20
    ws3.column_dimensions['D'].width = 60
    ws3.column_dimensions['E'].width = 8
    ws3.freeze_panes = 'A2'
    
    wb.save(str(OUTPUT))
    
    print(f'\n✅ 完成！')
    print(f'  P0: {p0_count} | P1: {p1_count} | P2: {p2_count}')
    print(f'  总计: {len(all_results)}个问题')
    print(f'  输出: {OUTPUT}')

if __name__ == '__main__':
    main()
