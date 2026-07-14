# -*- coding: utf-8 -*-
"""四川融策宣传册 v4 清新风 —— 完整版"""
from __future__ import annotations
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from PIL.Image import LANCZOS
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

OUT = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
WK = Path("work/sichuan_rongce_brochure/v4_assets")
DOCX = OUT / "四川融策宣传册_完善稿_v4_清新风.docx"
PDF = OUT / "四川融策宣传册_完善稿_v4_清新风.pdf"

NW = "#1A365D"  # navy
TL = "#3A7B8A"  # teal
GD = "#D4A574"  # gold
MU = "#718096"  # muted
IK = "#2D3748"  # ink
WH = "#FFFFFF"  # white
BG = "#FAF8F4"  # warm white
W, H = 1653, 2339

def rgb(s):
    h = s.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
def rga(s, a=255):
    return (*rgb(s), a)

def font(size, bold=False):
    for p in [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()

def draw(d, text, x, y, f, fill, w=None, gap=6, align="left"):
    mw = w or 9999
    lines = []
    for para in text.split("\n"):
        cur = ""
        for ch in para:
            if d.textlength(cur + ch, font=f) <= mw:
                cur += ch
            else:
                if cur: lines.append(cur)
                cur = ch
        if cur: lines.append(cur)
    for line in lines:
        xx = x
        if align == "center" and w:
            xx = x + (w - d.textlength(line, font=f)) // 2
        d.text((xx, y), line, font=f, fill=fill)
        y += f.size + gap
    return y

def left_bar(d):
    d.rectangle([0, 0, 18, H], fill=rga(NW))
    d.rectangle([18, 0, 24, H], fill=rga(GD))

def bottom_bar(d, txt=None):
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    tx = txt or "\u516c\u5f00\u516c\u6b63  \u7528\u5fc3\u670d\u52a1  \u8bda\u4fe1\u4e3a\u672c  \u670d\u52a1\u81f3\u4e0a  \u8ffd\u6c42\u5353\u8d8a"
    d.text(((W-d.textlength(tx, font=font(18)))//2, H-42), tx, font=font(18), fill=rga(WH, 180))

def page_header(d, title, en, y=120):
    d.text((160, y), title, font=font(52, True), fill=NW)
    d.text((160, y+68), en, font=font(22), fill=rga(GD))
    d.line([160, y+105, 600, y+105], fill=rga(GD), width=3)

def page_num_box(d, n, x=880, y=120):
    d.rounded_rectangle([x, y, x+570, 580], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((x+80, y+60), f"0{n}", font=font(80, True), fill=rga(TL, 60))
    d.text((x+80, y+160), "RONGCE", font=font(40, True), fill=rga(TL, 40))

# =========== P1: Cover ===========
def p1_cover():
    img = Image.new("RGB", (W, H), rgb(BG))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, 280], fill=rga(NW))
    d.rectangle([0, 280, W, 290], fill=rga(GD))
    d.ellipse([W-480, -120, W-80, 280], fill=rga(TL, 40))
    d.text((130, 70), "SICHUAN", font=font(80, True), fill=rga(WH, 200))
    d.text((130, 155), "RONGCE", font=font(108, True), fill=WH)
    d.text((130, 340), "\u8c0b\u4e13\u4e1a\u4e4b\u7b56  \u878d\u54c1\u8d28\u4e4b\u7cbe", font=font(48, True), fill=NW)
    d.text((130, 420), "\u653f\u5e9c\u5ba1\u8ba1\u4e0e\u5de5\u7a0b\u54a8\u8be2\u7efc\u5408\u670d\u52a1\u673a\u6784", font=font(30), fill=MU)
    d.line([130, 500, 500, 500], fill=rga(GD), width=5)
    d.rectangle([0, H-60, W, H], fill=rga(NW))
    d.text((130, H-45), "\u5ba1\u8ba1 \u00b7 \u7ee9\u6548 \u00b7 \u8d22\u653f\u76d1\u7763 \u00b7 \u5de5\u7a0b\u54a8\u8be2 \u00b7 \u6570\u5b57\u5316\u5206\u6790", font=font(20), fill=rga(WH, 200))
    return img

# =========== P2: About ===========
def p2_about():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_header(d, "\u5173\u4e8e\u878d\u7b56", "ABOUT US")
    # Pain point opening
    y = draw(d, "\u8d22\u653f\u8d44\u91d1\u89c4\u6a21\u5927\u3001\u9879\u76ee\u7c7b\u578b\u591a\u3001\u7ba1\u7406\u94fe\u6761\u957f\u2014\u2014\n\u5982\u4f55\u786e\u4fdd\u6bcf\u4e00\u5206\u94b1\u90fd\u82b1\u5728\u5200\u53e3\u4e0a\uff1f", 160, 290, font(36), IK, 620, 12)
    lines = [
        "\u878d\u7b56\u7684\u56de\u7b54\uff1a\u7528\u4e13\u4e1a\u5224\u65ad\u3001\u6570\u636e\u65b9\u6cd5\u548c\u73b0\u573a\u7ecf\u9a8c\uff0c",
        "\u628a\u5ba1\u8ba1\u4ece\u201c\u67e5\u8d26\u201d\u5347\u7ea7\u4e3a\u201c\u6cbb\u7406\u4f53\u68c0\u201d\u3002",
        "\u6211\u4eec\u4e0d\u53ea\u5173\u6ce8\u201c\u6709\u6ca1\u6709\u53d1\u7968\u201d\uff0c",
        "\u66f4\u5173\u6ce8\u201c\u5408\u4e0d\u5408\u7406\u201d\u3001\u201c\u503c\u4e0d\u503c\u5f97\u201d\u3001\u201c\u80fd\u4e0d\u80fd\u66f4\u597d\u201d\u3002"
    ]
    for line in lines:
        y = draw(d, line, 160, y, font(26), MU, 620, 6)
    # Data cards
    for i, (num, label) in enumerate([("20+", "\u5e74\u79ef\u7d2f"), ("400+", "\u5bb6\u5ba2\u6237"), ("50\u4eba", "\u56e2\u961f"), ("3\u7701", "\u8986\u76d6")]):
        x = 880 + i * 153
        d.rounded_rectangle([x, 160, x+135, 280], radius=12, fill=rga(TL, 20), outline=rga(TL, 60), width=2)
        tw = d.textlength(num, font=font(38, True))
        d.text((x+(135-tw)//2, 187), num, font=font(38, True), fill=TL)
        d.text((x+(135-d.textlength(label, font=font(18)))//2, 240), label, font=font(18), fill=MU)
    y = max(y, 450) + 30
    for item in [
        "\u59cb\u4e8e 2000 \u5e74\uff0c\u56db\u5ddd\u7701\u5185\u8f83\u65e9\u6210\u7acb\u7684\u4f1a\u8ba1\u5e08\u4e8b\u52a1\u6240\u4e4b\u4e00",
        "\u957f\u671f\u670d\u52a1\u8d22\u653f\u3001\u5ba1\u8ba1\u3001\u6559\u80b2\u3001\u6c11\u653f\u3001\u4ea4\u901a\u3001\u56fd\u8d44\u3001\u533b\u4fdd\u7b49\u9886\u57df",
        "\u5f62\u6210\u5ba1\u8ba1\u3001\u7ee9\u6548\u3001\u8d22\u653f\u76d1\u7763\u3001\u5de5\u7a0b\u54a8\u8be2\u534f\u540c\u53d1\u5c55\u7684\u4e1a\u52a1\u683c\u5c40",
        "\u575a\u6301\u5ba2\u89c2\u516c\u6b63\u3001\u5b9e\u4e8b\u6c42\u662f\u3001\u8d28\u91cf\u4f18\u5148\uff0c\u91cd\u89c6\u8bc1\u636e\u95ed\u73af\u548c\u6574\u6539\u843d\u5730"
    ]:
        d.ellipse([160, y+6, 172, y+18], fill=rga(GD))
        y = draw(d, item, 190, y, font(26), IK, 600, 12)
    bottom_bar(d)
    return img

# =========== P3: Method ===========
def p3_method():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_header(d, "\u6838\u5fc3\u65b9\u6cd5", "OUR APPROACH")
    # Three lines
    for i, (t, txt) in enumerate([
        ("\u8d44\u91d1\u7ebf", "\u6bcf\u4e00\u7b14\u94b1\u4ece\u54ea\u91cc\u6765\u3001\n\u5230\u54ea\u91cc\u53bb\u3001\u82b1\u5f97\u503c\u4e0d\u503c\u3002\n\u8d44\u91d1\u6d41\u3001\u652f\u4ed8\u6d41\u3001\n\u7968\u636e\u6d41\u4ea4\u53c9\u6838\u9a8c\u3002"),
        ("\u9879\u76ee\u7ebf", "\u4ece\u7acb\u9879\u3001\u62db\u6807\u3001\u5408\u540c\u3001\n\u65bd\u5de5\u5230\u9a8c\u6536\u3001\u7ed3\u7b97\uff0c\n\u5168\u94fe\u6761\u7a7f\u900f\u6838\u67e5\u3002"),
        ("\u8d23\u4efb\u7ebf", "\u51b3\u7b56\u7a0b\u5e8f\u3001\u5c97\u4f4d\u804c\u8d23\u3001\n\u5185\u63a7\u5236\u5ea6\u3001\u6574\u6539\u843d\u5b9e\u3002\n\u628a\u95ee\u9898\u8ffd\u6eaf\u5230\u4eba\u3002"),
    ]):
        x = 160 + i * 380
        d.rounded_rectangle([x, 280, x+340, 520], radius=16, fill=rga(WH), outline=rga(TL, 60), width=2)
        d.rounded_rectangle([x, 280, x+340, 380], radius=16, fill=rga(TL, 30))
        d.text((x+32, 310), t, font=font(32, True), fill=NW)
        draw(d, txt, x+32, 410, font(24), MU, 340-64, 8)
    # Five steps
    steps = ["\u5236\u5ea6\u5ba1\u67e5", "\u6570\u636e\u6838\u9a8c", "\u73b0\u573a\u6838\u67e5", "\u8bc1\u636e\u95ed\u73af", "\u6574\u6539\u5efa\u8bae"]
    for i, step in enumerate(steps):
        x = 200 + i * 270
        d.ellipse([x, 600, x+80, 680], fill=rga(NW if i%2==0 else TL))
        tw = d.textlength(str(i+1), font=font(32, True))
        d.text((x+40-tw//2, 618), str(i+1), font=font(32, True), fill=WH)
        draw(d, step, x+100, 618, font(26), IK, 150, 6)
        if i < 4:
            d.line([x+80, 640, x+190, 640], fill=rga(GD), width=3)
    # Quote box
    d.rounded_rectangle([160, 780, 1490, 920], radius=14, fill=rga("#E8E0D4", 80))
    draw(d, "\u6211\u4eec\u5173\u6ce8\u7684\u4e0d\u662f\u51fa\u5177\u4e00\u4efd\u62a5\u544a\uff0c\u800c\u662f\u5e2e\u52a9\u59d4\u6258\u65b9\u770b\u6e05\u95ee\u9898\u3001\n\u5398\u6e05\u8d23\u4efb\u3001\u627e\u5230\u6539\u8fdb\u8def\u5f84\u3002", 200, 820, font(28), IK, 1250, 10, "center")
    bottom_bar(d)
    return img

# =========== Service page template ===========
def service_page(title, en, sub, items, n):
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_header(d, title, en, 100)
    draw(d, sub, 160, 240, font(30), MU, 620, 10)
    page_num_box(d, n)
    # Decorative taglines in the right box
    taglines = {
        1: ["\u4ece\u5355\u9879\u5230\u5168\u8fc7\u7a0b", "\u4ece\u7ed3\u679c\u5230\u5236\u5ea6"],
        2: ["\u4ece\u6709\u6ca1\u6709\u5230\u503c\u4e0d\u503c", "\u8d26\u8868\u5408\u540c\u9879\u76ee\u8d2f\u901a"],
        3: ["\u4ece\u82b1\u591a\u5c11\u5230\u6548\u679c\u600e\u6837", "\u8ba9\u8d22\u653f\u8d44\u91d1\u66f4\u503c"],
        4: ["\u4ece\u6982\u7b97\u5230\u7ed3\u7b97\u5168\u5173", "\u5de5\u7a0b\u9020\u4ef7\u4e0e\u7ee9\u6548\u534f\u540c"],
        5: ["\u6570\u636e\u6269\u5927\u8986\u76d6\u9762", "\u7ecf\u9a8c\u6c89\u6dc0\u4e3a\u53ef\u590d\u7528\u5de5\u5177"],
    }.get(n, [])
    for i, t in enumerate(taglines):
        draw(d, t, 1020, 360 + i*60, font(28, True), rga(TL, 80), 400, 6)
    # Service items
    y = 330
    for i, item in enumerate(items):
        clr = rga(NW if i%2==0 else TL)
        d.rounded_rectangle([160, y, 1580, y+48], radius=8, fill=rga(WH), outline=rga(GD, 80), width=1)
        d.rectangle([160, y, 168, y+48], fill=clr)
        y = draw(d, item[0], 185, y+8, font(26, True), IK, 1350, 6) + 10
        for sub in item[1]:
            d.ellipse([185, y+8, 193, y+16], fill=rga(GD))
            y = draw(d, sub, 205, y, font(22), MU, 1300, 6) + 10
    bottom_bar(d)
    return img

def p4_services():
    return service_page("\u670d\u52a1\u4f53\u7cfb", "SERVICE SYSTEM",
        "\u4ece\u5355\u9879\u5ba1\u8ba1\u5230\u5168\u8fc7\u7a0b\u54a8\u8be2\uff0c\u4ece\u7ed3\u679c\u8bc4\u4ef7\u5230\u5236\u5ea6\u4f18\u5316\uff0c\u4ece\u4eba\u5de5\u6838\u67e5\u5230\u6570\u636e\u5316\u8bc6\u522b\u3002",
        [("\u653f\u5e9c\u5ba1\u8ba1", ["\u7ecf\u6d4e\u8d23\u4efb\u5ba1\u8ba1", "\u4e13\u9879\u8d44\u91d1\u5ba1\u8ba1", "\u8d22\u653f\u76d1\u7763\u68c0\u67e5", "\u5de5\u7a0b\u51b3\u7b97\u8d22\u52a1\u5ba1\u8ba1"]),
         ("\u9884\u7b97\u7ee9\u6548\u7ba1\u7406", ["\u4e8b\u524d\u8bc4\u4f30 \u00b7 \u76ee\u6807\u5ba1\u6838", "\u8fd0\u884c\u76d1\u63a7 \u00b7 \u91cd\u70b9\u8bc4\u4ef7", "\u7ed3\u679c\u5e94\u7528"]),
         ("\u5de5\u7a0b\u54a8\u8be2", ["\u9884\u7b97\u7f16\u5236 \u00b7 \u8d22\u653f\u8bc4\u5ba1", "\u7ed3\u7b97\u5ba1\u6838 \u00b7 \u5168\u8fc7\u7a0b\u5de5\u7a0b\u54a8\u8be2"]),
         ("\u91c7\u8d2d\u5ba1\u8ba1", ["\u91c7\u8d2d\u7a0b\u5e8f\u5408\u89c4", "\u56f4\u6807\u4e32\u6807\u7ebf\u7d22", "\u5408\u540c\u5c65\u7ea6\u6838\u67e5"]),
         ("\u7ba1\u7406\u54a8\u8be2", ["\u5185\u63a7\u5efa\u8bbe \u00b7 \u8d44\u4ea7\u7ba1\u7406", "\u6574\u6539\u63d0\u5347 \u00b7 \u6d41\u7a0b\u4f18\u5316"])], 1)

def p5_gov():
    return service_page("\u653f\u5e9c\u5ba1\u8ba1", "GOVERNMENT AUDIT",
        "\u4ece\u201c\u6709\u6ca1\u6709\u201d\u5230\u201c\u5bf9\u4e0d\u5bf9\u201d\u518d\u5230\u201c\u503c\u4e0d\u503c\u201d\u2014\u2014\u628a\u8d26\u8868\u3001\u5408\u540c\u3001\u9879\u76ee\u3001\u8d44\u91d1\u3001\u8d44\u4ea7\u3001\u8d23\u4efb\u8d2f\u901a\u6838\u67e5\u3002",
        [("\u7ecf\u6d4e\u8d23\u4efb\u5ba1\u8ba1", ["\u91cd\u5927\u51b3\u7b56\u4e0e\u8d44\u91d1\u8d44\u4ea7", "\u9879\u76ee\u5efa\u8bbe\u4e0e\u5185\u63a7\u98ce\u9669", "\u5ec9\u653f\u98ce\u9669\u6392\u67e5"]),
         ("\u4e13\u9879\u8d44\u91d1\u5ba1\u8ba1", ["\u7533\u62a5\u5206\u914d\u3001\u62e8\u4ed8\u4f7f\u7528", "\u7ee9\u6548\u4e0e\u7ed3\u4f59\u6c89\u6dc0\u5168\u94fe\u6761"]),
         ("\u8d22\u653f\u76d1\u7763\u68c0\u67e5", ["\u9884\u7b97\u6267\u884c\u3001\u8d22\u7ecf\u7eaa\u5f8b", "\u653f\u5e9c\u91c7\u8d2d\u4e0e\u4f1a\u8ba1\u4fe1\u606f\u8d28\u91cf"]),
         ("\u5de5\u7a0b\u51b3\u7b97\u8d22\u52a1\u5ba1\u8ba1", ["\u5efa\u8bbe\u6210\u672c\u5f52\u96c6", "\u8d44\u91d1\u6765\u6e90\u6838\u5b9e", "\u8d44\u4ea7\u4ea4\u4ed8\u4e0e\u5c3e\u5de5\u5c3e\u6b3e"])], 2)

def p6_perf():
    return service_page("\u9884\u7b97\u7ee9\u6548\u7ba1\u7406", "BUDGET PERFORMANCE",
        "\u60a8\u5173\u5fc3\u7684\u4e0d\u53ea\u662f\u201c\u82b1\u4e86\u591a\u5c11\u94b1\u201d\uff0c\u66f4\u662f\u201c\u6548\u679c\u600e\u4e48\u6837\u201d\u2014\u2014\u8ba9\u8d22\u653f\u8d44\u91d1\u4ece\u201c\u82b1\u4e86\u6ca1\u6709\u201d\u8d70\u5411\u201c\u82b1\u5f97\u503c\u4e0d\u503c\u201d\u3002",
        [("\u4e8b\u524d\u7ee9\u6548\u8bc4\u4f30", ["\u5fc5\u8981\u6027\u3001\u53ef\u884c\u6027\u5206\u6790", "\u8d22\u653f\u627f\u53d7\u80fd\u529b\u4e0e\u9884\u671f\u7ee9\u6548"]),
         ("\u7ee9\u6548\u76ee\u6807\u5ba1\u6838", ["\u76ee\u6807\u5b8c\u6574\u6027\u3001\u6307\u6807\u53ef\u8861\u91cf\u6027", "\u9884\u7b97\u5339\u914d\u6027\u5ba1\u6838"]),
         ("\u7ee9\u6548\u8fd0\u884c\u76d1\u63a7", ["\u6267\u884c\u8fdb\u5ea6\u3001\u8d44\u91d1\u652f\u4ed8", "\u4ea7\u51fa\u504f\u5dee\u4e0e\u98ce\u9669\u9884\u8b66"]),
         ("\u91cd\u70b9\u7ee9\u6548\u8bc4\u4ef7", ["\u653f\u7b56\u8bc4\u4ef7\u3001\u90e8\u95e8\u6574\u4f53\u8bc4\u4ef7", "\u9879\u76ee\u652f\u51fa\u4e0e\u4e13\u9879\u8d44\u91d1\u8bc4\u4ef7"]),
         ("\u7ed3\u679c\u5e94\u7528", ["\u6574\u6539\u6e05\u5355\u4e0e\u9884\u7b97\u6302\u94a9", "\u7ba1\u7406\u5236\u5ea6\u4f18\u5316\u5efa\u8bae"])], 3)

def p7_eng():
    return service_page("\u5de5\u7a0b\u54a8\u8be2\u4e0e\u8d22\u653f\u8bc4\u5ba1", "ENGINEERING CONSULTING",
        "\u4ece\u6982\u7b97\u5230\u7ed3\u7b97\uff0c\u5e2e\u60a8\u628a\u597d\u6bcf\u4e00\u9053\u5173\u2014\u2014\u5de5\u7a0b\u9020\u4ef7\u3001\u5408\u540c\u5c65\u7ea6\u3001\u8d44\u91d1\u652f\u4ed8\u548c\u9879\u76ee\u7ee9\u6548\u534f\u540c\u7ba1\u7406\u3002",
        [("\u9884\u7b97\u7f16\u5236\u4e0e\u8d22\u653f\u8bc4\u5ba1", ["\u5de5\u7a0b\u91cf\u3001\u5b9a\u989d\u5957\u7528\u6838\u9a8c", "\u6750\u6599\u4ef7\u683c\u3001\u63aa\u65bd\u8d39\u3001\u53d6\u8d39\u6807\u51c6"]),
         ("\u6e05\u5355\u53ca\u62db\u6807\u63a7\u5236\u4ef7", ["\u63d0\u5347\u62db\u6807\u6587\u4ef6\u548c\u63a7\u5236\u4ef7\u7f16\u5236\u8d28\u91cf"]),
         ("\u7ed3\u7b97\u5ba1\u6838", ["\u5408\u540c\u6761\u6b3e\u3001\u53d8\u66f4\u7b7e\u8bc1\u6838\u67e5", "\u9690\u853d\u5de5\u7a0b\u3001\u73b0\u573a\u5de5\u7a0b\u91cf\u6838\u9a8c"]),
         ("\u5168\u8fc7\u7a0b\u5de5\u7a0b\u54a8\u8be2", ["\u9879\u76ee\u524d\u671f\u3001\u62db\u91c7\u3001\u5b9e\u65bd", "\u9a8c\u6536\u3001\u7ed3\u7b97\u3001\u7ee9\u6548\u8bc4\u4ef7\u534f\u540c"])], 4)

def p8_digital():
    return service_page("\u6570\u5b57\u5316\u5ba1\u8ba1\u80fd\u529b", "DIGITAL AUDIT CAPABILITIES",
        "\u7528\u6570\u636e\u6269\u5927\u8986\u76d6\u9762\u3001\u63d0\u9ad8\u53d1\u73b0\u7387\u3001\u589e\u5f3a\u8bc1\u636e\u8d28\u91cf\u2014\u2014\u628a\u5ba1\u8ba1\u7ecf\u9a8c\u6c89\u6dc0\u4e3a\u53ef\u590d\u7528\u7684\u6570\u636e\u5de5\u5177\u3002",
        [("\u6570\u636e\u6807\u51c6", ["\u8d22\u52a1\u3001\u9884\u7b97\u3001\u652f\u4ed8\u3001\u5408\u540c", "\u91c7\u8d2d\u3001\u8d44\u4ea7\u3001\u5de5\u7a0b\u9879\u76ee\u5b57\u6bb5\u6574\u7406"]),
         ("\u89c4\u5219\u6a21\u578b", ["\u91cd\u590d\u652f\u4ed8\u3001\u8d85\u9884\u7b97\u6267\u884c\u8bc6\u522b", "\u4f9b\u5e94\u5546\u5f02\u5e38\u3001\u8d44\u91d1\u6c89\u6dc0\u68c0\u6d4b"]),
         ("\u7a7f\u900f\u6838\u67e5", ["\u7591\u70b9\u6765\u6e90\u3001\u6838\u67e5\u8def\u5f84", "\u4f50\u8bc1\u6750\u6599\u3001\u5f71\u54cd\u91d1\u989d\u3001\u6574\u6539\u5efa\u8bae"]),
         ("\u62a5\u544a\u590d\u6838", ["\u91d1\u989d\u6c47\u603b\u6821\u9a8c\u3001\u53e3\u5f84\u4e00\u81f4\u6027", "\u9644\u8868\u95ed\u73af\u3001\u7ed3\u8bba\u4f9d\u636e\u53ef\u8ffd\u6eaf"])], 5)

# =========== P9: Experience (improved) ===========
def p9_experience():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    pag
def p9_experience():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_header(d, "代表经验", "REPRESENTATIVE EXPERIENCE")
    draw(d, "长期服务省、市、县多级财政和主管部门，覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。", 160, 240, font(30), MU, 620, 10)
    # Service territory box
    d.rounded_rectangle([880, 130, 1450, 550], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((960, 180), "服务版图", font=font(36, True), fill=rga(TL, 60))
    d.text((960, 230), "SERVICE TERRITORY", font=font(18), fill=rga(TL, 40))
    # Territory items with decorative dots and lines
    terr = [("成都总部", "四川省会，核心枢纽"), ("阿坝州办事处", "川西高原，覆盖藏区"),
            ("西藏办事处", "高原地区拓展"), ("覆盖川、藏、黔", "三省联动，跨区域服务")]
    for i, (name, desc) in enumerate(terr):
        y0 = 280 + i * 60
        d.ellipse([960, y0, 990, y0+30], fill=rga(GD))
        draw(d, name, 1010, y0, font(26, True), IK, 300, 2)
        draw(d, desc, 1010, y0+28, font(18), MU, 300, 2)
        if i < 3:
            d.line([975, y0+30, 975, y0+60], fill=rga(GD, 100), width=2)
    # Clients section
    d.text((160, 390), "代表客户", font=font(32, True), fill=NW)
    d.text((160, 428), "REPRESENTATIVE CLIENTS", font=font(18), fill=rga(GD))
    clients = ["四川省财政厅", "省公安厅交警总队", "省民政厅", "省农业农村厅", "省教育厅",
        "省退役军人事务厅", "省药品监督管理局", "达州市财政局", "绵阳市财政局",
        "宜宾市财政局", "什邡市财政局", "德阳市财政局", "康定市财政局", "九寨沟县财政局"]
    yv = 470
    for i, c in enumerate(clients):
        if i % 5 == 0 and i > 0:
            yv += 40
        cx = 160 + (i % 5) * 280
        d.rounded_rectangle([cx, yv, cx+240, yv+36], radius=6, fill=rga(WH), outline=rga(GD, 80), width=1)
        tw = d.textlength(c, font=font(20))
        d.text((cx+120-tw//2, yv+5), c, font=font(20), fill=NW)
    # Industry coverage note
    d.text((160, 710), "覆盖领域：财政、审计、教育、民政、交通、国资、医保、药监、公安", font=font(22), fill=MU)
    bottom_bar(d)
    return img

# =========== P10: Contact (improved) ===========
def p10_contact():
    img = Image.new("RGB", (W, H), rgb(BG))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, H], fill=rga(NW))
    d.rectangle([0, 0, W, 280], fill=rga(NW))
    d.rectangle([0, 280, W, 290], fill=rga(GD))
    d.text((130, 80), "合作价值", font=font(52, True), fill=WH)
    d.text((130, 150), "VALUE PROPOSITION", font=font(22), fill=rga(GD, 200))
    d.line([130, 190, 500, 190], fill=rga(GD), width=3)
    draw(d, "我们交付的不只是一份报告，更是一套可执行的改进方案。", 130, 260, font(30), WH, 600, 10)
    # Contact card
    d.rounded_rectangle([130, 340, 750, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "联系融策", 180, 370, font(36, True), GD, 500, 8)
    for i, (icon, line) in enumerate([("\u260e", "028-87659276"), ("\u2709", "scrccpa@163.com"),
                                       ("\u25b6", "www.scrccpa.com"), ("\u2302", "成都市金牛区金周路595号")]):
        d.text((180, 440+i*48), icon, font=font(26), fill=GD)
        d.text((215, 438+i*48), line, font=font(26), fill=WH)
    # Commitment card
    d.rounded_rectangle([880, 340, 1530, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "合作承诺", 930, 380, font(36, True), GD, 500, 8)
    for i, item in enumerate(["客观公正  实事求是", "质量优先  证据闭环",
                               "问题有依据  结论可解释", "建议能落地  整改有追踪"]):
        d.ellipse([930, 460+i*48, 946, 476+i*48], fill=rga(GD))
        d.text((965, 458+i*48), item, font=font(26), fill=WH)
    # Bottom decorative elements
    d.ellipse([W-300, H-400, W-50, H-150], fill=rga(GD, 15))
    d.ellipse([W-250, H-350, W-100, H-200], fill=rga(TL, 15))
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text((130, H-40), "四川融策会计师事务所有限公司", font=font(20), fill=rga(GD, 180))
    return img

# =========== PDF Generation ===========
def generate_pdf(images):
    c = canvas.Canvas(str(PDF), pagesize=A4)
    pw, ph = A4
    for img_path in images:
        im = Image.open(img_path)
        im = im.resize((int(pw), int(ph)-20), LANCZOS)
        tmp = WK / "tmp_render.png"
        im.save(tmp)
        c.drawImage(str(tmp), 0, 10, width=pw, height=ph-20, preserveAspectRatio=False)
        tmp.unlink(missing_ok=True)
        c.showPage()
    c.save()

# =========== DOCX Generation ===========
def generate_docx(images):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    lbls = ["封面", "关于融策", "核心方法", "服务体系", "政府审计",
            "预算绩效", "工程咨询", "数字化", "代表经验", "合作价值"]
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v4清新风"
    for idx, (lbl, img_path) in enumerate(zip(lbls, images), 1):
        if idx > 1: doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(lbl)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(22); r.font.bold = True
        doc.add_picture(str(img_path), width=Cm(17.0))
    doc.save(DOCX)

# =========== Main ===========
def main():
    OUT.mkdir(parents=True, exist_ok=True)
    WK.mkdir(parents=True, exist_ok=True)
    makers = [p1_cover, p2_about, p3_method, p4_services,
              p5_gov, p6_perf, p7_eng, p8_digital,
              p9_experience, p10_contact]
    print("Rendering 10 pages...")
    images = []
    for i, maker in enumerate(makers, 1):
        out = WK / f"page_{i:02d}.png"
        if not out.exists():
            print(f"  Page {i}...")
            maker().save(out, quality=95)
        images.append(out)
    print("Generating PDF...")
    generate_pdf(images)
    print("Generating DOCX...")
    generate_docx(images)
    print(f"DONE\n{DOCX}\n{PDF}")

if __name__ == "__main__":
    main()
