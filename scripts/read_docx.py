# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os

base = r'C:\Users\scrccpa\Desktop\成本测算=护理学院\人社培训、认定重要文件'
for fname in os.listdir(base):
    if not fname.endswith('.docx'):
        continue
    print(f'\n===== {fname} =====')
    doc = Document(os.path.join(base, fname))
    for p in doc.paragraphs:
        if p.text.strip():
            try:
                print(p.text)
            except:
                pass
    for ti, t in enumerate(doc.tables):
        print(f'\n--- Table {ti+1} ---')
        for r in t.rows[:20]:
            try:
                print(' | '.join(c.text for c in r.cells))
            except:
                pass
