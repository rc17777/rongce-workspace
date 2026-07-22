# -*- coding: utf-8 -*-
import json, re
from collections import Counter
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ROOT = Path(r"C:\Users\scrccpa\Desktop\报告\2.部门预算项目绩效自评复核报告20260722\成都市郫都区投资促进局2025年度部门预算项目")
REPORT = next(ROOT.glob("*.docx"))
WORKPAPER = next(ROOT.glob("*.xlsx"))
OUT = ROOT.parent / "成都市郫都区投资促进局2025年度部门预算项目绩效自评复核结果.xlsx"

def vb(val):
    return "" if val is None else str(val).strip()

def font_info(run):
    f = run.font
    ea = run._element.rPr.rFonts.get(qn("w:eastAsia")) if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    return {"text": run.text, "font": f.name or "", "east_asia": ea or "", "size_pt": round(f.size.pt,1) if f.size else None, "bold": bool(f.bold)}

def para_info(p, idx):
    ppr = p._p.pPr
    lr = ppr.spacing.get(qn("w:lineRule")) if ppr is not None and ppr.spacing is not None else ""
    am = {WD_ALIGN_PARAGRAPH.CENTER:"居中", WD_ALIGN_PARAGRAPH.LEFT:"左对齐", WD_ALIGN_PARAGRAPH.RIGHT:"右对齐", WD_ALIGN_PARAGRAPH.JUSTIFY:"两端对齐"}
    a = am.get(p.alignment, "未显式设置")
    return {"index": idx, "text": p.text.strip(), "style": p.style.name if p.style else "", "alignment": a, "line_spacing": str(p.paragraph_format.line_spacing) if p.paragraph_format.line_spacing else "", "line_rule": lr, "runs": [font_info(r) for r in p.runs if r.text.strip()]}

doc = Document(REPORT)
paras = [para_info(p,i+1) for i,p in enumerate(doc.paragraphs) if p.text.strip()]
tabs = []
for ti, t in enumerate(doc.tables, 1):
    tabs.append({"table": ti, "rows": [[vb(c.text) for c in row.cells] for row in t.rows]})
secs = [{"pw": round(s.page_width.cm,2), "ph": round(s.page_height.cm,2)} for s in doc.sections]

wb = load_workbook(WORKPAPER, data_only=False)
sheets = []
for s in wb.worksheets:
    rows = [[vb(c.value) for c in row] for row in s.iter_rows()]
    rows = [r for r in rows if any(v not in ("","") for v in r)]
    sheets.append({"sheet": s.title, "rows": rows})

all_report_text = "\n".join(p["text"] for p in paras)
all_tab_text = "\n".join(" | ".join(r) for t in tabs for r in t["rows"])
all_wp_text = "\n".join(" | ".join(r) for s in sheets for r in s["rows"])

# Quick dump for analysis
for p in paras:
    t = p["text"]
    if any(k in t for k in ["一、","二、","三、","四、","五、","2025","得分","复核","自评","偏离","问题","建议","投资"]):
        fonts = "/".join(sorted({(r["east_asia"] or r["font"]) for r in p["runs"] if (r["east_asia"] or r["font"])}))
        print(f"[{p['index']}] align={p['alignment']} font={fonts or '--'} | {t[:130]}")
for s in sheets:
    print(f"\nSheet [{s['sheet']}] ({len(s['rows'])} rows)")
    for i,r in enumerate(s["rows"]):
        print(f"  R{i+1}: {' | '.join(r[:10])}")
for t in tabs:
    if t["table"] <= 3:
        print(f"\nTable [{t['table']}] ({len(t['rows'])} rows)")
        for r in t["rows"][:3]:
            print(f"  {' | '.join(r[:8])}")
