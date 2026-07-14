# -*- coding: utf-8 -*-
"""四川融策宣传册 v5 —— 第二部分（P7-P10 + PDF+DOCX生成）"""
from __future__ import annotations
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from PIL.Image import LANCZOS
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

OUT = Path(r"C:\Users\scrccpa\Desktop\数据化改革")
WK = Path("work/sichuan_rongce_brochure/v5_assets")
DOCX = OUT / "四川融策宣传册_完善稿_v5_第一性原理版.docx"
PDF = OUT / "四川融策宣传册_完善稿_v5_第一性原理版.pdf"

NW = "#1A365D"
TL = "#3A7B8A"
GD = "#D4A574"
MU = "#718096"
IK = "#2D3748"
WH = "#FFFFFF"
BG = "#FAF8F4"
W, H = 1653, 2339

def rgb(s):
    h = s.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
def rga(s, a=255):
    return (*rgb(s), a)

def font(size, bold=False):
    for p in [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()

def draw(d, text, x, y, f, fill, w=None, gap=6, align="left"):
    mw = w or 9999
    lines = []
    for para in text.split("\n"):
        cur = ""
        for ch in para:
            if d.textlength(cur + ch, font=f) <= mw:
                cur += ch
            else:
                if cur: lines.append(cur)
                cur = ch
        if cur: lines.append(cur)
    for line in lines:
        xx = x
        if align == "center" and w:
            xx = x + (w - d.textlength(line, font=f)) // 2
        d.text((xx, y), line, font=f, fill=fill)
        y += f.size + gap
    return y

def left_bar(d):
    d.rectangle([0, 0, 18, H], fill=rga(NW))
    d.rectangle([18, 0, 24, H], fill=rga(GD))

def bottom_bar(d):
    tx = "公开公正  用心服务  诚信为本  服务至上  追求卓越"
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text(((W-d.textlength(tx, font=font(18)))//2, H-42), tx, font=font(18), fill=rga(WH, 180))

def page_title(d, title, en, y=120):
    d.text((160, y), title, font=font(52, True), fill=NW)
    d.text((160, y+68), en, font=font(22), fill=rga(GD))
    d.line([160, y+105, 600, y+105], fill=rga(GD), width=3)

def service_page(title, en, sub, items, n):
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, title, en, 100)
    draw(d, sub, 160, 240, font(30), MU, 620, 10)
    d.rounded_rectangle([880, 120, 1450, 580], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((960, 170), f"0{n}", font=font(80, True), fill=rga(TL, 60))
    d.text((960, 270), "RONGCE", font=font(40, True), fill=rga(TL, 40))
    tags = {1:["从单项到全过程","从结果到制度"],2:["从有没有","到值不值"],
            3:["让资金从花了","走向花得值"],4:["从概算到结算","全关贯通"],
            5:["数据扩大覆盖面","经验沉淀为工具"]}.get(n,[])
    for i, t in enumerate(tags):
        draw(d, t, 1020, 360+i*60, font(28, True), rga(TL, 80), 400, 6)
    y = 330
    for i, item in enumerate(items):
        clr = rga(NW if i%2==0 else TL)
        d.rounded_rectangle([160, y, 1580, y+48], radius=8, fill=rga(WH), outline=rga(GD, 80), width=1)
        d.rectangle([160, y, 168, y+48], fill=clr)
        y = draw(d, item[0], 185, y+8, font(26, True), IK, 1350, 6) + 10
        for sub in item[1]:
            d.ellipse([185, y+8, 193, y+16], fill=rga(GD))
            y = draw(d, sub, 205, y, font(22), MU, 1300, 6) + 10
    bottom_bar(d)
    return img

# ======= P7: Engineering Consulting =======
def p7_eng():
    return service_page("工程咨询与财政评审", "ENGINEERING CONSULTING",
        "从概算到结算，帮您把好每一道关——工程造价、合同履约、资金支付和项目绩效协同管理。",
        [("预算编制与财政评审", ["工程量、定额套用核验", "材料价格、措施费、取费标准"]),
         ("清单及招标控制价", ["提升招标文件和控制价编制质量"]),
         ("结算审核", ["合同条款、变更签证核查", "隐蔽工程、现场工程量核验"]),
         ("全过程工程咨询", ["项目前期、招采、实施", "验收、结算、绩效评价协同"])], 4)

# ======= P8: Digital Audit =======
def p8_digital():
    return service_page("数字化审计能力", "DIGITAL AUDIT CAPABILITIES",
        "用数据扩大覆盖面、提高发现率、增强证据质量——把审计经验沉淀为可复用的数据工具。",
        [("数据标准", ["财务、预算、支付、合同", "采购、资产、工程项目字段整理"]),
         ("规则模型", ["重复支付、超预算执行识别", "供应商异常、资金沉淀检测"]),
         ("穿透核查", ["疑点来源、核查路径", "佐证材料、影响金额、整改建议"]),
         ("报告复核", ["金额汇总校验、口径一致性", "附表闭环、结论依据可追溯"])], 5)

# ======= P9: Experience =======
def p9_experience():
    img = Image.new("RGB", (W, H), rgb(WH))
    d = ImageDraw.Draw(img, "RGBA")
    left_bar(d)
    page_title(d, "代表经验", "REPRESENTATIVE EXPERIENCE")
    draw(d, "长期服务省、市、县多级财政和主管部门，覆盖预算绩效、财政监督、专项资金评价、工程决算和管理咨询。", 160, 240, font(30), MU, 620, 10)
    d.rounded_rectangle([880, 130, 1450, 550], radius=20, fill=rga(TL, 12), outline=rga(TL, 30), width=2)
    d.text((960, 180), "服务版图", font=font(36, True), fill=rga(TL, 60))
    d.text((960, 230), "SERVICE TERRITORY", font=font(18), fill=rga(TL, 40))
    terr = [("成都总部","四川省会，核心枢纽"),("阿坝州办事处","川西高原，覆盖藏区"),
            ("西藏办事处","高原地区拓展"),("覆盖川、藏、黔","三省联动，跨区域服务")]
    for i, (nm, dc) in enumerate(terr):
        y0 = 280 + i*60
        d.ellipse([960, y0, 990, y0+30], fill=rga(GD))
        draw(d, nm, 1010, y0, font(26, True), IK, 300, 2)
        draw(d, dc, 1010, y0+28, font(18), MU, 300, 2)
        if i < 3: d.line([975, y0+30, 975, y0+60], fill=rga(GD, 100), width=2)
    d.text((160, 390), "代表客户", font=font(32, True), fill=NW)
    d.text((160, 428), "REPRESENTATIVE CLIENTS", font=font(18), fill=rga(GD))
    clients = ["四川省财政厅","省公安厅交警总队","省民政厅","省农业农村厅","省教育厅",
        "省退役军人事务厅","省药品监督管理局","达州市财政局","绵阳市财政局",
        "宜宾市财政局","什邡市财政局","德阳市财政局","康定市财政局","九寨沟县财政局"]
    yv = 470
    for i, c in enumerate(clients):
        if i % 5 == 0 and i > 0: yv += 40
        cx = 160 + (i%5)*280
        d.rounded_rectangle([cx, yv, cx+240, yv+36], radius=6, fill=rga(WH), outline=rga(GD, 80), width=1)
        tw = d.textlength(c, font=font(20))
        d.text((cx+120-tw//2, yv+5), c, font=font(20), fill=NW)
    d.text((160, 710), "覆盖领域：财政、审计、教育、民政、交通、国资、医保、药监、公安", font=font(22), fill=MU)
    bottom_bar(d)
    return img

# ======= P10: Contact =======
def p10_contact():
    img = Image.new("RGB", (W, H), rgb(BG))
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle([0, 0, W, H], fill=rga(NW))
    d.rectangle([0, 0, W, 280], fill=rga(NW))
    d.rectangle([0, 280, W, 290], fill=rga(GD))
    d.text((130, 80), "合作价值", font=font(52, True), fill=WH)
    d.text((130, 150), "VALUE PROPOSITION", font=font(22), fill=rga(GD, 200))
    d.line([130, 190, 500, 190], fill=rga(GD), width=3)
    draw(d, "我们交付的不只是一份报告，更是一套可执行的改进方案。", 130, 260, font(30), WH, 600, 10)
    d.rounded_rectangle([130, 340, 750, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "联系融策", 180, 370, font(36, True), GD, 500, 8)
    for i, (icon, line) in enumerate([("\u260e","028-87659276"),("\u2709","scrccpa@163.com"),
                                       ("\u25b6","www.scrccpa.com"),("\u2302","成都市金牛区金周路595号")]):
        d.text((180, 440+i*48), icon, font=font(26), fill=GD)
        d.text((215, 438+i*48), line, font=font(26), fill=WH)
    d.rounded_rectangle([880, 340, 1530, 620], radius=16, fill=rga(WH, 30), outline=rga(GD, 80), width=2)
    draw(d, "合作承诺", 930, 380, font(36, True), GD, 500, 8)
    for i, item in enumerate(["客观公正  实事求是","质量优先  证据闭环",
                               "问题有依据  结论可解释","建议能落地  整改有追踪"]):
        d.ellipse([930, 460+i*48, 946, 476+i*48], fill=rga(GD))
        d.text((965, 458+i*48), item, font=font(26), fill=WH)
    d.ellipse([W-300, H-400, W-50, H-150], fill=rga(GD, 15))
    d.ellipse([W-250, H-350, W-100, H-200], fill=rga(TL, 15))
    d.rectangle([0, H-50, W, H], fill=rga(NW))
    d.text((130, H-40), "四川融策会计师事务所有限公司", font=font(20), fill=rga(GD, 180))
    return img

# ======= PDF =======
def generate_pdf(images):
    c = canvas.Canvas(str(PDF), pagesize=A4)
    pw, ph = A4
    for img_path in images:
        im = Image.open(img_path).resize((int(pw), int(ph)-20), LANCZOS)
        tmp = WK / "tmp_render.png"
        im.save(tmp)
        c.drawImage(str(tmp), 0, 10, width=pw, height=ph-20, preserveAspectRatio=False)
        tmp.unlink(missing_ok=True)
        c.showPage()
    c.save()

# ======= DOCX =======
def generate_docx(images):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    lbls = ["封面","为什么是融策","方法论","服务体系","政府审计","预算绩效","工程咨询","数字化","代表经验","合作价值"]
    doc = Document()
    doc.core_properties.title = "四川融策宣传册完善稿v5第一性原理版"
    for idx, (lbl, img_path) in enumerate(zip(lbls, images), 1):
        if idx > 1: doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2); sec.right_margin = Cm(1.2)
        p = doc.add_paragraph()
        r = p.add_run(lbl)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(22); r.font.bold = True
        doc.add_picture(str(img_path), width=Cm(17.0))
    doc.save(DOCX)

# ======= Main =======
def main():
    OUT.mkdir(parents=True, exist_ok=True)
    WK.mkdir(parents=True, exist_ok=True)
    # Import all page functions from part1
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    from rongce_v5_part1 import p1_cover, p2_about, p3_method, p4_services, p5_gov, p6_perf
    makers = [p1_cover, p2_about, p3_method, p4_services, p5_gov, p6_perf,
              p7_eng, p8_digital, p9_experience, p10_contact]
    print("Rendering 10 pages...")
    images = []
    for i, maker in enumerate(makers, 1):
        out = WK / f"page_{i:02d}.png"
        if not out.exists():
            print(f"  Page {i}...")
            maker().save(out, quality=95)
        images.append(out)
    print("Generating PDF...")
    generate_pdf(images)
    print("Generating DOCX...")
    generate_docx(images)
    print(f"DONE\n{DOCX}\n{PDF}")

if __name__ == "__main__":
    main()