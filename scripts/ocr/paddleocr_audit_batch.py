#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
审计案例PDF批量OCR处理脚本
使用PaddleOCR + 版面分析
输出Obsidian兼容的Markdown格式
"""

import os
import sys
import re
import json
import time
from pathlib import Path
from datetime import datetime

# 设置编码
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ============================================================
# 配置区域 - 根据你的环境修改
# ============================================================

# PDF源文件目录（支持多个）
# 使用glob查找实际路径（处理编码问题）
import glob
import fnmatch

user_profile = os.path.expanduser('~')
desktop = os.path.join(user_profile, 'Desktop')

SOURCE_DIRS = []
if os.path.exists(desktop):
    for item in os.listdir(desktop):
        item_path = os.path.join(desktop, item)
        if os.path.isdir(item_path):
            count = sum(1 for r, d, f2 in os.walk(item_path) for ff in f2 if ff.lower().endswith('.pdf'))
            if 80 <= count <= 100:  # 审计观察90个，经济责任审计93个
                SOURCE_DIRS.append(item_path)
                print(f"[INFO] 选中: [{count}个] 桌面\\{item}")

if not SOURCE_DIRS:
    print("[ERROR] 未找到审计案例目录！")
    print("[INFO] 回退：取桌面含PDF最多的前两个目录")
    dirs_with_counts = []
    for item in os.listdir(desktop):
        p = os.path.join(desktop, item)
        if os.path.isdir(p):
            cnt = sum(1 for r, d, f2 in os.walk(p) for ff in f2 if ff.lower().endswith('.pdf'))
            if cnt > 0:
                dirs_with_counts.append((cnt, p, item))
    dirs_with_counts.sort(reverse=True)
    for count, p, item in dirs_with_counts[:2]:
        SOURCE_DIRS.append(p)
        print(f"[INFO] 回退选中: [{count}个] 桌面\\{item}")

# Obsidian输出目录
OUTPUT_DIR = r"C:\Users\scrccpa\Documents\Obsidian Vault\审计案例库-OCR"

# PaddleOCR配置（适配v3.x版本）
PADDLEOCR_CONFIG = {
    "lang": "ch",                 # 中文
}

# 批处理配置
BATCH_CONFIG = {
    "max_workers": 4,             # 并行处理数量（根据CPU核心数调整）
    "retry_times": 2,             # 失败重试次数
    "timeout": 300,               # 单个PDF超时（秒）
}

# ============================================================
# 安装依赖说明
# ============================================================
"""
首次运行前需要安装：

pip install paddlepaddle paddleocr pymupdf pdf2image Pillow

# 如果使用GPU（需要CUDA）：
pip install paddlepaddle-gpu paddleocr pymupdf pdf2image Pillow

# 还需要安装poppler（pdf2image依赖）：
# Windows: 下载 https://github.com/oschwartz10612/poppler-windows/releases
#          解压到 D:\poppler 并添加到PATH
# macOS: brew install poppler
# Linux: apt-get install poppler-utils
"""

# ============================================================
# OCR引擎初始化
# ============================================================

def init_ocr():
    """初始化PaddleOCR引擎"""
    try:
        from paddleocr import PaddleOCR
        print("[INFO] 正在初始化PaddleOCR引擎（首次运行会自动下载模型）...")
        print("[INFO] 模型下载约需2-5分钟，请耐心等待...")
        
        ocr = PaddleOCR(**PADDLEOCR_CONFIG)
        print("[INFO] PaddleOCR引擎初始化完成！")
        return ocr
    except ImportError:
        print("[ERROR] PaddleOCR未安装！")
        print("[INFO] 请运行: pip install paddlepaddle paddleocr pymupdf pdf2image Pillow")
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] 初始化失败: {e}")
        sys.exit(1)

# ============================================================
# PDF处理函数
# ============================================================

def get_pdf_files():
    """获取所有PDF文件"""
    pdf_files = []
    for source_dir in SOURCE_DIRS:
        if os.path.exists(source_dir):
            for root, dirs, files in os.walk(source_dir):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_files.append({
                            'path': os.path.join(root, file),
                            'name': file,
                            'folder': os.path.basename(root)
                        })
    return pdf_files

def extract_images_from_pdf(pdf_path, output_dir):
    """将PDF转换为图片"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        image_paths = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 提高分辨率以获得更好的OCR效果
            mat = fitz.Matrix(2, 2)  # 2x缩放
            pix = page.get_pixmap(matrix=mat)
            
            img_path = os.path.join(output_dir, f"page_{page_num:03d}.png")
            pix.save(img_path)
            image_paths.append(img_path)
        
        doc.close()
        return image_paths
    except Exception as e:
        print(f"[ERROR] PDF转图片失败: {e}")
        return []

def ocr_image(ocr_engine, image_path):
    """对单张图片进行OCR"""
    try:
        result = ocr_engine.ocr(image_path, cls=True)
        
        if result and result[0]:
            text_lines = []
            for line in result[0]:
                if line:
                    text = line[1][0]  # 提取文字
                    confidence = line[1][1]  # 置信度
                    if confidence > 0.7:  # 过滤低置信度结果
                        text_lines.append(text)
            return "\n".join(text_lines)
        return ""
    except Exception as e:
        print(f"[ERROR] OCR处理失败: {e}")
        return ""

def classify_by_filename(filename):
    """基于文件名进行智能分类（与v3脚本一致）"""
    
    # 经济责任审计（最高优先级）
    if any(k in filename for k in ['经济责任', '领导干部', '主要领导', '党政领导', '离任', '任中', '任期']):
        return '经济责任审计'
    
    # 预算执行审计
    if any(k in filename for k in ['预算执行', '预算编制', '决算', '财政收支', '部门预算', '预算管理']):
        return '预算执行审计'
    
    # 工程审计
    if any(k in filename for k in ['工程', '建设项目', '招投标', '竣工', '造价', '施工', '基建', '工程结算', '概算']):
        return '工程审计'
    
    # 国企审计
    if any(k in filename for k in ['国有企业', '国资', '国企改革', '国有资本', '国有资产', '市属企业', '央企', '国企领导', '境外投资']):
        return '国企审计'
    
    # 金融审计
    if any(k in filename for k in ['金融', '银行', '保险', '证券', '投融资', '债务', '专项债', '债券', '信贷', '不良资产', '金融风险', '影子银行', '普惠金融', '绿色金融']):
        return '金融审计'
    
    # 资源环境审计
    if any(k in filename for k in ['资源', '环境', '生态', '污染', '绿色', '能源', '碳中和', 'ESG', '生态文明', '海洋生态', '矿产资源', '地热']):
        return '资源环境审计'
    
    # 信息系统审计
    if any(k in filename for k in ['信息', '数据', '数字化', '智慧审计', '网络安全', '信息系统', '大数据', '人工智能', 'AI', '大模型', '数智']):
        return '信息系统审计'
    
    # 绩效审计
    if any(k in filename for k in ['绩效', '效益', '效率', '效果', '投入产出', '绩效目标', '绩效指标', '绩效监控', '绩效评价']):
        return '绩效审计'
    
    # 政策落实审计
    if any(k in filename for k in ['政策落实', '跟踪审计', '政策措施', '决策部署', '重大政策', '政策效果', '两新', '两重']):
        return '政策落实审计'
    
    # 社保民生审计
    if any(k in filename for k in ['社保', '民生', '养老', '医疗', '医保', '就业', '住房', '救助', '福利', '工伤', '共同富裕', '看病贵', '托育']):
        return '社保民生审计'
    
    # 农业农村审计
    if any(k in filename for k in ['农业', '农村', '乡村', '粮食', '耕地', '涉农', '农民', '农产品', '农业补贴', '脱贫攻坚']):
        return '农业农村审计'
    
    # 教科文卫审计
    if any(k in filename for k in ['教育', '学校', '高校', '科研', '科技', '文化', '卫生', '医院', '医疗', '基础科学', '研发', '科技创新']):
        return '教科文卫审计'
    
    # 内部审计
    if any(k in filename for k in ['内部审计', '内审', '内部控制', '风险管理', '内控', '公司治理', '风险防控', '合规管理']):
        return '内部审计'
    
    # 专项资金审计
    if any(k in filename for k in ['专项', '资金', '补助', '转移支付', '项目资金']):
        return '专项资金审计'
    
    return '其他审计'

def extract_key_findings(text):
    """提取关键审计发现"""
    findings = []
    patterns = [
        r"([一二三四五六七八九十]+[、\.].*?问题.*?[。；])",
        r"(发现.*?问题[：:]\s*.*?[。；])",
        r"(主要问题[：:]\s*.*?[。；])",
        r"(存在.*?问题[：:]\s*.*?[。；])",
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        for match in matches[:5]:
            clean = match.strip().replace('\n', ' ')
            if 20 < len(clean) < 500:
                findings.append(clean)
    
    return list(set(findings))[:10]

def extract_recommendations(text):
    """提取审计建议"""
    recommendations = []
    patterns = [
        r"(建议[：:]\s*.*?[。；])",
        r"(审计建议[：:]\s*.*?[。；])",
        r"(提出以下建议[：:]\s*.*?[。；])",
        r"(对策[：:]\s*.*?[。；])",
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        for match in matches[:5]:
            clean = match.strip().replace('\n', ' ')
            if 20 < len(clean) < 500:
                recommendations.append(clean)
    
    return list(set(recommendations))[:10]

def create_markdown(pdf_info, text, scene):
    """创建Obsidian Markdown文件"""
    findings = extract_key_findings(text)
    recommendations = extract_recommendations(text)
    
    safe_name = re.sub(r'[\\/*?:"<>|]', "_", pdf_info['name'])
    safe_name = safe_name.replace('.pdf', '.md')
    
    # 计算文本统计
    text_length = len(text)
    word_count = len(text.replace(" ", "").replace("\n", ""))
    
    md_content = f"""---
created: {datetime.now().strftime('%Y-%m-%d %H:%M')}
tags:
  - 审计案例
  - {scene}
  - OCR
source: {pdf_info['path']}
folder: {pdf_info['folder']}
ocr_engine: PaddleOCR
---

# {pdf_info['name'].replace('.pdf', '')}

## 基本信息

- **来源**: {pdf_info['path']}
- **场景分类**: {scene}
- **原始文件夹**: {pdf_info['folder']}
- **OCR时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}
- **OCR引擎**: PaddleOCR
- **文本长度**: {text_length} 字符
- **字数统计**: {word_count} 字

## 案例摘要

{text[:1000].replace(chr(10), ' ')}...

## 审计发现

"""
    
    if findings:
        for i, finding in enumerate(findings, 1):
            md_content += f"{i}. {finding}\n"
    else:
        md_content += "（未能自动提取关键发现，请手动补充）\n"
    
    md_content += "\n## 审计建议\n\n"
    
    if recommendations:
        for i, rec in enumerate(recommendations, 1):
            md_content += f"{i}. {rec}\n"
    else:
        md_content += "（未能自动提取建议，请手动补充）\n"
    
    md_content += f"""
## OCR完整文本

<details>
<summary>点击展开完整文本内容（{text_length}字符）</summary>

```
{text}
```

</details>

## 质量说明

> ⚠️ **OCR质量提示**
> 
> 本文本由PaddleOCR自动识别生成，可能存在以下问题：
> - 数字识别错误（金额、日期等关键数据请核对原文）
> - 表格结构丢失（表格内容可能变为纯文本）
> - 特殊符号识别错误
> - 部分文字漏识别或错识别
> 
> **建议**：关键数据请务必与原始PDF核对！

## 关联案例

- 

## 备注

- 适合用于大模型训练的数据场景: {scene}
- 建议标注标签: #{scene} #审计案例 #OCR
- 校对优先级: {'高' if scene in ['经济责任审计', '信息系统审计', '国企审计'] else '中'}

"""
    
    return safe_name, md_content

# ============================================================
# 主处理流程
# ============================================================

def process_single_pdf(ocr_engine, pdf_info, temp_dir):
    """处理单个PDF"""
    print(f"  处理: {pdf_info['name'][:50]}...")
    
    try:
        # 1. PDF转图片
        image_paths = extract_images_from_pdf(pdf_info['path'], temp_dir)
        if not image_paths:
            return None, "PDF转图片失败"
        
        # 2. OCR识别
        all_text = []
        for img_path in image_paths:
            text = ocr_image(ocr_engine, img_path)
            if text:
                all_text.append(text)
            # 清理临时图片
            try:
                os.remove(img_path)
            except:
                pass
        
        full_text = "\n\n".join(all_text)
        
        if not full_text.strip():
            return None, "OCR未识别到文本"
        
        # 3. 分类
        scene = classify_by_filename(pdf_info['name'])
        
        # 4. 生成Markdown
        safe_name, md_content = create_markdown(pdf_info, full_text, scene)
        
        return {
            'scene': scene,
            'safe_name': safe_name,
            'content': md_content,
            'text_length': len(full_text)
        }, None
        
    except Exception as e:
        return None, str(e)

def main():
    print("=" * 60)
    print("审计案例PDF批量OCR处理工具")
    print("引擎: PaddleOCR")
    print("=" * 60)
    
    # 初始化OCR
    ocr_engine = init_ocr()
    
    # 获取PDF文件
    print("\n[1/4] 扫描PDF文件...")
    pdf_files = get_pdf_files()
    print(f"找到 {len(pdf_files)} 个PDF文件")
    
    if not pdf_files:
        print("未找到PDF文件！")
        return
    
    # 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    temp_dir = os.path.join(OUTPUT_DIR, "_temp_images")
    os.makedirs(temp_dir, exist_ok=True)
    
    # 处理统计
    scene_stats = {}
    processed = 0
    failed = 0
    failed_files = []
    
    print("\n[2/4] 开始OCR处理...")
    print(f"提示: 共{len(pdf_files)}个文件，预计需要 {len(pdf_files)*2}~{len(pdf_files)*5} 分钟")
    print("")
    
    start_time = time.time()
    
    for i, pdf_info in enumerate(pdf_files, 1):
        # 断点续传：检查是否已处理
        safe_name = re.sub(r'[\\/*?:"\u003c\u003e|]', "_", pdf_info['name'])
        safe_name = safe_name.replace('.pdf', '.md')
        scene = classify_by_filename(pdf_info['name'])
        scene_dir = os.path.join(OUTPUT_DIR, scene)
        md_path = os.path.join(scene_dir, safe_name)
        
        if os.path.exists(md_path):
            print(f"[{i}/{len(pdf_files)}] {pdf_info['name'][:40]}... [SKIP] 已处理")
            processed += 1
            scene_stats[scene] = scene_stats.get(scene, 0) + 1
            continue
        
        print(f"[{i}/{len(pdf_files)}] {pdf_info['name'][:40]}... ", end="")
        
        result, error = process_single_pdf(ocr_engine, pdf_info, temp_dir)
        
        if result:
            # 保存Markdown（给Obsidian）
            scene_dir = os.path.join(OUTPUT_DIR, result['scene'])
            os.makedirs(scene_dir, exist_ok=True)
            
            md_path = os.path.join(scene_dir, result['safe_name'])
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(result['content'])
            
            # 同时输出Word版本
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'SimSun'
                style.font.size = Pt(10.5)
                style.paragraph_format.line_spacing = 1.5
                
                # 标题
                doc.add_heading(result['content'].split('\n')[0].replace('# ', ''), level=1)
                
                # 提取正文（去掉markdown语法）
                lines = result['content'].split('\n')
                in_front_matter = False
                for line in lines:
                    # 跳过YAML front matter
                    if line.strip() == '---':
                        in_front_matter = not in_front_matter
                        continue
                    if in_front_matter:
                        continue
                    # 跳过空行
                    if not line.strip():
                        continue
                    # 处理标题
                    if line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('1. '):
                        doc.add_paragraph(line[3:], style='List Number')
                    else:
                        # 普通段落
                        doc.add_paragraph(line)
                
                # 保存Word文档
                word_dir = os.path.join(OUTPUT_DIR, 'word', result['scene'])
                os.makedirs(word_dir, exist_ok=True)
                word_name = result['safe_name'].replace('.md', '.docx')
                word_path = os.path.join(word_dir, word_name)
                doc.save(word_path)
            except Exception as e:
                pass  # Word输出失败不影响主流程
            
            scene_stats[result['scene']] = scene_stats.get(result['scene'], 0) + 1
            processed += 1
            print(f"[OK] {result['scene']} ({result['text_length']}字)")
        else:
            failed += 1
            failed_files.append(f"{pdf_info['name']}: {error}")
            print(f"[FAIL] {error}")
        
        # 每10个文件显示进度
        if i % 10 == 0:
            elapsed = time.time() - start_time
            avg_time = elapsed / i
            remaining = avg_time * (len(pdf_files) - i)
            print(f"      进度: {i}/{len(pdf_files)} | 已用: {elapsed/60:.1f}分钟 | 预计剩余: {remaining/60:.1f}分钟")
    
    # 清理临时目录
    try:
        import shutil
        shutil.rmtree(temp_dir)
    except:
        pass
    
    # 生成分类索引
    print("\n[3/4] 生成分类索引...")
    index_content = f"""---
created: {datetime.now().strftime('%Y-%m-%d %H:%M')}
tags:
  - 审计案例库
  - 索引
  - OCR
---

# 审计案例库OCR索引

> OCR处理完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}
> OCR引擎: PaddleOCR

## 分类统计

| 审计场景 | 案例数量 |
|:---------|:---------|
"""
    
    for scene, count in sorted(scene_stats.items(), key=lambda x: x[1], reverse=True):
        index_content += f"| [[{scene}]] | {count} |\n"
    
    index_content += f"""
## 处理统计

- **成功处理**: {processed}
- **失败**: {failed}
- **总计**: {len(pdf_files)}
- **处理时间**: {(time.time()-start_time)/60:.1f} 分钟

## 失败文件列表

"""
    if failed_files:
        for f in failed_files:
            index_content += f"- ❌ {f}\n"
    else:
        index_content += "无\n"
    
    index_content += "\n## 快速导航\n\n"
    for scene in sorted(scene_stats.keys()):
        index_content += f"- [[{scene}]]\n"
    
    # 写入索引
    index_path = os.path.join(OUTPUT_DIR, "00-索引.md")
    with open(index_path, 'w', encoding='utf-8') as f:
        f.write(index_content)
    
    # 生成质量报告
    print("\n[4/4] 生成质量报告...")
    report = f"""# OCR处理质量报告

## 处理概况

- **引擎**: PaddleOCR
- **处理时间**: {(time.time()-start_time)/60:.1f} 分钟
- **平均每个文件**: {(time.time()-start_time)/len(pdf_files)/60:.1f} 分钟

## 质量注意事项

### PaddleOCR已知局限

1. **表格识别**: 表格可能识别为纯文本，行列关系丢失
2. **数字精度**: 金额数字可能识别错误，务必人工核对
3. **版面分析**: 复杂版面（分栏、图文混排）识别效果一般
4. **印章/手写**: 印章覆盖的文字、手写批注识别困难

### 校对建议

| 优先级 | 检查项 | 说明 |
|:-------|:-------|:-----|
| P0 | 金额数字 | 涉及资金的数据必须核对 |
| P0 | 日期时间 | 审计期间、报告日期等 |
| P1 | 人名职务 | 领导干部姓名、职务 |
| P1 | 单位名称 | 被审计单位名称 |
| P2 | 法规条文 | 引用的法规条款号 |
| P3 | 普通正文 | 大段描述性文字 |

## 建议后续操作

1. **人工校对**: 建议至少校对P0级内容
2. **表格重建**: 对于关键表格，手动重建结构
3. **补充元数据**: 添加案例标签、难度等级等
4. **质量评估**: 标记识别质量等级（A/B/C/D）

"""
    
    report_path = os.path.join(OUTPUT_DIR, "00-OCR质量报告.md")
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report)
    
    # 输出总结
    print("\n" + "=" * 60)
    print("OCR处理完成！")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"成功: {processed}, 失败: {failed}")
    print(f"总用时: {(time.time()-start_time)/60:.1f} 分钟")
    print("\n场景分布:")
    for scene, count in sorted(scene_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {scene}: {count}")
    print("=" * 60)
    print("\n⚠️  重要提示:")
    print("PaddleOCR识别结果需要人工校对，特别是金额数字！")
    print("请查看 00-OCR质量报告.md 了解详细注意事项。")

if __name__ == "__main__":
    main()
