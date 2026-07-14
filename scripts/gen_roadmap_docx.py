# -*- coding: utf-8 -*-
"""生成融策Agent 1+3+5+N 增强路线 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1C, 0x35, 0x5E)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x37, 0x8F, 0xCF)
    return h

def add_para(text, bold=False, size=10.5, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    return p

def add_bullet(text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.space_after = Pt(2)
    return p

def shade_cell(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, size=9, color=None, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9, color=(255, 255, 255))
        shade_cell(cell, '1C355E')
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_text(cell, str(val), size=9)
            if r % 2 == 0:
                shade_cell(cell, 'EBF0F7')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para('融策Agent', bold=True, size=32, color=(0x1C, 0x35, 0x5E), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('"1+3+5+N" 增强路线', bold=True, size=24, color=(0x2E, 0x75, 0xB6), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

doc.add_paragraph()
add_para('—— 基于融策"1+5+N"业务体系与国资委穿透式监管', size=12, color=(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('"三库四引擎"技术架构的融合方案', size=12, color=(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_para(f'版本：v1.0', size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(f'日期：{datetime.date.today().strftime("%Y年%m月%d日")}', size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('四川融策会计师事务所 / 四川融策工程咨询公司', size=11, color=(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TOC placeholder
# ══════════════════════════════════════════════════════════════
add_heading_styled('目  录', 1)
add_para('（请在 Word 中右键此处 → 更新域 → 插入自动目录）', size=9, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 1: 体系定义
# ══════════════════════════════════════════════════════════════
add_heading_styled('一、体系定义', 1)

add_para('融策Agent "1+3+5+N" 体系是两个成熟框架的有机融合：', size=10.5)
add_bullet('1 = 统一数据底座：主数据标准、数据连接器、ODS/DWD/DWS/ADS四层数据架构')
add_bullet('3 = 三库四引擎：国资委穿透式监管AI技术标准（规则库·模型库·知识库 / 规则引擎·ML引擎·大模型引擎·图计算引擎）')
add_bullet('5 = 五库业务资产层：审计项目对象库、主题数据库、操作指引库、法规案例库、审计整改库')
add_bullet('N = N个审计应用场景：经责审计、采购审计、工程审计、绩效评价、资产清查、专项债、收支审计...')

doc.add_paragraph()
add_para('核心逻辑：5库是"存"的维度（业务全景地图），三库四引擎是"算"的维度（AI导航引擎），N个场景是"用"的维度。三层正交，合在一起才完整。', bold=True, size=10.5, color=(0x1C, 0x35, 0x5E))

# Architecture diagram description
add_heading_styled('1.1 全景架构', 2)

add_para('N 个审计应用场景', bold=True, size=10)
add_para('  经责审计 │ 采购审计 │ 工程审计 │ 绩效评价 │ 资产清查 │ 专项债 │ 收支审计', size=9, color=(0x66, 0x66, 0x66))
add_para('          ↓ 每个场景加载对应的规则+模型+知识+模板', size=9, color=(0x66, 0x66, 0x66))

add_para('3 — 三库四引擎（计算执行层）', bold=True, size=10)
add_para('  📋规则库→⚙️规则引擎  │  🤖模型库→🧠ML引擎', size=9)
add_para('  📚知识库→💬大模型引擎  │         →🕸️图计算引擎', size=9)
add_para('          ↓ 从五库提取规则/模型/知识，返回分析结果', size=9, color=(0x66, 0x66, 0x66))

add_para('5 — 五库（业务资产层）', bold=True, size=10)
add_para('  ①对象库 → ②主题库 → ③指引库 → ④法规案例库 → ⑤整改库', size=9)
add_para('          ↓ 业务资产：存什么、审什么、怎么改', size=9, color=(0x66, 0x66, 0x66))

add_para('1 — 统一数据底座', bold=True, size=10)
add_para('  主数据标准 │ 数据连接器 │ ODS贴源层 │ DWD明细层 │ DWS主题层 │ ADS应用层', size=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 2: 实施路线图总览
# ══════════════════════════════════════════════════════════════
add_heading_styled('二、实施路线图总览', 1)

add_para('四个阶段，12-16周完成全部建设：', size=10.5)

make_table(
    ['阶段', '主题', '时间', '核心产出', '状态标记'],
    [
        ['Phase 1', '核心引擎补缺', '第1-3周', 'ML引擎 + 规则引擎 + 模型注册', '🔴 本次启动'],
        ['Phase 2', '架构整合贯通', '第4-7周', 'RAG引擎 + Orchestrator + 数据连接器', '🟡 待Phase 1完成'],
        ['Phase 3', '业务库体系化', '第8-12周', '五库从概念落地为系统模块', '🟡 待Phase 2完成'],
        ['Phase 4', '扩展与生态', '第13-16周', '图计算 + 监管报送 + MLOps', '🟢 远期规划'],
    ],
    [2.5, 2.5, 2.0, 5.0, 2.5]
)

add_heading_styled('2.1 依赖关系', 2)
add_para('Phase 1（ML引擎 + 规则引擎 + 模型注册）是所有后续阶段的基建，必须先完成。', size=10)
add_para('Phase 2（RAG + Orchestrator + 数据连接器）依赖Phase 1的引擎能力。', size=10)
add_para('Phase 3（五库体系化）依赖Phase 2的Orchestrator调度能力。', size=10)
add_para('Phase 4（图计算 + MLOps）依赖Phase 1的模型注册和Phase 3的对象库数据。', size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 3: Phase 1
# ══════════════════════════════════════════════════════════════
add_heading_styled('三、Phase 1：核心引擎补缺（第1-3周）', 1)
add_para('目标：三库四引擎中缺口最大的组件从0到1，可独立运行。', bold=True, size=10.5)

add_heading_styled('3.1 ML引擎 — engine/ml_engine.py', 2)
add_para('当前状态：只有 TF-IDF、Benford、基础统计，缺少系统化机器学习能力。', size=10, color=(0x88, 0x33, 0x33))
add_para('目标状态：统一推理接口，封装6+模型，支持批量/单条推理。', size=10, color=(0x33, 0x88, 0x33))

make_table(
    ['模型', '业务场景', '技术方案', '优先级'],
    [
        ['Isolation Forest', '通用异常检测（报销/采购/费用）', 'sklearn.ensemble.IsolationForest', '🔴 P0'],
        ['Z-Score + IQR', '报价异常偏离检测', 'numpy/scipy', '🔴 P0'],
        ['DBSCAN', '供应商聚类/围标分组识别', 'sklearn.cluster.DBSCAN', '🟡 P1'],
        ['LOF（局部离群因子）', '非正态分布数据异常检测', 'sklearn.neighbors.LocalOutlierFactor', '🟡 P1'],
        ['时序异常检测', '收入/费用波动异常', 'statsmodels + 滑动窗口', '🟡 P1'],
        ['统计汇总', '分布校验、同比/环比偏离', 'scipy.stats', '🔴 P0'],
    ],
    [3.5, 5.5, 4.5, 1.5]
)

add_heading_styled('3.2 规则引擎实装 — 改造 api/risk_engine.py', 2)
add_para('当前状态：框架完整（Flask Blueprint + 异步任务），但业务逻辑全是 TODO 占位。', size=10, color=(0x88, 0x33, 0x33))
add_para('目标状态：规则可编译执行、可组合、可追踪，对接真实检测脚本。', size=10, color=(0x33, 0x88, 0x33))

add_para('实施内容：', bold=True, size=10)
add_bullet('规则条件编译：自然语言条件 → 可执行JSON表达式')
add_bullet('规则分层：red（红线，不可突破）/ yellow（预警，可调参）/ green（提示，可自定义）')
add_bullet('规则链：多条规则组合判定（AND / OR / 加权评分）')
add_bullet('对接真实检测脚本：L1报价规律、L3文本雷同、L4图片哈希、L5元数据、L7打印机、Benford')
add_bullet('规则效果追踪：命中率 / 误报率 / 覆盖度统计')

add_para('数据库变更（business_rules 表新增字段）：', bold=True, size=10)
add_bullet('rule_level — red / yellow / green')
add_bullet('rule_scope — group / subsidiary / custom')
add_bullet('executable — JSON条件表达式')
add_bullet('hit_count — 命中次数统计')

add_heading_styled('3.3 模型注册中心 — engine/model_registry.py', 2)
add_para('当前状态：模型散落在各 skill 脚本中，无统一管理。', size=10, color=(0x88, 0x33, 0x33))
add_para('目标状态：所有模型可发现、可查询、可调度。', size=10, color=(0x33, 0x88, 0x33))
add_bullet('registry.json — 模型元数据索引（名称/类型/路径/训练数据哈希/准确率/状态）')
add_bullet('ModelRegistry 类：list / get / register / evaluate / deprecate')
add_bullet('对接 Orchestrator（Phase 2 时使用）')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 4: Phase 2
# ══════════════════════════════════════════════════════════════
add_heading_styled('四、Phase 2：架构整合贯通（第4-7周）', 1)
add_para('目标：各引擎协同工作，形成完整的"数据→分析→输出"流水线。', bold=True, size=10.5)

add_heading_styled('4.1 RAG引擎 — engine/rag_engine.py', 2)
add_para('当前：知识库是SQLite+文件，只能精确匹配搜索。目标：语义检索 + 自动增强Agent推理上下文。', size=10)
add_bullet('向量化管道：sentence-transformers (all-MiniLM-L6-v2, 80MB) 本地 embedding → numpy 余弦检索')
add_bullet('RAG接口：retrieve(query) / augment_prompt(user_query) / ask(query)')
add_bullet('数据源：knowledge/ 下所有文章 + Function/Knowledge/knowledge.db')
add_bullet('集成点：Agent对话增强、报告自动引用法规、问题定性辅助匹配案例')

add_heading_styled('4.2 Orchestrator 调度中枢 — engine/orchestrator.py', 2)
add_para('当前：各能力独立调用。目标：一次分析请求 → 多引擎自动协同 → 聚合结果。', size=10)
add_bullet('场景配置（scenarios/）：procurement.yaml / engineering.yaml / financial.yaml / performance.yaml')
add_bullet('每个场景定义：启用规则 + 调度模型 + 检索知识域 + 输出报告模板')
add_bullet('分析流程：加载场景 → 并行调度引擎 → 聚合+交叉验证 → 风险评分 → 生成报告')

add_heading_styled('4.3 数据连接器 — engine/data_connector.py', 2)
add_bullet('导入模板库：科目余额表 / 序时账 / 合同台账 / 银行流水 / 投标文件包')
add_bullet('智能识别：上传文件 → 自动匹配模板 → 标准化导入')
add_bullet('为后续API直连预留接口')

add_heading_styled('4.4 大模型多路由 — LLM/router.py', 2)
add_bullet('SQL生成/简单查询 → DeepSeek V4 Flash（快+便宜）')
add_bullet('数据分析/复杂推理 → DeepSeek V4 Pro（主力）')
add_bullet('政策解读/长文生成 → DeepSeek V4 Pro')
add_bullet('知识库向量化 → sentence-transformers（本地）')

add_heading_styled('4.5 角色视图 — 改造前端 templates/', 2)
add_bullet('集团审计部视图：全集团风险热力图 + 穿透追溯 + 监管报送')
add_bullet('子公司风控视图：业务审核工作台 + 合同/付款合规校验 + 整改工单')
add_bullet('一线审计人员视图：数据上传→自动分析→标红异常 + 底稿生成 + 取证单')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 5: Phase 3
# ══════════════════════════════════════════════════════════════
add_heading_styled('五、Phase 3：业务库体系化（第8-12周）', 1)
add_para('目标：五库从概念落地为可操作的系统模块。', bold=True, size=10.5)

add_heading_styled('5.1 ① 审计项目对象库 — engine/object_repo.py', 2)
add_bullet('核心实体：被审计单位 / 审计项目 / 资金包 / 合同 / 供应商')
add_bullet('实体关系建模 + CRUD管理界面')
add_bullet('为Phase 4图计算引擎提供关系数据')

add_heading_styled('5.2 ② 主题数据库 — engine/theme_repo.py', 2)
add_bullet('按审计类型预置分析模板：经责审计 / 采购审计 / 工程审计 / 绩效评价')
add_bullet('每个主题 = 数据模板 + 分析模型 + 输出报告 → 一键加载')

add_heading_styled('5.3 ③ 操作指引库 — 增强 Function/Knowledge/', 2)
add_bullet('结构化审计程序库：审计类型 → 阶段 → 程序编号 → 审计目标 → 操作步骤')
add_bullet('对接 RAG 检索：输入审计场景 → 自动推荐对应程序')

add_heading_styled('5.4 ④ 法规案例库 — 增强 RAG 管道', 2)
add_bullet('法规索引：按发文机关/文号/日期/主题分类')
add_bullet('案例匹配：输入问题描述 → 检索相似案例 → 推荐定性依据')
add_bullet('问题定性辅助：发现异常 → 检索历史定性 → 引用法条 → 取证建议')

add_heading_styled('5.5 ⑤ 审计整改库 — 增强现有模块', 2)
add_bullet('问题台账增强：整改计划 / 截止日期 / 责任人 / 状态流转')
add_bullet('整改跟踪：措施 + 佐证材料 + 进度 + 到期提醒')
add_bullet('复核销号：复核意见 + 通过/退回 + 统计分析（完成率/超期/分布/周期）')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 6: Phase 4
# ══════════════════════════════════════════════════════════════
add_heading_styled('六、Phase 4：扩展与生态（第13-16周）', 1)

add_heading_styled('6.1 图计算引擎 — engine/graph_engine.py', 2)
add_bullet('Neo4j Community 部署 + 配置')
add_bullet('首期图模型：股权穿透图 / 招投标关系图 / 资金链路图')
add_bullet('API：query_penetration / find_path / detect_communities')

add_heading_styled('6.2 监管报送模板', 2)
add_bullet('自动生成：国资委穿透式监管自查报告 / 审计发现问题汇总表 / 整改情况统计表')
add_bullet('模板引擎：Jinja2 + 数据自动填充')
add_bullet('导出格式：Word/PDF/Excel')

add_heading_styled('6.3 MLOps 管线', 2)
add_bullet('模型训练数据版本管理（DVC）')
add_bullet('模型 A/B 测试框架')
add_bullet('生产模型 PSI 漂移监控')
add_bullet('自动再训练触发（PSI > 0.25）')

add_heading_styled('6.4 多租户支持', 2)
add_bullet('每个审计项目独立数据隔离')
add_bullet('客户专属知识库/规则库')
add_bullet('项目级权限控制')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 7: 目录结构
# ══════════════════════════════════════════════════════════════
add_heading_styled('七、融策Agent 目录结构目标', 1)

add_para('projects/data-analysis-agent/', bold=True, size=10)
add_bullet('engine/ — 新建：核心引擎层（ml_engine / rule_engine / model_registry / rag_engine / orchestrator / data_connector / graph_engine / object_repo / theme_repo）')
add_bullet('scenarios/ — 新建：审计场景配置（procurement.yaml / engineering.yaml / financial.yaml）')
add_bullet('models/ — 新建：模型文件存储（registry.json + .pkl 模型文件）')
add_bullet('templates/import/ — 新建：数据导入模板（科目余额表 / 序时账 / 合同台账 / 银行流水）')
add_bullet('api/ — 现有：保持，risk_engine.py 对接 engine/rule_engine')
add_bullet('agent/ — 现有：保持，agent.py 对接 Orchestrator')
add_bullet('Function/ — 现有：保持，Knowledge/ 增加 audit_procedures 表')
add_bullet('LLM/ — 现有 + router.py（多模型路由）')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 8: 考核对标
# ══════════════════════════════════════════════════════════════
add_heading_styled('八、融策Agent 与国资委 DRP 考核对标', 1)

make_table(
    ['国资委考核项', '分值', '融策Agent 对应模块', '实施阶段'],
    [
        ['基础能力建设（系统覆盖度、数据贯通率）', '2分', '统一数据底座 + 对象库 + 数据连接器', 'Phase 2-3'],
        ['风险防控效果（预警准确率≥70%）', '3分', 'ML引擎 + 规则引擎 + Orchestrator', 'Phase 1-2'],
        ['制度流程适配（规则嵌入度、整改闭环率）', '2分', '操作指引库 + 角色视图 + 整改库闭环', 'Phase 2-3'],
        ['创新应用加分（AI大模型、图计算）', '1分', 'RAG引擎 + 图计算引擎 + MLOps', 'Phase 2-4'],
    ],
    [3.5, 1.0, 5.5, 2.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# SECTION 9: 里程碑
# ══════════════════════════════════════════════════════════════
add_heading_styled('九、里程碑与交付物', 1)

make_table(
    ['里程碑', '时间节点', '交付物', '验收标准'],
    [
        ['M1', '第3周末', 'ML引擎 + 规则引擎实装 + 模型注册中心', '上传CSV → 自动标红异常记录，规则可编译执行'],
        ['M2', '第7周末', 'RAG引擎 + Orchestrator + 数据连接器 + 角色视图', '一次分析请求 → 多引擎自动协同输出报告'],
        ['M3', '第12周末', '五库全部可用 + 场景配置完整', '选择审计类型 → 一键加载全套模板和分析引擎'],
        ['M4', '第16周末', '图计算引擎 + 监管报送 + MLOps + 多租户', '股权穿透图可视化 + 自动生成监管报告'],
    ],
    [1.5, 2.0, 5.0, 4.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════
add_para('— 文档结束 —', size=9, color=(0xAA, 0xAA, 0xAA), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'生成日期：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}', size=8, color=(0xAA, 0xAA, 0xAA), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('四川融策会计师事务所 · 融策右护卫（OpenClaw AI助手）', size=8, color=(0xAA, 0xAA, 0xAA), alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
output_path = r'D:\openclaw-workspace\output\融策Agent_1+3+5+N_增强路线.docx'
doc.save(output_path)
print(f'Word doc generated: {output_path}')
