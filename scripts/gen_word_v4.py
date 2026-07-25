#!/usr/bin/env python3
"""
融策制度汇编 Word 生成器 v4.0 — 企业级图文并茂版
- 品牌配色：深蓝/青绿/铜金/暖灰
- 封面含Logo
- 页眉含公司名+Logo
- 页脚含页码
- 图表嵌入
- 精美排版
"""
import os, re, glob, io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ─── 路径配置 ───
SRC_DIR = r"C:\Users\scrccpa\.openclaw\workspace\output\新制度体系"
CHART_DIR = r"C:\Users\scrccpa\.openclaw\workspace\output\新制度体系\charts"
LOGO_PATH = r"C:\Users\scrccpa\.openclaw\skills\huashu-design\assets\rongce-brand\logo.png"
OUT_DIR = r"C:\Users\scrccpa\Desktop"

# ─── 品牌配色 ───
DEEP_BLUE   = RGBColor(0x0A, 0x1F, 0x3F)
TEAL        = RGBColor(0x1A, 0x5C, 0x6E)
COPPER      = RGBColor(0xC5, 0x95, 0x5C)
WARM_GRAY   = RGBColor(0xF5, 0xF2, 0xEC)
DARK_TEXT    = RGBColor(0x2D, 0x2D, 0x2D)
MUTED_TEXT   = RGBColor(0x99, 0x99, 0x99)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = RGBColor(0xE8, 0xED, 0xF2)

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
         'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
         'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
         'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
         'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}

# ─── 分类映射（文件归类） ───
CATEGORIES = {
    "人力资源篇": {
        "files": ["01-薪酬管理制度.md", "02-绩效考核管理制度.md", "03-员工手册.md",
                  "10-招聘与入职管理制度.md", "11-培训与发展管理制度.md",
                  "12-职级晋升管理制度.md", "36-高管经营目标责任与绩效考核办法.md",
                  "09-股东会议事规则.md", "38-岗位说明书汇编.md"],
        "output": "融策制度汇编-人力资源篇.docx",
        "chart": "01-hr-flow.png",
        "desc": "涵盖薪酬管理、绩效考核、员工行为规范、招聘入职、培训发展、职级晋升、高管考核、公司治理及岗位说明书等9项核心制度。"
    },
    "财务管理篇": {
        "files": ["06-财务报销管理制度.md", "13-项目收入确认与回款管理制度.md",
                  "14-预算管理制度.md", "15-资金管理制度.md",
                  "16-固定资产管理制度.md", "37-可分配利润核算细则.md"],
        "output": "融策制度汇编-财务管理篇.docx",
        "chart": "02-fin-flow.png",
        "desc": "涵盖费用报销、收入确认与回款、预算管理、资金管理、固定资产及可分配利润核算等6项财务制度。"
    },
    "业务质控篇": {
        "files": ["04-项目管理规范.md", "17-业务承接与合同管理制度.md",
                  "18-客户关系管理制度.md", "19-业务分包管理制度.md",
                  "20-投标管理制度.md", "34-项目独立核算与分润制度.md",
                  "35-跨部门协同与交叉营销奖励办法.md",
                  "07-审计质量控制制度.md", "08-造价咨询质量控制制度.md",
                  "21-三级复核实施细则.md", "22-执业责任追究制度.md"],
        "output": "融策制度汇编-业务质控篇.docx",
        "chart": "03-biz-qc-flow.png",
        "desc": "涵盖项目管理、业务承接与合同、客户关系、业务分包、投标管理、项目分润、跨部门协同、审计质控、造价质控、三级复核及执业责任追究等11项制度。"
    },
    "行政综合篇": {
        "files": ["05-制度发布与版本管理规范.md", "23-信息安全与保密管理制度.md",
                  "24-印章与证照管理制度.md", "25-档案管理制度.md",
                  "26-公司章程-会计师事务所.md", "27-公司章程-工程咨询公司.md",
                  "28-办公场所管理制度.md", "29-采购管理制度.md",
                  "30-数智化建设管理制度.md", "31-业务拓展与创新管理制度.md",
                  "32-风险管理制度.md", "33-党建工作制度.md"],
        "output": "融策制度汇编-行政综合篇.docx",
        "chart": "04-admin-flow.png",
        "desc": "涵盖制度发布管理、信息安全、印章证照、档案管理、公司章程、办公场所、采购、数智化建设、业务创新、风险管理及党建等12项制度。"
    }
}

# ─── 辅助函数 ───
def set_font(run, name='宋体', size=Pt(11), bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(),
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color_hex)
    shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')

def add_horizontal_line(doc, color='1A5C6E', width_cm=14):
    """Add a decorative horizontal line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '8')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_header_footer(doc, is_first_page_special=True):
    """Add branded header and footer to all sections."""
    for section in doc.sections:
        # Different first page
        section.different_first_page_header_footer = is_first_page_special
        
        # Default header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        
        # Add separator line to header
        pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1A5C6E')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # Try adding logo to header
        if os.path.exists(LOGO_PATH):
            try:
                run = hp.add_run()
                run.add_picture(LOGO_PATH, width=Inches(0.4))
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run2 = hp.add_run('  四川融策')
                set_font(run2, '微软雅黑', Pt(8), color=TEAL)
            except:
                pass
        
        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        
        # Add top border to footer
        pPr = fp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '4')
        top.set(qn('w:color'), 'CCCCCC')
        pBdr.append(top)
        pPr.append(pBdr)
        
        run = fp.add_run('— ')
        set_font(run, '微软雅黑', Pt(8), color=MUTED_TEXT)
        # Page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        fp._element.append(fldChar1)
        fp._element.append(instrText)
        fp._element.append(fldChar2)
        run = fp.add_run(' —')
        set_font(run, '微软雅黑', Pt(8), color=MUTED_TEXT)

def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

def add_cover_page(doc, category_name, subtitle=None):
    """Professional cover page with logo."""
    # Background color paragraph (full page effect)
    bg = doc.add_paragraph()
    bg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bg.paragraph_format.space_before = Pt(80)
    bg.paragraph_format.space_after = Pt(0)
    
    # Logo
    if os.path.exists(LOGO_PATH):
        run = bg.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.2))
    
    doc.add_paragraph()
    
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('四川融策')
    set_font(run, '微软雅黑', Pt(30), bold=True, color=DEEP_BLUE)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SICHUAN RONGCE')
    set_font(run, 'Arial', Pt(10), color=TEAL)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('━' * 24)
    set_font(run, '微软雅黑', Pt(11), color=COPPER)
    
    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('管理制度汇编')
    set_font(run, '微软雅黑', Pt(26), bold=True, color=DEEP_BLUE)
    
    if category_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f'—— {category_name} ——')
        set_font(run, '微软雅黑', Pt(16), color=TEAL)
    
    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('━' * 24)
    set_font(run, '微软雅黑', Pt(11), color=COPPER)
    
    for _ in range(5):
        doc.add_paragraph()
    
    # Company info at bottom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策会计师事务所有限公司  ·  四川融策工程咨询有限公司')
    set_font(run, '微软雅黑', Pt(10), color=MUTED_TEXT)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('发布日期：2026年7月22日')
    set_font(run, '微软雅黑', Pt(9), color=MUTED_TEXT)
    
    doc.add_page_break()

def add_chart_page(doc, chart_path, category_name, description):
    """Add a chart page with the flowchart embedded."""
    if not os.path.exists(chart_path):
        return
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'{category_name} · 制度体系概览')
    set_font(run, '微软雅黑', Pt(18), bold=True, color=DEEP_BLUE)
    
    add_horizontal_line(doc, 'C5955C')
    
    # Description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(description)
    set_font(run, '宋体', Pt(10), color=DARK_TEXT)
    
    # Add chart image
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(5.8))
    except Exception as e:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('[流程图待生成]')
        set_font(run, '微软雅黑', Pt(10), color=MUTED_TEXT)
    
    doc.add_page_break()

def add_section_title(doc, title, number=None):
    """Styled section title for individual policies."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    
    if number:
        run = p.add_run(f'第{number}篇')
        set_font(run, '微软雅黑', Pt(9), color=COPPER)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, '微软雅黑', Pt(16), bold=True, color=DEEP_BLUE)
    
    # Decorative underline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('━' * 20)
    set_font(run, '微软雅黑', Pt(8), color=COPPER)

def parse_md_to_docx(doc, md_content):
    """Parse markdown content into beautifully formatted docx."""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        rows_data = [r for r in table_rows if r and not all(re.match(r'^-{2,}$', c.strip()) or c.strip() in ('---', ':---', '---:', ':---:') for c in r)]
        if rows_data:
            table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.style = 'Table Grid'
            for ri, row in enumerate(rows_data):
                for ci, cell_text in enumerate(row):
                    if ci >= len(rows_data[0]):
                        break
                    cell = table.cell(ri, ci)
                    cell.text = ''
                    cp = cell.paragraphs[0]
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(2)
                    run = cp.add_run(cell_text.strip())
                    set_font(run, '微软雅黑', Pt(9))
                    if ri == 0:
                        set_cell_shading(cell, '0A1F3F')
                        run.font.color.rgb = WHITE
                        run.bold = True
                    elif ri % 2 == 0:
                        set_cell_shading(cell, 'F5F2EC')
        doc.add_paragraph()
        table_rows = []
        in_table = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Metadata block
        if stripped.startswith('> **制度编号**') or stripped.startswith('> **版本**') or \
           stripped.startswith('> **编制日期**') or stripped.startswith('> **适用范围**') or \
           stripped.startswith('> **优先级**') or stripped.startswith('> **替代说明**') or \
           stripped.startswith('> **审核人**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped.lstrip('> '))
            run = p.add_run(text)
            set_font(run, '微软雅黑', Pt(9), color=MUTED_TEXT)
            i += 1
            continue
        
        # ---
        if stripped == '---':
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '3')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            # Add left accent bar
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '1A5C6E')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(stripped[2:])
            set_font(run, '微软雅黑', Pt(15), bold=True, color=DEEP_BLUE)
            i += 1
            continue
        
        if stripped.startswith('## '):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            set_font(run, '微软雅黑', Pt(13), bold=True, color=TEAL)
            i += 1
            continue
        
        if stripped.startswith('### '):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[4:])
            set_font(run, '微软雅黑', Pt(11.5), bold=True, color=DEEP_BLUE)
            i += 1
            continue
        
        if stripped.startswith('#### '):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[5:])
            set_font(run, '微软雅黑', Pt(11), bold=True)
            i += 1
            continue
        
        # Empty line
        if not stripped:
            flush_table()
            i += 1
            continue
        
        # Also flush if new table has different column count
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # If already accumulating and column count changes, flush first
            if table_rows and len(cells) != len(table_rows[-1]):
                flush_table()
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        
        # Ordered list
        if re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            text = stripped
            parts = re.split(r'(\*\*.+?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, '宋体', Pt(11), bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run, '宋体', Pt(11))
            i += 1
            continue
        
        # Unordered list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            text = stripped[2:]
            parts = re.split(r'(\*\*.+?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, '宋体', Pt(11), bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run, '宋体', Pt(11))
            i += 1
            continue
        
        # Normal paragraph with inline bold
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = Cm(0.74)
        parts = re.split(r'(\*\*.+?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, '宋体', Pt(11), bold=True)
            else:
                run = p.add_run(part)
                set_font(run, '宋体', Pt(11))
        i += 1
    
        # Trailing table
        flush_table()

def add_toc(doc, entries):
    """Styled table of contents."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('目    录')
    set_font(run, '微软雅黑', Pt(18), bold=True, color=DEEP_BLUE)
    
    add_horizontal_line(doc, 'C5955C')
    
    for idx, (fname, title) in enumerate(entries, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        # Number
        run = p.add_run(f'{idx:02d}  ')
        set_font(run, 'Arial', Pt(10), color=COPPER)
        # Title
        run = p.add_run(title)
        set_font(run, '微软雅黑', Pt(10.5))
    
    doc.add_page_break()

def generate_volume(cat_name, info):
    """Generate one categorized volume with full branding."""
    doc = Document()
    set_page_margins(doc)
    
    add_cover_page(doc, cat_name)
    
    # Read files
    entries = []
    existing_files = []
    for fname in info['files']:
        path = os.path.join(SRC_DIR, fname)
        if os.path.exists(path):
            existing_files.append(fname)
            with open(path, encoding='utf-8') as f:
                first_line = f.readline().strip().lstrip('# ')
                entries.append((fname, first_line or fname))
    
    # Add flowchart if available
    chart_path = os.path.join(CHART_DIR, info['chart'])
    add_chart_page(doc, chart_path, cat_name, info['desc'])
    
    # TOC
    add_toc(doc, entries)
    
    # Content
    for idx, fname in enumerate(existing_files, 1):
        path = os.path.join(SRC_DIR, fname)
        with open(path, encoding='utf-8') as f:
            content = f.read()
        
        first_line = content.split('\n')[0].strip().lstrip('# ')
        add_section_title(doc, first_line, idx)
        parse_md_to_docx(doc, content)
        doc.add_page_break()
    
    # Add headers and footers
    add_header_footer(doc)
    
    out_path = os.path.join(OUT_DIR, info['output'])
    doc.save(out_path)
    print(f'  ✓ {info["output"]} ({len(existing_files)}项制度)')
    return len(existing_files)

def generate_full_volume():
    """Generate complete single-volume with max branding."""
    doc = Document()
    set_page_margins(doc)
    
    # ─── 封面 ───
    # Spacer
    for _ in range(3):
        doc.add_paragraph()
    
    # Logo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.6))
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策')
    set_font(run, '微软雅黑', Pt(34), bold=True, color=DEEP_BLUE)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SICHUAN RONGCE')
    set_font(run, 'Arial', Pt(11), color=TEAL)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Decorative
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 28)
    set_font(run, '微软雅黑', Pt(12), color=COPPER)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('公司制度体系')
    set_font(run, '微软雅黑', Pt(28), bold=True, color=DEEP_BLUE)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（完整版）')
    set_font(run, '微软雅黑', Pt(16), color=TEAL)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 28)
    set_font(run, '微软雅黑', Pt(12), color=COPPER)
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('共37项制度 · 七大类 · 2026年7月22日发布')
    set_font(run, '微软雅黑', Pt(11), color=TEAL)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('四川融策会计师事务所有限公司  ·  四川融策工程咨询有限公司')
    set_font(run, '微软雅黑', Pt(10), color=MUTED_TEXT)
    
    doc.add_page_break()
    
    # ─── 总览图 ───
    chart_main = os.path.join(CHART_DIR, '00-architecture.png')
    if os.path.exists(chart_main):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run('制度体系总览')
        set_font(run, '微软雅黑', Pt(20), bold=True, color=DEEP_BLUE)
        
        add_horizontal_line(doc, 'C5955C')
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_main, width=Inches(6.0))
        
        doc.add_page_break()
    
    # ─── 目录 ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('目    录')
    set_font(run, '微软雅黑', Pt(20), bold=True, color=DEEP_BLUE)
    
    add_horizontal_line(doc, 'C5955C')
    
    global_counter = 0
    for cat_name, info in CATEGORIES.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(cat_name)
        set_font(run, '微软雅黑', Pt(13), bold=True, color=TEAL)
        
        for fname in info['files']:
            path = os.path.join(SRC_DIR, fname)
            if os.path.exists(path):
                global_counter += 1
                with open(path, encoding='utf-8') as f:
                    title = f.readline().strip().lstrip('# ')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.2)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f'{global_counter:02d}  {title}')
                set_font(run, '微软雅黑', Pt(10))
    
    doc.add_page_break()
    
    # ─── 正文 ───
    global_counter = 0
    for cat_name, info in CATEGORIES.items():
        # Category divider page
        for _ in range(6):
            doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cat_name)
        set_font(run, '微软雅黑', Pt(24), bold=True, color=DEEP_BLUE)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 20)
        set_font(run, '微软雅黑', Pt(10), color=COPPER)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info['desc'])
        set_font(run, '宋体', Pt(10), color=DARK_TEXT)
        
        # Category chart
        chart_path = os.path.join(CHART_DIR, info['chart'])
        if os.path.exists(chart_path):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p.add_run()
                run.add_picture(chart_path, width=Inches(5.5))
            except:
                pass
        
        doc.add_page_break()
        
        # Each policy
        for fname in info['files']:
            path = os.path.join(SRC_DIR, fname)
            if not os.path.exists(path):
                continue
            
            global_counter += 1
            with open(path, encoding='utf-8') as f:
                content = f.read()
            
            first_line = content.split('\n')[0].strip().lstrip('# ')
            add_section_title(doc, first_line, global_counter)
            parse_md_to_docx(doc, content)
            doc.add_page_break()
    
    add_header_footer(doc)
    
    out_path = os.path.join(OUT_DIR, '融策公司制度体系（完整版）.docx')
    doc.save(out_path)
    print(f'  ✓ 融策公司制度体系（完整版）.docx ({global_counter}项制度)')

if __name__ == '__main__':
    print('═' * 50)
    print('  融策制度汇编 v4.0 — 图文并茂版')
    print('═' * 50)
    print()
    
    total = 0
    for cat_name, info in CATEGORIES.items():
        existing = [f for f in info['files'] if os.path.exists(os.path.join(SRC_DIR, f))]
        print(f'[{cat_name}] {len(existing)}项制度')
        count = generate_volume(cat_name, info)
        total += count
    
    print(f'\n分册合计: {total}项')
    print('\n生成完整版...')
    generate_full_volume()
    
    print(f'\n✅ 全部完成！桌面文件:')
    for f in sorted(glob.glob(os.path.join(OUT_DIR, '融策制度*'))):
        size_kb = os.path.getsize(f) / 1024
        print(f'  📄 {os.path.basename(f)} ({size_kb:.0f} KB)')
