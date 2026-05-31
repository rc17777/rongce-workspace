# -*- coding: utf-8 -*-
"""提取投标文件全部内容"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

base = r'D:\openclaw-workspace\temp\tender_analysis'
out = []

def read_docx(path):
    try:
        import docx
        doc = docx.Document(path)
        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        # Also read tables
        for table in doc.tables:
            text += '\n\n--- TABLE ---\n'
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text += ' | '.join(cells) + '\n'
        return text
    except Exception as e:
        return f"[ERROR reading DOCX: {e}]"

def read_pdf(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = ''
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + '\n'
            return text
    except ImportError:
        return "[pdfplumber not installed]"
    except Exception as e:
        return f"[ERROR reading PDF: {e}]"

for root, dirs, files in os.walk(base):
    for f in sorted(files):
        fp = os.path.join(root, f)
        rel = os.path.relpath(fp, base)
        out.append(f'\n{"="*80}\nFILE: {rel}\n{"="*80}')
        
        ext = f.lower()
        if ext.endswith('.docx'):
            txt = read_docx(fp)
        elif ext.endswith('.pdf'):
            txt = read_pdf(fp)
        else:
            txt = "[unknown format]"
        
        if txt:
            out.append(txt[:5000])  # limit per file

final = '\n'.join(out)
with open(r'D:\openclaw-workspace\temp\tender_analysis\_extracted.txt', 'w', encoding='utf-8') as fw:
    fw.write(final)

print(f'Extracted {len(files)} files')
print(f'Total chars: {len(final)}')
print('Saved to _extracted.txt')
