#!/usr/bin/env python3
"""桌面Markdown → Word转换"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DESKTOP = os.path.expanduser(r"~\Desktop")
FILES = [
    "从这里开始-Prompt学习路线.md",
    "Prompt速查卡.md",
    "Prompt工程权威资料汇总.md",
]

def md_to_docx(md_path, docx_path):
    """简陋但够用的MD→DOCX转换器"""
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()
    
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    for line in text.split("\n"):
        line = line.rstrip()
        
        # 跳过空行
        if not line.strip():
            continue
        
        # 标题
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[2:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[3:], level=3)
        elif line.startswith("#### "):
            p = doc.add_heading(line[4:], level=4)
        # 分隔线
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.size = Pt(8)
        # 无序列表
        elif re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            p = doc.add_paragraph(text, style='List Bullet')
        # 有序列表
        elif re.match(r"^\d+[\.\)]\s+", line):
            text = re.sub(r"^\d+[\.\)]\s+", "", line)
            p = doc.add_paragraph(text, style='List Number')
        # 代码块标记 - 跳过
        elif line.strip().startswith("```"):
            continue
        # 缩进代码
        elif line.startswith("    ") or line.startswith("\t"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
        # 普通段落
        else:
            p = doc.add_paragraph()
            
            # 内联格式处理
            # 粗体 **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # 处理内联代码 `text`
                    code_parts = re.split(r'(`.*?`)', part)
                    for cp in code_parts:
                        if cp.startswith("`") and cp.endswith("`"):
                            run = p.add_run(cp[1:-1])
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                        else:
                            p.add_run(cp)
    
    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f"  ✅ {os.path.basename(docx_path)} ({size_kb:.1f} KB)")

if __name__ == "__main__":
    print("=" * 50)
    print("  Markdown → Word 转换")
    print("=" * 50)
    for fname in FILES:
        md_path = os.path.join(DESKTOP, fname)
        docx_path = md_path.replace(".md", ".docx")
        if os.path.exists(md_path):
            print(f"\n  📄 {fname}")
            md_to_docx(md_path, docx_path)
        else:
            print(f"\n  ⚠️ 未找到: {fname}")
    print(f"\n完成。")
