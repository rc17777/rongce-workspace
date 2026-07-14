import os, re
from docx import Document

base = r"C:\Users\scrccpa\Desktop\教科院\教科院内控制度"
out = r"D:\openclaw-workspace\教科院内控分析"

files = [
    "2025.11.5-1.23.docx",
    "2026.1.23—3.13 副本(1).docx",
    "2026.3.13至今.docx",
    "内控制度2024.6.14-11.18.docx",
    "内控制度2024.11.8-3.17.docx",
    "内控制度2025.3.17-4.11 - 副本.docx",
    "内控制度2025.4.11-5.27.docx",
    "内控制度2025.5.27-7.11.docx",
    "内控制度2025.7.11-10.13 - 副本.docx",
    "内控制度2025.10.13-11.5.docx",
]

for f in files:
    fpath = os.path.join(base, f)
    name = os.path.splitext(f)[0]
    print(f"Processing: {f}")
    try:
        doc = Document(fpath)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if "Heading" in style or "heading" in style:
                    level = re.findall(r'\d+', style)
                    lv = level[0] if level else "1"
                    lines.append(f"{'#' * int(lv)} {text}")
                else:
                    lines.append(text)
        
        # Also extract tables
        for i, table in enumerate(doc.tables):
            lines.append(f"\n--- 表格 {i+1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        
        outpath = os.path.join(out, f"{name}.md")
        with open(outpath, 'w', encoding='utf-8') as wf:
            wf.write('\n'.join(lines))
        print(f"  -> {outpath} ({len(lines)} lines)")
    except Exception as e:
        print(f"  ERROR: {e}")

print("\nAll done!")
