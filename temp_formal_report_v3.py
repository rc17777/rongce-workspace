# -*- coding: utf-8 -*-
"""恩阳医养园PPP可用性付费报告 — v3.0 终审修正版"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# PARAMETERS
# ============================================================
CAPITAL_TOTAL = 102_298_077.79
CAPITAL_AMORT = 80_000_000.00
CAPITAL_LUMP = 22_298_077.79
RATE = 0.0799
YEARS = 18

A = CAPITAL_AMORT * RATE * (1+RATE)**YEARS / ((1+RATE)**YEARS - 1)

# Capital amortization schedule
cap_remaining = CAPITAL_AMORT
cap_schedule = []
for y in range(YEARS):
    interest = cap_remaining * RATE
    principal = A - interest
    if y == YEARS - 1:
        principal = cap_remaining
        interest = A - principal
    cap_remaining -= principal
    if cap_remaining < 0: cap_remaining = 0
    cap_schedule.append((principal, interest, A, max(cap_remaining, 0)))

# Bank: original表2逐项数据, scaled to bank-confirmed total 710,822,119.32
bank_original = [
    (31_686_745.84, 381_300_000.00),
    (26_809_671.03, 381_300_000.00),
    (24_811_756.73, 381_250_000.00),
    (21_351_219.32, 380_250_000.00),
    (20_263_854.17, 379_250_000.00),
    (20_265_833.33, 378_250_000.00),
    (20_162_465.30, 377_250_000.00),
    (20_111_770.82, 376_250_000.00),
    (23_022_951.39, 372_250_000.00),
    (22_871_874.99, 368_250_000.00),
    (22_617_395.82, 364_250_000.00),
    (22_414_618.06, 360_250_000.00),
    (27_148_298.64, 351_250_000.00),
    (27_728_124.99, 341_250_000.00),
    (35_070_729.16, 323_250_000.00),
    (34_158_229.18, 305_250_000.00),
    (311_968_819.45,         0),
    (          0.00,         0),
]
BANK_ORIG_SUM = sum(b[0] for b in bank_original)
BANK_REPORT_SUM = 710_822_119.32
BANK_SCALE = BANK_REPORT_SUM / BANK_ORIG_SUM  # ≈0.997695

bank_schedule = [(round(amt * BANK_SCALE, 2), rem) for amt, rem in bank_original]
bank_total = sum(b[0] for b in bank_schedule)

# Op costs (v3: 2023/2024 from original 表3, 2025 from audit, 2026+ based on 2025 actual)
OP_2023 = 5_768_117.42
OP_2024 = 6_890_208.16
OP_2025 = 5_713_241.86
OP_FUTURE = 5_710_000.00
op_costs = [OP_2023, OP_2024, OP_2025] + [OP_FUTURE]*15

# Income (v3: 2023/2024 from original 表3, 2025 from audit, 2026+ based on 2025 actual)
INC_2023 = 4_285_070.31
INC_2024 = 6_229_093.45
INC_2025 = 5_564_017.88
INC_FUTURE = 5_560_000.00
incomes = [INC_2023, INC_2024, INC_2025] + [INC_FUTURE]*15

YEAR_LABELS = list(range(2023, 2041))

yearly = []
for i in range(YEARS):
    yr = YEAR_LABELS[i]
    principal, interest, pmt, remaining = cap_schedule[i]
    lump = CAPITAL_LUMP if i == YEARS - 1 else 0
    total_cap = pmt + lump
    bank = bank_schedule[i][0]
    op = op_costs[i]
    inc = incomes[i]
    avail = total_cap + bank + op - inc
    yearly.append((yr, principal, interest, pmt, lump, total_cap, remaining, bank, op, inc, avail))

grand_total = sum(y[10] for y in yearly)
cap_total = sum(y[5] for y in yearly)
op_cost_total = sum(op_costs)
income_total = sum(incomes)

print(f"PMT={A:,.2f}  CapSum={cap_total:,.2f}  BankSum={bank_total:,.2f}  OpSum={op_cost_total:,.2f}  IncSum={income_total:,.2f}  Total={grand_total:,.2f}")

# ============================================================
# BUILD WORD DOCUMENT
# ============================================================
doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def scf(cell, text, fn='宋体', sz=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = fn; run.font.size = sz; run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color: run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0

def scs(cell, c):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def ap(text, fn='宋体', sz=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None, sa=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    if indent: p.paragraph_format.first_line_indent = indent
    p.paragraph_format.space_after = sa
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = fn; run.font.size = sz; run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    return p

def mt(headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]
        scf(c,h,fn='微软雅黑',sz=Pt(9),bold=True,color=RGBColor(0xFF,0xFF,0xFF))
        scs(c,'0A1F3F')
    for i,rd in enumerate(rows):
        for j,v in enumerate(rd):
            c=t.rows[i+1].cells[j]
            scf(c,str(v),fn='宋体',sz=Pt(8.5))
            if i%2==1: scs(c,'F5F2EC')
    if cw:
        for row in t.rows:
            for j,w in enumerate(cw): row.cells[j].width=Cm(w)
    return t

# ===== COVER =====
for _ in range(6): doc.add_paragraph()
ap('巴中市恩阳医养园PPP项目', fn='方正小标宋简体', sz=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ap('可用性付费测算结果报告', fn='方正小标宋简体', sz=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4): doc.add_paragraph()
ap('川融策咨询〔2025〕第490号', sz=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6): doc.add_paragraph()
ap('四 川 融 策 会 计 师 事 务 所 有 限 公 司', fn='微软雅黑', sz=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ap('Sichuan Rongce Accounting Firm Co., Ltd', sz=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(20))
ap('2026年7月29日', sz=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===== BODY =====
ap('川融策咨询〔2025〕第  号', align=WD_ALIGN_PARAGRAPH.RIGHT, sz=Pt(10))
doc.add_paragraph()
ap('巴中市恩阳医养园PPP项目', fn='黑体', sz=Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ap('可用性付费测算结果报告', fn='黑体', sz=Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
ap('巴中市恩阳区卫生健康局：', indent=Cm(0.74))
ap('我们接受委托，对巴中市恩阳医养园PPP项目2022年10月28日至2040年12月21日可用性付费进行测算。本次测算以《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）中审定的建筑安装工程费用485,098,077.79元为基数，运营维护成本及第三方收入已根据项目公司2025年度运营审计数据进行修正，现将测算结果报告如下：', indent=Cm(0.74))

# ===== 一、项目概况 =====
ap('一、项目概况', fn='黑体', sz=Pt(14), bold=True)
ap('2015年12月21日，巴中市恩阳医养园PPP项目经巴中市恩阳区人民政府（恩府〔2015〕208号）批复同意实施，由巴中市恩阳区卫生健康局按规定有序开展。', indent=Cm(0.74))
ap('2016年7月巴中市恩阳医养园PPP项目协议经区人民政府审定同意，巴中市恩阳区卫生健康局于2016年12月与恩阳区人民医院、湖南省第五工程有限公司签订PPP协议，成立项目管理公司，注册资本金10000万元，承担该项目的投资、建设、运营维护及移交。', indent=Cm(0.74))
ap('2017年7月19日项目公司与湖南省第五工程有限公司签订施工合同，项目建筑总面积118,515.13㎡，包括院内绿化、场地及路面硬化、室内管网等附属工程及设备购置等。该项目实际于2017年5月17日开工，2022年3月完工，2022年3月28日进入调试运营期，同年10月27日通过竣工验收，正式进入运营阶段。', indent=Cm(0.74))
ap('根据2021年10月15日签订的《巴中市恩阳医养园PPP项目合同之补充合同（2021年）》，项目合作期调整为23年，其中建设期不超过3年（实际建设期5年），运营期为18年（2022年10月28日至2040年12月21日）。', indent=Cm(0.74))

# ===== 二、测算依据 =====
ap('二、测算依据', fn='黑体', sz=Pt(14), bold=True)
ap('（一）政策依据', fn='楷体', sz=Pt(12), bold=True)
ap('（1）《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）', indent=Cm(0.74))
ap('（2）财政部关于印发《政府和社会资本合作项目财政承受能力论证指引》的通知（财金〔2015〕21号）', indent=Cm(0.74))
ap('（3）四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）', indent=Cm(0.74))

ap('（二）项目资料依据', fn='楷体', sz=Pt(12), bold=True)
for i, item in enumerate([
    '《恩阳医养园PPP项目协议》（2016年12月27日）',
    '《巴中市恩阳医养园PPP项目合同》（2017年）',
    '《巴中市恩阳医养园PPP项目合同之补充合同》（2017年/2018年/2021年）',
    '《关于巴中市恩阳区人民医院一期工程建设项目可行性研究报告的批复》（恩区发改行审〔2015〕83号）',
    '《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）',
    '《借款还本付息情况报告》（贵阳银行股份有限公司成都分行出具）',
    '项目公司2025年度财务报表及运营成本/收入明细账（2025年1-12月，共6册XLS）',
    '2022-2025年度运营期绩效考核结果（恩阳区卫健局出具）',
], 1):
    ap(f'（{i}）{item}', indent=Cm(0.74))

# ===== 三、测算说明 =====
ap('三、测算说明', fn='黑体', sz=Pt(14), bold=True)
ap('（一）项目回报机制识别', fn='楷体', sz=Pt(12), bold=True)
ap('一是政策方面：根据《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）："项目回报机制主要分为使用者付费、可行性缺口补助和政府付费等支付方式。其中：使用者付费，是指由最终消费用户直接付费购买公共产品和服务。可行性缺口补助，是指使用者付费不足以满足社会资本或项目公司成本回收和合理回报，而由政府以财政补贴、股本投入、优惠贷款和其他优惠政策的形式，给予社会资本或项目公司的经济补助。政府付费，是指政府直接付费购买公共产品和服务，主要包括可用性付费、使用量付费和绩效付费。"', indent=Cm(0.74))
ap('四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）规定："运营补贴。指项目运营期间根据不同的付费模式政府应当承担的直接付费责任。其中：政府付费模式下，政府承担全部运营补贴支出责任；可行性缺口补助（政府补助）模式下，政府承担部分运营补贴支出责任；使用者付费模式下，政府不承担运营补贴支出责任。"', indent=Cm(0.74))
ap('二是项目资料方面：项目实施方案、PPP合同文本中均明确项目回报机制为"可行性缺口补助"。2016年12月27日巴中市恩阳区卫生健康局与巴中市恩阳区人民医院、湖南省第五工程有限公司签订《恩阳医养园PPP项目协议》，协议约定"自运营日起，甲方按照本协议的约定及时、足额地向乙方支付可行性缺口补贴，且区政府将本项目的可行性缺口补贴纳入跨年度的财政预算，并提请人大决议"。', indent=Cm(0.74))

ap('（二）回报计算方法选用', fn='楷体', sz=Pt(12), bold=True)
ap('根据2017年7月巴中市恩阳区卫生健康局与项目公司签订《巴中市恩阳医养园PPP项目合同之补充合同》，合同约定计算公式：', indent=Cm(0.74))
ap('A = P×k×(1+k)^n / ((1+k)^n-1) + 实际融资成本 + 运营维护成本×(1+K) - 第三方收入', indent=Cm(0.74))
ap('其中A为运营期内各年政府运营补贴金额。', indent=Cm(0.74))

ap('（1）项目资本金回报：P×k×(1+k)^n / ((1+k)^n-1)', indent=Cm(0.74))
ap('P=项目资本金实际投入额102,298,077.79元（其中恩阳区人民医院投入2,000万元、社会资本投入82,298,077.79元）；k=社会资本中标年回报率7.99%；n=运营期18年。测算方案调整为：8,000万元按7.99%年回报率等额本息方式在18年运营期内逐年偿还（年还款额A={:,.2f}元）；剩余22,298,077.79元（含恩阳区人民医院投入的2,000万元）于运营期最后一年（2040年）一次性支付，不计算回报利息。'.format(A), indent=Cm(0.74))

ap('（2）实际融资成本：项目运营期间向贵阳银行股份有限公司成都分行借款产生的利息及本金合计710,822,119.32元（数据来源：贵阳银行成都分行《借款还本付息情况报告》，验证依据：2025年项目公司财务费用明细）。按银行实际还款计划逐年列示于下表2。原借款合同利率6.95%/年，后经三次下调至5.0%/年（2024年11月→6.45%，2025年11月→5.5%，2026年7月→5.0%），累计发放贷款38,280万元。', indent=Cm(0.74))

ap('（3）运营维护成本：2023年度（2022.10.28-2023.12.21）实际发生5,768,117.42元；2024年度（2023.12.21-2024.12.21）实际发生6,890,208.16元；2025年度实际发生5,713,241.86元（主营业务成本4,474,027.75元+管理费用1,239,214.11元，数据来源：项目公司2025年度运营审计XLS明细账4季度损益结转合计，不含年末调整项）；2026-2040年度暂按2025年度实际值四舍五入取整5,710,000.00元/年代入公式估算。', indent=Cm(0.74))

ap('（4）第三方收入：2023年度实际发生4,285,070.31元；2024年度实际发生6,229,093.45元（含第三经营年度2024.10.28-2024.12.31期间收入387,927.58元）；2025年度实际发生5,564,017.88元（主营业务收入4,961,142.62元+其他业务收入602,875.26元，数据来源：同上）；2026-2040年度暂按2025年度实际值四舍五入取整5,560,000.00元/年代入公式估算。', indent=Cm(0.74))

# ===== 四、测算过程 =====
ap('四、测算过程', fn='黑体', sz=Pt(14), bold=True)
ap('（一）每年项目资本金回报测算情况', fn='楷体', sz=Pt(12), bold=True)
ap(f'项目资本金回报测算详见表1（单位：元）：', indent=Cm(0.74))

t1h = ['年', '等额本息\n本金', '等额本息\n利息', '等额本息\n小计', '一次性支付', '资本金回报\n合计', '剩余(等额\n本息部分)']
t1r = []
for y in yearly:
    yr, pr, interest, pmt, lump, total_cap, remaining = y[0], y[1], y[2], y[3], y[4], y[5], y[6]
    t1r.append([str(yr), f'{pr:,.2f}', f'{interest:,.2f}', f'{pmt:,.2f}',
                f'{lump:,.2f}' if lump else '-', f'{total_cap:,.2f}',
                f'{remaining:,.2f}' if remaining > 0 else '0.00'])
t1r.append(['合计', f'{sum(y[1] for y in yearly):,.2f}', f'{sum(y[2] for y in yearly):,.2f}',
            f"{sum(y[3] for y in yearly):,.2f}", f'{CAPITAL_LUMP:,.2f}', f'{cap_total:,.2f}', '-'])
mt(t1h, t1r, cw=[1.1, 2.3, 2.3, 2.3, 2.2, 2.3, 2.8])
doc.add_paragraph()

# ===== 四（二）融资 =====
ap('（二）每年实际融资本息测算情况', fn='楷体', sz=Pt(12), bold=True)
ap(f'项目融资本息见表2（单位：元；银行还款计划合计710,822,119.32元）：', indent=Cm(0.74))

t2h = ['年', '实际融资本息', '剩余融资金额', '适用利率']
t2r = []
for i, y in enumerate(yearly):
    yr = y[0]; ba = y[7]; br = bank_schedule[i][1]
    rate_str = {0:'6.95%',1:'6.45%',2:'5.5%/5.0%',3:'5.0%',4:'5.0%',5:'5.0%',6:'5.0%',
                7:'5.0%',8:'5.0%',9:'5.0%',10:'5.0%',11:'5.0%',12:'5.0%',13:'5.0%',
                14:'5.0%',15:'5.0%',16:'5.0%',17:'-'}.get(i,'5.0%')
    t2r.append([str(yr), f'{ba:,.2f}', f'{br:,.2f}' if br > 0 else '-', rate_str])
t2r.append(['合计', f'{bank_total:,.2f}', '-', '—'])
mt(t2h, t2r, cw=[1.5, 4.0, 4.0, 2.0])
doc.add_paragraph()
ap('注：第18年（2040年度）实际融资本息为0元，系贵阳银行全部贷款已于第17年（2039年12月）结清，故2040年度无融资还款义务。', indent=Cm(0.74), sz=Pt(10))

# ===== 四（三）（四）（五）=====
ap('（三）运营维护成本', fn='楷体', sz=Pt(12), bold=True)
ap('2023年度实际发生额：5,768,117.42元（覆盖期2022.10.28-2023.12.21，约14个月，含第一经营年度及第二经营年度开始两个月）。', indent=Cm(0.74))
ap('2024年度实际发生额：6,890,208.16元（覆盖期2023.12.21-2024.12.21，12个月，含第二经营年度剩余及第三经营年度2024.10.28-2024.12.31期间431,456.22元）。', indent=Cm(0.74))
ap('2025年度实际发生额：5,713,241.86元（覆盖期2024.12.21-2025.12.21，全年。其中主营业务成本4,474,027.75元、管理费用1,239,214.11元，来源于项目公司2025年度运营审计XLS明细账4个季度损益结转合计）。', indent=Cm(0.74))
ap('2026-2040年度：暂按2025年度实际值四舍五入取整5,710,000.00元/年代入公式估算（共15年）。', indent=Cm(0.74))

ap('（四）第三方收入', fn='楷体', sz=Pt(12), bold=True)
ap('2023年度实际发生额：4,285,070.31元。2024年度实际发生额：6,229,093.45元（含第三经营年度2024.10.28-2024.12.31期间收入387,927.58元）。2025年度实际发生额：5,564,017.88元（主营业务收入4,961,142.62元+其他业务收入602,875.26元，来源同上）。2026-2040年度：暂按5,560,000.00元/年代入公式估算。', indent=Cm(0.74))

ap('（五）考核结果', fn='楷体', sz=Pt(12), bold=True)
ap('根据2018年3月7日补充协议，每年绩效考核分数≥80分时可用性付费支付比例为100%。截止2026年6月30日，已开展的四个运营年度考核分数均达到80分以上（2022年度93分、2023年度94分、2024年度95分、2025年度93.5分），可用性付费均按100%支付。', indent=Cm(0.74))

# ===== 五、测算结果 =====
ap('五、测算结果', fn='黑体', sz=Pt(14), bold=True)
ap('根据PPP合同约定，运营期内每年末支付一次可用性费用。项目可用性付费汇总见表3（单位：元）：', indent=Cm(0.74))

t3h = ['年', '资本金回报', '实际融资本息', '运营维护成本', '第三方收入', '可用性付费']
t3r = []
for y in yearly:
    yr, _, _, _, _, total_cap, _, bank, op, inc, avail = y
    t3r.append([str(yr), f'{total_cap:,.2f}', f'{bank:,.2f}', f'{op:,.2f}', f'{inc:,.2f}', f'{avail:,.2f}'])
t3r.append(['合计', f'{cap_total:,.2f}', f'{bank_total:,.2f}', f'{op_cost_total:,.2f}', f'{income_total:,.2f}', f'{grand_total:,.2f}'])
mt(t3h, t3r, cw=[1.2, 2.8, 2.8, 2.8, 2.8, 2.8])
doc.add_paragraph()

ap(f'测算结论：巴中市恩阳医养园PPP项目运营期（2022年10月28日至2040年12月21日）18年可用性付费总额为{grand_total:,.2f}元，约{grand_total/1e8:.2f}亿元。', indent=Cm(0.74), bold=True)

ap('备注：', fn='楷体', sz=Pt(10), bold=True)
ap('1.计费期限：除2023年度为2022年10月28日至2023年12月21日（约14个月）外，其余年度均为上一年度12月21日至当年12月21日。', sz=Pt(10))
ap('2.资本金测算方案说明：8,000万元按等额本息公式（P=80,000,000，k=7.99%，n=18）逐年偿还，18年等额本息合计153,544,445.28元（本金80,000,000.00元+利息73,544,445.28元）；剩余22,298,077.79元（含恩阳区人民医院2,000万元）于第18年一次性支付。', sz=Pt(10))
ap('3.融资本息数据来源：贵阳银行成都分行《借款还本付息情况报告》（合计数710,822,119.32元）。逐年数据按银行还款计划逐项列示，均已按银行报告确认总数进行了比例调整。2025年第3季度项目公司财务费用明细计提贵阳银行贷款利息6,291,985.59元（本金38,125万元，年利率约6.6%），与银行还款计划数基本吻合。', sz=Pt(10))
ap('4.运维成本与第三方收入数据来源：①2023-2024年度取自原征求意见稿表3实际发生数；②2025年度取自项目公司2025年度运营审计XLS明细账（主营业务成本+管理费用=运维成本；主营业务收入+其他业务收入=第三方收入），取4个标准季度损益结转合计值，不含年末调整项；③2026-2040年度以2025年度实际值为基础四舍五入取整估算。随着运营年度增加，建议逐步追加实际数据替代估算值。', sz=Pt(10))
ap('5.利率变动风险：银行借款利率已从6.95%下调至5.0%，如后续LPR继续下行，实际融资本息可能进一步减少。', sz=Pt(10))
ap('6.本报告系征求意见稿。如需调整资本金回报计算方式、运维成本/第三方收入估算参数或融资还款计划，可进一步修订。', sz=Pt(10))

doc.add_paragraph()
doc.add_paragraph()
ap('四川融策会计师事务所有限公司', align=WD_ALIGN_PARAGRAPH.RIGHT, sz=Pt(12))
ap('二〇二六年七月二十九日', align=WD_ALIGN_PARAGRAPH.RIGHT, sz=Pt(12))

outpath = r'C:\Users\scrccpa\Desktop\恩阳医养园PPP项目可用性付费测算报告（修订版v3-20260729）.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
print(f'Size: {os.path.getsize(outpath):,} bytes')
print(f'CapTotal={cap_total:,.2f} BankTotal={bank_total:,.2f} OpTotal={op_cost_total:,.2f} IncTotal={income_total:,.2f}')
print(f'GrandTotal={grand_total:,.2f} (≈{grand_total/1e8:.2f}亿)')
