import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx')

print(f'Total paragraphs: {len(doc.paragraphs)}')

# TOC
print('\n=== TOC area ===')
for i, p in enumerate(doc.paragraphs[:80]):
    t = p.text.strip()
    if not t:
        continue
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold:
        print(f'  [{i}] {t[:100]}')

# Section boundaries
print('\n=== Section content boundaries ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and len(t) < 20:
        if any(kw in t for kw in ['人力资源篇', '财务管理篇', '业务部管理篇', '业务质控篇', '行政综合篇', '薪酬管理', '绩效考核', '员工手册', '项目管理规范', '业务承接', '印章', '信息安全', '制度发布', '档案管理']):
            print(f'  [{i}] {t[:100]}')
    # Show doc title blocks
    if is_bold and len(t) > 5 and t.startswith('RC-'):
        print(f'  [{i}] {t[:100]}')
