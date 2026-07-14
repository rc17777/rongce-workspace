# -*- coding: utf-8 -*-
"""批量提取校园餐制度/案例文件为纯文本，供深读。"""
import os, sys, zipfile, re, glob
sys.stdout.reconfigure(encoding='utf-8')

SRC = r"C:\Users\scrccpa\Desktop\若尔盖审计局提供\校园餐相关文件"
OUT = r"C:\Users\scrccpa\.openclaw\workspace\projects\若尔盖校园餐审计\raw_text"
os.makedirs(OUT, exist_ok=True)

log = []

def safe_name(full):
    rel = os.path.relpath(full, SRC)
    return rel.replace("\\", "__").replace("/", "__")

def write_txt(full, text):
    name = safe_name(full)
    base = os.path.splitext(name)[0] + ".txt"
    p = os.path.join(OUT, base)
    with open(p, "w", encoding="utf-8") as f:
        f.write(text or "")
    n = len(text or "")
    log.append((rel_of(full), n))
    print(f"[OK {n:>7}] {os.path.basename(full)}")

def rel_of(full):
    return os.path.relpath(full, SRC)

# ---------- PDF ----------
def extract_pdf(full):
    import fitz
    doc = fitz.open(full)
    parts = []
    for i, page in enumerate(doc):
        parts.append(page.get_text())
    doc.close()
    txt = "\n".join(parts)
    # if almost empty -> likely scanned image
    if len(txt.strip()) < 50:
        txt = "[[可能为扫描件/图片型PDF，无文本层，需OCR]]\n" + txt
    return txt

# ---------- DOCX ----------
def extract_docx(full):
    import docx
    d = docx.Document(full)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)

# ---------- XLSX ----------
def extract_xlsx(full):
    import openpyxl
    wb = openpyxl.load_workbook(full, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"=== Sheet: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)

# ---------- OFD (zip+xml) ----------
def extract_ofd(full):
    parts = []
    try:
        z = zipfile.ZipFile(full)
        # OFD text is in <TextCode> elements inside Content xml
        names = [n for n in z.namelist() if n.lower().endswith('.xml')]
        for n in names:
            try:
                data = z.read(n).decode('utf-8', errors='ignore')
            except Exception:
                continue
            # extract TextCode contents
            codes = re.findall(r'<ofd:TextCode[^>]*>(.*?)</ofd:TextCode>', data, re.S)
            if not codes:
                codes = re.findall(r'<TextCode[^>]*>(.*?)</TextCode>', data, re.S)
            if codes:
                parts.append(f"--- {n} ---")
                parts.extend(codes)
        z.close()
    except Exception as e:
        return f"[[OFD解析失败: {e}]]"
    txt = "\n".join(parts)
    if len(txt.strip()) < 30:
        return "[[OFD无可提取文本，可能为图片型]]\n" + txt
    return txt

TASKS = []
for full in glob.glob(os.path.join(SRC, "**", "*"), recursive=True):
    if os.path.isdir(full):
        continue
    ext = os.path.splitext(full)[1].lower()
    if ext in ('.docx', '.pdf', '.xlsx', '.ofd'):
        TASKS.append(full)

for full in TASKS:
    ext = os.path.splitext(full)[1].lower()
    try:
        if ext == '.pdf':
            t = extract_pdf(full)
        elif ext == '.docx':
            t = extract_docx(full)
        elif ext == '.xlsx':
            t = extract_xlsx(full)
        elif ext == '.ofd':
            t = extract_ofd(full)
        else:
            continue
        write_txt(full, t)
    except Exception as e:
        print(f"[ERR] {rel_of(full)} -> {e}")
        log.append((rel_of(full), -1))

print("\n==== SUMMARY ====")
for rel, n in sorted(log, key=lambda x: x[1]):
    print(f"{n:>8}  {rel}")
