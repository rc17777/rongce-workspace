import sys, os, json
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

path = r'C:\Users\scrccpa\Desktop\财政局=财务决算\1.马尔康城市环境质量提升（房屋建筑）工程（州级）项目\1.马尔康城市环境质量提升（房屋建筑）工程（州级）项目----竣工财务决算审核报告（征求意见稿20260720）\马尔康城市环境质量提升（房屋建筑）工程（州级）项目-竣工财务决算审核报告.docx'

doc = Document(path)
out = []
# iterate body elements in order
from docx.oxml.ns import qn
body = doc.element.body
tbl_idx = 0
para_idx = 0
tables = doc.tables
paras = doc.paragraphs

# Simple approach: dump paragraphs then tables with markers
for i, p in enumerate(paras):
    t = p.text.strip()
    if t:
        style = p.style.name if p.style else ''
        out.append(f'[P{i:03d}|{style}] {t}')

out.append('\n===== TABLES =====')
for ti, tbl in enumerate(tables):
    out.append(f'--- TABLE {ti} ({len(tbl.rows)} rows x {len(tbl.columns)} cols) ---')
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip().replace('\n',' ') for c in row.cells]
        # dedupe merged cells
        dedup = []
        for c in cells:
            if not dedup or c != dedup[-1]:
                dedup.append(c)
        out.append(f'R{ri}: ' + ' | '.join(dedup))

text = '\n'.join(out)
os.makedirs(r'C:\Users\scrccpa\.openclaw\workspace\projects\财政局财务决算复核', exist_ok=True)
with open(r'C:\Users\scrccpa\.openclaw\workspace\projects\财政局财务决算复核\report_text.txt','w',encoding='utf-8') as f:
    f.write(text)
print(f'chars={len(text)} paras={len(paras)} tables={len(tables)}')
print(text[:3000])
