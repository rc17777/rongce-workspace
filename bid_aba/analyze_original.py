import sys, os
sys.stdout.reconfigure(encoding='utf-8')
import docx

src = r'C:\Users\scrccpa\Desktop\阿坝州财政局竣工财务决算审核项目_投标文件_高端版.docx'
dst = r'D:\openclaw-workspace\bid_aba\work_base.docx'

doc = docx.Document(src)
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')

headings = [p for p in doc.paragraphs if p.style.name and 'Heading' in p.style.name]
print(f'Headings: {len(headings)}')
for h in headings:
    print(f'  [{h.style.name}] {h.text[:50]}')

chars = sum(len(p.text) for p in doc.paragraphs)
table_chars = 0
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            table_chars += len(c.text)
print(f'Para chars: {chars}')
print(f'Table chars: {table_chars}')
print(f'Total chars: {chars + table_chars}')

doc.save(dst)
print(f'Saved: {os.path.getsize(dst)} bytes')
