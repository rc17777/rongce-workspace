import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

pth = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'
doc = Document(pth)

para_count = len(doc.paragraphs)
table_count = len(doc.tables)
total_chars = sum(len(p.text) for p in doc.paragraphs)
# Also count table text
table_chars = 0
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            table_chars += len(cell.text)

headings = [p.style.name for p in doc.paragraphs if 'Heading' in str(p.style.name)]
img_count = len([p for p in doc.paragraphs if '图：' in p.text])

print(f'段落数: {para_count}')
print(f'表格数: {table_count}')
print(f'图片引用: {img_count}')
print(f'段落文字: {total_chars} chars')
print(f'表格文字: {table_chars} chars')
print(f'合计: {total_chars + table_chars} chars')
print(f'标题级别: {len(headings)}')
print(f'文件大小: {os.path.getsize(pth)} bytes')
print(f'预估字数: {(total_chars + table_chars) // 2} 字（按中文字符估算）')
