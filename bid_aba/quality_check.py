import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx')

# 把所有段落文本和标题收集起来
all_texts = []
heading_texts = []
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    all_texts.append(txt)
    if p.style.name and 'Heading' in p.style.name:
        heading_texts.append(txt)

full_text = '\n'.join(all_texts)

print("=" * 60)
print("一、错别字检查")
print("=" * 60)

typo_checks = [
    (r'融策', '✅ "融策" 出现次数'),
    (r'阿坝[州县]', '✅ "阿坝" 出现次数'),
    (r'财务会计所', '❌ 可能是"会计师事务所"的错写'),
    (r'注册会计[^师]', '❌ "注册会计师"后跟了奇怪的字'),
    (r'四川融策工程咨询', '✅ "四川融策工程咨询"'),
    (r'财政厅', '✅ "财政厅"'),
    (r'李开|周贤伟|赖德明|贺珊|谭维|廖彬微|佘红丽|吴博|吴胜男', '✅ 所有人员姓名出现'),
    # 常见错词
    (r'范围\s*内', '✅ "范围内"'),
    (r'范筹', '❌ 可能应为"范畴"'),
    (r'部署', '✅ "部署"'),
    (r'布署', '❌ 可能应为"部署"'),
    (r'的的', '❌ 重复字"的的"'),
    (r'了了', '❌ 重复字"了了"'),
    (r'，,', '❌ 中英文逗号混用'),
    (r'，\s*，', '❌ 重复逗号'),
    (r'。。', '❌ 重复句号'),
    (r'；;', '❌ 中英文分号混用'),
    (r'份额', '❌ 可能应为"份额"'),
    (r'其它', '✅ "其它"（虽可用"其他"，非错误）'),
    (r'即然', '❌ 应为"既然"'),
    (r'既使', '❌ 应为"即使"'),
    (r'做为', '❌ 应为"作为"'),
    (r'不尽', '❌ 可能应为"不仅"'),
    (r'不符和', '❌ 可能应为"不符合"'),
    (r'因地置宜', '❌ 应为"因地制宜"'),
    (r'以经', '❌ 可能应为"已经"'),
    (r'年今', '❌ 可能应为"今年"'),
    (r'于[0-9]', '✅ "于年份"格式'),
    (r'大约[0-9]', '✅ "大约"后接数字'),
    (r'概算[^执]', '⚠ "概算"后字异常'),
    (r'四川融策会计师事务所[^限]', '⚠ 公司名后异常'),
    (r'报告差错率[^不]', '⚠ "报告差错率"后异常'),
    (r'客户满意度[^不]', '⚠ "客户满意度"后异常'),
    (r'及时交付率[^1]', '⚠ "及时交付率"后异常'),
]

for pattern, desc in typo_checks:
    matches = re.findall(pattern, full_text)
    if '❌' in desc:
        print(f"{desc}: 找到{len(matches)}处")
        if matches:
            for m in matches[:3]:
                # Find context
                idx = full_text.find(m)
                ctx = full_text[max(0,idx-15):idx+len(m)+15]
                print(f"  ...{ctx}...")
    else:
        print(f"{desc}: {len(matches)}次")

print()
print("=" * 60)
print("二、逻辑一致性检查")
print("=" * 60)

# 1. 检查数字/金额的前后一致性
# 2. 检查法规引用年份是否正确
law_years = {
    '预算法': '2014',
    '招投标法': '2017',
    '政府采购法': '2014',
    '政府投资条例': '2019',
    '财政部令第81号': '2016',
    '503号': '2016',
    '504号': '2016',
    '724号': '2003',
    '中发〔2018〕34号': '2018',
}

print("\n【法规引用年份核查】")
for kw, expected_year in law_years.items():
    # Find the context where this keyword appears
    positions = [m.start() for m in re.finditer(re.escape(kw), full_text)]
    if positions:
        year_in_context = False
        for pos in positions[:3]:
            ctx = full_text[max(0,pos-30):pos+60]
            # Check if year mentioned
            year_match = re.search(r'(19\d\d|20\d\d)', ctx)
            if year_match:
                print(f"  {kw}: 附近有年份 {year_match.group()} (预期:{expected_year})")
                year_in_context = True
                break
        if not year_in_context:
            print(f"  {kw}: 出现在文本中（预期年份:{expected_year}）")

# 3. 检查评分标准对应关系
print("\n【评分标准覆盖检查】")
score_items = {
    '实施方案': ['实施方案', '项目理解', '审核方法', '审核流程', '管理制度', '质量保证', '机构设置', '服务承诺'],
    '人员配备': ['项目负责人', '团队', '人员', '配备'],
    '报价': ['报价', '下浮', '收费'],
    '履约能力': ['业绩', '类似项目', '履约能力', '合同'],
}

for item, keywords in score_items.items():
    found = []
    for kw in keywords:
        if kw in full_text:
            found.append(kw)
    print(f"  {item}: {len(found)}/{len(keywords)}关键词找到 — {found}")

# 4. 检查人名出现位置
print("\n【人员姓名分布检查】")
names = ['李开', '周贤伟', '赖德明', '贺珊', '谭维', '廖彬微', '佘红丽', '吴博', '吴胜男']
for name in names:
    count = full_text.count(name)
    contexts = []
    for m in re.finditer(re.escape(name), full_text):
        start = max(0, m.start() - 20)
        end = min(len(full_text), m.end() + 20)
        contexts.append(full_text[start:end].replace('\n', ' ').strip())
    print(f"  {name}: {count}次 — {' | '.join(contexts[:2])}")

# 5. 检查地区名称
print("\n【地区名称一致性检查】")
regions = ['阿坝县', '汶川县', '金川县', '黑水县', '理县', '九寨沟县', '巴塘县', '色达县', '道孚县', '石渠县']
for r in regions:
    c = full_text.count(r)
    print(f"  {r}: {c}次")

print()
print("=" * 60)
print("三、图表与内容匹配检查")
print("=" * 60)

# Check image references
print("\n【图片分布】")
img_sections = {
    'fig1-audit-process.drawio.png': ['流程图', '审核流程', '三阶段'],
    'fig2-org-chart.drawio.png': ['组织结构', '组织架构', '组织机构'],
    'fig3-quality-control.drawio.png': ['复核', '质量控制', '质量'],
    'fig4-confidentiality.drawio.png': ['保密', '信息安全'],
    'fig5-fund-flow.drawio.png': ['资金', '资金流向', '资金管理'],
    'fig6-cost-audit.drawio.png': ['造价', '工程', '成本'],
}

for img_name in ['fig1-audit-process.drawio.png', 'fig2-org-chart.drawio.png', 
                 'fig3-quality-control.drawio.png', 'fig4-confidentiality.drawio.png',
                 'fig5-fund-flow.drawio.png', 'fig6-cost-audit.drawio.png']:
    # Check if image is near matching content
    related_kws = img_sections.get(img_name, [])
    kw_found = any(kw in full_text for kw in related_kws)
    print(f"  {img_name}: 相关关键词在正文中{'✅' if kw_found else '❌'}")
    
    # Check covers
covers = ['cover1', 'cover2', 'cover3', 'cover4', 'cover5', 'cover6', 'cover7',
          'cover5-ch6', 'cover5-ch7', 'cover6-ch8', 'cover6-ch9', 'cover7-ch10', 'cover7-ch11']
print("\n【章节封面图】")
for cover in covers:
    # Check if cover image file exists
    import os
    exists = os.path.exists(f'D:\\openclaw-workspace\\bid_aba\\work_dir\\{cover}.drawio.png')
    print(f"  {cover}.drawio.png: {'✅ 文件存在' if exists else '❌ 文件缺失'}")

print()
print("=" * 60)
print("四、格式问题检查")
print("=" * 60)

# Check for missing fields (XXX placeholders)
xxx_count = len(re.findall(r'XXX|xxx|【待填】|________', full_text))
print(f"\n【占位符检查】: {'✅ 无占位符' if xxx_count == 0 else f'❌ 找到{xxx_count}处占位符（XXX等）'}")

# Check table count
print(f"\n【表格统计】: {len(doc.tables)} 张表格")
for i, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.rows[0].cells) if t.rows else 0
    # Get first cell of first row
    header = t.rows[0].cells[0].text[:30] if t.rows else ''
    print(f"  表{i+1}: {rows}行×{cols}列 - 首字段: {header}")

# Check paragraph count
print(f"\n【段落统计】: {len([p for p in doc.paragraphs if p.text.strip()])} 非空段落")

print()
print("=" * 60)
print("五、整体数据")
print("=" * 60)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f"总字符数: {total}")
print(f"估算字数: {int(total/2)}")
print(f"文件大小: {os.path.getsize('D:\\openclaw-workspace\\bid_aba\\阿坝州财政局竣工财务决算审核_投标方案.docx') / 1024:.0f} KB")
