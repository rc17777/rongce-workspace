import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

doc = Document(PTH)

chapter_covers = [
    ('一、项目理解与总体思路', 'cover1-intro.drawio.png'),
    ('二、审核依据与政策解读', 'cover2-laws.drawio.png'),
    ('三、审核范围与审核内容', 'cover3-scope.drawio.png'),
    ('四、审核方法与技术路线', 'cover4-methods.drawio.png'),
    ('五、审核程序与进度安排', 'cover4-methods.drawio.png'),
    ('六、重点难点分析与对策', 'cover5-quality.drawio.png'),
    ('七、审计管理制度与质量保证', 'cover5-quality.drawio.png'),
    ('八、公司概况与服务能力', 'cover6-company.drawio.png'),
    ('九、项目团队配备', 'cover6-company.drawio.png'),
    ('十、类似业绩与履约能力', 'cover7-performance.drawio.png'),
    ('十一、服务承诺与保障措施', 'cover7-performance.drawio.png'),
]

inserted = 0
for heading_text, cover_file in chapter_covers:
    img_path = os.path.join(IMG, cover_file)
    if not os.path.exists(img_path):
        print(f"  NOT FOUND: {img_path}")
        continue
    
    found = False
    for p in doc.paragraphs:
        if heading_text in p.text:
            # Use python-docx to add image before this paragraph
            # Trick: get the element and insert a new paragraph before it
            parent = p._element.getparent()
            idx = list(parent).index(p._element)
            
            # Add image paragraph using python-docx's run.add_picture but via direct XML
            from docx.oxml import OxmlElement
            
            # Create the image paragraph
            img_para = OxmlElement('w:p')
            
            # Add empty pPr
            pPr = OxmlElement('w:pPr')
            img_para.append(pPr)
            
            # Add run with picture
            run_elem = OxmlElement('w:r')
            
            # Add drawing
            dw_elem = OxmlElement('w:drawing')
            
            # wp:inline with proper namespace
            import lxml.etree as etree
            NSWP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            inline = etree.SubElement(dw_elem, f'{{{NSWP}}}inline')
            inline.set('distT', '0')
            inline.set('distB', '0')
            inline.set('distL', '0')
            inline.set('distR', '0')
            
            # extent
            ext = etree.SubElement(inline, f'{{{NSWP}}}extent')
            cx = int(6.5 * 914400)
            cy = int(6.5 * 300 / 1100 * 914400)
            ext.set('cx', str(cx))
            ext.set('cy', str(cy))
            
            # effectExtent
            eff = etree.SubElement(inline, f'{{{NSWP}}}effectExtent')
            eff.set('l', '0'); eff.set('t', '0'); eff.set('b', '0'); eff.set('r', '0')
            
            # docPr
            dp = etree.SubElement(inline, f'{{{NSWP}}}docPr')
            dp.set('id', str(100 + inserted))
            dp.set('name', f'Picture {inserted+1}')
            
            # graphic
            NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            graphic = etree.SubElement(inline, f'{{{NS_A}}}graphic')
            
            # graphicData
            gd = etree.SubElement(graphic, f'{{{NS_A}}}graphicData')
            gd.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
            
            # pic:pic
            NS_PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
            pic = etree.SubElement(gd, f'{{{NS_PIC}}}pic')
            
            # nvPicPr
            nv = etree.SubElement(pic, f'{{{NS_PIC}}}nvPicPr')
            cNvPr = etree.SubElement(nv, f'{{{NS_PIC}}}cNvPr')
            cNvPr.set('id', '0'); cNvPr.set('name', f'Picture {inserted+1}')
            etree.SubElement(nv, f'{{{NS_PIC}}}cNvPicPr')
            
            # blipFill
            bf = etree.SubElement(pic, f'{{{NS_PIC}}}blipFill')
            blip = etree.SubElement(bf, f'{{{NS_A}}}blip')
            NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            
            # Add image relationship
            rId = f'rIdC{inserted+1}'
            blip.set(f'{{{NS_R}}}embed', rId)
            stretch = etree.SubElement(bf, f'{{{NS_A}}}stretch')
            etree.SubElement(stretch, f'{{{NS_A}}}fillRect')
            
            # spPr
            sp = etree.SubElement(pic, f'{{{NS_PIC}}}spPr')
            xfrm = etree.SubElement(sp, f'{{{NS_A}}}xfrm')
            off = etree.SubElement(xfrm, f'{{{NS_A}}}off')
            off.set('x', '0'); off.set('y', '0')
            ext = etree.SubElement(xfrm, f'{{{NS_A}}}ext')
            ext.set('cx', str(cx)); ext.set('cy', str(cy))
            prst = etree.SubElement(sp, f'{{{NS_A}}}prstGeom')
            prst.set('prst', 'rect')
            
            run_elem.append(dw_elem)
            img_para.append(run_elem)
            
            # Add relationship to document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
            
            part = doc.part
            img_rel = part.relate_to(img_path, rel_type, is_external=False)
            
            # Update the r:embed to use actual relationship ID
            blip.set(f'{{{NS_R}}}embed', img_rel.rId)
            
            # Insert before heading
            parent.insert(idx, img_para)
            found = True
            inserted += 1
            print(f"  ✅ Inserted {cover_file} before '{heading_text[:20]}...'")
            break
    
    if not found:
        print(f"  ❌ Heading not found: '{heading_text}'")

doc.save(PTH)
print(f'Saved. size={os.path.getsize(PTH)}')
print(f'Total covers inserted: {inserted}')
