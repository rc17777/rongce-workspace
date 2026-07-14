import sys, os, json
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

doc = Document()

# Styles
doc.styles['Normal'].font.name = '仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.styles['Normal'].paragraph_format.line_spacing = 1.5

def H1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def H2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def H3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def P(text, bold=False, fs=12, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(fs)
    r.font.name = '仿宋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def T(hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hds):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255,255,255)
        r.font.size = Pt(fs)
        r.font.name = 'Microsoft YaHei'
        el = OxmlElement('w:shd'); el.set(qn('w:fill'), hc); el.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(el)
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(ct))
            r.font.size = Pt(fs)
            r.font.name = 'Microsoft YaHei'
            if ri % 2 == 1:
                el = OxmlElement('w:shd'); el.set(qn('w:fill'), 'F2F4F4'); el.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(el)
    doc.add_paragraph()

def I(fname, w=6.0, cap=None):
    fp = os.path.join(IMG, fname)
    if os.path.exists(fp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(fp, width=Inches(w))
        if cap:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(f'图：{cap}')
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(128,128,128)
        doc.add_paragraph()

# ========== COVER ==========
for _ in range(6):
    doc.add_paragraph()
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('2026年阿坝州财政局州本级\n建设项目竣工财务决算审核项目')
r.font.size = Pt(22); r.bold = True; r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

cp2 = doc.add_paragraph(); cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cp2.add_run('应  标  方  案')
r2.font.size = Pt(18); r2.font.name = '黑体'
r2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()
cp3 = doc.add_paragraph(); cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = cp3.add_run('四川融策会计师事务所有限公司\n四川融策工程咨询有限公司')
r3.font.size = Pt(14)

cp4 = doc.add_paragraph(); cp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = cp4.add_run('二〇二六年六月')
r4.font.size = Pt(14)

doc.add_page_break()

# ========== TOC ==========
doc.add_heading('目  录', level=0)
toc = [
    ('一、项目理解与总体思路', '1'),
    ('二、审核依据与政策解读', '5'),
    ('三、审核范围与审核内容', '11'),
    ('四、审核方法与技术路线', '20'),
    ('五、审核程序与进度安排', '27'),
    ('六、重点难点分析与对策', '33'),
    ('七、审计管理制度与质量保证', '40'),
    ('八、公司概况与服务能力', '48'),
    ('九、项目团队配备', '56'),
    ('十、类似业绩与履约能力', '60'),
    ('十一、服务承诺与保障措施', '65'),
]
for item, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{item}  {"·"*15}  {page}')
    r.font.size = Pt(12)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

doc.add_page_break()

print('Cover + TOC done')
doc.save(PTH)
print(f'Saved initial: {os.path.getsize(PTH)}')
