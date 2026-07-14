import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

doc = Document(PTH)
body = doc.element.body

chapter_covers = [
    ('十、类似业绩与履约能力', 'cover7-ch10.drawio.png'),
    ('十一、服务承诺与保障措施', 'cover7-ch11.drawio.png'),
]

for heading_text, cover_file in chapter_covers:
    img_path = os.path.join(IMG, cover_file)
    if not os.path.exists(img_path):
        print(f"NOT FOUND: {img_path}")
        continue
    
    heading_elem = None
    for p in doc.paragraphs:
        if heading_text in p.text:
            heading_elem = p._element
            break
    if heading_elem is None:
        print(f"NOT FOUND: {heading_text}")
        continue
    
    idx = list(body).index(heading_elem)
    # Remove old shared cover image before this heading
    for j in range(1, 5):
        if idx >= j:
            candidate = body[idx-j]
            tag = candidate.tag.split('}')[-1] if '}' in candidate.tag else candidate.tag
            if tag == 'p':
                dw = candidate.find('.//' + qn('w:drawing'))
                if dw is not None:
                    body.remove(candidate)
                    break
    
    new_idx = list(body).index(heading_elem)
    temp_para = doc.add_paragraph()
    temp_run = temp_para.add_run()
    temp_run.add_picture(img_path, width=Inches(6.5))
    img_elem = temp_para._element
    body.remove(img_elem)
    body.insert(new_idx, img_elem)
    
    blank = doc.add_paragraph()
    blank_elem = blank._element
    body.remove(blank_elem)
    body.insert(new_idx+1, blank_elem)
    print(f"✅ {cover_file} before '{heading_text[:20]}...'")

doc.save(PTH)
print(f'Done! Size: {os.path.getsize(PTH)}')
