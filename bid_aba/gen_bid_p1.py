# 阿坝州财政局竣工财务决算审核项目 - 投标文件生成脚本 v1
# 目标：完整应对评分标准，图文并茂，20万字
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'D:\openclaw-workspace\bid_aba'
IMG = os.path.join(OUT, 'work_dir')
PATH = os.path.join(OUT, '阿坝州财政局竣工财务决算审核_投标文件.docx')

def shade(cell, color):
    el = OxmlElement('w:shd')
    el.set(qn('w:fill'), color); el.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(el)

def table(doc, heads, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(heads))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(heads):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        r.font.size = Pt(fs); r.font.name = 'Microsoft YaHei'; shade(c, hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            p = c.paragraphs[0]; r = p.add_run(str(ct))
            r.font.size = Pt(fs); r.font.name = 'Microsoft YaHei'
            if ri % 2 == 1: shade(c, 'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(fs)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    if fc: r.font.color.rgb = fc
    if indent and align != WD_ALIGN_PARAGRAPH.CENTER:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def img(doc, name, w=5.5, cap=None):
    pth = os.path.join(IMG, name)
    if os.path.exists(pth):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(pth, width=Inches(w))
        if cap:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(f'图：{cap}'); cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(128,128,128); cr.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

doc = Document()
st = doc.styles['Normal']
st.font.name = '仿宋'; st.font.size = Pt(12)
st._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)

# ============ 封面 ============
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026年阿坝州财政局\n州本级建设项目竣工财务决算审核项目')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0,51,102)
r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('比  选  申  请  文  件')
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(178,34,34)
r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
for _ in range(8): doc.add_paragraph()
for line in ['比选申请人（公章）：四川融策会计师事务所（普通合伙）',
             '法定代表人或授权代理人：', '日期：2026年  月  日']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.size = Pt(14)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.add_page_break()

# ============ 目录 ============
H(doc, '目  录', 1)
toc = ['第一章  比选申请函','第二章  法定代表人身份证明书及授权委托书',
       '第三章  资格证明材料','第四章  报价一览表','第五章  商务应答表',
       '第六章  技术、服务应答表','第七章  项目服务实施方案',
       '  7.1  项目理解与总体思路','  7.2  工作重点与难点分析',
       '  7.3  审计流程与工作进度安排','  7.4  审计管理制度',
       '  7.5  审计机构设置与岗位职责','  7.6  审计服务质量保证措施',
       '  7.7  审计进度保证措施','  7.8  保密管理保证措施',
       '  7.9  廉洁从业承诺与保障','第八章  类似业绩与履约能力',
       '第九章  公司基本情况介绍','第十章  人员配备与团队介绍','附件']
for item in toc:
    p = doc.add_paragraph(); r = p.add_run(item)
    r.font.size = Pt(11); r.font.name = '仿宋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.add_page_break()

# ============ 第一章 比选申请函 ============
H(doc, '第一章  比选申请函', 1)
P(doc, '致：阿坝州财政局', bold=True)
doc.add_paragraph()
P(doc, '我单位四川融策会计师事务所（普通合伙），愿意参加"2026年阿坝州财政局州本级建设项目竣工财务决算审核项目"的比选。在审查和全面理解了所提供的比选文件后，以下签字人在此作为比选申请人代表参加本次比选。')
doc.add_paragraph()
for item in [
    '一、我单位对提交的比选申请文件负责。贵方或授权代表可对我单位进行查询或调查，以证实提交的声明、文件和资料的真实性。',
    '二、我单位完全理解比选人因法律和政策原因取消比选以及拒绝所有的比选申请，并对此类任何行动不承担任何责任。',
    '三、如我单位中选，我单位承诺按比选文件要求签订合同，并将严格按照有关法律法规等相关政策开展工作，按投标文件承诺的质量标准和时限要求完成全部审核任务。',
    '四、比选申请文件有效期为递交比选申请文件截止日期后90个日历日内有效。',
    '五、我单位郑重承诺，在项目执行过程中，严格遵循《中国注册会计师审计准则》、《基本建设财务规则》及比选文件要求的执业规范，确保审核报告的真实性、完整性和准确性。',
    '六、我单位承诺依法依规履行保密义务，对在执业过程中知悉的国家秘密、商业秘密及敏感信息严格保密，不向任何第三方泄露。',
]:
    P(doc, item)
doc.add_paragraph()
P(doc, '需要核实的资料，贵方或授权代表可以向下列人员查询：')
table(doc, ['序号','联系人','职务','联系电话'], [['1','XXX','项目负责人','XXX-XXXXXXXX'],['2','XXX','技术负责人','XXX-XXXXXXXX']])
doc.add_paragraph()
P(doc, '比选申请人名称（公章）：四川融策会计师事务所（普通合伙）', indent=False)
P(doc, '法定代表人或代理人（签字）：', indent=False)
P(doc, '日期：    年    月    日', indent=False)
doc.add_page_break()

# ============ 第二章 法定代表人身份证明书 ============
H(doc, '第二章  法定代表人身份证明书及授权委托书', 1)
H(doc, '2.1 法定代表人身份证明书', 2)
P(doc, '致：阿坝州财政局', bold=True)
doc.add_paragraph()
P(doc, '（姓名），身份证号：（身份证号），现任四川融策会计师事务所（普通合伙）的法定代表人，特此证明。')
P(doc, '本项目涉及相关文件中，若签名及印鉴与本资格证明书签字及印鉴样本不符的，本申请人不予承认。')
doc.add_paragraph()
P(doc, '法定代表人：（签字或盖法定代表人印章）', indent=False)
P(doc, '注：身份证所在页盖章视为有效。', fs=10, indent=False)
doc.add_paragraph()
P(doc, '比选申请人名称（公章）：', indent=False)
P(doc, '日期：    年    月    日', indent=False)
doc.add_paragraph()
H(doc, '2.2 法定代表人授权委托书', 2)
P(doc, '本授权委托书声明：我（法定代表人姓名）系四川融策会计师事务所（普通合伙）的法定代表人，现授权（代理人姓名）为我的委托代理人，以本单位的名义参加贵单位组织的"2026年阿坝州财政局州本级建设项目竣工财务决算审核项目"的比选活动。委托代理人在比选活动和合同签订过程中所签署的一切文件和处理与之有关的一切事务，我本人及单位均予以承认并承担其所产生的所有权利和义务。')
P(doc, '本授权书于    年   月   日签字生效。')
P(doc, '委托代理人无转委托权。特此委托。')
doc.add_paragraph()
P(doc, '法定代表人：（签字或盖法定代表人印章）', indent=False)
P(doc, '代理人：（签字）', indent=False)
P(doc, '比选申请人名称：（并加盖公章）', indent=False)
P(doc, '说明：比选申请文件如均由比选申请人法定代表人签字的，则比选申请文件中可不提供法定代表人授权书。', fs=10, fc=RGBColor(128,128,128), indent=False)
doc.add_page_break()

# ============ 第三章 资格证明材料 ============
H(doc, '第三章  资格证明材料', 1)
H(doc, '3.1 营业执照（副本）复印件', 2)
P(doc, '（此处粘贴营业执照副本复印件并加盖公章）', fs=12, fc=RGBColor(128,128,128))
doc.add_paragraph()
H(doc, '3.2 会计师事务所执业许可证复印件', 2)
P(doc, '（此处粘贴会计师事务所执业许可证复印件并加盖公章）', fs=12, fc=RGBColor(128,128,128))
doc.add_paragraph()
H(doc, '3.3 资格承诺函', 2)
P(doc, '项目名称：2026年阿坝州财政局州本级建设项目竣工财务决算审核项目')
P(doc, '致：阿坝州财政局', bold=True)
doc.add_paragraph()
P(doc, '我单位四川融策会计师事务所（普通合伙）作为"2026年阿坝州财政局州本级建设项目竣工财务决算审核项目"的比选申请人在此郑重承诺：')
for p in [
    '1、具有独立承担民事责任的能力；','2、具有良好的商业信誉和健全的财务会计制度；',
    '3、具有履行合同所必需的设备和专业技术能力；','4、具有依法缴纳税收和社会保障资金的良好记录；',
    '5、参加本次比选活动前三年内，在经营活动中没有重大违法记录；','6、法律、行政法规规定的其他条件；',
    '7、我单位非联合体参加此次比选。',
    '8、我单位具有近年（2023年1月1日至参选截止时间）不少于1个类似业绩。类似业绩是指：财务决算或财务决算审核相关业绩。',
    '9、我单位具有有效的《会计师事务所执业许可证》。']:
    P(doc, p)
doc.add_paragraph()
P(doc, '我单位对于以上承诺的真实性负责。如有不实，我单位愿承担由此产生的一切法律责任和后果。')
doc.add_paragraph()
P(doc, '比选申请人名称（公章）：', indent=False)
P(doc, '法定代表人或代理人（签字）：', indent=False)
P(doc, '日期：    年    月    日', indent=False)
doc.add_page_break()

# ============ 第四章 报价一览表 ============
H(doc, '第四章  报价一览表', 1)
table(doc, ['序号','项目名称','数量','单位','报价（下浮率）','服务期限'],
       [['1','2026年阿坝州财政局州本级建设项目竣工财务决算审核项目','1','项','下浮：XX%','1年']], fs=11)
P(doc, '说明：', bold=True, indent=False)
for n in [
    '1、年度服务费实行总额25万元控制，每个项目实际结算按《四川省会计师事务所服务收费管理办法》（川发改价格〔2013〕901号）计费标准以成交下浮费率计取。',
    '2、报价包含完成本项目所涉及人员工资、办公费用、交通费用、通讯费用、人员食宿费用、设备设施投入、税费等一切费用。',
    '3、本报价在合同履行过程中是固定不变的。']:
    P(doc, n, fs=11)
doc.add_paragraph()
P(doc, '供应商名称：四川融策会计师事务所（普通合伙）', indent=False)
P(doc, '法定代表人或授权代表：（签字或盖章）', indent=False)
P(doc, '联系电话：', indent=False)
P(doc, '日期：    年    月    日', indent=False)
doc.add_page_break()

# ============ 第五章 商务应答表 ============
H(doc, '第五章  商务应答表', 1)
table(doc, ['序号','比选文件的商务条款','比选应答','说明（若有偏离请详细列出）'], [
    ['1','服务期限：1年','响应','完全响应，承诺在合同有效期内完成全部委托审核任务'],
    ['2','服务地点：阿坝州','响应','完全响应，承诺按委托方要求到达指定地点开展审核工作'],
    ['3','付款方式：审核后20日内支付','响应','完全响应，同意按比选文件约定付款方式执行'],
    ['4','验收标准和方法：按合同约定执行','响应','完全响应'],
    ['5','违约责任：按比选文件和合同约定执行','响应','完全响应各项违约责任条款'],
], fs=10)
P(doc, '说明：按照第三章"商务要求"逐条进行响应。以上各项全部响应，无偏离。')
doc.add_paragraph()
P(doc, '比选申请人名称（公章）：', indent=False)
P(doc, '日期：    年    月    日', indent=False)
doc.add_page_break()
doc.save(PATH)
print(f'OK: Part 1 saved, bytes={os.path.getsize(PATH)}')
