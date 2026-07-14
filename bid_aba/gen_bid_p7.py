import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

doc = Document(PTH)

# ===== 扩展：公司业务特色 =====
H(doc, '9.6 公司业务特色与优势领域', 2)
P(doc, '四川融策会计师事务所经过多年深耕，形成了以下鲜明业务特色：')
P(doc, '（一）政府审计全链条服务。我单位在政府审计领域形成了从审计目标设定、方案编制、现场实施、底稿编制、报告出具到后续跟踪的全链条服务能力，能够为各级政府提供一站式审计服务。特别是在绩效评价、资产清查、专项资金审计、预算执行审计等领域积累了丰富的实务经验。')
P(doc, '（二）工程审计与财务审计深度融合。不同于仅从事财务审计或仅从事工程造价审核的单一型事务所，我单位同时具备政府审计资质和工程咨询服务能力，能够实现工程审核与财务审核的无缝对接。在竣工财务决算审核中，工程造价审核和财务审计同一团队完成，避免了不同机构之间的信息传递损失和结论冲突问题，审核质量和效率显著提升。')
P(doc, '（三）民族地区服务经验丰富。我单位在阿坝州、甘孜州、凉山州等民族地区均有成功的服务案例，熟悉民族地区政府投资项目管理特点和审计要求，了解高原地区建设项目审核的特殊注意事项。')
P(doc, '（四）数字化审计能力。我单位积极探索AI技术在审计工作中的应用，研发了"1+3+5+N"数字化审计平台，能够通过数据分析辅助识别异常、智能底稿生成提高工作效率、数据交叉比对提升审核精确度。')

# 扩展：公司质量管理体系详细介绍
H(doc, '9.7 公司质量管理体系', 2)
P(doc, '我单位建立了以ISO9001质量管理体系为基础、以三级复核制度为核心的全面质量管理体系，从制度建设、过程控制、人员管理、技术保障四个方面确保服务质量。')

T(doc, ['管理维度','核心措施','执行标准','考核指标','持续改进'], [
    ['制度建设','三级复核、项目立项审批、风险评估、重大问题报告','符合执业准则和ISO9001标准','制度执行率100%','每年度制度评审修订'],
    ['过程控制','准备阶段→实施阶段→报告阶段三阶段控制','项目全流程管理规范','各阶段成果文件齐全、签字完备','季度质量检查，问题整改'],
    ['人员管理','执业资格准入、年度培训、绩效考核、淘汰机制','人均≥40学时/年，持证率100%','培训考核合格率≥95%','培训内容和方式持续优化'],
    ['技术保障','造价软件、审计系统、数据安全','正版化、安全等级保护','软件覆盖率100%','技术更新迭代年度评估'],
], hc='2980B9')

P(doc, '在质量管理体系运行过程中，我单位注重客户反馈和持续改进。每个项目完成后15日内进行客户回访，了解客户对服务的满意度及改进建议。回访结果纳入项目组的绩效考核，作为质量评价的重要依据。我单位每年开展管理评审，对质量管理体系的运行效果进行全面评价，确定改进方向。')

# 扩展：常见问题处理流程
H(doc, '7.13 常见审核问题与处理流程', 2)
P(doc, '在竣工财务决算审核实践中，经常遇到一些共性问题。我单位针对这些常见问题制定了标准化处理流程，确保问题处理的规范性和一致性。')

T(doc, ['常见问题类型','具体表现','处理流程','法律依据'],[
    ['工程量争议','对结算申报的工程量与实际工程量存在差异','现场重新核实→查阅施工记录和监理日志→三方会商确认→出具书面核定意见','《建设工程工程量清单计价规范》'],
    ['材料价格争议','合同约定材料价格与实际采购价格存在较大差异','核实采购合同和发票→查阅同期造价信息→市场询价→综合确定合理价格','施工合同约定、造价信息期刊'],
    ['设计变更手续不完备','变更签证缺少审批签字或事后补签','核实变更必要性和真实性→查阅设计变更通知书和会议纪要→确认变更审批程序','《建设工程勘察设计管理条例》'],
    ['费用标准争议','对建设单位管理费、监理费等费用的计取标准存在分歧','核查费用计提基数是否正确→比对规定标准→计算正确金额','《基本建设项目建设成本管理规定》'],
    ['资产移交不规范','竣工后资产未及时办理移交手续','核查资产清单与实物是否一致→督促办理移交手续→形成书面移交记录','《基本建设财务规则》'],
    ['资金管理问题','专项资金被挪用、结余资金未及时处理','资金流向调查→逐笔核实→提出整改建议→跟踪整改落实情况','《预算法》、《基本建设财务规则》'],
], hc='E74C3C')

P(doc, '对于上述常见问题，我单位审核人员按照标准化流程进行处理，确保问题处理的合规性和一致性。同时，对于特别复杂或争议较大的问题，启动专家咨询机制，征询法律、造价、技术等领域专家的意见，确保问题处理的专业性和公正性。')

# 扩展：信息化手段应用
H(doc, '7.14 信息化手段在审核中的应用', 2)
P(doc, '我单位积极应用信息化手段提升审核工作效率和质量，具体体现在以下方面：')
P(doc, '（一）电子数据采集与分析。对于采用会计电算化系统和工程管理信息系统的项目，我单位支持直接读取电子账套数据和工程管理数据进行审核分析，减少人工重复录入，提高审核效率和准确性。通过建立数据分析模型，快速识别异常数据和可疑交易，为确定审核重点提供依据。')
P(doc, '（二）造价辅助审核系统。我单位配备了广联达、斯维尔等主流造价软件，能够对电子工程量清单进行快速校核，自动计算工程量并生成比对分析报告。对于结构复杂的工程，利用BIM模型进行三维可视化审核，提高工程量核实的精确度。')
P(doc, '（三）电子工作底稿系统。我单位建立了标准化的电子工作底稿系统，审核人员在系统中按照统一模板编制工作底稿，实现底稿编制、复核、审批、归档的全流程电子化管理。系统内置了行业常用的审核程序表单和复核清单，提高底稿编制的规范性和效率。')
P(doc, '（四）远程审核技术。对于因交通、气候等原因无法进行现场审核的偏远项目，我单位支持通过视频会议、文件网络传输、远程屏幕共享等方式进行远程审核。对于需要有形证据的审核事项，采用现场照片、视频记录等方式替代部分现场核实工作。远程审核能够有效降低差旅成本、缩短审核周期。')
P(doc, '（五）大数据辅助决策。我单位建立了历史项目数据库，积累了各类项目的造价指标、费用标准、常见问题等信息资源。在开展新项目审核时，可以利用历史数据进行横向和纵向比对分析，快速识别异常数据，提高审核工作的针对性和有效性。')

doc.save(PTH)
print(f'OK: Part 7 (bulk expansion) complete, bytes={os.path.getsize(PTH)}')
