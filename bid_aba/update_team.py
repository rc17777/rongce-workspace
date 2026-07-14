import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

doc = Document(PTH)

# ===== 替换人员配置表 =====
H(doc, '9.6 项目团队人员名单', 2)
P(doc, '我单位拟为本项目配备以下专业审核团队，全部人员均为我单位正式在册员工，具有相应的执业资格和丰富的实务经验：')

# 团队表
T(doc, ['序号','职务','姓名','职称证书/注册证书','项目职责'],[
    ['1','项目负责人','李开','高级会计师/注册会计师','全面负责项目组织实施和质量把控，审核报告终审签字'],
    ['2','团队专业技术人员','周贤伟','中级会计师/注册会计师','工程造价审核、工程资料审查、现场核实'],
    ['3','团队专业技术人员','赖德明','注册会计师','财务收支审核、资金合规审查、工作底稿编制'],
    ['4','团队专业技术人员','贺珊','注册会计师','财务收支审核、竣工决算报表审核、报告撰写'],
    ['5','团队其他成员','谭维','中级会计师','资料审核、数据分析、底稿编制'],
    ['6','团队其他成员','廖彬微','初级会计师','资料整理、数据录入、现场配合'],
    ['7','团队其他成员','佘红丽','初级会计师','凭证审查、往来款核对、函证工作'],
    ['8','团队其他成员','吴博','初级会计师','工程资料审核、工程量核对、现场拍照取证'],
    ['9','团队其他成员','吴胜男','初级会计师','资料管理、档案整理、后勤保障'],
], hc='2980B9')

P(doc, '以上团队人员均具备相应的执业资格和实务经验，项目负责人李开具有高级会计师职称和注册会计师执业资格，主持完成过多项政府投资项目竣工财务决算审核工作，专业功底扎实，组织协调能力强。')

# ===== 替换业绩表 =====
H(doc, '10.5 阿坝州及民族地区业绩详情', 2)
P(doc, '我单位在阿坝州及周边民族地区积累了丰富的竣工财务决算审核服务经验，以下为近年来在阿坝州和甘孜州的主要业绩：')

T(doc, ['年份','项目名称','采购人/项目单位','区域'],[
    ['2026年','阿坝县财政局采购政府性投资项目竣工财务决算机构','阿坝县财政局','阿坝州'],
    ['2025年','阿坝州文化中心项目竣工财务决算审核','阿坝州国有资产投资管理有限公司','阿坝州'],
    ['2024年','汶川县财政局政府性投资项目竣工财务决算审核','汶川县财政局','阿坝州'],
    ['2024年','金川县财政局绩效评价、资产清查、财务决算、审计咨询服务','金川县财政局','阿坝州'],
    ['2024年','黑水县财政局采购项目竣工财务决算服务机构','黑水县财政局','阿坝州'],
    ['2024年','阿坝县财政局采购政府性投资项目竣工财务决算机构','阿坝县财政局','阿坝州'],
    ['2024-2026年','理县财政资金绩效评价监督检查及财务决算审核服务','理县财政局','阿坝州'],
    ['2023年','九寨沟县财政局绩效评价、财务决算咨询服务采购','九寨沟县财政局','阿坝州'],
    ['2025年','巴塘县财政局2024年财务决算采购项目','巴塘县财政局','甘孜州'],
    ['2025年','色达县2025年至2027年财决服务采购项目','色达县财政局','甘孜州'],
    ['2024年','道孚县财政局采购政府性投资项目竣工财务决算机构','道孚县财政局','甘孜州'],
    ['2024年','石渠县财政局工程竣工财务决算项目审核服务采购','石渠县财政局','甘孜州'],
], hc='1ABC9C')

P(doc, '上述业绩充分体现了：一是我单位在阿坝州政府采购市场具有较强的竞争力和服务能力，服务范围覆盖阿坝州多个县（市）。二是我单位熟悉阿坝州各级政府投资项目管理的特点和财务决算审核的要求，具备丰富的本地化服务经验。三是服务内容涵盖竣工财务决算审核、绩效评价、资产清查、财政监督检查等多类型业务，具有综合服务能力。四是在甘孜州等民族地区也有成功的服务经验，对高原民族地区项目审核有深入的理解和丰富的实践经验。')
P(doc, '特别是2024年在金川县、汶川县、黑水县、理县、阿坝县等县同时开展竣工财务决算审核服务，充分说明了我单位在阿坝州的资源调配能力和项目管理能力。2023年起服务的九寨沟县项目合同持续执行，体现了我单位的服务质量和客户满意度。2025年又中标巴塘县、色达县等甘孜州项目，业务辐射能力进一步增强。')

doc.save(PTH)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'Done. size={os.path.getsize(PTH)} chars={total} words={int(total/2)}')
