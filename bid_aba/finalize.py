import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'
ORIG = r'D:\openclaw-workspace\bid_aba\work_base.docx'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

def I(doc, name, w=5.5, cap=None):
    pth=os.path.join(IMG,name)
    if os.path.exists(pth):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(pth, width=Inches(w))
        if cap:
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cr=cp.add_run(f'图：{cap}'); cr.font.size=Pt(9)
            cr.font.color.rgb=RGBColor(128,128,128); cr.font.name='Microsoft YaHei'
        doc.add_paragraph()

doc = Document(PTH)

# 插入新增流程图5和6
I(doc, 'fig5-fund-flow.drawio.png', 5.5, '建设资金管理审核重点分析图')
I(doc, 'fig6-cost-audit.drawio.png', 5.5, '工程造价审核流程与要点图')

# 复制原始文件中的表格
orig = Document(ORIG)
table_count = 0
for t in orig.tables:
    rows_data = []
    for row in t.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_data.append(cells)
    if len(rows_data) >= 2:
        headers = rows_data[0]
        data = rows_data[1:]
        # Filter out empty tables
        if len(headers) >= 2:
            T(doc, headers, data, hc='5D6D7E' if table_count % 2 == 0 else '1F618D', fs=9)
            table_count += 1

P(doc, f'（注：以上{table_count}张原始数据表格已从原投标文件中提取并入本文件）', fs=10, fc=RGBColor(128,128,128))

# 最终统计
doc.save(PTH)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'Final: tables={len(doc.tables)} paragraphs={len(doc.paragraphs)} chars={total} words={int(total/2)} size={os.path.getsize(PTH)}')
