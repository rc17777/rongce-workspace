import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'
if not os.path.exists(PTH):
    print(f'ERROR: {PTH} not found.')
    sys.exit(1)

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

doc = Document(PTH)

# ========== 第十章 人员配备与团队介绍 ==========
H(doc, '第十章  人员配备与团队介绍', 1)

H(doc, '10.1 项目团队配置', 2)
P(doc, '针对本项目的特点和审核需求，我单位选派政治素质高、专业能力强、职业道德好的专业技术人员组成项目工作组。团队成员搭配合理、优势互补，能够全面覆盖审计、会计、工程造价、数据分析等各专业领域，确保高质量完成审核任务。')

P(doc, '项目核心团队配置一览表', bold=True, fs=11, indent=False)
T(doc, ['岗位','姓名','执业资格','职称','专业背景','从业年限'],[
    ['项目负责人','XXX','注册会计师','高级会计师','审计学','15年'],
    ['审核组长','XXX','注册会计师','会计师','会计学','8年'],
    ['技术负责人','XXX','注册造价工程师','高级工程师','土木工程','12年'],
    ['财务审核1','XXX','注册会计师','会计师','会计学','6年'],
    ['财务审核2','XXX','注册会计师','会计师','财务管理','5年'],
    ['工程审核1','XXX','注册造价工程师','工程师','工程管理','7年'],
    ['工程审核2','XXX','二级造价师','工程师','工程造价','4年'],
    ['现场核实1','XXX','助理会计师','助理','会计学','3年'],
    ['现场核实2','XXX','会计从业','助理','财务管理','2年'],
    ['资料分析1','XXX','数据分析师','中级','统计学','4年'],
    ['资料分析2','XXX','数据分析师','初级','信息系统','3年'],
], hc='27AE60')

H(doc, '10.2 项目负责人简历', 2)
P(doc, '（以下为拟派项目负责人简历）', fs=11, fc=RGBColor(128,128,128))
doc.add_paragraph()
P(doc, '姓名：XXX   |   执业资格：注册会计师   |   职称：高级会计师')
P(doc, '从业年限：XX年   |   学历：XX   |   专业：审计学')
doc.add_paragraph()
P(doc, '个人简介：XX，注册会计师、高级会计师。从事审计工作XX年，在竣工财务决算审核、经济责任审计、政府投资项目审计等领域具有丰富的执业经验。先后担任XX市、XX州等地竣工财务决算审核项目负责人，熟悉政府投资项目管理要求和审核标准。曾参与XX省财政厅、XX市审计局等委托的多个大中型政府投资项目的竣工财务决算审核工作，累计审核项目投资总额超过XX亿元。')
doc.add_paragraph()
T(doc, ['主要业绩','委托方','合同金额','完成情况'],[
    ['XX市2024年度政府投资项目竣工财务决算审核','XX市财政局','XX万元','已完成'],
    ['XX州本级2023-2024年建设项目建设项目财务决算审核','XX州财政局','XX万元','已完成'],
    ['XX县2023年基本建设项目建设项目竣工财务决算审核','XX县财政局','XX万元','已完成'],
], hc='2980B9')

H(doc, '10.3 团队专业技术人员简历', 2)
P(doc, '（以下为拟派团队专业技术成员简历）', fs=11, fc=RGBColor(128,128,128))
doc.add_paragraph()

# Person 2
P(doc, 'XXX，注册会计师，会计师职称。毕业于XX大学会计学专业，从业X年。长期从事政府审计工作，在专项资金审计、预算执行审计方面经验丰富。曾参与多个省市级政府投资项目的财务审核工作，熟悉《基本建设财务规则》等法规要求。', bold=False)

doc.add_paragraph()
# Person 3
P(doc, 'XXX，注册造价工程师，高级工程师职称。毕业于XX大学土木工程专业，从业X年。曾就职于XX建筑工程公司、XX造价咨询公司，具有丰富的工程造价审核经验。擅长市政、房建等领域的工程量审核和造价分析。', bold=False)

doc.add_paragraph()
# Person 4
P(doc, 'XXX，注册会计师，会计师职称。毕业于XX大学财务管理专业，从业X年。精通政府会计准则和预算管理制度，在账务核对、支出合规性审查方面经验丰富。', bold=False)

doc.add_paragraph()
# Person 5
P(doc, 'XXX，注册会计师，会计师事务所执业满5年。擅长财务分析、资金流向追踪、往来款项清理等，在多个竣工财务决算审核项目中承担重要审核任务。', bold=False)

H(doc, '10.4 在职证明材料', 2)
P(doc, '（此处附拟派团队人员的在职证明、执业资格证书复印件、社保缴纳证明等材料）', fs=11, fc=RGBColor(128,128,128))
doc.add_paragraph()
P(doc, '上述拟派团队人员均为我单位正式在职员工，依法签订了劳动合同并缴纳社会保险。相关人员的执业资格证书均在有效期内，符合执业要求。')
doc.add_page_break()

# ========== 附件 ==========
H(doc, '附件', 1)

H(doc, '附件一：营业执照副本复印件', 2)
P(doc, '（此处粘贴营业执照副本复印件并加盖公章）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件二：会计师事务所执业许可证复印件', 2)
P(doc, '（此处粘贴会计师事务所执业许可证复印件并加盖公章）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件三：业绩证明材料', 2)
P(doc, '（此处粘贴合同协议书复印件，按业绩一览表顺序排列）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件四：人员资质证明材料', 2)
P(doc, '（此处粘贴项目团队人员在职证明、执业资格证书和职称证书复印件）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件五：公司管理制度汇编', 2)
P(doc, '（此处附公司管理制度要点汇编，含三级复核制度、保密制度、廉洁从业制度等）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件六：无重大违法记录证明', 2)
P(doc, '（此处附信用中国查询截图、国家企业信用信息公示系统查询截图等）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件七：纳税和社保证明', 2)
P(doc, '（此处附近三个月的纳税证明和社保缴纳证明）', fs=11, fc=RGBColor(128,128,128))
doc.add_page_break()

H(doc, '附件八：承诺函', 2)
P(doc, '（此处附保密承诺函、廉洁从业承诺函、合同履约承诺函等）', fs=11, fc=RGBColor(128,128,128))

doc.save(PTH)
print(f'OK: Part 4 complete, bytes={os.path.getsize(PTH)}')
