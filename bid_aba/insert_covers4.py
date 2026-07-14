import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
NPTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

doc = Document(PTH)

chapter_covers = [
    ('一、项目理解与总体思路', 'cover1-intro.drawio.png'),
    ('二、审核依据与政策解读', 'cover2-laws.drawio.png'),
    ('三、审核范围与审核内容', 'cover3-scope.drawio.png'),
    ('四、审核方法与技术路线', 'cover4-methods.drawio.png'),
    ('五、审核程序与进度安排', 'cover4-methods.drawio.png'),
    ('六、重点难点分析与对策', 'cover5-quality.drawio.png'),
    ('七、审计管理制度与质量保证', 'cover5-quality.drawio.png'),
    ('八、公司概况与服务能力', 'cover6-company.drawio.png'),
    ('九、项目团队配备', 'cover6-company.drawio.png'),
    ('十、类似业绩与履约能力', 'cover7-performance.drawio.png'),
    ('十一、服务承诺与保障措施', 'cover7-performance.drawio.png'),
]

# Strategy: add all images at end of document using add_picture (which handles rels properly)
# Then move the paragraph elements before the heading

body = doc.element.body

for heading_text, cover_file in chapter_covers:
    img_path = os.path.join(IMG, cover_file)
    if not os.path.exists(img_path):
        print(f"NOT FOUND: {img_path}")
        continue
    
    # Find heading
    heading_elem = None
    for p in doc.paragraphs:
        if heading_text in p.text:
            heading_elem = p._element
            break
    
    if heading_elem is None:
        print(f"NOT FOUND: {heading_text}")
        continue
    
    # Add image to end of document using proper add_picture
    # This ensures relationship is created correctly
    temp_para = doc.add_paragraph()
    temp_run = temp_para.add_run()
    temp_pic = temp_run.add_picture(img_path, width=Inches(6.5))
    
    # Now move this paragraph before the heading
    img_elem = temp_para._element
    body.remove(img_elem)
    idx = list(body).index(heading_elem)
    body.insert(idx, img_elem)
    
    # Add blank paragraph after
    blank = doc.add_paragraph()
    blank_elem = blank._element
    body.remove(blank_elem)
    idx2 = list(body).index(heading_elem)
    body.insert(idx2, blank_elem)
    
    print(f"✅ {cover_file} before '{heading_text[:20]}...'")

try:
    doc.save(NPTH)
    print(f'Done! Size: {os.path.getsize(NPTH)}')
except Exception as e:
    print(f'Error: {e}')
