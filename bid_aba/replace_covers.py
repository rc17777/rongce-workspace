import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

doc = Document(PTH)
body = doc.element.body

# Replace the shared covers with chapter-specific ones
chapter_covers = [
    ('六、重点难点分析与对策', 'cover5-ch6.drawio.png'),
    ('七、审计管理制度与质量保证', 'cover5-ch7.drawio.png'),
    ('八、公司概况与服务能力', 'cover6-ch8.drawio.png'),
    ('九、项目团队配备', 'cover6-ch9.drawio.png'),
]

for heading_text, cover_file in chapter_covers:
    img_path = os.path.join(IMG, cover_file)
    if not os.path.exists(img_path):
        print(f"NOT FOUND: {img_path}")
        continue
    
    # Find heading element
    heading_elem = None
    for p in doc.paragraphs:
        if heading_text in p.text:
            heading_elem = p._element
            break
    
    if heading_elem is None:
        print(f"NOT FOUND: {heading_text}")
        continue
    
    # Find the image paragraph before this heading (the old shared cover)
    idx = list(body).index(heading_elem)
    # The image should be at idx-2 (image + blank paragraph before heading)
    if idx >= 2:
        prev = body[idx-2]
        tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
        if tag == 'p':
            # Check if it has a drawing element (image)
            drawing_elem = prev.find('.//' + qn('w:drawing'))
            if drawing_elem is not None:
                # Remove old image paragraph
                body.remove(prev)
                # Also remove blank
                if idx > 0 and body[idx-1] is heading_elem:
                    pass  # blank is at idx-1 now
            
            # Re-find heading index after removal
            new_idx = list(body).index(heading_elem)
            
            # Add new cover image at the correct position
            temp_para = doc.add_paragraph()
            temp_run = temp_para.add_run()
            temp_pic = temp_run.add_picture(img_path, width=Inches(6.5))
            
            img_elem = temp_para._element
            body.remove(img_elem)
            body.insert(new_idx, img_elem)
            
            # Add blank after image
            blank = doc.add_paragraph()
            blank_elem = blank._element
            body.remove(blank_elem)
            body.insert(new_idx + 1, blank_elem)
            
            print(f"✅ Replaced with {cover_file} before '{heading_text[:20]}...'")
            continue
    
    # Fallback: just insert before heading
    print(f"  Inserting new cover before '{heading_text[:20]}...'")
    # Remove old image if exists (search for image before heading)
    for j in range(1, 4):
        if idx >= j:
            candidate = body[idx-j]
            tag2 = candidate.tag.split('}')[-1] if '}' in candidate.tag else candidate.tag
            if tag2 == 'p':
                dw = candidate.find('.//' + qn('w:drawing'))
                if dw is not None:
                    body.remove(candidate)
                    break
    
    new_idx2 = list(body).index(heading_elem)
    temp_para2 = doc.add_paragraph()
    temp_run2 = temp_para2.add_run()
    temp_run2.add_picture(img_path, width=Inches(6.5))
    img_elem2 = temp_para2._element
    body.remove(img_elem2)
    body.insert(new_idx2, img_elem2)
    
    blank2 = doc.add_paragraph()
    blank_elem2 = blank2._element
    body.remove(blank_elem2)
    body.insert(new_idx2+1, blank_elem2)
    print(f"  ✅ Inserted {cover_file}")

doc.save(PTH)
print(f'Done! Size: {os.path.getsize(PTH)}')
