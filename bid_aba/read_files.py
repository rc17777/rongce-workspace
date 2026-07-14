import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

for fname, label in [(r'C:\Users\scrccpa\Desktop\人员.docx', '人员'), (r'C:\Users\scrccpa\Desktop\业绩.docx', '业绩')]:
    doc = Document(fname)
    print(f'=== {label} ===')
    print(f'Paragraphs: {len(doc.paragraphs)} Tables: {len(doc.tables)}')
    for p in doc.paragraphs:
        if p.text.strip():
            print(f'  P: {p.text[:200]}')
    for ti, t in enumerate(doc.tables):
        print(f'  Table {ti}: {len(t.rows)} rows x {len(t.rows[0].cells) if t.rows else 0} cols')
        for ri, r in enumerate(t.rows):
            cells = [c.text.strip()[:30] for c in r.cells]
            print(f'    Row{ri}: {cells}')
    print()
