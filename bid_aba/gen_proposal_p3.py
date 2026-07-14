import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def I(doc, name, w=5.5, cap=None):
    pth=os.path.join(IMG,name)
    if os.path.exists(pth):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(pth, width=Inches(w))
        if cap:
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cr=cp.add_run(f'图：{cap}'); cr.font.size=Pt(9)
            cr.font.color.rgb=RGBColor(128,128,128); cr.font.name='Microsoft YaHei'
        doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

doc = Document(PTH)

# ======== 第四章 审核方法 ========
print('Ch4...')
H(doc, '四、审核方法与技术路线', 1)

H(doc, '4.1 审核方法体系概述', 2)
P(doc, '竣工财务决算审核是一项综合性很强的专业工作，需要综合运用多种审核方法才能确保审核质量。我单位根据多年实践经验和行业最佳实践，建立了完整的审核方法体系。该体系以风险导向审计方法为统领，以数据分析方法为支撑，以现场核实方法为保障，以证据评价方法为基准，确保审核工作的科学性、系统性和严谨性。')
P(doc, '审核方法的选择应当遵循以下原则：一是目标导向原则，即根据审核目标选择最适宜的审核方法；二是成本效益原则，即在保证审核质量的前提下选择效率较高的审核方法；三是风险适配原则，即审核方法的严格程度应当与风险评估结果相匹配，高风险领域采用更严格的审核方法。')

H(doc, '4.2 八种专业审核方法详解', 2)

T(doc, ['方法名称','方法描述','适用场景','操作要点','所需资料','输出成果'],[
    ['全面核查法','对报审资料进行全面系统性审阅核对','首次审核项目、高风险项目、总投资规模较大项目','逐项检查不遗漏，纸质与电子数据交叉验证','全部会计账簿和凭证、工程合同和结算书、竣工决算报表','审核问题清单、差异分析表'],
    ['数据分析法','运用定量分析技术发现异常和趋势','数据量大需快速定位问题；资金流量分析；造价审核','建立行业标准偏差阈值（20%），识别偏离异常数据','财务电子账套、工程量清单、造价信息','异常数据识别表、趋势分析图'],
    ['现场核实法','实地查看、测量、盘点验证账实相符','工程量核实、设备资产盘点、隐蔽工程审核','随机抽样+重点检查结合，拍照录像留存','竣工图纸、设备清单、验收报告','现场核实记录表、现场照片、测量数据'],
    ['函证询证法','向第三方发函核实特定事项','银行存款核实、债权债务确认、往来款项核对','回函率须达90%以上，未回函执行替代测试程序','往来账清单、银行账户信息','函证汇总表、回函统计分析'],
    ['专家咨询法','征询专业领域专家意见','技术争议、质量认定、特殊专业问题','出具书面专家意见，明确责任边界','技术鉴定报告、专业意见书','专家咨询意见书、技术鉴定报告'],
    ['穿行测试法','选取业务从起点追踪至终点','内部控制评价、资金流程审核、制度执行检查','选取重大业务3-5笔，覆盖主要类型','内控制度文件、业务审批记录、账务凭证','穿行测试记录、内控评价表'],
    ['符合性测试','检查内部控制设计合理性和执行有效性','制度执行情况评价、内控有效性验证','采用询问、观察、检查、重新执行等方法','内控制度、业务流程记录','内控测试记录、缺陷汇总表'],
    ['实质性测试','对账户余额和交易进行检查','成本核算审核、支出合规性审核、重大错报查证','细节测试+分析程序，样本量匹配风险','会计账簿和凭证、合同协议、验收文件','实质性测试底稿、差异明细表'],
], hc='1ABC9C')

H(doc, '4.3 审核方法的组合运用策略', 2)
P(doc, '在实际审核工作中，上述八种方法并非完全割裂使用，而是需要根据项目特点灵活组合，形成审核策略矩阵。以下为我单位针对不同类型审核事项的方法组合策略：')

T(doc, ['审核事项类型','风险评估等级','方法组合策略','审核深度要求'],[
    ['资金使用合规性','高','全面核查法+符合性测试+实质性测试','逐笔核对大额支出，抽查小额支出，测试内控制度'],
    ['工程造价真实性','高','全面核查法+数据分析法+现场核实法+专家咨询法','工程量全面核算，异常数据深入分析，实地核实'],
    ['招投标合规性','中','符合性测试+穿行测试+全面核查法','测试招标程序，穿行测试招标全过程，重点检查围标串标'],
    ['设备采购合规性','中','函证询证法+现场核实法+穿行测试','函证供货商价格，现场盘点设备，测试采购流程'],
    ['内控制度有效性','中','符合性测试+穿行测试+数据分析法','测试关键控制点，穿行测试核心业务流程'],
    ['债权债务清理','低','函证询证法+数据分析法','对主要债权债务方发函，分析账龄和异常变动'],
    ['资产管理完整性','低','现场核实法+全面核查法','盘点资产，核对清单与实物'],
], hc='2E86C1')

I(doc, 'fig6-cost-audit.drawio.png', 5.5, '工程造价审核流程与要点图')

H(doc, '4.4 工程造价审核技术方法', 2)
P(doc, '工程造价审核是竣工财务决算审核的技术核心，需要运用专业的工程造价审核技术方法。我单位目前配备了广联达、斯维尔等主流造价软件，拥有一支经验丰富的工程造价审核团队，能够熟练运用以下专业审核技术：')
P(doc, '（一）定额审核法。以国家和地方颁发的现行建筑工程预算定额、费用定额、材料预算价格等为依据，对工程结算中套用的定额子目和取费标准进行逐项审核。审核时要重点关注：定额套用是否正确，是否存在高套定额的情况；定额换算是否合理，是否存在不合理换算的情况；定额中已经包含的工作内容是否又单独列为工作内容重复计取。')
P(doc, '（二）清单审核法。以工程量清单计价规范为依据，对各工程量清单项目的特征描述是否准确、综合单价是否合理、工程量计算是否准确等进行审核。工程量清单计价是目前建设工程最主要的计价方式，在审核中需要重点关注清单项目的特征描述是否与施工图纸和实际施工一致，综合单价是否包含合同约定的全部工作内容和风险费用。')
P(doc, '（三）指标对比法。通过与同类项目的造价指标进行比较分析，快速判断工程造价的总体合理性。我单位积累了丰富的建设项目造价数据库，能够为各类型建设项目的造价审核提供可靠的对比参考。对于超过同类项目造价指标平均值的项目，需要进行深入分析，找出差异原因。')

H(doc, '4.5 财务审核技术方法', 2)
P(doc, '财务审核技术方法主要应用于建设资金管理、会计核算和竣工财务决算编制等环节的审核：')
P(doc, '（一）账户分析技术。通过对被审核项目银行账户的对账单、余额调节表和资金流水进行分析，核实资金收支的真实性和完整性。重点关注资金流向是否存在异常、大额资金支付是否符合审批程序、是否存在资金外循环等情形。')
P(doc, '（二）凭证检查技术。对会计凭证的规范性、完整性和真实性进行检查。检查凭证的附件是否齐全、审批签字是否完整、业务内容是否合法合规。对于大额现金支付凭证、不合规发票、虚假发票等违反财务管理规定的凭证，要进行详细登记和进一步追查。')
P(doc, '（三）账龄分析技术。对应收应付款项的账龄进行分析，识别长期挂账的债权债务，分析长期挂账的原因和潜在风险。对于三年以上的往来款项，要求被审核单位提供说明，核实是否存在坏账或者应核销未核销的情况。')
P(doc, '（四）趋势分析技术。对建设项目的各项财务指标进行纵向和横向比较分析，识别异常变动。例如，对各年度的建设成本、管理费用、资金到位率等指标进行趋势分析，发现异常变动领域作为重点审核方向。')

doc.save(PTH)
print(f'Ch4 done. size={os.path.getsize(PTH)}')
