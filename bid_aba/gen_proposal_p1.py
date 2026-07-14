import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def I(doc, name, w=5.5, cap=None):
    pth=os.path.join(IMG,name)
    if os.path.exists(pth):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(pth, width=Inches(w))
        if cap:
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cr=cp.add_run(f'图：{cap}'); cr.font.size=Pt(9)
            cr.font.color.rgb=RGBColor(128,128,128); cr.font.name='Microsoft YaHei'
        doc.add_paragraph()

doc = Document()
doc.styles['Normal'].font.name='仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')

for s in ['Normal','Heading 1','Heading 2','Heading 3']:
    if s in doc.styles:
        doc.styles[s].paragraph_format.line_spacing=1.5

# --- COVER ---
for _ in range(6):
    doc.add_paragraph()
cp = doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=cp.add_run('2026年阿坝州财政局州本级\n建设项目竣工财务决算审核项目'); r.font.size=Pt(22); r.bold=True; r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体'); r.font.color.rgb=RGBColor(0,0,0)
doc.add_paragraph()
cp2 = doc.add_paragraph(); cp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=cp2.add_run('应  标  方  案'); r2.font.size=Pt(18); r2.font.name='黑体'; r2._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
for _ in range(4):
    doc.add_paragraph()
cp3 = doc.add_paragraph(); cp3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3=cp3.add_run('四川融策会计师事务所有限公司\n四川融策工程咨询有限公司'); r3.font.size=Pt(14); r3.font.name='仿宋'; r3._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
doc.add_paragraph()
cp4 = doc.add_paragraph(); cp4.alignment=WD_ALIGN_PARAGRAPH.CENTER
r4=cp4.add_run('二〇二六年六月'); r4.font.size=Pt(14)
doc.add_page_break()

# ==== TOC PAGE ====
doc.add_heading('目  录', level=0)
toc_items = [
    ('一、项目理解与总体思路', '1'),
    ('二、审核依据与政策解读', '3'),
    ('三、审核范围与审核内容', '8'),
    ('四、审核方法与技术路线', '15'),
    ('五、审核程序与进度安排', '22'),
    ('六、重点难点分析与对策', '29'),
    ('七、审计管理制度与质量保证', '36'),
    ('八、公司概况与服务能力', '43'),
    ('九、项目团队配备', '52'),
    ('十、类似业绩与履约能力', '58'),
    ('十一、服务承诺与保障措施', '63'),
    ('附件', '66'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{item}  {"·"*20}  {page}')
    r.font.size = Pt(12)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')

doc.add_page_break()
print(f'Cover+TOC done. size={0}')

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

# ======== 第一章 项目理解 ========
print('Ch1...')
H(doc, '一、项目理解与总体思路', 1)

H(doc, '1.1 项目背景认知', 2)
P(doc, '受阿坝州财政局委托，我单位有幸参与2026年度阿坝州州本级建设项目竣工财务决算审核项目的应标工作。竣工财务决算审核是基本建设项目管理链条中的关键环节，也是最终判定建设项目投资效益、资金使用合规性及资产管理规范性的重要手段。通过对建设项目的竣工财务决算进行全面、系统、深入的审核，能够准确反映建设项目的实际投资规模和资金使用情况，及时发现和纠正项目建设管理中存在的问题，有效防范财政资金风险，为财政部门批复项目竣工财务决算提供可靠依据。')
P(doc, '根据党中央、国务院关于全面实施预算绩效管理的决策部署和深化财税体制改革的总体要求，加强政府投资项目管理、规范基本建设财务行为、提高财政资金使用效益已成为各级财政部门的重要职责。竣工财务决算审核作为基本建设财务管理的重要环节，对于确保财政资金安全、提高投资效益、防范廉政风险具有不可替代的重要意义。尤其是在阿坝州这样的民族地区、高原地区，政府投资项目的建设管理具有独特的地理环境特点、社会文化特点和行政管理特点，更需要专业、审慎、细致的财务决算审核工作来保障财政资金的安全与效益。')

H(doc, '1.2 阿坝州基本建设管理特征分析', 2)
P(doc, '阿坝藏族羌族自治州地处青藏高原东南缘，属于四川省西北部，是全国十大藏区的重要组成部分，也是长江、黄河上游重要的生态屏障和水源涵养地。全州辖马尔康市和理县、茂县、松潘县、九寨沟县、若尔盖县、红原县、阿坝县、壤塘县、金川县、小金县、黑水县、汶川县共12个县，国土面积8.42万平方公里，约占四川省总面积的17.3%，总人口约92万人。')
P(doc, '阿坝州基本建设管理具有以下显著特征：')
P(doc, '第一，高原施工条件特殊。阿坝州平均海拔3000米以上，其中超过4000米的山峰近200座，属于典型的高原气候区，年均气温仅7.8摄氏度，昼夜温差大，日照充足但紫外线辐射强，每年11月至次年4月为霜冻期，冬季漫长寒冷。大部分地区属于季节性冻土层，对工程建设有着显著影响，施工作业窗口期相对较短，通常为每年5月至10月。这种特殊的施工条件直接影响了建设项目的施工组织设计、工期安排、人员配置和成本构成，在竣工财务决算审核中需特别关注高原地区施工特殊费用的计算与实际执行情况。')
P(doc, '第二，资金来源结构多元。阿坝州建设项目资金来源主要包括中央预算内投资、省级财政配套资金、地方自筹资金以及对口援建资金四大部分。资金来源渠道多元，管理要求各异，资金拨付使用的规范性和合规性是竣工财务决算审核的重点关注领域。特别是中央财政对民族地区的转移支付资金和浙江省对口支援藏区资金的使用管理，因其资金来源的特殊性，需要更加严格的审核。')
P(doc, '第三，项目类型覆盖面广。州本级建设项目涵盖了市政基础设施（道路、桥梁、供水、供热等）、公共服务设施（学校、医院、文化场馆等）、生态保护工程（湿地保护、水土保持、植被恢复等）、交通水利工程（公路、灌溉、防洪等）、信息化工程（电子政务、数字乡村等）以及灾后恢复重建项目多个领域，各类型项目的建设管理要求和财务核算特点差异较大，对审核人员的专业复合能力提出了较高要求。')
P(doc, '第四，建设单位管理水平参差不齐。阿坝州各建设单位的专业管理能力存在明显差异，部分单位的财务管理规范性和工程资料完整度较好，但也有部分单位在档案资料的规范性和完整性方面存在不足，特别是村镇级实施的项目，在财务核算的规范性上可能需要更多关注。这种差异要求在审核方案设计上必须采取"因项目制宜"的差异化策略，对财务管理薄弱环节的重点关注和加强审核。')

T(doc, ['维度','阿坝州特征','对审核工作的影响','应对策略'],[
    ['地理环境','平均海拔3000米，高原气候，施工窗口短（5-10月）','施工特殊费用审核需重点关注','比照定额核定高原施工增加费，核实冬季施工措施费'],
    ['资金来源','中央+省级+地方+对口援建，多来源结构','需要逐项核验各渠道资金使用是否符合各自管理规定','建立资金来源分类审核台账，逐项核对拨款文件和资金使用规定'],
    ['项目类型','市政、公共服务、生态、交通水利、信息化、灾后重建','不同类型项目审核重点差异大','编制分类型审核指南，针对性设置审核检查清单'],
    ['管理水平','建设单位水平参差，部分单位档案不完整','工作量可能超出预期，灵活调配资源','预留充足的现场审核时间，制定资料缺失替代程序'],
    ['气候因素','冬季漫长，且高原反应可能影响工作效率','合理安排现场审核行程，避开极端天气','制定高原工作安全预案，购买高原出行保险'],
], hc='2E86C1')

H(doc, '1.3 总体审核思路', 2)
P(doc, '基于对项目背景的深入理解和阿坝州基本建设管理特征的全面分析，我单位提出以下总体审核思路：')
P(doc, '第一，坚持依法审核、依规办事。以财政部令第81号《基本建设财务规则》、《基本建设项目竣工财务决算管理暂行办法》（财建〔2016〕503号）和《基本建设项目建设成本管理规定》（财建〔2016〕504号）为核心依据，严格执行各项法规政策的规定，确保审核工作有法可依、有章可循。')
P(doc, '第二，坚持风险导向、突出重点。在全面了解项目基本情况的基础上，通过对项目建设管理各环节的风险评估，动态确定审核重点领域和重点事项，将有限的审核资源配置到高风险领域。依托阿坝州建设项目的特点，将高原施工特殊费用、多来源资金使用、工程变更签证、材料设备采购、建设单位管理费等作为重点审核事项。')
P(doc, '第三，坚持数据支撑、方法科学。充分运用数据分析、穿行测试、实质性测试等专业方法，以事实和数据为支撑开展审核工作。对工程造价审核运用造价信息比对法、市场询价验证法和历史数据对比法，对财务审核运用大额支出逐笔核查、异常数据分析和相关指标横向比对等专业技术手段。')
P(doc, '第四，坚持协同高效、全程跟踪。发挥我单位"财务审计+工程咨询"双轮驱动的专业优势，实现造价审核和财务审计协同联动，避免两阶段审核脱节。同时建立与阿坝州财政局、项目单位和相关各方的有效沟通机制，确保审核信息及时传递、问题及时处理、成果及时交付。')

T(doc, ['审核维度','核心目标','关键措施','预期成果'],[
    ['财务合规性','确保建设资金使用合法合规','会计凭证核查、资金流向追踪、重点支出分析','资金使用合规性评价报告'],
    ['工程造价真实性','确保工程造价的准确完整','工程量复核、单价审查、变更签证核实、材料价格比对','工程造价审核明细表'],
    ['内控有效性','评价建设管理内控制度执行','符合性测试、穿行测试、内部控制评价','内部控制评价报告'],
    ['资产管理完整性','确保资产账实相符、移交规范','实物盘查、资产核查、移交手续检查','资产交付清单与差异性分析'],
    ['资金结余合规性','确保结余资金处理方式合规','结余核实、应缴财政核对','结余资金认定意见'],
], hc='1ABC9C')

doc.save(PTH)
print(f'Ch1 done. size={os.path.getsize(PTH)}')
