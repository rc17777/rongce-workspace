import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'
if not os.path.exists(PTH):
    print(f'ERROR: {PTH} not found. Run gen_bid_p1.py and gen_bid_p2.py first.')
    sys.exit(1)

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

doc = Document(PTH)

# ========== 第八章 类似业绩与履约能力 ==========
H(doc, '第八章  类似业绩与履约能力', 1)

P(doc, '我单位自成立以来，长期深耕政府审计和工程咨询领域，在竣工财务决算审核、预算执行审计、专项资金审计、经济责任审计等方面积累了丰富的执业经验。以下为近年完成的类似业绩：')

P(doc, '类似业绩一览表', bold=True, fs=11, indent=False)
T(doc, ['序号','项目名称','委托方','合同金额','完成时间'],[
    ['1','XX市财政局2024年度政府投资项目竣工财务决算审核','XX市财政局','XX万元','2024年'],
    ['2','XX县2023-2024年建设项目竣工财务决算审核','XX县财政局','XX万元','2024年'],
    ['3','XX区2024年财政评审项目','XX区财政局','XX万元','2024年'],
    ['4','XX州本级政府采购项目审计服务','XX州财政局','XX万元','2023-2024年'],
    ['5','XX市审计局2024年政府投资项目审计','XX市审计局','XX万元','2024年'],
    ['6','XX县2023年基本建设项目竣工财务决算审核','XX县财政局','XX万元','2023年'],
], hc='2980B9')

P(doc, '我单位承诺以上业绩材料真实有效，可提供合同协议书复印件供核实。')
doc.add_paragraph()
P(doc, '比选申请人名称（公章）：', indent=False)
P(doc, '日期：    年    月    日', indent=False)
doc.add_page_break()

# ========== 第九章 公司基本情况介绍 ==========
H(doc, '第九章  公司基本情况介绍', 1)

H(doc, '9.1 公司概况', 2)
P(doc, '四川融策会计师事务所（普通合伙）（以下简称"融策"或"本所"）是经四川省财政厅批准设立的专业审计服务机构，持有《会计师事务所执业许可证》。本所办公场所位于四川省成都市高新区，拥有现代化办公场地500余平方米，配备了完整的办公自动化系统和专业的审计软件系统。')
P(doc, '本所业务范围涵盖：财务报表审计、经济责任审计、竣工财务决算审计、专项资金审计、预算执行审计、绩效评价、资产清查、内部控制审计与咨询、政府采购审计、招投标代理、工程造价咨询、税务咨询等全方位专业服务。依托高水平的专业团队和丰富的执业经验，本所已累计完成各类审计、审核、咨询项目数千个，客户遍及四川省各级政府部门、企事业单位和社会团体。')

P(doc, '公司基本信息一览表', bold=True, fs=11, indent=False)
T(doc, ['项目','内容'],[
    ['公司名称','四川融策会计师事务所（普通合伙）'],
    ['执业资质','会计师事务所执业许可证'],
    ['成立时间','XXXX年'],
    ['注册地址','四川省成都市高新区'],
    ['办公面积','500+平方米'],
    ['专业团队','注册会计师X名，造价工程师X名，中级以上职称X名'],
    ['主营业务','政府审计、竣工财务决算、绩效评价、资产清查、工程咨询等'],
    ['核心优势','政府审计+工程咨询双资质，跨专业一体化服务能力'],
], hc='1F618D')

H(doc, '9.2 公司发展历程', 2)
P(doc, '本所自成立以来，始终秉持"专业、客观、公正、诚信"的执业理念，深耕政府审计和工程咨询领域，逐步发展成为四川省具有较大规模和影响力的专业服务机构。')
P(doc, 'XXXX年，本所经四川省财政厅批准正式成立，获得会计师事务所执业资格。成立之初便确立了"聚焦政府审计，服务公共财政"的战略定位，专注于为各级政府部门提供高质量的审计服务。')
P(doc, 'XXXX年，本所业务拓展至工程咨询领域，同步开展工程造价咨询、财政评审、招标代理等服务，形成了"审计+咨询"双轮驱动的业务格局。')
P(doc, 'XXXX年，本所通过ISO9001质量管理体系认证，建立了标准化的质量控制体系。同年开始探索AI技术在审计工作中的应用，成为四川省较早开展数字审计探索的事务所之一。')
P(doc, 'XXXX年，本所完成团队升级，引进多名注册会计师、造价工程师和高级会计师，专业技术人员达到XX人。建立了"1+3+5+N"数字化审计平台，应用AI技术提升审计效率和质量。')

H(doc, '9.3 公司核心竞争优势', 2)
P(doc, '本所的核心竞争力植根于多年的专业化深耕和持续创新，主要体现在以下方面：')
P(doc, '（一）专业团队优势。本所拥有一支高素质的专业团队，核心成员具有注册会计师、造价工程师、注册评估师、高级会计师等多项执业资格，专业背景涵盖会计、审计、工程造价、工程管理、法律、信息技术等多个领域。团队中有多人具备科研事业单位、国有企业、行政机关等多元从业背景，能够从多视角理解和把握项目需求。')
P(doc, '（二）执业经验优势。本所长期服务于各级政府部门，在竣工财务决算审核、预算执行审计、专项资金审计、绩效评价等领域积累了丰富的执业经验。对政府投资项目管理流程、财政资金管理规范、政府采购法规等有着深入的理解和准确的把握，能够快速准确地识别和解决各类复杂问题。')
P(doc, '（三）质量控制优势。本所建立了完善的"三级复核"质量控制体系，从项目经理一级复核到质控部二级复核再到技术负责人三级复核，层层把关、环环相扣。同时建立了项目风险评估制度、重大问题报告制度、客户回访制度等配套制度，形成了全流程质量保障体系。')
P(doc, '（四）技术手段优势。本所积极拥抱数字化转型，在四川省同行业中率先探索AI技术在审计工作中的应用。自主研发了"1+3+5+N"数字化审计平台，集成了智能底稿生成、数据分析、异常检测、报告辅助等核心功能，可以大幅提升审核工作的效率和质量。在工程造价审核方面，配备了广联达、斯维尔等主流造价软件，能够满足各类复杂项目的审核需要。')
P(doc, '（五）本地化服务优势。本所立足于四川省，对各市州特别是民族地区的财政管理体制、项目特点、审计要求等有着深入的了解和丰富的服务经验。在阿坝州、甘孜州、凉山州等民族地区均有成功服务案例，熟悉高原地区项目审核的特殊要求和注意事项。')

H(doc, '9.4 公司管理制度体系', 2)
P(doc, '本所建立了完善的内部管理制度体系，涵盖以下核心领域：')
T(doc, ['序号','管理领域','主要内容','执行标准'],[
    ['1','质量控制','三级复核制度、质量考核、执业规范','ISO9001标准'],
    ['2','人事管理','招聘、培训、考核、薪酬、晋升','制度化、规范化'],
    ['3','财务管理','预算管理、成本核算、收支审批','企业会计准则'],
    ['4','保密管理','保密承诺、涉密管控、信息安全','国家保密规定'],
    ['5','档案管理','底稿归档、借阅登记、保管期限','行业规范'],
    ['6','廉洁从业','廉洁承诺、礼品登记、举报机制','职业道德规范'],
    ['7','信息系统','网络安全、数据备份、软件正版化','等级保护要求'],
    ['8','应急管理','突发事件应急预案、业务连续性','ISO22301标准'],
], hc='8E44AD')

P(doc, '上述管理制度体系的建立和有效执行，确保本所各项业务工作规范、有序、高效运行，为高质量完成委托任务提供了可靠保障。')

H(doc, '9.5 公司信息安全管理', 2)
P(doc, '针对政府审计项目中涉密信息多的特点，本所将信息安全作为公司管理的重中之重。技术层面，建立了内网办公系统，涉密数据使用专用加密存储设备，办公网络实行内外网物理隔离，USB接口实行授权管理。管理层面，建立了全员保密责任制度，入职签署保密协议，定期开展保密教育和检查，泄密行为实行一票否决制。法律层面，保密条款写入服务合同和劳动合同，违反保密义务的承担相应法律责任。')

doc.save(PTH)
print(f'OK: Part 3 complete, bytes={os.path.getsize(PTH)}')
