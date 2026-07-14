import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

doc = Document(PTH)
body = doc.element.body

# Replace cover4-methods before ch4 and ch5 with individual covers
chapter_covers = [
    ('四、审核方法与技术路线', 'cover4-ch4.drawio.png'),
    ('五、审核程序与进度安排', 'cover4-ch5.drawio.png'),
]

for heading_text, cover_file in chapter_covers:
    img_path = os.path.join(IMG, cover_file)
    
    # Find heading
    heading_elem = None
    for p in doc.paragraphs:
        if heading_text in p.text:
            heading_elem = p._element
            break
    if heading_elem is None:
        print(f"NOT FOUND: {heading_text}")
        continue
    
    # Find the shared cover (cover4-methods) before this heading
    idx = list(body).index(heading_elem)
    for j in range(1, 5):
        if idx >= j:
            candidate = body[idx - j]
            # Check if this paragraph contains an image
            drawings = candidate.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if not drawings:
                drawings = candidate.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            if not drawings:
                drawings = candidate.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            if drawings:
                body.remove(candidate)
                break
    
    # Insert new cover
    new_idx = list(body).index(heading_elem)
    temp_para = doc.add_paragraph()
    temp_run = temp_para.add_run()
    temp_run.add_picture(img_path, width=Inches(6.5))
    img_elem = temp_para._element
    body.remove(img_elem)
    body.insert(new_idx, img_elem)
    
    # Add spacing
    blank = doc.add_paragraph()
    blank_elem = blank._element
    body.remove(blank_elem)
    body.insert(new_idx + 1, blank_elem)
    
    print(f"✅ {cover_file} before '{heading_text}'")

doc.save(PTH)
print(f'Done! Size: {os.path.getsize(PTH)}')
