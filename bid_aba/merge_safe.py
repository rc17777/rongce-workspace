import sys, os, copy, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
ORIG = r'D:\openclaw-workspace\bid_aba\work_base.docx'

# 读取原始高端版的所有段落
orig_doc = Document(ORIG)
orig_texts = []
for p in orig_doc.paragraphs:
    t = p.text.strip()
    if len(t) > 30:
        orig_texts.append(t)

print(f"Read {len(orig_texts)} original paragraphs")

# 按关键词分类
sections = {}

def add_to_section(name, texts):
    if name not in sections:
        sections[name] = []
    sections[name].extend(texts)

# Classify each paragraph
for txt in orig_texts:
    classified = False
    for section_name, keywords in [
        ('阿坝州区域发展与项目特征', ['阿坝藏族羌族', '阿坝州作为', '阿坝州总人口', '阿坝州地貌', '阿坝州地广人稀', '阿坝州的城镇化', '阿坝州经济', '阿坝州的脱贫攻坚', '阿坝州民族', '阿坝州是四川省', '阿坝州的气候']),
        ('高原施工专题', ['高原施工', '高原增加费', '冬季施工', '霜冻', '高海拔', '施工作业窗口', '季节性冻土']),
        ('法律法规专题', ['法律法规', '法规体系', '法律依据', '政策文件', '国务院', '财政部令', '川发改', '川财投']),
        ('基本建设财务规则体系', ['基本建设财务规则', '财政部令第81号', '财务管理规则', '建设成本管理', '竣工财务决算管理暂行办法', '503号', '504号']),
        ('政府投资项目概算管理', ['概算管理', '概算调整', '概算执行', '初步设计概算', '概算审批', '超概算']),
        ('工程造价管理与控制', ['工程造价', '工程量清单', '计价', '定额', '招标控制价', '投标报价', '造价管理', '成本控制', '量价分离']),
        ('工程价款支付与结算', ['工程价款', '工程款支付', '质保金', '缺陷责任', '价款结算', '竣工结算', '结算审核']),
        ('建设资金管理与审计', ['资金管理', '建设资金', '现金流', '会计核算', '账务处理', '资金拨付', '专项资金', '财务核算']),
        ('竣工财务决算编制与审核', ['竣工财务决算', '决算报表', '决算编制', '决算审核', '审核报告', '工作底稿', '审核证据', '决算报告', '待摊投资', '转出投资', '基建收入']),
        ('资产移交与档案管理', ['资产移交', '档案管理', '档案专项', '资产交付', '转固定资产', '资产管理']),
        ('物资采购管理', ['物资采购', '设备采购', '材料采购', '招投标', '比选', '采购管理']),
        ('合同管理', ['合同管理', '合同签订', '合同履行', '合同变更', '合同纠纷', '工程索赔', '反索赔']),
        ('质量安全管理', ['质量管理', '安全管理', '安全生产', '质量监督', '工程质量', '施工安全', '工程监理']),
        ('竣工验收与交付', ['竣工验收', '验收报告', '验收标准', '工程验收', '分项验收', '综合验收']),
        ('专项审核领域', ['PPP', '对口援建', '援建资金', '生态保护', '民族地区', '扶贫项目']),
        ('项目监督与绩效评价', ['绩效评价', '后评价', '监督', '投诉', '信息公开', '绩效管理', '预算绩效']),
        ('建设管理与内控', ['内部控制', '内控制度', '审批程序', '岗位责任', '不相容', '建设管理', '项目管理']),
        ('项目管理创新', ['创新', '信息化', '数字化', 'BIM', '智慧工地']),
        ('数据与信息化', ['数据分析', '数据库', '造价指标', '信息化管理', '数据采集']),
    ]:
        if any(kw in txt for kw in keywords):
            add_to_section(section_name, [txt])
            classified = True
            break
    
    if not classified:
        add_to_section('建设管理与内控', [txt])

# Deduplicate
for k in sections:
    seen = set()
    uniq = []
    for t in sections[k]:
        if t[:80] not in seen:
            seen.add(t[:80])
            uniq.append(t)
    sections[k] = uniq

print(f"Classified into {len(sections)} sections:")
for k, v in sections.items():
    print(f"  {k}: {len(v)} paragraphs")

# 现在重建整个文档
doc = Document(PTH)
body = doc.element.body

# 查找所有"附"相关段落并确定位置
# Find the "附" heading and mark elements to remove
found = []
for i, el in enumerate(body):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    if tag == 'p':
        pPr = el.find(qn('w:pPr'))
        if pPr is not None:
            ps = pPr.find(qn('w:pStyle'))
            if ps is not None:
                val = ps.get(qn('w:val'))
                if val == 'Heading1':
                    t = el.find(qn('w:r'))
                    if t is not None:
                        tt = t.find(qn('w:t'))
                        if tt is not None and tt.text and '附：项目建设管理专题分析' in tt.text:
                            # Found appendix heading - mark all paragraphs from here to next H1
                            found.append(('appendix_start', i, el))
                            for j in range(i+1, len(list(body))):
                                el2 = list(body)[j]
                                tag2 = el2.tag.split('}')[-1] if '}' in el2.tag else el2.tag
                                if tag2 == 'p':
                                    pPr2 = el2.find(qn('w:pPr'))
                                    if pPr2 is not None:
                                        ps2 = pPr2.find(qn('w:pStyle'))
                                        if ps2 is not None and ps2.get(qn('w:val')) == 'Heading1':
                                            break
                                found.append(('appendix_content', j, el2))
                            break

print(f"Found appendix: {len(found)} elements")

if found:
    appendix_start_idx = found[0][1]
    # Extract the appendix section elements
    to_remove = [f[2] for f in found]
    
    # Now for each classified section, find the matching chapter and insert content
    # Map sections to chapter insertion points
    chapter_map = {
        '阿坝州区域发展与项目特征': '一、项目理解与总体思路',
        '基本建设财务规则体系': '三、审核范围与审核内容',
        '政府投资项目概算管理': '三、审核范围与审核内容',
        '工程造价管理与控制': '三、审核范围与审核内容',
        '工程价款支付与结算': '三、审核范围与审核内容',
        '建设资金管理与审计': '四、审核方法与技术路线',
        '竣工财务决算编制与审核': '三、审核范围与审核内容',
        '资产移交与档案管理': '三、审核范围与审核内容',
        '专项审核领域': '四、审核方法与技术路线',
        '项目监督与绩效评价': '四、审核方法与技术路线',
        '建设管理与内控': '七、审计管理制度与质量保证',
        '合同管理': '三、审核范围与审核内容',
        '物资采购管理': '四、审核方法与技术路线',
        '高原施工专题': '六、重点难点分析与对策',
        '竣工验收与交付': '三、审核范围与审核内容',
        '法律法规专题': '二、审核依据与政策解读',
        '项目管理创新': '八、公司概况与服务能力',
        '数据与信息化': '五、审核程序与进度安排',
        '质量安全管理': '七、审计管理制度与质量保证',
    }
    
    # Find each chapter heading element
    chapter_els = {}
    for i, el in enumerate(body):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                ps = pPr.find(qn('w:pStyle'))
                if ps is not None:
                    val = ps.get(qn('w:val'))
                    if val == 'Heading1':
                        t = el.find(qn('w:r'))
                        if t is not None:
                            tt = t.find(qn('w:t'))
                            if tt is not None and tt.text:
                                chapter_els[tt.text] = el
    
    print(f"Found {len(chapter_els)} chapter headings")
    
    # Create sub-headings and content paragraphs for each section
    # Use python-docx to add paragraphs at end, then move via XML insertion
    
    for section_name, paras in sections.items():
        chapter_name = chapter_map.get(section_name)
        if not chapter_name or chapter_name not in chapter_els:
            print(f"  SKIP: {section_name} -> {chapter_name}")
            continue
        
        # Find where to insert: before the NEXT chapter heading after target
        insert_before = None
        found_ch = False
        for i, el in enumerate(body):
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p':
                pPr = el.find(qn('w:pPr'))
                if pPr is not None:
                    ps = pPr.find(qn('w:pStyle'))
                    if ps is not None:
                        val = ps.get(qn('w:val'))
                        if val == 'Heading1':
                            t = el.find(qn('w:r'))
                            if t is not None:
                                tt = t.find(qn('w:t'))
                                if tt is not None and tt.text:
                                    if found_ch:
                                        insert_before = el
                                        break
                                    if tt.text == chapter_name:
                                        found_ch = True
        
        if insert_before is None:
            print(f"  SKIP: no insert point for {section_name} after {chapter_name}")
            continue
        
        # Create section sub-heading using python-docx (add to end, then move)
        H2_elem = doc.add_paragraph()
        H2_run = H2_elem.add_run(f'{section_name}')
        H2_run.bold = True
        H2_run.font.size = Pt(14)
        H2_run.font.name = '黑体'
        H2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # Move H2 before insert_before
        h2_xml = H2_elem._element
        body.remove(h2_xml)
        body.insert(list(body).index(insert_before), h2_xml)
        
        # Add content paragraphs
        for para_text in paras:
            p_elem = doc.add_paragraph()
            p_run = p_elem.add_run(para_text)
            p_run.font.size = Pt(11)
            p_run.font.name = '仿宋'
            p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            p_elem.paragraph_format.first_line_indent = Cm(0.74)
            p_elem.paragraph_format.space_after = Pt(6)
            
            p_xml = p_elem._element
            body.remove(p_xml)
            # Insert before insert_before
            ins_idx = list(body).index(insert_before)
            body.insert(ins_idx, p_xml)
        
        print(f"  ✅ {section_name} ({len(paras)} paras) -> {chapter_name}")
    
    # Remove appendix
    for el in to_remove:
        try:
            body.remove(el)
        except:
            pass
    
    print(f"✅ Appendix removed")
else:
    print("⚠ No appendix found in document")

doc.save(PTH)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'Saved! Size: {os.path.getsize(PTH)} chars: {total} words: {int(total/2)}')
