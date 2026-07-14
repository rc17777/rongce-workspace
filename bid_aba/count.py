import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
doc = Document(r"D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx")
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f"Paras:{len(doc.paragraphs)} Tables:{len(doc.tables)} Chars:{total} Words:{int(total/2)}")
