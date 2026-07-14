import sys, os
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx import Document

src = r'D:\openclaw-workspace\bid_aba\work_base.docx'
base = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'

# Read original paragraphs
src_doc = Document(src)
paras = [p for p in src_doc.paragraphs if p.text.strip()]

# Open base doc and append content
doc = Document(base)

# Find what section we can expand based on original content
# Original has no headings, just continuous text. Let's add structured headings
# and insert big chunks under appropriate sections.

# Add heading for original content organization
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

H(doc, '7.17 项目建设管理专题分析', 3)
P(doc, '以下为基于我单位多年项目经验，对阿坝州建设项目建设管理相关专题的系统分析。')

# Group paragraphs by content themes
sections = {
    '阿坝州区域发展概况与项目特征': [],
    '基本建设财务管理法规体系': [],
    '政府投资项目概算管理': [],
    '工程造价管理与控制': [],
    '工程价款支付与结算': [],
    '建设资金管理与审计': [],
    '竣工财务决算编制与审核': [],
    '资产移交与档案管理': [],
    '专项审核领域（PPP/对口援建/生态）': [],
    '项目监督与绩效评价': [],
}

current_section = None
for p in paras:
    txt = p.text.strip()
    if not txt:
        continue
    # Simple content-based classification
    if any(kw in txt[:100] for kw in ['阿坝藏族羌族', '阿坝州地貌', '阿坝州作为', '阿坝州总人口', '阿坝州地广人稀', '阿坝州的城镇化', '阿坝州的脱贫攻坚']):
        current_section = '阿坝州区域发展概况与项目特征'
    elif any(kw in txt[:80] for kw in ['基本建设财务规则', '基本建设项目竣工财务决算管理暂行办法', '基本建设项目建设成本管理规定', '四川省基本建设财务管理规定', '政府投资项目条例', 'PPP项目']):
        current_section = '基本建设财务管理法规体系'
    elif any(kw in txt[:80] for kw in ['概算管理', '概算调整', '概算执行']):
        current_section = '政府投资项目概算管理'
    elif any(kw in txt[:80] for kw in ['工程造价', '量价分离', '定额', '计价', '工程量清单', '招标控制价']):
        current_section = '工程造价管理与控制'
    elif any(kw in txt[:80] for kw in ['工程价款', '工程款支付', '质保金', '缺陷责任']):
        current_section = '工程价款支付与结算'
    elif any(kw in txt[:80] for kw in ['资金管理', '建设资金', '现金流', '会计核算', '账务处理']):
        current_section = '建设资金管理与审计'
    elif any(kw in txt[:80] for kw in ['竣工财务决算', '决算报表', '决算编制', '决算审核', '审核报告', '工作底稿', '审核证据']):
        current_section = '竣工财务决算编制与审核'
    elif any(kw in txt[:80] for kw in ['资产移交', '档案管理', '档案专项', '竣工验收']):
        current_section = '资产移交与档案管理'
    elif any(kw in txt[:80] for kw in ['对口援建', 'PPP', '生态保护', '环保', '民族传统']):
        current_section = '专项审核领域（PPP/对口援建/生态）'
    elif any(kw in txt[:80] for kw in ['绩效评价', '后评价', '监督', '投诉', '信息公开']):
        current_section = '项目监督与绩效评价'
    elif any(kw in txt[:100] for kw in ['工程索赔', '反索赔', '工程保险', '安全生产', '工程监理', '合同管理', '合同纠纷', '竣工财务决算审核与']):
        current_section = '专项审核领域（PPP/对口援建/生态）'
    
    if current_section and current_section in sections:
        sections[current_section].append(txt)

# Add each section with heading
for section_name, paras_list in sections.items():
    if not paras_list:
        continue
    H(doc, section_name, 3)
    
    # If long paragraph, try to make it into multiple paragraphs
    for p_text in paras_list:
        # Split long blocks
        if len(p_text) > 500:
            # Split into logical segments
            segs = []
            for sep in ['。', '；', '！', '？']:
                parts = p_text.split(sep)
                if len(parts) > 3:
                    cur = ''
                    for part in parts:
                        if len(cur) + len(part) < 300:
                            cur += part + sep
                        else:
                            segs.append(cur)
                            cur = part + sep
                    if cur:
                        segs.append(cur)
                    break
            else:
                segs = [p_text]
            
            for seg in segs:
                if seg.strip():
                    P(doc, seg.strip(), fs=11)
        else:
            P(doc, p_text, fs=11)

doc.save(base)
print(f'OK: Merged original content. Size={os.path.getsize(base)} bytes')

# Quick char count
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'Total chars: {total}, approx words: {int(total/2)}')
