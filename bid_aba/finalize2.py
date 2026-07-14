import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

doc = Document(PTH)

# === 大量补充内容 ===
H(doc, '7.18 阿坝州建设项目建设管理特征分析', 3)
P(doc, '阿坝藏族羌族自治州地处青藏高原东南缘，是全国十大藏区的重要组成部分，也是长江、黄河上游重要的生态屏障。全州辖马尔康市和12个县，国土面积8.42万平方公里，总人口约92万人。自国家实施西部大开发战略以来，特别是"十四五"规划实施以来，阿坝州基础设施建设进入了历史最快发展期，州本级建设项目涵盖了市政基础设施、公共服务设施、生态保护工程、交通水利工程以及信息化工程等众多领域。')
P(doc, '从资金来源看，阿坝州建设项目资金主要来源于中央预算内投资、省级财政配套、地方自筹资金以及对口援建资金，资金来源结构多元，多头管理特征明显。中央预算内投资主要用于重大基础设施建设，省级配套资金侧重民生领域项目，地方自筹资金则体现地方发展需求，对口援建资金特别是浙江省对口支援藏区资金在阿坝州发挥了重要的补充作用。这种多元化的资金来源结构给竣工财务决算审核带来了一定难度，需要审核人员对各渠道资金的使用管理要求都有深入的了解。')
P(doc, '从项目管理角度看，阿坝州建设项目管理具有以下特点：一是各行业行政主管部门各负其责，管理标准和程序存在差异；二是各建设单位管理水平参差不齐，档案资料的完整性和规范性差异较大；三是对口援建项目既有地方特色又有援建方要求，需特别关注资金使用的合规性问题；四是受高原气候影响，施工窗口期短暂，工程进度管理和资金拨付节奏与内地不同。')

# 扩展工程审核专题
H(doc, '7.19 工程量审核专题', 3)
P(doc, '工程量的准确核实是竣工财务决算审核的基础环节。竣工财务决算审核中涉及的工程量主要包括土石方工程量、混凝土及钢筋混凝土工程量、装饰装修工程量、安装工程量和措施项目工程量等。不同类型工程量的核实方法各有特点，需要审核人员具备相应的专业技术知识。')
P(doc, '土石方工程量的核实是工程量审核中最常见也最容易出现争议的环节。土石方工程量包括场地平整、基坑开挖、沟槽开挖、路基填筑等，核实的主要依据包括原始地面标高测量记录、竣工测量数据、设计图纸断面图等。在核实过程中，需要关注以下几个要点：原始地面标高的真实性，是否存在人为调整原始标高以增加工程量的情况；回填土方的来源，是否利用了开挖土方进行回填，还是从场外另行购土回填；土石方的分类是否准确，各类土石方的单价差异较大，分类错误会导致工程造价的虚增或虚减。')

P(doc, '混凝土及钢筋混凝土工程量的核实需要重点关注构件的几何尺寸是否与图纸一致、配筋情况是否符合设计要求、混凝土标号是否与图纸一致、现浇与预制构件是否区分清楚等。核实的主要方法包括：对主要构件进行实际测量，与设计图纸进行比对；对隐蔽工程的钢筋用量进行抽查核验，比对设计配筋表和实际配筋记录；对混凝土浇筑记录与工程量申报数据进行比对分析。')

P(doc, '装饰装修工程量的核实需要重点关注各类饰面材料的实际施工面积是否与申报一致、不同材质交界处的计算规则是否正确、吊顶工程的龙骨和面层是否区分计算、地面工程的找平层和面层是否分别计算等。核实的主要方法包括：对主要空间进行实测实量，与竣工图纸进行比对；对异形造型的展开面积进行专项核算。')

P(doc, '安装工程量的核实涉及给排水、电气、暖通、消防等多个专业，需要重点关注设备型号和规格是否与图纸一致、管线走向和长度是否符合设计要求、设备安装数量是否与图纸和竣工验收清单相符、系统的联动调试是否完成等。核实的主要方法是现场清点设备、核对图纸和验收文件，对关键设备进行开箱检验。')

# 扩展：材料价格审核
H(doc, '7.20 材料价格审核专题', 3)
P(doc, '材料价格在工程造价中占比最高，通常在总造价的60%-70%左右，因此材料价格的审核是工程造价审核的重点。材料价格审核的主要方法包括：')
P(doc, '第一，造价信息比对法。以工程造价管理部门发布的同期造价信息为基准，对申报的材料价格进行比对分析，超出造价信息价格一定幅度的（通常为10%-15%），要求提供采购合同和发票进行核实。对于造价信息中未包含的新型材料或特殊材料，参考同类项目采购价格或向生产厂家询价确定合理价格。')
P(doc, '第二，采购凭证核查法。对大宗材料和特殊材料的采购合同和发票进行逐项核查，核实采购时间、采购数量、采购单价、供货单位等信息，判断采购价格的合理性。对存在疑问的采购价格，向供货单位进行函证核实。')
P(doc, '第三，市场询价验证法。对于价格信息不透明或价格波动较大的材料，通过电话咨询、网上查询、实地走访等方式进行市场询价，获取2-3家供应商的报价作为参考，综合确定材料的合理市场价。')
P(doc, '第四，历史数据对比法。将申报的材料价格与我单位历史项目数据库中的同类材料价格进行横向对比，发现异常价格数据，作为重点审核线索。对于明显偏离市场价格的材料价格，进一步追查原因。')

# 总结性内容
H(doc, '7.21 综合服务方案总结', 3)
P(doc, '综上所述，我单位为服务阿坝州财政局州本级建设项目竣工财务决算审核项目，制定了全面、细致、可操作的实施方案。方案遵循法规依据齐全、审核方法科学、流程节点清晰、质量标准严格、保密措施到位、廉洁要求明确的核心原则，覆盖了从项目准备到成果交付的全过程。')
P(doc, '在项目实施过程中，我单位将充分发挥在政府审计和工程咨询领域积累的专业优势和经验优势，灵活运用全面核查法、数据分析法、现场核实法、函证询证法、专家咨询法、穿行测试法、符合性测试法和实质性测试法等专业方法，确保审核质量满足委托方的要求和期待。')
P(doc, '我单位郑重承诺，将严格按照比选文件和合同约定，以最高的专业标准、最严的质量要求、最好的服务态度，高质量完成阿坝州财政局委托的竣工财务决算审核任务，为阿坝州财政资金的安全高效使用提供有力保障。')

doc.save(PTH)
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'FINAL: tables={len(doc.tables)} paras={len(doc.paragraphs)} chars={total} words={int(total/2)} size={os.path.getsize(PTH)}')
