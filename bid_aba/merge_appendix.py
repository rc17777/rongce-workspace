import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'

doc = Document(PTH)
body = doc.element.body

# 找到所有段落，按heading归类
class ParaGroup:
    def __init__(self, heading_elem, heading_text):
        self.heading = heading_elem
        self.text = heading_text
        self.children = []

# 扫描所有段落元素
all_paras = list(body)
groups = []
current_group = None
in_appendix = False

for el in all_paras:
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    if tag != 'p':
        groups.append(ParaGroup(None, '[non-para]'))
        groups[-1].children.append(el)
        continue
    
    # Check if heading
    pPr = el.find(qn('w:pPr'))
    is_heading = False
    heading_text = ''
    if pPr is not None:
        pStyle = pPr.find(qn('w:style'))
        if pStyle is not None:
            val = pStyle.get(qn('w:val'))
            if val and 'Heading' in val:
                is_heading = True
                r_elem = el.find(qn('w:r'))
                if r_elem is not None:
                    t_elem = r_elem.find(qn('w:t'))
                    if t_elem is not None and t_elem.text:
                        heading_text = t_elem.text
    
    if is_heading:
        groups.append(ParaGroup(el, heading_text))
        current_group = groups[-1]
    else:
        if current_group is not None:
            current_group.children.append(el)
        else:
            groups.append(ParaGroup(None, ''))
            groups[-1].children.append(el)
            current_group = groups[-1]

print(f"Found {len(groups)} groups")

# 定义专题到章节的映射
topic_to_chapter = {
    '阿坝州区域发展与项目特征': '一、项目理解与总体思路',
    '基本建设财务规则体系': '三、审核范围与审核内容',
    '政府投资项目概算管理': '三、审核范围与审核内容',
    '工程造价管理与控制': '三、审核范围与审核内容',
    '工程价款支付与结算': '三、审核范围与审核内容',
    '建设资金管理与审计': '四、审核方法与技术路线',
    '竣工财务决算编制与审核': '三、审核范围与审核内容',
    '资产移交与档案管理': '三、审核范围与审核内容',
    '专项审核领域': '四、审核方法与技术路线',
    '项目监督与绩效评价': '四、审核方法与技术路线',
    '建设管理与内控': '七、审计管理制度与质量保证',
    '合同管理': '三、审核范围与审核内容',
    '物资采购管理': '四、审核方法与技术路线',
    '高原施工专题': '六、重点难点分析与对策',
    '竣工验收与交付': '三、审核范围与审核内容',
    '法律法规专题': '二、审核依据与政策解读',
    '项目管理创新': '八、公司概况与服务能力',
    '数据与信息化': '五、审核程序与进度安排',
}

# 找到"附"的位置和每个目标章节的位置
appendix_group = None
appendix_idx = -1
chapter_map = {}  # chapter_text -> (group_idx, group)

for i, g in enumerate(groups):
    if g.heading is not None:
        text = g.text.strip()
        if '附：项目建设管理专题分析' in text:
            appendix_group = g
            appendix_idx = i
        elif text in topic_to_chapter.values() or any(text.startswith(ch) for ch in ['一、','二、','三、','四、','五、','六、','七、','八、','九、','十、','十一、']):
            chapter_map[text] = (i, g)

print(f"Appendix found: {appendix_group is not None} at index {appendix_idx}")
print(f"Chapter groups: {list(chapter_map.keys())[:5]}...")

# Now extract the appendix's children (the 18 topic groups with their content)
# and move them to their respective chapters

# The appendix group contains: the heading + 18 topic headings + their content
# We need to find each topic heading within the appendix children

if appendix_group is not None:
    # Get all elements from appendix heading to end of document
    appendix_start = appendix_idx
    
    # Each topic within the appendix is: heading2 + paragraphs
    # We need to find which topics are attached to which paragraphs
    
    # Find the boundary of appendix: from heading to next heading1 or end
    appendix_end = len(groups)
    for i in range(appendix_start + 1, len(groups)):
        g = groups[i]
        if g.heading is not None and g.heading.get(qn('w:pPr')) is not None:
            pStyle = g.heading.find(qn('w:pPr')).find(qn('w:style'))
            if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                if '附' not in g.text and '报价' not in g.text:
                    appendix_end = i
                    break
    
    # Collect all elements that belong to appendix
    appendix_els = []
    for gi in range(appendix_start, appendix_end):
        g = groups[gi]
        appendix_els.append(g.heading)
        for child in g.children:
            appendix_els.append(child)
    
    # Now find each topic's content
    # Topics are heading2 elements
    topic_start = {}
    current_topic = None
    current_start = 0
    
    for ai, el in enumerate(appendix_els):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag != 'p':
            continue
        pPr = el.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:style'))
            if pStyle is not None:
                val = pStyle.get(qn('w:val'))
                if val == 'Heading2':
                    r_elem = el.find(qn('w:r'))
                    if r_elem is not None:
                        t_elem = r_elem.find(qn('w:t'))
                        if t_elem is not None and t_elem.text:
                            # Found a topic
                            if current_topic:
                                topic_start[current_topic] = (current_start, ai)
                            current_topic = t_elem.text
                            current_start = ai
    
    if current_topic:
        topic_start[current_topic] = (current_start, len(appendix_els))
    
    print(f"\nFound {len(topic_start)} topics:")
    for topic, (s, e) in topic_start.items():
        chapter = topic_to_chapter.get(topic, 'UNMAPPED')
        print(f"  {topic} -> {chapter} ({e-s} elements)")
    
    # Now move elements to chapters
    # Strategy: find each chapter's location, insert topic content after the last child
    moved_count = 0
    for topic, (start, end) in topic_start.items():
        chapter_name = topic_to_chapter.get(topic)
        if not chapter_name or chapter_name not in chapter_map:
            print(f"  SKIP: {topic} -> {chapter_name}")
            continue
        
        ch_idx, ch_group = chapter_map[chapter_name]
        
        # Find where to insert: after the last child of this chapter group
        # Actually, we want to insert after the chapter's last content paragraph
        # The chapter group may span multiple groups. Find the last group before the next chapter.
        
        # Find the last group that belongs to this chapter
        last_idx = ch_idx
        for gi in range(ch_idx + 1, len(groups)):
            if groups[gi].heading is not None and groups[gi].text in chapter_map:
                break
            last_idx = gi
        
        # The elements to insert after are the children of groups[last_idx]
        last_group = groups[last_idx]
        if last_group.children:
            insert_after = last_group.children[-1]
        else:
            insert_after = last_group.heading
        
        parent = insert_after.getparent()
        insert_idx = list(parent).index(insert_after) + 1
        
        # Create topic sub-heading
        topic_heading = copy.deepcopy(appendix_els[start])
        # Add the topic heading
        parent.insert(insert_idx, topic_heading)
        insert_idx += 1
        
        # Add the content paragraphs
        for ei in range(start + 1, end):
            el_to_copy = appendix_els[ei]
            # Check for drawing elements - skip images (they're in the old doc already elsewhere)
            tag = el_to_copy.tag.split('}')[-1] if '}' in el_to_copy.tag else el_to_copy.tag
            if tag == 'p':
                drawing = el_to_copy.find('.//' + qn('w:drawing'))
                if drawing is not None:
                    continue  # skip image paragraphs
            
            copied = copy.deepcopy(el_to_copy)
            body.insert(insert_idx, copied)
            insert_idx += 1
        
        moved_count += 1
        print(f"  MOVED: {topic} ({end-start} el) -> after {chapter_name}")
        
        # Update all group indices since we modified the body
        # (This is fragile but should work for sequential processing)
    
    # Now remove the entire appendix section
    # All appendix elements are still in the body. Remove them.
    for el in appendix_els:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p' or tag == 'tbl' or tag == 'sdt':
            drawing = el.find('.//' + qn('w:drawing'))
            if drawing is not None:
                continue  # keep images
            try:
                body.remove(el)
            except:
                pass
    
    # Remove the appendix heading
    try:
        if appendix_group and appendix_group.heading is not None:
            body.remove(appendix_group.heading)
    except:
        pass
    
    print(f"\nTotal moved: {moved_count} / 18 topics")

doc.save(PTH)
print(f'Saved. Size: {os.path.getsize(PTH)}')

# Count chars
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'Total chars: {total}')
