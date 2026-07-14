import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标方案.docx'
ORIG = r'D:\openclaw-workspace\bid_aba\work_base.docx'

doc = Document(PTH)
orig = Document(ORIG)

def P(doc, text, bold=False, fs=12, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

# Add heading
H(doc, '附：项目建设管理专题分析（补充资料）', 1)
P(doc, '以下内容为我单位根据多年实务经验，对阿坝州建设项目管理的各专题进行的系统分析，作为前述实施方案的详细补充。')

# Extract all paragraph text from original document, organized by content themes
orig_texts = []
for p in orig.paragraphs:
    t = p.text.strip()
    if len(t) > 30:  # Skip very short lines
        orig_texts.append(t)

# Group into sections
sections = {
    '阿坝州区域发展与项目特征': [],
    '基本建设财务规则体系': [],
    '政府投资项目概算管理': [],
    '工程造价管理与控制': [],
    '工程价款支付与结算': [],
    '建设资金管理与审计': [],
    '竣工财务决算编制与审核': [],
    '资产移交与档案管理': [],
    '专项审核领域': [],
    '项目监督与绩效评价': [],
    '建设管理与内控': [],
    '合同管理': [],
    '物资采购管理': [],
    '高原施工专题': [],
    '质量安全管理': [],
    '竣工验收与交付': [],
    '法律法规专题': [],
    '项目管理创新': [],
    '数据与信息化': [],
}

for txt in orig_texts:
    kw_matched = False
    for section_name, keywords in [
        ('阿坝州区域发展与项目特征', ['阿坝藏族羌族', '阿坝州作为', '阿坝州总人口', '阿坝州地貌', '阿坝州地广人稀', '阿坝州的城镇化', '阿坝州经济', '阿坝州的脱贫攻坚', '阿坝州民族', '阿坝州是四川省', '阿坝州的气候']),
        ('高原施工专题', ['高原施工', '高原增加费', '冬季施工', '霜冻', '高海拔', '施工作业窗口', '季节性冻土']),
        ('法律法规专题', ['法律法规', '法规体系', '法律依据', '政策文件', '国务院', '财政部令', '财建', '川发改', '川财投']),
        ('基本建设财务规则体系', ['基本建设财务规则', '财政部令第81号', '财务管理规则', '建设成本管理', '基本建设项目竣工', '竣工财务决算管理暂行办法', '503号', '504号']),
        ('政府投资项目概算管理', ['概算管理', '概算调整', '概算执行', '初步设计概算', '概算审批', '超概算']),
        ('工程造价管理与控制', ['工程造价', '工程量清单', '计价', '定额', '招标控制价', '投标报价', '造价管理', '成本控制', '量价分离']),
        ('工程价款支付与结算', ['工程价款', '工程款支付', '质保金', '缺陷责任', '价款结算', '竣工结算', '结算审核']),
        ('建设资金管理与审计', ['资金管理', '建设资金', '现金流', '会计核算', '账务处理', '资金拨付', '专项资金', '财务核算']),
        ('竣工财务决算编制与审核', ['竣工财务决算', '决算报表', '决算编制', '决算审核', '审核报告', '工作底稿', '审核证据', '决算报告', '待摊投资', '转出投资', '基建收入']),
        ('资产移交与档案管理', ['资产移交', '档案管理', '档案专项', '资产交付', '转固定资产', '资产管理']),
        ('物资采购管理', ['物资采购', '设备采购', '材料采购', '招投标', '比选', '采购管理']),
        ('合同管理', ['合同管理', '合同签订', '合同履行', '合同变更', '合同纠纷', '工程索赔', '反索赔']),
        ('质量安全管理', ['质量管理', '安全管理', '安全生产', '质量监督', '工程质量', '施工安全', '工程监理']),
        ('竣工验收与交付', ['竣工验收', '验收报告', '验收标准', '工程验收', '分项验收', '综合验收']),
        ('专项审核领域', ['PPP', '对口援建', '援建资金', '生态保护', '民族地区', '扶贫项目']),
        ('项目监督与绩效评价', ['绩效评价', '后评价', '监督', '投诉', '信息公开', '绩效管理', '预算绩效']),
        ('建设管理与内控', ['内部控制', '内控制度', '审批程序', '岗位责任', '不相容', '建设管理', '项目管理']),
        ('项目管理创新', ['创新', '信息化', '数字化', 'BIM', '智慧工地']),
        ('数据与信息化', ['数据分析', '数据库', '造价指标', '信息化管理', '数据采集']),
    ]:
        if any(kw in txt for kw in keywords):
            sections[section_name].append(txt)
            kw_matched = True
            break
    
    if not kw_matched:
        # Try to classify by looking at first ~100 chars
        if '工程' in txt[:100] and ('资金' in txt or '建设' in txt):
            sections['建设管理与内控'].append(txt)
        else:
            # Add to catch-all
            if len(txt) > 100:
                sections['建设管理与内控'].append(txt)

# Insert each section
for section_name in sections:
    para_list = sections[section_name]
    if not para_list:
        continue
    
    # Deduplicate
    seen = set()
    unique = []
    for t in para_list:
        if t[:80] not in seen:
            seen.add(t[:80])
            unique.append(t)
    
    if not unique:
        continue
    
    H(doc, section_name, 2)
    
    for txt in unique:
        if len(txt) > 800:
            # Split into reasonable paragraphs
            parts = []
            current = ''
            for sent_sep in ['。', '；', '！', '？']:
                sentences = txt.split(sent_sep)
                if len(sentences) > 1:
                    for sent in sentences:
                        if len(current) + len(sent) < 400:
                            current += sent + sent_sep
                        else:
                            if current.strip():
                                parts.append(current)
                            current = sent + sent_sep
                    if current.strip():
                        parts.append(current)
                    break
            else:
                parts = [txt]
            
            for part in parts:
                if part.strip():
                    P(doc, part.strip(), fs=11)
        else:
            P(doc, txt, fs=11)

n_paras_merged = len(doc.paragraphs)
doc.save(PTH)

# Count total chars
total = 0
for p in doc.paragraphs:
    for r in p.runs:
        total += len(r.text)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for rn in p.runs:
                    total += len(rn.text)

print(f'Done! size={os.path.getsize(PTH)} chars={total} words={int(total/2)}')
print(f'Para count: {n_paras_merged}')
for sn, pl in sections.items():
    if pl:
        print(f'  {sn}: {len(pl)} paragraphs')
