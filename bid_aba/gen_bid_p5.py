import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PTH = r'D:\openclaw-workspace\bid_aba\阿坝州财政局竣工财务决算审核_投标文件.docx'
IMG = r'D:\openclaw-workspace\bid_aba\work_dir'

def shade(c, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'),color); el.set(qn('w:val'),'clear')
    c._tc.get_or_add_tcPr().append(el)

def T(doc, hds, rows, hc='1F618D', fs=9):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hds):
        c=t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        r.font.size=Pt(fs); r.font.name='Microsoft YaHei'; shade(c,hc)
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; r=p.add_run(str(ct))
            r.font.size=Pt(fs); r.font.name='Microsoft YaHei'
            if ri%2==1: shade(c,'F2F4F4')
    doc.add_paragraph()

def P(doc, text, bold=False, fs=12, fc=None, align=None, sa=6, indent=True):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(fs)
    r.font.name='仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    if fc: r.font.color.rgb=fc
    if indent and align!=WD_ALIGN_PARAGRAPH.CENTER: p.paragraph_format.first_line_indent=Cm(0.74)
    return p

def H(doc, text, level=1):
    h=doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return h

def I(doc, name, w=5.5, cap=None):
    pth=os.path.join(IMG,name)
    if os.path.exists(pth):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(pth, width=Inches(w))
        if cap:
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cr=cp.add_run(f'图：{cap}'); cr.font.size=Pt(9)
            cr.font.color.rgb=RGBColor(128,128,128); cr.font.name='Microsoft YaHei'
        doc.add_paragraph()

doc = Document(PTH)

# ====== 扩展内容：审核方法详细说明 ======
# Insert after 7.3
H(doc, '7.7 审核方法体系详解', 2)
P(doc, '在本项目的审核过程中，我单位将综合运用以下八种专业审核方法，形成科学、系统的审核方法体系。各种方法相互配合、相互印证，确保审核结论的可靠性和充分性。')

P(doc, '八种专业审核方法对照表', bold=True, fs=11, indent=False)
T(doc, ['序号','方法名称','方法描述','适用场景','操作要点'],[
    ['1','全面核查法','对竣工财务决算报告及所附全部会计账簿、凭证、合同协议、竣工图纸等资料进行全面系统查阅','适用于所有审核项目，特别是初次审核的项目','逐项核对，不遗漏任何重要资料，做好核查记录'],
    ['2','数据分析法','通过对会计数据、工程数据、预算数据进行定量分析，发现异常变动和异常数值','适用于资金管理、造价审核等数据量大需快速定位问题领域','建立比对基准线，识别偏离度超过20%的数据，重点关注异常值'],
    ['3','现场核实法','审核人员到项目现场进行实地查看、测量、盘点，验证账实是否相符','适用于工程量核实、资产盘点、隐蔽工程审核等','随机抽样+重点核实结合，拍摄现场照片作为佐证材料'],
    ['4','函证询证法','向被审核单位以外的相关方发函核实债权债务、银行存款等事项的真实性','适用于债权确认、债务核实、银行存款核对等','函证回函率≥90%，未回函的执行替代程序'],
    ['5','专家咨询法','针对专业性较强的问题征询相关领域专家意见','适用于技术争议、质量认定、造价纠纷等','出具书面专家意见，明确专家责任，咨询结果作为参考依据'],
    ['6','穿行测试法','选取若干笔重大经济业务从起点追踪至财务报表反映终点','适用于内控评价、资金流程审核等','选取3-5笔代表性业务，覆盖主要业务类型和金额段'],
    ['7','符合性测试','评价内部控制设计合理性和执行有效性的测试','适用于制度执行审核、内控评价等','采用询问、观察、检查、重新执行等方法评价控制有效性'],
    ['8','实质性测试','对各账户余额和交易类别进行测试以确认金额真实准确','适用于成本核算审核、支出合规性审核等','采用细节测试+分析程序，样本量满足统计推断要求'],
], hc='8E44AD')

P(doc, '在实际审核过程中，上述八种方法并非孤立使用，而是根据项目特点、审核重点和风险评估结果灵活组合运用。例如，在审查建设成本时，首先运用数据分析法对成本构成进行总体分析，识别异常领域；然后运用全面核查法对异常领域进行逐项审核；对发现的问题运用现场核实法进行实地验证；必要时运用函证询证法向相关方核实。通过方法组合，实现审核效率与质量的统一。')

# ====== 扩展内容：阿坝州项目审核特殊事项 ======
H(doc, '7.8 阿坝州项目审核特殊事项', 2)
P(doc, '阿坝州作为高原民族地区，其基本建设项目具有鲜明的区域特征，在竣工财务决算审核中需要重点关注以下特殊事项：')
P(doc, '第一，高原施工增加费的审核。阿坝州平均海拔3000米以上，大部分地区属于高海拔施工区域。根据四川省建筑工程计价定额的规定，高原施工增加费是指在海拔2000米以上地区施工，因气压低、空气稀薄、气候寒冷等特殊条件导致人工和机械效率降低而增加的费用。审核时应重点关注：高原施工增加费的计算基数是否正确，费率是否符合规定标准，是否存在虚报高原施工天数套取资金问题，高原施工增加费和冬季施工费是否重复计取。')
P(doc, '第二，冬季施工费的审核。阿坝州冬季漫长寒冷，11月至次年4月为霜冻期，大部分地区属于季节性冻土层。冬季施工费是指在规定冬季施工期间施工时，因采取防寒保温措施而增加的费用。审核时应当关注：冬季施工的起止日期是否符合规定，冬季施工措施费的计算是否合理，实际采取的冬季施工措施与费用是否匹配。')
P(doc, '第三，材料运输费的审核。阿坝州地广人稀、交通条件相对落后，多数建设项目所需的建筑材料需要从成都、绵阳等地区长途运输。材料运输费用在工程造价中占比较高。审核时应关注：材料运距是否合理，运费标准是否符合当地市场水平，是否存在因运输距离虚报增加费用的问题。')
P(doc, '第四，生态保护相关费用的审核。阿坝州作为长江、黄河上游生态屏障，建设项目环保要求严格。审核时应关注：环保设施是否按规定建设并正常运行，环评批复要求的环保措施是否落实，生态补偿费用是否按规定计提和使用。')
P(doc, '第五，民族文化保护费用的审核。阿坝州是藏族羌族聚居区，建设项目涉及民族传统建筑风貌保护的，审核时应关注：保护方案是否经相关部门批准，保护措施是否落实到位，保护费用是否合理。')

I(doc, 'fig4-confidentiality.drawio.png', 5.5, '保密管理体系架构图')

# ====== 扩展内容：合同履约保障措施 ======
H(doc, '7.9 合同履约保障措施', 2)
P(doc, '我单位承诺，如中选本项目，将严格履行合同义务，确保项目顺利完成。具体的合同履约保障措施如下：')

T(doc, ['保障领域','具体保障措施','违约后果'],[
    ['质量保障','严格执行三级复核制度，所有审核报告经技术负责人审定后出具','经审核认定未达质量要求的，按审减服务费的5%核减费用；重大质量问题，项目费用不予支付'],
    ['进度保障','按倒排工期计划推进，建立进度预警机制，超出时限2000元/天处罚','累计3个项目未按时完成，委托方有权清退并解除合同'],
    ['保密保障','全员签署保密承诺，涉密信息专管专用','泄密行为承担法律和经济责任'],
    ['廉洁保障','签署廉洁从业承诺书，主动接受委托方监督','违规行为严肃处理，情况严重的解除合同'],
    ['沟通保障','指定专人对接委托方，每周以书面形式报告工作进展','未及时报告的视为违约，涉及重大事项未报告的承担相应责任'],
    ['人员保障','项目负责人和核心成员全程不更换，确需变更的须经委托方书面同意','擅自更换人员的按违约处理'],
], hc='E74C3C')

P(doc, '我单位理解并接受比选文件中关于合同履约的各项条款，保证在合同履行过程中严格遵照执行，如因我单位原因造成违约，愿意承担相应的违约责任和处罚。')

doc.save(PTH)
print(f'OK: Part 5 (expansion) complete, bytes={os.path.getsize(PTH)}')
