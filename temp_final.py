# -*- coding: utf-8 -*-
"""PPP可用性付费报告 — v4.0 措辞严谨修正版"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# PARAMETERS
# ============================================================
CAPITAL_AMORT = 80_000_000.00
CAPITAL_LUMP = 22_298_077.79
RATE = 0.0799
YEARS = 18
A = CAPITAL_AMORT * RATE * (1+RATE)**YEARS / ((1+RATE)**YEARS - 1)

# Capital schedule
cr = CAPITAL_AMORT; cs = []
for y in range(YEARS):
    interest = cr * RATE; principal = A - interest
    if y == YEARS-1: principal = cr; interest = A - principal
    cr -= principal
    if cr < 0: cr = 0
    cs.append((principal, interest, A, max(cr, 0)))

# Bank: original table2 structure, scaled to bank-confirmed total 710,822,119.32
bo = [(31686745.84,381300000),(26809671.03,381300000),(24811756.73,381250000),
      (21351219.32,380250000),(20263854.17,379250000),(20265833.33,378250000),
      (20162465.30,377250000),(20111770.82,376250000),(23022951.39,372250000),
      (22871874.99,368250000),(22617395.82,364250000),(22414618.06,360250000),
      (27148298.64,351250000),(27728124.99,341250000),(35070729.16,323250000),
      (34158229.18,305250000),(311968819.45,0),(0,0)]
sf = 710822119.32 / sum(b[0] for b in bo)
bs = [(round(a*sf,2), r) for a,r in bo]

# Op costs
OP = [5768117.42, 6890208.16, 5713241.86] + [5710000]*15
# Income
INC = [4285070.31, 6229093.45, 5564017.88] + [5560000]*15

yl = list(range(2023,2041))
yd = []
for i in range(YEARS):
    pr, interest, pmt, rem = cs[i]
    lump = CAPITAL_LUMP if i==YEARS-1 else 0
    tc = pmt + lump
    bk = bs[i][0]; op = OP[i]; inc = INC[i]
    avail = tc + bk + op - inc
    yd.append((yl[i], pr, interest, pmt, lump, tc, rem, bk, op, inc, avail))

gt = sum(y[10] for y in yd)
ct = sum(y[5] for y in yd)
bt = sum(b[0] for b in bs)
ot = sum(OP); it = sum(INC)

print(f"A={A:,.2f} Cap={ct:,.2f} Bank={bt:,.2f} Op={ot:,.2f} Inc={it:,.2f} Total={gt:,.2f}")

# ============================================================
doc = Document()
for s in doc.sections:
    s.page_width=Cm(21); s.page_height=Cm(29.7)
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2); s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)

sty=doc.styles['Normal']; sty.font.name='宋体'; sty.font.size=Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); sty.paragraph_format.line_spacing=1.5

def sc(cell, text, fn='宋体', sz=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=align
    r=p.add_run(text); r.font.name=fn; r.font.size=sz; r.font.bold=bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'),fn)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1

def scs(cell, c):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}"/>'))

def pa(text, fn='宋体', sz=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None, sa=Pt(0)):
    p=doc.add_paragraph(); p.alignment=align
    if indent: p.paragraph_format.first_line_indent=indent
    p.paragraph_format.space_after=sa; p.paragraph_format.line_spacing=1.5
    r=p.add_run(text); r.font.name=fn; r.font.size=sz; r.font.bold=bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'),fn)
    return p

def tb(headers, rows, cw=None):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; sc(c,h,'微软雅黑',Pt(9),True,color=RGBColor(0xFF,0xFF,0xFF)); scs(c,'0A1F3F')
    for i,rd in enumerate(rows):
        for j,v in enumerate(rd):
            c=t.rows[i+1].cells[j]; sc(c,str(v),'宋体',Pt(8.5))
            if i%2==1: scs(c,'F5F2EC')
    if cw:
        for row in t.rows:
            for j,w in enumerate(cw): row.cells[j].width=Cm(w)
    return t

# ===== COVER =====
for _ in range(6): doc.add_paragraph()
pa('巴中市恩阳医养园PPP项目','方正小标宋简体',Pt(22),True,WD_ALIGN_PARAGRAPH.CENTER)
pa('可用性付费测算结果报告','方正小标宋简体',Pt(22),True,WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4): doc.add_paragraph()
pa('川融策咨询〔2025〕第490号',sz=Pt(14),align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6): doc.add_paragraph()
pa('四 川 融 策 会 计 师 事 务 所 有 限 公 司','微软雅黑',Pt(16),True,WD_ALIGN_PARAGRAPH.CENTER)
pa('Sichuan Rongce Accounting Firm Co., Ltd',sz=Pt(10),align=WD_ALIGN_PARAGRAPH.CENTER,sa=Pt(20))
pa('2026年7月29日',sz=Pt(14),align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===== BODY =====
pa('川融策咨询〔2025〕第  号',align=WD_ALIGN_PARAGRAPH.RIGHT,sz=Pt(10))
doc.add_paragraph()
pa('巴中市恩阳医养园PPP项目','黑体',Pt(14),True,WD_ALIGN_PARAGRAPH.CENTER)
pa('可用性付费测算结果报告','黑体',Pt(14),True,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
pa('巴中市恩阳区卫生健康局：',indent=Cm(0.74))
pa('我们接受委托，对巴中市恩阳医养园PPP项目2022年10月28日至2040年12月21日可用性付费进行测算。本次测算以《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）审定的建筑安装工程费用485,098,077.79元为基数。2023-2024年度运维成本及第三方收入沿用川融策咨询〔2025〕第490号征求意见稿（以下简称"征求意见稿"）表3数据，2025年度数据经本所依据项目公司原始账套独立验证后修正。现将测算结果报告如下：',indent=Cm(0.74))

# ===== 一 =====
pa('一、项目概况','黑体',Pt(14),True)
pa('2015年12月21日，巴中市恩阳医养园PPP项目经巴中市恩阳区人民政府（恩府〔2015〕208号）批复同意实施，由巴中市恩阳区卫生健康局按规定有序开展。',indent=Cm(0.74))
pa('2016年7月项目协议经区人民政府审定同意，巴中市恩阳区卫生健康局于2016年12月与恩阳区人民医院、湖南省第五工程有限公司签订PPP协议，成立项目管理公司（注册资本金10,000万元，截至测算基准日实缴资本金102,298,077.79元），承担该项目的投资、建设、运营维护及移交。',indent=Cm(0.74))
pa('2017年7月19日项目公司与湖南省第五工程有限公司签订施工合同，项目建筑总面积118,515.13\u33a1，包括院内绿化、场地及路面硬化、室内管网等附属工程及设备购置等。项目实际于2017年5月17日开工，2022年3月完工，同年3月28日进入调试运营期，10月27日通过竣工验收，次日（2022年10月28日）正式进入运营阶段。',indent=Cm(0.74))
pa('根据2021年10月15日《巴中市恩阳医养园PPP项目合同之补充合同（2021年）》，项目合作期调整为23年，运营期为18年（2022年10月28日至2040年12月21日）。',indent=Cm(0.74))

# ===== 二 =====
pa('二、测算依据','黑体',Pt(14),True)
pa('（一）政策依据','楷体',Pt(12),True)
pa('（1）《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）',indent=Cm(0.74))
pa('（2）财政部关于印发《政府和社会资本合作项目财政承受能力论证指引》的通知（财金〔2015〕21号）',indent=Cm(0.74))
pa('（3）四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）',indent=Cm(0.74))
pa('（二）项目资料依据','楷体',Pt(12),True)
for i,item in enumerate([
    '《恩阳医养园PPP项目协议》（2016年12月27日签订）',
    '《巴中市恩阳医养园PPP项目合同》（2017年签订）',
    '《巴中市恩阳医养园PPP项目合同之补充合同》（2017年/2018年/2021年）',
    '《关于巴中市恩阳区人民医院一期工程建设项目可行性研究报告的批复》（恩区发改行审〔2015〕83号）',
    '《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）——验证竣工结算额',
    '《借款还本付息情况报告》（贵阳银行股份有限公司成都分行出具）——验证融资本息合计数',
    '项目公司2025年度财务报表及运营成本/收入明细账（6册XLS，经本所独立验证）',
    '2022-2025年度运营期绩效考核结果（恩阳区卫健局出具，本所尚未核验原件）',
],1):
    pa(f'（{i}）{item}',indent=Cm(0.74))

# ===== 三 =====
pa('三、测算说明','黑体',Pt(14),True)
pa('（一）项目回报机制识别','楷体',Pt(12),True)
pa('一是政策方面：根据财金〔2014〕113号文，项目回报机制分为使用者付费、可行性缺口补助和政府付费三种。川财金〔2017〕91号文进一步规定，可行性缺口补助模式下政府承担部分运营补贴支出责任。二是项目资料方面：PPP合同及实施方案均明确本项目回报机制为"可行性缺口补助"，2016年12月27日签订的《恩阳医养园PPP项目协议》约定区政府将可行性缺口补贴纳入跨年度财政预算并提请人大决议。',indent=Cm(0.74))

pa('（二）回报计算方法选用','楷体',Pt(12),True)
pa('根据2017年7月签订的补充合同，计算公式为：A = P\u00d7k\u00d7(1+k)^n / ((1+k)^n-1) + 实际融资成本 + 运营维护成本\u00d7(1+K) \u2212 第三方收入。其中A为各年政府运营补贴，K为运维成本加成系数（本项目实际操作中K取0，即运维成本不加成）。',indent=Cm(0.74))

pa('（1）项目资本金回报：P=实缴资本金102,298,077.79元；k=社会资本中标年回报率7.99%；n=18年。测算方案：其中8,000万元以等额本息方式在18年内逐年偿还（年还款额{:,.2f}元），剩余22,298,077.79元于第18年一次性支付、不计算回报利息。'.format(A),indent=Cm(0.74))

pa('（2）实际融资成本：贵阳银行成都分行《借款还本付息情况报告》确认18年合计710,822,119.32元。本所尚未取得银行逐年还款计划明细，逐年分配沿用征求意见稿表2的还款结构，按银行确认总数710,822,119.32/原表合计712,464,358.22\u22480.997695比例统一缩放。该分配方式经2025年度项目公司财务费用明细交叉验证（第3季度计提贷款利息6,291,985.59元，本金38,125万元），与还款计划中对应年度数值基本吻合。原合同利率6.95%/年，经三次下调至5.0%/年。',indent=Cm(0.74))

pa('（3）运营维护成本：2023-2024年度分别为5,768,117.42元、6,890,208.16元（沿用征求意见稿表3；该数据与原报告正文按经营年度划分的成本描述存在口径差异，原因为2023年度支付期覆盖约14个月，与原正文中12个月经营年度的切分不一致——本所未取得2023-2024年原始账套进行独立验证）；2025年度经独立验证为5,713,241.86元（主营业务成本4,474,027.75元+管理费用1,239,214.11元，取自项目公司2025年度XLS明细账4季度损益结转合计，不含年末调整项）；2026-2040年度暂以2025年度独立验证值取整5,710,000.00元/年估算。',indent=Cm(0.74))

pa('（4）第三方收入：2023-2024年度分别为4,285,070.31元、6,229,093.45元（沿用征求意见稿表3）；2025年度经独立验证为5,564,017.88元（主营业务收入4,961,142.62元+其他业务收入602,875.26元）；2026-2040年度暂以5,560,000.00元/年估算。',indent=Cm(0.74))

# ===== 四 =====
pa('四、测算过程','黑体',Pt(14),True)

# --- (一) ---
pa('（一）每年项目资本金回报测算情况','楷体',Pt(12),True)
pa('项目资本金回报详见表1（单位：元）：',indent=Cm(0.74))

t1h=['年','等额本息\n本金','等额本息\n利息','等额本息\n小计','一次性支付','资本金回报\n合计','剩余\n(等额本息)']
t1r=[]
for y in yd:
    yr,pr,interest,pmt,lump,tc,rem,*_=y
    t1r.append([str(yr),f'{pr:,.2f}',f'{interest:,.2f}',f'{pmt:,.2f}',
                f'{lump:,.2f}'if lump else'-',f'{tc:,.2f}',
                f'{rem:,.2f}'if rem>0 else'0.00'])
sp=sum(y[1]for y in yd); si=sum(y[2]for y in yd); sa=sum(y[3]for y in yd)
t1r.append(['合计',f'{sp:,.2f}',f'{si:,.2f}',f'{sa:,.2f}',f'{CAPITAL_LUMP:,.2f}',f'{ct:,.2f}','-'])
tb(t1h,t1r,cw=[1.1,2.3,2.3,2.3,2.2,2.3,2.8])
doc.add_paragraph()

# --- (二) ---
pa('（二）每年实际融资本息测算情况','楷体',Pt(12),True)
pa('项目融资本息见表2（单位：元）。逐年数据以征求意见稿表2还款结构为基础，按银行报告确认总数710,822,119.32元等比缩放：',indent=Cm(0.74))

t2h=['年','实际融资本息','剩余融资金额','适用利率']
t2r=[]
rates={0:'6.95%',1:'6.45%',2:'5.5%/5.0%',3:'5.0%',4:'5.0%',5:'5.0%',6:'5.0%',
       7:'5.0%',8:'5.0%',9:'5.0%',10:'5.0%',11:'5.0%',12:'5.0%',13:'5.0%',
       14:'5.0%',15:'5.0%',16:'5.0%',17:'-'}
for i,y in enumerate(yd):
    yr,_,_,_,_,_,_,bk,_,_,_ = y; br=bs[i][1]; rs=rates.get(i,'5.0%')
    t2r.append([str(yr),f'{bk:,.2f}',f'{br:,.2f}'if br>0 else'-',rs])
t2r.append(['合计',f'{bt:,.2f}','-','\u2014'])
tb(t2h,t2r,cw=[1.5,4.0,4.0,2.0])
doc.add_paragraph()
pa('注1：逐年融资本息系以银行确认总数710,822,119.32元为锚，按征求意见稿表2还款结构等比缩放，非银行原始逐项数据。',indent=Cm(0.74),sz=Pt(10))
pa('注2：第18年（2040年度）融资本息为0元，系贵阳银行贷款已于第17年（2039年12月）全部结清。',indent=Cm(0.74),sz=Pt(10))

# --- (三)(四)(五) ---
pa('（三）运营维护成本','楷体',Pt(12),True)
pa('2023年度：5,768,117.42元（覆盖期2022.10.28-2023.12.21，约14个月。沿用征求意见稿表3，本所未取得2023年度原始账套独立验证。该数值表3与原报告正文按经营年度划分的成本描述存在口径差异——正文描述第一经营年度5,341,344.03元、第二经营年度5,948,956.93元，二者与表3支付年度数值因期限切分不同而无法直接对应）。',indent=Cm(0.74))
pa('2024年度：6,890,208.16元（覆盖期2023.12.21-2024.12.21，12个月。同上，沿用征求意见稿表3）。',indent=Cm(0.74))
pa('2025年度\u3010独立验证\u3011：5,713,241.86元（覆盖期2024.12.21-2025.12.21。计算公式：主营业务成本4,474,027.75元+管理费用1,239,214.11元=5,713,241.86元。其中主营业务成本=1,016,316.83(Q1)+667,176.24(Q2)+1,476,847.53(Q3)+1,313,687.15(Q4)；管理费用=558,779.11(Q1)+287,198.91(Q2)+213,914.32(Q3)+179,321.77(Q4)。以上季度数值均取自项目公司2025年度XLS明细账损益结转账面金额，不含年末调整项）。',indent=Cm(0.74))
pa('2026-2040年度：暂以2025年度独立验证值取整5,710,000.00元/年估算（共15年）。',indent=Cm(0.74))

pa('（四）第三方收入','楷体',Pt(12),True)
pa('2023年度：4,285,070.31元（沿用征求意见稿表3）。',indent=Cm(0.74))
pa('2024年度：6,229,093.45元（沿用征求意见稿表3，含第三经营年度2024.10.28-2024.12.31期间387,927.58元）。',indent=Cm(0.74))
pa('2025年度\u3010独立验证\u3011：5,564,017.88元（计算公式：主营业务收入4,961,142.62元+其他业务收入602,875.26元=5,564,017.88元。其中主营收入=1,160,281.89(Q1)+1,210,580.37(Q2)+1,407,481.80(Q3)+1,182,798.56(Q4)；其他收入=166,207.80(Q1)+29,264.39(Q2)+173,959.48(Q3)+233,443.59(Q4)。来源同上，均取自2025年度XLS明细账损益结转账面金额）。',indent=Cm(0.74))
pa('2026-2040年度：暂以2025年度独立验证值取整5,560,000.00元/年估算（共15年）。',indent=Cm(0.74))

pa('（五）考核结果','楷体',Pt(12),True)
pa('根据2018年3月7日补充协议，年度绩效考核分数\u226580分时可用性付费支付比例为100%。截至2026年6月30日，已完成的四个运营年度考核结果（据恩阳区卫健局出具，本所未核验原件）：2022年度93分、2023年度94分、2024年度95分、2025年度93.5分，均\u226580分，对应支付比例100%。',indent=Cm(0.74))

# ===== 五 =====
pa('五、测算结果','黑体',Pt(14),True)
pa('项目可用性付费汇总见表3（单位：元）：',indent=Cm(0.74))

t3h=['年','资本金回报','实际融资本息','运营维护成本','第三方收入','可用性付费']
t3r=[]
for y in yd:
    yr,_,_,_,_,tc,_,bk,op,inc,avail=y
    t3r.append([str(yr),f'{tc:,.2f}',f'{bk:,.2f}',f'{op:,.2f}',f'{inc:,.2f}',f'{avail:,.2f}'])
t3r.append(['合计',f'{ct:,.2f}',f'{bt:,.2f}',f'{ot:,.2f}',f'{it:,.2f}',f'{gt:,.2f}'])
tb(t3h,t3r,cw=[1.2,2.8,2.8,2.8,2.8,2.8])
doc.add_paragraph()

pa(f'测算结论：巴中市恩阳医养园PPP项目运营期（2022年10月28日至2040年12月21日）18年可用性付费总额为{gt:,.2f}元（约{gt/1e8:.2f}亿元）。',indent=Cm(0.74),bold=True)

# ===== 备注 =====
pa('备注：','楷体',Pt(10),True)
pa('1.计费期限：除2023年度覆盖2022年10月28日至2023年12月21日（约14个月）外，其余年度均为上一年度12月21日至当年12月21日。2023年度资本金回报按全年等额值计算（非按14/12比例调整），与原征求意见稿处理方式一致。',sz=Pt(10))
pa(f'2.资本金方案：实缴资本金102,298,077.79元。测算分为两部分\u2014\u20148,000万元按P={CAPITAL_AMORT:,.0f}、k=7.99%、n=18年等额本息逐年偿还，18年合计153,544,445.28元（本金{CAPITAL_AMORT:,.0f}元+利息73,544,445.28元）；22,298,077.79元于第18年一次性支付，不计利息。',sz=Pt(10))
pa('3.融资本息取证层级：①合计数710,822,119.32元\u2014\u2014贵阳银行成都分行《借款还本付息情况报告》确认（一级证据）；②逐年分配\u2014\u2014沿用征求意见稿表2还款结构等比缩放（二级推导，非银行逐项原始数据）；③2025年Q3财务费用计提利息6,291,985.59元与还款计划估值基本吻合（侧面验证）。如需正式报告，建议取得银行逐年还款明细表替换逐年推导值。',sz=Pt(10))
pa('4.运维成本与第三方收入取证层级：①2023-2024年度\u2014\u2014沿用征求意见稿表3（二级证据，本所尚未取得原始账套独立验证）；②2025年度\u2014\u2014项目公司6册XLS明细账4季度损益结转（一级证据，可追溯至凭证行）；③2026-2040年度\u2014\u2014基于2025年度一级证据取整估算（三级估算）。',sz=Pt(10))
pa('5.本报告系征求意见稿。正式出具前建议补充以下原始资料以提升数据独立验证覆盖率：①2023-2024年度项目公司运营成本/收入明细账；②贵阳银行逐年还款计划明细表；③恩阳区卫健局绩效考核正式文件。',sz=Pt(10))
pa('6.利率变动风险：贵阳银行贷款利率已从6.95%经三次下调至5.0%，未来LPR变动将影响实际融资本息。',sz=Pt(10))

doc.add_paragraph(); doc.add_paragraph()
pa('四川融策会计师事务所有限公司',align=WD_ALIGN_PARAGRAPH.RIGHT,sz=Pt(12))
pa('二\u3007二六年七月二十九日',align=WD_ALIGN_PARAGRAPH.RIGHT,sz=Pt(12))

outpath = r'C:\Users\scrccpa\Desktop\恩阳医养园PPP项目可用性付费测算报告（修订版v4-20260729）.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
print(f'Size: {os.path.getsize(outpath):,} bytes')
print(f'Total={gt:,.2f} ({gt/1e8:.2f}亿)')
