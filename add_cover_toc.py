import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

COMPLETE = r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx'

doc = Document(COMPLETE)
body = doc.element.body

# ============================================
# Step 1: Remove old cover/TOC (first ~25 paragraphs)
# ============================================
to_remove = []
for i in range(min(25, len(doc.paragraphs))):
    to_remove.append(doc.paragraphs[i]._element)

for elem in to_remove:
    body.remove(elem)

print(f'Removed {len(to_remove)} old cover paragraphs')

# ============================================
# Step 2: Create new cover page
# ============================================
def add_cover_para(text, bold=False, size=12, color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='微软雅黑', space_before=0, space_after=0):
    """Insert a new paragraph at the beginning of the document"""
    p_elem = etree.SubElement(body, qn('w:p'))
    # Add to beginning
    body.insert(0, p_elem)
    
    pPr = etree.SubElement(p_elem, qn('w:pPr'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        jc.set(qn('w:val'), 'center')
    elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
        jc.set(qn('w:val'), 'left')
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        jc.set(qn('w:val'), 'right')
    
    if space_before:
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), str(space_before))
    if space_after:
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:after'), str(space_after))
    
    r = etree.SubElement(p_elem, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    
    if bold:
        etree.SubElement(rPr, qn('w:b'))
    
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(size * 2))  # half-points
    
    if color:
        c = etree.SubElement(rPr, qn('w:color'))
        c.set(qn('w:val'), color)
    
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    
    return p_elem

# Blank lines at top
for _ in range(4):
    add_cover_para('', size=12)

# Title
add_cover_para('四川融策', bold=True, size=26, color='0A1F3F')
add_cover_para('━' * 30, size=10, color='C5955C')
add_cover_para('公司制度体系', bold=True, size=22, color='0A1F3F')
add_cover_para('（完整版）', size=14, color='1A5C6E')
add_cover_para('━' * 30, size=10, color='C5955C')
add_cover_para('', size=12)
add_cover_para('四川融策会计师事务所有限公司', size=12)
add_cover_para('四川融策工程咨询有限公司', size=12)
add_cover_para('二〇二六年七月', size=12)

# Page break
pb = etree.Element(qn('w:p'))
r = etree.SubElement(pb, qn('w:r'))
br = etree.SubElement(r, qn('w:br'))
br.set(qn('w:type'), 'page')
body.insert(0, pb)  # insert page break at start, pushing cover to separate page

print('Cover page added')

# ============================================
# Step 3: Insert TOC with section entries
# ============================================
# TOC after cover (before first content paragraph, which is now 薪酬管理制度)

# Find first content doc
first_doc_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text or ''
    if not t:
        continue
    try:
        is_bold = any(r.bold for r in p.runs if r.bold)
    except AttributeError:
        continue
    if '四川融策薪酬管理制度' in t and is_bold:
        first_doc_idx = i
        break

if first_doc_idx:
    first_elem = doc.paragraphs[first_doc_idx]._element
    
    # TOC heading
    toc_heading = etree.Element(qn('w:p'))
    pPr = etree.SubElement(toc_heading, qn('w:pPr'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'center')
    r = etree.SubElement(toc_heading, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    etree.SubElement(rPr, qn('w:b'))
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '36')  # 18pt
    c = etree.SubElement(rPr, qn('w:color'))
    c.set(qn('w:val'), '0A1F3F')
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    t = etree.SubElement(r, qn('w:t'))
    t.text = '目    录'
    t.set(qn('xml:space'), 'preserve')
    first_elem.addprevious(toc_heading)
    
    # Subtitle
    sub = etree.Element(qn('w:p'))
    r2 = etree.SubElement(sub, qn('w:r'))
    t2 = etree.SubElement(r2, qn('w:t'))
    t2.text = '制度体系总览'
    t2.set(qn('xml:space'), 'preserve')
    first_elem.addprevious(sub)
    
    # Overview text
    overview_text = '四川融策公司制度体系涵盖五大篇章，共38项管理制度及配套文件，形成覆盖人力资源管理、财务管理、业务运营管理、业务质量控制、行政综合管理的完整内控体系。'
    ov = etree.Element(qn('w:p'))
    r3 = etree.SubElement(ov, qn('w:r'))
    t3 = etree.SubElement(r3, qn('w:t'))
    t3.text = overview_text
    t3.set(qn('xml:space'), 'preserve')
    first_elem.addprevious(ov)
    
    # Empty line
    empty_p = etree.Element(qn('w:p'))
    first_elem.addprevious(empty_p)
    
    # Section entries
    section_info = [
        ('人力资源篇', '9项制度 · 17份岗位说明书 · 薪酬绩效体系'),
        ('财务管理篇', '6项制度 · 报销标准 · 预算资金管理 · 利润核算'),
        ('业务部管理篇', '6项制度 · 组织架构 · 客户管理 · 投标调配 · 经营分析'),
        ('业务质控篇', '11项制度 · 三级复核 · 质量控制 · 执业责任'),
        ('行政综合篇', '12项制度 · 信息保密 · 印章证照 · 档案管理 · 公司章程'),
    ]
    
    for label, desc in section_info:
        # Section title
        sec_p = etree.Element(qn('w:p'))
        pPr = etree.SubElement(sec_p, qn('w:pPr'))
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '120')
        r = etree.SubElement(sec_p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        etree.SubElement(rPr, qn('w:b'))
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), '24')  # 12pt
        c = etree.SubElement(rPr, qn('w:color'))
        c.set(qn('w:val'), '0A1F3F')
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        t = etree.SubElement(r, qn('w:t'))
        t.text = label
        t.set(qn('xml:space'), 'preserve')
        first_elem.addprevious(sec_p)
        
        # Description
        desc_p = etree.Element(qn('w:p'))
        r2 = etree.SubElement(desc_p, qn('w:r'))
        rPr2 = etree.SubElement(r2, qn('w:rPr'))
        sz2 = etree.SubElement(rPr2, qn('w:sz'))
        sz2.set(qn('w:val'), '20')  # 10pt
        c2 = etree.SubElement(rPr2, qn('w:color'))
        c2.set(qn('w:val'), '666666')
        t2 = etree.SubElement(r2, qn('w:t'))
        t2.text = f'    {desc}'
        t2.set(qn('xml:space'), 'preserve')
        first_elem.addprevious(desc_p)
    
    print(f'TOC updated with {len(section_info)} sections')

# ============================================
# Step 4: Add section boundary markers
# ============================================
# Insert "人力资源篇", "财务管理篇", etc. before each section's first document

section_boundaries = {
    '四川融策薪酬管理制度': ('人力资源篇', '—— 人力资源篇 ——'),
    '四川融策财务报销管理制度': ('财务管理篇', '—— 财务管理篇 ——'),
    'RC-BIZ-003': ('业务部管理篇', '—— 业务部管理篇 ——'),
    '四川融策项目管理规范': ('业务质控篇', '—— 业务质控篇 ——'),
    '四川融策制度发布与版本管理规范': ('行政综合篇', '—— 行政综合篇 ——'),
}

markers_added = 0
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if not t:
        continue
    try:
        is_bold = any(r.bold for r in p.runs if r.bold)
    except AttributeError:
        continue
    
    if not is_bold:
        continue
    
    for trigger, (section_name, section_subtitle) in section_boundaries.items():
        if trigger in t:
            elem = p._element
            
            # Page break before section
            pb = etree.Element(qn('w:p'))
            r = etree.SubElement(pb, qn('w:r'))
            br = etree.SubElement(r, qn('w:br'))
            br.set(qn('w:type'), 'page')
            elem.addprevious(pb)
            
            # Section title
            sec_p = etree.Element(qn('w:p'))
            pPr = etree.SubElement(sec_p, qn('w:pPr'))
            jc = etree.SubElement(pPr, qn('w:jc'))
            jc.set(qn('w:val'), 'center')
            sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:before'), '200')
            sp.set(qn('w:after'), '200')
            
            r = etree.SubElement(sec_p, qn('w:r'))
            rPr = etree.SubElement(r, qn('w:rPr'))
            etree.SubElement(rPr, qn('w:b'))
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), '36')  # 18pt
            c = etree.SubElement(rPr, qn('w:color'))
            c.set(qn('w:val'), '0A1F3F')
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            t = etree.SubElement(r, qn('w:t'))
            t.text = section_subtitle
            t.set(qn('xml:space'), 'preserve')
            elem.addprevious(sec_p)
            
            markers_added += 1
            print(f'  Added boundary: {section_name}')
            break

print(f'Added {markers_added} section boundaries')

# ============================================
# Step 5: Save
# ============================================
doc.save(COMPLETE)

import os
print(f'\nSaved: {COMPLETE}')
print(f'Size: {os.path.getsize(COMPLETE):,} bytes')
print(f'Paragraphs: {len(doc.paragraphs)}')
