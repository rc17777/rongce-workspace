import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from docx import Document
from docx.oxml.ns import qn

COMPLETE = r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx'

doc = Document(COMPLETE)
body = doc.element.body

# ==========================================
# Step 1: Find boundaries
# ==========================================

# Find 业务质控篇 section heading (deep in doc, after TOC)
biz_qc_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '业务质控篇' and i > 500:
        biz_qc_idx = i
        break

print(f'业务质控篇 section: paragraph {biz_qc_idx}')

# The inserted block starts right after 财务管理篇 ends.
# Find where the first RC-BIZ overview item or BIZ doc heading appears
# by walking backwards from 业务质控篇

insert_start = None
for i in range(biz_qc_idx - 1, 0, -1):
    t = doc.paragraphs[i].text.strip()
    # Look for the overview items which start with 【
    if t.startswith('【RC-BIZ-00'):
        # This is a reversed overview item - walk back to find first
        continue
    # Look for the actual BIZ doc content: "业务部经营数据统计与分析制度"
    if t == '业务部经营数据统计与分析制度':
        continue
    # Check for text that's from 财务管理篇
    if '可分配利润核算细则' in t or '财务管理篇' in t:
        # Found end of FIN section - next paragraph should be start of inserted content
        # But we need to walk forward to find the actual start (the paragraph AFTER FIN section ends)
        # Since content is reversed, RC-BIZ-008's last chapter is right after FIN
        # Walk forward from here
        for j in range(i + 1, biz_qc_idx):
            t2 = doc.paragraphs[j].text.strip()
            if t2 and ('RC-BIZ' in t2 or '业务部' in t2 or '【RC-BIZ' in t2):
                # We might need to go back further - check if there's a page break
                # Look for empty paragraphs or page breaks before this
                start_j = j
                for k in range(j - 1, max(0, j - 5), -1):
                    tk = doc.paragraphs[k].text.strip()
                    if not tk:
                        start_j = k
                    else:
                        break
                insert_start = start_j
                break
        if insert_start:
            break

if insert_start is None:
    print('ERROR: Could not find insertion start')
    sys.exit(1)

insert_end = biz_qc_idx - 1

print(f'Insert start: paragraph {insert_start}')
print(f'Insert end: paragraph {insert_end}')
print(f'Block size: {insert_end - insert_start + 1} paragraphs')

# Verify first and last
print(f'\nFirst 5:')
for i in range(insert_start, min(insert_start + 5, insert_end + 1)):
    t = doc.paragraphs[i].text.strip()
    print(f'  [{i}] {t[:90] if t else "(empty)"}')
print(f'\nLast 5:')
for i in range(max(insert_start, insert_end - 4), insert_end + 1):
    t = doc.paragraphs[i].text.strip()
    print(f'  [{i}] {t[:90] if t else "(empty)"}')

# ==========================================
# Step 2: Extract, remove, reverse, re-insert
# ==========================================

# Collect paragraphs
inserted_elements = []
for i in range(insert_start, insert_end + 1):
    inserted_elements.append(doc.paragraphs[i]._element)

# Remove them all
for elem in inserted_elements:
    body.remove(elem)

print(f'\nRemoved {len(inserted_elements)} paragraphs')

# Find 业务质控篇 again (indices shifted after removal)
biz_qc_elem = None
for p in doc.paragraphs:
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '业务质控篇' and p._element.getparent() is body:
        biz_qc_elem = p._element
        break

if biz_qc_elem is None:
    print('ERROR: 业务质控篇 not found after removal')
    sys.exit(1)

# Reverse the list to restore original order
inserted_elements.reverse()
# Now inserted_elements is in original document order

# Insert in forward order using addprevious
# addprevious(X for target): inserts X immediately before target
# Forward iteration: [A,B,C] → addprevious(A)→[A,target], addprevious(B)→[A,B,target], addprevious(C)→[A,B,C,target] ✓
for elem in inserted_elements:
    biz_qc_elem.addprevious(copy.deepcopy(elem))

print(f'Re-inserted {len(inserted_elements)} paragraphs in correct order')

# ==========================================
# Step 3: Verify
# ==========================================
doc.save(COMPLETE)

import os
print(f'Saved: {COMPLETE}')
print(f'Size: {os.path.getsize(COMPLETE):,} bytes')

# Quick sanity check
doc2 = Document(COMPLETE)
print(f'\nVerification - section TOC entries:')
for i, p in enumerate(doc2.paragraphs[:80]):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and any(x in t for x in ['人力资源篇', '财务管理篇', '业务部管理篇', '业务质控篇', '行政综合篇']):
        print(f'  TOC [{i}] {t}')

# Check that the inserted section content is in correct order
print(f'\nVerification - BIZ section headings in body:')
found_biz = False
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '业务部管理篇':
        found_biz = True
    if found_biz and is_bold and ('RC-BIZ' in t or '业务部' in t or '【RC-BIZ' in t):
        print(f'  [{i}] {t[:90]}')
    if is_bold and t == '业务质控篇' and found_biz:
        print(f'  [{i}] {t[:90]}')
        break
