"""Recalculate with corrected parameters:
1. Bank repayment: use bank schedule (710,822,119.32)
2. Capital 8000万 at 7.99%, 18yr equal installment
3. Remaining ~2229.8万 lump sum in year 18
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

# ===== Parameters =====
capital_equity = 80_000_000.00       # 8000万按7.99%等额本息
capital_lumpsum = 22_298_077.79      # 剩余2229.8万最后一年一次支付
rate_k = 0.0799                       # 7.99%
n_years = 18

# ===== 1. Calculate 8000万 等额本息 =====
# A = P * k * (1+k)^n / ((1+k)^n - 1)
factor = (1 + rate_k) ** n_years
annual_payment = capital_equity * rate_k * factor / (factor - 1)
print(f"8000万等额本息: 每年 = {annual_payment:,.2f}")

# Year by year breakdown
remaining = capital_equity
equity_schedule = []
total_interest = 0
total_principal = 0

for yr in range(1, n_years + 1):
    interest = remaining * rate_k
    principal = annual_payment - interest
    # Last year adjustment to zero out
    if yr == n_years:
        principal = remaining
        interest = annual_payment - principal
    
    remaining -= principal
    total_interest += interest
    total_principal += principal
    
    equity_schedule.append({
        'year': yr,
        'payment': round(annual_payment, 2),
        'principal': round(principal, 2),
        'interest': round(interest, 2),
        'remaining': max(0, round(remaining, 2))
    })

print(f"8000万: 总还款={total_principal+total_interest:,.2f}, 本金={total_principal:,.2f}, 利息={total_interest:,.2f}")

# ===== 2. Bank repayment schedule (from bank report) =====
bank_schedule = [
    (1,  2023, 31_686_745.84,  1_500_000.00,  30_186_745.84),
    (2,  2024, 26_809_671.03,  0,             26_809_671.03),
    (3,  2025, 24_811_756.73,  50_000.00,     24_761_756.73),
    (4,  2026, 21_351_219.32,  1_000_000.00,  20_351_219.32),
    (5,  2027, 20_263_854.17,  1_000_000.00,  19_263_854.17),
    (6,  2028, 20_265_833.33,  1_000_000.00,  19_265_833.33),
    (7,  2029, 20_162_465.30,  1_000_000.00,  19_162_465.30),
    (8,  2030, 23_073_645.82,  4_000_000.00,  19_073_645.82),
    (9,  2031, 22_870_868.06,  4_000_000.00,  18_870_868.06),
    (10, 2032, 22_719_375.01,  4_000_000.00,  18_719_375.01),
    (11, 2033, 22_465_312.49,  4_000_000.00,  18_465_312.49),
    (12, 2034, 22_262_534.71,  4_000_000.00,  18_262_534.71),
    (13, 2035, 26_996_215.28,  9_000_000.00,  17_996_215.28),
    (14, 2036, 27_575_625.01,  10_000_000.00, 17_575_625.01),
    (15, 2037, 34_918_645.83,  18_000_000.00, 16_918_645.83),
    (16, 2038, 34_006_145.84,  18_000_000.00, 16_006_145.84),
    (17, 2039, 308_582_205.55, 302_250_000.00, 6_332_205.55),
    (18, 2040, 0,              0,              0),
]

# ===== 3. Operational cost and third-party income =====
# From original report
op_cost_schedule = [
    5_768_117.42,    # 2023
    6_890_208.16,    # 2024
    6_090_636.00,    # 2025 onward (estimated avg)
]
third_income_schedule = [
    4_285_070.31,    # 2023
    6_229_093.45,    # 2024
    5_060_000.00,    # 2025 onward (estimated avg)
]

# ===== 4. Combine into annual availability payment =====
print(f"\n{'='*100}")
print(f"年度可用性付费测算（修订版）")
print(f"{'='*100}")
print(f"{'年':<6} {'资本金回报':>14} {'融资本息':>14} {'运维成本':>12} {'第三方收入':>12} {'可用性付费':>14}")
print(f"{'-'*6} {'-'*14} {'-'*14} {'-'*12} {'-'*12} {'-'*14}")

total_equity_return = 0
total_bank = 0
total_op = 0
total_income = 0
total_avail = 0

years_data = []

for i in range(n_years):
    yr_label = 2023 + i
    
    # Equity return
    if i == n_years - 1:  # Last year: lumpsum
        equity_return = equity_schedule[i]['payment'] + capital_lumpsum
        equity_note = f"等额{equity_schedule[i]['payment']:,.0f}+一次性{capital_lumpsum:,.0f}"
    else:
        equity_return = equity_schedule[i]['payment']
        equity_note = ""
    
    # Bank
    bank_total = bank_schedule[i][2]
    if bank_total == 0:
        bank_total = 0
    
    # Op cost
    if i < 2:
        op = op_cost_schedule[i]
    else:
        op = op_cost_schedule[2]
    
    # Third income
    if i < 2:
        income = third_income_schedule[i]
    else:
        income = third_income_schedule[2]
    
    avail = equity_return + bank_total + op - income
    
    total_equity_return += equity_return
    total_bank += bank_total
    total_op += op
    total_income += income
    total_avail += avail
    
    years_data.append({
        'year': yr_label,
        'equity': equity_return,
        'bank': bank_total,
        'op': op,
        'income': income,
        'avail': avail,
        'equity_note': equity_note
    })
    
    print(f"{yr_label:<6} {equity_return:>14,.2f} {bank_total:>14,.2f} {op:>12,.2f} {income:>12,.2f} {avail:>14,.2f}")

print(f"{'-'*6} {'-'*14} {'-'*14} {'-'*12} {'-'*12} {'-'*14}")
print(f"{'合计':<6} {total_equity_return:>14,.2f} {total_bank:>14,.2f} {total_op:>12,.2f} {total_income:>12,.2f} {total_avail:>14,.2f}")

print(f"\n汇总:")
print(f"  资本金回报(8000万等额本息): {sum(e['payment'] for e in equity_schedule):,.2f}")
print(f"  资本金一次性支付(第18年):    {capital_lumpsum:,.2f}")
print(f"  资本金回报合计:              {total_equity_return:,.2f}")
print(f"  融资本息合计:                {total_bank:,.2f}")
print(f"  运维成本合计:                {total_op:,.2f}")
print(f"  第三方收入合计:              {total_income:,.2f}")
print(f"  可用性付费总额:              {total_avail:,.2f}")

# ===== 5. Generate docx =====
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

DARK_BLUE = RGBColor(0x0A, 0x1F, 0x3F)
TEAL = RGBColor(0x1A, 0x5C, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xCC, 0x00, 0x00)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = DARK_BLUE
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = TEAL

def add_p(text, bold=False, size=10.5, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0A1F3F"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else '')
            run.font.size = Pt(8)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F2EC"/>')
                cell._element.get_or_add_tcPr().append(shading)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return table

# ===== COVER =====
for _ in range(6):
    doc.add_paragraph()
add_p('巴中市恩阳医养园PPP项目', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
add_p('可用性付费测算结果报告', bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
doc.add_paragraph()
add_p('（修订版 · 按银行还款计划 + 资本金8,000万元等额本息）', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=TEAL)
for _ in range(4):
    doc.add_paragraph()
add_p('四川融策会计师事务所有限公司', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
add_p('二〇二六年七月二十九日', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
doc.add_page_break()

# ===== 测算说明 =====
add_h('一、测算说明', 1)

add_p('本报告根据巴中市恩阳区卫生健康局委托，对恩阳医养园PPP项目2022年10月28日至2040年12月21日可用性付费进行测算。测算参数依据PPP合同体系及委托方最新确认的还款方案。', size=10.5)

add_h('1.1 关键参数', 2)
params = [
    ['参数', '取值', '依据'],
    ['建安工程费用（结算审定）', '485,098,077.79元', '中宏审〔2024〕1213号竣工结算审核报告'],
    ['资本金总额（实际投入）', '102,298,077.79元', '项目公司实际投入'],
    ['其中：按7.99%等额本息', '80,000,000.00元（8,000万元）', '委托方确认方案'],
    ['其中：第18年一次性支付', '22,298,077.79元', '委托方确认方案'],
    ['资本金合理利润率', '7.99%', '2017补充合同（中标年回报率）'],
    ['资本金回报方式', '等额本息法，18年', '2017补充合同公式'],
    ['银行融资本息', '按贵阳银行实际还款计划', '借款还本付息情况报告'],
    ['银行融资总额', '382,800,000.00元', '贵阳银行5笔贷款合同'],
    ['融资利率（实际执行）', '6.95%→6.45%→5.5%→5.0%', '银行利率调整通知'],
    ['融资利率（合同上限）', '7.99%', '2017补充合同第一条'],
    ['运营维护成本', '前2年实际+后续年均564万元估算', '项目公司财务资料'],
    ['第三方收入', '前2年实际+后续年均506万元估算', '项目公司财务资料'],
    ['运营期', '18年（2022.10.28-2040.12.21）', '2021补充合同'],
    ['绩效考核', '四年均≥93分，支付比例100%', '卫健局考核结果'],
]
add_table(params[0], params[1:], [4, 6, 5.5])

add_h('1.2 计算公式', 2)
add_p('年度可用性付费 = 资本金回报（等额本息） + 银行融资本息（据实） + 运营维护成本 - 第三方收入', bold=True)
add_p('')
add_p('其中资本金回报（8,000万元部分）按等额本息公式：')
add_p('  A = P × k × (1+k)^n / [(1+k)^n - 1]', size=10)
add_p('  P = 80,000,000,  k = 7.99%,  n = 18')
add_p(f'  每年等额还款 = {annual_payment:,.2f}元')
add_p('')
add_p('剩余22,298,077.79元在运营期第18年（2040年度）一次性支付。')

# ===== 资本金回报测算 =====
add_h('二、资本金回报测算（8,000万元等额本息）', 1)

equity_rows = []
for e in equity_schedule:
    equity_rows.append([
        f"第{e['year']}年 (2022+{e['year']})",
        f"{e['payment']:,.2f}",
        f"{e['principal']:,.2f}",
        f"{e['interest']:,.2f}",
        f"{e['remaining']:,.2f}"
    ])
equity_rows.append([
    '合计',
    f"{sum(e['payment'] for e in equity_schedule):,.2f}",
    f"{sum(e['principal'] for e in equity_schedule):,.2f}",
    f"{sum(e['interest'] for e in equity_schedule):,.2f}",
    '0'
])

add_table(
    ['年度', '年还款额（元）', '归还本金（元）', '支付利息（元）', '剩余本金（元）'],
    equity_rows,
    [3, 3, 3, 3, 3]
)

add_p(f'8,000万元等额本息18年合计：{sum(e["payment"] for e in equity_schedule):,.2f}元', bold=True)
add_p(f'加第18年一次性支付：{capital_lumpsum:,.2f}元', bold=True)
add_p(f'资本金回报总计：{sum(e["payment"] for e in equity_schedule) + capital_lumpsum:,.2f}元', bold=True, color=DARK_BLUE)

# ===== 银行还款 =====
add_h('三、银行融资本息（贵阳银行实际还款计划）', 1)

bank_rows = []
for b in bank_schedule:
    yr, yr_label, total_pay, principal, interest = b
    bank_rows.append([
        f"第{yr}年（{yr_label}年度）",
        f"{total_pay:,.2f}",
        f"{principal:,.2f}",
        f"{interest:,.2f}"
    ])
bank_total_sum = sum(b[2] for b in bank_schedule)
bank_principal_sum = sum(b[3] for b in bank_schedule)
bank_interest_sum = sum(b[4] for b in bank_schedule)
bank_rows.append(['合计', f'{bank_total_sum:,.2f}', f'{bank_principal_sum:,.2f}', f'{bank_interest_sum:,.2f}'])

add_table(
    ['年度', '还款总额（元）', '归还本金（元）', '支付利息（元）'],
    bank_rows,
    [4, 3.5, 3.5, 3.5]
)

add_p(f'银行融资本息18年合计：{bank_total_sum:,.2f}元', bold=True)
add_p(f'其中本金：{bank_principal_sum:,.2f}元，利息：{bank_interest_sum:,.2f}元')

# ===== 可用性付费汇总 =====
add_h('四、18年可用性付费测算汇总', 1)

avail_rows = []
for d in years_data:
    avail_rows.append([
        str(d['year']),
        f"{d['equity']:,.2f}",
        f"{d['bank']:,.2f}",
        f"{d['op']:,.2f}",
        f"{d['income']:,.2f}",
        f"{d['avail']:,.2f}",
    ])
avail_rows.append([
    '合计',
    f"{total_equity_return:,.2f}",
    f"{total_bank:,.2f}",
    f"{total_op:,.2f}",
    f"{total_income:,.2f}",
    f"{total_avail:,.2f}",
])

add_table(
    ['年度', '资本金回报（元）', '融资本息（元）', '运维成本（元）', '第三方收入（元）', '可用性付费（元）'],
    avail_rows,
    [1.5, 2.8, 2.8, 2.5, 2.5, 2.8]
)

add_p('')
add_p(f'18年可用性付费总额：{total_avail:,.2f}元（约{total_avail/100000000:.2f}亿元）', bold=True, size=12, color=DARK_BLUE)
add_p(f'年均可用性付费：约{total_avail/18:,.0f}元', bold=True)

# ===== 构成分析 =====
add_h('五、可用性付费构成分析', 1)

comp_rows = [
    ['构成项目', '金额（元）', '占比'],
    ['资本金回报（8000万等额本息）', f"{sum(e['payment'] for e in equity_schedule):,.2f}", f"{sum(e['payment'] for e in equity_schedule)/total_avail*100:.1f}%"],
    ['资本金一次性支付（第18年）', f"{capital_lumpsum:,.2f}", f"{capital_lumpsum/total_avail*100:.1f}%"],
    ['银行融资本息', f"{total_bank:,.2f}", f"{total_bank/total_avail*100:.1f}%"],
    ['运营维护成本', f"{total_op:,.2f}", f"{total_op/total_avail*100:.1f}%"],
    ['减：第三方收入', f"-{total_income:,.2f}", f"-{total_income/total_avail*100:.1f}%"],
    ['可用性付费总额', f"{total_avail:,.2f}", '100.0%'],
]
add_table(comp_rows[0], comp_rows[1:], [5, 4, 2.5])

# ===== 与征求意见稿对比 =====
add_h('六、与征求意见稿（川融策咨询〔2025〕第490号）对比', 1)

compare_rows = [
    ['项目', '征求意见稿', '本修订版', '差异', '说明'],
    ['资本金基数', '102,298,077.79', '102,298,077.79', '0', '一致'],
    ['资本金回报方式', '82,298,077.79全部\n等额本息', '8,000万等额本息\n+2,229.8万第18年一次支付', '方式调整', '委托方最新确认'],
    ['其中：恩阳医院出资', '2,000万（只还本）', '并入一次性支付', '—', '已整合'],
    ['资本金回报合计', '177,955,158.81', f'{total_equity_return:,.2f}', f'{total_equity_return-177955158.81:,.2f}', ''],
    ['融资本息合计', '712,464,358.22', f'{total_bank:,.2f}', f'{total_bank-712464358.22:,.2f}', '按银行还款计划'],
    ['可用性付费总额', '909,053,854.85', f'{total_avail:,.2f}', f'{total_avail-909053854.85:,.2f}', ''],
]
add_table(compare_rows[0], compare_rows[1:], [3, 3.5, 3.5, 2.5, 3])

# ===== 结论 =====
add_h('七、结论', 1)

add_p('本修订版测算根据委托方最新确认的还款方案，采用以下参数：', size=10.5)
add_p('1. 资本金8,000万元按7.99%年利率、18年等额本息方式计算年度回报；', size=10.5)
add_p('2. 剩余资本金22,298,077.79元在运营期第18年一次性支付；', size=10.5)
add_p('3. 银行融资本息按贵阳银行实际还款计划（借款还本付息情况报告）执行。', size=10.5)
add_p('')
add_p(f'经测算，恩阳医养园PPP项目18年运营期（2022年10月-2040年12月）可用性付费总额为：', size=11)
add_p(f'{total_avail:,.2f}元', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
add_p(f'（约{total_avail/100000000:.2f}亿元）', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=TEAL)
add_p('')
add_p('以上测算结果供委托方参考使用。', size=10.5)

doc.add_paragraph()
doc.add_paragraph()
add_p('四川融策会计师事务所有限公司', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, color=DARK_BLUE)
add_p('二〇二六年七月二十九日', size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, color=DARK_BLUE)

# ===== Save =====
output = r"C:\Users\scrccpa\Desktop\恩阳医养园PPP项目可用性付费测算报告（修订版）.docx"
doc.save(output)
import os
print(f"\nSaved: {output}")
print(f"Size: {os.path.getsize(output):,} bytes")
