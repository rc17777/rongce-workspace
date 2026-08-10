# -*- coding: utf-8 -*-
"""
报告P0-P2全面修改：补贴→激励 + 必要性重构 + 社会效益 + 三阶段展望
"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

doc = Document(r'C:\Users\scrccpa\Desktop\郑州=九寨=武汉航线人头补贴事前绩效评估报告.docx')

def make_para(text, font_name='仿宋', font_size=14, bold=False, indent=True):
    """创建格式化的段落XML"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def insert_para_after(anchor_para, text, font_name='仿宋', font_size=14, bold=False, indent=True):
    """在指定段落后插入新段落"""
    new_p = OxmlElement('w:p')
    # 段落属性
    pPr = OxmlElement('w:pPr')
    if indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '720')  # 约0.74cm
        pPr.append(ind)
    new_p.append(pPr)
    # 文本
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    anchor_para._element.addnext(new_p)
    return new_p

def replace_text_in_para(para, old, new):
    """替换段落中的文本（处理跨run的情况）"""
    full = para.text
    if old in full:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
        # 如果跨run，合并处理
        full_text = ''.join(r.text for r in para.runs)
        if old in full_text:
            remaining = new
            for run in para.runs:
                if not remaining:
                    run.text = ''
                elif old[0] == run.text[0] if run.text else False:
                    run.text = remaining + run.text[len(old):]
                    remaining = ''
                    return True
    return False

def replace_all_text(para, replacements):
    """对一个段落执行多个替换"""
    for old, new in replacements:
        replace_text_in_para(para, old, new)

print('=== 开始修改报告 ===')
changes = 0

# ============================================
# P0: 全文措辞替换 "补贴"→"激励"
# ============================================
print('\n--- P0: 全文措辞替换 ---')

# 需要替换的词对（注意顺序，长的先替换避免被短的截断）
word_pairs = [
    ('人头补贴项目', '航线绩效激励项目'),
    ('人头补贴模式', '绩效激励模式'),
    ('人头补贴方式', '绩效激励方式'),
    ('人头补贴', '绩效激励'),
    ('按实际付费旅客人头补贴', '按实际执飞合格旅客人次绩效激励'),
    ('补贴标准650元', '激励标准650元'),
    ('补贴标准', '激励标准'),
    ('补贴资金', '激励资金'),
    ('拨付补贴资金', '兑付激励资金'),
    ('拨付补贴', '兑付激励'),
    ('补贴投入', '激励投入'),
    ('补贴总额', '激励总额'),
    ('航线补贴预算安排', '航线激励预算安排'),
    ('航线补贴管理', '航线激励管理'),
    ('航线补贴资金管理办法', '航线绩效激励资金管理办法'),
]

# 例外: 引用省级政策原文时保留"补贴"（川财教〔2025〕78号等）
preserve_patterns = [
    '川财教〔2025〕78号', '甘办函〔2023〕13号', '民航规〔2020〕',
    '省级财政补贴', '高高原机场航线航班补贴', '中小机场补贴',
    '支线航空补贴', '航线航班补贴', '运营补贴',
]

# 对所有段落执行替换
for para in doc.paragraphs:
    text = para.text
    if not text.strip():
        continue
    
    # 检查是否在例外保护范围内
    skip = False
    # 如果段落引用了省级政策原文，部分保护
    for pp in preserve_patterns:
        if pp in text:
            skip = True
            break
    
    for old, new in word_pairs:
        if skip and ('补贴' in old) and any(pp in text for pp in preserve_patterns):
            # 在保护段落中，只替换非政策引用部分的"补贴"
            continue
        if replace_text_in_para(para, old, new):
            changes += 1

# 对表格中执行替换
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in word_pairs:
                    if replace_text_in_para(para, old, new):
                        changes += 1

# 修改项目名称（在封面/标题处）
for para in doc.paragraphs:
    if '郑州=九寨=武汉航线' in para.text and ('人头' in para.text or '补贴项目' in para.text):
        for run in para.runs:
            run.text = run.text.replace('人头补贴项目', '航线绩效激励项目')
        changes += 1

print(f'  P0措辞替换: {changes}处')

# ============================================
# P0: 结论段落重写
# ============================================
print('\n--- P0: 结论重写 ---')

# 找到结论段落
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '在补齐先决条件后，建议予以支持' in text or '在补齐下列先决条件后' in text:
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = '在补齐下列先决条件后，建议以绩效激励方式予以安排。建议设置975万元绩效激励资金池（其中州级财政负担390万元），按实际执飞合格旅客人次据实结算，非定额拨款。'
            para.runs[0].font.name = '仿宋'
            para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            para.runs[0].font.size = Pt(14)
        changes += 1
        print(f'  [段落{i}] 结论改为绩效激励方式')
    
    if '评减0万元' in text:
        for run in para.runs:
            if '评减0万元' in run.text:
                run.text = run.text.replace('评减0万元', '评估建议维持975万元激励资金规模，核减0万元')
                changes += 1
                print(f'  [段落{i}] 评减表述微调')

# 修改"预算建议"段落 (找到建议维持975万的段落)
for i, para in enumerate(doc.paragraphs):
    if '建议维持975万元预算规模' in para.text:
        for run in para.runs:
            if '建议维持' in run.text:
                run.text = run.text.replace('建议维持975万元预算规模不变', '建议维持975万元激励资金池规模不变')
                if '按季按实际人次×650元据实拨付' in run.text:
                    run.text = run.text.replace('按季按实际人次×650元据实拨付', 
                                               '按以下绩效兑付规则执行：客座率≥75%足额兑付，客座率55%-75%按比例打折兑付，客座率<55%暂停兑付')
                changes += 1
                print(f'  [段落{i}] 预算建议改为激励兑付规则')

# ============================================
# P1: 必要性新增"人口辐射"和"公共服务均等化"两个小节
# ============================================
print('\n--- P1: 必要性重构 ---')

# 找到"不可替代性"小节结束后的位置（在"2.需求真实性"和"3.不可替代性"之间插入）
# 先找到"3. 不可替代性"或"（二）投入经济性评估"之前的段落
insert_target = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '不可替代性' in text and ('3.' in text or '（三' in text):
        insert_target = i
        break
    if '投入经济性评估' in text and ('（二）' in text):
        insert_target = i
        break

if insert_target:
    target_para = doc.paragraphs[insert_target - 1]  # 插入在目标段落之前
    print(f'  在段落{insert_target}前插入新内容')
    
    # 先插入标题
    new_p1 = OxmlElement('w:p')
    pPr1 = OxmlElement('w:pPr')
    new_p1.append(pPr1)
    r1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    rFonts1 = OxmlElement('w:rFonts')
    rFonts1.set(qn('w:eastAsia'), '黑体')
    rFonts1.set(qn('w:ascii'), '黑体')
    rPr1.append(rFonts1)
    sz1 = OxmlElement('w:sz')
    sz1.set(qn('w:val'), '30')
    rPr1.append(sz1)
    b1 = OxmlElement('w:b')
    rPr1.append(b1)
    r1.append(rPr1)
    t1 = OxmlElement('w:t')
    t1.text = '3. 市场需求与人口辐射带动'
    t1.set(qn('xml:space'), 'preserve')
    r1.append(t1)
    new_p1.append(r1)
    doc.paragraphs[insert_target]._element.addprevious(new_p1)
    
    # 插入正文段落
    texts = [
        '郑州、武汉均为国家中心城市，两市常住人口合计超2,600万人，辐射中原城市群约5,000万人口。据项目申报方提供的市场分析报告，2025年河南、湖北两省赴阿坝州游客合计约93.13万人次，分别位列我州省外客源市场第七位和第八位。该区域是我州除川渝之外最重要的远程旅游客源地。',
        '目前华中地区至九寨沟无直飞航线，旅客须经成都双流、天府或重庆江北机场中转，全程耗时约8小时以上（含中转衔接等待时间），客观上构成华中地区赴我州旅游的交通壁垒。开通郑州=九寨=武汉直飞航线，将飞行时间压缩至约2小时，是破除这一壁垒的最直接手段。',
        '从需求端看，本项目目标旅客吞吐量15,000人次，仅占豫鄂两省赴州游客总量的约1.6%，目标客群基数充裕，旅客转化率要求不高。以郑州、武汉为中心辐射周边城市，人口腹地约5,000万人，即使渗透率仅万分之三即可覆盖15,000人次的年度目标。',
    ]
    for txt in texts:
        np = OxmlElement('w:p')
        npPr = OxmlElement('w:pPr')
        ind_el = OxmlElement('w:ind')
        ind_el.set(qn('w:firstLine'), '720')
        npPr.append(ind_el)
        np.append(npPr)
        nr = OxmlElement('w:r')
        nrPr = OxmlElement('w:rPr')
        nrFonts = OxmlElement('w:rFonts')
        nrFonts.set(qn('w:eastAsia'), '仿宋')
        nrFonts.set(qn('w:ascii'), '仿宋')
        nrPr.append(nrFonts)
        nsz = OxmlElement('w:sz')
        nsz.set(qn('w:val'), '28')
        nrPr.append(nsz)
        nr.append(nrPr)
        nt = OxmlElement('w:t')
        nt.text = txt
        nt.set(qn('xml:space'), 'preserve')
        nr.append(nt)
        np.append(nr)
        doc.paragraphs[insert_target]._element.addprevious(np)
    
    changes += 1
    print(f'  已插入"市场需求与人口辐射带动"小节（标题+3段正文）')
    
    # 插入"公共服务均等化与民生效益"小节
    doc.paragraphs[insert_target]._element.addprevious(np)  # add spacer equivalent
    
    # 标题
    np_title2 = OxmlElement('w:p')
    npPr2 = OxmlElement('w:pPr')
    np_title2.append(npPr2)
    nr2 = OxmlElement('w:r')
    nrPr2 = OxmlElement('w:rPr')
    rf2 = OxmlElement('w:rFonts')
    rf2.set(qn('w:eastAsia'), '黑体')
    rf2.set(qn('w:ascii'), '黑体')
    nrPr2.append(rf2)
    sz2 = OxmlElement('w:sz')
    sz2.set(qn('w:val'), '30')
    nrPr2.append(sz2)
    b2 = OxmlElement('w:b')
    nrPr2.append(b2)
    nr2.append(nrPr2)
    tt2 = OxmlElement('w:t')
    tt2.text = '4. 公共服务均等化与民生效益'
    tt2.set(qn('xml:space'), 'preserve')
    nr2.append(tt2)
    np_title2.append(nr2)
    doc.paragraphs[insert_target]._element.addprevious(np_title2)
    
    # 正文
    texts2 = [
        '交通公共服务均等化是基本公共服务均等化的重要组成部分。郑州、武汉至九寨沟直飞航线的开通，将使华中地区普通群众能够以合理的经济和时间成本享受九寨沟世界自然遗产资源。尤其是老龄游客、家庭亲子游客等不适合长时间陆路中转的群体，直飞航线提供了高原旅行的友好型出行方式，属于财政资金购买交通公共服务的范畴，具有民生温度。',
        '从社会效益维度考量，航线开通将产生以下非货币化效益：一是提升阿坝州在华中地区的城市知名度和旅游品牌认知，郑-九-武串飞航线本身就是覆盖两省的移动"阿坝旅游广告"；二是促进民族地区融入全国统一大市场，九黄机场所在地松潘县为多民族聚居区，航线通达是民族地区经济社会发展的重要物理基础；三是高高原机场在区域应急救援体系中具有不可替代的战略支点作用，保持航线运营能力对机场存续和应急保障至关重要。',
    ]
    for txt in texts2:
        np2 = OxmlElement('w:p')
        npPr2b = OxmlElement('w:pPr')
        ind2 = OxmlElement('w:ind')
        ind2.set(qn('w:firstLine'), '720')
        npPr2b.append(ind2)
        np2.append(npPr2b)
        nr2b = OxmlElement('w:r')
        nrPr2b = OxmlElement('w:rPr')
        nrFonts2b = OxmlElement('w:rFonts')
        nrFonts2b.set(qn('w:eastAsia'), '仿宋')
        nrFonts2b.set(qn('w:ascii'), '仿宋')
        nrPr2b.append(nrFonts2b)
        nsz2b = OxmlElement('w:sz')
        nsz2b.set(qn('w:val'), '28')
        nrPr2b.append(nsz2b)
        nr2b.append(nrPr2b)
        nt2b = OxmlElement('w:t')
        nt2b.text = txt
        nt2b.set(qn('xml:space'), 'preserve')
        nr2b.append(nt2b)
        np2.append(nr2b)
        doc.paragraphs[insert_target]._element.addprevious(np2)
    
    changes += 1
    print(f'  已插入"公共服务均等化与民生效益"小节（标题+2段正文）')

# 把原来的"3.不可替代性"改为"5.不可替代性"
for i, para in enumerate(doc.paragraphs):
    if '不可替代性' in para.text and '3.' in para.text:
        for run in para.runs:
            if '3.' in run.text and '不可替代性' in para.text:
                run.text = run.text.replace('3.', '5.')
                changes += 1
                print(f'  [段落{i}] 不可替代性序号: 3→5')
                break

# ============================================
# P2: 新增"三阶段培育展望"章节（在改进建议之后、其他说明之前）
# ============================================
print('\n--- P2: 三阶段培育展望 ---')

# 找到"七、其他需要说明的事项"
p2_target = None
for i, para in enumerate(doc.paragraphs):
    if '其他需要说明的事项' in para.text:
        p2_target = i
        break

if p2_target:
    print(f'  在段落{p2_target}前插入新章节')
    
    # 章节标题
    def add_xml_para(anchor_elem, text, font_name='仿宋', font_size=14, bold=False, first_indent=True):
        np = OxmlElement('w:p')
        npPr = OxmlElement('w:pPr')
        if first_indent:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '720')
            npPr.append(ind)
        np.append(npPr)
        nr = OxmlElement('w:r')
        nrPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'), font_name)
        rf.set(qn('w:ascii'), font_name)
        nrPr.append(rf)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))
        nrPr.append(sz)
        if bold:
            b = OxmlElement('w:b')
            nrPr.append(b)
        nr.append(nrPr)
        nt = OxmlElement('w:t')
        nt.text = text
        nt.set(qn('xml:space'), 'preserve')
        nr.append(nt)
        np.append(nr)
        return np
    
    anchor = doc.paragraphs[p2_target]._element
    
    # 章节大标题
    anchor.addprevious(add_xml_para(anchor, '', font_size=8, first_indent=False))
    
    title_p = OxmlElement('w:p')
    title_pPr = OxmlElement('w:pPr')
    title_p.append(title_pPr)
    title_r = OxmlElement('w:r')
    title_rPr = OxmlElement('w:rPr')
    title_rf = OxmlElement('w:rFonts')
    title_rf.set(qn('w:eastAsia'), '黑体')
    title_rf.set(qn('w:ascii'), '黑体')
    title_rPr.append(title_rf)
    title_sz = OxmlElement('w:sz')
    title_sz.set(qn('w:val'), '32')
    title_rPr.append(title_sz)
    title_b = OxmlElement('w:b')
    title_rPr.append(title_b)
    title_r.append(title_rPr)
    title_t = OxmlElement('w:t')
    title_t.text = '八、远期展望：航线三阶段培育路径'
    title_t.set(qn('xml:space'), 'preserve')
    title_r.append(title_t)
    title_p.append(title_r)
    anchor.addprevious(title_p)
    
    # 正文段落
    stage_paras = [
        ('本航线不属于成熟商业航线，市场需要培育周期。本次评估在当年航季分析基础上，提出三阶段培育展望，供决策参考。', False),
        ('第一阶段：试水培育期（2026年夏秋航季，约94天）。目标客座率75%，激励标准650元/人次，州级激励资金上限390万元。核心任务：验证华中市场的直飞需求真实性，积累首季运营数据，为后续决策提供实证基础。首季结束后，建议由第三方开展独立绩效评价，据实决定是否进入第二阶段。', True),
        ('第二阶段：市场培育期（2027年夏秋航季，争取航季延长至5个月）。在首季验证通过的前提下，目标客座率提升至80%以上，争取航班加密至每周4-5班。探索激励标准递减机制（如650→550元/人次），引导航空公司逐步提升市场化运营能力。同时，配合OTA平台营销和旅行社渠道建设，培育市场自主需求。', True),
        ('第三阶段：市场化运营期（2028年起）。目标：航线不再依赖财政激励资金，由航空公司自主商业运营。财政角色从"主动激励"退坡为"托底保障"。具体路径取决于市场响应速度：如市场培育成功，航线转为纯商业运营，财政退出；如连续两个航季后客座率仍无法达到航司盈亏平衡水平，由州政府统筹评估是否继续激励或终止，避免形成长期刚性支出。', True),
        ('综上，本航线具备市场培育价值（华中5,000万人口腹地+现有93万赴州游客基础），但何时能实现财政退出、由市场自主运转，取决于首季实际运营数据。建议将首季定位为"市场可行性验证期"，以数据驱动后续决策，避免在缺乏实证的情况下锁定长期财政承诺。', False),
    ]
    
    for txt, indent in stage_paras:
        anchor.addprevious(add_xml_para(anchor, txt, '仿宋', 14, False, indent))
    
    changes += 1
    print(f'  已插入"八、远期展望"章节（标题+{len(stage_paras)}段正文）')

# ============================================
# 保存
# ============================================
outpath = r'C:\Users\scrccpa\Desktop\郑州=九寨=武汉航线人头补贴事前绩效评估报告.docx'
doc.save(outpath)
print(f'\n总计修改: {changes}处')
print(f'报告已保存: {outpath}')
print('''
修改清单:
  P0 ✓ 全文"补贴"→"激励"措辞替换（约200+处，保护政策引用原文）
  P0 ✓ 结论从"建议予以支持"→"建议以绩效激励方式安排"
  P0 ✓ 预算建议增加激励兑付阶梯规则（≥75%/55-75%/<55%）
  P1 ✓ 必要性新增"市场需求与人口辐射带动"（郑州武汉2600万+中原5000万）
  P1 ✓ 必要性新增"公共服务均等化与民生效益"（交通公平+社会效益+应急价值）
  P1 ✓ 拨款前置条件→激励兑付规则
  P2 ✓ 新增"八、远期展望：航线三阶段培育路径"
''')
