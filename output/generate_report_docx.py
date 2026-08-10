# -*- coding: utf-8 -*-
"""
郑州=九寨=武汉航线人头补贴项目 事前绩效评估报告 Word生成
融策右护卫 | 2026-07-25
融合三模型评审意见（qwen公文/sonnet审计/fable幕僚）
"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

import matplotlib
matplotlib.use('Agg')
matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
matplotlib.rcParams['axes.unicode_minus'] = False
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

OUTDIR = r'C:\Users\scrccpa\Desktop'
CHARTS_DIR = r'C:\Users\scrccpa\.openclaw\workspace\temp\charts'
os.makedirs(CHARTS_DIR, exist_ok=True)

# ============================================================
# 融策品牌色
# ============================================================
DEEP_BLUE = '#0A1F3F'
TEAL = '#1A5C6E'
GOLD = '#C5955C'
WARM_GRAY = '#F5F2EC'
LIGHT_BLUE = '#2E86AB'
SOFT_RED = '#C0392B'
ORANGE = '#E67E22'
GREEN = '#27AE60'

# ============================================================
# 图表1: 多情景效益对比图
# ============================================================
def chart_multi_scenario():
    scenarios = ['55%\n(保守)', '65%\n(基准)', '75%\n(目标)', '85%\n(乐观)']
    passengers = [11220, 13260, 15000, 17340]
    subsidy = [729.3, 861.9, 975, 1127.1]
    consumption = [3927, 4641, 5250, 6069]

    fig, ax1 = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(scenarios))
    width = 0.32

    bars1 = ax1.bar(x - width, subsidy, width, color=DEEP_BLUE, edgecolor='white', linewidth=0.5)
    bars2 = ax1.bar(x + width, consumption, width, color=GOLD, edgecolor='white', linewidth=0.5)
    ax2 = ax1.twinx()
    line = ax2.plot(x, passengers, 'o-', color=TEAL, linewidth=2.5, markersize=8, markerfacecolor=TEAL)

    ax1.set_xticks(x)
    ax1.set_xticklabels(scenarios, fontsize=10)
    ax1.set_ylabel('金额（万元）', fontsize=10, color=DEEP_BLUE)
    ax2.set_ylabel('旅客量（人次）', fontsize=10, color=TEAL)
    ax1.set_title('多情景补贴-消费-旅客量对比', fontsize=13, fontweight='bold', color=DEEP_BLUE, pad=15)

    for bar, val in zip(bars1, subsidy):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 30, f'{val:.0f}万',
                ha='center', va='bottom', fontsize=8, color=DEEP_BLUE, fontweight='bold')
    for bar, val in zip(bars2, consumption):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 30, f'{val:.0f}万',
                ha='center', va='bottom', fontsize=8, color=GOLD, fontweight='bold')

    ax1.legend([bars1, bars2], ['补贴投入', '消费拉动'], loc='upper left', fontsize=9, frameon=False)
    ax2.legend(['旅客量'], loc='upper right', fontsize=9, frameon=False)
    ax1.set_ylim(0, 7500)
    ax1.grid(axis='y', alpha=0.2, color=DEEP_BLUE)
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'chart_scenarios.png')
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ============================================================
# 图表2: 运营成本构成饼图
# ============================================================
def chart_cost_pie():
    labels = ['航油\n51.0%', '折旧\n23.7%', '人工\n6.4%', '维修\n5.4%', '起降\n3.4%', '其他\n10.1%']
    sizes = [51.0, 23.7, 6.4, 5.4, 3.4, 10.1]
    colors = [DEEP_BLUE, TEAL, GOLD, '#7FB3D8', '#AAB7B8', '#D5DBDB']
    explode = (0.05, 0, 0, 0, 0, 0)

    fig, ax = plt.subplots(figsize=(6, 4.5))
    wedges, texts, autotexts = ax.pie(sizes, explode=explode, labels=labels, colors=colors,
                                       autopct='', startangle=90, pctdistance=0.6,
                                       textprops={'fontsize': 9})
    for i, (w, s) in enumerate(zip(wedges, sizes)):
        ang = (w.theta2 - w.theta1) / 2 + w.theta1
        x = np.cos(np.deg2rad(ang))
        y = np.sin(np.deg2rad(ang))
        ax.annotate(f'{s}%', xy=(x*0.73, y*0.73), ha='center', va='center',
                    fontsize=9, fontweight='bold', color='white' if i < 2 else DEEP_BLUE)

    ax.set_title('航线单班运营成本构成（武汉=九寨，加权）', fontsize=12, fontweight='bold', color=DEEP_BLUE, pad=15)
    path = os.path.join(CHARTS_DIR, 'chart_cost_pie.png')
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ============================================================
# 图表3: 历史人均补贴对标
# ============================================================
def chart_historical_benchmark():
    routes = ['本项目\n650', '北京=九寨\n522', '杭州=九寨\n1,299', '武汉=九寨(23)\n2,273', '成都=九寨(23)\n667']
    values = [650, 522, 1299, 2273, 667]
    colors = [GOLD, DEEP_BLUE, DEEP_BLUE, SOFT_RED, DEEP_BLUE]

    fig, ax = plt.subplots(figsize=(8, 3.8))
    bars = ax.barh(routes, values, color=colors, edgecolor='white', linewidth=0.8, height=0.6)

    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 30, bar.get_y() + bar.get_height()/2, f'{val}元/人',
                va='center', fontsize=10, fontweight='bold', color=DEEP_BLUE)

    ax.set_xlabel('人均补贴（元/人次）', fontsize=10, color=DEEP_BLUE)
    ax.set_title('省内各航线人均补贴对标（2023年数据）', fontsize=12, fontweight='bold', color=DEEP_BLUE, pad=12)
    ax.set_xlim(0, 2800)
    ax.axvline(x=650, color=GOLD, linestyle='--', linewidth=1.5, alpha=0.7)
    ax.text(650, 4.6, '本项目', ha='center', fontsize=8, color=GOLD, fontweight='bold')
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)
    ax.grid(axis='x', alpha=0.15, color=DEEP_BLUE)
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'chart_benchmark.png')
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ============================================================
# 图表4: 航司盈亏敏感性分析
# ============================================================
def chart_breakeven():
    fig, ax = plt.subplots(figsize=(8, 3.8))
    fares = np.arange(400, 1000, 10)
    cost_per_pax = 1252
    subsidy_per_pax = 650
    profit = fares + subsidy_per_pax - cost_per_pax

    ax.fill_between(fares, 0, profit, where=(profit >= 0), color=GREEN, alpha=0.3, label='盈利区间')
    ax.fill_between(fares, profit, 0, where=(profit < 0), color=SOFT_RED, alpha=0.3, label='亏损区间')
    ax.plot(fares, profit, color=DEEP_BLUE, linewidth=2.5)
    ax.axhline(y=0, color=GOLD, linestyle='--', linewidth=1.5, alpha=0.7)

    be_fare = cost_per_pax - subsidy_per_pax
    ax.axvline(x=be_fare, color=GOLD, linestyle=':', linewidth=1.5, alpha=0.8)
    ax.annotate(f'盈亏平衡\n票价={be_fare:.0f}元', xy=(be_fare, 0), xytext=(be_fare+120, 150),
                arrowprops=dict(arrowstyle='->', color=DEEP_BLUE, lw=1.5),
                fontsize=10, color=DEEP_BLUE, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.3', facecolor=WARM_GRAY, edgecolor=DEEP_BLUE, alpha=0.8))

    ax.set_xlabel('平均票价（元）', fontsize=10, color=DEEP_BLUE)
    ax.set_ylabel('航司单客利润（元）', fontsize=10, color=DEEP_BLUE)
    ax.set_title('航司盈亏平衡分析（含650元/人补贴）', fontsize=12, fontweight='bold', color=DEEP_BLUE, pad=12)
    ax.legend(loc='upper left', fontsize=9, frameon=False)
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)
    ax.grid(alpha=0.15, color=DEEP_BLUE)
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'chart_breakeven.png')
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ============================================================
# 生成所有图表
# ============================================================
print('生成图表...')
chart_paths = {
    'scenarios': chart_multi_scenario(),
    'cost_pie': chart_cost_pie(),
    'benchmark': chart_historical_benchmark(),
    'breakeven': chart_breakeven(),
}
print(f'  已生成 {len(chart_paths)} 张图表')

# ============================================================
# Word文档生成
# ============================================================
print('生成Word报告...')
doc = Document()

# -- 页面设置 --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
style.paragraph_format.line_spacing = 1.5

# ============================================================
# 辅助函数
# ============================================================
def add_title(text, level=0):
    """level: 0=报告标题, 1=部分标题, 2=节标题, 3=目标题"""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '方正小标宋简体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        p.paragraph_format.space_after = Pt(20)
    elif level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.name = '楷体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '仿宋'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body(text, indent=True, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.name = '仿宋'
        run.font.size = Pt(14)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run = p.add_run(text)
    run.font.name = '仿宋'
    run.font.size = Pt(14)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    return p

def add_image(img_path, width=Inches(5.5), caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    if caption:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(caption)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2.font.name = '仿宋'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{DEEP_BLUE.replace("#","")}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '仿宋'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F2EC"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table

# ============================================================
# 封面信息
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_title('郑州=九寨=武汉航线人头补贴项目')
add_title('事前绩效评估报告')

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('委托单位：', '阿坝藏族羌族自治州财政局'),
    ('评估机构：', '四川融策会计师事务所有限公司'),
    ('报告文号：', '融策绩评〔2026〕第  号'),
    ('评估基准日：', '2026年7月25日'),
    ('报告密级：', '内部·呈州政府决策参考'),
]
for label, val in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(14)
    run.font.name = '仿宋'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run = p.add_run(val)
    run.font.size = Pt(14)
    run.font.name = '仿宋'
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

doc.add_page_break()

# ============================================================
# 报告摘要（决策要点）
# ============================================================
add_title('报告摘要（决策要点）', level=1)

add_body('阿坝州财政局委托四川融策会计师事务所有限公司（以下简称"我所"），对郑州=九寨=武汉航线人头补贴项目进行事前绩效评估。项目申报补贴总额975万元（申请省级财政585万元、州级财政390万元），以650元/人次人头补贴模式，委托西部航空执飞，计划2026年8月上旬至10月24日执飞150架次，目标旅客吞吐量15,000人次。')

add_body('我所依据《预算法》、《四川省预算绩效评估管理办法》（川财绩〔2025〕12号）等规定，采用成本效益分析法、比较法、因素分析法、最低成本法等方法，对项目立项必要性、投入经济性、绩效目标合理性、实施方案可行性和筹资合规性进行了独立评估。', indent=True)

doc.add_paragraph()
add_body('【评估结论】', indent=False, bold_prefix='')
add_body('在补齐先决条件后，建议予以支持。评估建议维持975万元预算规模。补贴标准650元/人次在成本推导维度上逻辑自洽，投入产出具备较好经济性基础。', indent=True)

doc.add_paragraph()
add_body('【核心数据】', indent=False, bold_prefix='')
add_table(
    ['指标', '数值', '说明'],
    [
        ['州级实际新增负担', '390万元', '975万中省级585万待省厅批复'],
        ['补贴标准', '650元/人次', '基于运营成本1,252元/人推导'],
        ['目标旅客量', '15,000人次', '150架次×约100人/班（75%客座率）'],
        ['理论投入产出比', '约1:5.38', '3500÷650（数学恒等式，需注意消费数据可核验性）'],
        ['航司盈亏平衡票价', '602元', '含补贴后的航司盈亏临界点'],
        ['州级净增占存量比', '26%', '390万÷存量补贴预算1,500万'],
    ],
    col_widths=[4, 3, 8.5]
)

doc.add_paragraph()
add_body('【资金拨付前置条件】', indent=False, bold_prefix='')

conditions = [
    ['□ 条件一', '取得民航西南地区管理局航线审批批复文件'],
    ['□ 条件二', '州文体旅游局、西部航空、九黄机场三方正式签署客运量奖励协议'],
    ['□ 条件三', '补充提供人均旅游消费3,500元/人的原始调研数据文件（注：现有提供的省文旅厅通知文件经信息化提取后，内容与文件名不一致）'],
    ['□ 条件四', '以650元/人次为参数重新提供正确的客座率-票价-补贴测算矩阵（注：现有测算表内置参数为100元/人）'],
]
add_table(
    ['序号', '条件内容'],
    conditions,
    col_widths=[2, 13.5]
)

doc.add_paragraph()
add_body('【主要风险提示】', indent=False, bold_prefix='')
add_body('一是历史运营数据存在内部矛盾：项目申报方可行性报告称"历史客座率均超过80%"，但2023年实际运营数据显示武汉=九寨航线客座率仅36.63%，两组数据来自同一提供方，需请航线开发工作专班统一口径并确认。二是州级航线补贴管理办法尚未出台，专项资金管理制度化框架不完整。三是航线审批及协议签署尚未完成，项目法律关系尚待建立。', indent=True)

doc.add_page_break()

# ============================================================
# 第一部分：评估工作概述
# ============================================================
add_title('一、评估工作概述', level=1)

add_title('（一）评估目的', level=2)
add_body('受阿坝藏族羌族自治州财政局委托，我所对郑州=九寨=武汉航线人头补贴项目开展事前绩效评估，为州政府航线补贴预算安排和项目立项决策提供第三方独立参考。本项目评估基准日为2026年7月25日。')

add_title('（二）评估依据', level=2)
add_body('1.《中华人民共和国预算法》；')
add_body('2.《四川省预算绩效评估管理办法》（川财绩〔2025〕12号）；')
add_body('3.《四川省省级财政高高原机场航线航班补贴资金管理办法》（川财教〔2025〕78号）；')
add_body('4.《四川省省级财政专项资金管理办法》（川府发〔2023〕2号）；')
add_body('5. 关于成立阿坝州航线开发工作专班的通知（2024年3月15日）；')
add_body('6. 中国注册会计师相关执业准则。')

add_title('（三）评估原则', level=2)
add_body('本次评估遵循客观独立、绩效导向、证据驱动、审慎保守原则，秉持厉行节约要求，从严审核项目投入。')

add_title('（四）评估方法', level=2)
add_body('以成本效益分析法为主线，辅以比较法（甘孜州同类航线补贴、九黄机场历史航线、省内各机场补贴规模对照）、因素分析法（客座率、人均消费、航油价格等关键变量敏感性分析）、最低成本法（替代补贴模式比选）和公众评判法（相关方意见征询）。')

add_title('（五）评估程序', level=2)
add_body('我所于2026年7月23日接受委托，向项目申报单位发出72项资料需求清单，先后接收两批共69份文件资料。随后逐一核验文件可读性与完整性，对扫描件采用信息化手段提取数据并交叉验证，与项目单位和州财政局进行沟通确认，在此基础上开展分析评估并形成本报告。')

doc.add_page_break()

# ============================================================
# 第二部分：项目基本情况
# ============================================================
add_title('二、项目基本情况', level=1)

add_title('（一）项目概况', level=2)
add_body('郑州=九寨=武汉航线人头补贴项目，由四川九寨黄龙机场有限责任公司（以下简称"九黄机场公司"）申报，属新增项目。航线规划为郑州→九寨黄龙机场→武汉串飞，采用A319高高原机型，每周3班（周三、五、日），执飞期2026年8月上旬至10月24日（约94天），共150架次，目标旅客吞吐量15,000人次。')

add_body('项目申报补贴总额975万元，其中申请省级财政补贴585万元、州级财政配套390万元。补贴方式为按实际付费旅客人头补贴，标准为650元/人次。项目申报方预期通过航线运营，在目标客座率75%下，拉动旅游消费5,250万元，理论投入产出比约1:5.38。', indent=True)

add_title('（二）项目背景', level=2)
add_body('川青铁路通车后，成都至九寨沟交通时间大幅压缩，九黄机场成都航线客流从2023年13.56万人次骤降至2025年2.17万人次，降幅达84%。在此背景下，机场实施战略转向，从省内短途转向省外远程客源市场。河南、湖北两省常住人口合计超过1.6亿，2025年两省赴阿坝州游客约93.13万人次，系九黄机场最重要的远程客源市场。西部航空曾于2014年至2016年执行郑州=九寨航线，具备执飞基础条件。')

add_title('（三）绩效目标', level=2)
add_body('项目申报方提交了航线人头补贴财政资金绩效评价表，设定了产出指标（执飞架次、旅客吞吐量、补贴覆盖天数等）和效益指标（旅游消费拉动、投入产出比、满意度等）。经审核，该评价表整体框架规范，但存在部分指标以"已完成"时态表述尚未执行的项目目标，建议在项目正式实施前调整为事前预期式表述。')

add_title('（四）资金安排', level=2)

add_table(
    ['资金来源', '金额（万元）', '占比', '拨付方式'],
    [
        ['省级财政补贴（川财教〔2025〕78号）', '585', '60%', '待省厅批复后按季结算'],
        ['州级财政配套', '390', '40%', '州级预算安排'],
        ['合计', '975', '100%', '—'],
    ],
    col_widths=[6, 2.5, 2, 5]
)

doc.add_page_break()

# ============================================================
# 第三部分：评估内容及分析
# ============================================================
add_title('三、评估内容及分析', level=1)

# --- （一）立项必要性 ---
add_title('（一）立项必要性评估', level=2)

add_body('1. 政策符合性', indent=True, bold_prefix='')
add_body('本项目契合国家关于"支持高高原支线机场发展"的交通战略，符合四川省高高原机场航线补贴政策框架（川财教〔2025〕78号）。州政府于2024年3月成立了航线开发工作专班（常务副州长任组长），明确赋予"负责新开航线有关财政支持政策的制定和实施"职能，项目具有组织合法性基础。需指出，州级层面尚未出台航线补贴管理办法及实施细则，建议在项目实施前制定并发布。')

add_body('2. 需求真实性', indent=True, bold_prefix='')
add_body('豫鄂两省1.6亿人口基础为公开统计数据，可独立验证。两省赴州游客数据（约93.13万人次）来源于项目申报方市场分析报告，该数据口径（全州）与景区门票口径（九寨沟景区2023-2025年湖北籍游客共56.34万人次）存在差异，建议统一口径后确认。历史上西部航空确曾执飞郑州=九寨航线（2014-2016年），利害关系声明已确认上述事实。需要关注：项目申报方可行性报告中称"郑州、武汉航线历史平均客座率均超过80%"，但2023年实际运营数据（武汉=九寨、川航执飞）显示客座率仅36.63%，两组数据存在显著差异。建议请航线开发工作专班核实确认，并在此基础上重新审视客流预测假设。')

add_body('3. 不可替代性', indent=True, bold_prefix='')
add_body('川青铁路对省内短途航线冲击效应明显（成都=九寨客流下降84%），验证了"聚焦省外远程航线、避免与高铁正面竞争"策略的合理性。郑州、武汉至九寨沟陆路+航空中转约8小时，直飞约2小时，时间优势显著。目前九黄机场省外直飞航线以北京、杭州为主，华中市场存在空白。项目申报方提供的替代方案比选内容较为简略，建议补充系统的替代方案论证。')

# --- （二）投入经济性 ---
add_title('（二）投入经济性评估', level=2)

add_body('【补贴标准定价依据】', indent=False, bold_prefix='')
add_body('我所从以下维度对650元/人次补贴标准进行了分析：', indent=True)

add_body('第一，成本推导。郑州=九寨往返单班运营成本24.26万元、武汉=九寨往返25.83万元，加权平均单班成本约25.05万元。按200座/往返（100人×2段）计算，单客运营成本约1,252元。扣除票价约650元后，差额约602元。补贴650元覆盖运营差额及营销费用，定价逻辑基本自洽。航线运营成本构成以航油（约占51%）和飞机折旧（约占24%）为主。但需指出，航油单价（10,834元/吨）和小时油耗（2.8吨）等底层参数系项目申报方提供，未经独立第三方核实。', indent=True)

add_body('第二，市场对标。省内各航线2023年人均补贴区间为522元/人次至2,273元/人次，本项目650元处于合理中位区间。但需注意：上述对标数据为存量航线的实际执行人均补贴，系补贴总额除以实际客运量的结果值，与本项目事前定价标准属于不同口径，对标仅作参考。', indent=True)

add_body('第三，政策框架参照。省级政策（川财教〔2025〕78号）按班次设定补贴上限（始发航线每班次≤8万元），按本项目150人/往返（75座×2段）折算约533元/人次（不含省级以下叠加）。甘孜州政策（甘办函〔2023〕13号）同样为按班次补贴模式（稻城≤8万+州级≤8万=≤16万/班次），按100座/班折算约1,600元/人，但该数值系我所折算推导，并非政策原文的直接规定。本项目650元/人次的价格水平在折算后的省级框架上限以内。但阿坝州目前采用"人头补贴"模式，与省级"班次补贴"方法论存在差异。', indent=True)

# 成本饼图
add_image(chart_paths['cost_pie'], width=Inches(4.5), caption='图1：航线单班运营成本构成（武汉=九寨，加权平均）')

add_body('【多情景效益分析】', indent=False, bold_prefix='')
add_body('我所基于四档客座率假设，对项目补贴投入与预期消费拉动进行了多情景测算：', indent=True)

add_table(
    ['情景', '客座率', '旅客量（人次）', '补贴投入（万元）', '消费拉动（万元）'],
    [
        ['保守', '55%', '11,220', '729.3', '3,927'],
        ['基准', '65%', '13,260', '861.9', '4,641'],
        ['目标', '75%', '15,000', '975.0', '5,250'],
        ['乐观', '85%', '17,340', '1,127.1', '6,069'],
    ],
    col_widths=[2, 2, 3, 3, 3]
)

add_body('需要说明：上述"消费拉动"系依据项目申报方调研转述的人均消费3,500元/人计算（3,500×旅客量）。人均消费3,500元的推导依据为省文旅厅统计口径下游客在阿坝州日均消费约1,732.28元、停留约2天。但该数据的一手来源文件（文件名标注为"关于反馈四川省各市州2024年地方接待国内游客数据有关情况的通知"）经信息化提取后，实际内容为关于补贴申报的另一份文件（川文旅发〔2025〕47号）。人均消费3,500元/人作为效益测算基石数据，其可核验性受限于一手来源文件的供给情况，建议在使用该数据时明确标注其来源。', indent=True)

# 多情景图表
add_image(chart_paths['scenarios'], width=Inches(5.5), caption='图2：多情景补贴投入-消费拉动-旅客量对比')

add_body('【航司盈亏分析】', indent=False, bold_prefix='')
add_body('基于运营成本1,252元/人、补贴650元/人，航司盈亏平衡票价为602元/人。项目申报方引流方案设计进港票价500元、出港800元，按55%∶45%进出港比例加权计算约635元/人，略高于盈亏平衡点。', indent=True)

# 盈亏平衡图
add_image(chart_paths['breakeven'], width=Inches(5.5), caption='图3：航司盈亏平衡分析（含650元/人补贴）')

add_body('【历史对标】', indent=False, bold_prefix='')
add_body('2023年武汉=九寨航线（川航执飞）人均补贴2,273元/人次、客座率36.63%。本项目650元/人次为其28.6%，补贴效率理论上有显著改善空间。但需关注：①上述航线为保底补贴模式，与本项目人头补贴模式不同；②36.63%的客座率反映华中至九寨航线需求基础可能低于预期，75%的目标需要较强的市场营销支撑。', indent=True)

# 历史对标图
add_image(chart_paths['benchmark'], width=Inches(5.5), caption='图4：省内各航线人均补贴对标（2023年数据）')

add_body('【敏感性分析】', indent=False, bold_prefix='')
add_body('对三个关键变量进行了敏感性分析：', indent=True)

add_table(
    ['变量', '变动幅度', '对补贴总额影响', '对航司利润影响'],
    [
        ['客座率', '±10个百分点', '±130万元', '±188万元（因客流变化）'],
        ['人均消费', '±500元/人', '不影响补贴总额', '消费拉动±750万元（非财政口径）'],
        ['航油价格', '±20%', '不影响补贴总额', '±188万元（影响航司成本）'],
    ],
    col_widths=[3, 3, 4, 5.5]
)

# --- （三）绩效目标合理性 ---
add_title('（三）绩效目标合理性评估', level=2)
add_body('项目申报方提交的绩效目标表设置了产出、效益、满意度等多维度指标，框架规范。经审核，主要建议两处修正：一是将"武汉航线平均客座率超80%、郑州旺季客座率稳定90%以上"等表述调整为事前预期式（如"目标客座率75%"），避免将预期目标表述为已实现状态。二是建议将投入产出比指标的表述从单纯的人均消费÷人均补贴，扩展为包含税收回流、就业带动等可量化财政净效益的综合指标体系。')

# --- （四）实施方案可行性 ---
add_title('（四）实施方案可行性评估', level=2)
add_body('1. 执飞技术条件', indent=True, bold_prefix='')
add_body('西部航空CCAR-121运行规范确认九寨黄龙机场在其A319-133型批准使用机场列表中（类别含正常使用及备降）。RNP AR正式运行批复文件已提供。地面服务代理协议、航油供应保障合同已签订。总体判断，执飞技术条件具备。')

add_body('2. 航线审批及协议签署', indent=True, bold_prefix='')
add_body('客运量奖励协议（三方协议）约定：甲方州文体旅游局、乙方西部航空、丙方九黄机场。补贴标准650元/人次，保证金400万元（乙方按两航线分别支付）。协议第十一条5.2款明确："获得民航管理局批准是本协议生效的必要条件"。截至评估基准日，我所收到上述审批文件。协议各方尚未正式签署（我所收到的协议文件签章页为空白）。上述两事项构成项目启动的前置条件。')

add_body('3. 监管与止损机制', indent=True, bold_prefix='')
add_body('项目设计了客流三方核对机制（机场统计→州文旅局复核→州财政局终审）和防作弊机制。止损方面，设置了55%客座率黄色预警、75%存续考核门槛、连续两月低于55%约谈退出等量化指标。但操作细则较为简略，建议在正式协议中嵌入具体的执行条款和资金追回法律路径。')

# --- （五）筹资合规性 ---
add_title('（五）筹资合规性评估', level=2)
add_body('资金来源方面，州级财政补贴390万元拟通过州本级预算安排。省级补贴585万元拟依据川财教〔2025〕78号申报，但需另行走省财政厅审批程序，具体获批金额以省厅批复为准。公平性方面，九黄机场公司声明与其他现有执飞航空公司无利益冲突和关联关系，现有航线均有补贴安排。筹资合规性方面，需指出州级航线补贴管理办法尚未出台，补贴审批、资金拨付、绩效监控等环节缺乏制度化规范。建议尽快制定发布，将航线补贴管理纳入常态化制度轨道。')

# ---- 补贴模式比选 ----
add_title('（六）补贴模式比选建议（专题分析）', level=2)
add_body('鉴于省级政策框架采用"按班次补贴"模式（川财教〔2025〕78号），与阿坝州本项目拟采用的"按人头补贴"模式在方法论上存在差异，我所进行了两种模式的对比分析：')

add_table(
    ['比较维度', '人头补贴模式（现行方案）', '班次补贴模式（省政策框架）'],
    [
        ['补贴计算方式', '按实际付费旅客人次×650元/人', '按实际执飞班次×不超过省级上限'],
        ['与省级政策对齐', '方法论不同，需额外解释', '完全对齐省级政策框架'],
        ['预算确定性', '受客座率影响，存在波动', '按班次数确定，预算刚性强'],
        ['航司激励方向', '鼓励尽可能多拉客', '鼓励执飞+拉客并重'],
        ['审计合规便利性', '需证明人头统计准确性', '直接套用省级标准，审计风险较低'],
        ['省厅认可可能性', '需单独沟通确认', '与全省其他市州一致'],
    ],
    col_widths=[3.5, 5, 6]
)

add_body('建议：首季试运行可维持人头补贴模式（便于与历史数据对标），第二年起探索向班次补贴模式过渡。或采用"双约束"机制：每班次实际补贴上限=MIN（实际人次×650元，省定上限×州级叠加系数）。')

doc.add_page_break()

# ============================================================
# 第四部分：评估结论
# ============================================================
add_title('四、评估结论', level=1)

add_title('（一）总体结论', level=2)
add_body('在补齐下列先决条件后，建议予以支持。评估建议维持975万元预算规模（其中州级财政负担390万元），评减0万元。', indent=True)

add_body('本结论基于以下已完成的核验工作：', indent=True)
add_body('——650元/人次补贴标准在成本推导维度上逻辑自洽，处于省内航线补贴合理区间；', indent=True)
add_body('——航线填补华中市场空白、差异化竞争高铁的定位合理；', indent=True)
add_body('——航空公司执飞资质基本具备，串飞航线运行模式具有创新性；', indent=True)
add_body('——理论投入产出具备较好经济性基础（但效益测算的消费端数据可核验性受限）。', indent=True)

add_title('（二）资金拨付前置条件', level=2)
add_body('以下条件全部满足后，方可拨付补贴资金：', indent=True)

add_table(
    ['序号', '前置条件', '当前状态', '负责单位'],
    [
        ['1', '取得民航西南地区管理局航线审批批复文件', '尚未取得', '西部航空/九黄机场'],
        ['2', '三方客运量奖励协议正式签署（含量化止损条款和资金追回路径）', '协议文本完成但未签署', '州文体旅游局/西部航空/九黄机场'],
        ['3', '补充提供人均消费3,500元/人的原始调研数据文件，或调整为可独立核验的消费参数', '现有提供的文件经核验内容不一致', '九黄机场/专班'],
        ['4', '以650元/人次为参数重新提供正确的补贴测算矩阵（现有测算表参数为100元/人）', '需重新提供', '九黄机场'],
    ],
    col_widths=[1, 7, 3.5, 4]
)

add_title('（三）预算建议', level=2)
add_body('建议维持975万元预算规模不变（州级负担390万元、省级申请585万元）。主要理由：650元/人定价标准本身在成本推导和省内对标维度上具备合理基础，评减重点不在于金额本身而在于拨付前提条件的满足。待四项前置条件全部满足后，按季按实际人次×650元据实拨付。')

add_title('（四）主要风险提示', level=2)

add_table(
    ['风险等级', '风险描述', '建议应对'],
    [
        ['高', '航线审批及协议签署尚未完成，项目法律关系待建立', '在审批到位、协议签署前，不启动资金拨付程序'],
        ['高', '州级航线补贴管理办法缺失，资金管理缺乏制度化框架', '建议在首笔补贴拨付前完成管理办法制定'],
        ['中', '历史客座率数据存在内部矛盾（80% vs 36.63%），客流预测假设存在不确定性', '建议以最保守客座率假设设定首季目标，季末根据实际数据调整下季预算'],
        ['中', '人均消费数据的一手来源文件无法独立核验，效益测算的消费端存在不确定性', '标注数据来源说明；建议委托第三方开展独立的游客消费调研'],
        ['中', '补贴模式与省级框架存在方法论差异（人头vs班次）', '建议与省财政厅沟通确认模式的合规性；探索向班次模式过渡'],
    ],
    col_widths=[1.8, 6, 7]
)

# ============================================================
# 利益相关方分析（新增）
# ============================================================
add_title('（五）利益相关方角色分析', level=2)
add_body('我所注意到，州文体旅游局在本项目中同时承担三重角色：作为航线开发工作专班办公室挂靠单位（政策制定）、作为客运量奖励协议甲方（合同签约）、以及作为"航空+旅游"产品开发的职能单位（项目受益）。建议在制度设计中增设防火墙机制，如将协议甲方调整为州财政局或增加财政局为联合甲方，将专班办公室主任由州政府副秘书长担任而非文体旅游局局长，以避免"运动员兼裁判"的角色冲突。')

doc.add_page_break()

# ============================================================
# 第五部分：发现的问题
# ============================================================
add_title('五、发现的主要问题', level=1)

add_body('（一）核心效益数据可核验性受限。人均旅游消费3,500元/人作为效益测算的基石数据，其一手的省文旅厅通知文件经信息化提取后实际内容与文件名不一致。该数据目前仅能依据测算底稿的二手转述。', indent=True)

add_body('（二）项目申报文件存在内部数据矛盾。可行性报告中"历史客座率均超80%"与统计表中2023年武汉=九寨航线客座率36.63%存在显著差异，两组数据均来源于同一提供方。建议在项目实施前，由航线开发工作专班核实并统一口径。', indent=True)

add_body('（三）制度化管理框架不完整。州级层面尚未出台航线补贴管理办法及实施细则，补贴审批、资金拨付、绩效监控、监督问责等环节缺乏制度化规范。', indent=True)

add_body('（四）项目协议签署及审批尚未完成。三方协议未正式签署，协议约定的生效条件（民航审批）尚未满足，项目在法律层面尚不具备启动条件。', indent=True)

add_body('（五）补贴测算工具存在参数错误。项目申报方提供的客座率票价测算表中，补贴参数为100元/人次而非政策约定的650元/人次，两者存在6.5倍偏差，该测算表不可用于决策支持。', indent=True)

# ============================================================
# 第六部分：改进建议
# ============================================================
add_title('六、改进建议', level=1)

add_title('（一）近期建议（出具报告后30日内）', level=2)
add_body('1. 加快推进航线审批和协议签署，确保项目在法律框架完备的前提下启动。', indent=True)
add_body('2. 补充提供正确的消费数据来源文件；如无法提供，建议采用更保守的消费参数重新测算效益，或委托第三方开展独立游客消费调研。', indent=True)
add_body('3. 重新制作以650元/人次为补贴参数的客座率-补贴测算矩阵。', indent=True)
add_body('4. 将量化止损条款、资金追回机制等嵌入正式协议。', indent=True)

add_title('（二）中期建议（项目实施前）', level=2)
add_body('1. 制定并发布《阿坝州航线补贴资金管理办法》，将航线补贴管理纳入制度化轨道。', indent=True)
add_body('2. 修正绩效目标表的表述口径，将事后式评价调整为事前目标式表述。', indent=True)
add_body('3. 与省财政厅沟通确认人头补贴模式的合规性，评估是否需要向班次补贴模式过渡。', indent=True)
add_body('4. 完善利益相关方角色设置，增设甲方约束机制。', indent=True)

add_title('（三）长期建议（培育期）', level=2)
add_body('1. 建立航线补贴年度绩效跟踪评价机制，首季结束后根据实际运营数据动态调整补贴策略。', indent=True)
add_body('2. 研究补贴递减和市场化退出路径，避免形成长期刚性支出。', indent=True)
add_body('3. 积累OTA平台搜索数据、旅行社订团数据等需求端数据，建立客流预测的独立验证体系。', indent=True)

# ============================================================
# 第七部分：其他需要说明的事项
# ============================================================
add_title('七、其他需要说明的事项', level=1)

add_body('（一）评估范围受限说明。受限于项目申报方提供资料的完整性，下列事项本所未能独立核实：①航空公司内部测算文件及董事会决议（申报方表示"无法提供"）；②补贴递减与动态退出方案（申报方表示"协议一年一签，无法固定三年递减模式"）；③机场旺季保障能力独立评估（相关文件为空）；④OTA平台搜索数据及旅行社市场调研数据。上述受限事项对评估结论的完整性存在一定影响。')

add_body('（二）责任声明。本报告依据项目申报方提供的文件资料编制，我所对资料的完整性和真实性不承担责任。本报告的评估结论基于截至评估基准日（2026年7月25日）可获取的资料，若后续出现影响评估结论的重大事项，应重新评估。')

add_body('（三）使用限制。本报告仅供阿坝州财政局及州政府航线补贴决策使用。未经我所书面同意，不得将本报告全部或部分内容用于其他目的。')

doc.add_page_break()

# ============================================================
# 附件清单
# ============================================================
add_title('附件清单', level=1)

add_table(
    ['附件编号', '附件名称', '说明'],
    [
        ['附件1', '结论-数据-来源-推导 证据链对照表', 'Excel格式，含4个工作表/44条逐项溯源'],
        ['附件2', '政策文件逐件核验表', '11份核心政策/法律文件核验结果'],
        ['附件3', '缺失资料及补充建议清单', '10项关键缺失资料及补救建议'],
        ['附件4', '全部审核文件清单', '两方提供共69个文件的完整目录'],
        ['附件5', '绩效目标表（建议修正版）', '调整为事前预期式表述'],
        ['附件6', '项目单位征求意见及采纳情况说明', '待被评估单位反馈后补充'],
    ],
    col_widths=[2, 7, 6.5]
)

doc.add_paragraph()
doc.add_paragraph()

# 落款
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('四川融策会计师事务所有限公司')
run.font.size = Pt(14)
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('二〇二六年七月二十五日')
run.font.size = Pt(14)
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# ============================================================
# 保存
# ============================================================
outpath = os.path.join(OUTDIR, '郑州=九寨=武汉航线人头补贴事前绩效评估报告.docx')
doc.save(outpath)
print(f'\n报告已保存: {outpath}')
print(f'文档页数（估算）: 13-15页（含封面/图表/表格）')
print(f'嵌入图表: {len(chart_paths)}张')
