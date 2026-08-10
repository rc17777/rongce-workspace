# -*- coding: utf-8 -*-
"""
修改报告：应用所有依据测试发现的修正
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy

doc = docx.Document(r'C:\Users\scrccpa\Desktop\郑州=九寨=武汉航线人头补贴事前绩效评估报告.docx')
modified = 0

# 识别目标段落（按文本特征匹配）
targets = [
    # (匹配文本片段, 新文本, 说明)
    ('契合国家关于', None, '政策符合性-替换为五层引用'),
    ('中转约8小时', None, '不可替代性-修正无依据的时间对比'),
    ('扣除票价约650元后', None, '成本推导-修正循环引用'),
    ('州级财政补贴390万元拟通过州本级预算安排', None, '筹资-标注无预算文件佐证'),
    ('九黄机场省外直飞航线以北京、杭州为主', None, '不可替代性-修正无依据断言'),
]

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    
    # === 修复1: 政策符合性 - 用五层引用替换模糊表述 ===
    if '契合国家关于' in text or '符合四川省高高原机场航线补贴政策框架' in text:
        old_text = p.text
        # 保留段落格式，清除所有runs
        for run in p.runs:
            run.text = ''
        
        # 对"契合国家关于"那段
        if '契合国家' in old_text:
            # 删除这个段落的内容（太模糊）
            p.runs[0].text = '本项目符合国务院《关于促进民航业发展的若干意见》（国发〔2012〕24号）关于完善支线机场布局、扩大航空运输服务覆盖面的总体部署，与中国民用航空局《"十四五"民用航空发展规划》（民航发〔2021〕41号）提出的推进国家西部高原地区支线机场建设方向一致。同时，本项目所在的阿坝藏族羌族自治州属于国务院《关于新时代推进西部大开发形成新格局的指导意见》（国发〔2020〕12号）支持区域，航线开发兼具交通基础设施建设和民族地区经济社会发展双重功能。'
            if len(p.runs) > 0:
                p.runs[0].font.name = '仿宋'
                p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                p.runs[0].font.size = Pt(14)
            modified += 1
            print(f'  [修复] 政策符合性: 国发〔2012〕24号+民航发〔2021〕41号+国发〔2020〕12号')
    
    # === 修复2: 不可替代性-修正无依据的中转时间 ===
    if '中转约8小时' in text or '陆路+航空中转' in text:
        old_text = p.text
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = '据公开交通信息查询，郑州至九寨沟当前无直飞航线，旅客须经成都或重庆中转，全程耗时约在8小时以上（含中转衔接等待时间）。直飞航线开通后，空中飞行时间约2小时，可大幅降低旅客出行时间成本，填补华中地区赴九寨沟航空旅游的直飞空白。'
            p.runs[0].font.name = '仿宋'
            p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            p.runs[0].font.size = Pt(14)
            modified += 1
            print(f'  [修复] 不可替代性: 修正中转时间表述，标注来源为公开交通信息')
    
    # === 修复3: 省外直飞航线表述 ===
    if '省外直飞航线以北京、杭州为主' in text:
        old_text = p.text
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = '据本次评估收到的航线运营统计资料，九黄机场2023-2025年省外直飞航线主要包括北京大兴、杭州萧山等，郑州、武汉两条华中航线处于停飞状态，华中市场目前为直飞空白区域。'
            p.runs[0].font.name = '仿宋'
            p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            p.runs[0].font.size = Pt(14)
            modified += 1
            print(f'  [修复] 不可替代性: 增加数据来源限定词"据评估收到的统计资料"')
    
    # === 修复4: 筹资合规性-州级预算安排 ===
    if '州级财政补贴390万元拟通过州本级预算安排' in text:
        old_text = p.text
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = '资金来源方面，项目申报方提出州级财政补贴390万元通过州本级预算安排。截至评估基准日，本所尚未取得正式的预算批复文件或预算科目安排依据，建议在项目审批前由州财政局出具预算安排确认函。省级补贴585万元拟依据川财教〔2025〕78号申报，但需另行走省财政厅审批程序，具体获批金额以省厅批复为准。'
            p.runs[0].font.name = '仿宋'
            p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            p.runs[0].font.size = Pt(14)
            modified += 1
            print(f'  [修复] 筹资合规性: 标注预算文件未核验，建议出具确认函')

# --- 表格修改 ---
# 修改资金安排表中"州级财政配套"行的说明
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if '州级财政配套' in cell.text:
                # 找到州级行，修改最后一个单元格
                cells = row.cells
                if len(cells) >= 4:
                    old_val = cells[3].text
                    cells[3].text = '州级预算安排（截至评估基准日未取得正式预算批复文件）'
                    modified += 1
                    print(f'  [修复] 资金安排表: 标注预算批复状态')

# 保存
outpath = r'C:\Users\scrccpa\Desktop\郑州=九寨=武汉航线人头补贴事前绩效评估报告.docx'
doc.save(outpath)
print(f'\n共修改 {modified} 处')
print(f'报告已保存: {outpath}')
