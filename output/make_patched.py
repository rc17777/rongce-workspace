# -*- coding: utf-8 -*-
"""
局部修改版：基于现有docx，修复P0致命问题后可应急使用
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil

# 复制原文件
src = r'C:\Users\scrccpa\Desktop\郑州=九寨=武汉航线人头补贴事前绩效评估报告.docx'
dst = r'C:\Users\scrccpa\Desktop\郑州=九寨=武汉航线绩效评估报告_局部修改版.docx'
shutil.copy(src, dst)
doc = Document(dst)

changes = 0

def set_run_font(run, name='仿宋', size=14, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def make_xml_para(text, font_name='仿宋', font_size=14, bold=False, indent=True):
    np = OxmlElement('w:p')
    npPr = OxmlElement('w:pPr')
    if indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '720')
        npPr.append(ind)
    np.append(npPr)
    nr = OxmlElement('w:r')
    nrPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), font_name)
    rf.set(qn('w:ascii'), font_name)
    nrPr.append(rf)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    nrPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        nrPr.append(b)
    nr.append(nrPr)
    nt = OxmlElement('w:t')
    nt.text = text
    nt.set(qn('xml:space'), 'preserve')
    nr.append(nt)
    np.append(nr)
    return np

print('=== 局部修改版 ===')

# ===== P0-1: 修复"航线航线"重复 =====
print('P0-1: 修复"航线航线"重复...')
for para in doc.paragraphs:
    for run in para.runs:
        if '航线航线' in run.text:
            run.text = run.text.replace('航线航线', '航线')
            changes += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if '航线航线' in run.text:
                        run.text = run.text.replace('航线航线', '航线')
                        changes += 1
print(f'  修复{changes}处')

# ===== P0-2: 修复章节编号（八→七，七→八） =====
print('P0-2: 修复章节编号...')
prev_changes = changes
for para in doc.paragraphs:
    for run in para.runs:
        if '八、远期展望' in run.text or '八、远期展望' in run.text:
            run.text = run.text.replace('八、远期展望', '七、远期展望（初步思考，不作为评估结论组成部分）')
            changes += 1
        if '七、其他需要说明的事项' in run.text or '七、其他需要说明的事项' in run.text:
            run.text = run.text.replace('七、其他需要说明的事项', '八、其他需要说明的事项')
            changes += 1
        if '八、远期展望：航线三阶段培育路径' in run.text:
            run.text = run.text.replace('八、远期展望：航线三阶段培育路径', '七、远期展望：航线三阶段培育路径（初步思考，不作为评估结论组成部分）')
            changes += 1
print(f'  修复{changes - prev_changes}处')

# ===== P0-3: 给法规加书名号 =====
print('P0-3: 补书名号...')
prev = changes
for para in doc.paragraphs:
    for run in para.runs:
        t = run.text
        if '预算法' in t and '《预算法》' not in t and '中华人民共和国' not in t:
            run.text = t.replace('预算法', '《中华人民共和国预算法》')
            changes += 1
            break

# 检查评估依据中的法规
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t == '1.中华人民共和国预算法' or t == '1. 中华人民共和国预算法':
        for run in para.runs:
            run.text = '1.《中华人民共和国预算法》'
        changes += 1
    if '2.四川省预算绩效评估管理办法' in t and '《' not in t:
        for run in para.runs:
            if '2.' in run.text and '四川省预算绩效评估管理办法' in run.text:
                run.text = run.text.replace('四川省预算绩效评估管理办法', '《四川省预算绩效评估管理办法》（川财绩〔2025〕12号）')
        changes += 1
    if '3.四川省省级财政高高原机场' in t:
        for run in para.runs:
            if '四川省省级财政' in run.text:
                run.text = run.text.replace(
                    '四川省省级财政高高原机场航线航班补贴资金管理办法',
                    '《四川省省级财政高高原机场航线航班补贴资金管理办法》（川财教〔2025〕78号）')
        changes += 1
    if '4.四川省省级财政专项资金管理办法' in t and '《' not in t:
        for run in para.runs:
            if '四川省省级财政专项资金管理办法' in run.text:
                run.text = run.text.replace('四川省省级财政专项资金管理办法', '《四川省省级财政专项资金管理办法》（川府发〔2023〕2号）')
        changes += 1
print(f'  修复{changes - prev}处')

# ===== P0-4: 在摘要顶部加数据一致性声明 =====
print('P0-4: 添加数据一致性声明...')
# 找到报告摘要后的第一个段落
target = None
for i, para in enumerate(doc.paragraphs):
    if '川融策专审' in para.text:
        target = para._element
        break

if target:
    warning_box = make_xml_para(
        '【重要提示】报告出具方已注意到报告存在以下待核实事项，在以下事项核实完毕前，本报告的评估结论仅供预审参考，不作为正式拨款依据：'
        '①目标旅客吞吐量15,000人次与目标客座率75%无法同时自洽——报告引用的A319座位数存在不一致（100座/75座），'
        '两种口径下按75%客座率计算的实际人次分别为11,250人次和8,438人次，均低于15,000。'
        '建议由航线开发工作专班在项目实施前确认A319实际座位布局和客流预测口径，统一全报告数据。'
        '②本报告出具日（2026年7月25日）三方协议尚未正式签署，民航审批尚未完成，协议以民航审批为生效条件。'
        '如前述条件未全部满足，项目不具备法律基础，本评估结论自动失效。',
        font_name='楷体', font_size=12, bold=True)
    target.addprevious(warning_box)
    changes += 1
    print('  已添加')

# ===== P0-5: 填补"资金安排"空章节 =====
print('P0-5: 填补空章节...')
for i, para in enumerate(doc.paragraphs):
    if '（四）资金安排' in para.text and (i+1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip()):
        target_elem = para._element
        fill_text = '根据项目申报方提交的方案，航线绩效激励资金总额975万元，其中申请省级财政585万元（依据川财教〔2025〕78号）、州级财政配套390万元（拟通过州本级预算安排）。省级资金需另行走省财政厅审批程序；州级资金截至评估基准日尚未取得正式预算批复文件，建议在项目审批前由州财政局出具预算安排确认函。'
        np = make_xml_para(fill_text)
        target_elem.addnext(np)
        changes += 1
        break

# 填补"（四）主要风险提示"
for i, para in enumerate(doc.paragraphs):
    if '（四）主要风险提示' in para.text:
        # Check if next paragraph is empty
        nxt = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if nxt and not nxt.text.strip():
            target_elem = para._element
            fill_text = '一是数据可靠性风险：申报方可行性报告"历史客座率均超80%"与2023年实际运营数据36.63%存在实质性矛盾，客流预测假设的可靠性存疑。二是法律基础风险：三方协议以民航审批为生效条件，审批未完成前项目不具备法律基础。三是预算合规风险：州级390万元配套资金无正式预算批复，存在无法到位的可能。四是运营风险：历史武汉=九寨航线进出港比约28:1（进港2,317人次、出港82人次），高度单向客流可能导致航司实际亏损远超盈亏模型预测，航线可持续性存疑。'
            np = make_xml_para(fill_text)
            target_elem.addnext(np)
            changes += 1
            break

print(f'  填补空章节完成')

# ===== P0-6: 统一术语：补充"绩效激励"定义 =====
print('P0-6: 补充术语定义...')
for i, para in enumerate(doc.paragraphs):
    if '评估原则' in para.text:
        target_elem = para._element
        def_text = '【术语说明】本报告所称"绩效激励"，指财政资金不与航空公司亏损挂钩、不与固定班次捆绑，而按实际达成且经三方核验的合格旅客吞吐量据实结算的引导性财政支出。其与第三方协议中"客运量奖励"一词指向同一笔资金，与省级政策文件（川财教〔2025〕78号）中"航线航班补贴"在方法论上存在差异（省级按班次、本州按人次）。本报告在非引用政策原文的语境下统一使用"绩效激励"一词。'
        np = make_xml_para(def_text, font_name='楷体', font_size=12, indent=True)
        target_elem.addnext(np)
        changes += 1
        break

print(f'  已添加术语定义')

# ===== 保存 =====
doc.save(dst)
print(f'\n总计修改: {changes}处')
print(f'局部修改版已保存: {dst}')
