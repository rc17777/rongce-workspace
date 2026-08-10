# -*- coding: utf-8 -*-
"""融策制度v2.0 第一批：回款考核+员工绩效+部门绩效"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

outdir = r'C:\Users\scrccpa\.openclaw\workspace\output'

def make_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)
    style = doc.styles['Normal']
    style.font.name = '宋体'; style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    return doc

def ap(doc, text, bold=False, align=None, fs=12, fn='宋体', indent=None, sa=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text); r.font.size = Pt(fs); r.font.name = fn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fn); r.bold = bold
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    sizes = {0: 16, 1: 14, 2: 12}
    r.font.size = Pt(sizes.get(level, 12)); r.bold = True
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def clause(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num}  '); r.font.size = Pt(12); r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.bold = True
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = '仿宋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def make_table(doc, headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9); r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(9); r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t

def add_sign_page(doc, title):
    doc.add_page_break()
    heading(doc, f'《{title}》签收确认表', 1)
    ap(doc, '')
    make_table(doc, ['序号','姓名','岗位','签收日期','签字','声明'],
        [[str(i),'','','','','本人确认已阅读并理解本制度全部内容'] for i in range(1,9)])
    ap(doc, '注：本表由人力资源统一保管，作为制度已告知员工的书面证据。', indent=0.74, fn='仿宋')

# ═══ 文件1：回款考核 v2.0 ═══
doc1 = make_doc()
heading(doc1, '项目回款与清欠专项考核办法', 0)
ap(doc1, '制度编号：RC-FIN-006  版本：V2.0  修订日期：2026年7月', align=WD_ALIGN_PARAGRAPH.CENTER, fs=10)
ap(doc1, '')
heading(doc1, '第一章  总  则', 1)
clause(doc1, '第一条', '为强化项目回款管理，建立"以回款论英雄"的考核导向，根据公司《项目收入确认与回款管理制度》（RC-FIN-002），制定本办法。')
clause(doc1, '第二条', '本办法适用于公司全部业务项目的回款考核。')
clause(doc1, '第三条', '双清定义：\n（一）"清收"：清理和回收已完工项目的应收账款，主要责任人为项目经理。\n（二）"清欠"：清理和追回其他应收款、挂账款等，主要责任人为财务部。')
clause(doc1, '第四条', '考核原则：以奖为主、以罚为辅；量化到人、责任到岗；按月统计、季度考核。')
clause(doc1, '第五条', '【民主程序】本办法经公司全体职工大会讨论通过。全体员工须签署《制度签收确认表》。')

heading(doc1, '第二章  考核对象与指标', 1)
clause(doc1, '第六条', '考核对象：项目经理（第一责任人）、部门负责人（督促责任人）、财务部清欠岗（执行责任人）。')
clause(doc1, '第七条', '项目经理考核指标：\n（一）项目回款完成率 = 实际回款 ÷ 应收目标 × 100%。应收目标以合同约定付款节点计算；合同未约定的，按行业惯例：交付后90天内≥70%，180天内≥95%。\n（二）长期挂账消减数。')
clause(doc1, '第八条', '部门负责人：部门综合回款率 + 超9月挂账消减率。')
clause(doc1, '第九条', '财务部：年度清欠完成率 + 超1年挂账销号率。')

heading(doc1, '第三章  考核奖罚标准', 1)
clause(doc1, '第十条', '项目经理奖惩标准：')
make_table(doc1, ['回款完成率','奖惩','方式'],
    [['≥100%','奖励','超出部分每1%奖项目奖金2%，封顶30%'],
     ['90%-100%','不奖不罚','——'],
     ['75%-90%','轻度','当季个人绩效系数下调一级（如B→C）'],
     ['60%-75%','中度','绩效系数下调一级+项目奖金扣减20%'],
     ['<60%','重度','绩效系数定为D级(0.5)，取消该项目全部奖金']])
ap(doc1, '')
ap(doc1, '说明：绩效系数下调仅影响当季绩效工资，不影响基本工资。扣减后薪酬不低于当地最低工资标准。', indent=0.74, fn='仿宋')
clause(doc1, '第十一条', '长期挂账消减奖励：每催回一笔超1年挂账，奖励回款金额的2%，单笔上限5,000元。')
clause(doc1, '第十二条', '部门负责人奖惩：')
make_table(doc1, ['回款率','奖惩','方式'],
    [['≥95%','奖励','部门绩效系数上浮0.1，封顶1.2'],
     ['85%-95%','不奖不罚','——'],
     ['70%-85%','处罚','部门绩效系数下调0.1'],
     ['50%-70%','较重','系数下调0.2+绩效面谈'],
     ['<50%','严重','系数下调0.3+总经理约谈']])
ap(doc1, '')
clause(doc1, '第十三条', '财务部清欠奖惩：')
make_table(doc1, ['清欠完成率','奖惩','方式'],
    [['≥100%','奖励','每高1%奖800元，封顶15,000元'],
     ['80%-100%','不奖不罚','——'],
     ['60%-80%','处罚','个人绩效系数下调一级'],
     ['40%-60%','较重','下调两级'],
     ['<40%','严重','定为D级，调离清欠岗位']])
ap(doc1, '')
clause(doc1, '第十四条', '【赔偿上限】因项目经理明显怠于催收（连续90天无催收记录）导致坏账的：\n（一）该项目回款完成率按零分计算；\n（二）赔偿每月不超过月工资20%，扣除后不低于最低工资标准；\n（三）赔偿总额以半年工资为上限；\n（四）坏账认定须经财务部初审→总经理办公会审议→外部审计认可。')
clause(doc1, '第十五条', '年度绩效扣罚上限为绩效工资总额的50%（不含基本工资），且不低于最低工资标准。')

heading(doc1, '第四章  考核程序', 1)
clause(doc1, '第十六条', '每年1月15日前编制《年度双清目标清单》。')
clause(doc1, '第十七条', '每月5日前出具《项目回款月报》。排名仅发总经理和部门负责人，不公开通报。')
clause(doc1, '第十八条', '每季度首月10日前计算奖惩，总经理审批后执行。')
clause(doc1, '第十九条', '岗位变动的按实际任职时间分段计算。')

heading(doc1, '第五章  配套措施', 1)
clause(doc1, '第二十条', '催收工具：财务部统一制作《催收函》模板。逾期超180天可协调律师发函。')
clause(doc1, '第二十一条', '催收记录：项目经理每周五前汇总本周催收记录并更新台账。')
clause(doc1, '第二十二条', '"暂停合作"通道：逾期超180天、催收3次无果的政府客户，可申请暂停新业务合作。必要时向上级主管部门反映。')
clause(doc1, '第二十三条', '红黄灯预警（仅发总经理和部门负责人）：🟢≥90%正常 | 🟡60%-90%部门负责人约谈 | 🔴<60%总经理约谈。')

heading(doc1, '第六章  附  则', 1)
clause(doc1, '第二十四条', '本办法与RC-FIN-002并行执行。')
clause(doc1, '第二十五条', '本办法自2026年  月  日起施行。')
clause(doc1, '第二十六条', '本办法由财务部负责解释。')
add_sign_page(doc1, '项目回款与清欠专项考核办法')
doc1.save(os.path.join(outdir, '融策-项目回款与清欠专项考核办法-v2.0.docx'))
print('✅ 回款考核v2.0')

# ═══ 文件2：员工绩效 v2.0 ═══
doc2 = make_doc()
heading(doc2, '融策公司员工绩效考核办法', 0)
ap(doc2, '制度编号：RC-HR-005  版本：V2.0  修订日期：2026年7月', align=WD_ALIGN_PARAGRAPH.CENTER, fs=10)
ap(doc2, '')
heading(doc2, '第一章  总  则', 1)
clause(doc2, '第一条', '为建立"业绩导向、量化到人"的绩效考核体系，制定本办法。')
clause(doc2, '第二条', '原则：公平公正公开；量化优先；鼓励差异化。')
clause(doc2, '第三条', '适用于除总经理外全体在岗员工。')
clause(doc2, '第四条', '【民主程序】本办法经全体职工大会讨论通过。')

heading(doc2, '第二章  考核体系', 1)
clause(doc2, '第五条', '实行"季度绩效+年度KPI"双层考核。')
clause(doc2, '第六条', '绩效基数随公司年度目标动态调整。')
clause(doc2, '第七条', '业务部门绩效基数上浮10%。')

heading(doc2, '第三章  季度考核——标准评级制', 1)
clause(doc2, '第八条', '季度考核采用"3+1"结构：3项定量(70%)+1项重点工作(30%)。')
clause(doc2, '第九条', '每季度首月10日前确认指标和目标值。')
clause(doc2, '第十条', '标准评级制（达标即评，非强制分布）：')
make_table(doc2, ['等级','分数','系数','说明'],
    [['A（优秀）','≥90分','1.2','显著超出预期'],
     ['B（良好）','75-89分','1.0','达到预期'],
     ['C（合格）','60-74分','0.8','基本达到'],
     ['D（待改进）','<60分','0.5','未达要求']])
ap(doc2, '')
clause(doc2, '第十一条', '【D级评定要求】D级须逐项说明扣分理由并附客观材料。无客观依据的D级，人力资源不予核定。')
clause(doc2, '第十二条', '【待岗培训】连续两季D级→待岗培训（≤2个月）：\n（一）培训期间发基本工资，暂停绩效；\n（二）合格→恢复原岗原薪；\n（三）不合格→协商调岗，协商不成按《劳动合同法》处理。')
clause(doc2, '第十三条', '考核流程：员工自评→部门负责人评分→C级以下面谈（A/B级群体通报）→人力资源汇总→总经理审批。')

heading(doc2, '第四章  年度KPI', 1)
clause(doc2, '第十四条', '年度KPI奖金 = 基数 × 公司营收完成率 × 部门得分/100 × 个人得分/100。')
clause(doc2, '第十五条', '连续两季A级须附突出业绩书面材料，经总经理审批。')

heading(doc2, '第五章  考核纪律', 1)
clause(doc2, '第十六条', '考核不规范情形：全员同分无说明、C级以下无面谈、D级无客观材料→一周内补考核。')
clause(doc2, '第十七条', '申诉：3个工作日内书面提出。申诉期不影响执行，成功后追溯补发。禁止打击报复。')

heading(doc2, '第六章  附  则', 1)
clause(doc2, '第十八条', '本办法自2026年  月  日起施行。')
clause(doc2, '第十九条', '由人力资源负责解释。')
add_sign_page(doc2, '员工绩效考核办法')
doc2.save(os.path.join(outdir, '融策-员工绩效考核办法-v2.0.docx'))
print('✅ 员工绩效v2.0')

# ═══ 文件3：部门绩效 v2.0 ═══
doc3 = make_doc()
heading(doc3, '部门绩效考核办法', 0)
ap(doc3, '制度编号：RC-HR-006  版本：V2.0  修订日期：2026年7月', align=WD_ALIGN_PARAGRAPH.CENTER, fs=10)
ap(doc3, '')
heading(doc3, '第一章  总  则', 1)
clause(doc3, '第一条', '建立公司经营业绩与部门绩效联动机制。')
clause(doc3, '第二条', '核心逻辑：公司业绩→公司系数；部门交付→部门系数；部门总额=基数×人数×公司系数×部门系数；部门内按个人绩效差异化分配。')
clause(doc3, '第三条', '适用于审计业务部、工程咨询部、财务部、行政综合部。')
clause(doc3, '第四条', '【民主程序】经全体职工大会讨论通过。')

heading(doc3, '第二章  公司业绩系数', 1)
clause(doc3, '第五条', '公司业绩系数 = 营收完成率得分×70% + 回款完成率得分×30%。（v2.0砍掉利润指标）')
clause(doc3, '第六条', '得分规则：≥120%→120分，100%→100分，80%→80分，<60%→按实际。')
clause(doc3, '第七条', '系数映射：≥110→1.3，100→1.0，80→0.7，<70→0.3-0.5。线性插值。')

heading(doc3, '第三章  部门业绩系数', 1)
clause(doc3, '第八条', '部门业绩系数 = 交付(50%) + 管理(30%) + 贡献(20%)。')
clause(doc3, '第九条', '业务部门指标：')
make_table(doc3, ['维度','指标','权重','规则'],
    [['交付','报告一次通过率','20%','≥95%→100，每降1%扣2'],
     ['','交付准时率','15%','≥90%→100，每降1%扣2'],
     ['','有效投诉率','15%','0次→100，1次→80，≥3次→0'],
     ['管理','回款完成率','15%','≥100%→100，每降1%扣1'],
     ['','成本控制率','15%','≤预算→100，超5%→80，超15%→0'],
     ['贡献','客户续约率','10%','≥80%→100，每降5%扣10'],
     ['','知识沉淀','10%','完成计划→100，80%→80']])
ap(doc3, '')
clause(doc3, '第十条', '职能部门指标：')
make_table(doc3, ['维度','指标','权重','规则'],
    [['交付','工作差错率','20%','0次→100，1次→80，≥3次→0'],
     ['','任务按时完成率','15%','≥95%→100，每降1%扣2'],
     ['','制度合规率','15%','100%→100，每项违规扣10'],
     ['管理','预算偏差','15%','≤15%→100，15-25%→80，>25%→0'],
     ['','关键事件','15%','业务部门提供正/负面事件，综合评价'],
     ['贡献','创新提案','10%','≥2项→100，1项→60'],
     ['','协作评价','10%','360度定性评价']])
ap(doc3, '')
clause(doc3, '第十一条', '部门系数映射：≥95→1.2，85→1.0，75→0.9，<65→0.5。')

heading(doc3, '第四章  部门绩效总额', 1)
clause(doc3, '第十二条', '部门季度总额 = 基数 × 人数 × 公司系数 × 部门系数。')
clause(doc3, '第十三条', '【v2.0变更】季度全额发放，不再预留25%清算。年度激励通过年终奖另行安排。')

heading(doc3, '第五章  部门内部分配', 1)
clause(doc3, '第十四条', '个人绩效 = 总额 × (个人系数×岗位权重) ÷ ∑(全员系数×权重)。权重：负责人1.5/骨干1.2/普通1.0/试用0.6。')
clause(doc3, '第十五条', '【试用期保底】试用期员工薪酬不低于合同约定80%且不低于最低工资。')
clause(doc3, '第十六条', '分配流程：负责人提出→面谈→人事备案→总经理审批。')

heading(doc3, '第六章  附  则', 1)
clause(doc3, '第十七条', '自2026年  月  日起施行。')
clause(doc3, '第十八条', '由人力资源负责解释。')
add_sign_page(doc3, '部门绩效考核办法')
doc3.save(os.path.join(outdir, '融策-部门绩效考核办法-v2.0.docx'))
print('✅ 部门绩效v2.0')
print('\n🎉 第一批3份完成')
