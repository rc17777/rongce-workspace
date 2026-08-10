# -*- coding: utf-8 -*-
"""融策制度v2.0 第二批：投标管理+经营激励+岗位职责卡"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

outdir = r'C:\Users\scrccpa\.openclaw\workspace\output'

def make_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)
    style = doc.styles['Normal']
    style.font.name = '宋体'; style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    return doc

def ap(doc, text, bold=False, align=None, fs=12, fn='宋体', indent=None, sa=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text); r.font.size = Pt(fs); r.font.name = fn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fn); r.bold = bold
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    sizes = {0: 16, 1: 14, 2: 12}
    r.font.size = Pt(sizes.get(level, 12)); r.bold = True
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def clause(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num}  '); r.font.size = Pt(12); r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.bold = True
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = '仿宋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def make_table(doc, headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9); r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(9); r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t

def add_sign_page(doc, title):
    doc.add_page_break()
    heading(doc, f'《{title}》签收确认表', 1)
    ap(doc, '')
    make_table(doc, ['序号','姓名','岗位','签收日期','签字','声明'],
        [[str(i),'','','','','本人确认已阅读并理解本制度全部内容'] for i in range(1,9)])
    ap(doc, '注：本表由人力资源统一保管。', indent=0.74, fn='仿宋')

# ═══ 文件4：投标管理 v2.0 ═══
doc4 = make_doc()
heading(doc4, '融策公司投标管理办法', 0)
ap(doc4, '制度编号：RC-BIZ-003  版本：V2.0  修订日期：2026年7月', align=WD_ALIGN_PARAGRAPH.CENTER, fs=10)
ap(doc4, '')
heading(doc4, '第一章  总  则', 1)
clause(doc4, '第一条', '为规范投标流程，预防废标风险，提升投标文件质量，制定本办法。')
clause(doc4, '第二条', '适用于公司全部投标项目。')
clause(doc4, '第三条', '投标工作实行"专人负责、双人复核、责任追溯"原则。')
clause(doc4, '第四条', '【民主程序】本办法经全体职工大会讨论通过。')

heading(doc4, '第二章  投标组织', 1)
clause(doc4, '第五条', '投标工作组：投标负责人+资信编制人+技术编制人+报价编制人。')
clause(doc4, '第六条', '标准编标流程（倒排工期）：')
make_table(doc4, ['节点','距截标','责任人','产出'],
    [['标前评审','T-7天','投标负责人','可行性评估表'],
     ['分工启动','T-6天','投标负责人','任务分派单'],
     ['初稿完成','T-4天','各编制人','资信/技术/报价初稿'],
     ['交叉复核','T-3天','复核人','复核意见清单'],
     ['修改定稿','T-2天','各编制人','定稿版'],
     ['终审签字','T-1天','投标负责人','签字确认单'],
     ['封装递交','T日','资信编制人','密封文件+回执']])
ap(doc4, '')
clause(doc4, '第七条', '【紧急投标快速通道】距截标≤5个工作日的：\n（一）标前评审改为电话确认；\n（二）编标和复核可并行；\n（三）交叉复核和终审签字不可省略。')

heading(doc4, '第三章  交叉复核', 1)
clause(doc4, '第八条', '交叉复核：资信由投标负责人复核；技术由未参与编制的同级以上人员复核；报价由财务负责人复核。复核人须签字留底。')
clause(doc4, '第九条', '提交前须逐项核对《投标文件提交检查清单》（30项）。')

heading(doc4, '第四章  基础资料管理', 1)
clause(doc4, '第十条', '行政综合岗负责投标基础资料库（资质/人员/业绩/荣誉四类），每季度更新。')
clause(doc4, '第十一条', '【证书管理责任】\n（一）证书持有人为第一责任人，须到期前30天完成续期；\n（二）行政综合岗须到期前60天发出预警；\n（三）持证人未续期导致废标→持证人主要责任；\n（四）行政未预警导致遗漏→行政次要责任。')

heading(doc4, '第五章  投标激励与事故处理', 1)
clause(doc4, '第十二条', '投标计件奖励：')
make_table(doc4, ['类型','编标费','说明'],
    [['简单（询价/谈判）','500元','≤50页'],
     ['一般（磋商）','800元','50-100页'],
     ['较难（公开招标-服务）','1200元','>100页'],
     ['复杂（公开招标-工程）','1500元','全套']])
ap(doc4, '')
ap(doc4, '中标项目×1.5倍。分配：资信30%+技术40%+报价20%+统筹10%。', indent=0.74, fn='仿宋')
clause(doc4, '第十三条', '【投标事故处理】\n（一）一类事故（废标/落标）：\n  · 直接编制人：书面检查+当季绩效系数下调一级+取消该项目编标费；\n  · 复核人：批评教育+当季绩效系数下调半级；\n  · 年度内累计两次的：暂停投标编制资格一个季度，期间负责检查清单复核和资料库维护。\n（二）二类事故（未影响结果但有瑕疵）：\n  · 直接编制人：当季绩效系数下调半级；\n  · 复核人：口头警告。')

heading(doc4, '第六章  附  则', 1)
clause(doc4, '第十四条', '自2026年  月  日起施行。')
clause(doc4, '第十五条', '由工程咨询部负责解释。')
add_sign_page(doc4, '投标管理办法')
doc4.save(os.path.join(outdir, '融策-投标管理办法-v2.0.docx'))
print('✅ 投标管理v2.0')

# ═══ 文件5：经营激励 v2.0 ═══
doc5 = make_doc()
heading(doc5, '融策公司经营激励办法', 0)
ap(doc5, '制度编号：RC-BIZ-004  版本：V2.0  修订日期：2026年7月', align=WD_ALIGN_PARAGRAPH.CENTER, fs=10)
ap(doc5, '')
heading(doc5, '第一章  总  则', 1)
clause(doc5, '第一条', '鼓励全员参与经营开发，建立"全员经营"文化。')
clause(doc5, '第二条', '涵盖两类激励：项目介绍激励（全员适用）+ 经营开发激励（经营人员适用）。')
clause(doc5, '第三条', '【民主程序】本办法经全体职工大会讨论通过。')

heading(doc5, '第二章  项目介绍激励（全员适用）', 1)
clause(doc5, '第四条', '任何员工提供有效项目信息并促成中标的，按以下标准奖励：')
make_table(doc5, ['合同金额','奖励标准','说明'],
    [['10万以下','3%','最低500元'],
     ['10-50万','2%','——'],
     ['50-100万','1.5%','——'],
     ['100万以上','1%','封顶50,000元']])
ap(doc5, '')
clause(doc5, '第五条', '兑现条件：\n（一）投标前书面报备并获确认立项；\n（二）成功中标并签订合同；\n（三）首笔款项到账后一次性发放。')
clause(doc5, '第六条', '多人联合介绍的，报备时书面约定分配比例；未约定的均分。')
clause(doc5, '第七条', '公司领导层和专职经营人员不适用本章。')

heading(doc5, '第三章  经营开发考核（经营人员适用）', 1)
clause(doc5, '第八条', '经营人员实行目标责任制。年初下达年度经营目标（新签合同额+回款率）。')
clause(doc5, '第九条', '经营目标考核：\n（一）完成率低于50%→绩效面谈+改进计划；\n（二）连续两年低于50%→年度绩效不高于C级+协商调整岗位或不再续签聘任协议。')
clause(doc5, '第十条', '经费管理：超预算部分从经营奖金中抵扣：\n（一）超15%以内→抵扣超支额的30%；\n（二）超15%以上→抵扣超支额的50%；\n（三）客观原因（如招标文件费上涨等）经总经理认定后不适用。')
clause(doc5, '第十一条', '一标一奖：')
make_table(doc5, ['项目类型','计提标准','备注'],
    [['政府审计/绩效评价','1%',''],
     ['工程咨询','1.2%','含财政评审'],
     ['专项债/资产清查等','1.5%',''],
     ['新客户首次中标','×1.5',''],
     ['新业务领域','×2.0','']])
ap(doc5, '')
clause(doc5, '第十二条', '一标一奖分配：信息获取人20%+主要编制人30%+客户维护人40%+部门统筹10%。')
clause(doc5, '第十三条', '【量化扣减规则】项目出现以下情形的：\n（一）因团队工作失误导致亏损→奖金×0.5；亏损超合同额30%→取消；\n（二）因客户或不可抗力导致亏损→不扣减；\n（三）客户有效投诉（经核实属服务质量问题）→奖金×0.7；\n（四）扣减决定须经总经理办公会审议并书面告知。')
clause(doc5, '第十四条', '首笔回款到账后发放。')

heading(doc5, '第四章  投标经费管理', 1)
clause(doc5, '第十五条', '单项投标经费预算：公开招标800元/项目，磋商/谈判500元/项目，异地据实报销差旅。')
clause(doc5, '第十六条', '年度投标经费总额年初预算单列，经营开发部负责人为控制第一责任人。')

heading(doc5, '第五章  附  则', 1)
clause(doc5, '第十七条', '自2026年  月  日起施行。')
clause(doc5, '第十八条', '由经营开发部负责解释。')
add_sign_page(doc5, '经营激励办法')
doc5.save(os.path.join(outdir, '融策-经营激励办法-v2.0.docx'))
print('✅ 经营激励v2.0')

# ═══ 文件6：岗位职责卡 v2.0 ═══
doc6 = make_doc()
heading(doc6, '融策公司岗位职责卡', 0)
ap(doc6, '制度编号：RC-HR-004  版本：V2.0  修订日期：2026年7月', align=WD_ALIGN_PARAGRAPH.CENTER, fs=10)
ap(doc6, '')
ap(doc6, '【使用说明】每岗位一张卡片，任职人和直接上级共同确认签字。岗位变动时重新签署。', indent=0.74, fn='仿宋')
ap(doc6, '')

roles = [
    {'title':'岗位1：总经理','report':'执行董事/股东会','subs':'各部门负责人',
     'duties':['全面主持公司经营管理工作，对年度经营目标负总责',
               '审批年度经营计划和财务预算',
               '审批重大项目立项、投标报价和合同签订',
               '主持月度经营例会和季度绩效考核评审会',
               '审批制度文件和对外公文',
               '分管财务部和人力资源',
               '对接重要客户和主管部门',
               '审批应收账款催收升级方案',
               '制定公司年度发展战略'],
     'kpi':'营业收入完成率、净利润、应收账款周转天数'},
    {'title':'岗位2：审计业务部经理','report':'总经理','subs':'各项目经理/审计组长',
     'duties':['全面负责审计业务部管理，对业绩和执业质量负总责',
               '分配项目，审核成员配置，把控进度和预算',
               '审核审计报告初稿，对结论和法规引用复核',
               '组织部门业务培训和执业资格考试',
               '审核审计方案和取证计划',
               '跟踪部门回款进度，督促项目经理催收',
               '配合总经理对接政府审计客户',
               '组织季度业务复盘和案例分析',
               '签署P1级项目审计报告'],
     'kpi':'部门营收完成率、报告一次通过率、部门回款率'},
    {'title':'岗位3：工程咨询部经理','report':'总经理','subs':'各项目负责人/造价工程师',
     'duties':['全面负责工程咨询部管理，对业绩和交付质量负总责',
               '组织投标文件编制和审核',
               '审核工程预算、结算和全过程咨询成果',
               '调配工程师资源，监控工期和成本',
               '对接财政评审中心和建设单位',
               '跟踪部门回款进度',
               '组织技术培训和继续教育',
               '建立维护工程咨询基础数据库'],
     'kpi':'部门营收完成率、中标率、项目毛利率、部门回款率'},
    {'title':'岗位4：项目经理/审计组长','report':'部门经理','subs':'项目组成员',
     'duties':['对所负责项目的质量、进度、成本和回款负第一责任',
               '编制项目实施方案和取证计划',
               '组织开展现场审计/检查，指导成员取证',
               '撰写报告初稿，对数据和结论准确性负责',
               '与客户保持日常沟通',
               '项目完工后15个工作日内完成归档和收费申请',
               '跟踪回款：交付后45天首次提醒，之后每60天跟进',
               '每周五前汇总本周催收记录并更新台账',
               '填报考勤和工时，控制项目成本'],
     'kpi':'项目毛利率、报告一次合格率、回款完成率'},
    {'title':'岗位5：财务负责人','report':'总经理','subs':'出纳',
     'duties':['全面负责财务管理和会计核算',
               '编制年度财务预算和月度资金计划',
               '每月5日前出具财务报表和账龄分析表',
               '管理回款台账，标记红黄绿灯预警',
               '每月组织回款通报（仅发总经理和部门负责人）',
               '编制季度双清考核结果表',
               '负责税收筹划和纳税申报',
               '审核费用报销和对外付款',
               '协调外部审计、税务检查和银行事务'],
     'kpi':'财务报表准时率、资金计划准确率、年度清欠完成率'},
    {'title':'岗位6：审计助理','report':'项目经理/审计组长','subs':'——',
     'duties':['按方案和取证计划执行现场取证',
               '编制审计工作底稿',
               '完成数据核对、分析和整理',
               '协助撰写报告和整理附件',
               '完成项目归档',
               '记录工作日志和工时',
               '参加培训，完成继续教育'],
     'kpi':'底稿合格率、任务按时完成率'},
    {'title':'岗位7：造价工程师','report':'工程咨询部经理/项目负责人','subs':'——',
     'duties':['独立完成预算、结算、全过程咨询文件编制',
               '工程量计算、清单编制、组价分析',
               '参与财政评审核对',
               '收集整理材料价格信息',
               '协助编制投标文件技术标和报价标',
               '按进度计划交付成果'],
     'kpi':'成果文件一次合格率、交付准时率'},
    {'title':'岗位8：行政综合岗','report':'总经理','subs':'——',
     'duties':['【核心职责】负责公司行政事务和办公环境管理',
               '管理公司资质证照、执业许可证年检更新',
               '维护人员资质库（执业证书/职称证书/继续教育）',
               '负责投标基础资料库更新维护',
               '【辅助职责】（允许根据实际情况调整优先级）',
               '组织会议、培训和团建活动',
               '管理印章使用和合同归档',
               '办公用品采购和固定资产管理',
               '协助招聘、入职和考勤统计'],
     'kpi':'资质证照零过期、投标资料完整率'},
]

for role in roles:
    doc6.add_page_break()
    heading(doc6, role['title'], 1)
    ap(doc6, '')
    ap(doc6, f'汇报上级：{role["report"]}', indent=0.74, fn='仿宋')
    ap(doc6, f'下属岗位：{role["subs"]}', indent=0.74, fn='仿宋')
    ap(doc6, '')
    ap(doc6, '【岗位职责】', bold=True, indent=0.74)
    for i, duty in enumerate(role['duties'], 1):
        ap(doc6, f'{i}. {duty}', indent=1.5, fn='仿宋')
    ap(doc6, '')
    ap(doc6, f'【KPI指标】{role["kpi"]}', bold=True, indent=0.74)
    ap(doc6, '')
    ap(doc6, '任职人签字：______________    日期：______________', indent=0.74)
    ap(doc6, '直接上级签字：______________    日期：______________', indent=0.74)

add_sign_page(doc6, '岗位职责卡')
doc6.save(os.path.join(outdir, '融策-岗位职责卡（8岗位）-v2.0.docx'))
print('✅ 岗位职责卡v2.0')
print('\n🎉 第二批3份完成')
