# -*- coding: utf-8 -*-
"""生成融策版《项目回款双清专项考核办法》Word文档"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_para(text, bold=False, alignment=None, font_size=12, font_name=None, space_after=6, first_line_indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name or '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or '宋体')
    run.bold = bold
    return p

def add_heading_text(text, level=1):
    """添加标题样式段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 0:
        run.font.size = Pt(18)
        run.bold = True
    elif level == 1:
        run.font.size = Pt(15)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

# ══════════════════════════════════════════════
# 正文开始
# ══════════════════════════════════════════════

# 红头文号
add_para('四川融策会计师事务所有限公司', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, font_name='宋体')
add_para('四川融策工程咨询有限公司', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, font_name='宋体')

add_para('')

# 文号行
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('融策发〔2026〕  号')
run.font.size = Pt(14)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

add_para('')

# 标题
add_heading_text('关于印发《项目回款与清欠专项考核办法》的通知', level=0)

add_para('')

# 主送
add_para('公司各部门：', first_line_indent=0.74)

# 正文引言
intro = (
    '为进一步压实项目回款责任，加速资金回笼，降低应收账款存量，'
    '杜绝"签了合同不管钱、干了活没人收钱"的管理盲区，'
    '根据《RC-FIN-002 项目收入确认与回款管理制度》相关规定，'
    '结合公司项目制运营实际，制定了《项目回款与清欠专项考核办法》，'
    '经公司办公会议审议通过，现予印发，请遵照执行。'
)
add_para(intro, first_line_indent=0.74)

add_para('')

# 落款
add_para('四川融策会计师事务所有限公司', alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('四川融策工程咨询有限公司', alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('2026年  月  日', alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# 分页
doc.add_page_break()

# ══════════════════════════════════════════════
# 办法正文
# ══════════════════════════════════════════════

add_heading_text('项目回款与清欠专项考核办法', level=0)

add_para('')

add_heading_text('第一章  总  则', level=1)

clauses = [
    ('第一条', '为强化项目回款管理，建立"以回款论英雄"的考核导向，切实解决公司应收账款居高不下、部分项目长期挂账的问题，根据公司《项目收入确认与回款管理制度》（RC-FIN-002），制定本办法。'),
    ('第二条', '本办法适用于公司全部业务项目的回款考核，覆盖审计业务部、工程咨询部及财务部涉及回款和清欠职责的全体人员。'),
    ('第三条', '双清定义：\n（一）"清收"：指清理和回收已完工项目的应收账款，主要责任人为项目经理（审计组长/工程项目负责人）。\n（二）"清欠"：指清理和追回其他应收款、挂账款、代垫款等，主要责任人为财务部相关岗位。'),
    ('第四条', '双清专项考核坚持以下原则：\n（一）以奖为主、以罚为辅，正向激励优先；\n（二）量化到人、责任到岗，杜绝"大锅饭"；\n（三）按月通报、季度考核、年度清算。'),
]

for num, text in clauses:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{num}  ')
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

add_heading_text('第二章  考核对象与指标', level=1)

clauses2 = [
    ('第五条', '考核对象分为三类：\n（一）项目经理（含审计组长、工程项目负责人）——项目回款第一责任人；\n（二）部门负责人（审计业务部经理、工程咨询部经理）——部门整体回款督促责任人；\n（三）财务部清欠岗——协助清欠和长期挂账清理责任人。'),
    ('第六条', '项目经理考核指标：\n（一）项目回款完成率 = 考核期内该项目实际回款金额 ÷ 该项目当期应收目标金额 × 100%。\n应收目标金额以合同约定的付款节点计算，无明确节点的按以下规则确定：\n  · 项目完工交付后60天内应收不低于合同总额的70%；\n  · 项目完工交付后180天内应收不低于合同总额的95%；\n  · 质保金/尾款按合同约定时间计算。\n（二）长期挂账消减数：考核期内完成催收并到账的超1年挂账项目数量。'),
    ('第七条', '部门负责人考核指标：\n（一）部门综合回款率 = 部门管辖全部项目实际回款总额 ÷ 部门全部项目应收目标总额 × 100%；\n（二）部门超9月挂账项目消减率 = 考核期内部门消减的超9月挂账项目数 ÷ 年初部门超9月挂账项目总数 × 100%。'),
    ('第八条', '财务部考核指标：\n（一）年度清欠完成率 = 年度已清理其他应收款金额 ÷ 年度清欠目标金额 × 100%；\n（二）超1年挂账销号率 = 年度内完成销号处理的超1年挂账笔数 ÷ 年初超1年挂账总笔数 × 100%。'),
]

for num, text in clauses2:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{num}  ')
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

add_heading_text('第三章  考核奖罚标准', level=1)

add_para('第九条  项目经理项目回款奖惩标准：', bold=True, first_line_indent=0.74)

# 表格1
table1 = doc.add_table(rows=6, cols=4, style='Table Grid')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['回款完成率区间', '奖惩', '计算方式', '上限']
data1 = [
    ['≥100%', '奖励', '超出100%部分，每高1个百分点奖励项目奖金的2%', '项目奖金总额的30%'],
    ['90%～100%', '不奖不罚', '——', '——'],
    ['75%～90%', '轻度处罚', '每低1个百分点扣减项目奖金的1%', '扣至项目奖金的15%'],
    ['60%～75%', '中度处罚', '每低1个百分点扣减项目奖金的1.5%', '扣至项目奖金的35%'],
    ['＜60%', '重度处罚', '取消该项目全部奖金，项目经理本年度不得承接新项目', '——'],
]

for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        cell = table1.rows[r+1].cells[c]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('')
add_para('第十条  长期挂账消减专项奖励：', bold=True, first_line_indent=0.74)
add_para('项目经理每成功催收回一笔超1年挂账项目款项，额外奖励该笔回款金额的2%，单笔奖励上限5,000元。', first_line_indent=0.74)

add_para('')
add_para('第十一条  部门负责人奖惩标准：', bold=True, first_line_indent=0.74)

table2 = doc.add_table(rows=6, cols=4, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['部门综合回款率', '奖惩', '计算方式', '上限']
data2 = [
    ['≥95%', '奖励', '超出95%部分，每高1个百分点奖励500元', '封顶10,000元'],
    ['85%～95%', '不奖不罚', '——', '——'],
    ['70%～85%', '处罚', '每低1个百分点扣罚200元', '上限扣罚4,000元'],
    ['50%～70%', '较重处罚', '每低1个百分点扣罚300元', '上限扣罚8,000元'],
    ['＜50%', '严重处罚', '扣罚10,000元，年度不得评优', '——'],
]

for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        cell = table2.rows[r+1].cells[c]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('')
add_para('第十二条  财务部清欠奖惩标准：', bold=True, first_line_indent=0.74)

table3 = doc.add_table(rows=6, cols=4, style='Table Grid')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['年度清欠完成率', '奖惩', '计算方式', '上限']
data3 = [
    ['≥100%', '奖励', '超出100%部分，每高1个百分点奖励800元', '封顶15,000元'],
    ['80%～100%', '不奖不罚', '——', '——'],
    ['60%～80%', '处罚', '每低1个百分点扣罚300元', '上限扣罚6,000元'],
    ['40%～60%', '较重处罚', '每低1个百分点扣罚500元', '上限扣罚12,000元'],
    ['＜40%', '严重处罚', '扣罚15,000元，调离清欠岗位', '——'],
]

for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for r, row_data in enumerate(data3):
    for c, val in enumerate(row_data):
        cell = table3.rows[r+1].cells[c]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('')
add_para('第十三条  扣罚以年度绩效工资为上限（扣至零为止）。奖励金额在年度绩效工资清算时一并发放。', first_line_indent=0.74)

add_heading_text('第四章  考核程序', level=1)

clauses3 = [
    ('第十四条', '每年1月15日前，财务部汇总上年度全部项目应收目标金额，按项目、按部门编制《年度双清目标清单》，经总经理审批后作为当年考核基准。'),
    ('第十五条', '每月5日前，财务部编制上月《项目回款月报》和《应收账款账龄分析表》，发送总经理及各部门负责人，并在公司月度例会上通报回款进度排名。'),
    ('第十六条', '每季度首月10日前，财务部按本办法第九条至第十二条标准，计算上季度各考核对象的奖惩金额，形成《季度双清考核结果表》，经总经理审批后执行。'),
    ('第十七条', '年度终了后20个工作日内，财务部汇总全年回款数据，完成年度双清考核清算：\n（一）季度考核中已发放的奖励不扣回；\n（二）季度考核中已执行的处罚不退还；\n（三）年度清算仅对全年综合完成率进行总奖罚核算，扣除季度已兑现部分后多退少补。'),
    ('第十八条', '考核期内岗位发生变动的，按实际任职时间分段计算。项目移交的，回款责任随项目转移，移交时须在移交清单中写明回款状态。'),
]

for num, text in clauses3:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{num}  ')
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

add_heading_text('第五章  配套措施', level=1)

clauses4 = [
    ('第十九条', '催收工具保障：财务部统一制作《催收函》模板，项目经理发起催收时可直接套用。对逾期超过90天的项目，财务部配合出具律师函版本，必要时协调外部律师介入。'),
    ('第二十条', '催收记录管理：所有催收行为（电话、微信、上门、发函）须留痕，项目经理应在催收完成后24小时内在回款台账中记录，作为考核依据。无催收记录的挂账项目，发生坏账时项目经理承担主要责任。'),
    ('第二十一条', '"以诉促收"通道：对逾期超过180天、催收3次以上无实质进展的项目，项目经理可申请启动法律催收程序。经总经理批准后，由财务部协调外部律师推进。通过诉讼回款的金额按正常考核计算，诉讼费用计入项目成本。'),
    ('第二十二条', '红黄灯预警：财务部每月出具的项目回款月报中标注——\n  🟢 绿灯：回款进度正常（≥计划进度的90%）\n  🟡 黄灯：回款滞后（计划进度的60%～90%），由部门负责人约谈项目经理\n  🔴 红灯：回款严重滞后（＜计划进度的60%），由总经理约谈项目经理及部门负责人'),
]

for num, text in clauses4:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{num}  ')
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

add_heading_text('第六章  附  则', level=1)

clauses5 = [
    ('第二十三条', '本办法与《RC-FIN-002 项目收入确认与回款管理制度》并行执行。RC-FIN-002侧重流程规范，本办法侧重量化激励；两者有冲突的，以本办法为准。'),
    ('第二十四条', '本办法自2026年  月  日起施行。'),
    ('第二十五条', '本办法由财务部负责解释。'),
]

for num, text in clauses5:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{num}  ')
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

add_para('')

# 附件标记
add_para('附件：1. 年度双清目标清单（模板）', first_line_indent=0.74)
add_para('        2. 季度双清考核结果表（模板）', first_line_indent=0.74)
add_para('        3. 项目回款月报（模板）', first_line_indent=0.74)

add_para('')
add_para('')

# 抄送
add_para('抄送：总经理，财务部，审计业务部，工程咨询部，存档。')

add_para('')
add_para('四川融策会计师事务所有限公司办公室         2026年  月  日印发')

# 保存
output_path = r'C:\Users\scrccpa\.openclaw\workspace\output\融策-项目回款与清欠专项考核办法.docx'
doc.save(output_path)
print(f'✅ 已保存: {output_path}')
