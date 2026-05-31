#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档保格式编辑工具库 (preserve_format.py)

基于 AI-Word-Skill 的 SOP 方法论：
  - 复制模板 → run 级别编辑 → 保存
  - 反模式：paragraph.text = ... / Document() 新建

核心心法：.docx 本质是 ZIP + OOXML，段落由多个 w:r (run) 组成，
每个 run 可携带独立的字体/字号/颜色等 rPr。直接操作 run.text 可保留格式。

作者 sgsss998 / AI-Word-Skill 贡献核心思想，融策左护法整合中文适配。
"""

from __future__ import annotations

import shutil
from copy import deepcopy
from contextlib import contextmanager
from typing import Callable, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# 上下文管理器
# ---------------------------------------------------------------------------

@contextmanager
def PreserveDoc(template_path: str, output_path: str):
    """保格式上下文管理器：自动复制模板、加载、保存。

    Usage:
        with PreserveDoc("template.docx", "output.docx") as doc:
            replace_all(doc, "{{name}}", "张三")
    """
    shutil.copy(template_path, output_path)
    doc = Document(output_path)
    try:
        yield doc
    finally:
        doc.save(output_path)


def copy_template(template_path: str, output_path: str) -> Document:
    """复制模板并返回 Document 对象。"""
    shutil.copy(template_path, output_path)
    return Document(output_path)


# ---------------------------------------------------------------------------
# 核心替换函数
# ---------------------------------------------------------------------------

def replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """在单个段落中替换文字（单 run 内），保留 run 格式。

    Returns:
        True 若找到并替换了 at least one occurrence.
    """
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    return False


def replace_cross_runs(paragraph, old_text: str, new_text: str) -> bool:
    """跨 run 替换：处理 Word 把一段文字拆到多个 run 的情况。

    典型场景：run[0]="某城" + run[1]="市"，此时单 run 内找不到"某城市"。

    Returns:
        True 若找到并替换。
    """
    full_text = ''.join(r.text for r in paragraph.runs)
    if old_text not in full_text:
        return False

    start = full_text.find(old_text)
    end = start + len(old_text)

    # 找出与 old_text 跨度重叠的 run 索引
    char_pos = 0
    affected = []
    for i, run in enumerate(paragraph.runs):
        r_start = char_pos
        r_end = char_pos + len(run.text)
        if r_start < end and r_end > start:
            affected.append(i)
        char_pos = r_end

    if not affected:
        return False

    # 合并 affected runs 的文本，替换，写回第一个 run，清空其余
    merged = ''.join(paragraph.runs[i].text for i in affected)
    merged = merged.replace(old_text, new_text, 1)
    paragraph.runs[affected[0]].text = merged
    for i in affected[1:]:
        paragraph.runs[i].text = ''
    return True


def rewrite_paragraph(paragraph, new_text: str) -> None:
    """整段重写，保留段落格式和首个 run 的字体 DNA。

    关键：保留 runs[0]，只改它的 text。清空其余 run。
    代价：如果原段落有"段落内部分加粗/变色"的格式，会被抹平。
    适用：段落整体替换、段落内排版一致性优先的场景。

    如果 paragraph.runs 为空，不做任何操作。
    """
    if not paragraph.runs:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ''


def replace_all(doc: Document, old_text: str, new_text: str) -> int:
    """全文档批量替换（段落 + 表格单元格）。

    Returns:
        替换命中次数。
    """
    count = 0
    # 正文段落
    for p in doc.paragraphs:
        for run in p.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                count += 1
    # 表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
    return count


# ---------------------------------------------------------------------------
# 插入段落
# ---------------------------------------------------------------------------

def insert_paragraph_deepcopy(
    doc: Document,
    template_index: int,
    new_text: str,
    anchor: Optional = None,
) -> None:
    """通过 deepcopy 模板段落 XML 插入新段落，保留格式。

    这是 python-docx 中插入段落最安全的方式。
    doc.add_paragraph() 常见问题：KeyError（样式不存在）、默认样式漂移。

    Args:
        doc: 文档对象
        template_index: 用作模板的段落索引（选一个版式正确的段落）
        new_text: 新段落的文字内容
        anchor: 插入位置锚点（docx Paragraph 对象），默认为文档末尾
    """
    template_el = doc.paragraphs[template_index]._element

    new_p = deepcopy(template_el)
    # 移除原有 run 元素
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)

    # 从模板复制第一个 run 的格式，设置新文字
    first_run = template_el.find(qn('w:r'))
    if first_run is not None:
        new_r = deepcopy(first_run)
        for t in new_r.findall(qn('w:t')):
            t.text = new_text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_p.append(new_r)

    # 插入到文档
    if anchor is not None:
        anchor._element.addnext(new_p)
    else:
        doc.element.body.append(new_p)


# ---------------------------------------------------------------------------
# 表格处理
# ---------------------------------------------------------------------------

def process_table_cells(
    doc: Document,
    cell_fn: Callable[[str], str],
) -> None:
    """遍历所有表格单元格，对每个单元格文字执行 cell_fn 变换。

    Args:
        doc: 文档对象
        cell_fn: 接收单元格纯文本，返回变换后的文本
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.runs:
                        old = p.runs[0].text if p.runs[0].text else ''
                        new = cell_fn(old)
                        if new != old:
                            rewrite_paragraph(p, new)


def get_table_cell_text(table, row_idx: int, col_idx: int) -> str:
    """获取表格指定单元格的纯文本。"""
    cell = table.rows[row_idx].cells[col_idx]
    return ''.join(r.text for p in cell.paragraphs for r in p.runs)


def set_table_cell_text(table, row_idx: int, col_idx: int, text: str) -> None:
    """设置表格指定单元格的文本（保留格式）。"""
    cell = table.rows[row_idx].cells[col_idx]
    for p in cell.paragraphs:
        if p.runs:
            rewrite_paragraph(p, text)
            return
    # 如果没有 run，则添加
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.runs[0].text = text if p.runs else ...


# ---------------------------------------------------------------------------
# 段落实用工具
# ---------------------------------------------------------------------------

def is_paragraph_empty(paragraph) -> bool:
    """判断段落是否为空（无文字内容）。"""
    return not paragraph.text.strip()


def get_paragraph_full_text(paragraph) -> str:
    """获取段落完整文本（拼接所有 run）。"""
    return ''.join(r.text for r in paragraph.runs)


# ---------------------------------------------------------------------------
# 验证
# ---------------------------------------------------------------------------

def verify_old_text_gone(doc: Document, old_text: str) -> int:
    """验证旧文字是否已全部清除，返回残留次数。"""
    count = 0
    for p in doc.paragraphs:
        if old_text in p.text:
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        count += 1
    return count


if __name__ == "__main__":
    print("preserve_format.py — 保格式编辑工具库")
    print("用法: from preserve_format import PreserveDoc, replace_all, ...")
    print()
    print("可用函数:")
    print("  copy_template / PreserveDoc — 复制模板创建副本")
    print("  replace_in_paragraph      — 单run内替换")
    print("  replace_cross_runs        — 跨run替换")
    print("  rewrite_paragraph         — 整段重写（保格式）")
    print("  replace_all               — 全文批量替换（段落+表格）")
    print("  insert_paragraph_deepcopy — deepcopy插入新段落")
    print("  process_table_cells       — 遍历表格单元格变换")
    print("  verify_old_text_gone      — 验证旧文字已清除")
