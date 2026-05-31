# -*- coding: utf-8 -*-
"""保格式编辑 vs 反模式 对比验证"""
import sys, os, shutil
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches
from preserve_format import copy_template, rewrite_paragraph, replace_all

out_dir = os.path.dirname(os.path.abspath(__file__))
tpl_path = os.path.join(out_dir, "test_template.docx")

# ---- 创建测试模板 ----
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Inches(0.35)

# 标题
title = doc.add_paragraph()
title.alignment = 1
tr = title.add_run('审计项目报告')
tr.font.name = 'Microsoft YaHei'
tr.font.size = Pt(22)
tr.bold = True

# 正文 p[1]
p = doc.add_paragraph()
p.add_run('本次审计覆盖{{项目名称}}，审计范围为2024年度财务报表。审计组进驻现场开展审计工作。')

# 正文 p[2] — 含有加粗的混合格式
p2 = doc.add_paragraph()
p2.add_run('根据审计发现，')
r = p2.add_run('合同管理环节存在重大缺陷')
r.bold = True
p2.add_run('，主要表现为合同签订流程不规范。')

# 表格
doc.add_paragraph('费用明细表：')
t = doc.add_table(rows=3, cols=3)
t.cell(0,0).text = '项目'; t.cell(0,1).text = '金额'; t.cell(0,2).text = '备注'
t.cell(1,0).text = '材料费'; t.cell(1,1).text = '{{金额}}'; t.cell(1,2).text = '待核实'
t.cell(2,0).text = '人工费'; t.cell(2,1).text = '{{工时}}'; t.cell(2,2).text = '待核实'

doc.save(tpl_path)
print(f"模板已创建: {tpl_path}")

# ---- 测试 1: 保格式模式 ----
print("\n=== 测试 1: 保格式编辑 ===")
doc1 = copy_template(tpl_path, os.path.join(out_dir, "test_preserve.docx"))
n = replace_all(doc1, '{{项目名称}}', '天府广场项目')
print(f"  replace_all 命中: {n}")

rewrite_paragraph(doc1.paragraphs[1],
    '本次审计覆盖天府广场项目，审计范围为2024年度财务报表。（已用rewrite_paragraph重写）')

p3_orig = doc1.paragraphs[2]  # 有加粗的段落（不修改它，验证保格式）
print(f"  段落[2]（原样保留）: {len(p3_orig.runs)} runs")

replace_all(doc1, '{{金额}}', '850,000.00')
replace_all(doc1, '{{工时}}', '320小时')
doc1.save(os.path.join(out_dir, "test_preserve.docx"))
print("  已保存: test_preserve.docx")

# ---- 测试 2: 反模式 ----
print("\n=== 测试 2: 反模式编辑 ===")
doc2 = copy_template(tpl_path, os.path.join(out_dir, "test_bad.docx"))
doc2.paragraphs[1].text = '本次审计覆盖天府广场项目，审计范围为2024年度财务报表。（已用paragraph.text=，格式破坏）'
doc2.save(os.path.join(out_dir, "test_bad.docx"))
print("  已保存: test_bad.docx")

# ---- 对比分析 ----
print("\n" + "="*60)
print("对比分析")
print("="*60)

doc1_check = Document(os.path.join(out_dir, "test_preserve.docx"))
doc2_check = Document(os.path.join(out_dir, "test_bad.docx"))

# 段落1 — 正文
p1_ok = doc1_check.paragraphs[1]
p1_bad = doc2_check.paragraphs[1]
print(f"\n【正文段落】")
print(f"  保格式: {len(p1_ok.runs)} runs")
for i, r in enumerate(p1_ok.runs):
    print(f"    run[{i}]: font={r.font.name}, size={r.font.size}, bold={r.bold}, text=\"{r.text[:50]}\"")
print(f"  反模式: {len(p1_bad.runs)} runs")
for i, r in enumerate(p1_bad.runs):
    print(f"    run[{i}]: font={r.font.name}, size={r.font.size}, bold={r.bold}, text=\"{r.text[:50]}\"")

# 段落2 — 含加粗
p2_ok = doc1_check.paragraphs[2]
p2_bad = doc2_check.paragraphs[2]
print(f"\n【含加粗段落】")
print(f"  保格式: {len(p2_ok.runs)} runs")
for i, r in enumerate(p2_ok.runs):
    print(f"    run[{i}]: bold={r.bold}, text=\"{r.text[:40]}\"")
print(f"  反模式: {len(p2_bad.runs)} runs")
for i, r in enumerate(p2_bad.runs):
    print(f"    run[{i}]: bold={r.bold}, text=\"{r.text[:40]}\"")

# 标题
t1 = doc1_check.paragraphs[0]
t2 = doc2_check.paragraphs[0]
print(f"\n【标题段落（未修改）】")
print(f"  保格式: {len(t1.runs)} runs, font={t1.runs[0].font.name}, size={t1.runs[0].font.size}, bold={t1.runs[0].bold}")
print(f"  反模式: {len(t2.runs)} runs, font={t2.runs[0].font.name}, size={t2.runs[0].font.size}, bold={t2.runs[0].bold}")

# 表格
for name, d in [("保格式", doc1_check), ("反模式", doc2_check)]:
    tbl = d.tables[0]
    c = tbl.rows[1].cells[1]
    full = ''.join(r.text for p in c.paragraphs for r in p.runs)
    print(f"\n【表格】{name}: 材料费金额 = \"{full}\"")

print("\n=== 验证完成 ===")
print("打开 test_preserve.docx 和 test_bad.docx 用 Word 对比查看。")
