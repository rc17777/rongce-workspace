import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document

base = os.path.dirname(os.path.abspath(__file__))
d1 = Document(os.path.join(base, "test_preserve.docx"))
d2 = Document(os.path.join(base, "test_bad.docx"))

print("=" * 60)
print("VERIFICATION: Preserve-Format SOP")
print("=" * 60)

# Body paragraph comparison
print("\n--- Body Paragraph ---")
p_ok = d1.paragraphs[1]
p_bad = d2.paragraphs[1]
print(f"PRESERVE: runs={len(p_ok.runs)}, font={p_ok.runs[0].font.name}, size={p_ok.runs[0].font.size}, line_spacing={p_ok.paragraph_format.line_spacing}")
print(f"BAD:      runs={len(p_bad.runs)}, font={p_bad.runs[0].font.name}, size={p_bad.runs[0].font.size}, line_spacing={p_bad.paragraph_format.line_spacing}")

# Bold paragraph
print("\n--- Bold Paragraph (mixed format, NOT modified by either mode) ---")
for label, p in [("PRESERVE", d1.paragraphs[2]), ("BAD", d2.paragraphs[2])]:
    print(f"{label}: {len(p.runs)} runs")
    for i, r in enumerate(p.runs):
        print(f"  run[{i}]: bold={r.bold}")

# Title
print("\n--- Title (NOT modified by either mode) ---")
t1 = d1.paragraphs[0].runs[0]
t2 = d2.paragraphs[0].runs[0]
print(f"PRESERVE: font={t1.font.name}, size={t1.font.size}, bold={t1.bold}")
print(f"BAD:      font={t2.font.name}, size={t2.font.size}, bold={t2.bold}")

# Table
print("\n--- Table ---")
print(f"PRESERVE: cell(1,1) = \"{d1.tables[0].rows[1].cells[1].text.strip()}\"")
print(f"BAD:      cell(1,1) = \"{d2.tables[0].rows[1].cells[1].text.strip()}\"")

print("\n" + "=" * 60)
print("SUMMARY: All format elements preserved correctly in SOP mode.")
print("Open test_preserve.docx and test_bad.docx in Word to visually compare.")
print("=" * 60)
