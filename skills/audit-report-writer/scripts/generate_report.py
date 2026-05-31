"""
审计报告Word文档生成器
生成标准格式的审计报告（专项审计/经济责任审计/管理咨询报告）
用法: 
  py generate_report.py --type special --config config.json --output report.docx
  py generate_report.py --type consult --config config.json --output report.docx
"""
import json
import argparse
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_cover(doc, config):
    """生成封面页"""
    # 空行
    for _ in range(4):
        doc.add_paragraph()

    # 委托单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.get('client', ''))
    run.font.size = Pt(18)
    run.font.name = '黑体'

    doc.add_paragraph()

    # 报告名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.get('title', ''))
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()
    doc.add_paragraph()

    # 报告信息
    info_items = [
        ('项目名称', config.get('project', '')),
        ('委托单位', config.get('client', '')),
        ('报告编号', config.get('report_no', '')),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 出具单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.get('firm', '四川融策会计师事务所有限公司'))
    run.font.size = Pt(16)
    run.font.name = '黑体'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.get('date', datetime.now().strftime('%Y年%m月%d日')))
    run.font.size = Pt(14)

    doc.add_page_break()


def add_section_title(doc, text, level=1):
    """添加章节标题"""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
    return p


def add_body_text(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
    p.paragraph_format.line_spacing = 1.5
    return p


def add_finding(doc, finding, index):
    """添加审计发现"""
    add_section_title(doc, f'问题{index}：{finding.get("title", "")}', level=2)

    sections = [
        ('问题描述', finding.get('description', '')),
        ('涉及金额/范围', finding.get('scope', '')),
        ('违反规定', finding.get('regulation', '')),
        ('原因分析', finding.get('cause', '')),
    ]

    for label, content in sections:
        if content:
            p = doc.add_paragraph()
            run_label = p.add_run(f'【{label}】')
            run_label.font.bold = True
            run_label.font.size = Pt(12)
            run_content = p.add_run(f' {content}')
            run_content.font.size = Pt(12)
            run_content.font.name = '宋体'
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

    # 整改建议
    suggestions = finding.get('suggestions', [])
    if suggestions:
        p = doc.add_paragraph()
        run_label = p.add_run('【整改建议】')
        run_label.font.bold = True
        run_label.font.size = Pt(12)
        for i, sug in enumerate(suggestions, 1):
            p_sug = doc.add_paragraph()
            run = p_sug.add_run(f'({i}) {sug}')
            run.font.size = Pt(12)
            run.font.name = '宋体'
            p_sug.paragraph_format.first_line_indent = Cm(0.74)

    # 空行分隔
    doc.add_paragraph()


def generate_special_audit(doc, config):
    """专项审计报告"""
    # 一、基本情况
    add_section_title(doc, '一、基本情况', level=1)
    basics = config.get('basics', {})
    if basics.get('basis'):
        add_body_text(doc, f'审计依据：{basics["basis"]}')
    if basics.get('scope'):
        add_body_text(doc, f'审计范围：{basics["scope"]}')
    if basics.get('method'):
        add_body_text(doc, f'审计方法：{basics["method"]}')
    if basics.get('overview'):
        add_body_text(doc, basics['overview'])

    # 二、审计发现
    add_section_title(doc, '二、审计发现', level=1)
    findings = config.get('findings', [])
    for i, finding in enumerate(findings, 1):
        add_finding(doc, finding, i)

    # 三、审计结论
    add_section_title(doc, '三、审计结论', level=1)
    conclusion = config.get('conclusion', '总体结论待补充。')
    add_body_text(doc, conclusion)

    # 四、整改建议
    add_section_title(doc, '四、整改建议', level=1)
    add_section_title(doc, '4.1 立即整改事项', level=2)
    immediate = config.get('immediate_suggestions', ['待补充。'])
    for sug in immediate:
        add_body_text(doc, f'• {sug}')

    add_section_title(doc, '4.2 限期整改事项', level=2)
    timed = config.get('timed_suggestions', ['待补充。'])
    for sug in timed:
        add_body_text(doc, f'• {sug}')

    add_section_title(doc, '4.3 持续改进建议', level=2)
    ongoing = config.get('ongoing_suggestions', ['待补充。'])
    for sug in ongoing:
        add_body_text(doc, f'• {sug}')


def generate_consulting(doc, config):
    """管理咨询报告"""
    # 一、项目背景
    add_section_title(doc, '一、项目背景', level=1)
    bg = config.get('background', {})
    if bg.get('purpose'):
        add_body_text(doc, f'委托事项：{bg["purpose"]}')
    if bg.get('scope'):
        add_body_text(doc, f'咨询范围：{bg["scope"]}')
    if bg.get('method'):
        add_body_text(doc, f'工作方法：{bg["method"]}')

    # 二、现状评估
    add_section_title(doc, '二、现状评估', level=1)
    assessment = config.get('assessment', '')
    add_body_text(doc, assessment)

    # 三、问题分析
    add_section_title(doc, '三、问题分析', level=1)
    findings = config.get('findings', [])
    for i, finding in enumerate(findings, 1):
        add_finding(doc, finding, i)

    # 四、改进建议
    add_section_title(doc, '四、改进建议', level=1)
    suggestions = config.get('consulting_suggestions', [])
    for i, sug in enumerate(suggestions, 1):
        add_section_title(doc, f'建议{i}：{sug.get("title", "")}', level=2)
        for key in ['target', 'steps', 'expected', 'investment', 'timeline']:
            if sug.get(key):
                labels = {
                    'target': '目标状态', 'steps': '实施步骤',
                    'expected': '预期效果', 'investment': '投入估算',
                    'timeline': '时间表'
                }
                p = doc.add_paragraph()
                run_l = p.add_run(f'【{labels[key]}】')
                run_l.font.bold = True
                run_c = p.add_run(f' {sug[key]}')
                run_c.font.size = Pt(12)

    # 五、实施路线图
    add_section_title(doc, '五、实施路线图', level=1)
    roadmap = config.get('roadmap', [])
    for phase in roadmap:
        add_body_text(doc, f'• {phase}')


def add_issue_checklist(doc, findings):
    """生成问题清单附件"""
    doc.add_page_break()
    add_section_title(doc, '附件1：问题清单', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['序号', '问题描述', '涉及金额', '违反法规', '整改建议', '责任部门']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for idx, f in enumerate(findings, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = f.get('title', '')
        row.cells[2].text = f.get('scope', '')
        row.cells[3].text = f.get('regulation', '')
        suggestions = f.get('suggestions', [])
        row.cells[4].text = '; '.join(suggestions) if suggestions else ''
        row.cells[5].text = f.get('responsible', '')


def main():
    parser = argparse.ArgumentParser(description='审计报告Word文档生成器')
    parser.add_argument('--type', required=True, choices=['special', 'consult'],
                        help='报告类型: special=专项审计, consult=管理咨询')
    parser.add_argument('--config', required=True, help='JSON配置文件路径')
    parser.add_argument('--output', default='audit_report.docx', help='输出文件路径')
    args = parser.parse_args()

    with open(args.config, 'r', encoding='utf-8') as f:
        config = json.load(f)

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 生成封面
    create_cover(doc, config)

    # 生成正文
    if args.type == 'special':
        generate_special_audit(doc, config)
    elif args.type == 'consult':
        generate_consulting(doc, config)

    # 生成问题清单
    findings = config.get('findings', [])
    if findings:
        add_issue_checklist(doc, findings)

    doc.save(args.output)
    print(f'✅ 审计报告已生成: {args.output}')


if __name__ == '__main__':
    main()
