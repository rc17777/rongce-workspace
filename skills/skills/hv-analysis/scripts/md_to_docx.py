#!/usr/bin/env python3
"""
Markdown → DOCX converter (fallback for when weasyprint is broken).
Usage: python md_to_docx.py input.md output.docx [--title "TITLE"] [--author "AUTHOR"]
"""

import sys
import os
import re
import argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import markdown


def md_to_docx(md_path, docx_path, title="Report", author="AI Assistant"):
    """Convert Markdown file to DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    # Add title page
    for _ in range(6):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 82, 118)
    
    doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run('横纵分析报告')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run(f'作者：{author}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Page break
    doc.add_page_break()
    
    # Now process the markdown content line by line
    lines = md_text.split('\n')
    in_table = False
    table_data = []
    in_code_block = False
    code_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(44, 62, 80)
                code_lines = []
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Tables
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Check if it's a separator line (---|---)
            if re.match(r'^\|[\s\-:]+\|$', line.strip()):
                i += 1
                continue
            
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
            
            # Check if next line continues the table
            if i + 1 < len(lines) and not (lines[i + 1].strip().startswith('|') and lines[i + 1].strip().endswith('|')):
                # Table ended, render it
                if table_data:
                    rows_count = len(table_data)
                    cols_count = max(len(r) for r in table_data) if table_data else 0
                    
                    table = doc.add_table(rows=rows_count, cols=cols_count)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < cols_count:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text
                                # Bold for header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                    # Blue background for header
                                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A5276"/>')
                                    cell._tc.get_or_add_tcPr().append(shading)
                    
                    doc.add_paragraph()  # spacing
                
                in_table = False
                table_data = []
            
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r'^#\s+', '', line))
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 82, 118)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r'^##\s+', '', line))
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 82, 118)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r'^###\s+', '', line))
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 82, 118)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r'^####\s+', '', line))
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 82, 118)
        elif line.strip() == '---':
            # Horizontal rule - skip in DOCX (add spacing)
            doc.add_paragraph()
        elif line.strip() == '':
            # Empty line - adds paragraph break in markdown
            pass
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            # Bullet point
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            text = re.sub(r'^[\*\-\+]\s+', '', line.strip())
            # Process bold markers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(10.5)
        elif re.match(r'^\d+[\.\、]', line.strip()):
            # Numbered list
            p = doc.add_paragraph()
            p.style = doc.styles['List Number']
            text = re.sub(r'^\d+[\.\、]\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(10.5)
        elif line.strip().startswith('> '):
            # Blockquote
            p = doc.add_paragraph()
            text = re.sub(r'^>\s+', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Regular paragraph
            text = line.strip()
            if text:
                # Process bold markers
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                # Remove remaining * for italic
                text = re.sub(r'\*(.*?)\*', r'\1', text)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(10.5)
                p.paragraph_format.first_line_indent = Pt(21)  # First line indent
        
        i += 1
    
    # Save
    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")
    return docx_path


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('output', help='Output DOCX file')
    parser.add_argument('--title', default='Report', help='Report title')
    parser.add_argument('--author', default='AI Assistant', help='Author name')
    args = parser.parse_args()
    
    md_to_docx(args.input, args.output, args.title, args.author)
