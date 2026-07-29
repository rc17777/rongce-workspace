"""
Report Parser (报告文档解析器)

Extracts tables from Word/PDF audit reports into the generic TableModel.
Delegates actual document parsing to officecli-docx (Skill) and pdf (tool),
then converts raw output into structured TableDocument.

Architecture:
  Input:  audit_report.docx | audit_report.pdf
  Output: TableDocument (structured digital twin)

Supported formats:
  - .docx:  via officecli-docx (Skill) → extract tables
  - .pdf:   via pdf tool → text extraction → table reconstruction
  - .xlsx:  direct via openpyxl (separate code path)

Design principle: this module does NO accounting interpretation.
All domain semantics are injected by DomainAdapter (Layer 2).
"""

from __future__ import annotations

import json
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Optional

from .table_model import (
    TableDocument,
    TableSheet,
    extract_table_from_rows,
    detect_number_format,
)


# ---------------------------------------------------------------------------
# Main Parser
# ---------------------------------------------------------------------------

class ReportParser:
    """
    Parse audit report documents into structured TableDocument.
    """

    def __init__(self, source_path: str):
        self.source_path = Path(source_path)
        if not self.source_path.exists():
            raise FileNotFoundError(f"Source file not found: {source_path}")
        self.suffix = self.source_path.suffix.lower()

    def parse(self) -> TableDocument:
        """Parse the source document into a TableDocument."""
        if self.suffix == ".docx":
            return self._parse_docx()
        elif self.suffix == ".pdf":
            return self._parse_pdf()
        elif self.suffix in (".xlsx", ".xls"):
            return self._parse_xlsx()
        elif self.suffix == ".json":
            return self._parse_json_model()
        else:
            raise ValueError(f"Unsupported file format: {self.suffix}")

    # ---- DOCX Parsing ----

    def _parse_docx(self) -> TableDocument:
        """
        Parse .docx via officecli-docx (Skill).

        Strategy: use officecli batch to extract all tables,
        then convert to TableModel.
        """
        from .table_model import load_document

        # First try: if a pre-built model JSON exists alongside the docx
        json_path = self.source_path.with_suffix(".tables.json")
        if json_path.exists():
            try:
                return load_document(str(json_path))
            except Exception:
                pass

        # Fallback: extract tables via python-docx directly
        try:
            return self._extract_docx_tables_direct()
        except ImportError:
            pass

        # Last resort: empty document with metadata
        return TableDocument(
            source_path=str(self.source_path),
            source_type="docx",
            metadata={"parse_status": "needs_external_extraction"},
        )

    def _extract_docx_tables_direct(self) -> TableDocument:
        """Extract tables directly using python-docx (if installed)."""
        import docx

        doc = docx.Document(str(self.source_path))
        doc_model = TableDocument(
            source_path=str(self.source_path),
            source_type="docx",
        )

        # Extract all tables
        for i, table in enumerate(doc.tables):
            raw_rows = []
            for row in table.rows:
                raw_rows.append([cell.text for cell in row.cells])

            if not raw_rows:
                continue

            # Try to find a caption from preceding paragraphs
            caption = ""
            # Simple heuristic: look at preceding paragraph text

            sheet = extract_table_from_rows(
                name=f"Table_{i + 1}",
                raw_rows=raw_rows,
                caption=caption,
            )
            doc_model.sheets.append(sheet)

        # Extract paragraphs for context (useful for AI review)
        doc_model.metadata["paragraphs"] = [
            p.text for p in doc.paragraphs if p.text.strip()
        ]

        return doc_model

    # ---- PDF Parsing ----

    def _parse_pdf(self) -> TableDocument:
        """
        Parse .pdf document.

        PDF table extraction is inherently heuristic. The output should be
        reviewed by a domain adapter to correct classification errors.
        """
        doc_model = TableDocument(
            source_path=str(self.source_path),
            source_type="pdf",
        )

        # PDF parsing requires external tool (see pdf tool in OpenClaw)
        # For now, create a placeholder with metadata
        doc_model.metadata["parse_status"] = "needs_pdf_extraction"
        doc_model.metadata["note"] = (
            "PDF table extraction requires external processing. "
            "Use OpenClaw's pdf tool or a dedicated PDF table extractor."
        )

        return doc_model

    # ---- XLSX Parsing ----

    def _parse_xlsx(self) -> TableDocument:
        """Parse .xlsx workbook into TableDocument (one sheet per tab)."""
        try:
            import openpyxl
        except ImportError:
            return TableDocument(
                source_path=str(self.source_path),
                source_type="xlsx",
                metadata={"parse_status": "needs_openpyxl"},
            )

        wb = openpyxl.load_workbook(str(self.source_path), data_only=True)
        doc_model = TableDocument(
            source_path=str(self.source_path),
            source_type="xlsx",
        )

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Detect data range
            raw_rows = []
            for row in ws.iter_rows(values_only=True):
                raw_rows.append([str(c) if c is not None else "" for c in row])

            if not raw_rows:
                continue

            # Trim empty trailing rows
            while raw_rows and all(c == "" for c in raw_rows[-1]):
                raw_rows.pop()

            sheet = extract_table_from_rows(
                name=sheet_name,
                raw_rows=raw_rows,
                page_ref=sheet_name,
                caption=sheet_name,
            )
            doc_model.sheets.append(sheet)

        return doc_model

    # ---- JSON Model (pre-built) ----

    def _parse_json_model(self) -> TableDocument:
        """Load a pre-built TableDocument from JSON."""
        from .table_model import load_document
        return load_document(str(self.source_path))

    # ---- Convenience ----

    @staticmethod
    def from_raw_data(
        tables: dict[str, list[list[str]]],
        source_type: str = "manual",
    ) -> TableDocument:
        """
        Build a TableDocument from programmatically-supplied raw tables.
        
        Args:
            tables: {sheet_name: [[cell, cell, ...], ...]}
            source_type: label for the source
        """
        doc = TableDocument(source_type=source_type)

        for name, raw_rows in tables.items():
            sheet = extract_table_from_rows(name=name, raw_rows=raw_rows)
            doc.sheets.append(sheet)

        return doc


# ---------------------------------------------------------------------------
# CLI Entry Point
# ---------------------------------------------------------------------------

def main():
    """CLI for testing the parser."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python report_parser.py <report.docx|report.xlsx|report.pdf>")
        sys.exit(1)

    parser = ReportParser(sys.argv[1])
    doc = parser.parse()

    print(f"Source: {doc.source_path}")
    print(f"Sheets: {len(doc.sheets)}")
    for sheet in doc.sheets:
        print(f"  - {sheet.name}: {len(sheet.rows)} rows, {len(sheet.columns)} cols")

    # Save model
    output_path = Path(sys.argv[1]).with_suffix(".tables.json")
    from .table_model import save_document
    save_document(doc, str(output_path))
    print(f"Model saved to: {output_path}")


if __name__ == "__main__":
    main()
