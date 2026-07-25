# -*- coding: utf-8 -*-
import os, sys, re
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = r'C:\Users\scrccpa\.openclaw\workspace\output\新制度体系'
DESK = os.path.join(os.path.expanduser("~"), "Desktop")

FH = '黑体'
FF = '仿宋'
FS = '宋体'
SE = Pt(22)
SS = Pt(16)
SI = Pt(14)
SW = Pt(10.5)

DOCS = [
    ("00","00-制度体系架构.md","RC-SYS-000","制度体系架构","总纲"),
    ("01","01-薪酬管理制度.md","RC-HR-001","薪酬管理制度","人力资源"),
    ("02","02-绩效考核管理制度.md","RC-HR-002","绩效考核管理制度","人力资源"),
    ("03","03-员工手册.md","RC-HR-003","员工手册","人力资源"),
    ("04","04-项目管理规范.md","RC-BIZ-001","项目管理规范","业务质控"),
    ("05","05-制度发布与版本管理规范.md","RC-ADM-006","制度发布与版本管理规范","行政综合"),
    ("06","06-财务报销管理制度.md","RC-FIN-001","财务报销管理制度","财务管理"),
    ("07","07-审计质量控制制度.md","RC-QC-001","审计质量控制制度","业务质控"),
    ("08","08-造价咨询质量控制制度.md","RC-QC-002","造价咨询质量控制制度","业务质控"),
    ("09","09-股东会议事规则.md","RC-GOV-003","股东会议事规则","行政综合"),
    ("10","10-招聘与入职管理制度.md","RC-HR-004","招聘与入职管理制度","人力资源"),
    ("11","11-培训与发展管理制度.md","RC-HR-005","培训与发展管理制度","人力资源"),
    ("12","12-职级晋升管理制度.md","RC-HR-006","职级晋升管理制度","人力资源"),
    ("13","13-项目收入确认与回款管理制度.md","RC-FIN-002","项目收入确认与回款管理制度","财务管理"),
    ("14","14-预算管理制度.md","RC-FIN-003","预算管理制度","财务管理"),
    ("15","15-资金管理制度.md","RC-FIN-004","资金管理制度","财务管理"),
    ("16","16-固定资产管理制度.md","RC-FIN-005","固定资产管理制度","财务管理"),
    ("17","17-业务承接与合同管理制度.md","RC-BIZ-002","业务承接与合同管理制度","业务质控"),
    ("18","18-客户关系管理制度.md","RC-BIZ-003","客户关系管理制度","业务质控"),
    ("19","19-业务分包管理制度.md","RC-BIZ-004","业务分包管理制度","业务质控"),
    ("20","20-投标管理制度.md","RC-BIZ-005","投标管理制度","业务质控"),
    ("21","21-三级复核实施细则.md","RC-QC-003","三级复核实施细则","业务质控"),
    ("22","22-执业责任追究制度.md","RC-QC-004","执业责任追究制度","业务质控"),
    ("23","23-信息安全与保密管理制度.md","RC-ADM-002","信息安全与保密管理制度","行政综合"),
    ("24","24-印章与证照管理制度.md","RC-ADM-004","印章与证照管理制度","行政综合"),
    ("25","25-档案管理制度.md","RC-ADM-005","档案管理制度","行政综合"),
    ("26","26-公司章程-会计师事务所.md","RC-GOV-001","公司章程-会计师事务所","行政综合"),
    ("27","27-公司章程-工程咨询公司.md","RC-GOV-002","公司章程-工程咨询公司","行政综合"),
    ("28","28-办公场所管理制度.md","RC-ADM-001","办公场所管理制度","行政综合"),
    ("29","29-采购管理制度.md","RC-ADM-003","采购管理制度","行政综合"),
    ("30","30-数智化建设管理制度.md","RC-SPL-001","数智化建设管理制度","行政综合"),
    ("31","31-业务拓展与创新管理制度.md","RC-SPL-002","业务拓展与创新管理制度","行政综合"),
    ("32","32-风险管理制度.md","RC-SPL-003","风险管理制度","行政综合"),
    ("33","33-党建工作制度.md","RC-SPL-004","党建工作制度","行政综合"),
]
ASMS = {
    "人力资源篇": ["00","01","02","03","10","11","12"],
    "财务管理篇": ["06","13","14","15","16"],
    "业务质控篇": ["04","07","08","17","18","19","20","21","22"],
    "行政综合篇": ["05","09","23","24","25","26","27","28","29","30","31","32","33"],
}
def mods(t, d):
    if d=="02":
        t=t.replace("D等比例不低于5%（强制分布，避免老好人现象）","D等比例参考值为5%左右（指导性比例，非强制分布。如全员达标可无D等）")
        t=t.replace("360度协作评分","部门负责人协作评分"); t=t.replace("360度评价","部门负责人评价"); t=t.replace("360度评分","部门负责人评分")
        t=t.replace("| 季度考核 | 每季度一次 | 季度结束后10个工作日内完成 | 绩效工资清算 |","| 半年度考核 | 每半年一次 | 考核期结束后5个工作日内完成 | 绩效工资清算 |")
        t=t.replace("季度考核","半年度考核"); t=t.replace("连续两个季度考核为D等","连续两次半年度考核为D等")
        t=t.replace("第1~5个工作日：员工填写自评表","第1个工作日：员工填写自评表")
        t=t.replace("第6~8个工作日：项目经理/部门负责人评分","第2~3个工作日：部门负责人评分")
        t=t.replace("第9~10个工作日：人力资源部汇总统计","第4个工作日：人力资源部汇总统计")
        t=t.replace("第11~12个工作日：部门负责人进行绩效面谈\n第13~15个工作日：员工确认签字（如有异议启动申诉）\n第16~20个工作日：申诉处理（如有）","第5个工作日：绩效面谈、结果确认（如有异议启动申诉）")
        t=t.replace("**第四条** 本制度适用于公司全体在职员工（试用期员工参加考核但不与绩效工资挂钩，作为转正参考）。","**第四条** 本制度适用于公司全体在职员工（试用期员工参加考核但不与绩效工资挂钩，作为转正参考）。\n\n**第四条之一** 考核保护条款：\n1. 产假/流产假期间女员工不参加当期考核，绩效工资按产假前标准发放\n2. 病假连续缺勤超考核周期50%的员工不参加当期考核，绩效工资按实际出勤折算\n3. 工伤停工留薪期内员工不参加当期考核，待遇按《工伤保险条例》执行\n4. 上述保护期内员工不纳入等级分布基数统计")
        t=t.replace("4. 改进期满仍为D的，可依法协商解除劳动合同","4. 改进期满或调岗后仍为D的，依据《劳动合同法》第四十条处理：提前30日书面通知或支付代通知金，依法支付经济补偿金，事先通知工会")
        t=t.replace("3. 提供必要的培训或调岗支持","3. 提供必要的培训或调整工作岗位（调岗须协商一致）")
        t=t.replace("**第二十六条** 本制度经管理层审议通过后发布施行。","**第二十六条** 本制度经职工代表大会（或全体员工大会）讨论通过，经管理层批准后发布施行。")
        t=t.replace("| 团队协作 | 10% | 360度评价 |","| 团队协作 | 10% | 部门负责人评价 |")
    elif d=="03":
        t=t.replace("4. 考核不合格的，公司可延长试用期（总试用期不超过法定上限）或依法解除劳动合同","4. 考核不合格的，视为不符合录用条件，依据《劳动合同法》第三十九条第（一）项解除，需有书面考核记录")
        t=t.replace("- 晚婚（男25周岁/女23周岁以上初婚）增加7天（参照四川省规定）\n","")
        t=t.replace("**第四十五条** 公司不监控员工个人通讯内容。因业务需要，公司可以对工作邮箱、工作电脑中的工作文件进行合理管理，但应事先告知员工。","**第四十五条** 公司不监控员工个人通讯内容。因保护商业秘密和信息安全，公司可对工作邮箱和工作电脑中的工作文件进行合理管理。目的限于保护客户信息安全、防止商业秘密泄露。范围限于工作文件不涉及个人通讯。应事先书面告知员工。")
        t=t.replace("**第四十八条** 本手册经民主程序（职工代表大会或全体员工讨论）通过后发布施行。","**第四十八条** 本手册经民主程序通过后发布施行。形式：召开全体员工大会或职工代表大会，充分听取员工意见，修改完善后经全体员工签字确认。")
    elif d=="06":
        t=t.replace("3. 专票应在认证期限内（360天）送交财务部门认证抵扣","3. 专票应及时送交财务部门认证抵扣（不迟于次年企业所得税汇算清缴截止日）")
        t=t.replace("本制度由财务部负责解释和修订，经总经理批准后施行。","本制度由财务部负责解释和修订。费用标准的调整须经职工代表大会（或全体员工）讨论后，经总经理批准后施行。")
    elif d=="07":
        t=t.replace("4. 需要向监管部门报告的，由质控负责人决定并组织实施","4. 法律法规要求向监管部门报告的，应当依法及时报告，由质控负责人组织实施，不得隐瞒不报")
        t=t.replace("扣减当年全部绩效工资，降级使用","扣减绩效工资（每月不超过当月工资20%），降级使用"); t=t.replace("扣减当年50%绩效工资","扣减绩效工资（每月不超过当月工资20%）")
        t=t.replace("扣减当季绩效工资的50%~100%","扣减绩效工资（每月不超过当月工资20%）"); t=t.replace("扣减当季绩效工资的30%","扣减绩效工资（每月不超过当月工资15%）"); t=t.replace("扣减当季绩效工资的20%","扣减绩效工资（每月不超过当月工资10%）")
    elif d=="08":
        t=t.replace("扣减当年全部绩效工资，降级使用","扣减绩效工资（每月不超过当月工资20%），降级使用"); t=t.replace("扣减当年50%绩效工资","扣减绩效工资（每月不超过当月工资20%）")
        t=t.replace("扣减当季绩效工资的50%~100%","扣减绩效工资（每月不超过当月工资20%）"); t=t.replace("扣减当季绩效工资的30%","扣减绩效工资（每月不超过当月工资15%）"); t=t.replace("扣减当季绩效工资的20%","扣减绩效工资（每月不超过当月工资10%）")
    elif d=="10":
        t=t.replace("- 基本合格：可延长试用期（累计不超过法定上限），明确改进目标；","- 基本合格：明确改进目标给予额外辅导；仍不符合要求的在试用期内依法解除；")
    elif d=="13":
        t=t.replace("5. **形成坏账的**：对应项目奖金不予发放，已发放的予以追回。","5. **形成坏账的**：对应项目奖金尚未发放的不再发放；已经发放的不予追回。")
    elif d=="22":
        t=t.replace("扣减全年绩效奖金100%","扣减绩效工资（每月不超过当月工资20%）"); t=t.replace("扣减全年绩效奖金50%以上","扣减绩效工资（每月不超过当月工资20%）")
        t=t.replace("扣减全年绩效奖金30%~50%","扣减绩效工资（每月不超过当月工资15%）"); t=t.replace("扣减季度绩效奖金20%~30%","扣减绩效工资（每月不超过当月工资10%）"); t=t.replace("扣减全年绩效奖金50%","扣减绩效工资（每月不超过当月工资20%）")
        t=t.replace("③降级（降低岗位等级）","③调岗或降级（须经员工协商一致）"); t=t.replace("③降级或调岗","③调岗或降级（须经员工协商一致）")
        t=t.replace("**第五条** 多人共同负责的，根据各人的过错程度和职责分工确定各自责任比例。受上级指令违规操作的，执行人承担执行责任，指令人承担主要责任。","**第五条** 多人共同负责的，根据各人的过错程度和职责分工确定各自责任比例。受上级指令违规操作的，指令人承担主要责任。\n\n**第五条之一** 员工有权拒绝执行违反法律法规、职业道德准则的指令（如出具虚假报告等）。因拒绝违法指令遭受不利对待的，公司应予纠正并追责。")
    elif d=="12":
        t=t.replace("降级使用","调岗或降级（须经员工协商一致）")
    elif d=="14":
        t=t.replace("| 11月上旬 | 财务部发布下一年度预算编制指引 | 财务部 |\n| 11月中旬 | 各部门编制部门预算草案 | 各部门负责人 |\n| 11月下旬 | 财务部汇总审核，编制公司预算草案 | 财务部 |\n| 12月上旬 | 管理层审议预算草案 | 总经理+各部门负责人 |\n| 12月中旬 | 修改完善，形成正式预算方案 | 财务部 |\n| 12月下旬 | 总经理审批 / 股东会审批（重大预算） | 总经理/股东会 |\n| 12月底 | 下达各部门执行 | 财务部 |","| 第1天 | 财务部发布预算编制指引 | 财务部 |\n| 第2~3天 | 各部门编制预算草案 | 各部门负责人 |\n| 第4天 | 财务部汇总编制公司预算草案 | 财务部 |\n| 第5天 | 管理层审议批准 | 总经理 |\n| 第5天 | 下达各部门执行 | 财务部 |")
        t=t.replace("1. **增量预算法**：以历史数据为基础，结合业务增长预期进行调整（适用于稳定业务）；\n2. **零基预算法**：不参考历史数据，逐项论证支出必要性（适用于新项目、新业务）；\n3. **弹性预算法**：设定不同业务量情景下的预算方案（适用于收入不确定性较大的情况）。","采用**增量预算法**：以上年度实际收支为基础结合业务增长预期调整。公式：新年度预算=上年度实际数×（1±调整比例）。调整比例由财务部建议管理层审定。")
    elif d=="04":
        t=t.replace("| C类（小型） | <10万元 | 项目经理或高级审计员负责 | 二级复核 |","| C类（小型） | <10万元 | 项目经理或高级审计员负责 | 二级复核（项目经理+部门负责人） |")
    return t

def sf(run, fn, sz, b=False):
    run.font.name=fn; run.font.size=sz; run.bold=b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), fn)

def mkdoc():
    d=Document()
    s=d.sections[0]
    s.page_width=Cm(21); s.page_height=Cm(29.7)
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2)
    s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)
    s.header_distance=Cm(1.5); s.footer_distance=Cm(1.5)
    st=d.styles['Normal']
    st.font.name=FF; st.font.size=SI
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FF)
    st.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    return d

def hdr(sec, txt):
    h=sec.header; h.is_linked_to_previous=False
    p=h.paragraphs[0]; p.text=txt; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: r.font.size=Pt(9); r.font.name=FS; r.element.rPr.rFonts.set(qn('w:eastAsia'),FS)

def ftr(sec):
    f=sec.footer; f.is_linked_to_previous=False
    p=f.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r0=p.add_run('-- '); r0.font.size=Pt(9); r0.font.name=FS
    r0.element.rPr.rFonts.set(qn('w:eastAsia'),FS)
    p.add_run()._r.append(parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>'))
    p.add_run()._r.append(parse_xml('<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve"> PAGE </w:instrText>'))
    p.add_run()._r.append(parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>'))
    r1=p.add_run(' --'); r1.font.size=Pt(9); r1.font.name=FS
    r1.element.rPr.rFonts.set(qn('w:eastAsia'),FS)

def cover(d, t1, t2, sub, ver, dt):
    for _ in range(6):
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(0)
    for t in [t1,t2]:
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(t); sf(r,FH,Pt(26),True)
    d.add_paragraph()
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(sub); sf(r,FH,Pt(36),True)
    for _ in range(3): d.add_paragraph()
    for t in [f'版本：{ver}',dt]:
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(t); sf(r,FF,SS)
    d.add_page_break()

def toc(d):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('目    录'); sf(r,FH,SE,True)
    d.add_paragraph()
    p=d.add_paragraph()
    p.add_run()._r.append(parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>'))
    p.add_run()._r.append(parse_xml('<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
    p.add_run()._r.append(parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="separate"/>'))
    r=p.add_run('（请在Word中右键此处选择更新域以生成目录）'); r.font.color.rgb=RGBColor(128,128,128)
    p.add_run()._r.append(parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>'))
    d.add_page_break()

def parse_table(lines):
    rows=[]
    for i,line in enumerate(lines):
        if i==1: continue
        cells=[c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows

def add_tbl(d, rows):
    if not rows: return
    ncols=max(len(r) for r in rows)
    for r in rows:
        while len(r)<ncols: r.append('')
    tbl=d.add_table(rows=len(rows), cols=ncols)
    tbl.style='Table Grid'; tbl.alignment=1
    for i,row_data in enumerate(rows):
        for j,cell_text in enumerate(row_data):
            cell=tbl.cell(i,j); cell.text=''
            p=cell.paragraphs[0]
            p.paragraph_format.space_before=Pt(1)
            p.paragraph_format.space_after=Pt(1)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
            clean=re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            r=p.add_run(clean)
            if i==0:
                sf(r,FH,SW,True); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                shading=parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9E2F3" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                sf(r,FF,SW)
    d.add_paragraph()

def render(d, etype, content):
    if etype=='heading':
        level,text=content
        clean=re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        if level==1:
            p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(24); p.paragraph_format.space_after=Pt(12)
            r=p.add_run(clean); sf(r,FH,SE,True)
        elif level==2:
            p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(8)
            r=p.add_run(clean); sf(r,FH,SS,True)
        elif level==3:
            p=d.add_paragraph()
            p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)
            r=p.add_run(clean); sf(r,FH,SI,True)
        else:
            p=d.add_paragraph()
            p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(3)
            r=p.add_run(clean); sf(r,FH,SI)

    elif etype=='paragraph':
        p=d.add_paragraph()
        p.paragraph_format.first_line_indent=Cm(0.74)
        p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
        parts=re.split(r'(\*\*.*?\*\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r=p.add_run(part[2:-2]); sf(r,FF,SI,True)
            elif part:
                r=p.add_run(part); sf(r,FF,SI)

    elif etype=='table':
        rows=parse_table(content); add_tbl(d, rows)

    elif etype=='code':
        for line in content.split('\n'):
            p=d.add_paragraph()
            p.paragraph_format.left_indent=Cm(1)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
            r=p.add_run(line); r.font.name='Consolas'; r.font.size=Pt(10)

    elif etype=='blockquote':
        for line in content.split('\n'):
            if not line.strip(): continue
            p=d.add_paragraph()
            p.paragraph_format.left_indent=Cm(1)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
            parts=re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r=p.add_run(part[2:-2]); sf(r,FF,SW,True)
                elif part:
                    r=p.add_run(part); sf(r,FF,SW)

    elif etype=='ul':
        for indent,text in content:
            p=d.add_paragraph()
            p.paragraph_format.left_indent=Cm(0.74+indent*0.3)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
            px='* ' if indent==0 else '  - '
            parts=re.split(r'(\*\*.*?\*\*)', text)
            first=True
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    t=(px+part[2:-2]) if first else part[2:-2]
                    r=p.add_run(t); sf(r,FF,SI,True)
                elif part:
                    t=(px+part) if first else part
                    r=p.add_run(t); sf(r,FF,SI)
                first=False

    elif etype=='ol':
        for idx,(indent,text) in enumerate(content,1):
            p=d.add_paragraph()
            p.paragraph_format.left_indent=Cm(0.74+indent*0.3)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
            px=f'{idx}. '
            parts=re.split(r'(\*\*.*?\*\*)', text)
            first=True
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    t=(px+part[2:-2]) if first else part[2:-2]
                    r=p.add_run(t); sf(r,FF,SI,True)
                elif part:
                    t=(px+part) if first else part
                    r=p.add_run(t); sf(r,FF,SI)
                first=False

    elif etype=='hr':
        pass


def parse_md(text):
    lines=text.split('\n'); elements=[]; pos=0
    while pos<len(lines):
        line=lines[pos]
        if line.strip()=='': pos+=1; continue
        if line.strip() in('---','***','___'): elements.append(('hr','')); pos+=1; continue
        m=re.match(r'^(#{1,6})\s+(.*)', line)
        if m: elements.append(('heading',(len(m.group(1)), m.group(2).strip()))); pos+=1; continue
        if '|' in line and pos+1<len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[pos+1]):
            tl=[]
            while pos<len(lines) and '|' in lines[pos] and lines[pos].strip(): tl.append(lines[pos]); pos+=1
            elements.append(('table', tl)); continue
        if line.strip().startswith('```'):
            pos+=1; cl=[]
            while pos<len(lines) and not lines[pos].strip().startswith('```'): cl.append(lines[pos]); pos+=1
            if pos<len(lines): pos+=1
            elements.append(('code', '\n'.join(cl))); continue
        if line.strip().startswith('>'):
            ql=[]
            while pos<len(lines) and lines[pos].strip().startswith('>'): ql.append(lines[pos].strip().lstrip('>').strip()); pos+=1
            elements.append(('blockquote', '\n'.join(ql))); continue
        m_ul=re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m_ul:
            items=[]
            while pos<len(lines):
                mi=re.match(r'^(\s*)[-*]\s+(.*)', lines[pos])
                if mi: items.append((len(mi.group(1)), mi.group(2))); pos+=1
                elif lines[pos].strip()=='': pos+=1; break
                else: break
            elements.append(('ul', items)); continue
        m_ol=re.match(r'^(\s*)\d+[.、]\s*(.*)', line)
        if m_ol:
            items=[]
            while pos<len(lines):
                mi=re.match(r'^(\s*)\d+[.、]\s*(.*)', lines[pos])
                if mi: items.append((len(mi.group(1)), mi.group(2))); pos+=1
                elif lines[pos].strip()=='': pos+=1; break
                else: break
            elements.append(('ol', items)); continue
        pl=[]
        while pos<len(lines):
            l=lines[pos]
            if l.strip()=='' or l.strip() in('---','***','___') or l.strip().startswith('#') or l.strip().startswith('```') or l.strip().startswith('>') or re.match(r'\s*[-*]\s+',l) or re.match(r'\s*\d+[.、]\s*',l): break
            if '|' in l and pos+1<len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[pos+1]): break
            pl.append(l); pos+=1
        if pl: elements.append(('paragraph', ' '.join(pl)))
    return elements

def gen_one(doc_ids, outname, title_sub, version_str="V1.0"):
    print(f'Generating: {outname}')
    d=mkdoc()
    sec=d.sections[0]
    hdr(sec, '四川融策会计师/工程咨询有限公司')
    ftr(sec)
    
    cover(d, '四川融策会计师事务所有限公司', '四川融策工程咨询有限公司', title_sub, version_str, '2026年7月')
    toc(d)
    
    for did in doc_ids:
        info=[x for x in DOCS if x[0]==did]
        if not info: continue
        _, fname, dcode, dname, dgroup = info[0]
        fpath=os.path.join(SRC, fname)
        if not os.path.exists(fpath):
            print(f'  WARN: {fname} not found')
            continue
        with open(fpath, 'r', encoding='utf-8') as f:
            text=f.read()
        
        # Apply modifications
        text=mods(text, did)
        
        # Parse and render
        elements=parse_md(text)
        
        # Add doc header
        p=d.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(f'{dcode}  {dname}')
        sf(r,FH,SE,True)
        d.add_page_break()
        
        for etype, content in elements:
            render(d, etype, content)
        
        # Page break between docs (except last)
        if did!=doc_ids[-1]:
            d.add_page_break()
        
        print(f'  OK: {dcode} {dname}')
    
    outpath=os.path.join(DESK, outname)
    d.save(outpath)
    print(f'Saved: {outpath}')
    return outpath

if __name__=='__main__':
    all_ids=[d[0] for d in DOCS]
    
    print('=== Generating Complete Version ===')
    gen_one(all_ids, '融策公司制度体系（完整版）.docx', '制 度 体 系')
    
    for name, ids in ASMS.items():
        print(f'=== Generating Assembly: {name} ===')
        gen_one(ids, f'融策制度汇编-{name}.docx', f'制度汇编\\\\n{name}')
    
    print('=== ALL DONE ===')
