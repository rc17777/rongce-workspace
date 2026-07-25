import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

COMPLETE = r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx'

doc = Document(COMPLETE)

# Find the boundaries of the inserted 业务部管理篇 content
# Start: after 财务管理篇 (after 可分配利润核算细则)
# End: before 业务质控篇 section heading

body = doc.element.body

# Find 业务质控篇 section heading (the one deep in the doc, not in TOC)
biz_qc_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '业务质控篇' and i > 100:
        biz_qc_idx = i
        break

# Find end of 财务管理篇 - look for "四川融策可分配利润核算细则" last content
# The inserted content starts right after the 财务管理 section
# Find the first inserted element (PB + overview items start)
# The overview items have 【RC-BIZ-00X】 pattern

overview_start = None
overview_end = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('【RC-BIZ-00'):
        if overview_start is None:
            overview_start = i
        overview_end = i

print(f'Overview items: paragraphs {overview_start} to {overview_end}')
print(f'业务质控篇 section heading at: {biz_qc_idx}')

# The inserted content is from overview_start to biz_qc_idx-1
# But we need to include the page break before overview items
# Check if the paragraph before overview_start is empty (page break)
insert_start = overview_start
for i in range(overview_start - 1, overview_start - 5, -1):
    t = doc.paragraphs[i].text.strip()
    if not t:
        insert_start = i
    else:
        break

insert_end = biz_qc_idx - 1

print(f'Inserted content spans: paragraphs {insert_start} to {insert_end}')
print(f'Total paragraphs to fix: {insert_end - insert_start + 1}')

# Collect all paragraphs in the inserted range
inserted_paras = []
for i in range(insert_start, insert_end + 1):
    inserted_paras.append(doc.paragraphs[i]._element)

print(f'Collected {len(inserted_paras)} inserted paragraph elements')

# Now remove them all from the document
for elem in inserted_paras:
    body.remove(elem)

print('Removed all inserted paragraphs')

# Now re-insert in CORRECT order
# The correct order should be: page_break, overview_items(003→008), 003_content, 004_content, ..., 008_content
# But the paragraphs are currently in document order (003 first, 008 last) but with reversed chunks

# Since we collected them in document order (which happens to be reversed of what we want),
# and we know the original order was correct in the new doc, we just need to 
# insert them in NORMAL (non-reversed) order before 业务质控篇

target_elem = doc.paragraphs[biz_qc_idx - (insert_end - insert_start + 1)]._element
# Wait, after removal the indices shift. Let me re-find 业务质控篇.
biz_qc_elem = None
for p in doc.paragraphs:
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t == '业务质控篇' and p._element.getparent() is body:
        biz_qc_elem = p._element
        break

if biz_qc_elem is None:
    print('ERROR: Could not find 业务质控篇 after removal')
    sys.exit(1)

# Reverse the list so that addprevious inserts in correct order
# Remember: addprevious inserts immediately before target, pushing previous siblings back
# So to get [A, B, C, target], we need to insert in order: A, B, C
# But because each addprevious pushes earlier siblings:
#   addprevious(A) → [A, target]
#   addprevious(B) → [B, A, target]  (B pushes A back)
#   addprevious(C) → [C, B, A, target]  (C pushes B back)
# 
# So we need to insert in REVERSE order to get the correct final order
# But wait - the inserted_paras are already in the order they appear in the doc,
# which in the original new_doc was: PB, overview, RC-BIZ-003..., RC-BIZ-008...
# After addprevious with reversed, they became reversed.
# After removal and re-collection, they're in this reversed order.
# 
# If we now insert them in reverse order again:
# reversed(reversed_order) = original order
# Using addprevious:
#   - Original order is [RC-BIZ-008..., ..., RC-BIZ-003..., overview, PB]
#   - Reversed for insertion: [PB, overview, RC-BIZ-003..., ..., RC-BIZ-008]
#   - addprevious(PB) → [PB, target]
#   - addprevious(overview1) → [PB, overview1, target]
#   - ...
#   - addprevious(RC-BIZ-008_last) → [PB, overview1, ..., RC-BIZ-003..., ..., RC-BIZ-008_last, target]
# 
# YES! This should work. Insert in reversed order using addprevious.

for elem in reversed(inserted_paras):
    biz_qc_elem.addprevious(copy.deepcopy(elem))

print(f'Re-inserted {len(inserted_paras)} paragraphs in correct order')

# Save
doc.save(COMPLETE)

import os
print(f'Saved: {COMPLETE}')
print(f'Size: {os.path.getsize(COMPLETE):,} bytes')
