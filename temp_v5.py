# -*- coding: utf-8 -*-
"""PPP报告v5 - K因子修正版"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

CAPITAL_AMORT = 80_000_000.00; CAPITAL_LUMP = 22_298_077.79; RATE = 0.0799; YEARS = 18
K = 0.0799
A = CAPITAL_AMORT * RATE * (1+RATE)**YEARS / ((1+RATE)**YEARS - 1)

cr = CAPITAL_AMORT; cs = []
for y in range(YEARS):
    interest = cr * RATE; principal = A - interest
    if y == YEARS-1: principal = cr; interest = A - principal
    cr -= principal
    if cr < 0: cr = 0
    cs.append((principal, interest, A, max(cr, 0)))

bo = [(31686745.84,381300000),(26809671.03,381300000),(24811756.73,381250000),
      (21351219.32,380250000),(20263854.17,379250000),(20265833.33,378250000),
      (20162465.30,377250000),(20111770.82,376250000),(23022951.39,372250000),
      (22871874.99,368250000),(22617395.82,364250000),(22414618.06,360250000),
      (27148298.64,351250000),(27728124.99,341250000),(35070729.16,323250000),
      (34158229.18,305250000),(311968819.45,0),(0,0)]
sf = 710822119.32 / sum(b[0] for b in bo)
bs = [(round(a*sf,2), r) for a,r in bo]

# v5 FIX: apply K factor to 2025+ op costs
OP_RAW_2023 = 5341344.03; OP_RAW_2024_Y2 = 5948956.93; OP_RAW_2024_Y3 = 431456.22
OP_2023 = OP_RAW_2023 * (1+K)      # = 5,768,117.42
OP_2024 = (OP_RAW_2024_Y2 + OP_RAW_2024_Y3) * (1+K)  # = 6,890,208.16  (原文Y2已含K, Y3也含K, sum matches)
OP_2025 = 5713241.86 * (1+K)        # v5 FIX
OP_FUTURE = 5710000.00 * (1+K)      # v5 FIX
OP = [OP_2023, OP_2024, OP_2025] + [OP_FUTURE]*15

INC_2023 = 4285070.31; INC_2024 = 6229093.45; INC_2025 = 5564017.88; INC_FUTURE = 5560000
INC = [INC_2023, INC_2024, INC_2025] + [INC_FUTURE]*15

yl = list(range(2023,2041))
yd = []
for i in range(YEARS):
    pr, interest, pmt, rem = cs[i]
    lump = CAPITAL_LUMP if i==YEARS-1 else 0
    tc = pmt + lump; bk = bs[i][0]; op = OP[i]; inc = INC[i]
    avail = tc + bk + op - inc
    yd.append((yl[i], pr, interest, pmt, lump, tc, rem, bk, op, inc, avail))

gt = sum(y[10] for y in yd); ct = sum(y[5] for y in yd)
bt = sum(b[0] for b in bs); ot = sum(OP); it = sum(INC)

print(f"A={A:,.2f} Cap={ct:,.2f} Bank={bt:,.2f} Op={ot:,.2f} Inc={it:,.2f} Total={gt:,.2f} ({gt/1e8:.2f}yi)")

# ============================================================
doc = Document()
for s in doc.sections:
    s.page_width=Cm(21); s.page_height=Cm(29.7)
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2); s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)
sty=doc.styles['Normal']; sty.font.name='宋体'; sty.font.size=Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); sty.paragraph_format.line_spacing=1.5

def sc(cell, text, fn='宋体', sz=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=align
    r=p.add_run(text); r.font.name=fn; r.font.size=sz; r.font.bold=bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'),fn)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1
def scs(cell, c):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}"/>'))
def pa(text, fn='宋体', sz=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None, sa=Pt(0)):
    p=doc.add_paragraph(); p.alignment=align
    if indent: p.paragraph_format.first_line_indent=indent
    p.paragraph_format.space_after=sa; p.paragraph_format.line_spacing=1.5
    r=p.add_run(text); r.font.name=fn; r.font.size=sz; r.font.bold=bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'),fn)
    return p
def tb(headers, rows, cw=None):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; sc(c,h,'微软雅黑',Pt(9),True,color=RGBColor(0xFF,0xFF,0xFF)); scs(c,'0A1F3F')
    for i,rd in enumerate(rows):
        for j,v in enumerate(rd):
            c=t.rows[i+1].cells[j]; sc(c,str(v),'宋体',Pt(8.5))
            if i%2==1: scs(c,'F5F2EC')
    if cw:
        for row in t.rows:
            for j,w in enumerate(cw): row.cells[j].width=Cm(w)
    return t

# COVER
for _ in range(6): doc.add_paragraph()
pa('巴中市恩阳医养园PPP项目','方正小标宋简体',Pt(22),True,WD_ALIGN_PARAGRAPH.CENTER)
pa('可用性付费测算结果报告','方正小标宋简体',Pt(22),True,WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4): doc.add_paragraph()
pa('川融策咨询〔2025〕第490号',sz=Pt(14),align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6): doc.add_paragraph()
pa('四 川 融 策 会 计 师 事 务 所 有 限 公 司','微软雅黑',Pt(16),True,WD_ALIGN_PARAGRAPH.CENTER)
pa('Sichuan Rongce Accounting Firm Co., Ltd',sz=Pt(10),align=WD_ALIGN_PARAGRAPH.CENTER,sa=Pt(20))
pa('2026年7月29日',sz=Pt(14),align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# BODY
pa('川融策咨询〔2025〕第  号',align=WD_ALIGN_PARAGRAPH.RIGHT,sz=Pt(10))
doc.add_paragraph()
pa('巴中市恩阳医养园PPP项目','黑体',Pt(14),True,WD_ALIGN_PARAGRAPH.CENTER)
pa('可用性付费测算结果报告','黑体',Pt(14),True,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
pa('巴中市恩阳区卫生健康局：',indent=Cm(0.74))
pa('我们接受委托，对巴中市恩阳医养园PPP项目2022年10月28日至2040年12月21日可用性付费进行测算。本次测算以《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）审定的建筑安装工程费用485,098,077.79元为基数。现将测算结果报告如下（v5修订：修正2025年度起运维成本未按合同约定乘以(1+K)系数的问题）。',indent=Cm(0.74))

# 一
pa('一、项目概况','黑体',Pt(14),True)
pa('2015年12月21日，巴中市恩阳医养园PPP项目经巴中市恩阳区人民政府（恩府〔2015〕208号）批复同意实施。2016年12月，巴中市恩阳区卫生健康局与恩阳区人民医院、湖南省第五工程有限公司签订PPP协议，成立项目管理公司（注册资本金10,000万元，截至测算基准日实缴资本金102,298,077.79元）。',indent=Cm(0.74))
pa('2017年7月19日项目公司与湖南省第五工程有限公司签订施工合同，项目建筑总面积118,515.13\u33a1。项目实际于2017年5月17日开工，2022年3月完工，同年3月28日进入调试运营期，10月27日通过竣工验收，2022年10月28日正式进入运营阶段。',indent=Cm(0.74))
pa('根据2021年10月15日《巴中市恩阳医养园PPP项目合同之补充合同（2021年）》，项目合作期调整为23年，运营期为18年（2022年10月28日至2040年12月21日）。',indent=Cm(0.74))

# 二
pa('二、测算依据','黑体',Pt(14),True)
pa('（一）政策依据','楷体',Pt(12),True)
pa('（1）《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）',indent=Cm(0.74))
pa('（2）财政部关于印发《政府和社会资本合作项目财政承受能力论证指引》的通知（财金〔2015〕21号）',indent=Cm(0.74))
pa('（3）四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）',indent=Cm(0.74))
pa('（二）项目资料依据','楷体',Pt(12),True)
for i,item in enumerate([
    '《恩阳医养园PPP项目协议》（2016年12月27日签订）',
    '《巴中市恩阳医养园PPP项目合同》（2017年签订）',
    '《巴中市恩阳医养园PPP项目合同之补充合同》（2017年/2018年/2021年）',
    '《关于巴中市恩阳区人民医院一期工程建设项目可行性研究报告的批复》（恩区发改行审〔2015〕83号）',
    '《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）\u2014\u2014竣工结算额',
    '《借款还本付息情况报告》（贵阳银行股份有限公司成都分行出具）\u2014\u2014融资本息合计数',
    '项目公司2025年度运营成本/收入明细账（6册XLS，经本所独立验证）',
    '2022-2025年度运营期绩效考核结果（恩阳区卫健局出具）',
],1):
    pa(f'（{i}）{item}',indent=Cm(0.74))

# 三
pa('三、测算说明','黑体',Pt(14),True)
pa('（一）项目回报机制识别','楷体',Pt(12),True)
pa('PPP合同及实施方案明确本项目回报机制为"可行性缺口补助"。2016年12月27日《恩阳医养园PPP项目协议》约定区政府将可行性缺口补贴纳入跨年度财政预算并提请人大决议。',indent=Cm(0.74))

pa('（二）回报计算方法选用','楷体',Pt(12),True)
pa('根据2017年7月补充合同，计算公式：A = P\u00d7k\u00d7(1+k)^n / ((1+k)^n-1) + 实际融资成本 + 运营维护成本\u00d7(1+K) \u2212 第三方收入。合同约定K取值为社会资本中标年回报率7.99%。',indent=Cm(0.74))

pa('（1）项目资本金回报：P=实缴资本金102,298,077.79元；k=7.99%；n=18年。测算方案：8,000万元以等额本息方式在18年内逐年偿还（年还款额{:,.2f}元），剩余22,298,077.79元于第18年一次性支付、不计算回报利息。'.format(A),indent=Cm(0.74))

pa('（2）实际融资成本：贵阳银行成都分行《借款还本付息情况报告》确认18年融资本息合计710,822,119.32元（本所未取得逐年还款明细，逐年分配以征求意见稿表2还款结构为基础等比缩放，已标注取证层级）。合同利率自6.95%经三次下调至5.0%/年。',indent=Cm(0.74))

pa('（3）运营维护成本（含K=7.99%加成）：根据合同公式，运维成本须乘以(1+7.99%)后代入。2023年度运维成本=第一经营年度成本5,341,344.03\u00d71.0799=5,768,117.42元（与征求意见稿表3一致）。2024年度运维成本=第二经营年度5,948,956.93\u00d71.0799+第三经营年度(2个月)431,456.22\u00d71.0799=6,890,208.16元。2025年度经本所独立验证：税前成本5,713,241.86元\u00d71.0799=6,169,729.88元。2026-2040年度暂以2025年度税前成本取整5,710,000\u00d71.0799\uff1d{:,.2f}元/年估算。'.format(OP_FUTURE),indent=Cm(0.74))

pa('（4）第三方收入：不需乘以K系数。2023年度4,285,070.31元；2024年度6,229,093.45元（沿用征求意见稿表3）；2025年度经独立验证5,564,017.88元；2026-2040年度暂以5,560,000.00元/年估算。',indent=Cm(0.74))

# 四
pa('四、测算过程','黑体',Pt(14),True)

pa('（一）每年项目资本金回报测算情况','楷体',Pt(12),True)
pa('项目资本金回报详见表1（单位：元）：',indent=Cm(0.74))
t1h=['年','等额本息\n本金','等额本息\n利息','等额本息\n小计','一次性支付','资本金回报\n合计','剩余\n(等额本息)']
t1r=[]
for y in yd:
    yr,pr,interest,pmt,lump,tc,rem,*_=y
    t1r.append([str(yr),f'{pr:,.2f}',f'{interest:,.2f}',f'{pmt:,.2f}',
                f'{lump:,.2f}'if lump else'-',f'{tc:,.2f}',f'{rem:,.2f}'if rem>0 else'0.00'])
sp=sum(y[1]for y in yd); si=sum(y[2]for y in yd); sa=sum(y[3]for y in yd)
t1r.append(['合计',f'{sp:,.2f}',f'{si:,.2f}',f'{sa:,.2f}',f'{CAPITAL_LUMP:,.2f}',f'{ct:,.2f}','-'])
tb(t1h,t1r,cw=[1.1,2.3,2.3,2.3,2.2,2.3,2.8])
doc.add_paragraph()

pa('（二）每年实际融资本息测算情况','楷体',Pt(12),True)
pa('融资本息见表2（单位：元）。逐年数据以征求意见稿表2结构为基础，按银行报告确认总数等比缩放（非银行原始逐年数据）：',indent=Cm(0.74))
t2h=['年','实际融资本息','剩余融资金额','适用利率']
t2r=[]; rates={0:'6.95%',1:'6.45%',2:'5.5%/5.0%',3:'5.0%',4:'5.0%',5:'5.0%',6:'5.0%',7:'5.0%',8:'5.0%',9:'5.0%',10:'5.0%',11:'5.0%',12:'5.0%',13:'5.0%',14:'5.0%',15:'5.0%',16:'5.0%',17:'-'}
for i,y in enumerate(yd):
    yr,_,_,_,_,_,_,bk,_,_,_=y; br=bs[i][1]; rs=rates.get(i,'5.0%')
    t2r.append([str(yr),f'{bk:,.2f}',f'{br:,.2f}'if br>0 else'-',rs])
t2r.append(['合计',f'{bt:,.2f}','-','\u2014'])
tb(t2h,t2r,cw=[1.5,4.0,4.0,2.0])
doc.add_paragraph()
pa('注1：逐年数据系等比缩放推导值（缩放系数=710,822,119.32\u00f7712,464,358.22\u22480.997695），非银行原始逐年还款计划。合计数710,822,119.32元为银行确认数。',indent=Cm(0.74),sz=Pt(10))
pa('注2：第18年（2040年度）融资本息为0元，贷款已于第17年（2039年12月）全部结清。',indent=Cm(0.74),sz=Pt(10))

pa('（三）运营维护成本','楷体',Pt(12),True)
pa('合同公式为"运营维护成本\u00d7(1+K)"，K=7.99%。以下成本均已包含该加成：',indent=Cm(0.74))
pa('2023年度：5,768,117.42元（底层：第一经营年度成本5,341,344.03\u00d71.0799。覆盖期2022.10.28-2023.12.21约14个月。与征求意见稿表3一致，本所未取得2023年原始账套独立验证）。',indent=Cm(0.74))
pa('2024年度：6,890,208.16元（底层：第二经营年度5,948,956.93\u00d71.0799+第三经营年度(2个月)431,456.22\u00d71.0799=6,424,278.59+465,929.57。覆盖期2023.12.21-2024.12.21。沿用征求意见稿表3）。',indent=Cm(0.74))
pa('2025年度\u3010独立验证\u3011：6,169,729.88元（底层：2025年税前运维成本5,713,241.86\u00d71.0799=6,169,729.88。其中税前成本=主营业务成本Q1-Q4合计4,474,027.75元+管理费用Q1-Q4合计1,239,214.11元，均取自2025年度XLS明细账4季度损益结转合计值，不含年末调整项）。',indent=Cm(0.74))
pa(f'2026-2040年度：暂以2025年度税前成本取整5,710,000\u00d71.0799\uff1d{OP_FUTURE:,.2f}元/年估算（15年）。',indent=Cm(0.74))

pa('（四）第三方收入','楷体',Pt(12),True)
pa('合同公式中第三方收入不乘K系数，直接扣减：',indent=Cm(0.74))
pa('2023年度：4,285,070.31元（沿用征求意见稿表3）。2024年度：6,229,093.45元（同上）。',indent=Cm(0.74))
pa('2025年度\u3010独立验证\u3011：5,564,017.88元（主营业务收入Q1-Q4合计4,961,142.62元+其他业务收入Q1-Q4合计602,875.26元，来源同上）。',indent=Cm(0.74))
pa('2026-2040年度：暂以2025年度实际值取整5,560,000.00元/年估算（15年）。',indent=Cm(0.74))

pa('（五）考核结果','楷体',Pt(12),True)
pa('根据补充协议，绩效考核\u226580分时支付比例100%。截至2026年6月30日四个运营年度考核分数均\u226580分（2022:93/2023:94/2024:95/2025:93.5），对应支付比例100%。',indent=Cm(0.74))

# 五
pa('五、测算结果','黑体',Pt(14),True)
pa('可用性付费汇总见表3（单位：元）：',indent=Cm(0.74))
t3h=['年','资本金回报','实际融资本息','运维成本\n(含K)','第三方收入','可用性付费']
t3r=[]
for y in yd:
    yr,_,_,_,_,tc,_,bk,op,inc,avail=y
    t3r.append([str(yr),f'{tc:,.2f}',f'{bk:,.2f}',f'{op:,.2f}',f'{inc:,.2f}',f'{avail:,.2f}'])
t3r.append(['合计',f'{ct:,.2f}',f'{bt:,.2f}',f'{ot:,.2f}',f'{it:,.2f}',f'{gt:,.2f}'])
tb(t3h,t3r,cw=[1.2,2.8,2.8,2.8,2.8,2.8])
doc.add_paragraph()
pa(f'测算结论：巴中市恩阳医养园PPP项目运营期18年可用性付费总额为{gt:,.2f}元（约{gt/1e8:.2f}亿元）。',indent=Cm(0.74),bold=True)

# 备注
pa('备注：','楷体',Pt(10),True)
pa('1.计费期限：2023年度覆盖2022.10.28-2023.12.21（约14个月），其余年度均为12.21-12.21。2023年度资本金回报按全年等额值计算（非按比例调整），与原征求意见稿处理方式一致。',sz=Pt(10))
pa(f'2.资本金方案：实缴102,298,077.79元。8,000万元按P={CAPITAL_AMORT:,.0f}、k=7.99%、n=18等额本息，18年等额本息合计153,544,445.28元（本金{CAPITAL_AMORT:,.0f}元+利息73,544,445.28元）；22,298,077.79元于第18年一次性支付，不计利息。',sz=Pt(10))
pa('3.运维成本K系数：合同公式"运营维护成本\u00d7(1+K)"，K=7.99%。2023-2024年度数字来源于征求意见稿（已含K），2025年度起按本所独立验证的税前成本\u00d7(1+7.99%)计算。2026-2040年度税前成本以2025年度为基准取整估算。',sz=Pt(10))
pa('4.融资本息取证层级：(1)合计数710,822,119.32元\u2014\u2014贵阳银行《借款还本付息情况报告》确认（一级证据）；(2)逐年分配\u2014\u2014以征求意见稿表2结构等比缩放推导（二级推导，本所未取得银行逐年还款明细）；(3)2025年Q3财务费用计提利息6,291,985.59元与之基本吻合（侧面验证）。',sz=Pt(10))
pa('5.正式出具报告前建议补充：(1)2023-2024年度项目公司原始成本/收入明细账；(2)贵阳银行逐年还款计划明细表；(3)恩阳区卫健局绩效考核正式文件。',sz=Pt(10))
pa('6.本报告系征求意见稿（v5修订版）。v4\u2192v5主要修正：2025年度起运维成本补正合同约定的(1+7.99%)K系数。',sz=Pt(10))

doc.add_paragraph(); doc.add_paragraph()
pa('四川融策会计师事务所有限公司',align=WD_ALIGN_PARAGRAPH.RIGHT,sz=Pt(12))
pa('二\u3007二六年七月二十九日',align=WD_ALIGN_PARAGRAPH.RIGHT,sz=Pt(12))

outpath = r'C:\Users\scrccpa\Desktop\恩阳医养园PPP项目可用性付费测算报告（修订版v5-20260729）.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
print(f'Size: {os.path.getsize(outpath):,} bytes')
