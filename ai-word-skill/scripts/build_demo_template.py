#!/usr/bin/env python3
"""
生成演示用 Word 母版（融策业务场景）：
  1. 审计报告母版 — 审计报告-模板.docx
  2. 标书母版 — 标书-模板.docx

运行: python scripts/build_demo_template.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = Path(__file__).resolve().parent.parent / "demo"


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr, val in edge_data.items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, cn_font='仿宋_GB2312', en_font='Times New Roman', size=Pt(14), bold=False):
    """设置 run 的字体（中英文分别指定）"""
    run.font.size = size
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def add_paragraph_with_font(doc, text, cn_font='仿宋_GB2312', en_font='Times New Roman',
                             size=Pt(14), bold=False, alignment=None, space_after=Pt(6)):
    """添加一个带中英文字体的段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


# ============================================================
#  审计报告母版
# ============================================================

def build_audit_report_template():
    """生成审计报告母版 — 模拟政府审计报告格式"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # === 封面区域 ===
    add_paragraph_with_font(doc, '', size=Pt(14))
    add_paragraph_with_font(doc, '{{报告标题}}', cn_font='宋体', size=Pt(22),
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

    # 封面信息表
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ('被审计单位：', '{{被审计单位}}'),
        ('审计项目：', '{{审计项目}}'),
        ('审计期间：', '{{审计期间}}'),
        ('报告日期：', '{{报告日期}}'),
        ('项目负责人：', '{{项目负责人}}'),
        ('报告文号：', '{{报告文号}}'),
    ]
    for i, (label, value) in enumerate(info_rows):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]
        cell_label.width = Cm(3.5)
        cell_value.width = Cm(10)
        # label
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        set_run_font(run, cn_font='宋体', size=Pt(14))
        # value
        p = cell_value.paragraphs[0]
        run = p.add_run(value)
        set_run_font(run, cn_font='宋体', size=Pt(14))

    doc.add_page_break()

    # === 正文第一页 ===
    add_paragraph_with_font(doc, '一、审计基本情况', cn_font='黑体', size=Pt(16),
                            bold=True, space_after=Pt(12))
    add_paragraph_with_font(doc,
        '{{基本情况正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))

    add_paragraph_with_font(doc, '二、审计评价意见', cn_font='黑体', size=Pt(16),
                            bold=True, space_after=Pt(12))
    add_paragraph_with_font(doc,
        '{{评价意见正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))

    add_paragraph_with_font(doc, '三、审计发现的主要问题', cn_font='黑体', size=Pt(16),
                            bold=True, space_after=Pt(12))
    add_paragraph_with_font(doc,
        '{{问题正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))

    add_paragraph_with_font(doc, '四、审计建议', cn_font='黑体', size=Pt(16),
                            bold=True, space_after=Pt(12))
    add_paragraph_with_font(doc,
        '{{建议正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))

    # 落款
    add_paragraph_with_font(doc, '', size=Pt(14))
    add_paragraph_with_font(doc, '{{审计机构名称}}', cn_font='仿宋_GB2312',
                            size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_with_font(doc, '{{落款日期}}', cn_font='仿宋_GB2312',
                            size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    out_path = OUT_DIR / '审计报告-模板.docx'
    doc.save(str(out_path))
    print(f"[OK] 已生成: {out_path}")


# ============================================================
#  标书母版
# ============================================================

def build_bid_template():
    """生成标书母版 — 模拟工程项目投标文件格式"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # === 封面 ===
    add_paragraph_with_font(doc, '', size=Pt(14))
    add_paragraph_with_font(doc, '正  本', cn_font='宋体', size=Pt(28),
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
    add_paragraph_with_font(doc, '{{项目名称}}', cn_font='宋体', size=Pt(22),
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_paragraph_with_font(doc, '投  标  文  件', cn_font='宋体', size=Pt(24),
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ('投标人名称：', '{{投标人名称}}'),
        ('法定代表人：', '{{法定代表人}}'),
        ('投标日期：', '{{投标日期}}'),
        ('招标编号：', '{{招标编号}}'),
    ]
    for i, (label, value) in enumerate(info):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]
        cell_label.width = Cm(4)
        cell_value.width = Cm(9)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        set_run_font(run, cn_font='宋体', size=Pt(14))
        p = cell_value.paragraphs[0]
        run = p.add_run(value)
        set_run_font(run, cn_font='宋体', size=Pt(14))

    doc.add_page_break()

    # === 目录页 ===
    add_paragraph_with_font(doc, '目  录', cn_font='黑体', size=Pt(18),
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    add_paragraph_with_font(doc,
        '一、投标函及投标函附录\n'
        '二、法定代表人身份证明\n'
        '三、授权委托书\n'
        '四、投标保证金\n'
        '五、已标价工程量清单\n'
        '六、施工组织设计\n'
        '七、项目管理机构\n'
        '八、资格审查资料\n'
        '九、其他材料',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(12))

    doc.add_page_break()

    # === 投标函 ===
    add_paragraph_with_font(doc, '一、投标函及投标函附录', cn_font='黑体', size=Pt(16),
                            bold=True, space_after=Pt(12))
    add_paragraph_with_font(doc, '致：{{招标人名称}}', cn_font='仿宋_GB2312',
                            size=Pt(14), space_after=Pt(12))
    add_paragraph_with_font(doc,
        '{{投标函正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))

    # 投标函附录（表格）
    add_paragraph_with_font(doc, '投标函附录', cn_font='黑体', size=Pt(14),
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    table2 = doc.add_table(rows=5, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['序号', '条款内容', '约定内容']
    for j, h in enumerate(headers):
        p = table2.rows[0].cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, cn_font='黑体', size=Pt(12), bold=True)
    appendix_rows = [
        ('1', '项目经理', '{{项目经理}}'),
        ('2', '工期', '{{工期}}'),
        ('3', '缺陷责任期', '{{缺陷责任期}}'),
        ('4', '履约担保金额', '{{履约担保金额}}'),
    ]
    for i, (seq, clause, value) in enumerate(appendix_rows):
        for j, text in enumerate([seq, clause, value]):
            p = table2.rows[i + 1].cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, cn_font='仿宋_GB2312', size=Pt(12))

    doc.add_page_break()

    # === 施工组织设计（简化） ===
    add_paragraph_with_font(doc, '六、施工组织设计', cn_font='黑体', size=Pt(16),
                            bold=True, space_after=Pt(12))
    add_paragraph_with_font(doc, '6.1 工程概况', cn_font='楷体_GB2312', size=Pt(14),
                            bold=True, space_after=Pt(6))
    add_paragraph_with_font(doc,
        '{{工程概况正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))
    add_paragraph_with_font(doc, '6.2 施工总体部署', cn_font='楷体_GB2312', size=Pt(14),
                            bold=True, space_after=Pt(6))
    add_paragraph_with_font(doc,
        '{{施工部署正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))
    add_paragraph_with_font(doc, '6.3 主要施工方案', cn_font='楷体_GB2312', size=Pt(14),
                            bold=True, space_after=Pt(6))
    add_paragraph_with_font(doc,
        '{{施工方案正文}}',
        cn_font='仿宋_GB2312', size=Pt(14), space_after=Pt(6))

    # 落款
    add_paragraph_with_font(doc, '', size=Pt(14))
    add_paragraph_with_font(doc, '{{投标人名称}}（盖章）', cn_font='仿宋_GB2312',
                            size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_with_font(doc, '法定代表人或其委托代理人：（签字）', cn_font='仿宋_GB2312',
                            size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_with_font(doc, '{{投标日期}}', cn_font='仿宋_GB2312',
                            size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    out_path = OUT_DIR / '标书-模板.docx'
    doc.save(str(out_path))
    print(f"[OK] 已生成: {out_path}")


# ============================================================
#  主入口
# ============================================================

if __name__ == '__main__':
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_audit_report_template()
    build_bid_template()
    print("\n[完成] 两个母版模板已生成到 demo/ 目录")
