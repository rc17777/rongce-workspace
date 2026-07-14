#!/usr/bin/env python3
"""
融策业务适配 Demo — AI-Word-Skill 实战演示

基于 AI-Word-Skill 的核心 SOP，演示：
  场景1: 审计报告 — 母版副本 + 占位符替换 + 正文整段灌入
  场景2: 标书文件 — 母版副本 + 占位符替换 + 表格填充 + 正文灌入
  场景3: 对比实验 — SOP vs 整段赋值（看版式差异）

运行:
  python scripts/demo_rongce.py
  python scripts/demo_rongce.py --scene all          # 全部场景
  python scripts/demo_rongce.py --scene audit        # 仅审计报告
  python scripts/demo_rongce.py --scene bid           # 仅标书
  python scripts/demo_rongce.py --scene compare       # 仅对比实验
"""
from __future__ import annotations

import sys
import argparse
from pathlib import Path

# 确保能找到 core 模块
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core import (
    open_template, save_doc,
    rewrite_paragraph, rewrite_paragraphs_by_index,
    replace_placeholders, replace_all, replace_cross_runs,
    fill_table_cell,
)
from docx import Document
import shutil


ROOT = Path(__file__).resolve().parent.parent
DEMO_DIR = ROOT / "demo"
OUT_DIR = ROOT / "out"
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
#  场景1: 审计报告生成
# ============================================================

AUDIT_CONTENT = {
    # 占位符替换（封面信息 + 落款）
    "placeholders": {
        "{{报告标题}}": "关于XX单位2025年度预算执行与\n财政财务收支情况的审计报告",
        "{{被审计单位}}": "XX市XX局",
        "{{审计项目}}": "2025年度预算执行审计",
        "{{审计期间}}": "2025年1月—12月",
        "{{报告日期}}": "2026年5月11日",
        "{{项目负责人}}": "张三",
        "{{报告文号}}": "融策审报〔2026〕第088号",
        "{{审计机构名称}}": "四川融策会计师事务所",
        "{{落款日期}}": "2026年5月11日",
    },

    # 正文段落索引 → 内容（段落 7,9,11,13 分别是四个章节的正文）
    "body_paragraphs": {
        7: (
            "根据《中华人民共和国审计法》第十九条规定和2026年度审计工作计划安排，"
            "四川融策会计师事务所接受XX市审计局委托，派出审计组自2026年3月1日至4月30日，"
            "对XX市XX局（以下简称「XX局」）2025年度预算执行及财政财务收支情况进行了就地审计。"
            "XX局对其提供的财务会计资料及其他相关资料的真实性和完整性负责。"
            "我们的审计是依据《中国注册会计师审计准则》和《中华人民共和国国家审计准则》进行的。"
            "在审计过程中，我们结合XX局的实际情况，实施了包括查阅会计资料、检查实物资产、"
            "分析性复核等必要的审计程序。现将审计情况报告如下。"
        ),
        9: (
            "审计结果表明，XX局2025年度预算编制、执行和决算编报基本符合《预算法》及相关规定，"
            "财政财务收支基本真实、合规。内部控制制度基本健全，执行情况较好。"
            "但在预算执行管理、专项资金使用、资产管理等方面仍存在一些需要改进的问题。"
            "审计期间，XX局积极配合审计工作，对审计发现的问题进行了边审边改。"
        ),
        11: (
            "（一）预算执行管理方面。存在部分项目预算执行率偏低的问题，"
            "3个项目全年预算执行率低于50%，涉及金额325万元，资金使用效益有待提高。"
            "\n（二）专项资金使用方面。存在专项资金被挤占挪用的问题，"
            "将「乡村振兴示范项目」专项资金58万元用于单位日常公用经费支出，"
            "不符合专项资金专款专用的规定。"
            "\n（三）资产管理方面。存在固定资产账实不符的问题，"
            "抽查发现12项账面记载的固定资产（原值合计86万元）实际已无法使用或已处置但未办理报废手续；"
            "另有3辆公务用车长期闲置，账面价值45万元。"
            "\n（四）政府采购方面。存在拆分项目规避招标的问题，"
            "将同一品目的信息化设备采购拆分为4笔小额采购（每笔均未超过招标限额），"
            "合计金额168万元，规避了公开招标程序。"
        ),
        13: (
            "（一）强化预算执行管理。建立健全预算执行动态监控机制，"
            "对执行率偏低的项目及时分析原因并采取改进措施，切实提高财政资金使用效益。"
            "\n（二）规范专项资金管理。严格执行专项资金管理办法，"
            "做到专款专用、单独核算，严禁挤占挪用。对已挪用的58万元专项资金应限期归还原渠道。"
            "\n（三）加强资产管理。全面开展固定资产清查盘点，"
            "及时办理资产处置手续，做到账实相符。对闲置资产提出盘活利用方案。"
            "\n（四）规范采购行为。严格遵守政府采购法律法规，"
            "不得通过拆分项目等方式规避招标程序。"
        ),
    },
}


def demo_audit_report():
    """演示: 审计报告 — 母版副本 + 占位符 + 正文灌入"""
    print("\n" + "=" * 65)
    print("  场景1: 审计报告生成")
    print("=" * 65)

    template = DEMO_DIR / "审计报告-模板.docx"
    if not template.is_file():
        print("[ERROR] 母版不存在，请先运行: python scripts/build_demo_template.py")
        return

    # Step 1: 母版副本
    out_sop = OUT_DIR / "审计报告-SOP输出.docx"
    doc = open_template(template, out_sop)

    # Step 2: 占位符批量替换（封面信息、落款）
    n_ph = replace_placeholders(doc, AUDIT_CONTENT["placeholders"])
    print(f"  [占位符] 替换了 {n_ph} 处")

    # Step 3: 正文整段灌入（保留母版仿宋字体、行距）
    n_body = rewrite_paragraphs_by_index(doc, AUDIT_CONTENT["body_paragraphs"])
    print(f"  [正文]   灌入了 {n_body} 段")

    # Step 4: 保存
    save_doc(doc, out_sop)
    print(f"  [输出]   {out_sop}")
    print(f"\n  >> 用 Word 打开查看，版式应与母版一致")


# ============================================================
#  场景2: 标书生成
# ============================================================

BID_CONTENT = {
    "placeholders": {
        "{{项目名称}}": "XX市高新区市政道路\n改造提升工程（二期）",
        "{{投标人名称}}": "四川融策工程咨询有限公司",
        "{{法定代表人}}": "李四",
        "{{投标日期}}": "2026年5月11日",
        "{{招标编号}}": "GXQ-SZDL-2026-002",
        "{{招标人名称}}": "XX市高新区建设局",
    },
    "body_paragraphs": {
        # 投标函正文（段落索引 12）
        12: (
            "1. 我方已仔细研究了XX市高新区市政道路改造提升工程（二期）施工招标文件的全部内容，"
            "愿意以人民币（大写）叁仟捌佰陆拾万零伍仟元整（¥38,605,000.00）的投标总报价，"
            "工期365日历天，按合同约定实施和完成承包工程，修补工程中的任何缺陷，工程质量达到国家现行验收合格标准。"
            "\n2. 我方承诺在招标文件规定的投标有效期内不修改、撤销投标文件。"
            "\n3. 随同本投标函提交投标保证金一份，金额为人民币（大写）伍拾万元整（¥500,000.00）。"
            "\n4. 如我方中标：（1）我方承诺在收到中标通知书后，在中标通知书规定的期限内与你方签订合同；"
            "（2）随同本投标函递交的投标函附录属于合同文件的组成部分；"
            "（3）我方承诺按照招标文件规定向你方递交履约担保；"
            "（4）我方承诺在合同约定的期限内完成并移交全部合同工程。"
            "\n5. 我方在此声明，所递交的投标文件及有关资料内容完整、真实和准确。"
        ),
        # 工程概况（段落索引 19）
        19: (
            "本项目位于XX市高新区核心区域，建设内容包括："
            "科园南路（K0+000～K2+850）全长2.85km、科园北路（K0+000～K1+620）全长1.62km、"
            "创业路（K0+000～K1+140）全长1.14km，总计5.61km的市政道路改造提升。"
            "主要工程内容：路面铣刨加铺沥青混凝土面层、人行道铺装更新、"
            "雨污水管网改造、交通标志标线更新、绿化提升及智慧路灯系统安装。"
            "项目预算总投资约为4200万元，资金来源为区级财政资金。"
        ),
        # 施工部署（段落索引 21）
        21: (
            "本项目采用「分段平行施工、流水作业」的总体部署方案。"
            "将全线划分为3个施工工区：第一工区（科园南路K0+000～K1+400）、"
            "第二工区（科园南路K1+400～K2+850及科园北路）、"
            "第三工区（创业路全线）。各工区配备独立的施工管理团队和机械设备。"
            "总工期365日历天，其中施工准备期15天，主体工程施工期320天，竣工验收期30天。"
            "关键节点：2026年8月30日前完成全部管网改造，2027年2月28日前完成路面工程。"
        ),
        # 施工方案（段落索引 23）
        23: (
            "（1）路面工程：采用铣刨4cm原沥青面层+摊铺6cm AC-20C中粒式沥青砼下面层"
            "+4cm SMA-13沥青玛蹄脂碎石上面层的技术方案。"
            "\n（2）管网改造：雨污水管道采用HDPE双壁波纹管（SN8级），"
            "管径DN300-DN800，热熔连接。施工采用明挖法，沟槽支护采用拉森钢板桩。"
            "\n（3）人行道工程：采用透水混凝土基层+C40彩色压印混凝土面层的铺装方案，"
            "配套设置无障碍坡道和盲道。"
            "\n（4）智慧路灯：采用LED光源+单灯控制系统，集成环境监测、WiFi热点、"
            "信息发布等功能模块，实现「多杆合一」。"
        ),
    },
}


def demo_bid():
    """演示: 标书 — 母版副本 + 占位符 + 表格填充 + 正文灌入"""
    print("\n" + "=" * 65)
    print("  场景2: 标书文件生成")
    print("=" * 65)

    template = DEMO_DIR / "标书-模板.docx"
    if not template.is_file():
        print("[ERROR] 母版不存在，请先运行: python scripts/build_demo_template.py")
        return

    out_sop = OUT_DIR / "标书-SOP输出.docx"
    doc = open_template(template, out_sop)

    # Step 1: 占位符
    n_ph = replace_placeholders(doc, BID_CONTENT["placeholders"])
    print(f"  [占位符] 替换了 {n_ph} 处")

    # Step 2: 投标函附录表格填充
    table = doc.tables[1]  # 第二个表格（投标函附录）
    appendix_data = {
        (1, 2): "王工（一级建造师  川1512020202501234）",
        (2, 2): "365日历天",
        (3, 2): "24个月",
        (4, 2): "中标合同价的10%",
    }
    for (r, c), text in appendix_data.items():
        fill_table_cell(table, r, c, text)
    print(f"  [表格]   填充了投标函附录 {len(appendix_data)} 项")

    # Step 3: 正文灌入
    n_body = rewrite_paragraphs_by_index(doc, BID_CONTENT["body_paragraphs"])
    print(f"  [正文]   灌入了 {n_body} 段")

    save_doc(doc, out_sop)
    print(f"  [输出]   {out_sop}")
    print(f"\n  >> 用 Word 打开查看，版式应与母版一致")


# ============================================================
#  场景3: 对比实验 — SOP vs 整段赋值
# ============================================================

COMPARE_PARAGRAPHS = {
    0: "审计报告对比实验：SOP改写 vs 整段赋值（请看字体、段落间距的差异）",
    1: "尊敬的各位领导：",
    2: (
        "本次审计依据《中华人民共和国审计法》《中华人民共和国预算法》"
        "及其实施条例等法律法规，按照法定程序和审计准则执行。"
        "审计组通过查阅会计账簿、原始凭证、合同协议等资料，"
        "结合实地核查、询问访谈、数据分析等方式开展工作。"
    ),
    3: (
        "审计过程中，我们重点关注了预算编制与执行的合规性、"
        "专项资金使用的规范性、国有资产管理的有效性、"
        "政府采购程序的合法性以及内部控制制度的健全性。"
    ),
    4: (
        "通过审计，我们识别出若干需要关注的问题，"
        "并提出切实可行的整改建议，以期促进被审计单位规范管理、提高效能。"
    ),
}


def demo_compare():
    """演示: 生成一对对比文件 — SOP改写 vs 整段赋值"""
    print("\n" + "=" * 65)
    print("  场景3: SOP vs 整段赋值 对比实验")
    print("=" * 65)

    template = DEMO_DIR / "审计报告-模板.docx"
    if not template.is_file():
        print("[ERROR] 母版不存在，请先运行: python scripts/build_demo_template.py")
        return

    # --- SOP 版本 ---
    sop_out = OUT_DIR / "对比-SOP改写.docx"
    doc_sop = open_template(template, sop_out)
    for idx, text in COMPARE_PARAGRAPHS.items():
        if idx < len(doc_sop.paragraphs):
            rewrite_paragraph(doc_sop.paragraphs[idx], text)
    save_doc(doc_sop, sop_out)
    print(f"  [SOP]    {sop_out}")

    # --- 整段赋值版本 ---
    bad_out = OUT_DIR / "对比-整段赋值(版式异常).docx"
    doc_bad = open_template(template, bad_out)
    for idx, text in COMPARE_PARAGRAPHS.items():
        if idx < len(doc_bad.paragraphs):
            doc_bad.paragraphs[idx].text = text
    save_doc(doc_bad, bad_out)
    print(f"  [反模式] {bad_out}")

    print(f"\n  >> 用 Word 并排打开这两个文件，"
          f"对比字体（仿宋 vs 等线/宋体丢失）和行距差异")
    print(f"  >> SOP 版本保留了母版的仿宋字体和1.5倍行距")
    print(f"  >> 整段赋值版本字体可能回退到默认值，行距可能不一致")


# ============================================================
#  主入口
# ============================================================

SCENES = {
    "audit": demo_audit_report,
    "bid": demo_bid,
    "compare": demo_compare,
    "all": lambda: [f() for f in (demo_audit_report, demo_bid, demo_compare)],
}


def main():
    parser = argparse.ArgumentParser(description="融策业务 AI-Word-Skill Demo")
    parser.add_argument("--scene", choices=list(SCENES.keys()), default="all",
                        help="运行场景 (default: all)")
    args = parser.parse_args()

    print("=" * 65)
    print("  融策业务 AI-Word-Skill Demo")
    print("  基于 python-docx run 级保格式 SOP")
    print("=" * 65)

    SCENES[args.scene]()

    print("\n" + "=" * 65)
    print(f"  所有输出文件在: {OUT_DIR}")
    print("=" * 65)


if __name__ == "__main__":
    main()
