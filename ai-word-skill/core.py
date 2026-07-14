#!/usr/bin/env python3
"""
AI-Word-Skill 核心函数库
基于 python-docx 的 OOXML run 级别操作，保留 Word 母版格式。
来源: github.com/sgsss998/AI-Word-Skill (MIT)
适配: 融策会计师事务所 / 融策工程咨询公司
"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# ============================================================
#  L1: 单 run 内替换（最安全）
# ============================================================

def replace_in_paragraph(paragraph: Paragraph, old_text: str, new_text: str) -> bool:
    """在段落的单个 run 中查找并替换文本，完全保留 run 格式。"""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    return False


# ============================================================
#  L2: 跨 run 替换（处理 Word 拆分 run 的情况）
# ============================================================

def replace_cross_runs(paragraph: Paragraph, old_text: str, new_text: str) -> bool:
    """跨 run 替换：拼合段落 → 定位 → 合并写回首 run，清空其余。"""
    full_text = ''.join(r.text for r in paragraph.runs)
    if old_text not in full_text:
        return False

    start = full_text.find(old_text)
    end = start + len(old_text)

    char_pos = 0
    affected: list[int] = []
    for i, run in enumerate(paragraph.runs):
        r_start = char_pos
        r_end = char_pos + len(run.text)
        if r_start < end and r_end > start:
            affected.append(i)
        char_pos = r_end

    if not affected:
        return False

    merged = ''.join(paragraph.runs[i].text for i in affected)
    merged = merged.replace(old_text, new_text, 1)
    paragraph.runs[affected[0]].text = merged
    for i in affected[1:]:
        paragraph.runs[i].text = ''
    return True


# ============================================================
#  L3: 整段重写（保留首 run 格式 DNA）
# ============================================================

def rewrite_paragraph(paragraph: Paragraph, new_text: str) -> None:
    """整段替换文本，继承 runs[0] 的字体/字号/颜色等 rPr。"""
    if not paragraph.runs:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ''


# ============================================================
#  L4: 插入新段落（deepcopy 模板段落 XML）
# ============================================================

def copy_paragraph_element(template_paragraph: Paragraph) -> any:
    """深拷贝段落 XML 元素，清空所有 w:r，返回干净的段落骨架。"""
    new_p = deepcopy(template_paragraph._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    return new_p


def add_run_to_element(p_element, template_paragraph: Paragraph, text: str) -> None:
    """向段落 XML 元素中添加一个 run（拷贝母版第一个 run 的格式）。"""
    template_r = template_paragraph._element.find(qn('w:r'))
    if template_r is None:
        return
    new_r = deepcopy(template_r)
    for t in new_r.findall(qn('w:t')):
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    p_element.append(new_r)


def insert_paragraph_after(doc: Document, ref_paragraph: Paragraph,
                           template_paragraph: Paragraph, text: str) -> Paragraph:
    """在参考段落后插入一个新段落，继承母版格式。"""
    new_p_elem = copy_paragraph_element(template_paragraph)
    add_run_to_element(new_p_elem, template_paragraph, text)
    ref_paragraph._element.addnext(new_p_elem)
    return Paragraph(new_p_elem, doc)


# ============================================================
#  全文档替换（段落 + 表格）
# ============================================================

def replace_all(doc: Document, old: str, new: str) -> int:
    """全文档替换：遍历所有正文段落 + 表格单元格。"""
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            count += 1
    return count


# ============================================================
#  表格单元格文本写入
# ============================================================

def fill_table_cell(table, row_idx: int, col_idx: int, text: str) -> bool:
    """向表格指定单元格写入文本（保留格式）。"""
    try:
        cell = table.rows[row_idx].cells[col_idx]
        if cell.paragraphs and cell.paragraphs[0].runs:
            rewrite_paragraph(cell.paragraphs[0], text)
        else:
            cell.text = text
        return True
    except (IndexError, AttributeError):
        return False


# ============================================================
#  黄金路径：母版副本 + 改写 + 保存
# ============================================================

def open_template(template_path: str | Path, output_path: str | Path) -> Document:
    """复制母版到输出路径，返回 Document 对象。"""
    shutil.copy(str(template_path), str(output_path))
    return Document(str(output_path))


def copy_template_for_edit(template_path: str | Path, output_path: str | Path) -> Document:
    """open_template 的别名（保持兼容）。"""
    return open_template(template_path, output_path)


def save_doc(doc: Document, output_path: str | Path) -> None:
    """保存 Document 到指定路径。"""
    doc.save(str(output_path))


# ============================================================
#  便捷函数：按段落索引批量重写
# ============================================================

def rewrite_paragraphs_by_index(doc: Document, mapping: dict[int, str]) -> int:
    """按段落索引批量重写（继承首 run 格式）。返回成功数。"""
    count = 0
    for idx, text in mapping.items():
        if idx < len(doc.paragraphs):
            rewrite_paragraph(doc.paragraphs[idx], text)
            count += 1
    return count


def replace_placeholders(doc: Document, mapping: dict[str, str]) -> int:
    """批量占位符替换（{{KEY}} → VALUE），先单 run 再跨 run。返回成功数。"""
    count = 0
    for old, new in mapping.items():
        for p in doc.paragraphs:
            if replace_in_paragraph(p, old, new):
                count += 1
            elif replace_cross_runs(p, old, new):
                count += 1
    # 表格也扫一轮
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_paragraph(p, old, new):
                        count += 1
                    elif replace_cross_runs(p, old, new):
                        count += 1
    return count


print("[core.py] AI-Word-Skill 核心库已加载 — python-docx run 级保格式工具集")
