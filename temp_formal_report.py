"""Generate formal report following the original template exactly"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import math

# ===== Calculation =====
capital_equity = 80_000_000.00
capital_lumpsum = 22_298_077.79
rate_k = 0.0799
n_years = 18

factor = (1 + rate_k) ** n_years
annual_payment = capital_equity * rate_k * factor / (factor - 1)

# Equity schedule
remaining = capital_equity
equity_schedule = []
for yr in range(1, n_years + 1):
    interest = remaining * rate_k
    principal = annual_payment - interest
    if yr == n_years:
        principal = remaining
    remaining -= principal
    equity_schedule.append({
        'year': yr,
        'payment': round(annual_payment, 2),
        'principal': round(principal, 2),
        'interest': round(interest, 2),
        'remaining': max(0, round(remaining, 2))
    })

# Bank schedule (from bank report)
bank_schedule = [
    (2023, 31_686_745.84,  1_500_000.00,  30_186_745.84),
    (2024, 26_809_671.03,  0,             26_809_671.03),
    (2025, 24_811_756.73,  50_000.00,     24_761_756.73),
    (2026, 21_351_219.32,  1_000_000.00,  20_351_219.32),
    (2027, 20_263_854.17,  1_000_000.00,  19_263_854.17),
    (2028, 20_265_833.33,  1_000_000.00,  19_265_833.33),
    (2029, 20_162_465.30,  1_000_000.00,  19_162_465.30),
    (2030, 23_073_645.82,  4_000_000.00,  19_073_645.82),
    (2031, 22_870_868.06,  4_000_000.00,  18_870_868.06),
    (2032, 22_719_375.01,  4_000_000.00,  18_719_375.01),
    (2033, 22_465_312.49,  4_000_000.00,  18_465_312.49),
    (2034, 22_262_534.71,  4_000_000.00,  18_262_534.71),
    (2035, 26_996_215.28,  9_000_000.00,  17_996_215.28),
    (2036, 27_575_625.01,  10_000_000.00, 17_575_625.01),
    (2037, 34_918_645.83,  18_000_000.00, 16_918_645.83),
    (2038, 34_006_145.84,  18_000_000.00, 16_006_145.84),
    (2039, 308_582_205.55, 302_250_000.00, 6_332_205.55),
    (2040, 0,              0,              0),
]

op_cost_actual = [5_768_117.42, 6_890_208.16]
income_actual = [4_285_070.31, 6_229_093.45]
op_est = 6_090_636.00  # from original table (2025+)
income_est = 5_060_000.00

year_labels = list(range(2023, 2041))

# ===== DOCUMENT =====
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

DARK_BLUE = RGBColor(0x0A, 0x1F, 0x3F)
TEAL = RGBColor(0x1A, 0x5C, 0x6E)
GOLD = RGBColor(0xC5, 0x95, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

def add_h1(text):
    """Main heading - 一、二、三 etc"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = BLACK
    return p

def add_h2(text):
    """Sub heading - （一）（二）etc"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.color.rgb = BLACK
    return p

def add_body(text, indent=True):
    """Body paragraph"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_formula(text):
    """Formula line - centered"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    return p

def add_table(headers, rows, col_widths=None, bold_last=False):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    for i, row in enumerate(rows):
        is_last = bold_last and (i == len(rows) - 1)
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if val is not None else '')
            run.font.size = Pt(8.5)
            run.font.bold = is_last
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    
    doc.add_paragraph()
    return table

def fmt(n):
    """Format number with commas and 2 decimals"""
    return f"{n:,.2f}"

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('巴中市恩阳医养园PPP项目')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('可用性付费测算结果报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('川融策咨询〔2025〕第490号')
run.font.size = Pt(14)
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('四 川 融 策 会 计 师 事 务 所 有 限 公 司')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sichuan Rongce Accounting Firm Co., Ltd')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年7月11日')
run.font.size = Pt(14)
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run('目  录')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

toc_items = [
    '一、项目概况',
    '二、测算依据',
    '（一）政策依据',
    '（二）项目资料依据',
    '三、测算说明',
    '（一）项目回报机制识别',
    '（二）回报计算方法选用',
    '四、测算过程',
    '（一）每年项目资本金回报测算情况',
    '（二）每年实际融资本息测算情况',
    '（三）运营维护成本',
    '（四）第三方收入',
    '（五）考核结果',
    '五、测算结果',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# REPORT BODY
# ============================================================

# Title line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('巴中市恩阳医养园PPP项目')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('可用性付费测算结果报告')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

doc.add_paragraph()

# Recipient
add_body('巴中市恩阳区卫生健康局：', indent=False)
add_body('我们接受委托，对巴中市恩阳医养园PPP项目2022年10月28日-2040年12月21日可用性付费进行测算。本次测算基数以《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）中审定的建筑安装工程费用485,098,077.79元为准，现将结果报告如下：')

# ===== 一、项目概况 =====
add_h1('一、项目概况')

add_body('2015年12月21日，巴中市恩阳医养园PPP项目经巴中市恩阳区人民政府（恩府〔2015〕208号）批复同意实施，由巴中市恩阳区卫生健康局按规定有序开展。')
add_body('2016年7月巴中市恩阳医养园PPP项目协议经区人民政府审定同意，巴中市恩阳区卫生健康局于2016年12月与恩阳区人民医院、湖南省第五工程有限公司签订PPP协议，成立项目管理公司——巴中市恩阳区医养园项目经营管理有限公司，注册资本金10,000万元，承担该项目的投资、建设、运营维护及移交。')
add_body('2017年7月19日项目公司与湖南省第五工程有限公司签订施工合同，项目建筑总面积118,515.13㎡，包括院内绿化、场地及路面硬化、室内管网等附属工程及设备购置等。该项目实际于2017年5月17日开工，2022年3月完工，2022年3月28日进入调试运营期，同年10月27日通过竣工验收，正式进入运营阶段。')
add_body('项目合作期限根据2021年5月签订的《巴中市恩阳医养园PPP项目合同之补充合同》（三）调整为：合作期23年（含建设期，建设期不超过5年，运营期为18年；如遇建设期工期延长，运营期则相应顺延，运营年限保持不变），自监理工程师发出开工令之日起计算。')
add_body('项目回报机制为可行性缺口补助，社会资本不参与提供医疗服务盈利，其投资回报主要通过经营医院行政后勤管理（食堂、停车场、超市、保洁）等作为第三方收益，不足部分由政府可行性缺口补助补足。')

# ===== 二、测算依据 =====
add_h1('二、测算依据')
add_h2('（一）政策依据')
add_body('（1）《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）')
add_body('（2）财政部关于印发《政府和社会资本合作项目财政承受能力论证指引》的通知（财金〔2015〕21号）')
add_body('（3）四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）')

add_h2('（二）项目资料依据')
add_body('（1）《恩阳医养园PPP项目协议》（2016年12月27日）')
add_body('（2）《巴中市恩阳医养园PPP项目合同》（2017年6月26日）')
add_body('（3）《巴中市恩阳医养园PPP项目合同之补充合同》（2017年7月）')
add_body('（4）《巴中市恩阳医养园PPP项目合同之补充合同》（2018年3月7日）')
add_body('（5）《巴中市恩阳医养园PPP项目合同之补充合同》（2021年5月）')
add_body('（6）《关于巴中市恩阳区人民医院一期工程建设项目可行性研究报告的批复》（恩区发改行审〔2015〕83号）')
add_body('（7）《巴中市恩阳医养园PPP建设项目竣工结算审核报告》（中宏审〔2024〕1213号）')
add_body('（8）《巴中市恩阳医养园PPP项目借款还本付息情况报告》（贵阳银行成都分行）')

# ===== 三、测算说明 =====
add_h1('三、测算说明')

add_h2('（一）项目回报机制识别')
add_body('一是政策方面：根据《关于印发政府和社会资本合作模式操作指南（试行）的通知》（财金〔2014〕113号）："项目回报机制主要分为使用者付费、可行性缺口补助和政府付费等支付方式。其中：使用者付费，是指由最终消费用户直接付费购买公共产品和服务。可行性缺口补助，是指使用者付费不足以满足社会资本或项目公司成本回收和合理回报，而由政府以财政补贴、股本投入、优惠贷款和其他优惠政策的形式，给予社会资本或项目公司的经济补助。政府付费，是指政府直接付费购买公共产品和服务，主要包括可用性付费、使用量付费和绩效付费。"')
add_body('四川省财政厅关于印发《四川省政府与社会资本合作（PPP）项目财政承受能力论证办法》的通知（川财金〔2017〕91号）规定："运营补贴。指项目运营期间根据不同的付费模式政府应当承担的直接付费责任。其中：政府付费模式下，政府承担全部运营补贴支出责任；可行性缺口补助（政府补助）模式下，政府承担部分运营补贴支出责任；使用者付费模式下，政府不承担运营补贴支出责任。"')
add_body('二是项目资料方面：项目实施方案、PPP合同文本中均明确项目回报机制为"可行性缺口补助"。')
add_body('2016年12月27日巴中市恩阳区卫生健康局与巴中市恩阳区人民医院、湖南省第五工程有限公司签订《恩阳医养园PPP项目协议》，协议约定"自运营日起，甲方按照本协议的约定及时、足额地向乙方支付可行性缺口补贴，且区政府将本项目的可行性缺口补贴纳入跨年度的财政预算，并提请人大决议"。')

add_h2('（二）回报计算方法选用')
add_body('根据2017年7月巴中市恩阳区卫生健康局与项目公司签订《巴中市恩阳医养园PPP项目合同之补充合同》，合同计算公式为：')
add_formula('A = P×k×(1+k)^n / [(1+k)^n-1] + 实际融资本息 + 运营维护成本 - 第三方收入')
add_body('其中，A为运营期内各年政府运营补贴。各构成项目说明如下：')

add_body('（1）项目资本金回报：P×k×(1+k)^n / [(1+k)^n-1]', indent=False)
add_body('P——项目资本金（本项目按委托方确认方案，以实际投入资本金8,000万元按等额本息方式测算）；')
add_body('k——合理利润率（社会资本中标年回报率7.99%）；')
add_body('n——财政运营补贴周期（18年）；')
add_body('剩余资本金22,298,077.79元在运营期最后一年（第18年）一次性支付。')
add_body('（2）实际融资成本：为项目运营期间向贵阳银行股份有限公司借款产生的利息及本金。按照借款合同约定以及实际还本付息情况计算，18个运营年度合计710,822,119.32元。')
add_body('（3）运营维护成本：按照项目公司提供的财务资料数据进行核算，18个运营年度合计101,961,757.18元。')
add_body('（4）第三方收入：按照医养园项目公司提供的财务资料数据进行核算，18个运营年度合计91,474,163.76元。')

# ===== 四、测算过程 =====
add_h1('四、测算过程')

add_h2('（一）每年项目资本金回报测算情况')
add_body('按项目公司实际投入资本金102,298,077.79元。根据委托方确认的还款方案，其中8,000万元按合理利润率7.99%、运营期18年，以等额本息的方式测算；剩余22,298,077.79元在运营期第18年一次性支付。每年项目资本金回报见下表（单位：元）：')

# Table 1: Equity schedule
add_body('表1：项目资本金回报测算表', indent=False)
equity_header = ['年', '资本金投入', '本金（等额本息）', '合理利润率', '利润', '每年项目资本金回报', '剩余资本金']
equity_rows = []
for i, e in enumerate(equity_schedule):
    yr = 2023 + i
    if i == n_years - 1:
        ret = fmt(e['payment'] + capital_lumpsum)
    else:
        ret = fmt(e['payment'])
    # Combine the 8000万 portion
    equity_rows.append([
        str(yr),
        '102,298,077.79' if i == 0 else '',
        fmt(e['principal']),
        '7.99%',
        fmt(e['interest']),
        ret,
        fmt(e['remaining'])
    ])
# Add subtotal for equity
total_equity = sum(e['payment'] for e in equity_schedule) + capital_lumpsum
equity_rows.append([
    '小计',
    '',
    fmt(sum(e['principal'] for e in equity_schedule) + capital_lumpsum),
    '',
    fmt(sum(e['interest'] for e in equity_schedule)),
    fmt(total_equity),
    ''
])
add_table(equity_header, equity_rows, [1.3, 2.8, 2.8, 1.8, 2.8, 3, 2.8], bold_last=True)

add_h2('（二）每年实际融资本息测算情况')
add_body('除资本金外项目实际融资金额382,800,000.00元，按实际融资利率（据实浮动），运营期18年，以项目公司与贵阳银行协商约定的还款计划及还款利率测算。2021年6月30日，本项目向贵阳银行成都分行申请了39,750万元项目贷款，贷款年限为18年，原贷款利率为6.95%/年，现已调整为5%/年（2024年11月22日贷款利率调整为6.45%/年，2025年11月28日贷款利率调整为5.5%/年，2026年7月1日贷款利率调整为5.0%/年），累计发放贷款38,280万元。')
add_body('每年实际融资本息见下表（单位：元）：')

# Table 2: Bank schedule
add_body('表2：每年实际融资本息测算表', indent=False)
bank_header = ['年', '实际融资金额', '本金', '利率', '利息', '每年实际融资本息', '剩余融资金额']
bank_rows = []
bank_remain = 382_800_000.00
for i, b in enumerate(bank_schedule):
    yr, total_pay, principal, interest = b
    if i == 0:
        amt_str = '382,800,000.00'
    else:
        amt_str = ''
    
    # Determine rate for display
    if yr <= 2023:
        rate_str = '6.95%'
    elif yr == 2024:
        rate_str = '6.45%'
    elif yr == 2025:
        rate_str = '5.55%'
    else:
        rate_str = '5.00%'
    
    if principal > 0:
        bank_remain -= principal
    
    if yr == 2040:
        bank_remain_str = '-'
    else:
        bank_remain_str = fmt(max(0, bank_remain))
    
    bank_rows.append([
        str(yr),
        amt_str,
        fmt(principal),
        rate_str,
        fmt(interest),
        fmt(total_pay),
        bank_remain_str
    ])

bank_total = sum(b[1] for b in bank_schedule)
bank_rows.append([
    '小计',
    '382,800,000.00',
    fmt(sum(b[2] for b in bank_schedule)),
    '',
    fmt(sum(b[3] for b in bank_schedule)),
    fmt(bank_total),
    ''
])
add_table(bank_header, bank_rows, [1.2, 2.5, 2.3, 1.3, 2.3, 2.5, 2.5], bold_last=True)

add_h2('（三）运营维护成本')
add_body('根据项目公司提供的财务资料，经审核，第一经营年度运营成本金额为5,341,344.03元；第二经营年度运营成本金额5,948,956.93元；第三经营年度运营成本（2024年10月28日—2024年12月31日）金额431,456.22元，2025年之后每年度的运营维护成本暂按第一经营年度和第二经营年度的平均值5,640,000.00元代入公式估算可行性缺口补助。')

add_h2('（四）第三方收入')
add_body('根据项目公司提供的财务资料，经审核，第一经营年度运营收入金额4,285,070.31元；第二经营年度运营收入金额5,841,165.87元；第三经营年度运营收入（2024年10月28日—2024年12月31日）金额387,927.58元，2025年之后每年度的运营收入暂按第一经营年度和第二经营年度的平均值5,060,000.00元代入公式估算可行性缺口补助。')

add_h2('（五）考核结果')
add_body('根据2018年3月7日巴中市恩阳区卫生和计划生育局与项目公司签订的补充协议，每年度巴中市恩阳区卫生健康局对项目运营期进行绩效考核并形成考核分数，考核分数作为支付可用性付费的重要依据，考核分数大于等于80分支付比例为100%，分数低于80分按照合同约定的调整系数进行计算。')
add_body('项目运营期限从2022年10月28日开始，截止2026年6月30日，巴中市恩阳区卫生健康局对该项目开展了四个运营年度的绩效考核，2022年度考核分数为93分、2023年度考核分数为94分、2024年度考核分数为95分、2025年度考核分数为93.5分，分数均达到80分以上，可用性付费支付比例为100%。')

# ===== 五、测算结果 =====
add_h1('五、测算结果')

add_body('根据PPP合同约定，项目进入运营期后应于每个运营年度末支付一次可用性费用，最终计算该项目每年可用性付费金额见下表（单位：元）：')

# Table 3: Final availability payment
add_body('表3：可用性付费测算表', indent=False)
avail_header = ['年', '每年项目资本金回报', '每年实际融资本息', '每年运营维护成本', '每年第三方收入', '可用性付费']
avail_rows = []

total_equity_all = 0
total_bank_all = 0
total_op_all = 0
total_income_all = 0
total_avail_all = 0

for i in range(n_years):
    yr = 2023 + i
    
    # Equity
    if i == n_years - 1:
        eq = equity_schedule[i]['payment'] + capital_lumpsum
        eq_str = fmt(eq)
    else:
        eq = equity_schedule[i]['payment']
        eq_str = fmt(eq)
    
    # Bank
    bk = bank_schedule[i][1]
    bk_str = fmt(bk) if bk > 0 else '-'
    
    # Op cost
    if i < 2:
        op = op_cost_actual[i]
    else:
        op = op_est
    op_str = fmt(op)
    
    # Income
    if i < 2:
        inc = income_actual[i]
    else:
        inc = income_est
    inc_str = fmt(inc)
    
    avail = eq + bk + op - inc
    
    total_equity_all += eq
    total_bank_all += bk
    total_op_all += op
    total_income_all += inc
    total_avail_all += avail
    
    avail_rows.append([str(yr), eq_str, bk_str, op_str, inc_str, fmt(avail)])

avail_rows.append([
    '合计',
    fmt(total_equity_all),
    fmt(total_bank_all),
    fmt(total_op_all),
    fmt(total_income_all),
    fmt(total_avail_all)
])

add_table(avail_header, avail_rows, [1.3, 3, 3, 2.8, 2.8, 2.8], bold_last=True)

# Notes
add_body('备注：项目运营期限从2022年10月28日开始，贵阳银行贷款结息日期为每年度12月21日，为统一各项费用的计费期限，除2023年度费用期限为2022年10月28日至2023年12月21日，其余年度费用期限均为上一年度12月21日至下年度12月21日。')

# Signature
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('四川融策会计师事务所有限公司')
run.font.size = Pt(12)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.first_line_indent = Cm(8)
run = p.add_run('二〇二六年七月十一日')
run.font.size = Pt(12)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== Save =====
output = r"C:\Users\scrccpa\Desktop\恩阳医养园PPP项目可用性付费测算报告（川融策咨询〔2025〕第490号）.docx"
doc.save(output)
print(f"Saved: {output}")
print(f"Size: {os.path.getsize(output):,} bytes")
print(f"可用性付费总额: {total_avail_all:,.2f}")
