import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx')

print(f'Total paragraphs: {len(doc.paragraphs)}')

# Show key structure
print('\n=== Document Structure ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold:
        # Show section boundaries and document titles
        if t in ['人力资源篇', '财务管理篇', '业务部管理篇', '业务质控篇', '行政综合篇']:
            print(f'\n--- [{i}] {t} ---')
        elif t.startswith('四川融策') and ('制度' in t or '规范' in t or '手册' in t or '章程' in t or '办法' in t or '细则' in t or '规则' in t):
            print(f'  [{i}] {t[:80]}')
        elif t.startswith('RC-'):
            print(f'  [{i}] {t[:80]}')

# Count documents per section
print('\n=== Document Counts ===')
sections = {
    '人力资源篇': [],
    '财务管理篇': [],
    '业务部管理篇': [],
    '业务质控篇': [],
    '行政综合篇': [],
}
current_section = None
for p in doc.paragraphs:
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and t in sections:
        current_section = t
    if current_section and is_bold and (t.startswith('RC-') or (t.startswith('四川融策') and any(k in t for k in ['制度','规范','手册','章程','办法','细则','规则']))):
        if t not in sections[current_section]:
            sections[current_section].append(t)

for section, docs in sections.items():
    print(f'  {section}: {len(docs)} docs')
    for d in docs[:3]:
        print(f'    - {d[:70]}')
    if len(docs) > 3:
        print(f'    ... and {len(docs)-3} more')
print(f'\nTotal docs: {sum(len(v) for v in sections.values())}')
