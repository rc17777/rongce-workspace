import sys
sys.stdout.reconfigure(encoding='utf-8')

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

COMPLETE = r'C:\Users\scrccpa\Desktop\融策公司制度体系（完整版）.docx'
NEW = r'C:\Users\scrccpa\Desktop\融策制度汇编-业务部管理篇.docx'
OUTPUT = COMPLETE  # overwrite

print('Loading documents...')
doc = Document(COMPLETE)
new_doc = Document(NEW)

# ============================================
# Step 1: Find the insertion point in the body
# ============================================
# Find the paragraph that is the section heading "业务质控篇" in the main body
# (not the TOC entry at the beginning)
target_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # The section heading is bold and has exactly/similarly "业务质控篇"
    is_bold = any(r.bold for r in p.runs if r.bold)
    if is_bold and '业务质控篇' in t:
        target_idx = i

print(f'Found "业务质控篇" section heading at paragraph index {target_idx}')
target_elem = doc.paragraphs[target_idx]._element

# ============================================
# Step 2: Get content from new doc (skip cover page)
# ============================================
# We want everything starting from "业务部管理篇 · 制度体系概览" onwards
new_elements = []
started = False
for p in new_doc.paragraphs:
    t = p.text.strip()
    if not started and '业务部管理篇 · 制度体系概览' in t:
        started = True
        continue  # skip the overview header itself, we'll handle it differently
    if not started:
        continue
    new_elements.append(p._element)

print(f'Collected {len(new_elements)} paragraph elements from new doc')

# ============================================
# Step 3: Insert a page break before the new section
# ============================================
# Create a page break paragraph
page_break_p = copy.deepcopy(new_elements[0])  # use first new element as template
# Clear it and make it just a page break
for child in list(page_break_p):
    page_break_p.remove(child)
# Add a run with page break
r_elem = etree.SubElement(page_break_p, qn('w:r'))
br_elem = etree.SubElement(r_elem, qn('w:br'))
br_elem.set(qn('w:type'), 'page')

# Insert new elements BEFORE the target (reverse order for addprevious)
all_to_insert = [page_break_p] + new_elements
for elem in reversed(all_to_insert):
    target_elem.addprevious(copy.deepcopy(elem))

print(f'Inserted {len(all_to_insert)} elements into complete doc')

# ============================================
# Step 4: Update the TOC
# ============================================
# Find the TOC section - the manual "目    录" near the start
toc_target_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    is_bold = any(r.bold for r in p.runs if r.bold)
    # TOC entry for 业务质控篇 - first occurrence with bold
    if is_bold and t == '业务质控篇' and i < 50:
        toc_target_idx = i
        break

if toc_target_idx:
    print(f'Found TOC entry "业务质控篇" at index {toc_target_idx}')
    toc_elem = doc.paragraphs[toc_target_idx]._element
    
    # Create a new paragraph for "业务部管理篇" in TOC style
    # Copy style from adjacent TOC entry
    ref_p = doc.paragraphs[toc_target_idx]
    new_p_elem = copy.deepcopy(ref_p._element)
    # Clear and set text
    for child in list(new_p_elem):
        new_p_elem.remove(child)
    r_elem = etree.SubElement(new_p_elem, qn('w:r'))
    rpr = etree.SubElement(r_elem, qn('w:rPr'))
    b_elem = etree.SubElement(rpr, qn('w:b'))
    t_elem = etree.SubElement(r_elem, qn('w:t'))
    t_elem.text = '业务部管理篇'
    t_elem.set(qn('xml:space'), 'preserve')
    
    toc_elem.addprevious(new_p_elem)
    print('Added "业务部管理篇" to TOC')
else:
    print('WARNING: Could not find TOC entry for 业务质控篇')

# ============================================
# Step 5: Save
# ============================================
doc.save(OUTPUT)
print(f'\nSaved to {OUTPUT}')

# Verify
import os
size = os.path.getsize(OUTPUT)
print(f'File size: {size:,} bytes')
